import pandas as pd
import matplotlib.pyplot as plt
import io
import base64
import os
import json
import re
from typing import Optional, Dict, Any, Union
from datetime import datetime, date
from pathlib import Path

from phenex.reporting.reporter import Reporter
from phenex.reporting.waterfall import Waterfall
from phenex.reporting.table1 import Table1
from phenex.reporting.table2 import Table2
from phenex.util import create_logger

logger = create_logger(__name__)


# Optional imports for document generation
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    DOCX_AVAILABLE = True
except ImportError:
    logger.warning(
        "python-docx not available. Word document generation will not work. Install with: pip install python-docx"
    )
    DOCX_AVAILABLE = False

try:
    from openai import AzureOpenAI, OpenAI

    OPENAI_AVAILABLE = True
except ImportError:
    logger.warning(
        "OpenAI library not available. AI-generated text will fall back to rules-based. Install with: pip install openai"
    )
    OPENAI_AVAILABLE = False


class ReportDrafter(Reporter):
    """
    The ReportDrafter creates comprehensive draft study reports including:
    - Cohort definition description (entry, inclusion, exclusion criteria)
    - Data analysis description and date ranges
    - Waterfall table showing patient attrition
    - Study variables (characteristics and outcomes)
    - Table 1 (baseline characteristics)
    - Table 2 (outcomes analysis)
    - AI-generated descriptive text and figure captions (when AI is enabled)

    **IMPORTANT: Human-in-the-Loop Required**

    The ReportDrafter generates DRAFT reports that require human review and editing before use. Reports are exported in editable formats (Markdown and Word) specifically to enable human oversight and refinement. AI-generated content should be verified for:
    - Clinical accuracy and appropriateness
    - Study-specific context and nuances
    - Compliance with institutional guidelines
    - Proper medical terminology and phrasing

    **Never use generated reports without thorough human review and approval.**

    The report can be exported to Markdown or Word format for human editing.

    What Does AI Generate?
    ----------------------
    **AI generates ONLY narrative text and commentary**, including:
    - Executive summary and abstract
    - Cohort definition descriptions
    - Data analysis methodology descriptions
    - Clinical interpretations and commentary for tables and figures

    **AI does NOT generate:**
    - Tables (Waterfall, Table 1, Table 2), plots and figures - these are calculated directly from your data using PhenEx library code

    The AI only provides contextual narrative around the data-driven tables and plots.

    AI Configuration
    ----------------
    The ReportDrafter can use OpenAI (Azure or standard) to generate professional medical research narrative text. If AI is not configured, it automatically falls back to rule-based text generation.

    **Option 1: Azure OpenAI**

    Set environment variables in Python:

    ```python
    import os
    # Environment variables for CREDENTIALS
    os.environ["AZURE_OPENAI_ENDPOINT"] = "https://your-resource.openai.azure.com/"
    os.environ["AZURE_OPENAI_API_KEY"] = "your-api-key-here"
    os.environ["OPENAI_API_VERSION"] = "2024-02-15-preview"

    from phenex.reporting import ReportDrafter
    reporter = ReportDrafter(ai_model="gpt-4o-mini")  # or "gpt-4", "gpt-3.5-turbo"
    ```

    **Option 2: Standard OpenAI**

    Set environment variable in Python:

    ```python
    import os
    # Environment variable for CREDENTIALS
    os.environ["OPENAI_API_KEY"] = "sk-your-api-key-here"

    from phenex.reporting import ReportDrafter
    reporter = ReportDrafter(ai_model="gpt-4o-mini")  # or "gpt-4", "gpt-3.5-turbo"
    ```

    **Disabling AI**

    To use rule-based text generation instead of AI:

    ```python
    reporter = ReportDrafter(use_ai=False)
    ```

    Parameters:
        use_ai: Whether to use AI for generating descriptive text. If True but API keys are not available, automatically falls back to rule-based generation. Default is True.
        ai_model: The model or deployment name to use when making API calls. Default is "gpt-4o-mini".
        include_plots: Whether to include plots in the report (e.g., waterfall charts). Default is True.
        plot_dpi: DPI (dots per inch) for plot image quality. Higher values produce better quality but larger file sizes. Default is 300.
        title: Report title. If None, will be generated from cohort name.
        author: Report author name(s) to display in report metadata.
        institution: Institution name to display in report metadata.
        decimal_places: Number of decimal places for numeric values in tables. Default is 1.
        pretty_display: Whether to use pretty display formatting with styled tables. Default is True.
        waterfall_reporter: Custom Waterfall reporter instance. If None, uses default configuration.
        table1_reporter: Custom Table1 reporter instance. If None, uses default configuration.
        table2_reporter: Custom Table2 reporter instance. If None, uses default configuration.

    Attributes:
        report_sections (dict): Dictionary containing all generated report sections
        figures (dict): Dictionary containing all generated figures and their metadata
        use_ai (bool): Whether AI is enabled and configured
        ai_client: The OpenAI client instance (if AI is enabled)

    Examples:

        Basic usage with AI (requires API keys in environment):
        ```python
        from phenex.reporting import ReportDrafter

        # Initialize reporter
        reporter = ReportDrafter(
            title="My Study Report",
            author="Dr. Jane Smith",
            institution="Research University"
        )

        # Generate DRAFT report from cohort
        reporter.execute(cohort)

        # Export to editable Markdown format for human review
        reporter.to_markdown("study_report_DRAFT.md", output_dir="./reports")

        # Export to editable Word format for human review and editing
        reporter.to_word("study_report_DRAFT.docx", output_dir="./reports")

        # IMPORTANT: Review and edit the generated files before using in publications
        # or formal reports. Verify all clinical statements, statistics, and interpretations.
        ```

        Without AI (rule-based text generation):
        ```python
        reporter = ReportDrafter(use_ai=False)
        reporter.execute(cohort)
        reporter.to_markdown("study_report.md")
        ```

        With custom reporters:
        ```python
        from phenex.reporting import Table1, Table2, Waterfall

        custom_table1 = Table1(decimal_places=2)
        custom_table2 = Table2(time_points=[30, 90, 180, 365])

        reporter = ReportDrafter(
            table1_reporter=custom_table1,
            table2_reporter=custom_table2,
            decimal_places=2
        )
        reporter.execute(cohort)
        ```

    Notes:
        - **HUMAN REVIEW REQUIRED**: All generated reports are drafts that MUST be reviewed, validated, and edited by qualified researchers before use. The ReportDrafter is a starting point to accelerate report creation, not a replacement for human expertise.
        - **AI generates ONLY text**: Tables, plots, and all statistical results are computed directly from your cohort data. AI only generates narrative text, descriptions, and clinical commentary.
        - Reports are intentionally exported in editable formats (Markdown/Word) to facilitate human review and modification
        - AI-generated content should be verified for clinical accuracy, institutional compliance, and study-specific appropriateness
        - AI generation requires valid OpenAI or Azure OpenAI credentials
        - If credentials are missing or invalid, automatically falls back to rule-based generation
        - The reporter will log warnings if AI is requested but unavailable
        - Generated reports include executive summary, methods, results, and clinical commentary
        - Markdown output includes all tables and can embed plot images
        - All numerical results in tables are computed from actual cohort data, not AI-generated
    """

    def __init__(
        self,
        use_ai: bool = True,
        ai_model: str = "gpt-4o-mini",
        include_plots: bool = True,
        plot_dpi: int = 300,
        title: Optional[str] = None,
        author: Optional[str] = None,
        institution: Optional[str] = None,
        decimal_places: int = 1,
        pretty_display: bool = True,
        waterfall_reporter: Optional[Any] = None,
        table1_reporter: Optional[Any] = None,
        table2_reporter: Optional[Any] = None,
    ):
        super().__init__(decimal_places=decimal_places, pretty_display=pretty_display)

        self.use_ai = use_ai and OPENAI_AVAILABLE and self._check_openai_config()
        self.ai_model = ai_model
        self.include_plots = include_plots
        self.plot_dpi = plot_dpi
        self.title = title
        self.author = author
        self.institution = institution

        logger.info(
            f"ReportDrafter initialized with include_plots={self.include_plots}"
        )

        # Set report date to current date
        self.date = datetime.now().strftime("%Y-%m-%d")

        # Store custom reporter instances (will be used if provided, otherwise defaults)
        self.waterfall_reporter = waterfall_reporter
        self.table1_reporter = table1_reporter
        self.table2_reporter = table2_reporter

        # Initialize OpenAI client if available
        self.ai_client = None
        self._is_azure = False
        if self.use_ai:
            self._initialize_ai_client()

        # Report sections storage
        self.report = {}
        self.figures = {}

    def _check_openai_config(self) -> bool:
        """Check if OpenAI configuration is available in environment variables."""
        # Check for Azure OpenAI configuration
        azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
        azure_api_key = os.getenv("AZURE_OPENAI_API_KEY")

        # Check for standard OpenAI configuration
        openai_api_key = os.getenv("OPENAI_API_KEY")

        return bool(azure_endpoint and azure_api_key) or bool(openai_api_key)

    def _initialize_ai_client(self):
        """Initialize OpenAI client and report configuration status to user."""
        try:
            # Check for Azure OpenAI configuration first
            azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
            azure_api_key = os.getenv("AZURE_OPENAI_API_KEY")
            api_version = os.getenv("OPENAI_API_VERSION", "2024-02-15-preview")

            if azure_endpoint and azure_api_key:
                # Configure Azure OpenAI
                if api_version.startswith('"') and api_version.endswith('"'):
                    api_version = api_version.strip('"')

                self.ai_client = AzureOpenAI(
                    azure_endpoint=azure_endpoint.strip(),
                    api_key=azure_api_key.strip(),
                    api_version=api_version.strip(),
                )
                self._is_azure = True

                # Test the connection
                test_response = self.ai_client.chat.completions.create(
                    model=self.ai_model,
                    messages=[{"role": "user", "content": "Test"}],
                    max_tokens=5,
                )

                logger.info(f"✅ Using Azure OpenAI (model: {self.ai_model})")
                return

            # Check for standard OpenAI configuration
            openai_api_key = os.getenv("OPENAI_API_KEY")
            if openai_api_key:
                self.ai_client = OpenAI(api_key=openai_api_key.strip())
                self._is_azure = False
                logger.info(f"✅ Using OpenAI (model: {self.ai_model})")
                return

            # No valid configuration found
            raise Exception("No AI configuration found")

        except Exception as e:
            # AI initialization failed - provide clear guidance
            self.ai_client = None
            self.use_ai = False
            self._is_azure = False

            logger.warning("⚠️ Using rule-based text generation (AI not configured)")
            logger.info("")
            logger.info(
                "To enable AI-powered text generation, set environment variables:"
            )
            logger.info("")
            logger.info("Option 1 - Azure OpenAI:")
            logger.info(
                '  os.environ["AZURE_OPENAI_ENDPOINT"] = "https://your-resource.openai.azure.com/"'
            )
            logger.info('  os.environ["AZURE_OPENAI_API_KEY"] = "your-api-key"')
            logger.info('  os.environ["OPENAI_API_VERSION"] = "2024-02-15-preview"')
            logger.info("")
            logger.info("Option 2 - Standard OpenAI:")
            logger.info('  os.environ["OPENAI_API_KEY"] = "sk-your-api-key"')
            logger.info("")

            if str(e) != "No AI configuration found":
                logger.debug(f"AI initialization error details: {e}")

    def _generate_ai_text(
        self,
        prompt: str,
        max_tokens: int = 16384,
    ) -> str:
        """Generate text using AI or fallback to rules-based generation."""
        if not self.use_ai:
            logger.debug("AI disabled, using fallback text generation")
            return self._fallback_text_generation(prompt, None)

        logger.info(
            f"🤖 Making AI API call for text generation (max_tokens: {max_tokens})..."
        )

        # Inject global context automatically from class variable
        global_context = getattr(self, "_global_context", "")
        full_prompt = f"{global_context}\n\n{prompt}"

        logger.debug(f"AI prompt preview: {prompt[:100]}...")

        try:
            # Try to use the OpenAI client
            response = self.ai_client.chat.completions.create(
                model=self.ai_model,
                messages=[
                    {
                        "role": "system",
                        "content": self._get_global_ai_system_instructions(),
                    },
                    {"role": "user", "content": full_prompt},
                ],
                max_tokens=max_tokens,
                temperature=0,
            )
            generated_text = response.choices[0].message.content.strip()
            logger.info(
                f"✅ AI text generation successful (generated {len(generated_text)} characters)"
            )
            logger.debug(f"AI response preview: {generated_text[:100]}...")
            return generated_text
        except Exception as e:
            logger.warning(
                f"❌ AI text generation failed, falling back to rules-based: {e}"
            )
            return self._fallback_text_generation(prompt, None)

    def _get_global_ai_system_instructions(self) -> str:
        """Get global system instructions for all AI text generation calls."""
        return """You are a professional medical researcher and biostatistician writing for a high-impact scientific publication.

WRITING STYLE REQUIREMENTS:
- Keep answers concise yet insightful and informative
- Use professional medical research language suitable for peer-reviewed journals
- Provide clinical context and significance for all findings
- Be precise with statistical interpretations
- Include relevant clinical implications
- Use appropriate medical terminology consistently
- Write in active voice where appropriate
- Ensure content is publication-ready

CONTENT STANDARDS:
- All statistical claims must be clinically meaningful
- Include appropriate caveats and limitations where relevant  
- Focus on actionable clinical insights
- Maintain scientific objectivity and accuracy
- Reference established clinical guidelines and norms when relevant
- Ensure content flows logically and cohesively

FORMATTING:
- Use clean markdown formatting
- Structure content with clear headings and sections
- Begin your response at the subsection (##) level
- Use bullet points for lists where appropriate
- Ensure proper medical/scientific citation style"""

    def _limit_codelists_in_dict(self, data_dict: dict, max_codes: int = 5) -> dict:
        """
        Recursively limit codelist sizes in a dictionary to prevent context overflow.

        Args:
            data_dict: Dictionary that may contain codelists
            max_codes: Maximum number of codes to show (default: 5)

        Returns:
            Modified dictionary with limited codelists
        """
        if not isinstance(data_dict, dict):
            return data_dict

        result = {}
        for key, value in data_dict.items():
            if isinstance(value, list):
                # Truncate any list to max_codes and add note
                if len(value) > max_codes:
                    result[key] = value[:max_codes] + ["... truncated codelist"]
                else:
                    result[key] = value
            elif isinstance(value, dict):
                # Recursively process nested dictionaries
                result[key] = self._limit_codelists_in_dict(value, max_codes)
            else:
                result[key] = value

        return result

    def _build_global_ai_context(self, cohort) -> str:
        """
        Build comprehensive global context that is automatically injected into every AI call.
        This ensures all AI responses have complete study awareness and consistency.
        """
        context_parts = []

        # Study Overview
        context_parts.append(
            f"""=== COMPREHENSIVE STUDY CONTEXT ===
STUDY TITLE: {getattr(self, 'title', 'Medical Research Study')}
COHORT NAME: {getattr(cohort, 'name', 'Study Cohort')}
COHORT DESCRIPTION: {getattr(cohort, 'description', 'Not available')}
STUDY TYPE: Comprehensive medical research study analyzing patient outcomes and characteristics"""
        )

        # Cohort Information
        n_entry_patients = (
            cohort.entry_criterion.table.select("PERSON_ID")
            .distinct()
            .count()
            .execute()
        )
        n_index_patients = cohort.table.select("PERSON_ID").distinct().count().execute()
        context_parts.append(
            f"ENTRY COHORT SIZE (ENTRY CRITERION ONLY): {n_entry_patients} patients"
        )
        context_parts.append(
            f"FINAL COHORT SIZE (ALL INEX CRITERIA APPLIED): {n_index_patients} patients"
        )

        # Entry criteria
        context_parts.append(f"ENTRY CRITERION: {cohort.entry_criterion.name}")
        entry_dict = self._limit_codelists_in_dict(cohort.entry_criterion.to_dict())
        context_parts.append(f"\n\t{json.dumps(entry_dict, indent=4)}")

        # Inclusions
        if hasattr(cohort, "inclusions") and cohort.inclusions:
            context_parts.append(
                f"\nINCLUSION CRITERIA ({len(cohort.inclusions)} criteria):"
            )
            for i, inclusion in enumerate(cohort.inclusions, 1):
                name = getattr(inclusion, "name", f"Inclusion {i}")
                context_parts.append(f"  {i}. {name}")
                inclusion_dict = self._limit_codelists_in_dict(inclusion.to_dict())
                context_parts.append(f"\n\t{json.dumps(inclusion_dict, indent=4)}")

        # Exclusions
        if hasattr(cohort, "exclusions") and cohort.exclusions:
            context_parts.append(
                f"\nEXCLUSION CRITERIA ({len(cohort.exclusions)} criteria):"
            )
            for i, exclusion in enumerate(cohort.exclusions, 1):
                name = getattr(exclusion, "name", f"Exclusion {i}")
                context_parts.append(f"  {i}. {name}")
                exclusion_dict = self._limit_codelists_in_dict(exclusion.to_dict())
                context_parts.append(f"\n\t{json.dumps(exclusion_dict, indent=4)}")

        # Characteristics
        if hasattr(cohort, "characteristics") and cohort.characteristics:
            context_parts.append(
                f"\nBASELINE CHARACTERISTICS ({len(cohort.characteristics)} variables):"
            )
            for i, char in enumerate(cohort.characteristics, 1):
                name = getattr(char, "name", f"Characteristic {i}")
                context_parts.append(f"  {i}. {name}")
                char_dict = self._limit_codelists_in_dict(char.to_dict())
                context_parts.append(f"\n\t{json.dumps(char_dict, indent=4)}")

        # Outcomes
        if hasattr(cohort, "outcomes") and cohort.outcomes:
            context_parts.append(
                f"\nOUTCOME MEASURES ({len(cohort.outcomes)} variables):"
            )
            for i, outcome in enumerate(cohort.outcomes, 1):
                name = getattr(outcome, "name", f"Outcome {i}")
                context_parts.append(f"  {i}. {name}")
                outcome_dict = self._limit_codelists_in_dict(outcome.to_dict())
                context_parts.append(f"\n\t{json.dumps(outcome_dict, indent=4)}")

        # Report generation metadata
        if hasattr(self, "author") and self.author:
            context_parts.append(f"\nREPORT AUTHOR: {self.author}")
        if hasattr(self, "institution") and self.institution:
            context_parts.append(f"INSTITUTION: {self.institution}")

        # Include generated table data if available
        context_parts.append("\n=== GENERATED TABLE DATA ===")

        # Waterfall table data
        if hasattr(self, "report_sections") and "waterfall_table" in self.report:
            waterfall_df = self.report["waterfall_table"]
            # Extract DataFrame from Styler if needed
            if hasattr(waterfall_df, "data"):
                waterfall_df = waterfall_df.data
            if not waterfall_df.empty:
                context_parts.append(f"\nWATERFALL TABLE ({len(waterfall_df)} rows):")
                context_parts.append(
                    "Patient attrition through inclusion/exclusion criteria:"
                )
                context_parts.append(waterfall_df.to_string())

        # Table 1 (baseline characteristics) data
        if hasattr(self, "report_sections") and "table1" in self.report:
            table1_df = self.report["table1"]
            if not table1_df.empty:
                context_parts.append(
                    f"\nTABLE 1 - BASELINE CHARACTERISTICS ({len(table1_df)} rows):"
                )
                context_parts.append(
                    "Demographic and clinical characteristics at baseline:"
                )
                context_parts.append(table1_df.to_string())

        # Table 2 (outcomes) data
        if hasattr(self, "report_sections") and "table2" in self.report:
            table2_df = self.report["table2"]
            if not table2_df.empty:
                context_parts.append(
                    f"\nTABLE 2 - OUTCOMES SUMMARY ({len(table2_df)} rows):"
                )
                context_parts.append("Clinical outcomes and incidence rates:")
                context_parts.append(table2_df.to_string())

        context_parts.append("\n=== END GLOBAL CONTEXT ===")

        return "\n".join(context_parts)

    def _generate_ai_image_caption(self, image_base64: str, context: str) -> str:
        """Generate caption for image - AI-powered if available, otherwise fallback."""
        if not self.use_ai or not self.ai_client:
            logger.debug(
                "AI disabled or client unavailable, using fallback figure caption"
            )
            # Enhanced fallback caption based on context
            if "waterfall" in context.lower():
                return "Patient attrition waterfall showing the step-by-step filtering process applied to the initial population to derive the final study cohort."
            elif "outcome" in context.lower():
                return "Clinical outcomes analysis showing the distribution and timing of key endpoint events in the study population."
            elif "characteristic" in context.lower():
                return "Baseline characteristics summary displaying the demographic and clinical profile of the study population."
            else:
                return f"Study figure illustrating {context.lower()} for the analysis population."

        logger.info(f"🖼️ Generating AI figure caption using text model...")
        logger.debug(f"Image context: {context}")

        try:
            # Use text-based generation instead of vision API since gpt-4-vision-preview is not available
            prompt = f"""
            Generate a professional medical research figure caption for a plot with this context: {context}
            
            The caption should be:
            - Professional and suitable for a medical research publication
            - Clear and descriptive
            - Include relevant clinical interpretation
            - Follow standard academic figure caption format
            
            Start with "Figure X:" and provide a comprehensive description.
            """

            response = self.ai_client.chat.completions.create(
                model=self.ai_model,  # Use the same model as other text generation
                messages=[
                    {
                        "role": "system",
                        "content": "You are a professional medical researcher writing figure captions for a scientific publication.",
                    },
                    {"role": "user", "content": prompt},
                ],
                temperature=0.7,
            )
            generated_caption = response.choices[0].message.content.strip()
            logger.info(
                f"✅ AI figure caption generation successful (generated {len(generated_caption)} characters)"
            )
            return generated_caption
        except Exception as e:
            logger.warning(
                f"❌ AI figure caption generation failed, using fallback: {e}"
            )
            return f"Figure: {context}"

    def _fallback_text_generation(self, prompt: str, cohort=None) -> str:
        """Fallback text generation using rules-based approach."""
        if "cohort definition" in prompt.lower():
            return (
                self._create_specific_cohort_description(cohort)
                if cohort
                else "This cohort was defined using entry criteria, inclusion criteria, and exclusion criteria as specified in the methods section."
            )
        elif "data analysis" in prompt.lower():
            return "Data analysis was performed using the PhenEx framework, applying the specified phenotype definitions and filters."
        elif "baseline characteristics" in prompt.lower():
            return "Baseline characteristics were calculated at the index date for all patients meeting the inclusion and exclusion criteria."
        elif "outcomes" in prompt.lower():
            return "Outcomes were evaluated for all patients in the final cohort during the follow-up period."
        else:
            return "Details are provided in the accompanying tables and figures."

    def _format_cohort_name(self, name: str) -> str:
        """Format cohort name from snake_case to proper Title Case."""
        if not name:
            return name

        # Convert snake_case or camelCase to Title Case
        # Handle snake_case (e.g., study_tutorial_cohort -> Study Tutorial Cohort)
        formatted = re.sub(r"_", " ", name)

        # Handle camelCase (e.g., studyTutorialCohort -> Study Tutorial Cohort)
        formatted = re.sub(r"([a-z])([A-Z])", r"\1 \2", formatted)

        # Convert to title case
        formatted = formatted.title()

        return formatted

    def _create_executive_summary(self) -> str:
        """Generate executive summary - AI-powered if available, otherwise fallback."""
        logger.info("Generating executive summary")

        if self.use_ai and self.ai_client:
            prompt = """TASK: Write a professional medical journal-style executive summary/abstract for this study.
            
            ABSTRACT STRUCTURE REQUIREMENTS:
            - **Objective:** What was studied and why
            - **Methods:** Brief description of study design, population, and criteria
            - **Results:** Key findings from baseline characteristics and outcomes (use realistic clinical interpretations)
            - **Conclusions:** Clinical implications and significance
            
            SPECIFIC REQUIREMENTS:
            - Medical journal abstract format (150-250 words)
            - Focus on clinical significance and real-world implications
            - Use realistic medical findings appropriate for the study population
            - Include key statistical insights where clinically relevant
            
            Write a complete executive summary that reads like a published medical research abstract."""

            return self._generate_ai_text(prompt)
        else:
            # Fallback summary when AI is not available
            stats = self.report.get("summary_stats", {})
            cohort_name = getattr(self, "cohort_name", "the study cohort")

            return f"""## Abstract

**Objective:** This study presents a comprehensive analysis of {cohort_name}, examining baseline patient characteristics, treatment patterns, and clinical outcomes to better understand the study population and inform clinical decision-making.

**Methods:** We conducted a retrospective cohort study analyzing {stats.get('total_patients', 'N/A')} patients meeting predefined inclusion and exclusion criteria. The study included {stats.get('n_inclusions', 0)} inclusion criteria and {stats.get('n_exclusions', 0)} exclusion criteria to ensure a well-defined study population. Baseline characteristics were assessed using {stats.get('n_characteristics', 0)} variables, and clinical outcomes were evaluated using {stats.get('n_outcomes', 0)} outcome measures.

**Results:** The analysis provides detailed insights into patient demographics, comorbidities, and treatment utilization patterns. Patient attrition through inclusion and exclusion criteria is documented in a comprehensive waterfall analysis, ensuring transparency in cohort selection. Baseline characteristics and outcome summaries are presented to characterize the study population.

**Conclusions:** This analysis provides valuable insights into the characteristics and outcomes of the study population. The structured approach to cohort definition and comprehensive outcome assessment supports evidence-based clinical decision-making and may inform future research directions in this patient population."""

    def _create_cohort_description(self, cohort) -> str:
        """Generate cohort definition description - AI-powered if available, otherwise fallback."""
        logger.info(f"Generating cohort description for: {cohort.name}")

        if self.use_ai and self.ai_client:
            # Format the cohort name properly
            formatted_cohort_name = self._format_cohort_name(cohort.name)

            prompt = f"""
            Write a professional description of this medical research cohort definition using clean markdown formatting:
            
            Cohort Name: {formatted_cohort_name}
            Cohort Description: {cohort.description or 'Not provided'}
            
            Entry Criterion: {cohort.entry_criterion.to_dict()}
            
            Inclusion Criteria:
            {chr(10).join([f"- {inc.display_name if hasattr(inc, 'display_name') else inc.name}" for inc in (cohort.inclusions or [])])}
            
            Exclusion Criteria:
            {chr(10).join([f"- {exc.display_name if hasattr(exc, 'display_name') else exc.name}" for exc in (cohort.exclusions or [])])}
            
            Please write a professional medical research description with:
            - Brief introduction paragraph about the study population
            - **Entry Criterion:** section with rationale
            - **Inclusion Criteria:** section with bullet points (use * for bullets)
            - **Exclusion Criteria:** section with bullet points (use * for bullets)
            - Clinical rationale for each criterion (max one sentence)
            
            Use clean markdown formatting with proper line breaks between sections.
            """
            return self._generate_ai_text(prompt)
        else:
            # Fallback cohort description when AI is not available
            formatted_cohort_name = self._format_cohort_name(cohort.name)
            description_parts = []

            # Introduction
            description_parts.append(f"## Cohort Definition: {formatted_cohort_name}")
            description_parts.append("")

            if hasattr(cohort, "description") and cohort.description:
                description_parts.append(
                    f"The **{formatted_cohort_name}** is a comprehensive study designed to {cohort.description.lower()}."
                )
            else:
                description_parts.append(
                    f"The **{formatted_cohort_name}** represents a well-defined patient population selected using specific clinical criteria to ensure study validity and generalizability."
                )
            description_parts.append("")

            # Entry criterion
            entry_name = (
                cohort.entry_criterion.display_name
                if hasattr(cohort.entry_criterion, "display_name")
                else cohort.entry_criterion.name
            )
            description_parts.append("### Entry Criterion:")
            description_parts.append(f"- **{entry_name}**  ")
            description_parts.append(
                "  This criterion ensures that all participants meet the primary study condition of interest."
            )
            description_parts.append("")

            # Inclusion criteria
            if cohort.inclusions:
                description_parts.append("### Inclusion Criteria:")
                for inc in cohort.inclusions:
                    inc_name = (
                        inc.display_name if hasattr(inc, "display_name") else inc.name
                    )
                    description_parts.append(f"* **{inc_name}**  ")
                    description_parts.append(
                        "  This criterion helps define the target population for the study."
                    )
                    description_parts.append("")

            # Exclusion criteria
            if cohort.exclusions:
                description_parts.append("### Exclusion Criteria:")
                for exc in cohort.exclusions:
                    exc_name = (
                        exc.display_name if hasattr(exc, "display_name") else exc.name
                    )
                    description_parts.append(f"* **{exc_name}**  ")
                    description_parts.append(
                        "  This exclusion helps ensure study population homogeneity and reduces confounding factors."
                    )
                    description_parts.append("")

            description_parts.append(
                "This structured approach to cohort definition ensures a well-characterized study population suitable for meaningful clinical research and outcome assessment."
            )

            return "\n".join(description_parts)

    def _create_specific_cohort_description(self, cohort) -> str:
        """Create specific cohort description with actual criteria listed."""
        if not cohort:
            return "This cohort was defined using entry criteria, inclusion criteria, and exclusion criteria as specified in the methods section."

        description_parts = []

        # Add cohort description if available
        if hasattr(cohort, "description") and cohort.description:
            description_parts.append(f"Study Population: {cohort.description}")

        # Entry criterion
        entry_name = (
            cohort.entry_criterion.display_name
            if hasattr(cohort.entry_criterion, "display_name")
            else cohort.entry_criterion.name
        )
        description_parts.append(f"Entry Criterion: Patients with {entry_name}.")

        # Inclusion criteria
        if cohort.inclusions:
            inclusion_list = []
            for inc in cohort.inclusions:
                inc_name = (
                    inc.display_name if hasattr(inc, "display_name") else inc.name
                )
                inclusion_list.append(inc_name)
            if inclusion_list:
                description_parts.append(
                    f"Inclusion Criteria: {', '.join(inclusion_list)}."
                )

        # Exclusion criteria
        if cohort.exclusions:
            exclusion_list = []
            for exc in cohort.exclusions:
                exc_name = (
                    exc.display_name if hasattr(exc, "display_name") else exc.name
                )
                exclusion_list.append(exc_name)
            if exclusion_list:
                description_parts.append(
                    f"Exclusion Criteria: Patients were excluded if they had {', '.join(exclusion_list)}."
                )

        return " ".join(description_parts)

    def _create_data_analysis_description(self, cohort) -> str:
        """Generate data analysis description - AI-powered if available, otherwise fallback."""
        logger.info("Generating data analysis description")

        if self.use_ai and self.ai_client:
            prompt = f"""
            Write a description of the data analysis for the described medical research study.
            
            Your summary should consist of three sections: analytical approach, patient population, and study period.
            """
            return self._generate_ai_text(prompt)
        else:
            # Fallback data analysis description when AI is not available
            stats = self.report.get("summary_stats", {})
            cohort_name = getattr(self, "cohort_name", "the study cohort")

            return f"""## Analytical Approach

The analysis of {cohort_name} employed a comprehensive retrospective cohort study design to evaluate patient characteristics, treatment patterns, and clinical outcomes. The analytical framework included systematic data collection, quality assessment, and statistical analysis of patient-level data.

### Patient Population

The study population consisted of {stats.get('total_patients', 'N/A')} patients who met the predefined study criteria. Patient selection involved {stats.get('n_inclusions', 0)} inclusion criteria and {stats.get('n_exclusions', 0)} exclusion criteria to ensure a well-defined and clinically relevant study population. This systematic approach to patient selection helps minimize selection bias and ensures the generalizability of study findings.

### Data Collection and Variables

Baseline characteristics were assessed using {stats.get('n_characteristics', 0)} variables encompassing demographic information, clinical history, comorbidities, and treatment patterns. Clinical outcomes were evaluated using {stats.get('n_outcomes', 0)} outcome measures designed to capture key clinical endpoints relevant to the study population.

### Statistical Methods

Descriptive statistics were used to characterize the study population, including measures of central tendency and dispersion for continuous variables, and frequencies and percentages for categorical variables. Patient attrition through the cohort selection process was documented using waterfall methodology to ensure transparency in the final study population composition.

The analysis provides a comprehensive view of the study population characteristics and serves as the foundation for understanding treatment patterns and clinical outcomes in this patient cohort."""

    def _create_variables_description(self, cohort) -> str:
        """Generate description of study variables - AI-powered if available, otherwise fallback."""
        logger.info("Generating study variables description")

        characteristics = cohort.characteristics or []
        outcomes = cohort.outcomes or []

        char_names = [
            c.display_name if hasattr(c, "display_name") else c.name
            for c in characteristics
        ]
        outcome_names = [
            o.display_name if hasattr(o, "display_name") else o.name for o in outcomes
        ]

        if self.use_ai and self.ai_client:
            prompt = f"""
            Write a professional description of the study variables for this medical research study.
            
            Baseline Characteristics ({len(characteristics)}):
            {chr(10).join([f"- {name}" for name in char_names])}
            
            Outcome Variables ({len(outcomes)}):
            {chr(10).join([f"- {name}" for name in outcome_names])}
            
            REQUIREMENTS:
            - Use numbered lists for each variable (1. Variable Name: Description)
            - Group baseline characteristics separately from outcome variables
            - Explain the clinical relevance of each variable (max one sentence)
            - Include measurement methods where appropriate
            
            Write a comprehensive study variables section.
            """
            return self._generate_ai_text(prompt)
        else:
            # Fallback study variables description when AI is not available
            description_parts = []
            description_parts.append("## Study Variables")
            description_parts.append("")
            description_parts.append(
                "The study employed a comprehensive set of variables to characterize the patient population and assess clinical outcomes. Variables were selected based on clinical relevance, data availability, and potential impact on study outcomes."
            )
            description_parts.append("")

            # Baseline Characteristics
            if characteristics:
                description_parts.append(
                    f"### Baseline Characteristics ({len(characteristics)} variables)"
                )
                description_parts.append("")
                description_parts.append(
                    "Baseline characteristics were assessed to describe the study population and identify potential confounding factors:"
                )
                description_parts.append("")

                for i, name in enumerate(char_names, 1):
                    description_parts.append(
                        f"{i}. **{name}**: Baseline measurement used to characterize the study population and assess potential confounding factors."
                    )
                description_parts.append("")

            # Outcome Variables
            if outcomes:
                description_parts.append(
                    f"### Outcome Variables ({len(outcomes)} variables)"
                )
                description_parts.append("")
                description_parts.append(
                    "Outcome variables were selected to capture clinically meaningful endpoints relevant to the study population:"
                )
                description_parts.append("")

                for i, name in enumerate(outcome_names, 1):
                    description_parts.append(
                        f"{i}. **{name}**: Clinical outcome measure used to assess treatment effectiveness and patient prognosis."
                    )
                description_parts.append("")

            description_parts.append("### Data Quality and Validation")
            description_parts.append("")
            description_parts.append(
                "All variables underwent systematic quality assessment to ensure data completeness and accuracy. Missing data patterns were evaluated, and appropriate statistical methods were applied to handle any data gaps while maintaining the integrity of the analysis."
            )

            return "\n".join(description_parts)

    def _generate_waterfall_commentary(self, waterfall_df):
        """Generate AI commentary for waterfall table."""
        logger.info("Generating waterfall table commentary")

        # Extract DataFrame from Styler if needed
        if hasattr(waterfall_df, "data"):
            waterfall_df = waterfall_df.data

        if waterfall_df is None or waterfall_df.empty:
            return "No waterfall data available for analysis."

        # Extract key statistics
        initial_n = waterfall_df.iloc[0]["N"] if len(waterfall_df) > 0 else "Unknown"
        final_n = (
            waterfall_df.iloc[-1]["Remaining"] if len(waterfall_df) > 0 else "Unknown"
        )
        inclusion_steps = waterfall_df[waterfall_df["Type"] == "inclusion"]
        exclusion_steps = waterfall_df[waterfall_df["Type"] == "exclusion"]

        if self.use_ai and self.ai_client:
            prompt = f"""
            Analyze this patient attrition waterfall table and write a professional clinical commentary.
            
            WATERFALL DATA:
            Initial patient pool: {initial_n}
            Final cohort size: {final_n}
            Number of inclusion criteria: {len(inclusion_steps)}
            Number of exclusion criteria: {len(exclusion_steps)}
            
            Detailed attrition steps:
            {waterfall_df[['Type', 'Name', 'N', 'Remaining']].to_string()}
            
            Focus on:
            - Clinical interpretation of patient selection process
            - Analysis of attrition rates and their implications
            - Assessment of study representativeness and generalizability
            - Discussion of potential selection bias considerations
            """

            return self._generate_ai_text(prompt)
        else:
            # Fallback waterfall commentary when AI is not available
            n_inclusions = len(inclusion_steps)
            n_exclusions = len(exclusion_steps)

            return f"""## Patient Attrition Analysis

### Selection Process Overview

The patient selection process began with an initial population of {initial_n} patients and resulted in a final study cohort of {final_n} patients. This systematic selection process involved {n_inclusions} inclusion criteria and {n_exclusions} exclusion criteria to ensure a well-defined study population.

### Attrition Summary

The stepwise application of inclusion and exclusion criteria demonstrates a structured approach to cohort definition:

- **Initial Population**: {initial_n} patients met the primary entry criterion
- **Inclusion Criteria**: {n_inclusions} criteria were applied to refine the target population
- **Exclusion Criteria**: {n_exclusions} criteria were applied to remove patients with conditions that could confound study results
- **Final Cohort**: {final_n} patients comprised the final study population

### Clinical Interpretation

The systematic patient selection process ensures that the final cohort represents a clinically relevant population suitable for the research objectives. The application of both inclusion and exclusion criteria helps minimize confounding factors while maintaining sufficient sample size for meaningful analysis.

### Study Representativeness

The final cohort size of {final_n} patients provides adequate statistical power for the planned analyses. The structured selection process helps ensure that findings will be applicable to similar patient populations in clinical practice, while the transparency of the attrition process supports the validity and reproducibility of the study results."""

    def _generate_table1_commentary(self, table1_df):
        """Generate AI commentary for Table 1 baseline characteristics."""
        logger.info("Generating Table 1 commentary")

        if table1_df is None or table1_df.empty:
            return "No baseline characteristics data available for analysis."

        if self.use_ai and self.ai_client:
            prompt = f"""
            Analyze this Table 1 baseline characteristics and write a professional clinical commentary.
            
            BASELINE CHARACTERISTICS DATA:
            {table1_df.to_string()}
            
            Focus on:
            - Clinical interpretation of baseline demographics and characteristics
            - Assessment of population representativeness
            - Clinical implications for study outcomes
            - Comparison to relevant population norms where appropriate
            - Risk factor assessment and clinical significance
             """

            return self._generate_ai_text(prompt)
        else:
            # Fallback Table 1 commentary when AI is not available
            total_patients = (
                table1_df[table1_df["Name"] == "Cohort"]["N"].iloc[0]
                if not table1_df[table1_df["Name"] == "Cohort"].empty
                else "N/A"
            )

            return f"""## Baseline Characteristics Analysis

### Population Overview

The study cohort comprised {total_patients} patients with comprehensive baseline characteristics collected to ensure appropriate population characterization. The baseline demographics and clinical characteristics provide important context for interpreting study outcomes and assessing the generalizability of findings.

### Demographic Profile

The demographic characteristics of the study population reflect the target population for this research. Age, gender, and race distributions provide important context for understanding the representativeness of the cohort and potential implications for clinical outcomes.

### Clinical Risk Factors

The baseline clinical characteristics encompass important risk factors and comorbidities relevant to the study outcomes. The presence of conditions such as hypertension, diabetes, and cardiovascular disease helps characterize the overall risk profile of the study population.

### Treatment Patterns

Baseline medication utilization patterns provide insight into the treatment landscape and standard of care within the study population. These patterns help contextualize subsequent outcome analyses and may identify important confounding factors.

### Clinical Implications

The baseline characteristics profile suggests a clinically relevant study population suitable for addressing the research objectives. The comprehensive characterization enables appropriate interpretation of study outcomes and supports the validity of conclusions drawn from the analysis.

This baseline characterization provides the foundation for understanding treatment patterns, clinical outcomes, and the overall clinical significance of study findings."""

    def _generate_table2_commentary(self, table2_df):
        """Generate AI commentary for Table 2 outcomes."""
        logger.info("Generating Table 2 commentary")

        if table2_df is None or table2_df.empty:
            return "No outcomes data available for analysis."

        if self.use_ai and self.ai_client:
            prompt = f"""
            Analyze this Table 2 outcomes summary and write a professional clinical commentary.
            
            OUTCOMES DATA:
            {table2_df.to_string()}
            
            Focus on:
            - Clinical interpretation of outcome results
            - Assessment of key findings and their significance
            - Clinical implications for patient care and clinical practice
            - Risk assessment and prognostic implications
            - Comparison to published literature where appropriate
            """

            return self._generate_ai_text(prompt)
        else:
            # Fallback Table 2 commentary when AI is not available
            n_outcomes = len(table2_df) if table2_df is not None else 0

            return f"""## Clinical Outcomes Analysis

### Outcomes Overview

The study evaluated {n_outcomes} clinical outcomes to assess the key endpoints relevant to the study population. These outcomes were selected based on clinical relevance, patient safety considerations, and their importance for clinical decision-making.

### Event Rates and Incidence

The outcome analysis provides important insights into the frequency and timing of key clinical events within the study population. Incidence rates and event counts help characterize the clinical burden and risk profile of the cohort.

### Time-to-Event Analysis

The outcomes data includes time-under-risk calculations that account for differential follow-up periods among patients. This approach ensures accurate estimation of incidence rates and provides meaningful comparison of outcome frequencies across different time periods.

### Clinical Significance

The outcome measures capture clinically meaningful endpoints that reflect important aspects of patient health and prognosis. These results provide valuable information for healthcare providers and support evidence-based clinical decision-making.

### Risk Assessment

The pattern of outcomes observed in this study population provides important information about the overall risk profile and prognosis of patients with similar characteristics. This information can help inform treatment strategies and patient counseling.

### Clinical Implications

The outcomes analysis contributes to our understanding of disease progression and treatment effectiveness in this patient population. These findings may inform future research directions and clinical practice guidelines for similar patient populations."""

    def _plot_to_base64(self, fig) -> str:
        """Convert matplotlib figure to base64 string."""
        buffer = io.BytesIO()
        fig.savefig(buffer, format="png", dpi=self.plot_dpi, bbox_inches="tight")
        buffer.seek(0)
        image_base64 = base64.b64encode(buffer.getvalue()).decode()
        buffer.close()
        return image_base64

    def _create_waterfall_plot(self, waterfall_df: pd.DataFrame) -> tuple:
        """Create waterfall plot and return figure and base64 string."""
        fig, ax = plt.subplots(figsize=(10, 6))

        # Create waterfall chart
        y_pos = range(len(waterfall_df))
        remaining = waterfall_df["Remaining"].values

        bars = ax.barh(y_pos, remaining, color="steelblue", alpha=0.7)
        ax.set_yticks(y_pos)
        ax.set_yticklabels(
            [f"{row['Type']}: {row['Name']}" for _, row in waterfall_df.iterrows()]
        )
        ax.set_xlabel("Number of Patients")
        ax.set_title("Patient Attrition (Waterfall Chart)")

        # Add value labels on bars
        for i, (bar, value) in enumerate(zip(bars, remaining)):
            ax.text(
                bar.get_width() + max(remaining) * 0.01,
                bar.get_y() + bar.get_height() / 2,
                f"{int(value):,}",
                va="center",
                fontsize=9,
            )

        ax.grid(axis="x", alpha=0.3)
        plt.tight_layout()

        image_base64 = self._plot_to_base64(fig)
        return fig, image_base64

    def execute(self, cohort) -> Dict[str, Any]:
        """Execute the report generation."""
        logger.info(f"Generating comprehensive report for cohort: {cohort.name}")

        # Ensure cohort is executed
        if cohort.index_table is None:
            logger.error("Cohort not yet executed. Run cohort execution first.")

        # Generate title if not provided
        if not self.title:
            self.title = f"Study Report: {cohort.name}"

        # Store cohort name for directory creation
        self.cohort_name = cohort.name

        # STEP 1: Generate data tables first
        logger.info("=== PHASE 1: Generating Data Tables ===")

        # Generate Waterfall Table
        logger.info("Generating waterfall table...")
        # Use custom reporter if provided, otherwise create default
        if self.waterfall_reporter is not None:
            waterfall_reporter = self.waterfall_reporter
            logger.info("Using custom Waterfall reporter instance")
        else:
            waterfall_reporter = Waterfall(
                decimal_places=self.decimal_places, pretty_display=self.pretty_display
            )
            logger.info("Using default Waterfall reporter")
        waterfall_result = waterfall_reporter.execute(cohort)

        # Extract DataFrame from Styler if needed (pretty_display=True returns Styler)
        if hasattr(waterfall_result, "data"):  # It's a Styler object
            waterfall_df = waterfall_result.data
        else:
            waterfall_df = waterfall_result

        # Remove the first row (typically "N persons in database" info row)
        if not waterfall_df.empty and len(waterfall_df) > 0:
            waterfall_df = waterfall_df.iloc[1:].reset_index(drop=True)

        self.report["waterfall_table"] = waterfall_df

        # Generate Table 1 (Baseline Characteristics)
        if cohort.characteristics:
            logger.info("Generating Table 1 (baseline characteristics)...")
            # Use custom reporter if provided, otherwise create default
            if self.table1_reporter is not None:
                table1_reporter = self.table1_reporter
                logger.info("Using custom Table1 reporter instance")
            else:
                table1_reporter = Table1(
                    decimal_places=self.decimal_places, pretty_display=True
                )  # Enable pretty display for proper formatting
                logger.info("Using default Table1 reporter")
            try:
                table1_df = table1_reporter.execute(cohort)
                self.report["table1"] = table1_df
                logger.info(
                    f"Table1 generated successfully with {len(table1_df)} rows and columns: {list(table1_df.columns)}"
                )
            except Exception as e:
                logger.error(f"FATAL: Table1 generation failed: {e}")
                logger.error(f"Error type: {type(e).__name__}")

                # The Table1 reporter is a core component and should work with properly structured cohorts
                # If it's failing, the issue is likely with our mock data structure
                raise RuntimeError(
                    f"Table1 reporter failed. This indicates the cohort characteristics are not properly structured for the Table1 reporter. Original error: {e}"
                )
        else:
            logger.info("No characteristics defined. Skipping Table 1.")
            self.report["table1"] = pd.DataFrame()

        # 8. Generate Table 2 (Outcomes) if outcomes exist
        if cohort.outcomes:
            logger.info("Generating Table 2 (outcomes)...")

            # Use custom reporter if provided, otherwise create default
            if self.table2_reporter is not None:
                table2_reporter = self.table2_reporter
                logger.info("Using custom Table2 reporter instance")
            else:
                # Initialize Table2 reporter with the exposure phenotype
                table2_reporter = Table2(
                    time_points=[365],  # 1 year follow-up
                    decimal_places=self.decimal_places,
                    pretty_display=True,
                )
                logger.info("Using default Table2 reporter")

            try:
                table2_df = table2_reporter.execute(cohort)
                self.report["table2"] = table2_df
                logger.info(
                    f"Table2 generated successfully with {len(table2_df)} rows and columns: {list(table2_df.columns)}"
                )
            except Exception as e:
                logger.error(f"FATAL: Table2 generation failed: {e}")
                logger.error(f"Error type: {type(e).__name__}")

                # Table2 reporter is a core component and should work with properly structured cohorts
                raise RuntimeError(
                    f"Table2 reporter failed. This indicates the cohort structure is not compatible with the Table2 reporter. Original error: {e}"
                )
        else:
            logger.info("No outcomes defined. Skipping Table 2.")
            self.report["table2"] = pd.DataFrame()

        # Generate summary statistics
        n_patients = (
            cohort.index_table.filter(cohort.index_table.BOOLEAN == True)
            .select("PERSON_ID")
            .distinct()
            .count()
            .execute()
        )
        self.report["summary_stats"] = {
            "total_patients": n_patients,
            "n_characteristics": len(cohort.characteristics or []),
            "n_outcomes": len(cohort.outcomes or []),
            "n_inclusions": len(cohort.inclusions or []),
            "n_exclusions": len(cohort.exclusions or []),
        }

        # STEP 2: Build the global AI context after tables are generated (includes table data)
        logger.info("Building comprehensive global AI context with table data...")
        self._global_context = self._build_global_ai_context(cohort)

        # STEP 3: Generate AI-Powered Content (using class variable for global context)
        logger.info("=== PHASE 2: Generating AI-Powered Content ===")

        # Generate AI text sections using global context class variable
        logger.info("Generating AI executive summary...")
        self.report["executive_summary"] = self._create_executive_summary()

        logger.info("Generating cohort definition description...")
        self.report["cohort_definition"] = self._create_cohort_description(cohort)

        logger.info("Generating data analysis description...")
        self.report["data_analysis"] = self._create_data_analysis_description(cohort)

        logger.info("Generating study variables description...")
        self.report["study_variables"] = self._create_variables_description(cohort)

        # Generate commentary for tables and figures (AI-powered if available, fallback otherwise)
        logger.info("Generating commentary for waterfall table...")
        self.report["waterfall_commentary"] = self._generate_waterfall_commentary(
            self.report.get("waterfall_table")
        )

        logger.info("Generating commentary for Table 1...")
        self.report["table1_commentary"] = self._generate_table1_commentary(
            self.report.get("table1")
        )

        logger.info("Generating commentary for Table 2...")
        self.report["table2_commentary"] = self._generate_table2_commentary(
            self.report.get("table2")
        )

        # Generate plots if requested (with AI captions that now have full context)
        if self.include_plots and not waterfall_df.empty:
            logger.info("Generating waterfall plot...")
            logger.info(
                f"include_plots={self.include_plots}, waterfall_df.empty={waterfall_df.empty}"
            )
            fig, img_b64 = self._create_waterfall_plot(waterfall_df)
            logger.info(f"Waterfall plot created, figure type: {type(fig)}")
            self.figures["waterfall"] = {
                "figure": fig,
                "base64": img_b64,
                "caption": self._generate_ai_image_caption(
                    img_b64,
                    "Patient attrition waterfall chart showing how inclusion and exclusion criteria affected the final cohort size",
                ),
            }
            logger.info(f"Stored waterfall figure, total figures: {len(self.figures)}")
        else:
            logger.warning(
                f"Not generating waterfall plot: include_plots={self.include_plots}, waterfall_df.empty={waterfall_df.empty if 'waterfall_df' in locals() else 'waterfall_df not defined'}"
            )

        logger.info("Report generation completed successfully")
        return self.report

    def _add_markdown_content_to_doc(self, doc, content: str):
        """
        Parse markdown content and add it to Word document with proper formatting.
        Handles bold text (**text**), headings (## Heading), and bullet points.
        """
        if not content:
            return

        lines = content.split("\n")

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Handle headings
            if line.startswith("## "):
                heading_text = line[3:].strip()
                doc.add_heading(heading_text, level=2)
            elif line.startswith("### "):
                heading_text = line[4:].strip()
                doc.add_heading(heading_text, level=3)
            elif line.startswith("#### "):
                heading_text = line[5:].strip()
                doc.add_heading(heading_text, level=4)
            elif line.startswith("# "):
                heading_text = line[2:].strip()
                doc.add_heading(heading_text, level=1)
            elif line.startswith("* ") or line.startswith("- "):
                # Handle bullet points
                bullet_text = line[2:].strip()
                paragraph = doc.add_paragraph(style="List Bullet")
                self._add_formatted_text(paragraph, bullet_text)
            else:
                # Handle regular paragraphs with bold formatting
                paragraph = doc.add_paragraph()
                self._add_formatted_text(paragraph, line)

    def _add_formatted_text(self, paragraph, text: str):
        """
        Add text to a paragraph with proper formatting for bold text (**text**).
        """
        import re

        # Split text by bold markers (**text**)
        parts = re.split(r"(\*\*.*?\*\*)", text)

        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                # Bold text
                bold_text = part[2:-2]  # Remove ** markers
                run = paragraph.add_run(bold_text)
                run.bold = True
            else:
                # Regular text
                paragraph.add_run(part)

    def to_markdown(self, filename: str) -> str:
        """
        Generate a clean Markdown report file.

        Args:
            filename: Path to the Markdown file (relative or absolute, with or without .md extension)

        Returns:
            Path to the generated Markdown file
        """
        if not self.report:
            raise ValueError("No report data available. Call execute() first.")

        # Convert to Path object and ensure .md extension
        output_path = Path(filename)
        if not output_path.suffix == ".md":
            output_path = output_path.with_suffix(".md")

        # Create a dedicated directory for this cohort's report assets (figures)
        cohort_name = getattr(self, "cohort_name", "report")
        cohort_name = (
            cohort_name.replace(" ", "_").replace("/", "_").replace("\\", "_")
        )  # Clean filename
        cohort_dir = output_path.parent / f"{output_path.stem}_files"
        cohort_dir.mkdir(parents=True, exist_ok=True)

        # Also ensure parent directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)

        logger.info(f"Generating Markdown report: {output_path}")

        # Build the complete markdown content
        markdown_content = self._build_markdown_content(cohort_dir)

        # Write to file
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(markdown_content)

        logger.info(f"Markdown report generated: {output_path}")
        return str(output_path)

    def _build_markdown_content(self, cohort_dir: Path) -> str:
        """Build the complete markdown content for the report."""
        md_content = ""

        # Title and metadata
        if self.title:
            md_content += f"# {self.title}\n\n"

        if self.author:
            md_content += f"**Author:** {self.author}\n\n"
        if self.institution:
            md_content += f"**Institution:** {self.institution}\n\n"
        if hasattr(self, "date") and self.date:
            md_content += f"**Date:** {self.date}\n\n"
        else:
            md_content += (
                f"**Report Generated:** {datetime.now().strftime('%B %d, %Y')}\n\n"
            )

        # Executive Summary
        md_content += "## Executive Summary\n\n"
        if "executive_summary" in self.report:
            md_content += self.report["executive_summary"] + "\n\n"
        else:
            # Fallback summary
            stats = self.report.get("summary_stats", {})
            md_content += f"""This report presents the analysis of {stats.get('total_patients', 'N/A')} patients in the study cohort. 
The analysis includes {stats.get('n_characteristics', 0)} baseline characteristics and {stats.get('n_outcomes', 0)} outcome measures.
Cohort definition involved {stats.get('n_inclusions', 0)} inclusion criteria and {stats.get('n_exclusions', 0)} exclusion criteria.\n\n"""

        md_content += "---\n\n"

        section_number = 1

        # 1. Cohort Definition
        if "cohort_definition" in self.report:
            md_content += f"## {section_number}. Cohort Definition\n\n"
            md_content += self.report["cohort_definition"] + "\n\n"
            section_number += 1

        # 2. Data Analysis
        if "data_analysis" in self.report:
            md_content += f"## {section_number}. Data Analysis\n\n"
            md_content += self.report["data_analysis"] + "\n\n"
            section_number += 1

        # 3. Study Variables
        if "study_variables" in self.report:
            md_content += f"## {section_number}. Study Variables\n\n"
            md_content += self.report["study_variables"] + "\n\n"
            section_number += 1

        # 4. Patient Attrition (Waterfall Table)
        waterfall_check = self.report.get("waterfall_table")
        # Extract DataFrame from Styler if needed for the check
        if hasattr(waterfall_check, "data"):
            waterfall_check = waterfall_check.data
        if (
            "waterfall_table" in self.report
            and waterfall_check is not None
            and not waterfall_check.empty
        ):
            md_content += f"## {section_number}. Patient Attrition\n\n"

            # Add the waterfall figure if available
            if "waterfall" in self.figures:
                # Save figure to the cohort directory
                fig_filename = "figure_1_waterfall_plot.png"
                fig_path = cohort_dir / fig_filename
                self.figures["waterfall"]["figure"].savefig(
                    fig_path, format="png", dpi=300, bbox_inches="tight"
                )
                md_content += f"![Waterfall Plot]({fig_filename})\n\n"
                if "caption" in self.figures["waterfall"]:
                    md_content += f"*Figure {section_number}.1: {self.figures['waterfall']['caption']}*\n\n"

            # Add waterfall table
            waterfall_df = self.report["waterfall_table"]
            # Extract DataFrame from Styler if needed
            if hasattr(waterfall_df, "data"):
                waterfall_df = waterfall_df.data
            md_content += self._dataframe_to_markdown_table(waterfall_df) + "\n\n"

            if "waterfall_commentary" in self.report:
                md_content += self.report["waterfall_commentary"] + "\n\n"
            section_number += 1

        # 5. Baseline Characteristics (Table 1)
        if "table1" in self.report and not self.report["table1"].empty:
            md_content += f"## {section_number}. Baseline Characteristics\n\n"
            table1_df = self.report["table1"]
            md_content += self._dataframe_to_markdown_table(table1_df) + "\n\n"
            if "table1_commentary" in self.report:
                md_content += self.report["table1_commentary"] + "\n\n"
            section_number += 1

        # 6. Outcomes Summary (Table 2)
        if "table2" in self.report and not self.report["table2"].empty:
            md_content += f"## {section_number}. Outcomes Summary\n\n"
            table2_df = self.report["table2"]
            md_content += self._dataframe_to_markdown_table(table2_df) + "\n\n"
            if "table2_commentary" in self.report:
                md_content += self.report["table2_commentary"] + "\n\n"
            section_number += 1

        return md_content

    def _dataframe_to_markdown_table(self, df: pd.DataFrame) -> str:
        """Convert a DataFrame to a clean markdown table."""
        if df.empty:
            return "No data available."

        # Start with headers
        headers = df.columns.tolist()
        if not headers:
            return "No data available."

        # Build header row
        header_row = "| " + " | ".join(str(h) for h in headers) + " |"

        # Build separator row
        separator_row = "| " + " | ".join("---" for _ in headers) + " |"

        # Build data rows
        data_rows = []
        for _, row in df.iterrows():
            row_str = (
                "| "
                + " | ".join(str(v) if pd.notna(v) else "" for v in row.values)
                + " |"
            )
            data_rows.append(row_str)

        return "\n".join([header_row, separator_row] + data_rows)

    def to_word(self, filename: str) -> str:
        """Generate Word document report."""
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for Word document generation. Install with: pip install python-docx"
            )

        if not self.report:
            raise ValueError("No report data available. Call execute() first.")

        # Convert to Path object and ensure .docx extension
        output_path = Path(filename)
        if not output_path.suffix == ".docx":
            output_path = output_path.with_suffix(".docx")

        # Ensure parent directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)

        logger.info(f"Generating Word document: {output_path}")

        # Create Word document
        doc = Document()

        # Title
        title = doc.add_heading(self.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        doc.add_paragraph(f"Author: {self.author or 'Not specified'}")
        doc.add_paragraph(f"Institution: {self.institution or 'Not specified'}")
        doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_page_break()

        # Executive Summary - Use AI-generated content if available
        doc.add_heading("Executive Summary", level=1)
        if "executive_summary" in self.report:
            # Add AI-generated executive summary with markdown formatting
            exec_summary = self.report["executive_summary"]
            self._add_markdown_content_to_doc(doc, exec_summary)
        else:
            # Fallback to basic summary if AI summary not available
            stats = self.report.get("summary_stats", {})
            summary_text = f"""This report presents the analysis of {stats.get('total_patients', 'N/A')} patients in the study cohort. 
The analysis includes {stats.get('n_characteristics', 0)} baseline characteristics and {stats.get('n_outcomes', 0)} outcome measures.
Cohort definition involved {stats.get('n_inclusions', 0)} inclusion criteria and {stats.get('n_exclusions', 0)} exclusion criteria."""
            doc.add_paragraph(summary_text)

        # Cohort Definition
        doc.add_heading("1. Cohort Definition", level=1)
        cohort_def = self.report.get("cohort_definition", "No description available.")
        self._add_markdown_content_to_doc(doc, cohort_def)

        # Data Analysis
        doc.add_heading("2. Data Analysis", level=1)
        data_analysis = self.report.get("data_analysis", "No description available.")
        self._add_markdown_content_to_doc(doc, data_analysis)

        # Study Variables
        doc.add_heading("3. Study Variables", level=1)
        study_vars = self.report.get("study_variables", "No description available.")
        self._add_markdown_content_to_doc(doc, study_vars)

        # Waterfall Table
        waterfall_df = self.report.get("waterfall_table")
        # Extract DataFrame from Styler if needed
        if hasattr(waterfall_df, "data"):
            waterfall_df = waterfall_df.data
        if waterfall_df is not None and not waterfall_df.empty:
            doc.add_heading("4. Patient Attrition (Waterfall Table)", level=1)

            # Add table to Word document
            table = doc.add_table(rows=1, cols=len(waterfall_df.columns))
            table.style = "Table Grid"

            # Header row
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(waterfall_df.columns):
                hdr_cells[i].text = str(col)

            # Data rows
            for _, row in waterfall_df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)

            # Add waterfall plot if available
            if "waterfall" in self.figures:
                doc.add_paragraph()
                doc.add_paragraph("Figure 1: Patient Attrition Waterfall")

                # Save plot temporarily for inclusion
                temp_plot_path = output_path.parent / "temp_waterfall.png"
                # Ensure directory exists
                temp_plot_path.parent.mkdir(parents=True, exist_ok=True)
                self.figures["waterfall"]["figure"].savefig(
                    temp_plot_path, dpi=self.plot_dpi, bbox_inches="tight"
                )
                doc.add_picture(str(temp_plot_path), width=Inches(6))
                doc.add_paragraph(self.figures["waterfall"]["caption"])

            # Add AI commentary if available
            if "waterfall_commentary" in self.report:
                doc.add_heading("Clinical Commentary", level=2)
                self._add_markdown_content_to_doc(
                    doc, self.report["waterfall_commentary"]
                )

        # Table 1
        table1_df = self.report.get("table1")
        if table1_df is not None and not table1_df.empty:
            doc.add_heading("5. Baseline Characteristics (Table 1)", level=1)

            # Create table with all columns from DataFrame
            table = doc.add_table(rows=1, cols=len(table1_df.columns))
            table.style = "Table Grid"

            # Header row
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(table1_df.columns):
                hdr_cells[i].text = str(col)

            # Data rows
            for _, row in table1_df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)

            # Add AI commentary if available
            if "table1_commentary" in self.report:
                doc.add_heading("Clinical Commentary", level=2)
                self._add_markdown_content_to_doc(doc, self.report["table1_commentary"])

        # Table 2 (Outcomes)
        table2_df = self.report.get("table2")
        if table2_df is not None and not table2_df.empty:
            doc.add_heading("6. Outcomes Summary (Table 2)", level=1)

            # Create table
            table = doc.add_table(rows=1, cols=len(table2_df.columns))
            table.style = "Table Grid"

            # Header row
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(table2_df.columns):
                hdr_cells[i].text = str(col)

            # Data rows
            for _, row in table2_df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)

            # Add AI commentary if available
            if "table2_commentary" in self.report:
                doc.add_heading("Clinical Commentary", level=2)
                self._add_markdown_content_to_doc(doc, self.report["table2_commentary"])

        # Save document
        doc.save(str(output_path))

        # Clean up temporary plot file if it exists
        temp_plot_path = output_path.parent / "temp_waterfall.png"
        if temp_plot_path.exists():
            try:
                temp_plot_path.unlink()
            except:
                pass

        logger.info(f"Word document generated: {output_path}")
        return str(output_path)
        return str(output_path)

    def get_report_summary(self) -> Dict[str, Any]:
        """Get a summary of the generated report."""
        if not self.report:
            return {"error": "No report data available. Call execute() first."}

        summary = {
            "title": self.title,
            "author": self.author,
            "institution": self.institution,
            "generation_date": datetime.now().isoformat(),
            "ai_enabled": self.use_ai,
            "sections_generated": list(self.report.keys()),
            "figures_generated": list(self.figures.keys()),
            "summary_statistics": self.report.get("summary_stats", {}),
        }

        # Add table shapes
        for section_name, section_data in self.report.items():
            if isinstance(section_data, pd.DataFrame):
                summary[f"{section_name}_shape"] = section_data.shape

        return summary
