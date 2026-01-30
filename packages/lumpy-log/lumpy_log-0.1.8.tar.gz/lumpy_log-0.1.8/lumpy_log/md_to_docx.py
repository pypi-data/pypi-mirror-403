"""Custom Markdown to DOCX converter using python-docx.

This module provides a simple but effective markdown to Word document converter
with nice code block handling and no external markdown parsing dependencies.
"""

import re
import hashlib
from pathlib import Path
from typing import List, Tuple
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
# NEW: imports for mermaid rendering
import os
import tempfile
import urllib.request
import base64


def _get_cache_dir() -> Path:
    """Get or create the cache directory for code images."""
    package_dir = os.path.dirname(os.path.abspath(__file__))
    cache_dir = Path(".lumpy_cache") / "code_images"
    cache_dir.mkdir(parents=True, exist_ok=True)
    return cache_dir


def _get_cached_image(code: str, language: str, renderer: str = "hcti") -> Path:
    """Check if a cached image exists for this code block."""
    content = f"{renderer}:{language}:{code}"
    cache_key = hashlib.sha256(content.encode("utf-8")).hexdigest()

    cache_dir = _get_cache_dir()
    cached_file = cache_dir / f"{cache_key}.png"
    
    print(f"Looking for cached image at: {cached_file}")
    if cached_file.exists():
        print("Found cached image.")
        return cached_file
    return None


def _save_to_cache(code: str, language: str, image_bytes: bytes, renderer: str = "hcti") -> Path:
    """Save rendered image to cache."""
    content = f"{renderer}:{language}:{code}"
    cache_key = hashlib.sha256(content.encode("utf-8")).hexdigest()

    cache_dir = _get_cache_dir()
    cached_file = cache_dir / f"{cache_key}.png"

    with open(cached_file, "wb") as f:
        f.write(image_bytes)

    return cached_file


def _add_code_block_formatting(paragraph, background_color: Tuple[int, int, int] = (240, 240, 240), apply_text_color: bool = True):
    """Add background color and monospace formatting to a paragraph for code blocks.
    
    Args:
        paragraph: docx paragraph object
        background_color: RGB tuple for background color (default: light gray)
        apply_text_color: If True, apply text color to code blocks (default: True)
    """
    # Set monospace font with console-style formatting
    for run in paragraph.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)  # Slightly smaller for better code density
        if apply_text_color:
            run.font.color.rgb = RGBColor(30, 30, 30)
    
    # Add background color using shading
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '%02x%02x%02x' % background_color)
    paragraph._element.get_or_add_pPr().append(shading_elm)
    
    # Set left margin/indent
    paragraph.paragraph_format.left_indent = Inches(0.25)
    # Reduced line spacing for console appearance
    paragraph.paragraph_format.line_spacing = 1.0
    # Minimize spacing between code lines
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


# NEW: syntax-highlighted code block rendering (optional pygments)
def _add_highlighted_code_block(doc, code: str, language: str = ""):
    """Add a code block with syntax highlighting using pygments if available."""
    try:
        from pygments import lex
        from pygments.lexers import get_lexer_by_name, TextLexer
        from pygments.styles import get_style_by_name
        from pygments.token import Token
    except Exception:
        # Fallback to plain code block formatting
        for line in code.split('\n'):
            para = doc.add_paragraph(line or ' ', style='Normal')
            _add_code_block_formatting(para)
        return

    try:
        lexer = get_lexer_by_name(language) if language else TextLexer()
    except Exception:
        lexer = TextLexer()

    style = get_style_by_name('friendly')

    def _parse_style(style_str: str):
        attrs = {'color': None, 'bold': False, 'italic': False}
        if not style_str:
            return attrs
        for part in style_str.split():
            if part == 'bold':
                attrs['bold'] = True
            elif part == 'italic':
                attrs['italic'] = True
            elif part.startswith('#') or re.match(r'^[0-9a-fA-F]{6}$', part):
                hexval = part if part.startswith('#') else f'#{part}'
                attrs['color'] = hexval
        return attrs

    def _style_for(tt):
        t = tt
        while t and t not in style.styles:
            t = t.parent
        return _parse_style(style.styles.get(t, ''))

    # Build paragraphs line-by-line to keep shading per line
    paragraph = doc.add_paragraph(style='Normal')
    _add_code_block_formatting(paragraph, apply_text_color=False)

    for tok_type, tok_val in lex(code, lexer):
        pieces = tok_val.split('\n')
        for idx, piece in enumerate(pieces):
            if piece:
                run = paragraph.add_run(piece)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)  # Match reduced size from _add_code_block_formatting
                attrs = _style_for(tok_type)
                if attrs['bold']:
                    run.bold = True
                if attrs['italic']:
                    run.italic = True
                if attrs['color']:
                    hexval = attrs['color'].lstrip('#')
                    r = int(hexval[0:2], 16)
                    g = int(hexval[2:4], 16)
                    b = int(hexval[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
            if idx < len(pieces) - 1:
                # newline => start a new shaded paragraph
                paragraph = doc.add_paragraph(style='Normal')
                _add_code_block_formatting(paragraph, apply_text_color=False)


# NEW: mermaid rendering via Kroki
def _render_mermaid_and_insert(doc, mermaid_code: str, width_inches: float = 6.0):
    """Render mermaid to PNG via Kroki and insert into the document."""
    try:
        req = urllib.request.Request(
            url='https://kroki.io/mermaid/png',
            data=mermaid_code.encode('utf-8'),
            headers={'Content-Type': 'text/plain'}
        )
        with urllib.request.urlopen(req, timeout=30) as resp:
            img_bytes = resp.read()

        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
        try:
            tmp.write(img_bytes)
            tmp.close()
            doc.add_picture(tmp.name, width=Inches(width_inches))
        finally:
            os.unlink(tmp.name)
    except Exception:
        # Fallback: insert raw code block
        para = doc.add_paragraph(mermaid_code, style='Normal')
        _add_code_block_formatting(para)


def _render_code_with_playwright(html_content: str, width_inches: float = 6.0) -> bytes:
    """Render HTML to PNG using Playwright (local, no API credentials required).

    Returns PNG bytes on success, or None on failure/missing dependency.
    """
    debug = os.environ.get('LUMPY_DEBUG') == '1'
    
    try:
        from playwright.sync_api import sync_playwright
    except ImportError as e:
        if debug:
            print(f"[DEBUG] Playwright not available: {e}")
        return None

    width_px = max(int(width_inches * 96), 640)
    
    if debug:
        print(f"[DEBUG] Rendering code block with Playwright")
        print(f"[DEBUG] Width: {width_inches} inches ({width_px} px)")
        print(f"[DEBUG] HTML content length: {len(html_content)} chars")

    try:
        with sync_playwright() as p:
            browser_name = os.environ.get("LUMPY_PLAYWRIGHT_BROWSER", "chromium").lower()
            if debug:
                print(f"[DEBUG] Launching browser: {browser_name}")
            
            # Try to use system chromium first, then fall back to bundled
            launch_kwargs = {"args": ["--no-sandbox"]}
            
            # If system chromium exists, use it
            system_chromium = "/usr/bin/chromium-browser"
            if os.path.exists(system_chromium):
                if debug:
                    print(f"[DEBUG] Using system chromium: {system_chromium}")
                launch_kwargs["executable_path"] = system_chromium
            
            browser = {
                "chromium": p.chromium,
                "firefox": p.firefox,
                "webkit": p.webkit,
            }.get(browser_name, p.chromium).launch(**launch_kwargs)
            
            if debug:
                print(f"[DEBUG] Browser launched successfully")
            
            page = browser.new_page(viewport={"width": width_px, "height": 10})
            
            if debug:
                print(f"[DEBUG] Page created with initial viewport: {width_px}x10")
            
            try:
                page.set_content(html_content, wait_until="load")
                if debug:
                    print(f"[DEBUG] HTML content set on page")
            except Exception as e:
                if debug:
                    print(f"[DEBUG] Error setting page content: {e}")
                raise
            
            # Resize to the content height for a tight screenshot
            try:
                box = page.evaluate("""
                    () => {
                        const pre = document.querySelector('pre');
                        const rect = pre ? pre.getBoundingClientRect() : document.body.getBoundingClientRect();
                        return { 
                            width: Math.ceil(rect.width) + 40, 
                            height: Math.ceil(rect.height) + 40,
                            preFound: !!pre
                        };
                    }
                """)
                if debug:
                    print(f"[DEBUG] Content box measurements: {box}")
            except Exception as e:
                if debug:
                    print(f"[DEBUG] Error evaluating content size: {e}")
                box = {"width": width_px, "height": 400, "preFound": False}
            
            final_width = max(width_px, box.get("width", width_px))
            final_height = box.get("height", 400)
            
            if debug:
                print(f"[DEBUG] Setting final viewport to: {final_width}x{final_height}")
            
            page.set_viewport_size({"width": final_width, "height": final_height})
            
            # Add a small delay to ensure rendering is complete
            page.wait_for_load_state("networkidle")
            
            if debug:
                print(f"[DEBUG] Taking screenshot...")
            
            img_bytes = page.screenshot(full_page=True)
            
            if debug:
                print(f"[DEBUG] Screenshot captured: {len(img_bytes)} bytes")
            
            browser.close()
            return img_bytes
    except Exception as e:
        if debug:
            print(f"[DEBUG] Error in Playwright rendering: {type(e).__name__}: {e}")
            import traceback
            traceback.print_exc()
        return None


def _render_code_as_image_and_insert(doc, code: str, language: str = "", width_inches: float = 6.0, hcti_user_id: str = None, hcti_api_key: str = None, use_cache: bool = True):
    """Render code block to PNG (HCTI if available, otherwise local Playwright) and insert.
    Falls back to text rendering if both fail.
    """
    debug = os.environ.get('LUMPY_DEBUG') == '1'
    
    api_user_id = hcti_user_id or os.environ.get('HCTI_API_USER_ID')
    api_key = hcti_api_key or os.environ.get('HCTI_API_KEY')

    # Determine which renderer will be used
    will_use_hcti = bool(api_user_id and api_key)
    renderer = "hcti" if will_use_hcti else "playwright"
    
    if debug:
        print(f"\n[DEBUG] _render_code_as_image_and_insert called")
        print(f"[DEBUG] Language: {language}")
        print(f"[DEBUG] Renderer: {renderer}")
        print(f"[DEBUG] Code length: {len(code)} chars")
        print(f"[DEBUG] Use cache: {use_cache}")

    # Cache check first (with renderer-specific key)
    if use_cache:
        cached_image = _get_cached_image(code, language, renderer)
        if cached_image:
            if debug:
                print(f"[DEBUG] Using cached image: {cached_image}")
            try:
                doc.add_picture(str(cached_image), width=Inches(width_inches))
                return
            except Exception as e:
                if debug:
                    print(f"[DEBUG] Error adding cached picture: {e}")

    # Build HTML once (offline-friendly, no external assets)
    html_content = f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8" />
    <style>
        :root {{
            --bg: #f6f8fa;
            --fg: #24292f;
            --border: #d0d7de;
        }}
        * {{ box-sizing: border-box; }}
        body {{
            margin: 0;
            padding: 20px;
            font-family: 'Courier New', monospace;
            background: white;
        }}
        pre {{
            margin: 0;
            padding: 16px;
            background: var(--bg);
            border: 1px solid var(--border);
            border-radius: 6px;
            overflow-x: auto;
            color: var(--fg);
            font-size: 14px;
            line-height: 1.5;
            white-space: pre;
        }}
        code {{
            font-family: 'Courier New', monospace;
            font-size: 14px;
            line-height: 1.5;
        }}
    </style>
</head>
<body>
    <pre><code>{code.replace('<', '&lt;').replace('>', '&gt;')}</code></pre>
</body>
</html>'''

    if debug:
        print(f"[DEBUG] Generated HTML: {len(html_content)} chars")

    img_bytes = None

    # Try HCTI if credentials are present
    if api_user_id and api_key:
        if debug:
            print(f"[DEBUG] Attempting HCTI rendering...")
        try:
            import json

            data = json.dumps({
                'html': html_content,
                'css': '',
                'google_fonts': 'Courier New'
            }).encode('utf-8')

            credentials = base64.b64encode(f'{api_user_id}:{api_key}'.encode('utf-8')).decode('utf-8')
            req = urllib.request.Request(
                url='https://hcti.io/v1/image',
                data=data,
                headers={
                    'Content-Type': 'application/json',
                    'Authorization': f'Basic {credentials}'
                }
            )

            with urllib.request.urlopen(req, timeout=30) as resp:
                result = json.loads(resp.read().decode('utf-8'))
                img_url = result.get('url')
                if debug:
                    print(f"[DEBUG] HCTI response URL: {img_url}")
                if img_url:
                    with urllib.request.urlopen(img_url, timeout=30) as img_resp:
                        img_bytes = img_resp.read()
                        if debug:
                            print(f"[DEBUG] HCTI image downloaded: {len(img_bytes)} bytes")
        except Exception as e:
            if debug:
                print(f"[DEBUG] HCTI rendering failed: {type(e).__name__}: {e}")
            img_bytes = None

    # If no HCTI or it failed, try local Playwright rendering
    if img_bytes is None:
        if debug:
            print(f"[DEBUG] Attempting Playwright rendering...")
        img_bytes = _render_code_with_playwright(html_content, width_inches=width_inches)

    # If everything failed, fall back to text rendering
    if img_bytes is None:
        if debug:
            print(f"[DEBUG] All image rendering failed, falling back to text rendering")
        _add_highlighted_code_block(doc, code, language)
        return

    if debug:
        print(f"[DEBUG] Image rendering successful: {len(img_bytes)} bytes")

    # Persist to cache if enabled, otherwise use a temp file
    if use_cache:
        try:
            cached_path = _save_to_cache(code, language, img_bytes, renderer)
            if debug:
                print(f"[DEBUG] Cached to: {cached_path}")
            doc.add_picture(str(cached_path), width=Inches(width_inches))
            return
        except Exception as e:
            if debug:
                print(f"[DEBUG] Error caching image: {e}")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
    try:
        tmp.write(img_bytes)
        tmp.close()
        if debug:
            print(f"[DEBUG] Adding picture from temp file: {tmp.name}")
        doc.add_picture(tmp.name, width=Inches(width_inches))
    finally:
        os.unlink(tmp.name)


def _sanitize_xml_string(s: str) -> str:
    """Remove characters not allowed in XML 1.0 (python-docx requirement)."""
    # Remove control characters except for tab, newline, carriage return
    return ''.join(
        c for c in s
        if (
            c == '\t' or c == '\n' or c == '\r' or
            (0x20 <= ord(c) <= 0xD7FF) or
            (0xE000 <= ord(c) <= 0xFFFD) or
            (0x10000 <= ord(c) <= 0x10FFFF)
        )
    )


def markdown_to_docx(markdown_content: str, output_path: str, render_code_as_images: bool = False, hcti_user_id: str = None, hcti_api_key: str = None) -> bool:
    """Convert markdown content to DOCX file.
    
    Args:
        markdown_content: Markdown text to convert
        output_path: Path where to save the DOCX file
        render_code_as_images: If True, render code blocks as images using HCTI API or Playwright.
                              Note: Playwright rendering requires proper X11 environment and system dependencies.
                              Recommended to use HCTI API credentials for reliable image rendering.
        hcti_user_id: HCTI API User ID (optional, falls back to environment variable)
        hcti_api_key: HCTI API Key (optional, falls back to environment variable)
    
    Returns:
        True if conversion succeeded, False otherwise
    """
    # Sanitize input for XML compatibility
    markdown_content = _sanitize_xml_string(markdown_content)
    doc = Document()
    
    # Configure page size and margins
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4 height
    section.page_width = Inches(8.27)    # A4 width
    
    # Set narrow margins (0.5 inches)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # Track if we're in a code block
    in_code_block = False
    code_block_lines = []
    code_language = ""
    
    lines = markdown_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Check for code block start
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                in_code_block = False
                if code_block_lines:
                    code_content = '\n'.join(code_block_lines)
                    lang = (code_language or '').strip().lower()
                    if lang == 'mermaid':
                        _render_mermaid_and_insert(doc, code_content)
                    elif render_code_as_images:
                        _render_code_as_image_and_insert(doc, code_content, lang, hcti_user_id=hcti_user_id, hcti_api_key=hcti_api_key)
                    else:
                        _add_highlighted_code_block(doc, code_content, lang)
                code_block_lines = []
                code_language = ""
            else:
                # Start of code block
                in_code_block = True
                code_language = line.strip()[3:].strip()
                code_block_lines = []
            i += 1
            continue
        
        # If in code block, accumulate lines
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Handle headers
        header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if header_match:
            level = len(header_match.group(1))
            title = header_match.group(2).strip()
            style = f'Heading {level}'
            _add_formatted_paragraph(doc, title, style=style)
            i += 1
            continue
        
        # Handle horizontal rules
        if re.match(r'^[-_*]{3,}$', line.strip()):
            paragraph = doc.add_paragraph()
            pPr = paragraph._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '12')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue
        
        # Handle unordered lists
        list_match = re.match(r'^(\s*)[-*+]\s+(.+)$', line)
        if list_match:
            indent_level = len(list_match.group(1)) // 2
            list_text = list_match.group(2).strip()
            para = _add_formatted_paragraph(doc, list_text, style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.25 + indent_level * 0.25)
            i += 1
            continue
        
        # Handle ordered lists
        ordered_list_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if ordered_list_match:
            indent_level = len(ordered_list_match.group(1)) // 2
            list_text = ordered_list_match.group(2).strip()
            para = _add_formatted_paragraph(doc, list_text, style='List Number')
            para.paragraph_format.left_indent = Inches(0.25 + indent_level * 0.25)
            i += 1
            continue
        
        # Handle regular paragraphs
        # Accumulate consecutive non-empty lines as a paragraph
        paragraph_lines = [line.strip()]
        i += 1
        
        while i < len(lines) and lines[i].strip() and not re.match(r'^#+', lines[i]) and not re.match(r'^```', lines[i]) and not re.match(r'^(\s*)[-*+\d]\s', lines[i]):
            paragraph_lines.append(lines[i].strip())
            i += 1
        
        full_text = ' '.join(paragraph_lines)
        if full_text:
            _add_formatted_paragraph(doc, full_text, style='Normal')
    
    # Handle unclosed code block
    if in_code_block and code_block_lines:
        code_content = '\n'.join(code_block_lines)
        lang = (code_language or '').strip().lower()
        if lang == 'mermaid':
            _render_mermaid_and_insert(doc, code_content)
        else:
            _add_highlighted_code_block(doc, code_content, lang)
    
    # Save document
    try:
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error saving DOCX file: {e}")
        return False


def markdown_file_to_docx(markdown_file: str, output_file: str = None, render_code_as_images: bool = False, hcti_user_id: str = None, hcti_api_key: str = None) -> bool:
    """Convert a markdown file to DOCX.
    
    Args:
        markdown_file: Path to markdown file
        output_file: Path for output DOCX file (defaults to same name with .docx extension)
        render_code_as_images: If True, render code blocks as images using HCTI API
        hcti_user_id: HCTI API User ID (optional, falls back to environment variable)
        hcti_api_key: HCTI API Key (optional, falls back to environment variable)
    
    Returns:
        True if conversion succeeded, False otherwise
    """
    md_path = Path(markdown_file)
    
    if not md_path.exists():
        print(f"Error: Markdown file not found: {md_path}")
        return False
    
    if output_file is None:
        output_file = str(md_path.with_suffix('.docx'))
    
    try:
        content = md_path.read_text(encoding='utf-8')
        content = _sanitize_xml_string(content)
        return markdown_to_docx(content, output_file, render_code_as_images, hcti_user_id, hcti_api_key)
    except Exception as e:
        print(f"Error reading markdown file: {e}")
        return False


def _parse_inline_formatting(text: str) -> List[Tuple[str, bool, bool, bool]]:
    """Parse text for bold, italic, and inline code."""
    segments = []
    pos = 0
    for match in re.finditer(r'\*\*(.+?)\*\*|\*(.+?)\*|__(.+?)__|_(.+?)_|`(.+?)`', text):
        if match.start() > pos:
            segments.append((text[pos:match.start()], False, False, False))
        if match.group(1):
            segments.append((match.group(1), True, False, False))
        elif match.group(2):
            segments.append((match.group(2), False, True, False))
        elif match.group(3):
            segments.append((match.group(3), True, False, False))
        elif match.group(4):
            segments.append((match.group(4), False, True, False))
        elif match.group(5):
            segments.append((match.group(5), False, False, True))
        pos = match.end()
    if pos < len(text):
        segments.append((text[pos:], False, False, False))
    return segments if segments else [(text, False, False, False)]


def _add_formatted_paragraph(doc, text: str, style: str = 'Normal', is_code_inline: bool = False):
    """Add a paragraph with inline formatting (bold, italic, code)."""
    paragraph = doc.add_paragraph(style=style)
    segments = _parse_inline_formatting(text)
    for seg_text, is_bold, is_italic, is_inline_code in segments:
        run = paragraph.add_run(seg_text)
        run.bold = is_bold
        run.italic = is_italic
        if is_inline_code or is_code_inline:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(200, 0, 0)
    return paragraph
