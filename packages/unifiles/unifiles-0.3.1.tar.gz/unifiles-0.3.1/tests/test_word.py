"""Word 模块测试用例。"""

import pytest
from pathlib import Path
from docx import Document

from unifiles.word import read_docx, write_docx
from unifiles.exceptions import FileReadError, FileWriteError


def test_read_docx_success(tmp_path: Path):
    """测试成功读取 Word 文档。"""
    # 创建测试文件
    test_file = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("这是第一段。")
    doc.add_paragraph("这是第二段。")
    doc.add_paragraph("这是第三段。")
    doc.save(test_file)

    # 测试读取
    result = read_docx(str(test_file))
    assert isinstance(result, str)
    assert "第一段" in result
    assert "第二段" in result
    assert "第三段" in result
    # 验证段落之间用换行符分隔
    lines = result.split("\n")
    assert len(lines) == 3


def test_read_docx_empty_paragraphs(tmp_path: Path):
    """测试读取包含空段落的文档。"""
    test_file = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("第一段")
    doc.add_paragraph("")  # 空段落
    doc.add_paragraph("第二段")
    doc.save(test_file)

    result = read_docx(str(test_file))
    lines = result.split("\n")
    # 空段落应该被跳过
    assert len(lines) == 2
    assert "第一段" in result
    assert "第二段" in result


def test_read_docx_file_not_found():
    """测试文件不存在的情况。"""
    with pytest.raises(FileNotFoundError, match="文件不存在"):
        read_docx("nonexistent.docx")


def test_read_docx_invalid_format(tmp_path: Path):
    """测试无效格式文件。"""
    # 创建一个非 .docx 文件
    test_file = tmp_path / "test.txt"
    test_file.write_text("这不是 Word 文档")

    # python-docx 会抛出异常，应该被捕获为 FileReadError
    with pytest.raises(FileReadError):
        read_docx(str(test_file))


def test_write_docx_success(tmp_path: Path):
    """测试正常写入。"""
    test_file = tmp_path / "output.docx"
    content = "这是第一行\n这是第二行\n这是第三行"

    write_docx(content, str(test_file))

    # 验证文件已创建
    assert test_file.exists()

    # 验证内容
    result = read_docx(str(test_file))
    assert "第一行" in result
    assert "第二行" in result
    assert "第三行" in result


def test_write_docx_with_title(tmp_path: Path):
    """测试带标题写入。"""
    test_file = tmp_path / "output.docx"
    content = "这是文档内容。"
    title = "我的文档标题"

    write_docx(content, str(test_file), title=title)

    # 验证文件已创建
    assert test_file.exists()

    # 验证内容包含标题和正文
    result = read_docx(str(test_file))
    # 标题会被读取为段落，所以应该包含标题文本
    assert title in result or "我的文档" in result
    assert "文档内容" in result


def test_write_docx_without_title(tmp_path: Path):
    """测试不带标题写入。"""
    test_file = tmp_path / "output.docx"
    content = "这是内容。"

    write_docx(content, str(test_file))

    # 验证文件已创建
    assert test_file.exists()

    # 验证内容
    result = read_docx(str(test_file))
    assert "内容" in result


def test_write_docx_empty_content(tmp_path: Path):
    """测试写入空内容。"""
    test_file = tmp_path / "output.docx"
    content = ""

    write_docx(content, str(test_file))

    # 验证文件已创建
    assert test_file.exists()

    # 验证可以读取（即使内容为空）
    result = read_docx(str(test_file))
    assert isinstance(result, str)


def test_write_docx_invalid_content(tmp_path: Path):
    """测试无效内容格式。"""
    test_file = tmp_path / "output.docx"

    # 测试非字符串类型
    with pytest.raises(ValueError, match="内容格式无效"):
        write_docx(123, str(test_file))  # type: ignore

    with pytest.raises(ValueError, match="内容格式无效"):
        write_docx(None, str(test_file))  # type: ignore


def test_write_docx_multiline_content(tmp_path: Path):
    """测试多行内容写入。"""
    test_file = tmp_path / "output.docx"
    content = "第一行\n\n第二行\n第三行"

    write_docx(content, str(test_file))

    result = read_docx(str(test_file))
    lines = result.split("\n")
    # 验证多行内容被正确写入
    assert len(lines) >= 3
