"""Word 文档操作模块。

提供 Word 文档的读取和写入功能。
"""

from pathlib import Path

from docx import Document

from .exceptions import FileReadError, FileWriteError


def read_docx(file_path: str) -> str:
    """读取 Word 文档内容。

    Args:
        file_path: Word 文档路径

    Returns:
        文档的文本内容，段落之间用换行符分隔

    Raises:
        FileNotFoundError: 文件不存在
        FileReadError: 读取文件时发生错误

    Example:
        >>> content = read_docx("document.docx")
        >>> print(content)
        >>> # 输出文档的所有文本内容
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    try:
        document = Document(file_path)
        # 收集所有段落的文本
        paragraphs_text: list[str] = []
        for para in document.paragraphs:
            if para.text.strip():  # 跳过空段落
                paragraphs_text.append(para.text)
        # 使用换行符连接段落
        return "\n".join(paragraphs_text)
    except FileNotFoundError:
        raise
    except Exception as e:
        raise FileReadError(f"读取 Word 文档失败: {e}") from e


def write_docx(content: str, file_path: str, title: str | None = None) -> None:
    """将内容写入 Word 文档。

    创建新的 Word 文档并写入内容。如果提供了标题，会将标题作为文档标题添加。

    Args:
        content: 要写入的文本内容
        file_path: 输出 Word 文档路径
        title: 可选的文档标题

    Raises:
        ValueError: 内容格式无效
        PermissionError: 文件权限不足
        FileWriteError: 写入文件时发生错误

    Example:
        >>> # 写入简单内容
        >>> write_docx("Hello World", "output.docx")
        >>> # 写入带标题的内容
        >>> write_docx("This is the content.", "output.docx", title="My Document")
    """
    if not isinstance(content, str):
        raise ValueError(f"内容格式无效，期望 str，实际类型: {type(content)}")

    try:
        document = Document()

        # 如果提供了标题，添加为文档标题
        if title is not None:
            document.add_heading(title, level=0)

        # 将内容按换行符分割，每行作为一个段落
        lines = content.split("\n")
        for line in lines:
            document.add_paragraph(line)

        document.save(file_path)
    except PermissionError:
        raise
    except Exception as e:
        raise FileWriteError(f"写入 Word 文档失败: {e}") from e
