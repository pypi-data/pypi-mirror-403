# Copyright © 2024 Ahsan Saeed
# Licensed under the Apache License, Version 2.0
# See LICENSE and NOTICE files for details.

"""
Word export service for deterministic document export.

This module provides functionality to export structured content (markdown or plain text)
into Word (.docx) templates. The export process is strictly non-generative: it does
not call any LLMs and does not modify wording. It only maps existing content into
Word document structures.

Key features:
- Scalar placeholder replacement (e.g., {{document_id}}, {{title}})
- Block placeholder replacement with structured content ({{summary}}, {{proposal}}, etc.)
- Markdown parsing and conversion to Word structures (headings, lists, tables)
- Manual list rendering: bullet and numbered lists use deterministic glyph/number insertion
  (no Word numbering XML) for stable, cross-platform rendering
- Plain text mode fallback (paragraphs only, no structure inference)
- Combined markdown file export (exported_markdown_content.md) for content analysis

List Rendering Approach:
- Bullet lists: Manual glyph insertion with configurable glyph sets and indentation
  policies defined by ListRenderConfig. Glyphs are inserted directly into paragraph text
  with manual indentation. Maximum visual depth is configurable via ListRenderConfig.
  Deep nesting behavior beyond max_visual_depth is deterministic and policy-driven
  (clamp_last, cycle, or textual strategies).
- Numbered lists: Manual hierarchical numbering (1., 1.1., 1.1.1.) with Python-based
  counter tracking per list block. Numbers reset for each new list block.
- Both list types use configurable manual paragraph indentation (default: 0.25" per level,
  -0.25" hanging indent, configurable via ListRenderConfig)

List Semantics:
- List rendering is block-based, not AST-linked. Each markdown list block is processed
  independently and rendered as a separate visual block in Word.
- Nested lists are preserved visually (indentation and glyphs reflect nesting depth),
  not structurally (no Word list object relationships).
- Mixed bullet ↔ numbered lists are rendered as separate blocks by design. This is a
  deliberate stability decision to ensure deterministic output and cross-platform
  consistency.

Table Behavior:
- Inline markdown formatting inside table cells is treated as literal text.
- Tables prioritize structure and layout determinism over inline formatting.

Deep Nesting Behavior:
- Very deep nesting (beyond max_visual_depth) may reduce readability due to Word layout
  constraints. This is a limitation of Word document layout, not the export engine.
- The engine handles deep nesting deterministically according to the configured strategy
  (clamp_last, cycle, or textual), but visual clarity may degrade with extreme nesting.

The module attempts to respect content fidelity: text is preserved when structure
cannot be rendered (text-preserving fallback), with structural transformations
attempted where structurally safe (markdown syntax → Word objects).
"""
import json
import logging
import os
from dataclasses import dataclass
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import List, Optional, Dict, Any, Iterable, Tuple, Union

# Try to load .env file if python-dotenv is available
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    # python-dotenv not available, skip .env loading
    pass

from docx import Document  # type: ignore
from docx.document import Document as DocumentType  # type: ignore
from docx.text.paragraph import Paragraph  # type: ignore
from docx.shared import Inches  # type: ignore

from markdown_it import MarkdownIt  # type: ignore

from docx_template_export.models.export_config import ListRenderConfig
from docx_template_export.models.export_models import WordExportRequest
from docx_template_export.models.export_summary import (
    ExportSummary,
    ScalarReplacement,
    BlockReplacement,
    SkippedItem,
    FallbackEvent,
    WarningEvent,
)
from docx_template_export.services.output_path import get_output_dir_for_document, ensure_output_dir_exists

logger = logging.getLogger(__name__)


class Region(Enum):
    """Document region where content is located."""
    BODY = "BODY"
    HEADER = "HEADER"
    FOOTER = "FOOTER"


class Container(Enum):
    """Container type for content."""
    FLOW = "FLOW"  # Normal paragraphs
    TABLE = "TABLE"  # Table cells
    TEXTBOX = "TEXTBOX"  # Textboxes (w:txbxContent)


@dataclass
class FormattedRun:
    """
    Represents a text segment with inline formatting (bold, italic).
    
    This dataclass is used to preserve formatting information when parsing
    markdown inline elements. It allows the library to maintain bold/italic
    formatting when converting markdown to Word document structures.
    
    Attributes:
        text: The text content of this run segment.
        bold: If True, this text segment should be rendered as bold.
        italic: If True, this text segment should be rendered as italic.
            Both bold and italic can be True simultaneously (bold-italic).
    
    Note:
        FormattedRun objects are used internally during markdown parsing
        and are not part of the public API. They are converted to Word
        formatting when rendering to the document.
    """
    text: str
    bold: bool = False
    italic: bool = False


@dataclass
class MarkdownBlock:
    """
    Simple normalized representation of a markdown block.

    For bullet lists and numbered lists, ``items`` contains a list of
    (level, text, formatted_runs) tuples, where level starts at 1 for
    top-level items, 2 for nested items, etc. This allows us to preserve
    nested list structure and formatting when exporting to Word.
    
    For headings and paragraphs, ``formatted_runs`` contains a list of
    FormattedRun objects that preserve bold/italic formatting. If None,
    falls back to plain ``text``.
    """

    type: str  # "heading" | "paragraph" | "bullet_list" | "numbered_list" | "table"
    text: Optional[str] = None  # for headings and paragraphs (fallback if formatted_runs not available)
    formatted_runs: Optional[List[FormattedRun]] = None  # for headings and paragraphs with formatting
    level: Optional[int] = None  # for headings
    items: Optional[List[tuple[int, str, Optional[List[FormattedRun]]]]] = None  # for bullet_list and numbered_list (level, text, formatted_runs)
    header: Optional[List[str]] = None  # for table header row
    rows: Optional[List[List[str]]] = None  # for table body rows


def _create_markdown_parser() -> MarkdownIt:
    """
    Create a markdown-it parser instance configured for commonmark with
    tables and lists enabled.
    """
    md = MarkdownIt("commonmark").enable("table").enable("list")
    return md


def _parse_inline_formatting(inline_token) -> List[FormattedRun]:
    """
    Parse inline markdown tokens to extract text with formatting (bold, italic).
    
    Returns a list of FormattedRun objects representing text segments with
    their formatting applied.
    
    Note on newline behavior:
    - Newlines from <w:br/> and <w:cr/> are preserved textually as \\n characters
    - These newlines may not render as visual Word line breaks (Word requires <w:br/> XML)
    - This is intentional and deterministic: text is preserved, no text is lost
    - Visual line break rendering is not attempted to maintain deterministic behavior
    """
    runs: List[FormattedRun] = []
    if not inline_token or not hasattr(inline_token, 'children'):
        return runs
    
    children = inline_token.children or []
    
    def parse_children(start_idx: int, bold_context: bool = False, italic_context: bool = False) -> int:
        """
        Recursively parse children tokens, tracking bold/italic context.
        Returns the index after processing.
        """
        idx = start_idx
        while idx < len(children):
            child = children[idx]
            
            if child.type == "text":
                if child.content:
                    runs.append(FormattedRun(
                        text=child.content,
                        bold=bold_context,
                        italic=italic_context
                    ))
                idx += 1
            
            elif child.type == "code_inline":
                if child.content:
                    runs.append(FormattedRun(
                        text=child.content,
                        bold=bold_context,
                        italic=italic_context
                    ))
                idx += 1
            
            elif child.type == "softbreak":
                runs.append(FormattedRun(
                    text=" ",
                    bold=bold_context,
                    italic=italic_context
                ))
                idx += 1
            
            elif child.type == "hardbreak":
                runs.append(FormattedRun(
                    text="\n",
                    bold=bold_context,
                    italic=italic_context
                ))
                idx += 1
            
            elif child.type == "image":
                # Extract alt text from content or attrs
                alt_text = ""
                if child.content:
                    alt_text = child.content
                elif hasattr(child, 'attrs') and child.attrs:
                    # markdown-it stores alt text in attrs array: [alt, href, title]
                    if isinstance(child.attrs, (list, tuple)) and len(child.attrs) > 0:
                        alt_text = str(child.attrs[0]) if child.attrs[0] else ""
                if alt_text:
                    runs.append(FormattedRun(
                        text=alt_text,
                        bold=bold_context,
                        italic=italic_context
                    ))
                idx += 1
            
            elif child.type == "strong_open":
                # Enter bold context
                idx += 1
                # Check for nested em (bold-italic)
                if idx < len(children) and children[idx].type == "em_open":
                    idx += 1
                    idx = parse_children(idx, bold_context=True, italic_context=True)
                    if idx < len(children) and children[idx].type == "em_close":
                        idx += 1
                else:
                    idx = parse_children(idx, bold_context=True, italic_context=italic_context)
                
                # Skip strong_close
                if idx < len(children) and children[idx].type == "strong_close":
                    idx += 1
            
            elif child.type == "em_open":
                # Enter italic context
                idx += 1
                # Check for nested strong (bold-italic)
                if idx < len(children) and children[idx].type == "strong_open":
                    idx += 1
                    idx = parse_children(idx, bold_context=True, italic_context=True)
                    if idx < len(children) and children[idx].type == "strong_close":
                        idx += 1
                else:
                    idx = parse_children(idx, bold_context=bold_context, italic_context=True)
                
                # Skip em_close
                if idx < len(children) and children[idx].type == "em_close":
                    idx += 1
            
            elif child.type in ("strong_close", "em_close"):
                # Exit formatting context - return to caller
                return idx
            
            else:
                # Unknown token - skip
                idx += 1
        
        return idx
    
    parse_children(0)
    return runs


def parse_markdown_to_blocks(text: str) -> List[MarkdownBlock]:
    """
    Parse markdown text into a list of structural blocks for block-based rendering.
    
    This function is the primary markdown parser used by the block-based renderer.
    It converts markdown text into a normalized list of MarkdownBlock objects that
    can be rendered to Word document structures.
    
    Blocks supported:
        - heading: Markdown headings (H1-H6) with preserved formatting
        - paragraph: Plain text paragraphs with preserved formatting
        - bullet_list: Unordered lists with nested item support
        - numbered_list: Ordered lists with hierarchical numbering
        - table: Markdown tables with header and body rows
    
    Safety Guarantees:
        - This function MUST NOT change wording - it only discovers structure
        - Text content is preserved as text when structure cannot be parsed (text-preserving fallback)
        - Inline formatting (bold, italic) is preserved through FormattedRun objects where detected
        - Nested list structures are preserved with level information where supported
        - Continuation paragraphs within list items are preserved where detected
    
    Fallback Behavior:
        On any parsing error, the function falls back to plain paragraphs split
        on blank lines. This attempts to ensure export completion,
        prioritizing text preservation over structural complexity.
    
    Args:
        text: Markdown text to parse. Can be empty or None (returns empty list).
    
    Returns:
        List of MarkdownBlock objects representing the parsed structure.
        Blocks are in the same order as they appear in the markdown text.
        Empty list if input is empty or None.
    
    Example:
        ```python
        markdown = "# Heading\\n\\nParagraph text\\n\\n- Item 1\\n- Item 2"
        blocks = parse_markdown_to_blocks(markdown)
        # Returns: [
        #   MarkdownBlock(type="heading", level=1, ...),
        #   MarkdownBlock(type="paragraph", text="Paragraph text", ...),
        #   MarkdownBlock(type="bullet_list", items=[...], ...)
        # ]
        ```
    """
    if not text or not text.strip():
        return []

    try:
        md = _create_markdown_parser()
        tokens = md.parse(text)

        blocks: List[MarkdownBlock] = []

        i = 0
        while i < len(tokens):
            tok = tokens[i]

            # Headings
            if tok.type == "heading_open":
                level = int(tok.tag[1]) if tok.tag.startswith("h") and tok.tag[1:].isdigit() else None
                # Next token should be inline with content
                if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                    inline_token = tokens[i + 1]
                    heading_text = inline_token.content or ""
                    formatted_runs = _parse_inline_formatting(inline_token)
                    blocks.append(
                        MarkdownBlock(
                            type="heading",
                            text=heading_text,  # Keep for fallback
                            formatted_runs=formatted_runs if formatted_runs else None,
                            level=level,
                        )
                    )
                # Skip until heading_close
                while i < len(tokens) and tokens[i].type != "heading_close":
                    i += 1
                i += 1
                continue

            # Paragraphs (top-level)
            if tok.type == "paragraph_open":
                if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                    inline_token = tokens[i + 1]
                    para_text = inline_token.content or ""
                    if para_text.strip():
                        formatted_runs = _parse_inline_formatting(inline_token)
                        blocks.append(
                            MarkdownBlock(
                                type="paragraph",
                                text=para_text,  # Keep for fallback
                                formatted_runs=formatted_runs if formatted_runs else None,
                            )
                        )
                # Skip until paragraph_close
                while i < len(tokens) and tokens[i].type != "paragraph_close":
                    i += 1
                i += 1
                continue

            # Bullet lists (unordered) with nesting support
            if tok.type == "bullet_list_open":
                items: List[tuple[int, str, Optional[List[FormattedRun]]]] = []
                continuation_paragraphs: List[MarkdownBlock] = []
                
                def parse_bullet_list(start_idx: int, base_level: int) -> int:
                    """
                    Parse a bullet list starting at start_idx.
                    Returns index after the matching bullet_list_close.
                    base_level: nesting level (1 for top-level, 2 for nested, etc.)
                    """
                    nonlocal continuation_paragraphs
                    idx = start_idx
                    current_level = base_level
                    
                    while idx < len(tokens):
                        if tokens[idx].type == "bullet_list_close":
                            return idx + 1
                        
                        if tokens[idx].type == "bullet_list_open":
                            # Nested list - recursively parse it
                            idx = parse_bullet_list(idx + 1, current_level + 1)
                            continue
                        
                        if tokens[idx].type == "list_item_open":
                            # Collect all paragraphs in this list item.
                            # - First paragraph becomes the list item text
                            # - Subsequent paragraphs become "continuation paragraphs" that must visually belong
                            #   to the list item in Word (indented like the list level, but without glyph).
                            item_text_parts: List[str] = []
                            item_formatted_runs: Optional[List[FormattedRun]] = None
                            continuation_paras: List[tuple[str, Optional[List[FormattedRun]]]] = []

                            j = idx + 1
                            current_para_parts: List[str] = []
                            current_para_inline_tokens: List[Any] = []

                            def _concat_formatted_runs(inline_tokens: List[Any]) -> Optional[List[FormattedRun]]:
                                all_runs: List[FormattedRun] = []
                                for itok in inline_tokens:
                                    all_runs.extend(_parse_inline_formatting(itok))
                                return all_runs if all_runs else None

                            def _flush_current_paragraph() -> None:
                                nonlocal item_text_parts, item_formatted_runs, continuation_paras, current_para_parts, current_para_inline_tokens
                                if not current_para_parts:
                                    current_para_parts = []
                                    current_para_inline_tokens = []
                                    return
                                para_text = " ".join(p for p in current_para_parts if p)
                                if para_text.strip():
                                    formatted = _concat_formatted_runs(current_para_inline_tokens)
                                    if not item_text_parts:
                                        # First paragraph becomes list item text
                                        item_text_parts = list(current_para_parts)
                                        item_formatted_runs = formatted
                                    else:
                                        continuation_paras.append((para_text, formatted))
                                current_para_parts = []
                                current_para_inline_tokens = []

                            while j < len(tokens) and tokens[j].type != "list_item_close":
                                tj = tokens[j]

                                if tj.type == "paragraph_open":
                                    # Starting a new paragraph: flush any pending buffer first (deterministic).
                                    _flush_current_paragraph()

                                elif tj.type == "inline":
                                    if tj.content:
                                        current_para_parts.append(tj.content)
                                    current_para_inline_tokens.append(tj)

                                elif tj.type == "paragraph_close":
                                    # Paragraph ended: flush buffer now so we never drop the final paragraph.
                                    _flush_current_paragraph()

                                elif tj.type == "table_open":
                                    # Table inside list item: flush current paragraph, then parse the table
                                    _flush_current_paragraph()
                                    # Parse table using same logic as top-level table parsing
                                    table_header: List[str] = []
                                    table_rows: List[List[str]] = []
                                    
                                    table_j = j + 1
                                    while table_j < len(tokens) and tokens[table_j].type != "table_close":
                                        # Header row
                                        if tokens[table_j].type == "thead_open":
                                            table_j += 1
                                            while table_j < len(tokens) and tokens[table_j].type != "thead_close":
                                                if tokens[table_j].type == "tr_open":
                                                    table_j += 1
                                                    current_row: List[str] = []
                                                    while table_j < len(tokens) and tokens[table_j].type != "tr_close":
                                                        if tokens[table_j].type == "th_open":
                                                            if table_j + 1 < len(tokens) and tokens[table_j + 1].type == "inline":
                                                                cell_text = tokens[table_j + 1].content or ""
                                                                current_row.append(cell_text)
                                                            while table_j < len(tokens) and tokens[table_j].type != "th_close":
                                                                table_j += 1
                                                        table_j += 1
                                                    if current_row:
                                                        table_header = current_row
                                                table_j += 1
                                            table_j += 1
                                            continue
                                        
                                        # Body rows
                                        if tokens[table_j].type == "tbody_open":
                                            table_j += 1
                                            while table_j < len(tokens) and tokens[table_j].type != "tbody_close":
                                                if tokens[table_j].type == "tr_open":
                                                    table_j += 1
                                                    current_row: List[str] = []
                                                    while table_j < len(tokens) and tokens[table_j].type != "tr_close":
                                                        if tokens[table_j].type == "td_open":
                                                            if table_j + 1 < len(tokens) and tokens[table_j + 1].type == "inline":
                                                                cell_text = tokens[table_j + 1].content or ""
                                                                current_row.append(cell_text)
                                                            while table_j < len(tokens) and tokens[table_j].type != "td_close":
                                                                table_j += 1
                                                        table_j += 1
                                                    if current_row:
                                                        table_rows.append(current_row)
                                                table_j += 1
                                            table_j += 1
                                            continue
                                        
                                        table_j += 1
                                    
                                    # Create table block with list continuation semantic contract
                                    if table_header or table_rows:
                                        table_block = MarkdownBlock(
                                            type="table",
                                            header=table_header or [],
                                            rows=table_rows or [],
                                        )
                                        # Private attribute: semantic contract for rendering phase
                                        # Tables inside list items are rendered as text-grid with indentation
                                        setattr(table_block, "_list_continuation_level", current_level)
                                        continuation_paragraphs.append(table_block)
                                    
                                    # Skip past table_close
                                    j = table_j
                                    continue

                                elif tj.type == "bullet_list_open":
                                    # Nested list begins: flush current paragraph (if any) then stop collecting text.
                                    _flush_current_paragraph()
                                    break

                                j += 1

                            # Guard: ensure any pending paragraph content is flushed at list-item end.
                            _flush_current_paragraph()

                            # Add list item if we have text
                            item_text = " ".join(p for p in item_text_parts if p)
                            if item_text.strip():
                                logical_level = current_level
                                items.append((logical_level, item_text, item_formatted_runs))

                                # Add continuation paragraphs (mark them as list continuations)
                                # SEMANTIC CONTRACT: The _list_continuation_level private attribute is a contract
                                # between parsing and rendering phases. It indicates that this paragraph visually
                                # belongs to the list item (same indentation, no bullet/number glyph).
                                # This attribute must not be removed or renamed without updating both parsing
                                # (where it's set) and rendering (where it's consumed) to maintain list semantics.
                                for para_text, para_formatted_runs in continuation_paras:
                                    cont_block = MarkdownBlock(
                                        type="paragraph",
                                        text=para_text,
                                        formatted_runs=para_formatted_runs,
                                    )
                                    # Private attribute: semantic contract for rendering phase
                                    setattr(cont_block, "_list_continuation_level", logical_level)
                                    continuation_paragraphs.append(cont_block)
                            
                            # Skip to list_item_close, processing any nested lists
                            while idx < len(tokens) and tokens[idx].type != "list_item_close":
                                if tokens[idx].type == "bullet_list_open":
                                    idx = parse_bullet_list(idx + 1, current_level + 1)
                                    continue
                                idx += 1
                            idx += 1
                            continue
                        
                        idx += 1
                    
                    return idx
                
                # Parse the list starting from the next token
                i = parse_bullet_list(i + 1, 1)
                
                if items:
                    # Diagnostic: verify no empty list items (debug-only, never raises)
                    for item_idx, item in enumerate(items):
                        if isinstance(item, tuple) and len(item) >= 2:
                            item_text = item[1] if len(item) > 1 else ""
                            if not item_text or not item_text.strip():
                                logger.debug(
                                    "Empty bullet list item detected at index %d (level %d). "
                                    "This may indicate parsing logic issues.",
                                    item_idx,
                                    item[0] if len(item) > 0 else 0,
                                )
                    
                    blocks.append(
                        MarkdownBlock(
                            type="bullet_list",
                            items=items,
                        )
                    )
                    # Append continuation paragraphs immediately after the list
                    blocks.extend(continuation_paragraphs)
                continue

            # Numbered lists (ordered) with nesting support
            if tok.type == "ordered_list_open":
                items: List[tuple[int, str, Optional[List[FormattedRun]]]] = []
                continuation_paragraphs: List[MarkdownBlock] = []
                
                def parse_ordered_list(start_idx: int, base_level: int) -> int:
                    """
                    Parse an ordered list starting at start_idx.
                    Returns index after the matching ordered_list_close.
                    base_level: nesting level (1 for top-level, 2 for nested, etc.)
                    """
                    nonlocal continuation_paragraphs
                    idx = start_idx
                    current_level = base_level
                    
                    while idx < len(tokens):
                        if tokens[idx].type == "ordered_list_close":
                            return idx + 1
                        
                        if tokens[idx].type == "ordered_list_open":
                            # Nested list - recursively parse it
                            idx = parse_ordered_list(idx + 1, current_level + 1)
                            continue
                        
                        if tokens[idx].type == "list_item_open":
                            # Collect all paragraphs in this list item.
                            # - First paragraph becomes the list item text
                            # - Subsequent paragraphs become "continuation paragraphs" that must visually belong
                            #   to the list item in Word (indented like the list level, but without glyph).
                            item_text_parts: List[str] = []
                            item_formatted_runs: Optional[List[FormattedRun]] = None
                            continuation_paras: List[tuple[str, Optional[List[FormattedRun]]]] = []

                            j = idx + 1
                            current_para_parts: List[str] = []
                            current_para_inline_tokens: List[Any] = []

                            def _concat_formatted_runs(inline_tokens: List[Any]) -> Optional[List[FormattedRun]]:
                                all_runs: List[FormattedRun] = []
                                for itok in inline_tokens:
                                    all_runs.extend(_parse_inline_formatting(itok))
                                return all_runs if all_runs else None

                            def _flush_current_paragraph() -> None:
                                nonlocal item_text_parts, item_formatted_runs, continuation_paras, current_para_parts, current_para_inline_tokens
                                if not current_para_parts:
                                    current_para_parts = []
                                    current_para_inline_tokens = []
                                    return
                                para_text = " ".join(p for p in current_para_parts if p)
                                if para_text.strip():
                                    formatted = _concat_formatted_runs(current_para_inline_tokens)
                                    if not item_text_parts:
                                        # First paragraph becomes list item text
                                        item_text_parts = list(current_para_parts)
                                        item_formatted_runs = formatted
                                    else:
                                        continuation_paras.append((para_text, formatted))
                                current_para_parts = []
                                current_para_inline_tokens = []

                            while j < len(tokens) and tokens[j].type != "list_item_close":
                                tj = tokens[j]

                                if tj.type == "paragraph_open":
                                    # Starting a new paragraph: flush any pending buffer first (deterministic).
                                    _flush_current_paragraph()

                                elif tj.type == "inline":
                                    if tj.content:
                                        current_para_parts.append(tj.content)
                                    current_para_inline_tokens.append(tj)

                                elif tj.type == "paragraph_close":
                                    # Paragraph ended: flush buffer now so we never drop the final paragraph.
                                    _flush_current_paragraph()

                                elif tj.type == "table_open":
                                    # Table inside list item: flush current paragraph, then parse the table
                                    _flush_current_paragraph()
                                    # Parse table using same logic as top-level table parsing
                                    table_header: List[str] = []
                                    table_rows: List[List[str]] = []
                                    
                                    table_j = j + 1
                                    while table_j < len(tokens) and tokens[table_j].type != "table_close":
                                        # Header row
                                        if tokens[table_j].type == "thead_open":
                                            table_j += 1
                                            while table_j < len(tokens) and tokens[table_j].type != "thead_close":
                                                if tokens[table_j].type == "tr_open":
                                                    table_j += 1
                                                    current_row: List[str] = []
                                                    while table_j < len(tokens) and tokens[table_j].type != "tr_close":
                                                        if tokens[table_j].type == "th_open":
                                                            if table_j + 1 < len(tokens) and tokens[table_j + 1].type == "inline":
                                                                cell_text = tokens[table_j + 1].content or ""
                                                                current_row.append(cell_text)
                                                            while table_j < len(tokens) and tokens[table_j].type != "th_close":
                                                                table_j += 1
                                                        table_j += 1
                                                    if current_row:
                                                        table_header = current_row
                                                table_j += 1
                                            table_j += 1
                                            continue
                                        
                                        # Body rows
                                        if tokens[table_j].type == "tbody_open":
                                            table_j += 1
                                            while table_j < len(tokens) and tokens[table_j].type != "tbody_close":
                                                if tokens[table_j].type == "tr_open":
                                                    table_j += 1
                                                    current_row: List[str] = []
                                                    while table_j < len(tokens) and tokens[table_j].type != "tr_close":
                                                        if tokens[table_j].type == "td_open":
                                                            if table_j + 1 < len(tokens) and tokens[table_j + 1].type == "inline":
                                                                cell_text = tokens[table_j + 1].content or ""
                                                                current_row.append(cell_text)
                                                            while table_j < len(tokens) and tokens[table_j].type != "td_close":
                                                                table_j += 1
                                                        table_j += 1
                                                    if current_row:
                                                        table_rows.append(current_row)
                                                table_j += 1
                                            table_j += 1
                                            continue
                                        
                                        table_j += 1
                                    
                                    # Create table block with list continuation semantic contract
                                    if table_header or table_rows:
                                        table_block = MarkdownBlock(
                                            type="table",
                                            header=table_header or [],
                                            rows=table_rows or [],
                                        )
                                        # Private attribute: semantic contract for rendering phase
                                        # Tables inside list items are rendered as text-grid with indentation
                                        setattr(table_block, "_list_continuation_level", current_level)
                                        continuation_paragraphs.append(table_block)
                                    
                                    # Skip past table_close
                                    j = table_j
                                    continue

                                elif tj.type in ("ordered_list_open", "bullet_list_open"):
                                    # Nested list begins: flush current paragraph (if any) then stop collecting text.
                                    _flush_current_paragraph()
                                    break

                                j += 1

                            # Guard: ensure any pending paragraph content is flushed at list-item end.
                            _flush_current_paragraph()

                            # Add list item if we have text
                            item_text = " ".join(p for p in item_text_parts if p)
                            if item_text.strip():
                                logical_level = current_level
                                items.append((logical_level, item_text, item_formatted_runs))

                                # Add continuation paragraphs (mark them as list continuations)
                                # SEMANTIC CONTRACT: The _list_continuation_level private attribute is a contract
                                # between parsing and rendering phases. It indicates that this paragraph visually
                                # belongs to the list item (same indentation, no bullet/number glyph).
                                # This attribute must not be removed or renamed without updating both parsing
                                # (where it's set) and rendering (where it's consumed) to maintain list semantics.
                                for para_text, para_formatted_runs in continuation_paras:
                                    cont_block = MarkdownBlock(
                                        type="paragraph",
                                        text=para_text,
                                        formatted_runs=para_formatted_runs,
                                    )
                                    # Private attribute: semantic contract for rendering phase
                                    setattr(cont_block, "_list_continuation_level", logical_level)
                                    continuation_paragraphs.append(cont_block)
                            
                            # Skip to list_item_close, processing any nested lists
                            while idx < len(tokens) and tokens[idx].type != "list_item_close":
                                if tokens[idx].type in ("ordered_list_open", "bullet_list_open"):
                                    # Recursively parse nested list
                                    if tokens[idx].type == "ordered_list_open":
                                        idx = parse_ordered_list(idx + 1, current_level + 1)
                                    else:
                                        # ============================================================================
                                        # ⚠️ FRAGILE LOGIC: Nested Bullet List Parsing Inside Ordered Lists
                                        # ============================================================================
                                        # This function intentionally duplicates parts of bullet-list parsing logic
                                        # to prevent text loss when bullet lists appear inside ordered list items.
                                        #
                                        # WHY IT EXISTS:
                                        #   - Previously, nested bullet lists were skipped, causing guaranteed text loss
                                        #   - This violates the library's loss-less text guarantee
                                        #
                                        # WHY IT DUPLICATES LOGIC:
                                        #   - We cannot reuse parse_bullet_list() here because it creates separate
                                        #     MarkdownBlock instances, but we need to add items to the parent
                                        #     numbered_list's items list to preserve ordering and semantics
                                        #
                                        # BEHAVIORAL REQUIREMENTS:
                                        #   - Must stay behaviorally aligned with bullet-list parsing (same text extraction,
                                        #     same formatting handling, same nesting logic)
                                        #   - Must assign logical_level = parent_level + 1 to nested items
                                        #   - Must preserve all visible text and formatting exactly
                                        #
                                        # ⚠️ DO NOT MODIFY WITHOUT:
                                        #   - Full regression tests covering nested bullet→ordered scenarios
                                        #   - Verification that text loss is still impossible
                                        #   - Ensuring behavioral alignment with parse_bullet_list() is maintained
                                        #
                                        # This logic is fragile by necessity and must not be refactored or optimized
                                        # without comprehensive test coverage.
                                        # ============================================================================
                                        def parse_nested_bullet_for_ordered(start_bullet_idx: int, parent_level: int) -> int:
                                            """Parse nested bullet list and add items to the parent items list."""
                                            nonlocal items
                                            bullet_idx = start_bullet_idx + 1  # Skip bullet_list_open
                                            nested_level = parent_level + 1
                                            
                                            while bullet_idx < len(tokens):
                                                if tokens[bullet_idx].type == "bullet_list_close":
                                                    return bullet_idx + 1
                                                
                                                if tokens[bullet_idx].type == "bullet_list_open":
                                                    # Recursively handle deeper nesting
                                                    bullet_idx = parse_nested_bullet_for_ordered(bullet_idx, nested_level)
                                                    continue
                                                
                                                if tokens[bullet_idx].type == "list_item_open":
                                                    # Parse this bullet list item - collect ALL paragraphs to prevent text loss
                                                    # This must collect all visible text from all paragraphs within the list item,
                                                    # not just the first paragraph.
                                                    nested_item_parts: List[str] = []
                                                    nested_item_inline_tokens: List[Any] = []
                                                    nested_j = bullet_idx + 1
                                                    
                                                    # Collect all paragraphs within this list item (not just the first one)
                                                    while nested_j < len(tokens) and tokens[nested_j].type != "list_item_close":
                                                        if tokens[nested_j].type == "paragraph_open":
                                                            # Paragraph boundary - continue collecting (do not stop)
                                                            pass
                                                        elif tokens[nested_j].type == "inline":
                                                            # Collect inline content from all paragraphs
                                                            if tokens[nested_j].content:
                                                                nested_item_parts.append(tokens[nested_j].content)
                                                            nested_item_inline_tokens.append(tokens[nested_j])
                                                        elif tokens[nested_j].type == "paragraph_close":
                                                            # Paragraph ended - continue to next paragraph (do not stop)
                                                            pass
                                                        elif tokens[nested_j].type in ("bullet_list_open", "ordered_list_open"):
                                                            # Deeper nesting - recursively parse
                                                            if tokens[nested_j].type == "bullet_list_open":
                                                                nested_j = parse_nested_bullet_for_ordered(nested_j, nested_level)
                                                            else:
                                                                nested_j = parse_ordered_list(nested_j + 1, nested_level)
                                                            continue
                                                        nested_j += 1
                                                    
                                                    # Concatenate all collected text from all paragraphs deterministically
                                                    # DESIGN DECISION: Multiple paragraphs inside nested bullet items are intentionally
                                                    # flattened into a single text string. This avoids structural complexity and semantic
                                                    # coupling that would require new block types or continuation attributes. The single-space
                                                    # join attempts to preserve readability while attempting to ensure no text loss.
                                                    nested_item_text = " ".join(p for p in nested_item_parts if p)
                                                    if nested_item_text.strip():
                                                        nested_formatted_runs = None
                                                        if nested_item_inline_tokens:
                                                            # Apply formatting to all inline tokens (from all paragraphs)
                                                            all_runs: List[FormattedRun] = []
                                                            for itok in nested_item_inline_tokens:
                                                                all_runs.extend(_parse_inline_formatting(itok))
                                                            nested_formatted_runs = all_runs if all_runs else None
                                                        items.append((nested_level, nested_item_text, nested_formatted_runs))
                                                    
                                                    # Skip to list_item_close
                                                    while bullet_idx < len(tokens) and tokens[bullet_idx].type != "list_item_close":
                                                        if tokens[bullet_idx].type in ("bullet_list_open", "ordered_list_open"):
                                                            if tokens[bullet_idx].type == "bullet_list_open":
                                                                bullet_idx = parse_nested_bullet_for_ordered(bullet_idx, nested_level)
                                                            else:
                                                                bullet_idx = parse_ordered_list(bullet_idx + 1, nested_level)
                                                            continue
                                                        bullet_idx += 1
                                                    bullet_idx += 1
                                                    continue
                                                
                                                bullet_idx += 1
                                            
                                            return bullet_idx
                                        
                                        idx = parse_nested_bullet_for_ordered(idx, current_level)
                                    continue
                                idx += 1
                            idx += 1
                            continue
                        
                        idx += 1
                    
                    return idx
                
                # Parse the list starting from the next token
                i = parse_ordered_list(i + 1, 1)
                
                if items:
                    # Diagnostic: verify no empty list items (debug-only, never raises)
                    for item_idx, item in enumerate(items):
                        if isinstance(item, tuple) and len(item) >= 2:
                            item_text = item[1] if len(item) > 1 else ""
                            if not item_text or not item_text.strip():
                                logger.debug(
                                    "Empty numbered list item detected at index %d (level %d). "
                                    "This may indicate parsing logic issues.",
                                    item_idx,
                                    item[0] if len(item) > 0 else 0,
                                )
                    
                    blocks.append(
                        MarkdownBlock(
                            type="numbered_list",
                            items=items,
                        )
                    )
                    # Append continuation paragraphs immediately after the list
                    blocks.extend(continuation_paragraphs)
                continue

            # Tables
            if tok.type == "table_open":
                header: List[str] = []
                rows: List[List[str]] = []

                i += 1
                while i < len(tokens) and tokens[i].type != "table_close":
                    # Header row
                    if tokens[i].type == "thead_open":
                        i += 1
                        while i < len(tokens) and tokens[i].type != "thead_close":
                            if tokens[i].type == "tr_open":
                                i += 1
                                current_row: List[str] = []
                                while i < len(tokens) and tokens[i].type != "tr_close":
                                    if tokens[i].type == "th_open":
                                        # Expect inline next
                                        if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                                            cell_text = tokens[i + 1].content or ""
                                            current_row.append(cell_text)
                                        # Skip to th_close
                                        while i < len(tokens) and tokens[i].type != "th_close":
                                            i += 1
                                    i += 1
                                if current_row:
                                    header = current_row
                            i += 1
                        i += 1
                        continue

                    # Body rows
                    if tokens[i].type == "tbody_open":
                        i += 1
                        while i < len(tokens) and tokens[i].type != "tbody_close":
                            if tokens[i].type == "tr_open":
                                i += 1
                                current_row: List[str] = []
                                while i < len(tokens) and tokens[i].type != "tr_close":
                                    if tokens[i].type == "td_open":
                                        if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                                            cell_text = tokens[i + 1].content or ""
                                            current_row.append(cell_text)
                                        # Skip to td_close
                                        while i < len(tokens) and tokens[i].type != "td_close":
                                            i += 1
                                    i += 1
                                if current_row:
                                    rows.append(current_row)
                            i += 1
                        i += 1
                        continue

                    i += 1

                if header or rows:
                    blocks.append(
                        MarkdownBlock(
                            type="table",
                            header=header or [],
                            rows=rows or [],
                        )
                    )
                continue

            # Fenced code blocks and indented code blocks
            if tok.type in ("fence", "code_block"):
                code_content = tok.content or ""
                if code_content.strip():
                    blocks.append(
                        MarkdownBlock(
                            type="paragraph",
                            text=code_content,
                        )
                    )
                # Skip to the end of the code block token
                i += 1
                continue

            i += 1

        # If parsing produced no blocks but text is non-empty, fall back to paragraphs
        if not blocks:
            logger.debug("Markdown parsed but yielded no blocks; falling back to paragraphs.")
            raise ValueError("No markdown blocks produced")

        return blocks

    except Exception as e:
        # Fail-safe: treat content as plain text paragraphs separated by blank lines
        # SAFETY INVARIANT: Markdown parsing fallback attempts to preserve content as text.
        # If parsing fails, content is preserved as plain paragraphs rather than being discarded.
        # This attempts to maintain content fidelity and prevent silent data loss.
        # The fallback preserves text deterministically, attempting to ensure no content is dropped.
        logger.warning(f"Markdown parsing failed, falling back to plain paragraphs: {e}")
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        return [MarkdownBlock(type="paragraph", text=p) for p in paragraphs]


def _observe_tree_structure(
    tree: "BlockNode",
    summary: Optional[ExportSummary],
    token: str
) -> None:
    """
    Observe tree structure for nested tables and mixed lists (observability only).
    
    This function detects edge cases in the tree structure without changing
    any rendering behavior. Used for export summary and structured logging.
    
    SAFETY: This is read-only observation. No rendering logic is affected.
    
    Args:
        tree: DocumentNode to observe
        summary: Optional export summary to record observations
        token: Placeholder token being processed (for context)
    """
    if summary is None:
        return
    
    from docx_template_export.models.markdown_tree import (
        BlockNode,
        DocumentNode,
        TableNode,
        TableCellNode,
        ListNode,
        ListItemNode,
    )
    
    def has_nested_table(node: BlockNode, depth: int = 0) -> bool:
        """Check if tree contains tables nested inside table cells."""
        if isinstance(node, TableCellNode):
            # Check if this cell contains a table
            for child in node.children:
                if isinstance(child, TableNode):
                    return True
                if has_nested_table(child, depth + 1):
                    return True
        elif isinstance(node, (DocumentNode, ListItemNode)):
            for child in node.children:
                if has_nested_table(child, depth + 1):
                    return True
        elif isinstance(node, ListNode):
            for item in node.items:
                if has_nested_table(item, depth + 1):
                    return True
        return False
    
    def has_mixed_lists(node: "BlockNode") -> bool:
        """Check if tree contains mixed ordered and bullet lists in same context."""
        if isinstance(node, (DocumentNode, ListItemNode)):
            list_types = set()
            for child in node.children:
                if isinstance(child, ListNode):
                    # ListNode uses 'kind' attribute, not 'list_type'
                    list_types.add(child.kind)
                elif isinstance(child, ListItemNode):
                    # Check parent list type
                    pass
                if has_mixed_lists(child):
                    return True
            # If we find both ordered and bullet lists at same level
            if "ordered" in list_types and "bullet" in list_types:
                return True
        elif isinstance(node, ListNode):
            for item in node.items:
                if has_mixed_lists(item):
                    return True
        return False
    
    # Check for nested tables
    if isinstance(tree, DocumentNode) and has_nested_table(tree):
        summary.fallback_events.append(
            FallbackEvent(
                event_type="nested_table_detected",
                location=token,
                reason="Table found inside table cell (tree renderer handles this)",
            )
        )
        # Structured log: gated by enable_export_trace (summary is only created when enabled)
        logger.info(f"Tree observation: Nested table detected in {token} (tree renderer will handle)")
    
    # Check for mixed lists
    if isinstance(tree, DocumentNode) and has_mixed_lists(tree):
        summary.fallback_events.append(
            FallbackEvent(
                event_type="mixed_lists_detected",
                location=token,
                reason="Mixed ordered and bullet lists found in same context",
            )
        )
        # Structured log: gated by enable_export_trace (summary is only created when enabled)
        logger.info(f"Tree observation: Mixed lists detected in {token}")


def _extract_block_types_from_tree(tree: "BlockNode") -> List[str]:
    """
    Extract block types from a markdown tree for summary recording.
    
    This function traverses the tree and collects unique block types
    (heading, paragraph, bullet_list, numbered_list, table) for observability.
    
    Args:
        tree: DocumentNode to extract block types from
    
    Returns:
        List of block type strings (e.g., ["heading", "paragraph", "bullet_list"])
    """
    from docx_template_export.models.markdown_tree import (
        DocumentNode,
        HeadingNode,
        ParagraphNode,
        ListNode,
        ListItemNode,
        TableNode,
        TableRowNode,
        TableCellNode,
    )
    
    block_types = set()
    
    def collect_types(node: "BlockNode"):
        if isinstance(node, HeadingNode):
            block_types.add("heading")
        elif isinstance(node, ParagraphNode):
            block_types.add("paragraph")
        elif isinstance(node, ListNode):
            if node.kind == "bullet":
                block_types.add("bullet_list")
            elif node.kind == "ordered":
                block_types.add("numbered_list")
        elif isinstance(node, TableNode):
            block_types.add("table")
        
        # Recursively collect from children
        if isinstance(node, (DocumentNode, ListItemNode, TableCellNode)):
            for child in node.children:
                collect_types(child)
        elif isinstance(node, ListNode):
            for item in node.items:
                collect_types(item)
        elif isinstance(node, TableNode):
            for row in node.rows:
                for cell in row.cells:
                    for child in cell.children:
                        collect_types(child)
    
    collect_types(tree)
    
    # Return sorted list for deterministic output
    return sorted(list(block_types))


def parse_markdown_to_tree(
    markdown: str,
    *,
    enabled: bool = False,
    summary: Optional[ExportSummary] = None
) -> Optional["BlockNode"]:
    """
    Build a tree-based intermediate representation (IR) from markdown text.
    
    This function attempts to create an immutable tree structure that preserves order
    and nesting of markdown elements (lists, list items, tables, table cells) where
    structurally safe. The tree representation is used by the tree-based renderer,
    which is an experimental, feature-flagged component that runs in parallel to
    the block-based renderer.
    
    The tree IR attempts to preserve semantic structure without flattening or losing
    nesting information where supported. This may allow for more accurate rendering
    of complex nested structures (e.g., lists within lists, tables within list items),
    subject to DOCX constraints.
    
    Safety Guarantees:
        - Tree parsing failures gracefully fall back to block-based renderer
        - Text content is preserved as text when structure cannot be parsed (text-preserving fallback)
        - Inline formatting (bold, italic) is preserved through FormattedRun objects where detected
        - Tree depth and node count are validated before rendering (in tree renderer)
        - Parsing errors are caught and recorded in export summary (if enabled)
    
    Relationship to Block-Based Parser:
        This function is parallel to parse_markdown_to_blocks() and does not replace it.
        The block-based parser remains the default and primary path. The tree parser
        is feature-flagged and used only when explicitly enabled or when the block
        renderer encounters structures it cannot handle.
    
    Args:
        markdown: The markdown text to parse. Can be empty or None (returns empty DocumentNode).
        enabled: If False, immediately return None without parsing. If True, parse and build tree.
            This flag allows callers to conditionally enable tree parsing.
        summary: Optional export summary for observability. If provided and tree parsing fails,
            a FallbackEvent is recorded with event_type="tree_parse_exception".
    
    Returns:
        DocumentNode if enabled=True and parsing succeeds, containing the root of the tree.
        None if enabled=False or if parsing fails (fallback to block-based renderer).
        Empty DocumentNode (children=[]) if input is empty or None.
    
    Note:
        This is an experimental, feature-flagged component. The tree is built and may be
        used for rendering when enabled, but the block-based renderer remains the default
        and primary path. Tree parsing failures automatically fall back to block-based
        rendering to attempt export completion.
    """
    if not enabled:
        return None
    
    if not markdown or not markdown.strip():
        from docx_template_export.models.markdown_tree import DocumentNode
        return DocumentNode(children=[])
    
    try:
        from docx_template_export.models.markdown_tree import (
            BlockNode,
            DocumentNode,
            HeadingNode,
            ParagraphNode,
            ListNode,
            ListItemNode,
            TableNode,
            TableRowNode,
            TableCellNode,
        )
        
        md = _create_markdown_parser()
        tokens = md.parse(markdown)
        
        # Stack-based parser to build tree
        # Stack contains (node, parent) pairs where node is the current container
        # and parent is the parent node to add children to
        stack: List[tuple[BlockNode, BlockNode]] = []
        root = DocumentNode(children=[])
        current_parent: BlockNode = root
        
        i = 0
        while i < len(tokens):
            tok = tokens[i]
            
            # Headings
            if tok.type == "heading_open":
                level = int(tok.tag[1]) if tok.tag.startswith("h") and tok.tag[1:].isdigit() else 1
                # Next token should be inline with content
                if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                    inline_token = tokens[i + 1]
                    formatted_runs = _parse_inline_formatting(inline_token)
                    heading_node = HeadingNode(level=level, runs=formatted_runs)
                    current_parent.children.append(heading_node)
                # Skip until heading_close
                while i < len(tokens) and tokens[i].type != "heading_close":
                    i += 1
                i += 1
                continue
            
            # Paragraphs
            if tok.type == "paragraph_open":
                if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                    inline_token = tokens[i + 1]
                    formatted_runs = _parse_inline_formatting(inline_token)
                    if formatted_runs:  # Only add non-empty paragraphs
                        para_node = ParagraphNode(runs=formatted_runs)
                        current_parent.children.append(para_node)
                # Skip until paragraph_close
                while i < len(tokens) and tokens[i].type != "paragraph_close":
                    i += 1
                i += 1
                continue
            
            # Bullet lists
            if tok.type == "bullet_list_open":
                list_node = ListNode(kind="bullet", items=[])
                current_parent.children.append(list_node)
                # Push list onto stack
                stack.append((list_node, current_parent))
                current_parent = list_node
                i += 1
                continue
            
            if tok.type == "bullet_list_close":
                if stack:
                    _, current_parent = stack.pop()
                i += 1
                continue
            
            # Ordered lists
            if tok.type == "ordered_list_open":
                list_node = ListNode(kind="ordered", items=[])
                current_parent.children.append(list_node)
                # Push list onto stack
                stack.append((list_node, current_parent))
                current_parent = list_node
                i += 1
                continue
            
            if tok.type == "ordered_list_close":
                if stack:
                    _, current_parent = stack.pop()
                i += 1
                continue
            
            # List items
            if tok.type == "list_item_open":
                item_node = ListItemNode(children=[])
                # Add to current list (which should be at top of stack)
                if isinstance(current_parent, ListNode):
                    current_parent.items.append(item_node)
                else:
                    # Fallback: add as child if not in list context
                    current_parent.children.append(item_node)
                # Push item onto stack
                stack.append((item_node, current_parent))
                current_parent = item_node
                i += 1
                continue
            
            if tok.type == "list_item_close":
                if stack:
                    _, current_parent = stack.pop()
                i += 1
                continue
            
            # Tables
            if tok.type == "table_open":
                table_node = TableNode(rows=[])
                current_parent.children.append(table_node)
                # Push table onto stack
                stack.append((table_node, current_parent))
                current_parent = table_node
                i += 1
                continue
            
            if tok.type == "table_close":
                if stack:
                    _, current_parent = stack.pop()
                i += 1
                continue
            
            # Table rows
            if tok.type == "tr_open":
                row_node = TableRowNode(cells=[])
                # Add to current table
                if isinstance(current_parent, TableNode):
                    current_parent.rows.append(row_node)
                else:
                    # Fallback: add as child if not in table context
                    current_parent.children.append(row_node)
                # Push row onto stack
                stack.append((row_node, current_parent))
                current_parent = row_node
                i += 1
                continue
            
            if tok.type == "tr_close":
                if stack:
                    _, current_parent = stack.pop()
                i += 1
                continue
            
            # Table cells (header and data)
            if tok.type in ("th_open", "td_open"):
                cell_node = TableCellNode(children=[])
                # Add to current row
                if isinstance(current_parent, TableRowNode):
                    current_parent.cells.append(cell_node)
                else:
                    # Fallback: add as child if not in row context
                    current_parent.children.append(cell_node)
                # Push cell onto stack
                stack.append((cell_node, current_parent))
                current_parent = cell_node
                i += 1
                continue
            
            if tok.type in ("th_close", "td_close"):
                if stack:
                    _, current_parent = stack.pop()
                i += 1
                continue
            
            # Table structure tokens (thead_open, tbody_open, etc.) - skip, they're structural only
            if tok.type in ("thead_open", "thead_close", "tbody_open", "tbody_close"):
                i += 1
                continue
            
            # Handle inline content in table cells (paragraphs inside cells)
            if tok.type == "inline" and isinstance(current_parent, TableCellNode):
                formatted_runs = _parse_inline_formatting(tok)
                if formatted_runs:
                    para_node = ParagraphNode(runs=formatted_runs)
                    current_parent.children.append(para_node)
                i += 1
                continue
            
            # Fenced code blocks and indented code blocks
            if tok.type in ("fence", "code_block"):
                code_content = tok.content or ""
                if code_content.strip():
                    # Treat code blocks as paragraphs
                    para_node = ParagraphNode(runs=[FormattedRun(text=code_content)])
                    current_parent.children.append(para_node)
                i += 1
                continue
            
            i += 1
        
        return root
    
    except Exception as e:
        # SAFETY INVARIANT: Tree parsing failures are handled gracefully.
        # WHY: Tree parsing may encounter edge cases in markdown structure that the parser
        # cannot handle. Returning None triggers fallback to block-based renderer, ensuring
        # deterministic output and preventing document corruption.
        error_msg = f"Tree parsing failed: {e}"
        logger.warning(error_msg)
        
        # Record fallback event in summary (if enabled)
        if summary is not None:
            summary.fallback_events.append(
                FallbackEvent(
                    event_type="tree_parse_exception",
                    location=None,
                    reason=error_msg,
                )
            )
            logger.info(f"Fallback: {error_msg}")
        
        return None


def dump_markdown_tree(node: "BlockNode", indent: int = 0) -> str:
    """
    Generate a human-readable string representation of the markdown tree.
    
    This function creates a stable, indented text representation of the markdown
    tree structure. It is used for debugging, testing, and observability to
    visualize the tree structure without modifying it.
    
    The output format is deterministic and stable, making it suitable for
    test assertions and debugging output. Text content is truncated for
    readability (long paragraphs are truncated to 60 characters).
    
    Args:
        node: The BlockNode to dump. Can be any node type (DocumentNode, HeadingNode,
            ParagraphNode, ListNode, etc.). The function recursively traverses
            the tree structure.
        indent: Current indentation level (for recursive calls). Defaults to 0
            for the root node. Each nesting level adds 2 spaces of indentation.
    
    Returns:
        Multi-line string representation of the tree with indentation showing
        nesting structure. Each node type is represented with its key attributes
        (e.g., "Heading(1, \"Title\")", "List(bullet)", "Paragraph(\"text...\")").
    
    Example:
        ```python
        tree = parse_markdown_to_tree("# Title\\n\\nParagraph", enabled=True)
        print(dump_markdown_tree(tree))
        # Output:
        # Document
        #   Heading(1, "Title")
        #   Paragraph("Paragraph")
        ```
    
    Returns:
        Multi-line string representation of the tree structure
    """
    from docx_template_export.models.markdown_tree import (
        DocumentNode,
        HeadingNode,
        ParagraphNode,
        ListNode,
        ListItemNode,
        TableNode,
        TableRowNode,
        TableCellNode,
    )
    
    indent_str = "  " * indent
    lines: List[str] = []
    
    if isinstance(node, DocumentNode):
        lines.append(f"{indent_str}Document")
        for child in node.children:
            lines.append(dump_markdown_tree(child, indent + 1))
    
    elif isinstance(node, HeadingNode):
        # Extract text from runs for display
        text = "".join(run.text for run in node.runs)
        lines.append(f"{indent_str}Heading({node.level}, \"{text}\")")
    
    elif isinstance(node, ParagraphNode):
        # Extract text from runs for display
        text = "".join(run.text for run in node.runs)
        # Truncate long text for readability
        if len(text) > 60:
            text = text[:57] + "..."
        lines.append(f"{indent_str}Paragraph(\"{text}\")")
    
    elif isinstance(node, ListNode):
        lines.append(f"{indent_str}List({node.kind})")
        for item in node.items:
            lines.append(dump_markdown_tree(item, indent + 1))
    
    elif isinstance(node, ListItemNode):
        lines.append(f"{indent_str}Item")
        for child in node.children:
            lines.append(dump_markdown_tree(child, indent + 1))
    
    elif isinstance(node, TableNode):
        lines.append(f"{indent_str}Table")
        for row in node.rows:
            lines.append(dump_markdown_tree(row, indent + 1))
    
    elif isinstance(node, TableRowNode):
        lines.append(f"{indent_str}Row")
        for cell in node.cells:
            lines.append(dump_markdown_tree(cell, indent + 1))
    
    elif isinstance(node, TableCellNode):
        lines.append(f"{indent_str}Cell")
        for child in node.children:
            lines.append(dump_markdown_tree(child, indent + 1))
    
    else:
        lines.append(f"{indent_str}Unknown({type(node).__name__})")
    
    return "\n".join(lines)


def _iter_header_footer_paragraphs(doc: DocumentType) -> Iterable[Tuple[Paragraph, Region]]:
    """
    Iterate through all paragraphs in headers and footers, including table cells.
    
    This generator function traverses all sections in the document and yields
    paragraphs from headers and footers. It includes both flow paragraphs and
    paragraphs within table cells in headers/footers.
    
    This function is used during scalar placeholder replacement to ensure all
    document regions are processed. Headers and footers are processed separately
    from body content to maintain deterministic replacement order.
    
    Args:
        doc: Word document to traverse. Must have at least one section.
    
    Yields:
        Tuples of (paragraph, region) where:
        - paragraph: python-docx Paragraph object from header or footer
        - region: Region.HEADER or Region.FOOTER indicating the document region
    
    Note:
        This function iterates through all sections in the document. Most documents
        have a single section, but multi-section documents are supported.
    """
    for section in doc.sections:
        # Header paragraphs
        for p in section.header.paragraphs:
            yield (p, Region.HEADER)
        
        # Header table cells
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield (p, Region.HEADER)
        
        # Footer paragraphs
        for p in section.footer.paragraphs:
            yield (p, Region.FOOTER)
        
        # Footer table cells
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield (p, Region.FOOTER)


def _iter_textbox_runs(doc: DocumentType, region: Optional[Region] = None) -> Iterable[Tuple[List[Any], Region, Container]]:
    """
    Iterate through textbox runs using XML/XPath traversal.
    
    This generator function finds all textboxes in the specified document region(s)
    and yields their XML run elements. Textboxes are accessed via XPath because
    python-docx does not provide high-level access to textbox content.
    
    Textbox Structure:
        Textboxes use w:txbxContent elements that contain paragraphs and runs.
        The XPath query `.//w:txbxContent` finds both WPS (Word Processing Shape)
        and classic/VML textboxes, ensuring comprehensive coverage.
    
    Args:
        doc: Word document to traverse. Must have valid XML structure.
        region: Optional region to search. If None, searches all regions
            (BODY, HEADER, FOOTER) in a fixed order. If specified, searches
            only that region.
    
    Yields:
        Tuples of (runs_list, region, Container.TEXTBOX) where:
        - runs_list: List of XML run elements (w:r) from a single textbox.
            These are lxml elements, not python-docx Run objects.
        - region: Region where the textbox was found (BODY, HEADER, or FOOTER)
        - Container.TEXTBOX: Container type indicator (TEXTBOX for this function)
    
    Note:
        This function uses direct XML/XPath access because textboxes are not
        easily accessible through python-docx's high-level API. The XPath
        approach attempts to find textboxes regardless of their
        underlying XML structure (WPS vs VML), subject to DOCX constraints.
    """
    from lxml import etree
    
    # Namespace mapping for XPath
    ns = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
    }

    # When region is None, iterate all regions in a fixed order.
    if region is None:
        for reg in (Region.BODY, Region.HEADER, Region.FOOTER):
            for runs, textbox_region, container in _iter_textbox_runs(doc, reg):
                yield (runs, textbox_region, container)
        return
    
    # Determine which XML element(s) to search based on region
    if region == Region.BODY:
        # Search the main document body
        search_element = doc._body._element
        # Use lxml etree xpath directly (supports namespaces) - access parent class method from MRO
        xpath_method = etree._Element.xpath.__get__(search_element, etree._Element)
        for txbx_content in xpath_method('.//wps:txbx//w:txbxContent', namespaces=ns):
            # Collect all runs within this textbox
            runs_xpath = etree._Element.xpath.__get__(txbx_content, etree._Element)
            runs = runs_xpath('.//w:r', namespaces=ns)
            if runs:
                yield (runs, region, Container.TEXTBOX)
    elif region == Region.HEADER:
        # Search all header sections
        for section in doc.sections:
            header = section.header
            search_element = header._element
            # Use lxml etree xpath directly.
            # NOTE: `.//w:txbxContent` is a superset that covers both WPS and classic/VML textboxes.
            xpath_method = etree._Element.xpath.__get__(search_element, etree._Element)
            for txbx_content in xpath_method('.//w:txbxContent', namespaces=ns):
                # Collect all runs within this textbox
                runs_xpath = etree._Element.xpath.__get__(txbx_content, etree._Element)
                runs = runs_xpath('.//w:r', namespaces=ns)
                if runs:
                    yield (runs, region, Container.TEXTBOX)
    elif region == Region.FOOTER:
        # Search all footer sections
        for section in doc.sections:
            footer = section.footer
            search_element = footer._element
            # Use lxml etree xpath directly.
            # NOTE: `.//w:txbxContent` is a superset that covers both WPS and classic/VML textboxes.
            xpath_method = etree._Element.xpath.__get__(search_element, etree._Element)
            for txbx_content in xpath_method('.//w:txbxContent', namespaces=ns):
                # Collect all runs within this textbox
                runs_xpath = etree._Element.xpath.__get__(txbx_content, etree._Element)
                runs = runs_xpath('.//w:r', namespaces=ns)
                if runs:
                    yield (runs, region, Container.TEXTBOX)


def _iter_all_locations(doc: DocumentType) -> Iterable[Tuple[Union[Paragraph, List[Any]], Region, Container]]:
    """
    Yield all content locations in the document with location metadata.
    
    Returns tuples of (paragraph_or_runs, region, container):
    - paragraph_or_runs: Paragraph object for FLOW/TABLE, or list of XML run elements for TEXTBOX
    - region: BODY, HEADER, or FOOTER
    - container: FLOW (normal paragraphs), TABLE (table cells), or TEXTBOX (w:txbxContent)
    """
    # Body paragraphs (FLOW)
    for p in doc.paragraphs:
        # Skip paragraphs inside table cells (they are handled in BODY TABLE section)
        if _is_paragraph_in_table_cell(p):
            continue
        logger.debug("Traversing BODY FLOW paragraph")
        yield (p, Region.BODY, Container.FLOW)
    
    # Body table cells (TABLE)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    logger.debug("Traversing BODY TABLE cell paragraph")
                    yield (p, Region.BODY, Container.TABLE)
    
    # Header and footer paragraphs and table cells
    for p, region in _iter_header_footer_paragraphs(doc):
        # Determine if this is FLOW or TABLE
        # Check if paragraph's parent is a table cell
        parent = p._p.getparent()
        is_table = parent is not None and parent.tag.endswith('}tc')  # w:tc element
        container = Container.TABLE if is_table else Container.FLOW
        logger.debug(f"Traversing {region.value} {container.value} paragraph")
        yield (p, region, container)
    
    # Textboxes in all regions
    for region in [Region.BODY, Region.HEADER, Region.FOOTER]:
        for runs, textbox_region, container in _iter_textbox_runs(doc, region):
            logger.debug(f"Traversing {textbox_region.value} TEXTBOX with {len(runs)} runs")
            yield (runs, textbox_region, container)


def _is_block_expansion_allowed(region: Region, container: Container) -> bool:
    """
    Check if block expansion is allowed for the given location.
    
    Block expansion is only allowed in BODY FLOW and BODY TABLE.
    
    SAFETY INVARIANT: Block placeholders are restricted to BODY FLOW and BODY TABLE.
    WHY: Headers/footers are restricted to prevent layout instability. Textboxes only support
    scalar replacement due to Word XML structure limitations (w:txbxContent cannot contain
    block-level elements). This attempts to ensure deterministic output and prevent document corruption.
    """
    return region == Region.BODY and container in [Container.FLOW, Container.TABLE]


def _get_text_from_xml_run(run_elem) -> str:
    """
    Extract text from an XML run element (w:r).
    
    Returns the concatenated text from all w:t elements within the run.
    Attempts to preserve explicit line breaks (<w:br/> and <w:cr/>) as newline characters.
    
    Note:
    Line breaks (<w:br/>, <w:cr/>) are preserved as textual '\\n' characters.
    This library does not synthesize visual Word line breaks.
    Text is preserved deterministically without guessing layout.
    """
    from lxml import etree
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    # Iterate through direct children only (strict sibling order attempts to preserve document order)
    # This attempts to ensure deterministic ordering when <w:t>, <w:br/>, and formatting elements are interleaved.
    parts: List[str] = []
    
    # Iterate through direct children (not descendants) to preserve strict sibling order
    for elem in run_elem:
        tag = elem.tag
        if tag.endswith('}t'):  # w:t element (text node)
            parts.append(elem.text or "")
        elif tag.endswith('}br') or tag.endswith('}cr'):  # w:br or w:cr element (explicit line break)
            # Preserve explicit line breaks as newline characters for better textual fidelity
            parts.append("\n")
        # All other node types (formatting w:rPr, etc.) are ignored - only text and breaks are extracted
    
    return "".join(parts)


def _set_text_in_xml_run(run_elem, text: str) -> None:
    """
    Set text in an XML run element (w:r), preserving structure.
    
    This function safely updates text content in a Word XML run element without
    creating or removing XML nodes. It only mutates existing text nodes to ensure
    structural stability and deterministic behavior.
    
    Safety Guarantees:
        - XML node creation is forbidden (no new w:t elements created)
        - XML node removal is forbidden (no w:t elements deleted)
        - Structural stability is preserved (run element structure unchanged)
        - Only existing text nodes may be mutated (first w:t updated, others cleared)
    
    Why This Approach:
        Creating or removing XML nodes could affect document structure in
        non-deterministic ways. By only mutating existing nodes, we ensure
        that the document structure attempts to remain stable and predictable.
    
    Args:
        run_elem: XML run element (w:r) to update. Must be a valid lxml element.
        text: New text content to set. Replaces existing text in the first w:t element.
    
    Returns:
        None. The run element is modified in-place.
    
    Note:
        If the run has no w:t elements, no mutation is performed. This is
        intentional to preserve structural stability.
    """
    from lxml import etree
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    xpath_method = etree._Element.xpath.__get__(run_elem, etree._Element)
    text_elems = xpath_method('.//w:t', namespaces=ns)
    
    if text_elems:
        # Update first text element; do not remove nodes (stability invariant)
        text_elems[0].text = text
        for elem in text_elems[1:]:
            elem.text = ""
    else:
        # No existing w:t; do not create nodes (stability invariant)
        return


def _safe_replace_scalar_across_textbox_runs(
    run_elems: List[Any],
    placeholder: str,
    replacement: str,
    region: Region,
) -> None:
    """
    Safely replace a scalar placeholder inside textbox runs, supporting multi-run spans.
    
    This function handles scalar placeholder replacement in textboxes where the placeholder
    may span multiple XML runs (w:r elements). Textboxes have a different XML structure
    than regular paragraphs, requiring special handling to preserve structure.
    
    Safety Guarantees:
        - Replacement can span multiple runs (handles Word's run splitting)
        - No run insertion or deletion (structure preserved)
        - No formatting changes (w:rPr preserved, only text modified)
        - Only text content is modified (via _set_text_in_xml_run)
        - Ambiguous placeholders (multiple occurrences) are skipped
        - Unmappable spans are skipped (logged as warning)
    
    Why Special Handling for Textboxes:
        Textboxes use w:txbxContent elements that contain runs. The XML structure is
        different from regular paragraphs, and placeholders may be split across runs
        differently. This function attempts to ensure safe replacement without corrupting the
        textbox structure.
    
    Args:
        run_elems: List of XML run elements (w:r) from the textbox. These are lxml
            elements, not python-docx Run objects.
        placeholder: Placeholder text to replace (e.g., "{{title}}", "{{key}}").
            Must appear exactly once in the concatenated text.
        replacement: Replacement text to insert in place of the placeholder.
        region: Document region (BODY, HEADER, FOOTER) for logging context.
    
    Returns:
        None. The run elements are modified in-place via _set_text_in_xml_run.
    
    Note:
        If the placeholder is found more than once or the span cannot be mapped
        cleanly to runs, the content is left unchanged and a warning is logged.
        This attempts to ensure deterministic behavior and prevent document corruption.
    """
    if not run_elems:
        return

    # Read-only scan: Extract text from each run and build concatenated string
    run_texts = []
    for run in run_elems:
        text = _get_text_from_xml_run(run)
        run_texts.append(text)
    
    full_text = "".join(run_texts)
    
    # Detection: If placeholder not found, return
    if placeholder not in full_text:
        return
    
    # SAFETY INVARIANT: Ambiguous placeholders (appearing multiple times) are skipped.
    # WHY: If a placeholder appears more than once in the same location, we cannot
    # deterministically decide which occurrence to replace. Skipping attempts to ensure deterministic
    # output and prevent partial replacements that could corrupt document structure.
    # Detection: If placeholder found more than once, log debug and return
    # (This should not happen as we check in replace_scalar_placeholders, but defensive check)
    placeholder_count = full_text.count(placeholder)
    if placeholder_count != 1:
        logger.debug(
            "Skipping textbox placeholder %s in %s TEXTBOX (found %d times, ambiguous). Leaving unchanged.",
            placeholder,
            region.value,
            placeholder_count,
        )
        return
    
    # Span mapping: Determine start run index + offset and end run index + offset
    placeholder_start = full_text.find(placeholder)
    placeholder_end = placeholder_start + len(placeholder)
    
    # Map placeholder position to runs
    current_pos = 0
    start_run_idx = None
    start_offset = None
    end_run_idx = None
    end_offset = None
    
    for i, run_text in enumerate(run_texts):
        # INVARIANT: current_pos must advance on every iteration to maintain correct
        # position tracking. It represents the cumulative character position across all
        # previous runs in the concatenated text.
        run_start = current_pos
        run_end = current_pos + len(run_text)
        
        # Check if placeholder starts in this run
        if start_run_idx is None and run_start <= placeholder_start < run_end:
            start_run_idx = i
            start_offset = placeholder_start - run_start
        
        # Check if placeholder ends in this run
        if run_start < placeholder_end <= run_end:
            end_run_idx = i
            end_offset = placeholder_end - run_start
            break
        
        # Advance current_pos to the end of this run before next iteration
        # This attempts to ensure run_start and run_end are computed correctly for subsequent runs
        current_pos = run_end
    
    # Abort on uncertainty: If span cannot be mapped cleanly, log and return unchanged
    if start_run_idx is None or end_run_idx is None:
        logger.warning(
            "Skipping textbox placeholder %s in %s TEXTBOX (cannot map span cleanly). Leaving unchanged.",
            placeholder,
            region.value,
        )
        return
    
    # Safe rewrite: Modify text ONLY
    # Start run: prefix + replacement
    # Middle runs: set text to empty string
    # End run: suffix
    
    # Update start run
    start_run = run_elems[start_run_idx]
    start_run_text = run_texts[start_run_idx]
    start_prefix = start_run_text[:start_offset] if start_offset > 0 else ""
    start_new_text = start_prefix + replacement
    _set_text_in_xml_run(start_run, start_new_text)
    
    # Update middle runs (if any)
    for i in range(start_run_idx + 1, end_run_idx):
        middle_run = run_elems[i]
        _set_text_in_xml_run(middle_run, "")
    
    # Update end run (if different from start run)
    if end_run_idx > start_run_idx:
        end_run = run_elems[end_run_idx]
        end_run_text = run_texts[end_run_idx]
        end_suffix = end_run_text[end_offset:] if end_offset < len(end_run_text) else ""
        _set_text_in_xml_run(end_run, end_suffix)
    
    logger.debug(
        "Successfully replaced scalar placeholder %s across %d run(s) in %s TEXTBOX",
        placeholder,
        end_run_idx - start_run_idx + 1,
        region.value,
    )


def _replace_text_in_runs(
    runs: List[Any],
    placeholder: str,
    replacement: str,
) -> None:
    """
    Replace placeholder text across a sequence of runs, handling multi-run spans.
    
    This function safely replaces placeholder text that may span multiple XML runs.
    It attempts to preserve run count and styling, modifying only text content. This is critical
    for handling placeholders that Word has split across multiple runs (e.g., due to
    formatting changes or XML structure), where structurally safe.
    
    Safety Guarantees:
        - Run count is preserved (no runs added or removed)
        - Run styling is preserved (formatting properties unchanged)
        - Only text content is modified (via _set_text_in_xml_run)
        - Ambiguous placeholders (appearing multiple times) are skipped
    
    Ambiguous Placeholder Handling:
        If a placeholder appears more than once in the same location, it is skipped
        for safety. This attempts to ensure deterministic output and prevent partial replacements
        that could corrupt document structure. The skip is logged at debug level.
    
    Args:
        runs: List of python-docx Run objects to search and replace within.
            The runs are assumed to be from the same paragraph or location.
        placeholder: Placeholder token to replace (e.g., "{{title}}").
        replacement: Replacement text to insert in place of the placeholder.
    
    Returns:
        None. The runs are modified in-place via _set_text_in_xml_run.
    
    Note:
        This function uses XML-level text extraction and replacement to ensure
        consistency with Word's internal structure. It does not use python-docx's
        high-level text properties to avoid node creation/removal.
    """
    if not runs:
        return

    # Build concatenated text from existing XML text nodes only.
    # NOTE: We intentionally do NOT use `run.text = ...` anywhere in this function,
    # because python-docx may create/remove `w:t` nodes, which is forbidden.
    # We only mutate existing `w:t.text` via `_set_text_in_xml_run`.
    full_text = "".join(_get_text_from_xml_run(run._r) for run in runs)
    if not full_text or placeholder not in full_text:
        return

    # SAFETY INVARIANT: Ambiguous placeholders (appearing multiple times) are skipped.
    # WHY: If a placeholder appears more than once in the same location, we cannot
    # deterministically decide which occurrence to replace. Skipping attempts to ensure deterministic
    # output and prevent partial replacements that could corrupt document structure.
    # Idempotent + deterministic safety rule: If the placeholder appears more than once,
    # skip entirely (no partial edits).
    count = full_text.count(placeholder)
    if count != 1:
        logger.debug(
            "Skipping scalar placeholder %s across runs (found %d times, ambiguous). Leaving unchanged.",
            placeholder,
            count,
        )
        return

    idx = full_text.find(placeholder)
    end = idx + len(placeholder)

    new_texts: List[str] = []
    current_start = 0
    inserted = False

    # Compute the rewritten text per run without mutating anything yet.
    for run in runs:
        original = _get_text_from_xml_run(run._r)
        run_len = len(original)
        run_start = current_start
        run_end = current_start + run_len

        # No overlap with placeholder span
        if run_end <= idx or run_start >= end:
            new_texts.append(original)
        else:
            before_len = max(0, idx - run_start)
            after_len = max(0, run_end - end)

            before = original[:before_len] if before_len > 0 else ""
            after = original[run_len - after_len :] if after_len > 0 else ""

            middle = ""
            if not inserted:
                middle = replacement
                inserted = True

            new_texts.append(before + middle + after)

        current_start += run_len

    # Transactional guard: we only mutate runs that already have at least one `w:t`.
    # If any run would need changing but has no `w:t`, we skip entirely.
    from lxml import etree
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    for run, new_text in zip(runs, new_texts):
        current_text = _get_text_from_xml_run(run._r)
        if new_text == current_text:
            continue
        xpath_method = etree._Element.xpath.__get__(run._r, etree._Element)
        text_elems = xpath_method(".//w:t", namespaces=ns)
        if not text_elems:
            logger.debug(
                "Skipping scalar placeholder %s across runs (run has no w:t; no node creation allowed). Leaving unchanged.",
                placeholder,
            )
            return

    # Commit: text-only mutation of existing XML nodes.
    for run, new_text in zip(runs, new_texts):
        _set_text_in_xml_run(run._r, new_text)


def _replace_header_footer_textboxes(
    doc: DocumentType,
    concrete: Dict[str, str],
    summary: Optional[ExportSummary] = None
) -> None:
    """
    Replace scalar placeholders inside header and footer textboxes (Phase 1, runs first).
    
    This function processes header and footer textboxes before other regions to maintain
    deterministic replacement order. Textboxes require special handling because they
    use a different XML structure (w:txbxContent) than regular paragraphs.
    
    Safety Guarantees:
        - Read-only preflight: Text extraction only, no mutations until validation passes
        - Replace only if placeholder appears exactly once (ambiguous = skip)
        - Replacement must be non-empty (empty replacements are skipped)
        - Multi-run safe: Handles placeholders spanning multiple XML runs
        - Atomic commit: Only mutates w:t.text values (no node creation/deletion/relocation)
        - If unsafe conditions detected, do nothing (leave content unchanged)
    
    Processing Order:
        This function MUST run first (before body and header/footer flow) to maintain
        deterministic replacement order and prevent cross-region interference during
        document traversal.
    
    Args:
        doc: Word document to modify. Headers and footers are accessed via doc.sections.
        concrete: Dictionary mapping placeholder tokens (with braces, e.g., "{{title}}")
            to replacement values. Values are guaranteed non-empty strings (empty values
            are filtered before calling this function).
        summary: Optional export summary for observability. If provided and
            enable_export_trace=True, successful replacements are recorded as
            ScalarReplacement entries.
    
    Returns:
        None. Textboxes are modified in-place via _safe_replace_scalar_across_textbox_runs.
    
    Note:
        This function uses XPath to find textboxes because python-docx does not provide
        high-level access to textbox content. The XPath query finds both WPS and classic/VML
        textboxes to ensure comprehensive coverage.
    
    SAFETY INVARIANT: Textboxes only support scalar replacement because Word's textbox
    XML structure (w:txbxContent) does not support block-level content insertion.
    Block placeholders in textboxes are skipped to prevent document corruption.
    
    Args:
        doc: Word document to modify
        concrete: Dictionary of placeholder tokens to replacement values
        summary: Optional export summary for observability (only used if enable_export_trace is True)
    """
    # Use unified textbox discovery; HEADER/FOOTER textboxes are filtered here.
    for run_elems, region, container in _iter_textbox_runs(doc):
        if region not in (Region.HEADER, Region.FOOTER):
            continue
        if container is not Container.TEXTBOX:
            continue

        # Build concatenated text from XML runs for detection
        full_text = "".join(_get_text_from_xml_run(run) for run in run_elems)
        if not full_text:
            continue

        # Only replace tokens that are in the authoritative mapping (scalar placeholders)
        for token, replacement in concrete.items():
            if token not in full_text:
                continue

            # Never replace empty or null values
            if not replacement or not replacement.strip():
                continue

            # Must appear exactly once
            if full_text.count(token) != 1:
                continue

            # SAFE cross-run replacement (no structural mutation)
            _safe_replace_scalar_across_textbox_runs(
                run_elems,
                token,
                replacement,
                region,
            )
            
            # Record successful replacement in summary (if enabled)
            if summary is not None:
                value_preview = replacement[:50] + ("..." if len(replacement) > 50 else "")
                summary.scalar_replacements.append(
                    ScalarReplacement(
                        token=token,
                        region=region.value,
                        container=Container.TEXTBOX.value,
                        value_preview=value_preview,
                    )
                )
                logger.info(
                    f"Replaced scalar placeholder {token} in {region.value} {Container.TEXTBOX.value} "
                    f"(preview: {value_preview})"
                )


def _replace_header_footer_flow_and_tables(
    doc: DocumentType,
    concrete: Dict[str, str],
    summary: Optional[ExportSummary] = None
) -> None:
    """
    Replace scalar placeholders in header and footer flow text and table cells (Phase 3).
    
    This function processes header and footer paragraphs and table cells, explicitly
    skipping anything inside w:txbxContent (textboxes are handled separately in Phase 1).
    It runs after body scalar replacement to maintain deterministic processing order.
    
    Safety Invariants:
        - Headers and footers are processed separately from body to maintain
          deterministic replacement order and prevent cross-region interference
        - Textboxes are explicitly skipped (handled in Phase 1)
        - Only FLOW and TABLE containers are processed (TEXTBOX excluded)
        - Ambiguous placeholders (multiple occurrences) are skipped
    
    Processing Order:
        This function runs as Phase 3, after:
        1. Phase 1: Header/footer textboxes
        2. Phase 2: Body scalar replacements (FLOW, TABLE, TEXTBOX)
        3. Phase 3: Header/footer flow and tables (this function)
    
    This order attempts to ensure deterministic replacement behavior and prevent cross-region
    interference during document traversal, subject to DOCX constraints.
    
    Args:
        doc: Word document to modify. Headers and footers are accessed via doc.sections.
        concrete: Dictionary mapping placeholder tokens (with braces, e.g., "{{title}}")
            to replacement values. Values are guaranteed non-empty strings.
        summary: Optional export summary for observability. If provided and
            enable_export_trace=True, successful replacements are recorded as
            ScalarReplacement entries.
    
    Returns:
        None. Paragraphs and table cells in headers/footers are modified in-place
        via _replace_text_in_runs.
    
    Note:
        This function explicitly skips textboxes by checking for w:txbxContent ancestors
        using XPath. This attempts to ensure textboxes are only processed in Phase 1 and not
        processed again here.
    """
    from lxml import etree
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    for paragraph, region in _iter_header_footer_paragraphs(doc):
        # Explicitly skip any paragraph that is inside a textbox content subtree
        try:
            xpath_method = etree._Element.xpath.__get__(paragraph._p, etree._Element)
            if xpath_method("ancestor::w:txbxContent", namespaces=ns):
                continue
        except Exception:
            # If we can't determine safely, skip.
            continue

        if not paragraph.runs:
            continue

        # Determine container type (FLOW or TABLE)
        container = Container.TABLE if _is_paragraph_in_table_cell(paragraph) else Container.FLOW

        # Use XML text extraction for consistency with replacement logic (attempts to preserve <w:br/> and <w:cr/> as \n)
        para_text = "".join(_get_text_from_xml_run(run._r) for run in paragraph.runs)
        if not para_text:
            continue

        for token, replacement in concrete.items():
            if not replacement or not replacement.strip():
                continue
            if token in para_text:
                _replace_text_in_runs(paragraph.runs, token, replacement)
                # Refresh para_text using same XML extraction method for consistency
                para_text = "".join(_get_text_from_xml_run(run._r) for run in paragraph.runs)
                
                # Record successful replacement in summary (if enabled)
                if summary is not None:
                    value_preview = replacement[:50] + ("..." if len(replacement) > 50 else "")
                    summary.scalar_replacements.append(
                        ScalarReplacement(
                            token=token,
                            region=region.value,
                            container=container.value,
                            value_preview=value_preview,
                        )
                    )
                    logger.info(
                        f"Replaced scalar placeholder {token} in {region.value} {container.value} "
                        f"(preview: {value_preview})"
                    )


def replace_scalar_placeholders(
    doc: DocumentType,
    mapping: Dict[str, Optional[str]],
    summary: Optional[ExportSummary] = None
) -> None:
    """
    Replace scalar placeholders (e.g., {{document_id}}, {{title}}) throughout the document.
    
    This function performs simple text substitution of scalar placeholders in all
    document regions and container types. Scalar replacements are inline text
    substitutions that do not change document structure (no paragraphs added/removed).
    
    Supported Locations:
        - Body paragraphs and table cells (FLOW and TABLE containers)
        - Header paragraphs and table cells (HEADER region)
        - Footer paragraphs and table cells (FOOTER region)
        - Textboxes in body, headers, and footers (TEXTBOX container)
    
    Processing Order (Deterministic):
        1. Header and footer textboxes (XML-only, processed first)
        2. Body scalar replacements (FLOW and TABLE containers, including body textboxes)
        3. Header and footer non-textbox text (FLOW and TABLE containers)
    
    This order attempts to ensure deterministic replacement behavior and prevent cross-region
    interference during document traversal, subject to DOCX constraints.
    
    Safety Guarantees:
        - Only replaces inline text where found; never inserts or removes paragraphs
        - Attempts to preserve document structure and formatting where structurally safe
        - Handles placeholders spanning multiple XML runs (multi-run safe) where detected
        - Ambiguous placeholders (appearing multiple times in same location) are skipped
        - Empty or None replacement values result in empty string substitution
    
    Placeholder Provenance:
        Only placeholders present in the `mapping` parameter will be replaced.
        Any placeholder found in the document but not in `mapping` will be ignored
        (logged at debug level). This enforces authoritative placeholder registry
        and prevents accidental replacements.
    
    Ambiguous Placeholder Handling:
        If a placeholder appears multiple times in the same location (e.g., same
        paragraph or same textbox), it is skipped for safety. This attempts to ensure
        deterministic output and prevent partial replacements that could corrupt
        document structure. The skip is logged and recorded in export summary.
    
    Args:
        doc: Word document to modify. The document is modified in-place.
        mapping: Dictionary mapping placeholder keys (without braces) to replacement
            values. Keys like "document_id" will match placeholders like "{{document_id}}".
            Values can be None or empty string (results in empty replacement).
        summary: Optional export summary for observability. If provided and
            enable_export_trace=True, successful replacements are recorded as
            ScalarReplacement entries. Skipped replacements are logged but not
            recorded (they are recorded as SkippedItem in block replacement logic).
    
    Returns:
        None. The document is modified in-place.
    
    Example:
        ```python
        mapping = {
            "document_id": "DOC-12345",
            "title": "My Document",
            "author": "John Doe"
        }
        replace_scalar_placeholders(doc, mapping)
        # Replaces {{document_id}} with "DOC-12345", {{title}} with "My Document", etc.
        ```
    """
    # Build concrete placeholder -> replacement mapping
    # This enforces authoritative placeholder registry - only keys from mapping are replaced
    concrete: Dict[str, Optional[str]] = {}
    for key, value in mapping.items():
        token = f"{{{{{key}}}}}"  # e.g. document_id -> {{document_id}}
        concrete[token] = value

    if not concrete:
        return

    # Build set of allowed placeholders for provenance checking
    allowed_tokens = set(concrete.keys())

    # Phase 1 — HEADER & FOOTER TEXTBOXES (XML-only, FIRST)
    # SAFETY INVARIANT: Header/footer textboxes are processed first to maintain
    # deterministic replacement order and avoid cross-region interference.
    _replace_header_footer_textboxes(doc, {k: (v or "") for k, v in concrete.items()}, summary=summary)

    # Phase 2 — BODY SCALAR REPLACEMENT (Unchanged)
    for location_data, region, container in _iter_all_locations(doc):
        if region != Region.BODY:
            continue

        if container == Container.TEXTBOX:
            run_elems = location_data  # list of XML w:r elements
            if not run_elems:
                continue

            full_text = "".join(_get_text_from_xml_run(run) for run in run_elems)
            if not full_text:
                continue

            for token, replacement in concrete.items():
                if token not in full_text:
                    continue
                if not replacement or not replacement.strip():
                    logger.debug(
                        "Skipping %s TEXTBOX replacement for %s (empty replacement)",
                        region.value,
                        token,
                    )
                    continue
                # SAFETY INVARIANT: Ambiguous placeholders (appearing multiple times) are skipped.
                # WHY: If a placeholder appears more than once in the same location, we cannot
                # deterministically decide which occurrence to replace. Skipping attempts to ensure deterministic
                # output and prevent partial replacements that could corrupt document structure.
                if full_text.count(token) != 1:
                    logger.debug(
                        "Skipping %s TEXTBOX replacement for %s (ambiguous occurrence)",
                        region.value,
                        token,
                    )
                    continue

                _safe_replace_scalar_across_textbox_runs(run_elems, token, replacement, region)
                
                # Record successful replacement in summary (if enabled)
                if summary is not None:
                    value_preview = replacement[:50] + ("..." if len(replacement) > 50 else "")
                    summary.scalar_replacements.append(
                        ScalarReplacement(
                            token=token,
                            region=region.value,
                            container=Container.TEXTBOX.value,
                            value_preview=value_preview,
                        )
                    )
                    logger.info(
                        f"Replaced scalar placeholder {token} in {region.value} {Container.TEXTBOX.value} "
                        f"(preview: {value_preview})"
                    )

        else:
            paragraph = location_data  # Paragraph in BODY FLOW/TABLE
            if not paragraph.runs:
                continue
            # Use XML text extraction for consistency with replacement logic (attempts to preserve <w:br/> and <w:cr/> as \n)
            para_text = "".join(_get_text_from_xml_run(run._r) for run in paragraph.runs)
            if not para_text:
                continue

            for token, replacement in concrete.items():
                if not replacement or not replacement.strip():
                    continue
                if token in para_text:
                    logger.debug(f"Replacing scalar placeholder {token} in {region.value} {container.value}")
                    _replace_text_in_runs(paragraph.runs, token, replacement)
                    # Refresh para_text using same XML extraction method for consistency
                    para_text = "".join(_get_text_from_xml_run(run._r) for run in paragraph.runs)
                    
                    # Record successful replacement in summary (if enabled)
                    if summary is not None:
                        value_preview = replacement[:50] + ("..." if len(replacement) > 50 else "")
                        summary.scalar_replacements.append(
                            ScalarReplacement(
                                token=token,
                                region=region.value,
                                container=container.value,
                                value_preview=value_preview,
                            )
                        )
                        logger.info(
                            f"Replaced scalar placeholder {token} in {region.value} {container.value} "
                            f"(preview: {value_preview})"
                        )

    # Phase 3 — HEADER & FOOTER NON-TEXTBOX TEXT (skip w:txbxContent)
    # SAFETY INVARIANT: Header/footer flow and tables are processed after body
    # to maintain deterministic replacement order and document stability.
    _replace_header_footer_flow_and_tables(doc, concrete, summary=summary)


def _is_paragraph_in_table_cell(paragraph: Paragraph) -> bool:
    """
    Check if a paragraph is located inside a table cell.
    
    This function determines the container type of a paragraph by inspecting
    its XML parent element. Table cells have the XML tag w:tc, which is used
    to identify table cell context.
    
    Args:
        paragraph: python-docx Paragraph object to check.
    
    Returns:
        True if the paragraph is inside a table cell (w:tc element),
        False otherwise (paragraph is in body flow or header/footer).
    
    Note:
        This function uses XML inspection because python-docx does not
        provide a direct API to check if a paragraph is in a table cell.
        The check is based on the XML tag of the paragraph's parent element.
    """
    parent = paragraph._p.getparent()
    if parent is None:
        return False
    return parent.tag.endswith('}tc')  # w:tc element


def _get_cell_for_paragraph(paragraph: Paragraph):
    """
    Get the table cell object that contains this paragraph.
    
    This function retrieves the python-docx Cell object for a paragraph that
    is located inside a table cell. It uses the paragraph's parent relationship
    to navigate from the XML element to the Cell object.
    
    Args:
        paragraph: python-docx Paragraph object that is inside a table cell.
            The paragraph must have been determined to be in a table cell
            (e.g., via _is_paragraph_in_table_cell) before calling this function.
    
    Returns:
        python-docx Cell object if the paragraph is in a table cell,
        None if the paragraph is not in a table cell or parent cannot be determined.
    
    Note:
        This function accesses python-docx internals (_parent, _tc) to navigate
        from paragraph to cell. This is intentional and matches the library's
        approach of using internals for precise XML control.
    """
    parent = paragraph._parent
    if hasattr(parent, "_tc"):
        return parent
    return None


def _set_paragraph_formatted_text(paragraph: Paragraph, formatted_runs: Optional[List[FormattedRun]], fallback_text: str = "") -> None:
    """
    Set paragraph text with formatting, preserving paragraph-level styles.
    
    This function sets paragraph text while preserving all paragraph-level formatting
    (styles, spacing, numbering, template formatting). It clears existing run text
    but does not remove runs or modify paragraph structure.
    
    Safety Guarantees:
        - Attempts to preserve paragraph-level styles, spacing, numbering, template formatting
        - Only clears run text; does not remove runs or modify paragraph structure
        - Attempts to preserve bold/italic formatting from FormattedRun objects where supported
        - Falls back to plain text if formatted_runs are not available
    
    Why This Approach:
        Using paragraph.clear() would remove paragraph-level formatting and structure.
        By only clearing run text, we attempt to preserve paragraph properties while updating
        content. This attempts to ensure template formatting (styles, spacing, etc.) is maintained.
    
    Args:
        paragraph: python-docx Paragraph object to update.
        formatted_runs: Optional list of FormattedRun objects with bold/italic formatting.
            If provided, these are used to set formatted text. If None, fallback_text is used.
        fallback_text: Plain text to use if formatted_runs is None or empty.
            Defaults to empty string.
    
    Returns:
        None. The paragraph is modified in-place.
    
    Note:
        Line breaks (<w:br/>, <w:cr/>) in formatted runs are preserved as textual
        '\\n' characters. This library does not synthesize visual Word line breaks.
        Text is preserved deterministically without guessing layout.
    """
    # Clear existing run text only (preserve paragraph-level formatting: styles, spacing, numbering, etc.)
    for run in paragraph.runs:
        run.text = ""
    
    if formatted_runs:
        # Add formatted runs
        for run_data in formatted_runs:
            run = paragraph.add_run(run_data.text)
            run.bold = run_data.bold
            run.italic = run_data.italic
    elif fallback_text:
        # Use run-preserving fallback (mutate text only, not structure)
        if paragraph.runs:
            paragraph.runs[0].text = fallback_text
        else:
            paragraph.add_run(fallback_text)


def _heading_style_for_level(level: Optional[int]) -> str:
    """
    Get Word heading style name for a given heading level.
    
    This function maps markdown heading levels (1-6) to Word heading style names.
    Levels outside the valid range are clamped to valid values for safety.
    
    Args:
        level: Heading level (1-6 for H1-H6). Can be None or outside range (clamped).
    
    Returns:
        Word heading style name (e.g., "Heading 1", "Heading 2", etc.).
        Returns "Heading 1" if level is None or < 1.
        Returns "Heading 9" if level > 9 (Word supports up to Heading 9).
    """
    if level is None or level < 1:
        return "Heading 1"
    if level > 9:
        level = 9
    return f"Heading {level}"


def _try_apply_word_table_style(tbl_elem, doc_body):
    """
    Attempt to apply a Word built-in table style to the table element.
    
    This function tries to apply Word's built-in table styles, preferring "TableGrid"
    as the primary choice for cross-platform compatibility. This style is widely
    supported and provides consistent visual appearance across macOS and Windows Word.
    
    IMPORTANT: Even when a style is successfully applied, explicit XML borders must
    still be set (see _apply_xml_table_borders_fallback). Windows Word requires explicit
    border definitions in the XML, as style-based borders may not render correctly
    across all document templates and Word versions.
    
    Safety Guarantees:
        - Graceful failure: Returns None if style application fails (no exception)
        - Style preference: Prefers "TableGrid" for maximum compatibility
        - Deterministic: Same table structure attempts to try same styles in same order
    
    Args:
        tbl_elem: XML table element (w:tbl) to format. Must be a valid table element.
        doc_body: Document body element (python-docx internal: doc._body) required
            for creating Table object and applying styles. If None, returns None.
    
    Returns:
        Table object if a style was successfully applied, None otherwise.
        The Table object can be used for further formatting (e.g., header row emphasis).
    
    Note:
        This function uses python-docx's Table object to apply styles. However, style
        application alone is not sufficient for Windows Word compatibility - explicit
        XML borders must always be applied as a mandatory fallback.
    """
    if doc_body is None:
        return None
    
    try:
        from docx.table import Table
        table = Table(tbl_elem, doc_body)
        
        # Prefer "TableGrid" first for maximum cross-platform compatibility
        # This style is universally supported and provides consistent grid appearance
        for style_name in ['TableGrid', 'Light Grid', 'Medium Grid 1']:
            try:
                table.style = style_name
                return table
            except Exception:
                continue
        
        return None
    except Exception:
        return None


def _apply_header_row_emphasis(table_obj):
    """
    Apply bold formatting to header row (first row) of a table.
    
    This function makes all text in the first row bold to visually distinguish
    it as a header row. This is a common table formatting convention that
    improves readability.
    
    Args:
        table_obj: Table object from python-docx. Must have at least one row.
            If the table is empty, no changes are made.
    
    Returns:
        None. The table is modified in-place (first row text becomes bold).
    
    Note:
        This function modifies all runs in all paragraphs in all cells of the
        first row. It does not check if the row is already bold, making it
        idempotent (safe to call multiple times).
    """
    if len(table_obj.rows) > 0:
        header_row = table_obj.rows[0]
        for cell in header_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True


def _apply_xml_table_borders_fallback(tbl_elem):
    """
    Apply explicit XML borders as mandatory fallback for cross-platform compatibility.
    
    This function applies explicit XML border formatting to a table, ensuring all borders
    are explicitly defined in the DOCX XML. This is REQUIRED for Windows Word compatibility,
    as Windows Word requires explicit border definitions in the XML even when table styles
    are applied. macOS Word may render style-based borders correctly, but Windows Word
    requires explicit XML borders to display table formatting consistently.
    
    IMPORTANT: This function is called even when a table style is successfully applied,
    because style-based borders may not render correctly on Windows Word across all document
    templates and Word versions. Explicit XML borders ensure deterministic, cross-platform
    table formatting.
    
    Safety Guarantees:
        - Deterministic: Same table structure always gets same border formatting
        - Cross-platform: Explicit borders ensure Windows Word compatibility
        - Mandatory: Borders are always applied, even if style was applied successfully
    
    Why Explicit Borders Are Required:
        - Windows Word requires explicit <w:tblBorders> XML elements to display borders
        - macOS Word may render style-based borders, but this is not guaranteed across templates
        - Explicit XML borders provide deterministic, platform-independent table formatting
        - Without explicit borders, tables may appear unformatted (no borders) on Windows Word
    
    Args:
        tbl_elem: XML table element (w:tbl) to format. Must be a valid table element.
            The function creates or modifies w:tblPr and w:tblBorders elements as needed.
    
    Returns:
        None. The table element is modified in-place with border formatting.
    
    Note:
        Border properties are deterministic: single line, size 4 (0.5pt), auto color, no spacing.
        All borders (top, bottom, left, right, insideH, insideV) are explicitly defined.
        This creates a clean, visible grid that works consistently across macOS and Windows Word.
    """
    from docx.oxml.ns import qn
    
    # Find or create w:tblPr element
    tbl_pr = None
    for child in tbl_elem:
        if child.tag == qn('w:tblPr'):
            tbl_pr = child
            break
    
    if tbl_pr is None:
        # Create tblPr if it doesn't exist
        tbl_pr = tbl_elem.makeelement(qn('w:tblPr'))
        # Insert at the beginning
        tbl_elem.insert(0, tbl_pr)
    
    # Check if w:tblBorders already exists
    tbl_borders = None
    for child in tbl_pr:
        if child.tag == qn('w:tblBorders'):
            tbl_borders = child
            break
    
    # Always create/overwrite borders to ensure explicit definitions
    # This ensures Windows Word compatibility even if borders were partially defined
    if tbl_borders is None:
        # Create tblBorders element
        tbl_borders = tbl_pr.makeelement(qn('w:tblBorders'))
        tbl_pr.append(tbl_borders)
    else:
        # Clear existing borders to ensure deterministic, explicit definitions
        # This ensures all borders are explicitly set with consistent values
        for border_elem in list(tbl_borders):
            tbl_borders.remove(border_elem)
    
    # Define border properties: single line, size 4 (0.5pt), auto color, no spacing
    # These are deterministic values that work consistently across platforms
    border_attrs = {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '0',
        qn('w:color'): 'auto',
    }
    
    # Apply borders to all sides: top, bottom, left, right, insideH (horizontal), insideV (vertical)
    # All borders must be explicitly defined for Windows Word compatibility
    border_names = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
    for border_name in border_names:
        border_elem = tbl_borders.makeelement(qn(f'w:{border_name}'))
        for attr_name, attr_value in border_attrs.items():
            border_elem.set(attr_name, attr_value)
        tbl_borders.append(border_elem)


def _apply_default_table_formatting(tbl_elem, doc_body=None):
    """
    Apply default visible grid formatting to a Word table XML element.
    
    This function provides deterministic, cross-platform table formatting by:
    1. Attempting to apply Word built-in table style ("TableGrid" preferred)
    2. ALWAYS applying explicit XML borders as mandatory fallback
    
    IMPORTANT: Explicit XML borders are ALWAYS applied, even when a table style is
    successfully applied. This is required for Windows Word compatibility, as Windows
    Word requires explicit border definitions in the XML to display table formatting
    correctly. macOS Word may render style-based borders, but explicit XML borders ensure
    consistent rendering across all platforms and document templates.
    
    Formatting Strategy:
        1. Try built-in Word styles in order: "TableGrid" (preferred), "Light Grid", "Medium Grid 1"
        2. If style succeeds, make header row text bold
        3. ALWAYS apply explicit XML borders (mandatory for Windows Word compatibility)
        4. If style fails or doc_body is None, still apply explicit XML borders
    
    Safety Guarantees:
        - Deterministic: Same table structure always gets same formatting
        - Cross-platform: Explicit borders ensure Windows Word compatibility
        - Mandatory borders: Borders are always explicitly defined in XML
        - Style + borders: Style provides visual enhancement, borders provide compatibility
    
    Args:
        tbl_elem: XML table element (w:tbl) to format. Must be a valid table element.
        doc_body: Optional document body element (python-docx internal: doc._body)
            for style application via Table object. If None, skips style application
            but still applies explicit XML borders.
    
    Returns:
        None. The table element is modified in-place with formatting.
    
    Note:
        This function ensures tables have visible formatting on both macOS and Windows Word.
        The combination of table style (when available) and explicit XML borders provides
        maximum compatibility and consistent visual appearance across platforms.
    """
    # Try to apply table style first (prefers "TableGrid" for compatibility)
    table_obj = _try_apply_word_table_style(tbl_elem, doc_body)
    if table_obj is not None:
        _apply_header_row_emphasis(table_obj)
    
    # ALWAYS apply explicit XML borders, even if style was applied successfully
    # This is mandatory for Windows Word compatibility - Windows Word requires explicit
    # border definitions in the XML, even when table styles are applied. macOS Word may
    # render style-based borders correctly, but explicit XML borders ensure consistent
    # cross-platform rendering across all document templates and Word versions.
    _apply_xml_table_borders_fallback(tbl_elem)


def _attempt_render_nested_word_table(
    block,
    *,
    parent_cell_elem,
    insertion_index,
    render_context,
    summary: Optional[ExportSummary] = None
) -> bool:
    """
    Attempt to render a markdown table block as a real nested Word table
    inside a table cell, with strict validation and atomic commit.

    Returns True only when the nested table is built, validated, and inserted.
    Returns False (with INFO log) on any failure; must not raise.
    
    SAFETY INVARIANT: Nested table rendering uses atomic strategy (scratch → clone → insert)
    to attempt to prevent partial document mutation. If any validation fails, zero changes are made.
    This attempts to ensure document stability and prevent Word corruption.
    
    Args:
        block: MarkdownBlock representing the table
        parent_cell_elem: XML element of the parent table cell (w:tc)
        insertion_index: Index where table should be inserted
        render_context: Optional rendering context (unused, for future extensibility)
        summary: Optional export summary for observability (only used if enable_export_trace is True)
    """
    # INVARIANT:
    # Nested table rendering is strictly atomic:
    # - Build in scratch document
    # - Validate fully
    # - Clone XML
    # - Insert exactly once
    # Any failure MUST result in zero document mutation.
    
    from copy import deepcopy
    from docx import Document
    from docx.oxml.ns import qn
    from lxml import etree

    try:
        header = block.header or []
        rows = block.rows or []

        # Determine number of columns
        num_cols = len(header) if header else (len(rows[0]) if rows else 0)
        if num_cols <= 0:
            fallback_msg = "Nested table could not be rendered safely; falling back to text grid."
            logger.info(fallback_msg)
            # Record fallback event in summary (if enabled)
            if summary is not None:
                summary.fallback_events.append(
                    FallbackEvent(
                        event_type="nested_table_fallback",
                        location=None,
                        reason="Table has zero columns (invalid structure)",
                    )
                )
                logger.info(f"Fallback: {fallback_msg}")
            return False

        # Build the table in isolation using a scratch document
        scratch_doc = Document()
        scratch_table = scratch_doc.add_table(rows=0, cols=num_cols)

        # Populate header row (if any)
        if header:
            hdr_row = scratch_table.add_row()
            for idx, cell_text in enumerate(header[:num_cols]):
                hdr_row.cells[idx].text = str(cell_text) if cell_text else ""

        # Populate data rows
        for row in rows:
            data_row = scratch_table.add_row()
            for idx, cell_text in enumerate(row[:num_cols]):
                data_row.cells[idx].text = str(cell_text) if cell_text else ""

        # Apply formatting BEFORE cloning (using scratch document context for style resolution)
        _apply_default_table_formatting(scratch_table._tbl, scratch_doc._body)

        # Extract and clone the already-formatted w:tbl element
        tbl_elem = scratch_table._tbl
        cloned_tbl = deepcopy(tbl_elem)

        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        # Validate: root element must be a table
        if not cloned_tbl.tag.endswith("}tbl"):
            fallback_msg = "Nested table could not be rendered safely; falling back to text grid."
            logger.info(fallback_msg)
            # Record fallback event in summary (if enabled)
            if summary is not None:
                summary.fallback_events.append(
                    FallbackEvent(
                        event_type="nested_table_fallback",
                        location=None,
                        reason="Cloned element is not a valid table (tag validation failed)",
                    )
                )
                logger.info(f"Fallback: {fallback_msg}")
            return False

        # SAFETY INVARIANT: Disallow section properties or numbering XML in nested tables
        # WHY: These XML elements are document-level and cannot appear inside table cells.
        # Their presence would cause Word corruption. This validation attempts to prevent invalid
        # document structure and ensure deterministic output.
        disallowed_xpaths = [
            ".//w:sectPr",
            ".//w:numPr",
            ".//w:abstractNum",
        ]
        for xp in disallowed_xpaths:
            xpath_method = etree._Element.xpath.__get__(cloned_tbl, etree._Element)
            if xpath_method(xp, namespaces=ns):
                fallback_msg = "Nested table could not be rendered safely; falling back to text grid."
                logger.info(fallback_msg)
                # Record fallback event in summary (if enabled)
                if summary is not None:
                    summary.fallback_events.append(
                        FallbackEvent(
                            event_type="nested_table_fallback",
                            location=None,
                            reason=f"Disallowed XML element found in nested table: {xp}",
                        )
                    )
                    logger.info(f"Fallback: {fallback_msg}")
                return False

        # Validate parent container is exactly w:tc
        if parent_cell_elem is None or parent_cell_elem.tag != qn("w:tc"):
            logger.error(
                "Nested table parent is not w:tc; tag=%s",
                parent_cell_elem.tag if parent_cell_elem is not None else None,
            )
            fallback_msg = "Nested table could not be rendered safely; falling back to text grid."
            logger.info(fallback_msg)
            # Record fallback event in summary (if enabled)
            if summary is not None:
                summary.fallback_events.append(
                    FallbackEvent(
                        event_type="nested_table_fallback",
                        location=None,
                        reason=f"Invalid parent container (expected w:tc, got {parent_cell_elem.tag if parent_cell_elem else None})",
                    )
                )
                logger.info(f"Fallback: {fallback_msg}")
            return False

        # Atomic commit: ensure there are paragraph siblings, then insert at position.
        # The original placeholder paragraph has already been removed earlier.
        try:
            # Ensure at least one <w:p> exists to satisfy Word rendering expectations.
            has_any_para = any(child.tag == qn("w:p") for child in parent_cell_elem)
            if not has_any_para:
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls

                p_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
                parent_cell_elem.insert(insertion_index, p_xml)
                insertion_index += 1

            # Insert the validated table
            parent_cell_elem.insert(insertion_index, cloned_tbl)

            # Ensure there is also a paragraph AFTER the table so that the
            # cell structure is: <w:p/> <w:tbl/> <w:p/>.
            children = list(parent_cell_elem)
            tbl_index = children.index(cloned_tbl)
            if tbl_index == len(children) - 1:
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls

                p_xml_after = parse_xml(f'<w:p {nsdecls("w")}/>')
                parent_cell_elem.insert(tbl_index + 1, p_xml_after)
        except Exception as e:
            fallback_msg = "Nested table could not be rendered safely; falling back to text grid."
            logger.info(fallback_msg, exc_info=True)
            # Record fallback event in summary (if enabled)
            if summary is not None:
                summary.fallback_events.append(
                    FallbackEvent(
                        event_type="nested_table_fallback",
                        location=None,
                        reason=f"Exception during nested table insertion: {type(e).__name__}: {str(e)}",
                    )
                )
                logger.info(f"Fallback: {fallback_msg}")
            return False

        logger.info(
            "Nested Word table committed into table cell at index %s",
            insertion_index,
        )
        return True

    except Exception as e:
        fallback_msg = "Nested table could not be rendered safely; falling back to text grid."
        logger.info(fallback_msg, exc_info=True)
        # Record fallback event in summary (if enabled)
        if summary is not None:
            summary.fallback_events.append(
                FallbackEvent(
                    event_type="nested_table_fallback",
                    location=None,
                    reason=f"Exception during nested table rendering: {type(e).__name__}: {str(e)}",
                )
            )
            logger.info(f"Fallback: {fallback_msg}")
        return False


def replace_block_placeholder_with_content(
    doc: DocumentType,
    paragraph: Paragraph,
    placeholder_token: str,
    blocks: Optional[List[MarkdownBlock]],
    config: "ListRenderConfig",
    summary: Optional[ExportSummary] = None,
) -> None:
    """
    Replace a block placeholder (e.g., {{summary}}) with structured markdown content.
    
    This function attempts to replace a block placeholder paragraph with parsed markdown content
    (headings, paragraphs, lists, tables) where structurally safe. The placeholder paragraph
    is removed and replaced with new content blocks inserted at the captured position.
    
    Safety Invariants:
        - Block placeholders MUST occupy entire paragraphs (not mixed with other text)
        - Block expansion is restricted to BODY containers only (headers/footers/textboxes excluded)
        - Placeholder paragraph is removed before inserting new content
        - Insertion position is captured before removal to attempt exact placement
        - Empty or None blocks result in placeholder removal only (no content inserted)
    
    Why Entire Paragraph Requirement:
        Partial-paragraph replacement would create ambiguous layout and non-deterministic
        output. If a placeholder shares a paragraph with other text, we cannot
        deterministically decide what to preserve. This constraint attempts to ensure document
        stability and prevent Word corruption. Placeholders mixed with text are
        automatically skipped (logged as warning, recorded in summary).
    
    Container Support:
        - BODY FLOW: Attempts to support block types (headings, paragraphs, lists, tables) where supported
        - BODY TABLE: Attempts to support block types within table cells, subject to DOCX constraints
        - HEADER/FOOTER/TEXTBOX: Not supported (block expansion restricted to BODY only)
    
    Content Insertion:
        New content blocks are inserted at the captured position where the placeholder
        paragraph was located. For BODY FLOW, insertion is into the document body.
        For TABLE cells, insertion is into the cell element. Insertion order attempts to
        preserve the blocks list order, subject to DOCX constraints.
    
    Args:
        doc: Word document being modified. Used for creating new paragraphs and
            accessing document body structure.
        paragraph: The paragraph containing the block placeholder. This paragraph
            will be removed and replaced with new content. Must be in BODY region
            (FLOW or TABLE container).
        placeholder_token: The placeholder token as it appears in the template,
            including braces (e.g., "{{summary}}", "{{proposal}}"). This is used
            for validation to ensure the paragraph contains only this placeholder.
        blocks: List of MarkdownBlock objects to insert. Can be None or empty
            (placeholder is removed but no content inserted). Blocks are inserted
            in the order provided.
        config: ListRenderConfig for controlling list rendering behavior (indentation,
            glyphs, deep nesting strategy). Must not be None.
        summary: Optional export summary for observability. If provided and
            enable_export_trace=True, successful replacements are recorded as
            BlockReplacement entries. Skipped replacements are recorded as SkippedItem.
    
    Returns:
        None. The document is modified in-place. The placeholder paragraph is removed
        and replaced with new content blocks (if any).
    
    Raises:
        ValueError: If config is None (defensive guard).
    
    Example:
        ```python
        blocks = [
            MarkdownBlock(type="heading", level=1, text="Introduction"),
            MarkdownBlock(type="paragraph", text="This is the introduction content."),
            MarkdownBlock(type="bullet_list", items=[...])
        ]
        replace_block_placeholder_with_content(doc, para, "{{summary}}", blocks, config)
        # Attempts to replace {{summary}} paragraph with heading, paragraph, and bullet list
        ```
    """
    # INVARIANT:
    # Paragraph insertion is performed inline (no shared helper) to attempt to preserve
    # XML insertion order and avoid accidental cross-container mutation.
    # This duplication is intentional for determinism and safety.
    
    # Defensive guard: ensure config is never None
    if config is None:
        raise ValueError("config must not be None")
    
    # Determine container: BODY FLOW or TABLE CELL
    # Get cell reference and actual XML parent BEFORE removing paragraph
    actual_parent_elem = paragraph._p.getparent()
    cell = _get_cell_for_paragraph(paragraph)
    is_in_table_cell = cell is not None
    cell_elem = actual_parent_elem if is_in_table_cell and actual_parent_elem is not None else None  # type: ignore
    
    # SAFETY INVARIANT: Block placeholders must occupy entire paragraphs
    # WHY: Partial-paragraph replacement would create ambiguous layout and non-deterministic output.
    # The placeholder paragraph is removed and replaced with new content. If the placeholder
    # shares the paragraph with other text, we cannot deterministically decide what to preserve.
    # This constraint ensures document stability and prevents Word corruption.
    # Reconstruct full paragraph text to check if placeholder is the only content
    # Use XML text extraction for consistency with replacement logic (preserves <w:br/> and <w:cr/> as \n)
    para_text = "".join(_get_text_from_xml_run(run._r) for run in paragraph.runs)
    if para_text.strip() != placeholder_token:
        logger.warning(
            f"Block placeholder {placeholder_token} is not the only content in its paragraph. "
            f"Paragraph contains: {para_text[:100]}. "
            "Block placeholders must occupy the entire paragraph. Skipping replacement."
        )
        return
    
    # Capture insertion point BEFORE removing paragraph
    # This attempts to ensure new content is inserted at the captured position of the placeholder
    insertion_parent = None
    insertion_index = None
    cell_insertion_index = None
    if is_in_table_cell and cell_elem is not None:
        cell_insertion_index = list(cell_elem).index(paragraph._p)
    else:
        if actual_parent_elem is not None:
            insertion_parent = actual_parent_elem
            insertion_index = actual_parent_elem.index(paragraph._p)

    # Remove the placeholder paragraph from its parent container.
    # For BODY FLOW we remove it immediately; for TABLE CELL we may delay removal
    # to use it as an anchor for nested table insertion.
    placeholder_removed = False
    if not is_in_table_cell:
        paragraph._p.getparent().remove(paragraph._p)
        placeholder_removed = True
    
    # If no content blocks, we may still need to ensure the cell is not structurally empty
    if not blocks:
        if is_in_table_cell and cell_elem is not None:
            from lxml import etree
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            xpath_method = etree._Element.xpath.__get__(cell_elem, etree._Element)
            existing_paras = xpath_method(".//w:p", namespaces=ns)
            if not existing_paras:
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                p_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
                insert_at = cell_insertion_index if cell_insertion_index is not None else 0
                cell_elem.insert(insert_at, p_xml)
        return

    # Insert all new content at the correct position (BODY FLOW or TABLE CELL)
    # Track previous block type to determine if spacing should be applied
    previous_block_type = None
    is_first_block = True
    
    # BUG #2 FIX: Capture original cell insertion index for "first paragraph" detection
    original_cell_insertion_index = cell_insertion_index
    
    for block in blocks:
        # BUG #1 FIX: Capture current block type at start of iteration
        current_block_type = block.type
        # For table-cell blocks other than tables, ensure the placeholder paragraph has been removed.
        if is_in_table_cell and block.type != "table" and not placeholder_removed:
            parent_for_para = paragraph._p.getparent()
            if parent_for_para is not None:
                parent_for_para.remove(paragraph._p)
            placeholder_removed = True

        if block.type == "heading":
            text = block.text or ""
            style = _heading_style_for_level(block.level)

            # Create paragraph directly in container
            if is_in_table_cell and cell_elem is not None:
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                p_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
                insert_at = cell_insertion_index if cell_insertion_index is not None else len(cell_elem)
                cell_elem.insert(insert_at, p_xml)
                new_p = Paragraph(p_xml, cell)  # type: ignore[arg-type]
                cell_insertion_index = insert_at + 1
            else:
                # BODY FLOW: insert at captured position
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                p_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
                insertion_parent.insert(insertion_index, p_xml)
                new_p = Paragraph(p_xml, doc._body)
                insertion_index += 1
            
            new_p.style = style
            _set_paragraph_formatted_text(new_p, block.formatted_runs, text)

        elif block.type == "paragraph":
            text = block.text or ""
            
            # Create paragraph directly in container
            if is_in_table_cell and cell_elem is not None:
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                p_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
                insert_at = cell_insertion_index if cell_insertion_index is not None else len(cell_elem)
                
                # BUG #2 FIX: Use index comparison instead of DOM counting
                # A paragraph is first in cell if insertion index equals original placeholder index
                is_first_para_in_cell = (cell_insertion_index == original_cell_insertion_index)
                
                cell_elem.insert(insert_at, p_xml)
                new_p = Paragraph(p_xml, cell)  # type: ignore[arg-type]
                cell_insertion_index = insert_at + 1
            else:
                # BODY FLOW: insert at captured position
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                p_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
                insertion_parent.insert(insertion_index, p_xml)
                new_p = Paragraph(p_xml, doc._body)
                insertion_index += 1
                is_first_para_in_cell = False  # Not applicable for BODY FLOW
            
            # Keep existing style (body style from template) - no style change for paragraphs
            _set_paragraph_formatted_text(new_p, block.formatted_runs, text)

            # List continuation paragraphs: semantic contract from parsing phase
            # SEMANTIC CONTRACT: The _list_continuation_level private attribute is set during parsing
            # to indicate that this paragraph visually belongs to a list item. It must be indented
            # like the list level (same left_indent and hanging_indent) but without bullet/number
            # glyphs. This attribute must not be removed or renamed without updating both parsing
            # (where it's set) and rendering (where it's consumed) to maintain list semantics.
            continuation_level = getattr(block, "_list_continuation_level", None)
            if isinstance(continuation_level, int) and continuation_level > 0:
                max_depth = config.max_visual_depth if config.max_visual_depth is not None else 3
                indent_per_level = config.indent_inches_per_level
                hanging_indent = -config.hanging_indent_inches
                visual_level = min(continuation_level, max_depth)
                new_p.paragraph_format.left_indent = Inches(indent_per_level * visual_level)
                # Use same hanging indent as list items so text aligns with list body text (not the glyph)
                new_p.paragraph_format.first_line_indent = Inches(hanging_indent)
                
                # Neutralize Word default paragraph spacing for continuation paragraphs (spacing comes from markdown, not Word defaults)
                from docx.shared import Pt
                new_p.paragraph_format.space_before = Pt(0)
                new_p.paragraph_format.space_after = Pt(0)
            
            # HOOK #1 & #2: Apply spacing to paragraphs following lists or tables
            # BODY FLOW: apply if not first block and previous was list/table
            # TABLE CELL: apply if not first para in cell and previous was list/table
            # Do NOT apply to continuation paragraphs or between list items
            elif config.paragraph_spacing_before_pt > 0:
                should_apply_spacing = False
                if is_in_table_cell:
                    # TABLE CELL: not first para in cell AND previous block was list/table
                    should_apply_spacing = (
                        not is_first_para_in_cell and 
                        not is_first_block and 
                        previous_block_type in ("bullet_list", "numbered_list", "table")
                    )
                else:
                    # BODY FLOW: not first block AND previous block was list/table
                    should_apply_spacing = (
                        not is_first_block and 
                        previous_block_type in ("bullet_list", "numbered_list", "table")
                    )
                
                if should_apply_spacing:
                    from docx.shared import Pt
                    new_p.paragraph_format.space_before = Pt(config.paragraph_spacing_before_pt)

        elif block.type == "bullet_list":
            # Manual bullet rendering - no Word numbering XML
            max_depth = config.max_visual_depth if config.max_visual_depth is not None else 3
            indent_per_level = config.indent_inches_per_level
            hanging_indent = -config.hanging_indent_inches
            
            # Guard against empty bullet_glyphs
            bullet_glyphs = config.bullet_glyphs if config.bullet_glyphs else ("•",)
            if not bullet_glyphs:
                bullet_glyphs = ("•",)

            raw_items = block.items or []
            clamped_logged = False  # Track if we've logged clamping for this block
            for idx, item in enumerate(raw_items):
                # Support legacy (str), (level, text), and new (level, text, formatted_runs) formats
                if isinstance(item, tuple):
                    if len(item) == 3:
                        logical_level, item_text, formatted_runs = item
                    elif len(item) == 2:
                        logical_level, item_text = item
                        formatted_runs = None
                    else:
                        logical_level, item_text = 1, item[0] if item else ""
                        formatted_runs = None
                else:
                    logical_level, item_text = 1, item
                    formatted_runs = None

                # Normalize logical level (actual nesting from markdown)
                if logical_level < 1:
                    logical_level = 1
                
                # Compute visual level (clamped by max_visual_depth) once per item
                if logical_level <= max_depth:
                    visual_level = logical_level
                    bullet_glyph = bullet_glyphs[min(logical_level - 1, len(bullet_glyphs) - 1)]
                else:
                    # Handle depth beyond max_visual_depth based on strategy
                    if config.deep_bullet_strategy == "clamp_last":
                        visual_level = max_depth
                        bullet_glyph = bullet_glyphs[min(max_depth - 1, len(bullet_glyphs) - 1)]
                        if not clamped_logged:
                            logger.debug("List nesting level %d clamped to %d", logical_level, max_depth)
                            clamped_logged = True
                    elif config.deep_bullet_strategy == "cycle":
                        # Cycle through available glyphs
                        visual_level = max_depth
                        glyph_index = ((logical_level - 1) % len(bullet_glyphs))
                        bullet_glyph = bullet_glyphs[glyph_index]
                    elif config.deep_bullet_strategy == "textual":
                        # Use textual indicator for deep nesting
                        visual_level = max_depth
                        bullet_glyph = f"[{logical_level}]"
                    else:
                        # Default to clamp_last
                        visual_level = max_depth
                        bullet_glyph = bullet_glyphs[min(max_depth - 1, len(bullet_glyphs) - 1)]
                        if not clamped_logged:
                            logger.debug("List nesting level %d clamped to %d", logical_level, max_depth)
                            clamped_logged = True

                # Create paragraph directly in container
                if is_in_table_cell and cell_elem is not None:
                    from docx.oxml import parse_xml
                    from docx.oxml.ns import nsdecls
                    p_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
                    insert_at = cell_insertion_index if cell_insertion_index is not None else len(cell_elem)
                    cell_elem.insert(insert_at, p_xml)
                    new_p = Paragraph(p_xml, cell)  # type: ignore[arg-type]
                    cell_insertion_index = insert_at + 1
                else:
                    # BODY FLOW: insert at captured position
                    from docx.oxml import parse_xml
                    from docx.oxml.ns import nsdecls
                    p_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
                    insertion_parent.insert(insertion_index, p_xml)
                    new_p = Paragraph(p_xml, doc._body)
                    insertion_index += 1

                # Set indentation using visual_level (consistent with glyph selection)
                new_p.paragraph_format.left_indent = Inches(indent_per_level * visual_level)
                new_p.paragraph_format.first_line_indent = Inches(hanging_indent)
                
                # Neutralize Word default paragraph spacing for list items (spacing comes from markdown, not Word defaults)
                from docx.shared import Pt
                new_p.paragraph_format.space_before = Pt(0)
                new_p.paragraph_format.space_after = Pt(0)

                # Add bullet prefix as formatting-isolated run, then item text
                bullet_run = new_p.add_run(f"{bullet_glyph}  ")
                bullet_run.bold = False
                bullet_run.italic = False
                
                if formatted_runs:
                    for run_data in formatted_runs:
                        run = new_p.add_run(run_data.text)
                        run.bold = run_data.bold
                        run.italic = run_data.italic
                else:
                    new_p.add_run(item_text)

        elif block.type == "numbered_list":
            # Manual numbered list rendering - no Word numbering XML
            max_depth = config.max_visual_depth if config.max_visual_depth is not None else 3
            indent_per_level = config.indent_inches_per_level
            hanging_indent = -config.hanging_indent_inches
            
            # Track numbering per list block (resets for each block)
            # Use dictionary-based counters for dynamic depth (supports any nesting level)
            number_counters: Dict[int, int] = {}
            last_visual_level = 0

            raw_items = block.items or []
            clamped_logged = False  # Track if we've logged clamping for this block
            for idx, item in enumerate(raw_items):
                # Support legacy (str), (level, text), and new (level, text, formatted_runs) formats
                if isinstance(item, tuple):
                    if len(item) == 3:
                        logical_level, item_text, formatted_runs = item
                    elif len(item) == 2:
                        logical_level, item_text = item
                        formatted_runs = None
                    else:
                        logical_level, item_text = 1, item[0] if item else ""
                        formatted_runs = None
                else:
                    logical_level, item_text = 1, item
                    formatted_runs = None

                # Normalize logical level (actual nesting from markdown)
                if logical_level < 1:
                    logical_level = 1
                
                # Compute visual level (clamped by max_visual_depth) for rendering
                visual_level = min(logical_level, max_depth)
                if logical_level > max_depth and not clamped_logged:
                    logger.debug("List nesting level %d clamped to %d", logical_level, max_depth)
                    clamped_logged = True

                # Update numbering counters based on visual level hierarchy
                # (numbering uses visual_level, but tracks full logical structure)
                if visual_level <= last_visual_level:
                    # Reset deeper visual levels when going back up
                    for l in range(visual_level + 1, max_depth + 1):
                        if l in number_counters:
                            number_counters[l] = 0

                # Initialize counter for visual level if needed
                if visual_level not in number_counters:
                    number_counters[visual_level] = 0

                # Increment counter for current visual level
                number_counters[visual_level] += 1
                last_visual_level = visual_level

                # Build hierarchical number string (1., 1.1., 1.1.1.) up to visual level
                number_parts = []
                for l in range(1, visual_level + 1):
                    if l not in number_counters:
                        number_counters[l] = 1
                    number_parts.append(str(number_counters[l]))
                number_str = ".".join(number_parts) + "."

                # Create paragraph directly in container
                if is_in_table_cell and cell_elem is not None:
                    from docx.oxml import parse_xml
                    from docx.oxml.ns import nsdecls
                    p_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
                    insert_at = cell_insertion_index if cell_insertion_index is not None else len(cell_elem)
                    cell_elem.insert(insert_at, p_xml)
                    new_p = Paragraph(p_xml, cell)  # type: ignore[arg-type]
                    cell_insertion_index = insert_at + 1
                else:
                    # BODY FLOW: insert at captured position
                    from docx.oxml import parse_xml
                    from docx.oxml.ns import nsdecls
                    p_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
                    insertion_parent.insert(insertion_index, p_xml)
                    new_p = Paragraph(p_xml, doc._body)
                    insertion_index += 1

                # Set indentation using visual_level (consistent with numbering)
                new_p.paragraph_format.left_indent = Inches(indent_per_level * visual_level)
                new_p.paragraph_format.first_line_indent = Inches(hanging_indent)
                
                # Neutralize Word default paragraph spacing for list items (spacing comes from markdown, not Word defaults)
                from docx.shared import Pt
                new_p.paragraph_format.space_before = Pt(0)
                new_p.paragraph_format.space_after = Pt(0)

                # Add number prefix as formatting-isolated run, then item text
                number_run = new_p.add_run(f"{number_str}  ")
                number_run.bold = False
                number_run.italic = False
                
                if formatted_runs:
                    for run_data in formatted_runs:
                        run = new_p.add_run(run_data.text)
                        run.bold = run_data.bold
                        run.italic = run_data.italic
                else:
                    new_p.add_run(item_text)

        elif block.type == "table":
            header = block.header or []
            rows = block.rows or []

            # Determine number of columns
            num_cols = len(header) if header else (len(rows[0]) if rows else 0)
            if num_cols <= 0:
                continue

            # Check if this table is a list continuation (table inside a list item)
            continuation_level = getattr(block, "_list_continuation_level", None)
            if isinstance(continuation_level, int) and continuation_level > 0:
                # Table inside list item: render as text-grid with indentation (same as continuation paragraphs)
                max_depth = config.max_visual_depth if config.max_visual_depth is not None else 3
                indent_per_level = config.indent_inches_per_level
                hanging_indent = -config.hanging_indent_inches
                visual_level = min(continuation_level, max_depth)
                
                # Render header row as bold paragraph with indentation
                if header:
                    if is_in_table_cell and cell_elem is not None:
                        from docx.oxml import parse_xml
                        from docx.oxml.ns import nsdecls
                        header_text = " | ".join(str(cell_text) if cell_text else "" for cell_text in header)
                        p_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
                        insert_at = cell_insertion_index if cell_insertion_index is not None else len(cell_elem)
                        cell_elem.insert(insert_at, p_xml)
                        header_p = Paragraph(p_xml, cell)  # type: ignore[arg-type]
                        cell_insertion_index = insert_at + 1
                    else:
                        from docx.oxml import parse_xml
                        from docx.oxml.ns import nsdecls
                        header_text = " | ".join(str(cell_text) if cell_text else "" for cell_text in header)
                        p_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
                        insertion_parent.insert(insertion_index, p_xml)
                        header_p = Paragraph(p_xml, doc._body)
                        insertion_index += 1
                    
                    header_p.paragraph_format.left_indent = Inches(indent_per_level * visual_level)
                    header_p.paragraph_format.first_line_indent = Inches(hanging_indent)
                    from docx.shared import Pt
                    header_p.paragraph_format.space_before = Pt(0)
                    header_p.paragraph_format.space_after = Pt(0)
                    header_run = header_p.add_run(header_text)
                    header_run.bold = True
                
                # Render separator row
                if header and rows:
                    if is_in_table_cell and cell_elem is not None:
                        from docx.oxml import parse_xml
                        from docx.oxml.ns import nsdecls
                        p_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
                        insert_at = cell_insertion_index if cell_insertion_index is not None else len(cell_elem)
                        cell_elem.insert(insert_at, p_xml)
                        separator_p = Paragraph(p_xml, cell)  # type: ignore[arg-type]
                        cell_insertion_index = insert_at + 1
                    else:
                        from docx.oxml import parse_xml
                        from docx.oxml.ns import nsdecls
                        p_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
                        insertion_parent.insert(insertion_index, p_xml)
                        separator_p = Paragraph(p_xml, doc._body)
                        insertion_index += 1
                    
                    separator_p.paragraph_format.left_indent = Inches(indent_per_level * visual_level)
                    separator_p.paragraph_format.first_line_indent = Inches(hanging_indent)
                    from docx.shared import Pt
                    separator_p.paragraph_format.space_before = Pt(0)
                    separator_p.paragraph_format.space_after = Pt(0)
                    separator_p.add_run("------------------------")
                
                # Render data rows as paragraphs with indentation
                for row in rows:
                    row_text = " | ".join(str(cell_text) if cell_text else "" for cell_text in row)
                    if is_in_table_cell and cell_elem is not None:
                        from docx.oxml import parse_xml
                        from docx.oxml.ns import nsdecls
                        p_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
                        insert_at = cell_insertion_index if cell_insertion_index is not None else len(cell_elem)
                        cell_elem.insert(insert_at, p_xml)
                        row_p = Paragraph(p_xml, cell)  # type: ignore[arg-type]
                        cell_insertion_index = insert_at + 1
                    else:
                        from docx.oxml import parse_xml
                        from docx.oxml.ns import nsdecls
                        p_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
                        insertion_parent.insert(insertion_index, p_xml)
                        row_p = Paragraph(p_xml, doc._body)
                        insertion_index += 1
                    
                    row_p.paragraph_format.left_indent = Inches(indent_per_level * visual_level)
                    row_p.paragraph_format.first_line_indent = Inches(hanging_indent)
                    from docx.shared import Pt
                    row_p.paragraph_format.space_before = Pt(0)
                    row_p.paragraph_format.space_after = Pt(0)
                    row_p.add_run(row_text)
                
                continue

            # If placeholder is in a table cell, first attempt a safe nested Word table.
            # If that is not possible, deterministically fall back to text-grid rendering.
            if is_in_table_cell:
                # SAFETY INVARIANT: Single-table-only in table cell uses text-grid fallback
                # WHY: A nested-table-only cell can yield empty `cell.text` in python-docx, making
                # content invisible. Text-grid fallback attempts to ensure content is visible and accessible.
                # This attempts to preserve content fidelity and prevent silent data loss.
                # Stability rule (existing behavior):
                # If the entire replacement consists of only a single markdown table block in a table cell,
                # fall back to deterministic text-grid paragraphs so the cell has visible text content.
                attempt_nested = not (len(blocks) == 1 and blocks[0].type == "table")

                success = False
                if attempt_nested:
                    logger.info("Attempting nested Word table render inside table cell")
                    # Use the placeholder paragraph's current index as the insertion anchor
                    start_index_for_nested = cell_insertion_index if cell_insertion_index is not None else (
                        list(cell_elem).index(paragraph._p) if cell_elem is not None else 0
                    )
                    if cell_elem is not None:
                        success = _attempt_render_nested_word_table(
                            block,
                            parent_cell_elem=cell_elem,
                            insertion_index=start_index_for_nested,
                            render_context=None,
                            summary=summary,
                        )
                    logger.info("Nested table render returned: %s", success)
                    
                    # Record fallback if nested table rendering failed (if enabled)
                    if not success and summary is not None:
                        summary.fallback_events.append(
                            FallbackEvent(
                                event_type="nested_table_fallback",
                                location=placeholder_token,
                                reason="Nested table rendering failed, using text-grid fallback",
                            )
                        )
                        logger.info(f"Fallback: Nested table rendering failed for {placeholder_token}, using text-grid")
                else:
                    # Single-table-only replacement in a table cell: deterministic text-grid fallback.
                    fallback_msg = "Single-table-only replacement inside a table cell uses deterministic text-grid fallback by design (nested table skipped)."
                    logger.info(fallback_msg)
                    
                    # Record fallback event in summary (if enabled)
                    if summary is not None:
                        summary.fallback_events.append(
                            FallbackEvent(
                                event_type="nested_table_fallback",
                                location=placeholder_token,
                                reason="Single-table-only in table cell uses text-grid fallback by design (attempts to ensure visible content)",
                            )
                        )
                        logger.info(f"Fallback: {fallback_msg}")
                if success and cell_elem is not None:
                    # Remove the placeholder paragraph AFTER successful insertion
                    parent_for_para = paragraph._p.getparent()
                    if parent_for_para is not None and not placeholder_removed:
                        parent_for_para.remove(paragraph._p)
                        placeholder_removed = True

                    # Recompute insertion index from DOM so subsequent blocks follow the table subtree
                    children = list(cell_elem)
                    new_index = None
                    for idx, child in enumerate(children):
                        if child.tag.endswith("}tbl"):
                            new_index = idx + 1
                    if new_index is None:
                        new_index = len(children)
                    cell_insertion_index = new_index
                    continue

                # Nested table failed: now remove the placeholder paragraph and fall back to text-grid
                if is_in_table_cell and not placeholder_removed:
                    parent_for_para = paragraph._p.getparent()
                    if parent_for_para is not None:
                        parent_for_para.remove(paragraph._p)
                    placeholder_removed = True
                
                # Render header row as bold paragraph
                if header and cell_elem is not None:
                    from docx.oxml import parse_xml
                    from docx.oxml.ns import nsdecls
                    header_text = " | ".join(str(cell_text) if cell_text else "" for cell_text in header)
                    p_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
                    insert_at = cell_insertion_index if cell_insertion_index is not None else len(cell_elem)
                    cell_elem.insert(insert_at, p_xml)
                    header_p = Paragraph(p_xml, cell)  # type: ignore[arg-type]
                    cell_insertion_index = insert_at + 1
                    header_run = header_p.add_run(header_text)
                    header_run.bold = True
                
                # Render separator row (if present in markdown, typically after header)
                # Note: Markdown tables don't have explicit separator rows in the block structure,
                # but we add a visual separator for clarity
                if header and rows and cell_elem is not None:
                    from docx.oxml import parse_xml
                    from docx.oxml.ns import nsdecls
                    p_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
                    insert_at = cell_insertion_index if cell_insertion_index is not None else len(cell_elem)
                    cell_elem.insert(insert_at, p_xml)
                    separator_p = Paragraph(p_xml, cell)  # type: ignore[arg-type]
                    cell_insertion_index = insert_at + 1
                    separator_p.add_run("------------------------")
                
                # Render data rows as paragraphs
                if cell_elem is not None:
                    from docx.oxml import parse_xml
                    from docx.oxml.ns import nsdecls
                    for row in rows:
                        row_text = " | ".join(str(cell_text) if cell_text else "" for cell_text in row)
                        p_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
                        insert_at = cell_insertion_index if cell_insertion_index is not None else len(cell_elem)
                        cell_elem.insert(insert_at, p_xml)
                        row_p = Paragraph(p_xml, cell)  # type: ignore[arg-type]
                        cell_insertion_index = insert_at + 1
                        row_p.add_run(row_text)
                
                continue

            # Create table at the exact placeholder position (BODY FLOW) or append (fallback)
            # NOTE:
            # BODY FLOW tables are constructed manually at the XML level,
            # while nested tables are built via python-docx in a scratch document.
            # This asymmetry is intentional and currently safe, but must be preserved
            # unless both paths are refactored together.
            if insertion_parent is not None and insertion_index is not None:
                # BODY FLOW: Insert table at captured position
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls, qn
                from docx.table import Table
                
                # Get namespace map from parent element to ensure compatibility
                parent_nsmap = insertion_parent.nsmap if hasattr(insertion_parent, 'nsmap') else {}
                w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                
                # Create table XML structure with proper namespaces
                # Use the parent's namespace context
                tbl_elem = insertion_parent.makeelement(qn('w:tbl'))
                
                # Add table properties
                tbl_pr = insertion_parent.makeelement(qn('w:tblPr'))
                tbl_elem.append(tbl_pr)
                
                # Add table grid
                tbl_grid = insertion_parent.makeelement(qn('w:tblGrid'))
                for _ in range(num_cols):
                    grid_col = insertion_parent.makeelement(qn('w:gridCol'))
                    tbl_grid.append(grid_col)
                tbl_elem.append(tbl_grid)
                
                # Add header row if present
                if header:
                    tr = insertion_parent.makeelement(qn('w:tr'))
                    for cell_text in header[:num_cols]:
                        tc = insertion_parent.makeelement(qn('w:tc'))
                        p = insertion_parent.makeelement(qn('w:p'))
                        r = insertion_parent.makeelement(qn('w:r'))
                        t = insertion_parent.makeelement(qn('w:t'))
                        t.text = str(cell_text) if cell_text else ""
                        r.append(t)
                        p.append(r)
                        tc.append(p)
                        tr.append(tc)
                    tbl_elem.append(tr)
                
                # Add data rows
                for row in rows:
                    tr = insertion_parent.makeelement(qn('w:tr'))
                    for cell_text in row[:num_cols]:
                        tc = insertion_parent.makeelement(qn('w:tc'))
                        p = insertion_parent.makeelement(qn('w:p'))
                        r = insertion_parent.makeelement(qn('w:r'))
                        t = insertion_parent.makeelement(qn('w:t'))
                        t.text = str(cell_text) if cell_text else ""
                        r.append(t)
                        p.append(r)
                        tc.append(p)
                        tr.append(tc)
                    tbl_elem.append(tr)
                
                # Insert at captured position
                insertion_parent.insert(insertion_index, tbl_elem)
                insertion_index += 1
                
                # Wrap as Table object for style application
                table = Table(tbl_elem, doc._body)
                
                # Apply default table formatting (style first, then borders fallback)
                _apply_default_table_formatting(tbl_elem, doc._body)
            else:
                # Fallback: Create table at end of document (should not happen in normal flow)
                table = doc.add_table(rows=0, cols=num_cols)
                
                # Apply default table formatting (style first, then borders fallback)
                _apply_default_table_formatting(table._tbl, doc._body)

                if header:
                    hdr_cells = table.add_row().cells
                    for idx, cell_text in enumerate(header):
                        if idx < num_cols:
                            hdr_cells[idx].text = cell_text or ""

                for row in rows:
                    row_cells = table.add_row().cells
                    for idx, cell_text in enumerate(row):
                        if idx < num_cols:
                            row_cells[idx].text = cell_text or ""
        
        # BUG #1 FIX: Update tracking at end of loop iteration (after all rendering is complete)
        previous_block_type = current_block_type
        is_first_block = False


def _build_scalar_mapping(req: WordExportRequest) -> Dict[str, Optional[str]]:
    """
    Build scalar placeholder mapping from WordExportRequest.
    
    This helper function extracts scalar_fields from the request and returns
    them as a dictionary. It is used internally to prepare the mapping for
    scalar placeholder replacement.
    
    Args:
        req: WordExportRequest containing scalar_fields dictionary.
    
    Returns:
        Dictionary mapping placeholder keys to replacement values.
        Keys are used as-is (caller adds braces to form tokens like {{key}}).
    """
    """
    Build mapping from scalar placeholder keys to values from the request.
    Keys are without braces (e.g. 'document_id'), values are raw strings.
    """
    return req.scalar_fields


def _build_plaintext_blocks(content: str) -> List[MarkdownBlock]:
    """
    Convert plain text into paragraph-only blocks.
    Does not infer bullets or tables.
    """
    if not content or not content.strip():
        return []
    parts = [p.strip() for p in content.split("\n\n") if p.strip()]
    if not parts:
        parts = [content.strip()]
    return [MarkdownBlock(type="paragraph", text=p) for p in parts]


def export_to_word(
    template_path: Path,
    request: WordExportRequest,
    markdown_mode: bool,
    output_path: Path,
    config: ListRenderConfig = ListRenderConfig()
) -> Dict[str, Any]:
    """
    Main entry point for exporting structured content to a Word template.
    
    This function populates a Word template with scalar and block content from
    the provided request and saves the result to the specified output path.
    
    This function is strictly non-generative: it does not call any LLMs and
    does not alter wording. It attempts to map existing content into the DOCX structure,
    preserving text when structure cannot be rendered (text-preserving fallback).
    
    Processing Flow:
        1. Load Word template from template_path
        2. Attempt to replace scalar placeholders where found (body, headers, footers, textboxes)
        3. Attempt to replace block placeholders in BODY containers only (with markdown parsing if enabled)
        4. Save combined markdown file (if block content exists)
        5. Save export summary log file (if enable_export_trace=True)
        6. Save final Word document to output_path
    
    Safety Guarantees:
        - Deterministic output: attempts to produce identical outputs given identical inputs
        - Text-preserving fallback: content is preserved as text when structure cannot be rendered
        - Template safety: block placeholders must occupy entire paragraphs (mixed = skipped)
        - Container restrictions: block expansion limited to BODY containers only
        - Fallback mechanisms: parsing failures fall back to plain text
    
    Observability:
        When `enable_export_trace=True` (default), the export summary (audit trail)
        is automatically saved to a timestamped subdirectory under `logs/`:
        - If `EXPORT_LOG_DIR` environment variable is set: `{EXPORT_LOG_DIR}/logs/{timestamp}/`
        - Otherwise: `{output_dir}/logs/{timestamp}/`
        - Each export run gets its own timestamped directory (e.g., `20240124_143022/`)
        - Directory structure is created automatically if it doesn't exist
        - Example path: `logs/20240124_143022/export_summary_output.json`
    
    Args:
        template_path: Path to the Word template file (.docx). Must exist and be readable.
            The template should contain placeholders like {{title}}, {{summary}}, etc.
        request: WordExportRequest containing scalar_fields, block_fields, and
            enable_export_trace flag. Defines what content to insert into the template.
        markdown_mode: If True, block_fields are parsed as markdown and converted to
            Word structures (headings, lists, tables). If False, block_fields are
            treated as plain text (paragraphs only, no structure inference).
        output_path: Path where the generated Word document will be saved.
            If relative, resolved relative to OUTPUT_DIRECTORY environment variable (if set) or current working directory.
            If absolute, used as-is (OUTPUT_DIRECTORY is ignored).
            Parent directory is created automatically if it doesn't exist.
        config: ListRenderConfig for controlling list rendering behavior (indentation,
            glyphs, deep nesting strategy). Defaults to ListRenderConfig() with standard
            settings if not provided.
    
    Returns:
        Dictionary containing:
        - "word_file_path": String path to the generated DOCX file (same as output_path)
        - "markdown_files": List containing a single path to the combined markdown file
            (`exported_markdown_content.md`), or empty list if all blocks are empty.
            The markdown file contains all non-empty block contents separated by `---` lines.
        - "export_summary": (Optional) Dictionary representation of ExportSummary when
            enable_export_trace=True. Contains scalar_replacements, block_replacements,
            skipped_items, fallback_events, and warnings.
        - "export_summary_log_path": (Optional) String path to the saved JSON log file
            when enable_export_trace=True. The file contains the export summary for
            audit trail and debugging purposes.
    
    Raises:
        FileNotFoundError: If template_path does not exist.
        ValueError: If config is None (defensive guard in block replacement).
        Other exceptions may be raised by underlying libraries (python-docx, markdown-it)
            but are generally caught and handled gracefully with fallbacks.
    
    Example:
        ```python
        from pathlib import Path
        from docx_template_export.models.export_models import WordExportRequest
        from docx_template_export.services.word_export_service import export_to_word
        
        request = WordExportRequest(
            scalar_fields={"title": "My Document", "author": "John Doe"},
            block_fields={"summary": "# Introduction\\n\\nContent here..."},
            enable_export_trace=True
        )
        
        result = export_to_word(
            template_path=Path("template.docx"),
            request=request,
            markdown_mode=True,
            output_path=Path("output.docx")
        )
        
        print(f"Generated: {result['word_file_path']}")
        print(f"Summary log: {result.get('export_summary_log_path')}")
        ```
    """
    if not template_path.exists():
        raise FileNotFoundError(f"Word template not found at: {template_path}")

    # Resolve output_path: if relative, resolve relative to OUTPUT_DIRECTORY (if set) or cwd
    # If absolute, use as-is (OUTPUT_DIRECTORY is ignored for absolute paths)
    if not output_path.is_absolute():
        output_dir_env = os.getenv("OUTPUT_DIRECTORY")
        if output_dir_env:
            base_output_dir = Path(output_dir_env)
            if not base_output_dir.is_absolute():
                base_output_dir = Path.cwd() / base_output_dir
            output_path = base_output_dir.resolve() / output_path
        else:
            # Default behavior: resolve relative to current working directory
            output_path = Path.cwd() / output_path
        output_path = output_path.resolve()

    # Ensure output directory exists
    output_dir = output_path.parent
    ensure_output_dir_exists(output_dir)

    doc = Document(str(template_path))

    # Initialize export summary (optional, gated by enable_export_trace)
    # Summary tracks what happened during export without influencing rendering
    summary: Optional[ExportSummary] = None
    log_file_handler: Optional[logging.FileHandler] = None
    log_dir: Optional[Path] = None  # Store log directory for reuse when saving JSON
    if request.enable_export_trace:
        summary = ExportSummary()
        
        # Set up automatic file logging for structured logs
        try:
            # Determine base log directory: use environment variable if set, otherwise use output_dir/logs
            log_dir_env = os.getenv("EXPORT_LOG_DIR")
            if log_dir_env:
                base_log_dir = Path(log_dir_env)
                base_logs_dir = base_log_dir / "logs"
            else:
                # Default: create logs directory alongside output directory
                base_logs_dir = output_dir / "logs"
            
            # Generate timestamp for this export run
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            # Create timestamped subdirectory for this run
            log_dir = base_logs_dir / timestamp
            log_dir.mkdir(parents=True, exist_ok=True)
            
            # Generate log filename (without timestamp since it's in the directory name)
            output_filename = output_path.stem  # Get filename without extension
            log_filename = f"export_trace_{output_filename}.log"
            log_file_path = log_dir / log_filename
            
            # Create file handler for structured logs
            log_file_handler = logging.FileHandler(str(log_file_path), encoding="utf-8")
            log_file_handler.setLevel(logging.INFO)
            # Use a simple format for structured logs
            formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
            log_file_handler.setFormatter(formatter)
            
            # Add handler to logger (only for this export run)
            logger.addHandler(log_file_handler)
            logger.info("Export trace enabled: collecting summary and structured logs")
        except Exception as e:
            # Don't fail the export if log file setup fails
            logger.warning("Failed to set up structured log file: %s", str(e))
            log_file_handler = None
            log_dir = None

    # 1) Replace scalar placeholders
    scalar_mapping = _build_scalar_mapping(request)
    replace_scalar_placeholders(doc, scalar_mapping, summary=summary)

    # 2) Replace block placeholders with structured content
    # Enforce placeholder provenance: only replace what's explicitly in request.block_fields
    # Build authoritative placeholder registry
    block_placeholders: Dict[str, Optional[str]] = {
        f"{{{{{key}}}}}": value for key, value in request.block_fields.items()
    }
    allowed_block_tokens = set(block_placeholders.keys())

    # Accumulate markdown content for all non-empty blocks
    combined_markdown_parts: List[str] = []

    # ============================================================================
    # ⚠️ GLOBAL FREEZE: Block Placeholder Target Collection
    # ============================================================================
    # WHY TARGETS ARE FROZEN BEFORE MUTATION:
    #   - Document traversal must be single-pass and read-only to avoid non-deterministic
    #     behavior from DOM mutations during iteration
    #   - If we mutated the document while traversing, we could miss placeholders or
    #     encounter inconsistent DOM state
    #
    # WHY SINGLE-PASS TRAVERSAL IS REQUIRED:
    #   - python-docx iterators are not safe for concurrent modification
    #   - Block replacement removes/inserts paragraphs, which would invalidate iterators
    #   - Multiple passes would be non-deterministic if document structure changes
    #
    # DOM STABILITY ASSUMPTIONS:
    #   - This logic relies on the assumption that placeholder paragraphs exist and are
    #     accessible during the read-only traversal phase
    #   - After freezing targets, mutations happen in a controlled sequence per token
    #   - Each mutation is atomic and does not affect other frozen targets
    #
    # ⚠️ DO NOT MODIFY WITHOUT:
    #   - Regression tests covering multiple block placeholders in the same document
    #   - Tests with mixed BODY FLOW and BODY TABLE scenarios
    #   - Verification that all placeholders are found and replaced deterministically
    #   - Ensuring no placeholders are missed or processed multiple times
    #
    # Any change to this traversal or mutation logic requires comprehensive regression
    # tests to ensure determinism and correctness.
    # ============================================================================
    frozen_block_targets: Dict[str, List[Tuple[Paragraph, Region, Container]]] = {}
    for token in allowed_block_tokens:
        frozen_block_targets[token] = []
    
    # Single read-only traversal to collect all targets
    for location_data, region, container in _iter_all_locations(doc):
        # Block expansion is only allowed in BODY FLOW and BODY TABLE
        if not _is_block_expansion_allowed(region, container):
            # Check if this location contains any block token (for warning only)
            if container == Container.TEXTBOX:
                # For textboxes, check XML runs
                run_elems = location_data
                location_text = "".join(_get_text_from_xml_run(run) for run in run_elems)
            else:
                # For paragraphs, check paragraph text
                paragraph = location_data
                if not paragraph.runs:
                    continue
                # Use XML text extraction for consistency with replacement logic (attempts to preserve <w:br/> and <w:cr/> as \n)
                location_text = "".join(_get_text_from_xml_run(run._r) for run in paragraph.runs)
            
            # Check all allowed tokens for warnings
            # SAFETY INVARIANT: Block placeholders in textboxes, headers, or footers are skipped
            # because textboxes only support scalar replacement (Word XML structure limitation),
            # and headers/footers are restricted to prevent layout instability.
            # This attempts to ensure deterministic output and prevent document corruption.
            for token in allowed_block_tokens:
                if token in location_text:
                    skip_reason = ""
                    if container == Container.TEXTBOX:
                        warning_msg = (
                            f"Block placeholder {token} found inside TEXTBOX ({region.value}). "
                            "Textboxes are scalar-only. Block expansion is not supported there. Leaving unchanged."
                        )
                        skip_reason = "block_in_textbox"
                        # Warning logs remain unchanged (important for debugging)
                        logger.warning(warning_msg)
                    else:
                        warning_msg = (
                            f"Block placeholder {token} found in {region.value}/{container.value} "
                            "where block expansion is not allowed (only BODY FLOW and BODY TABLE). Leaving unchanged."
                        )
                        skip_reason = f"block_in_{region.value.lower()}"
                        # Warning logs remain unchanged (important for debugging)
                        logger.warning(warning_msg)
                    
                    # Record skipped item in summary (if enabled)
                    # Structured log: gated by enable_export_trace
                    if summary is not None:
                        summary.skipped_items.append(
                            SkippedItem(
                                token=token,
                                region=region.value,
                                container=container.value,
                                reason=skip_reason,
                            )
                        )
                        summary.warnings.append(
                            WarningEvent(
                                message=warning_msg,
                                context={"token": token, "region": region.value, "container": container.value},
                            )
                        )
                        logger.info(f"Skipped block placeholder {token} in {region.value}/{container.value}: {skip_reason}")
            continue
        
        # Only process BODY FLOW and BODY TABLE from here
        paragraph = location_data  # Must be a Paragraph for allowed locations
        if not paragraph.runs:
            continue
        # Use XML text extraction for consistency with replacement logic (attempts to preserve <w:br/> and <w:cr/> as \n)
        para_text = "".join(_get_text_from_xml_run(run._r) for run in paragraph.runs)
        
        # Check all allowed tokens and collect matches
        for token in allowed_block_tokens:
            if token in para_text:
                frozen_block_targets[token].append((paragraph, region, container))

    # Now process each token using frozen targets (document mutations happen here)
    for token, content in block_placeholders.items():
        if content and content.strip():
            combined_markdown_parts.append(content)

        # Get frozen targets for this token
        frozen_targets = frozen_block_targets.get(token, [])
        
        # Invariant: log if placeholder appears multiple times
        if len(frozen_targets) > 1:
            logger.debug(f"Block placeholder {token} found in {len(frozen_targets)} locations")
        
        # Process frozen targets (document mutations happen here)
        for paragraph, region, container in frozen_targets:
            # Verify paragraph still exists and contains token (defensive check)
            try:
                if not paragraph.runs:
                    continue
                # Use XML text extraction for consistency with replacement logic (attempts to preserve <w:br/> and <w:cr/> as \n)
                para_text = "".join(_get_text_from_xml_run(run._r) for run in paragraph.runs)
                if token not in para_text:
                    logger.debug(f"Block placeholder {token} no longer found in paragraph (may have been removed)")
                    continue
            except (AttributeError, RuntimeError):
                # Paragraph may have been removed by previous replacement
                logger.debug(f"Paragraph no longer accessible (may have been removed)")
                continue

            if not content or not content.strip():
                # Empty content: just remove the placeholder token and leave section empty
                logger.debug(f"Block placeholder {token} has empty content, removing placeholder only")
                replace_block_placeholder_with_content(doc, paragraph, token, blocks=None, config=config)
                
                # Record empty block replacement in summary (if enabled)
                if summary is not None:
                    summary.block_replacements.append(
                        BlockReplacement(
                            token=token,
                            region=region.value,
                            container=container.value,
                            block_types=["empty"],
                        )
                    )
                    logger.info(f"Replaced block placeholder {token} in {region.value}/{container.value} with empty content")
                continue

            # SAFETY INVARIANT: Block placeholders must occupy entire paragraphs
            # WHY: Partial-paragraph replacement would create ambiguous layout and non-deterministic output.
            # The placeholder paragraph is removed and replaced with new content. If the placeholder
            # shares the paragraph with other text, we cannot deterministically decide what to preserve.
            # This constraint ensures document stability and prevents Word corruption.
            if para_text.strip() != token:
                warning_msg = (
                    f"Block placeholder {token} is not the only content in its paragraph. "
                    f"Paragraph contains: {para_text[:100]}. "
                    "Block placeholders must occupy the entire paragraph. Skipping replacement."
                )
                # Warning logs remain unchanged (important for debugging)
                logger.warning(warning_msg)
                
                # Record skipped item in summary (if enabled)
                # Structured log: gated by enable_export_trace
                if summary is not None:
                    summary.skipped_items.append(
                        SkippedItem(
                            token=token,
                            region=region.value,
                            container=container.value,
                            reason="not_entire_paragraph",
                        )
                    )
                    summary.warnings.append(
                        WarningEvent(
                            message=warning_msg,
                            context={"token": token, "region": region.value, "container": container.value, "para_text_preview": para_text[:100]},
                        )
                    )
                    logger.info(f"Skipped block placeholder {token} in {region.value}/{container.value}: not_entire_paragraph")
                continue

            # Phase 2: Tree-based renderer (enabled by default, can be disabled via environment variable)
            # Check environment variable (default: True - tree renderer is enabled by default)
            USE_MARKDOWN_TREE_RENDERER = os.getenv("USE_MARKDOWN_TREE_RENDERER", "true").lower() not in ("false", "0", "no")
            
            if USE_MARKDOWN_TREE_RENDERER and markdown_mode:
                # Use tree-based rendering
                from docx_template_export.renderers.tree_renderer import (
                    render_markdown_tree_to_docx,
                )
                tree = parse_markdown_to_tree(content, enabled=True, summary=summary)
                if tree:
                    # Render tree (handles placeholder removal internally)
                    # Tree renderer observability: detect nested tables and mixed lists
                    _observe_tree_structure(tree, summary, token)
                    render_markdown_tree_to_docx(tree, doc, placeholder_paragraph=paragraph, config=config)
                    
                    # Record successful tree-based block replacement in summary (if enabled)
                    if summary is not None:
                        # Extract block types from tree for summary
                        block_types = _extract_block_types_from_tree(tree)
                        summary.block_replacements.append(
                            BlockReplacement(
                                token=token,
                                region=region.value,
                                container=container.value,
                                block_types=block_types,
                            )
                        )
                        logger.info(
                            f"Replaced block placeholder {token} in {region.value}/{container.value} "
                            f"using tree renderer (block types: {', '.join(block_types)})"
                        )
                    continue
                else:
                    # Tree parsing failed, fall back to block-based rendering
                    # SAFETY INVARIANT: Tree renderer is feature-flagged and experimental.
                    # WHY: The tree renderer provides advanced features (nested tables, mixed lists)
                    # but is still being validated. Fallback to block renderer attempts to ensure deterministic
                    # output and prevent document corruption if tree parsing encounters edge cases.
                    fallback_msg = f"Tree parsing failed for {token}, falling back to block-based rendering"
                    logger.warning(fallback_msg)
                    
                    # Record fallback event in summary (if enabled)
                    if summary is not None:
                        summary.fallback_events.append(
                            FallbackEvent(
                                event_type="tree_parse_failed",
                                location=token,
                                reason="Tree parsing failed, using block-based renderer fallback",
                            )
                        )
                        logger.info(f"Fallback: {fallback_msg}")
            
            # Default: Use existing block-based rendering
            if markdown_mode:
                blocks = parse_markdown_to_blocks(content)
            else:
                blocks = _build_plaintext_blocks(content)

            replace_block_placeholder_with_content(doc, paragraph, token, blocks=blocks, config=config, summary=summary)
            
            # Record successful block replacement in summary (if enabled)
            # Note: Empty blocks are already recorded inside replace_block_placeholder_with_content
            # Only record if we actually rendered blocks (not empty, and not tree renderer which already recorded)
            if summary is not None and blocks:
                block_types = [block.type for block in blocks]
                summary.block_replacements.append(
                    BlockReplacement(
                        token=token,
                        region=region.value,
                        container=container.value,
                        block_types=block_types,
                    )
                )
                logger.info(
                    f"Replaced block placeholder {token} in {region.value}/{container.value} "
                    f"using block renderer (block types: {', '.join(block_types)})"
                )

    # 3) Save combined markdown (single file) alongside the DOCX
    saved_markdown_files: List[str] = []
    if combined_markdown_parts:
        combined_markdown = "\n\n---\n\n".join(combined_markdown_parts)
        markdown_file = output_dir / "exported_markdown_content.md"
        try:
            with open(markdown_file, "w", encoding="utf-8") as f:
                f.write(combined_markdown)
            saved_markdown_files.append(str(markdown_file))
            logger.info("Saved combined markdown content to %s", markdown_file)
        except Exception as e:
            logger.warning("Failed to save combined markdown file %s: %s", markdown_file, str(e))

    # 4) Save populated document
    doc.save(str(output_path))
    logger.info("Exported Word summary to %s", output_path)
    if saved_markdown_files:
        logger.info("Saved %d markdown content file(s): %s", len(saved_markdown_files), ", ".join(saved_markdown_files))
    
    # 5) Emit export summary log and save to log directory (if enabled)
    result: Dict[str, Any] = {
        "word_file_path": str(output_path),
        "markdown_files": saved_markdown_files,
    }
    
    if summary is not None:
        # Add summary to result
        result["export_summary"] = summary.to_dict()
        
        # Emit structured summary log
        logger.info(
            "Export summary: %d scalar replacements, %d block replacements, %d skipped items, "
            "%d fallback events, %d warnings",
            len(summary.scalar_replacements),
            len(summary.block_replacements),
            len(summary.skipped_items),
            len(summary.fallback_events),
            len(summary.warnings),
        )
        
        # Save export summary to log directory (automatic when enable_export_trace=True)
        try:
            # Reuse log_dir if it was already created for structured log file
            # Otherwise, create it now (shouldn't happen, but defensive)
            if log_dir is None:
                # Determine base log directory: use environment variable if set, otherwise use output_dir/logs
                log_dir_env = os.getenv("EXPORT_LOG_DIR")
                if log_dir_env:
                    base_log_dir = Path(log_dir_env)
                    base_logs_dir = base_log_dir / "logs"
                    logger.debug("Using log directory from EXPORT_LOG_DIR environment variable: %s", base_logs_dir)
                else:
                    # Default: create logs directory alongside output directory
                    base_logs_dir = output_dir / "logs"
                    logger.debug("Using default log directory alongside output: %s", base_logs_dir)
                
                # Generate timestamp for this export run
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                
                # Create timestamped subdirectory for this run
                log_dir = base_logs_dir / timestamp
                log_dir.mkdir(parents=True, exist_ok=True)
            
            # Generate filename (without timestamp since it's in the directory name)
            output_filename = output_path.stem  # Get filename without extension
            log_filename = f"export_summary_{output_filename}.json"
            log_file_path = log_dir / log_filename
            
            # Save summary as JSON (automatically saved when enable_export_trace=True)
            with open(log_file_path, "w", encoding="utf-8") as f:
                json.dump(summary.to_dict(), f, indent=2, ensure_ascii=False)
            
            logger.info("Export summary saved to log directory: %s", log_file_path)
            result["export_summary_log_path"] = str(log_file_path)
        except Exception as e:
            # Don't fail the export if log saving fails
            logger.warning("Failed to save export summary to log directory: %s", str(e))
    
    # Clean up file handler if it was added (to avoid memory leaks and handler accumulation)
    if log_file_handler is not None:
        try:
            logger.removeHandler(log_file_handler)
            log_file_handler.close()
        except Exception:
            # Ignore errors during cleanup
            pass
    
    return result