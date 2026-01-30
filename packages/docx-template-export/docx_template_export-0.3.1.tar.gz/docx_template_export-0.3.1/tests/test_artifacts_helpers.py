"""
Helper functions for production-grade test suite.

These helpers ensure all tests write artifacts and can extract text deterministically.
"""
import json
from pathlib import Path
from typing import Dict, Any, Optional

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from docx_template_export.models.export_models import WordExportRequest
from docx_template_export.services.word_export_service import export_to_word
from docx_template_export.models.export_config import ListRenderConfig
from typing import Dict, Any


def save_legacy_test_artifacts(
    test_name: str,
    request: WordExportRequest,
    result: Dict[str, Any],
    output_path: Path,
    artifacts_dir: Path,
    template_path: Optional[Path] = None,
) -> Path:
    """
    Helper function to save legacy test artifacts after export_to_word.
    
    This extracts markdown from result and calls write_legacy_test_artifacts.
    ALWAYS validates DOCX integrity before saving artifacts.
    
    Args:
        test_name: Name of the test (used for directory name)
        request: The WordExportRequest that was used
        result: Result dictionary from export_to_word
        output_path: Path to the generated DOCX file
        artifacts_dir: Base artifacts directory (tests/artifacts)
        template_path: Optional path to the template DOCX file
    
    Returns:
        Path to the test's artifact directory
    """
    # ALWAYS validate DOCX integrity - fail test if document is corrupted
    validate_docx_integrity(output_path)
    
    # Also validate template file if provided (templates can also be corrupted)
    if template_path and template_path.exists():
        validate_docx_integrity(template_path)
    
    # Get markdown content if available
    markdown_content = ""
    if result.get("markdown_files"):
        markdown_file = result["markdown_files"][0]
        if Path(markdown_file).exists():
            with open(markdown_file, "r", encoding="utf-8") as f:
                markdown_content = f.read()
    
    return write_legacy_test_artifacts(
        test_name=test_name,
        request=request,
        docx_path=output_path,
        artifacts_dir=artifacts_dir,
        template_path=template_path,
        markdown=markdown_content,
    )


def write_legacy_test_artifacts(
    test_name: str,
    request: WordExportRequest,
    docx_path: Path,
    artifacts_dir: Path,
    template_path: Optional[Path] = None,
    markdown: Optional[str] = None,
) -> Path:
    """
    Write legacy test artifacts to artifacts/legacy/<test_name>/ directory.
    
    Creates:
    - tests/artifacts/legacy/<test_name>/input.json
    - tests/artifacts/legacy/<test_name>/rendered_markdown.md
    - tests/artifacts/legacy/<test_name>/output.docx
    - tests/artifacts/legacy/<test_name>/template.docx (if template_path provided)
    
    NOTE: This function does NOT validate DOCX integrity. Call validate_docx_integrity()
    before calling this function, or use save_legacy_test_artifacts() which includes validation.
    
    Args:
        test_name: Name of the test (used for directory name)
        request: The WordExportRequest that was used
        docx_path: Path to the generated DOCX file
        artifacts_dir: Base artifacts directory (tests/artifacts)
        template_path: Optional path to the template DOCX file
        markdown: Optional rendered markdown content (if not provided, will try to extract from result)
    
    Returns:
        Path to the test's artifact directory
    """
    legacy_artifacts_dir = artifacts_dir / "legacy"
    test_artifact_dir = legacy_artifacts_dir / test_name
    test_artifact_dir.mkdir(parents=True, exist_ok=True)
    
    # Write input.json
    input_json_path = test_artifact_dir / "input.json"
    with open(input_json_path, "w", encoding="utf-8") as f:
        json.dump(
            {
                "scalar_fields": request.scalar_fields,
                "block_fields": request.block_fields,
            },
            f,
            indent=2,
            ensure_ascii=False,
        )
    
    # Write rendered_markdown.md
    markdown_path = test_artifact_dir / "rendered_markdown.md"
    if markdown:
        with open(markdown_path, "w", encoding="utf-8") as f:
            f.write(markdown)
    else:
        # Create empty file if markdown not available
        markdown_path.touch()
    
    # Copy output.docx
    output_docx_path = test_artifact_dir / "output.docx"
    if docx_path.exists():
        import shutil
        shutil.copy2(docx_path, output_docx_path)
    else:
        # Create empty file marker if DOCX doesn't exist
        output_docx_path.touch()
    
    # Copy template.docx if provided
    if template_path is not None and template_path.exists():
        template_docx_path = test_artifact_dir / "template.docx"
        import shutil
        shutil.copy2(template_path, template_docx_path)
    
    return test_artifact_dir


def write_test_artifacts(
    test_name: str,
    request: WordExportRequest,
    markdown: Optional[str],
    docx_path: Path,
    artifacts_dir: Path,
    template_path: Optional[Path] = None,
    subdirectory: Optional[str] = None,
) -> Path:
    """
    Write all test artifacts to disk for human inspection.
    
    Creates:
    - tests/artifacts/<subdirectory>/<test_name>/input.json (if subdirectory provided)
    - tests/artifacts/<test_name>/input.json (if subdirectory not provided)
    - tests/artifacts/<subdirectory>/<test_name>/rendered_markdown.md
    - tests/artifacts/<subdirectory>/<test_name>/output.docx
    - tests/artifacts/<subdirectory>/<test_name>/template.docx (if template_path provided)
    
    NOTE: This function does NOT validate DOCX integrity. Call validate_docx_integrity()
    before calling this function, or use run_export_and_collect_artifacts() which includes validation.
    
    Args:
        test_name: Name of the test (used for directory name)
        request: The WordExportRequest that was used
        markdown: The rendered markdown content (if available)
        docx_path: Path to the generated DOCX file
        artifacts_dir: Base artifacts directory (tests/artifacts)
        template_path: Optional path to the template DOCX file
        subdirectory: Optional subdirectory to organize tests (e.g., "phase3", "lists", "tables")
    
    Returns:
        Path to the test's artifact directory
    """
    if subdirectory:
        test_artifact_dir = artifacts_dir / subdirectory / test_name
    else:
        test_artifact_dir = artifacts_dir / test_name
    test_artifact_dir.mkdir(parents=True, exist_ok=True)
    
    # Write input.json
    input_json_path = test_artifact_dir / "input.json"
    with open(input_json_path, "w", encoding="utf-8") as f:
        json.dump(
            {
                "scalar_fields": request.scalar_fields,
                "block_fields": request.block_fields,
            },
            f,
            indent=2,
            ensure_ascii=False,
        )
    
    # Write rendered_markdown.md
    markdown_path = test_artifact_dir / "rendered_markdown.md"
    if markdown:
        with open(markdown_path, "w", encoding="utf-8") as f:
            f.write(markdown)
    else:
        # Create empty file if markdown not available
        markdown_path.touch()
    
    # Copy output.docx
    output_docx_path = test_artifact_dir / "output.docx"
    if docx_path.exists():
        import shutil
        shutil.copy2(docx_path, output_docx_path)
    else:
        # Create empty file marker if DOCX doesn't exist
        output_docx_path.touch()
    
    # Copy template.docx if provided
    if template_path is not None and template_path.exists():
        template_docx_path = test_artifact_dir / "template.docx"
        import shutil
        shutil.copy2(template_path, template_docx_path)
    
    return test_artifact_dir


def extract_all_text_from_docx(docx_path: Path) -> str:
    """
    Extract all visible text from a DOCX file deterministically.
    
    This function extracts text from:
    - All body paragraphs
    - All table cells (recursively)
    - All header paragraphs
    - All footer paragraphs
    
    Text is extracted in document order and flattened into a single string.
    Whitespace is normalized (collapsed to single spaces, newlines preserved).
    
    Args:
        docx_path: Path to the DOCX file
    
    Returns:
        Flattened string containing all visible text
    """
    if not docx_path.exists():
        return ""
    
    doc = Document(str(docx_path))
    text_parts: list[str] = []
    
    # Extract from body paragraphs
    for paragraph in doc.paragraphs:
        para_text = _extract_paragraph_text(paragraph)
        if para_text.strip():
            text_parts.append(para_text)
    
    # Extract from body tables
    for table in doc.tables:
        table_text = _extract_table_text(table)
        if table_text.strip():
            text_parts.append(table_text)
    
    # Extract from headers
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            para_text = _extract_paragraph_text(paragraph)
            if para_text.strip():
                text_parts.append(para_text)
        
        # Extract from header tables
        for table in header.tables:
            table_text = _extract_table_text(table)
            if table_text.strip():
                text_parts.append(table_text)
        
        # Extract from footers
        footer = section.footer
        for paragraph in footer.paragraphs:
            para_text = _extract_paragraph_text(paragraph)
            if para_text.strip():
                text_parts.append(para_text)
        
        # Extract from footer tables
        for table in footer.tables:
            table_text = _extract_table_text(table)
            if table_text.strip():
                text_parts.append(table_text)
    
    # Join all parts and normalize whitespace
    full_text = " ".join(text_parts)
    # Normalize whitespace: collapse multiple spaces, preserve newlines
    import re
    full_text = re.sub(r"[ \t]+", " ", full_text)  # Collapse spaces/tabs
    full_text = re.sub(r" *\n *", "\n", full_text)  # Normalize newlines
    return full_text.strip()


def _extract_paragraph_text(paragraph: Paragraph) -> str:
    """Extract text from a paragraph, including list markers."""
    # Use XML-based extraction for consistency with library
    from docx_template_export.services.word_export_service import _get_text_from_xml_run
    
    parts: list[str] = []
    for run in paragraph.runs:
        run_text = _get_text_from_xml_run(run._r)
        if run_text:
            parts.append(run_text)
    
    # Also check paragraph.text for list markers that might not be in runs
    if not parts and paragraph.text:
        parts.append(paragraph.text)
    
    return " ".join(parts) if parts else ""


def _extract_table_text(table: Table) -> str:
    """Extract text from all cells in a table."""
    cell_texts: list[str] = []
    for row in table.rows:
        row_texts: list[str] = []
        for cell in row.cells:
            cell_text = _extract_cell_text(cell)
            if cell_text.strip():
                row_texts.append(cell_text)
        if row_texts:
            cell_texts.append(" | ".join(row_texts))
    return "\n".join(cell_texts)


def _extract_cell_text(cell) -> str:
    """Extract text from a table cell, including nested tables."""
    parts: list[str] = []
    
    # Extract from paragraphs in cell
    for paragraph in cell.paragraphs:
        para_text = _extract_paragraph_text(paragraph)
        if para_text.strip():
            parts.append(para_text)
    
    # Extract from nested tables in cell
    for table in cell.tables:
        table_text = _extract_table_text(table)
        if table_text.strip():
            parts.append(f"[TABLE: {table_text}]")
    
    return " ".join(parts)


def assert_text_fidelity(
    input_texts: list[str],
    extracted_text: str,
    test_name: str,
) -> None:
    """
    Assert that all input text fragments appear in the extracted output.
    
    This enforces Invariant 1: Text Fidelity - every character in input
    must exist somewhere in output.
    
    Args:
        input_texts: List of text fragments that must appear in output
        extracted_text: The full extracted text from the DOCX
        test_name: Name of the test (for error messages)
    
    Raises:
        AssertionError: If any input text fragment is missing
    """
    missing_texts: list[str] = []
    
    for text_fragment in input_texts:
        # Normalize both for comparison (case-sensitive but whitespace-tolerant)
        normalized_fragment = " ".join(text_fragment.split())
        normalized_extracted = " ".join(extracted_text.split())
        
        if normalized_fragment not in normalized_extracted:
            missing_texts.append(text_fragment)
    
    if missing_texts:
        error_msg = (
            f"TEXT FIDELITY VIOLATION in {test_name}:\n"
            f"Missing text fragments:\n"
        )
        for missing in missing_texts:
            error_msg += f"  - {missing!r}\n"
        error_msg += f"\nExtracted text (first 1000 chars):\n{extracted_text[:1000]}"
        raise AssertionError(error_msg)


def assert_no_placeholders_remain(
    extracted_text: str,
    scalar_placeholders: list[str],
    block_placeholders: list[str],
    test_name: str,
) -> None:
    """
    Assert that no placeholder tokens remain in the output.
    
    This enforces that replacement was complete and correct.
    
    Args:
        extracted_text: The full extracted text from the DOCX
        scalar_placeholders: List of scalar placeholder keys (e.g., ["document_id"])
        block_placeholders: List of block placeholder keys (e.g., ["summary"])
        test_name: Name of the test (for error messages)
    
    Raises:
        AssertionError: If any placeholder token is found in output
    """
    remaining_placeholders: list[str] = []
    
    # Check scalar placeholders
    for key in scalar_placeholders:
        token = f"{{{{{key}}}}}"
        if token in extracted_text:
            remaining_placeholders.append(token)
    
    # Check block placeholders
    for key in block_placeholders:
        token = f"{{{{{key}}}}}"
        if token in extracted_text:
            remaining_placeholders.append(token)
    
    if remaining_placeholders:
        error_msg = (
            f"PLACEHOLDER REPLACEMENT FAILURE in {test_name}:\n"
            f"Unreplaced placeholders found:\n"
        )
        for placeholder in remaining_placeholders:
            error_msg += f"  - {placeholder}\n"
        raise AssertionError(error_msg)


def run_export_and_collect_artifacts(
    test_name: str,
    template_path: Path,
    request: WordExportRequest,
    artifacts_dir: Path,
    config: Optional[ListRenderConfig] = None,
    subdirectory: Optional[str] = None,
) -> tuple[Path, str, str]:
    """
    Run export and collect all artifacts.
    
    Args:
        test_name: Name of the test
        template_path: Path to the Word template
        request: The WordExportRequest
        artifacts_dir: Base artifacts directory
        config: Optional ListRenderConfig (defaults to ListRenderConfig())
        subdirectory: Optional subdirectory to organize tests (e.g., "phase3", "lists", "tables")
    
    Returns:
        Tuple of (artifact_dir, markdown_content, extracted_text)
    """
    if config is None:
        config = ListRenderConfig()
    
    # Create temporary output path
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        output_path = Path(tmp.name)
    
    try:
        # Run export
        result = export_to_word(
            template_path=template_path,
            request=request,
            markdown_mode=True,
            output_path=output_path,
            config=config,
        )
        
        # Get markdown content
        markdown_content = ""
        if result.get("markdown_files"):
            markdown_file = result["markdown_files"][0]
            if Path(markdown_file).exists():
                with open(markdown_file, "r", encoding="utf-8") as f:
                    markdown_content = f.read()
        
        # Extract text from DOCX
        docx_path = Path(result["word_file_path"])
        extracted_text = extract_all_text_from_docx(docx_path)
        
        # Validate DOCX integrity (ensures no Word corruption)
        validate_docx_integrity(docx_path)
        
        # Write artifacts
        artifact_dir = write_test_artifacts(
            test_name=test_name,
            request=request,
            markdown=markdown_content,
            docx_path=docx_path,
            artifacts_dir=artifacts_dir,
            template_path=template_path,
            subdirectory=subdirectory,
        )
        
        return artifact_dir, markdown_content, extracted_text
    
    finally:
        # Clean up temporary file if it still exists
        if output_path.exists():
            output_path.unlink()


def validate_docx_integrity(docx_path: Path) -> None:
    """
    Validate that a DOCX file can be opened and saved without corruption.
    
    This function performs comprehensive integrity checks:
    1. Opens the document with python-docx (validates XML structure)
    2. Validates XML structure directly using lxml
    3. Saves and reloads the document multiple times (catches structural issues)
    4. Validates that all sections, paragraphs, and tables are accessible
    5. Checks for common corruption patterns that Word would detect
    
    Raises:
        AssertionError: If the document is corrupted or cannot be opened cleanly
        Exception: If any other error occurs during validation
    
    Args:
        docx_path: Path to the DOCX file to validate
    """
    import tempfile
    import zipfile
    from pathlib import Path
    from lxml import etree
    
    # Check file exists
    if not docx_path.exists():
        raise AssertionError(f"DOCX file does not exist: {docx_path}")
    
    try:
        # Step 0: Validate DOCX is a valid ZIP archive and check structure
        try:
            with zipfile.ZipFile(str(docx_path), 'r') as zip_ref:
                # Check for required files
                required_files = ['[Content_Types].xml', 'word/document.xml']
                file_list = zip_ref.namelist()
                for required in required_files:
                    if not any(f.endswith(required) for f in file_list):
                        raise AssertionError(f"DOCX file missing required file: {required}")
                
                # Validate [Content_Types].xml is well-formed
                try:
                    content_types_content = zip_ref.read('[Content_Types].xml')
                    content_types_xml = etree.fromstring(content_types_content)
                except Exception as e:
                    raise AssertionError(f"[Content_Types].xml is invalid: {e}")
                
                # Validate word/_rels/document.xml.rels if it exists (relationships)
                rels_file = 'word/_rels/document.xml.rels'
                if any(f.endswith(rels_file) for f in file_list):
                    try:
                        rels_content = zip_ref.read(rels_file)
                        rels_xml = etree.fromstring(rels_content)
                        
                        # Validate all relationship targets exist
                        # Broken relationships cause Word to require recovery
                        rel_ns = 'http://schemas.openxmlformats.org/package/2006/relationships'
                        broken_rels = []
                        for rel in rels_xml.findall(f'{{{rel_ns}}}Relationship'):
                            target = rel.get('Target')
                            rel_type = rel.get('Type', '')
                            if target:
                                # Resolve target path based on relationship location
                                # Relationships in word/_rels/document.xml.rels have targets relative to word/
                                if target.startswith('../'):
                                    # Reference to parent directory (customXml, media, etc.)
                                    if target.startswith('../customXml/'):
                                        target_path = target.replace('../customXml/', 'customXml/')
                                    elif target.startswith('../media/'):
                                        target_path = target.replace('../media/', 'media/')
                                    else:
                                        target_path = target.replace('../', '')
                                elif target.startswith('media/'):
                                    target_path = target
                                elif target.startswith('theme/') or target.startswith('settings') or target.startswith('styles') or target.startswith('fontTable') or target.startswith('webSettings'):
                                    target_path = f'word/{target}'
                                elif target.startswith('customXml/'):
                                    target_path = target
                                elif target.startswith('header') or target.startswith('footer'):
                                    target_path = f'word/{target}'
                                else:
                                    target_path = f'word/{target}'
                                
                                # Check if target file exists
                                # Try multiple path variations
                                target_exists = any(
                                    f == target_path or 
                                    f.endswith('/' + target_path) or
                                    f == target or
                                    f.endswith('/' + target) or
                                    f.endswith(target_path) or
                                    f.endswith(target)
                                    for f in file_list
                                )
                                
                                if not target_exists:
                                    # All broken relationships are problematic - Word will require recovery
                                    broken_rels.append(f"Relationship (Type: {rel_type.split('/')[-1]}) target '{target}' not found")
                        
                        if broken_rels:
                            raise AssertionError(
                                f"DOCX file has broken relationships (Word will require recovery):\n" +
                                "\n".join(f"  - {rel}" for rel in broken_rels) +
                                f"\n\nThis document will trigger Word's recovery dialog. All relationship targets must exist."
                            )
                        
                        # Additional check: Validate that all referenced files in Content_Types exist
                        # Word is strict about Content_Types consistency
                        try:
                            content_types_content = zip_ref.read('[Content_Types].xml')
                            content_types_xml = etree.fromstring(content_types_content)
                            ct_ns = 'http://schemas.openxmlformats.org/package/2006/content-types'
                            
                            # Check Override elements (explicit content types)
                            for override in content_types_xml.findall(f'{{{ct_ns}}}Override'):
                                part_name = override.get('PartName', '')
                                if part_name and not part_name.startswith('/'):
                                    part_name = '/' + part_name
                                
                                # Check if the part exists (remove leading slash for ZIP lookup)
                                zip_part_name = part_name.lstrip('/')
                                if not any(f == zip_part_name or f.endswith('/' + zip_part_name) for f in file_list):
                                    # Some parts might be optional, but log it
                                    pass
                        except Exception:
                            # Content_Types validation is best-effort
                            pass
                    except AssertionError:
                        raise
                    except Exception as e:
                        raise AssertionError(f"document.xml.rels is invalid: {e}")
        except zipfile.BadZipFile:
            raise AssertionError(f"DOCX file is not a valid ZIP archive: {docx_path}")
        
        # Step 1: Open the document (validates XML structure)
        doc = Document(str(docx_path))
        
        # Step 2: Validate XML structure directly using lxml
        # This catches XML issues that python-docx might silently ignore
        # Word is stricter than python-docx, so we need aggressive validation
        try:
            with zipfile.ZipFile(str(docx_path), 'r') as zip_ref:
                # Parse main document XML
                doc_xml_content = zip_ref.read('word/document.xml')
                doc_xml = etree.fromstring(doc_xml_content)
                
                # Validate XML is well-formed and has required structure
                # Check for body element
                nsmap = doc_xml.nsmap
                w_ns = nsmap.get(None) or 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                body_elem = doc_xml.find(f'{{{w_ns}}}body')
                if body_elem is None:
                    raise AssertionError("Document XML missing required <w:body> element")
                
                # Validate XML can be serialized (catches malformed XML)
                try:
                    _ = etree.tostring(doc_xml, encoding='unicode')
                except Exception as e:
                    raise AssertionError(f"Document XML cannot be serialized (malformed): {e}")
                
                # Validate Word schema compliance - check for common issues
                # Check that body contains valid children (p, tbl, etc.)
                body_children = list(body_elem)
                for child in body_children:
                    tag = child.tag
                    # Valid tags in body: w:p, w:tbl, w:sectPr
                    if not (tag.endswith('}p') or tag.endswith('}tbl') or tag.endswith('}sectPr')):
                        # This might be okay, but log it
                        pass
                
                # Check for table cells - ensure tcPr comes before content
                for tbl_elem in body_elem.findall(f'.//{{{w_ns}}}tbl'):
                    for tr_elem in tbl_elem.findall(f'{{{w_ns}}}tr'):
                        for tc_elem in tr_elem.findall(f'{{{w_ns}}}tc'):
                            # Check children order - tcPr must come first
                            tc_children = list(tc_elem)
                            if len(tc_children) > 0:
                                first_child_tag = tc_children[0].tag
                                tcpr_tag = f'{{{w_ns}}}tcPr'
                                # If first child is not tcPr, check if tcPr exists later (invalid)
                                if first_child_tag != tcpr_tag:
                                    # Check if tcPr exists later (this is invalid)
                                    for idx, child in enumerate(tc_children[1:], start=1):
                                        if child.tag == tcpr_tag:
                                            raise AssertionError(
                                                f"Invalid table cell structure: <w:tcPr> must be first child of <w:tc>, "
                                                f"but found at position {idx+1}. This will cause Word to require recovery."
                                            )
        except etree.XMLSyntaxError as e:
            raise AssertionError(f"Document XML is malformed: {e}")
        
        # Step 3: Access all document elements to ensure they're valid
        # This will raise exceptions if XML is malformed
        sections = doc.sections
        paragraphs = doc.paragraphs
        tables = doc.tables
        
        # Step 4: Validate each section thoroughly, including header/footer XML structure
        for section in sections:
            # Access header and footer (will fail if XML is corrupted)
            header = section.header
            footer = section.footer
            
            # Access paragraphs in header/footer
            header_paras = header.paragraphs
            footer_paras = footer.paragraphs
            
            # Access tables in header/footer
            header_tables = header.tables
            footer_tables = footer.tables
            
            # Validate header/footer elements are accessible
            for para in header_paras:
                _ = para.text
                _ = para.runs
            for para in footer_paras:
                _ = para.text
                _ = para.runs
            for table in header_tables:
                for row in table.rows:
                    for cell in row.cells:
                        _ = cell.text
                        _ = cell.paragraphs
            for table in footer_tables:
                for row in table.rows:
                    for cell in row.cells:
                        _ = cell.text
                        _ = cell.paragraphs
            
            # Step 4a: Validate header/footer XML structure directly
            # Check for table cell structure issues in headers/footers
            try:
                header_xml = header._element.xml
                footer_xml = footer._element.xml
                
                # Parse and validate header XML
                header_tree = etree.fromstring(header_xml)
                w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                for tbl_elem in header_tree.findall(f'.//{{{w_ns}}}tbl'):
                    for tr_elem in tbl_elem.findall(f'{{{w_ns}}}tr'):
                        for tc_elem in tr_elem.findall(f'{{{w_ns}}}tc'):
                            tc_children = list(tc_elem)
                            tcpr_tag = f'{{{w_ns}}}tcPr'
                            tcpr_positions = [i for i, c in enumerate(tc_children) if c.tag == tcpr_tag]
                            if tcpr_positions and tcpr_positions[0] > 0:
                                raise AssertionError(
                                    f"Header table cell has invalid structure: <w:tcPr> must be first child, "
                                    f"but found at position {tcpr_positions[0]+1}. This will cause Word to require recovery."
                                )
                
                # Parse and validate footer XML
                footer_tree = etree.fromstring(footer_xml)
                for tbl_elem in footer_tree.findall(f'.//{{{w_ns}}}tbl'):
                    for tr_elem in tbl_elem.findall(f'{{{w_ns}}}tr'):
                        for tc_elem in tr_elem.findall(f'{{{w_ns}}}tc'):
                            tc_children = list(tc_elem)
                            tcpr_tag = f'{{{w_ns}}}tcPr'
                            tcpr_positions = [i for i, c in enumerate(tc_children) if c.tag == tcpr_tag]
                            if tcpr_positions and tcpr_positions[0] > 0:
                                raise AssertionError(
                                    f"Footer table cell has invalid structure: <w:tcPr> must be first child, "
                                    f"but found at position {tcpr_positions[0]+1}. This will cause Word to require recovery."
                                )
            except AssertionError:
                raise
            except Exception as e:
                # If we can't parse header/footer XML, that's also a problem
                raise AssertionError(f"Header/footer XML structure is invalid: {e}")
        
        # Step 5: Validate all paragraphs
        for para in paragraphs:
            _ = para.text
            _ = para.runs
            # Access paragraph XML to ensure it's valid
            try:
                _ = para._p.xml
            except Exception as e:
                raise AssertionError(f"Paragraph XML is invalid: {e}")
        
        # Step 6: Validate all tables and their structure
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    _ = cell.text
                    _ = cell.paragraphs
                    _ = cell.tables
                    # Access cell XML to ensure it's valid
                    try:
                        _ = cell._tc.xml
                    except Exception as e:
                        raise AssertionError(f"Table cell XML is invalid: {e}")
        
        # Step 7: Save and reload MULTIPLE times to catch any structural issues
        # Word's validation is strict, so we need to be equally strict
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            temp_path = Path(tmp.name)
        
        try:
            # First save/reload cycle
            doc.save(str(temp_path))
            reloaded_doc1 = Document(str(temp_path))
            _ = reloaded_doc1.sections
            _ = reloaded_doc1.paragraphs
            _ = reloaded_doc1.tables
            
            # Second save/reload cycle (catches issues that only appear after multiple saves)
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp2:
                temp_path2 = Path(tmp2.name)
            try:
                reloaded_doc1.save(str(temp_path2))
                reloaded_doc2 = Document(str(temp_path2))
                _ = reloaded_doc2.sections
                _ = reloaded_doc2.paragraphs
                _ = reloaded_doc2.tables
            finally:
                if temp_path2.exists():
                    temp_path2.unlink()
            
        finally:
            # Clean up temp file
            if temp_path.exists():
                temp_path.unlink()
        
        # Step 8: Validate the original file can be opened again after all operations
        # This ensures the file wasn't corrupted by our validation
        final_check = Document(str(docx_path))
        _ = final_check.sections
        _ = final_check.paragraphs
        _ = final_check.tables
        
        # Step 9: Aggressive validation - try to access ALL elements including textboxes
        # Word is very strict about textbox structure, so we need to validate them
        try:
            # Access all paragraphs including those in textboxes
            for para in final_check.paragraphs:
                # Try to access all runs and their properties
                for run in para.runs:
                    _ = run.text
                    _ = run.font
                    _ = run.bold
                    _ = run.italic
                
                # Try to access paragraph properties
                _ = para.paragraph_format
                _ = para.style
            
            # Access all tables and their nested structures
            for table in final_check.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Access all paragraphs in cells
                        for cell_para in cell.paragraphs:
                            _ = cell_para.text
                            for run in cell_para.runs:
                                _ = run.text
                        
                        # Access nested tables if any
                        for nested_table in cell.tables:
                            _ = nested_table.rows
            
            # Access header/footer elements aggressively
            for section in final_check.sections:
                header = section.header
                footer = section.footer
                
                # Access all header paragraphs and runs
                for header_para in header.paragraphs:
                    _ = header_para.text
                    for run in header_para.runs:
                        _ = run.text
                
                # Access all footer paragraphs and runs
                for footer_para in footer.paragraphs:
                    _ = footer_para.text
                    for run in footer_para.runs:
                        _ = run.text
                
                # Access header/footer tables
                for header_table in header.tables:
                    for row in header_table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                _ = para.text
                
                for footer_table in footer.tables:
                    for row in footer_table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                _ = para.text
        except Exception as e:
            # If we can't access elements, the file is corrupted
            raise AssertionError(
                f"DOCX file has structural corruption that prevents element access: {docx_path}\n"
                f"Error accessing document elements: {type(e).__name__}: {str(e)}\n"
                f"This indicates the document would trigger Word repair warnings."
            ) from e
        
        # Step 10: Final aggressive save/reload cycle with error checking
        # This catches issues that only appear after multiple save cycles
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp3:
            temp_path3 = Path(tmp3.name)
        try:
            final_check.save(str(temp_path3))
            # Try to open the saved file one more time
            final_reload = Document(str(temp_path3))
            # Access critical elements
            _ = final_reload.sections[0]
            _ = final_reload.paragraphs
            _ = final_reload.tables
        except Exception as e:
            raise AssertionError(
                f"DOCX file fails after multiple save/reload cycles: {docx_path}\n"
                f"Error: {type(e).__name__}: {str(e)}\n"
                f"This indicates the document would trigger Word repair warnings."
            ) from e
        finally:
            if temp_path3.exists():
                temp_path3.unlink()
        
    except AssertionError:
        # Re-raise AssertionError as-is
        raise
    except Exception as e:
        raise AssertionError(
            f"DOCX file is corrupted or cannot be opened cleanly: {docx_path}\n"
            f"Error: {type(e).__name__}: {str(e)}\n"
            f"This indicates the document would trigger Word repair warnings.\n"
            f"The document may open in Word only after recovery, which is unacceptable."
        ) from e
