"""
Golden regression test for complex real-world scenarios.

This test uses a comprehensive input that exercises all major features:
- Scalar replacement in multiple locations
- Block markdown with headings, paragraphs, lists, tables
- Nested structures
- Mixed content types

This test is the long-term regression anchor and must pass for all releases.
"""
import pytest
from pathlib import Path
from docx import Document

from docx_template_export.models.export_models import WordExportRequest
from tests.test_artifacts_helpers import (
    run_export_and_collect_artifacts,
    assert_text_fidelity,
    assert_no_placeholders_remain,
)


@pytest.fixture
def artifacts_dir(test_output_dir):
    """Return artifacts directory for this test suite."""
    artifacts_base = test_output_dir.parent / "artifacts"
    artifacts_base.mkdir(exist_ok=True)
    return artifacts_base


@pytest.fixture
def golden_template(test_fixtures_dir):
    """Create comprehensive template for golden regression test."""
    template_path = test_fixtures_dir / "golden_template.docx"
    doc = Document()
    
    # Header
    header = doc.sections[0].header
    header.paragraphs[0].text = "Document: {{document_id}}"
    
    # Body - scalar fields
    doc.add_heading("Document Information", 1)
    doc.add_paragraph("Title: {{title}}")
    doc.add_paragraph("Author: {{author}}")
    doc.add_paragraph("Version: {{version}}")
    doc.add_paragraph("")
    
    # Body - block fields
    doc.add_heading("Executive Summary", 1)
    doc.add_paragraph("{{executive_summary}}")
    doc.add_paragraph("")
    
    doc.add_heading("Introduction", 1)
    doc.add_paragraph("{{introduction}}")
    doc.add_paragraph("")
    
    doc.add_heading("Main Content", 1)
    doc.add_paragraph("{{main_content}}")
    doc.add_paragraph("")
    
    doc.add_heading("Conclusion", 1)
    doc.add_paragraph("{{conclusion}}")
    doc.add_paragraph("")
    
    # Footer
    footer = doc.sections[0].footer
    footer.paragraphs[0].text = "Page {{page_number}}"
    
    template_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(template_path))
    return template_path


def test_golden_regression_comprehensive(artifacts_dir, golden_template):
    """
    Golden regression test with comprehensive real-world input.
    
    This test exercises:
    - Scalar replacement (header, body, footer)
    - Block markdown with headings
    - Block markdown with paragraphs
    - Block markdown with lists (nested)
    - Block markdown with tables
    - Mixed content types
    
    This test must pass for all releases and serves as the regression anchor.
    """
    request = WordExportRequest(
        scalar_fields={
            "document_id": "GOLDEN-2024-001",
            "title": "Golden Regression Test Document",
            "author": "Test Suite",
            "version": "1.0.0",
            "page_number": "1",
        },
        block_fields={
            "executive_summary": """# Executive Summary

This document demonstrates comprehensive markdown export capabilities.

Key points:
- Feature A: Working correctly
- Feature B: Validated
- Feature C: Tested thoroughly
""",
            "introduction": """# Introduction

This is the introduction section with **bold** and *italic* text.

The library ensures:
1. Text fidelity (no text loss)
2. Determinism (same input ? same output)
3. Observability (all artifacts written)
""",
            "main_content": """# Main Content

## Section 1: Lists

### Bullet Lists
- Level 1 item
  - Level 2 item
    - Level 3 item
  - Another level 2 item

### Numbered Lists
1. First numbered item
   1. Nested numbered item
   2. Another nested item
2. Second numbered item

## Section 2: Tables

| Component | Status | Notes |
|-----------|--------|-------|
| Parser    | OK     | Working correctly |
| Renderer  | OK     | Validated |
| Tests     | OK     | Comprehensive |

## Section 3: Mixed Content

This paragraph appears before a list.

- Mixed bullet item
- Another bullet item

This paragraph appears after a list.
""",
            "conclusion": """# Conclusion

This golden regression test validates:

1. **Text Fidelity**: All input text appears in output
2. **Determinism**: Same input produces identical output
3. **Observability**: All artifacts are written to disk

The test suite ensures these invariants are never violated.
""",
        },
    )
    
    artifact_dir, markdown, extracted_text = run_export_and_collect_artifacts(
        test_name="golden_regression_comprehensive",
        template_path=golden_template,
        request=request,
        artifacts_dir=artifacts_dir,
        subdirectory="golden_regression",
    )
    
    # Assert text fidelity - comprehensive check
    assert_text_fidelity(
        input_texts=[
            # Scalar fields
            "GOLDEN-2024-001",
            "Golden Regression Test Document",
            "Test Suite",
            "1.0.0",
            # Executive summary
            "Executive Summary",
            "This document demonstrates",
            "Feature A",
            "Feature B",
            "Feature C",
            # Introduction
            "Introduction",
            "bold",
            "italic",
            "Text fidelity",
            "Determinism",
            "Observability",
            # Main content - lists
            "Section 1",
            "Bullet Lists",
            "Level 1 item",
            "Level 2 item",
            "Level 3 item",
            "Numbered Lists",
            "First numbered item",
            "Nested numbered item",
            "Second numbered item",
            # Main content - tables
            "Section 2",
            "Component",
            "Status",
            "Notes",
            "Parser",
            "OK",
            "Working correctly",
            # Main content - mixed
            "Section 3",
            "Mixed Content",
            "This paragraph appears before",
            "Mixed bullet item",
            "This paragraph appears after",
            # Conclusion
            "Conclusion",
            "Text Fidelity",
            "Determinism",
            "Observability",
            "invariants are never violated",
        ],
        extracted_text=extracted_text,
        test_name="golden_regression_comprehensive",
    )
    
    # Assert no placeholders remain
    assert_no_placeholders_remain(
        extracted_text=extracted_text,
        scalar_placeholders=["document_id", "title", "author", "version", "page_number"],
        block_placeholders=["executive_summary", "introduction", "main_content", "conclusion"],
        test_name="golden_regression_comprehensive",
    )
    
    # Assert artifacts exist
    assert (artifact_dir / "input.json").exists(), "input.json must exist"
    assert (artifact_dir / "rendered_markdown.md").exists(), "rendered_markdown.md must exist"
    assert (artifact_dir / "output.docx").exists(), "output.docx must exist"
    
    # Assert determinism: run twice and compare
    artifact_dir2, markdown2, extracted_text2 = run_export_and_collect_artifacts(
        test_name="golden_regression_comprehensive_run2",
        template_path=golden_template,
        request=request,
        artifacts_dir=artifacts_dir,
        subdirectory="golden_regression",
    )
    
    assert markdown == markdown2, "Markdown must be identical across runs (determinism)"
    normalized1 = " ".join(extracted_text.split())
    normalized2 = " ".join(extracted_text2.split())
    assert normalized1 == normalized2, "Extracted text must be identical across runs (determinism)"
