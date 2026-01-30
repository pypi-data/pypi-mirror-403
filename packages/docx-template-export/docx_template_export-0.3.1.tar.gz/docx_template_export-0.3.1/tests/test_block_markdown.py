"""
Test suite for block markdown replacement in BODY FLOW and table cells.

Enforces:
- Invariant 1: Text Fidelity - all block content appears in output
- Invariant 2: Determinism - same markdown produces same structure
- Invariant 3: Observability - all artifacts written to disk
"""
import pytest
from pathlib import Path
from docx import Document

from docx_template_export.models.export_models import WordExportRequest
from tests.test_artifacts_helpers import (
    run_export_and_collect_artifacts,
    assert_text_fidelity,
    assert_no_placeholders_remain,
)


@pytest.fixture
def artifacts_dir(test_output_dir):
    """Return artifacts directory for this test suite."""
    artifacts_base = test_output_dir.parent / "artifacts"
    artifacts_base.mkdir(exist_ok=True)
    return artifacts_base


@pytest.fixture
def block_template_body(test_fixtures_dir):
    """Create template with block placeholders in BODY FLOW."""
    template_path = test_fixtures_dir / "block_template_body.docx"
    doc = Document()
    doc.add_paragraph("{{introduction}}")
    doc.add_paragraph("{{body}}")
    doc.add_paragraph("{{conclusion}}")
    template_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(template_path))
    return template_path


@pytest.fixture
def block_template_table_cell(test_fixtures_dir):
    """Create template with block placeholder in table cell."""
    template_path = test_fixtures_dir / "block_template_table_cell.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].paragraphs[0].text = "{{content}}"
    template_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(template_path))
    return template_path


def test_block_markdown_body_flow_paragraphs(artifacts_dir, block_template_body):
    """Test block markdown with paragraphs in BODY FLOW."""
    request = WordExportRequest(
        scalar_fields={},
        block_fields={
            "introduction": "This is the introduction paragraph.",
            "body": "This is the body paragraph with **bold** and *italic* text.",
            "conclusion": "This is the conclusion paragraph.",
        },
    )
    
    artifact_dir, markdown, extracted_text = run_export_and_collect_artifacts(
        test_name="block_markdown_body_flow_paragraphs",
        template_path=block_template_body,
        request=request,
        artifacts_dir=artifacts_dir,
        subdirectory="block_markdown",
    )
    
    # Assert text fidelity
    assert_text_fidelity(
        input_texts=[
            "This is the introduction paragraph.",
            "This is the body paragraph",
            "bold",
            "italic",
            "This is the conclusion paragraph.",
        ],
        extracted_text=extracted_text,
        test_name="block_markdown_body_flow_paragraphs",
    )
    
    # Assert no placeholders remain
    assert_no_placeholders_remain(
        extracted_text=extracted_text,
        scalar_placeholders=[],
        block_placeholders=["introduction", "body", "conclusion"],
        test_name="block_markdown_body_flow_paragraphs",
    )


def test_block_markdown_body_flow_headings(artifacts_dir, block_template_body):
    """Test block markdown with headings in BODY FLOW."""
    request = WordExportRequest(
        scalar_fields={},
        block_fields={
            "introduction": "# Introduction\n\nThis is the intro text.",
            "body": "## Main Section\n\nBody content here.\n\n### Subsection\n\nMore content.",
            "conclusion": "## Conclusion\n\nFinal thoughts.",
        },
    )
    
    artifact_dir, markdown, extracted_text = run_export_and_collect_artifacts(
        test_name="block_markdown_body_flow_headings",
        template_path=block_template_body,
        request=request,
        artifacts_dir=artifacts_dir,
        subdirectory="block_markdown",
    )
    
    # Assert text fidelity
    assert_text_fidelity(
        input_texts=[
            "Introduction",
            "This is the intro text",
            "Main Section",
            "Body content here",
            "Subsection",
            "More content",
            "Conclusion",
            "Final thoughts",
        ],
        extracted_text=extracted_text,
        test_name="block_markdown_body_flow_headings",
    )
    
    # Assert no placeholders remain
    assert_no_placeholders_remain(
        extracted_text=extracted_text,
        scalar_placeholders=[],
        block_placeholders=["introduction", "body", "conclusion"],
        test_name="block_markdown_body_flow_headings",
    )


def test_block_markdown_table_cell(artifacts_dir, block_template_table_cell):
    """Test block markdown inside table cell."""
    request = WordExportRequest(
        scalar_fields={},
        block_fields={
            "content": """# Heading in Cell

This is a paragraph inside a table cell.

- Bullet item one
- Bullet item two

1. Numbered item one
2. Numbered item two
""",
        },
    )
    
    artifact_dir, markdown, extracted_text = run_export_and_collect_artifacts(
        test_name="block_markdown_table_cell",
        template_path=block_template_table_cell,
        request=request,
        artifacts_dir=artifacts_dir,
        subdirectory="block_markdown",
    )
    
    # Assert text fidelity
    assert_text_fidelity(
        input_texts=[
            "Heading in Cell",
            "This is a paragraph inside a table cell",
            "Bullet item one",
            "Bullet item two",
            "Numbered item one",
            "Numbered item two",
        ],
        extracted_text=extracted_text,
        test_name="block_markdown_table_cell",
    )
    
    # Assert no placeholders remain
    assert_no_placeholders_remain(
        extracted_text=extracted_text,
        scalar_placeholders=[],
        block_placeholders=["content"],
        test_name="block_markdown_table_cell",
    )


def test_block_markdown_mixed_content(artifacts_dir, block_template_body):
    """Test block markdown with mixed content types."""
    request = WordExportRequest(
        scalar_fields={},
        block_fields={
            "introduction": """# Introduction

This is a paragraph with **bold** and *italic* text.

- First bullet
- Second bullet
""",
            "body": """## Body Section

| Column A | Column B |
|----------|----------|
| Value 1  | Value 2  |
""",
            "conclusion": "## Conclusion\n\nFinal paragraph.",
        },
    )
    
    artifact_dir, markdown, extracted_text = run_export_and_collect_artifacts(
        test_name="block_markdown_mixed_content",
        template_path=block_template_body,
        request=request,
        artifacts_dir=artifacts_dir,
        subdirectory="block_markdown",
    )
    
    # Assert text fidelity
    assert_text_fidelity(
        input_texts=[
            "Introduction",
            "bold",
            "italic",
            "First bullet",
            "Second bullet",
            "Body Section",
            "Column A",
            "Column B",
            "Value 1",
            "Value 2",
            "Conclusion",
            "Final paragraph",
        ],
        extracted_text=extracted_text,
        test_name="block_markdown_mixed_content",
    )
    
    # Assert no placeholders remain
    assert_no_placeholders_remain(
        extracted_text=extracted_text,
        scalar_placeholders=[],
        block_placeholders=["introduction", "body", "conclusion"],
        test_name="block_markdown_mixed_content",
    )


def test_block_markdown_determinism(artifacts_dir, block_template_body):
    """Test that same markdown produces identical output (Invariant 2: Determinism)."""
    request = WordExportRequest(
        scalar_fields={},
        block_fields={
            "introduction": "# Test\n\nContent here.",
            "body": "- Item 1\n- Item 2",
            "conclusion": "Final text.",
        },
    )
    
    # Run export twice
    artifact_dir1, markdown1, extracted_text1 = run_export_and_collect_artifacts(
        test_name="block_markdown_determinism_run1",
        template_path=block_template_body,
        request=request,
        artifacts_dir=artifacts_dir,
        subdirectory="block_markdown",
    )
    
    artifact_dir2, markdown2, extracted_text2 = run_export_and_collect_artifacts(
        test_name="block_markdown_determinism_run2",
        template_path=block_template_body,
        request=request,
        artifacts_dir=artifacts_dir,
        subdirectory="block_markdown",
    )
    
    # Assert markdown is identical
    assert markdown1 == markdown2, "Markdown output must be identical for same input"
    
    # Assert extracted text is identical (normalized)
    normalized1 = " ".join(extracted_text1.split())
    normalized2 = " ".join(extracted_text2.split())
    assert normalized1 == normalized2, "Extracted text must be identical for same input"
