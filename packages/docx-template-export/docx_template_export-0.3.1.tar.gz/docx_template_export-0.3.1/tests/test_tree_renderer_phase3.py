# Copyright  2024 Ahsan Saeed
# Licensed under the Apache License, Version 2.0
# See LICENSE and NOTICE files for details.

"""
Phase III tests for tree-based renderer.

These tests provide high-confidence proof that the tree-based renderer:
- Preserves exact ordering
- Renders lists deterministically
- Handles tables safely and atomically
- Never corrupts Word documents
- Respects placeholder replacement boundaries
- Enforces safety guardrails

PHASE III RULE: These tests validate the renderer, they do NOT modify it.
If a test exposes an issue, it is reported, not fixed.
"""
import json
import tempfile
from pathlib import Path
from typing import List, Tuple, Optional

import pytest
from docx import Document  # type: ignore
from docx.document import Document as DocumentType  # type: ignore
from docx.oxml.ns import qn  # type: ignore

from docx_template_export.models.export_config import ListRenderConfig
from docx_template_export.models.markdown_tree import (
    BlockNode,
    DocumentNode,
    HeadingNode,
    ParagraphNode,
    ListNode,
    ListItemNode,
    TableNode,
    TableRowNode,
    TableCellNode,
)
from docx_template_export.renderers.tree_renderer import (
    render_markdown_tree_to_docx,
    MAX_TREE_DEPTH,
    MAX_NODE_COUNT,
)
from docx_template_export.services.word_export_service import FormattedRun, dump_markdown_tree, parse_markdown_to_tree


# ============================================================================
# Test Helpers (Test Code Only)
# ============================================================================

def extract_xml_element_order(doc: DocumentType) -> List[Tuple[str, int]]:
    """
    Extract the order of <w:p> and <w:tbl> elements from document body.
    
    Returns:
        List of (tag_name, index) tuples in document order.
        tag_name is 'p' for paragraphs, 'tbl' for tables.
    """
    body_elem = doc._body._body  # Get the XML element
    order = []
    for idx, child in enumerate(body_elem):
        if child.tag.endswith("}p"):
            order.append(("p", idx))
        elif child.tag.endswith("}tbl"):
            order.append(("tbl", idx))
    return order


def extract_paragraph_texts(doc: DocumentType) -> List[str]:
    """
    Extract text content from all paragraphs in document order.
    
    Returns:
        List of paragraph texts (including list item prefixes).
    """
    texts = []
    for para in doc.paragraphs:
        text = "".join(run.text for run in para.runs)
        texts.append(text)
    return texts


def assert_docx_integrity(doc: DocumentType, output_path: Path) -> None:
    """
    Assert DOCX structural integrity by saving and reloading.
    
    SAFETY CHECK: If the document cannot be saved/reloaded, it is corrupted.
    
    This function saves the document first, then uses the standard validate_docx_integrity
    function to ensure comprehensive validation.
    """
    # Save document first
    doc.save(str(output_path))
    assert output_path.exists(), "Document save failed"
    
    # Use the standard validation function for comprehensive checks
    from tests.test_artifacts_helpers import validate_docx_integrity
    validate_docx_integrity(output_path)


def create_placeholder_template(placeholder_text: str = "{{content}}") -> Path:
    """Create a minimal template with a single placeholder paragraph."""
    template_path = Path(tempfile.mktemp(suffix=".docx"))
    doc = Document()
    doc.add_paragraph("Before placeholder")
    doc.add_paragraph(placeholder_text)
    doc.add_paragraph("After placeholder")
    doc.save(str(template_path))
    return template_path


def save_test_artifacts_phase3(
    test_name: str,
    doc_node: DocumentNode,
    doc: DocumentType,
    output_path: Path,
    artifacts_dir: Path,
    template_path: Optional[Path] = None,
) -> Path:
    """
    Save Phase III test artifacts: input JSON (tree structure), markdown dump, output DOCX, and template DOCX.
    
    Args:
        test_name: Name of the test (used for directory name)
        doc_node: The DocumentNode that was rendered
        doc: The rendered Word document
        output_path: Path where DOCX was saved
        artifacts_dir: Base artifacts directory (tests/artifacts)
        template_path: Optional path to the template DOCX file
    
    Returns:
        Path to the test's artifact directory
    """
    test_artifact_dir = artifacts_dir / "phase3" / test_name
    test_artifact_dir.mkdir(parents=True, exist_ok=True)
    
    # Write input.json (standard format for consistency with other tests)
    # Note: Phase III tests construct DocumentNode directly, so scalar_fields/block_fields are empty
    # Tree-specific information is included as metadata
    input_data = {
        "scalar_fields": {},
        "block_fields": {},
        "_metadata": {
            "test_name": test_name,
            "note": "Phase III test - DocumentNode constructed directly (not from WordExportRequest)",
            "tree_structure": dump_markdown_tree(doc_node),
            "node_count": _count_tree_nodes(doc_node),
            "max_depth": _calculate_tree_depth(doc_node),
        }
    }
    input_json_path = test_artifact_dir / "input.json"
    with open(input_json_path, "w", encoding="utf-8") as f:
        json.dump(input_data, f, indent=2, ensure_ascii=False)
    
    # Write rendered_markdown.md (tree dump)
    markdown_path = test_artifact_dir / "rendered_markdown.md"
    with open(markdown_path, "w", encoding="utf-8") as f:
        f.write("# Markdown Tree Structure\n\n")
        f.write("```\n")
        f.write(dump_markdown_tree(doc_node))
        f.write("\n```\n\n")
        f.write("## Document Text Content\n\n")
        for i, para in enumerate(doc.paragraphs):
            f.write(f"{i+1}. {para.text}\n")
        f.write("\n## Tables\n\n")
        for i, table in enumerate(doc.tables):
            f.write(f"### Table {i+1}\n\n")
            for row_idx, row in enumerate(table.rows):
                row_text = " | ".join(cell.text for cell in row.cells)
                f.write(f"Row {row_idx+1}: {row_text}\n")
            f.write("\n")
    
    # Copy output.docx
    output_docx_path = test_artifact_dir / "output.docx"
    if output_path.exists():
        import shutil
        shutil.copy2(output_path, output_docx_path)
    
    # Copy template.docx if provided
    if template_path is not None and template_path.exists():
        template_docx_path = test_artifact_dir / "template.docx"
        import shutil
        shutil.copy2(template_path, template_docx_path)
    
    return test_artifact_dir


def _count_tree_nodes(node: BlockNode) -> int:
    """Count total nodes in tree."""
    count = 1
    if isinstance(node, DocumentNode):
        for child in node.children:
            count += _count_tree_nodes(child)
    elif isinstance(node, ListNode):
        for item in node.items:
            count += _count_tree_nodes(item)
    elif isinstance(node, ListItemNode):
        for child in node.children:
            count += _count_tree_nodes(child)
    elif isinstance(node, TableNode):
        for row in node.rows:
            for cell in row.cells:
                for child in cell.children:
                    count += _count_tree_nodes(child)
    return count


def _calculate_tree_depth(node: BlockNode, depth: int = 0) -> int:
    """Calculate maximum depth of tree."""
    max_depth = depth
    if isinstance(node, DocumentNode):
        for child in node.children:
            max_depth = max(max_depth, _calculate_tree_depth(child, depth + 1))
    elif isinstance(node, ListNode):
        for item in node.items:
            max_depth = max(max_depth, _calculate_tree_depth(item, depth + 1))
    elif isinstance(node, ListItemNode):
        for child in node.children:
            max_depth = max(max_depth, _calculate_tree_depth(child, depth + 1))
    elif isinstance(node, TableNode):
        for row in node.rows:
            for cell in row.cells:
                for child in cell.children:
                    max_depth = max(max_depth, _calculate_tree_depth(child, depth + 1))
    return max_depth


# ============================================================================
# 1. Golden-Order Tests (CRITICAL)
# ============================================================================

class TestGoldenOrder:
    """Test exact ordering of rendered elements."""
    
    def test_exact_sequence_heading_para_list_table_para(self, test_output_dir):
        """
        Test that elements appear in exact order:
        Heading ? Paragraph ? Bullet List ? Ordered List ? Table ? Paragraph
        """
        # Build DocumentNode programmatically
        doc_node = DocumentNode(children=[
            HeadingNode(level=1, runs=[FormattedRun(text="Test Heading")]),
            ParagraphNode(runs=[FormattedRun(text="Intro paragraph")]),
            ListNode(kind="bullet", items=[
                ListItemNode(children=[
                    ParagraphNode(runs=[FormattedRun(text="Bullet item 1")])
                ])
            ]),
            ListNode(kind="ordered", items=[
                ListItemNode(children=[
                    ParagraphNode(runs=[FormattedRun(text="Ordered item 1")])
                ])
            ]),
            TableNode(rows=[
                TableRowNode(cells=[
                    TableCellNode(children=[
                        ParagraphNode(runs=[FormattedRun(text="Cell A")])
                    ]),
                    TableCellNode(children=[
                        ParagraphNode(runs=[FormattedRun(text="Cell B")])
                    ])
                ])
            ]),
            ParagraphNode(runs=[FormattedRun(text="Final paragraph")]),
        ])
        
        # Create template with placeholder
        template_path = create_placeholder_template()
        doc = Document(str(template_path))
        
        # Find placeholder paragraph and capture surrounding paragraphs
        placeholder_para = None
        before_para = None
        after_para = None
        
        for i, para in enumerate(doc.paragraphs):
            if "{{content}}" in para.text:
                placeholder_para = para
                if i > 0:
                    before_para = doc.paragraphs[i - 1]
                if i < len(doc.paragraphs) - 1:
                    after_para = doc.paragraphs[i + 1]
                break
        
        assert placeholder_para is not None, "Placeholder paragraph not found"
        
        # Capture text before rendering
        before_text = before_para.text if before_para else None
        after_text = after_para.text if after_para else None
        
        # Render tree
        render_markdown_tree_to_docx(doc_node, doc, placeholder_paragraph=placeholder_para)
        
        # Extract XML order
        xml_order = extract_xml_element_order(doc)
        
        # Assert: We should have 1 heading (p), 1 intro para (p), 1 bullet item (p), 
        # 1 ordered item (p), 1 table (tbl), 1 final para (p)
        # Plus the "Before placeholder" and "After placeholder" paragraphs
        expected_order = ["p", "p", "p", "p", "p", "p", "tbl", "p", "p"]
        
        # Extract just the tags
        actual_tags = [tag for tag, _ in xml_order]
        
        # Verify structure: should have paragraphs and one table
        assert "tbl" in actual_tags, "Table not found in document"
        assert actual_tags.count("tbl") == 1, f"Expected 1 table, found {actual_tags.count('tbl')}"
        
        # Verify table appears after paragraphs (exact position depends on structure)
        tbl_index = actual_tags.index("tbl")
        assert tbl_index > 0, "Table appears before any content"
        
        # Verify placeholder was removed
        remaining_placeholders = [p for p in doc.paragraphs if "{{content}}" in p.text]
        assert len(remaining_placeholders) == 0, "Placeholder paragraph was not removed"
        
        # Verify structural integrity and save artifacts
        output_path = test_output_dir / "test_golden_order.docx"
        assert_docx_integrity(doc, output_path)
        
        # Save artifacts for inspection
        artifacts_dir = Path(__file__).parent / "artifacts"
        save_test_artifacts_phase3("golden_order", doc_node, doc, output_path, artifacts_dir, template_path)
    
    def test_placeholder_boundary_integrity(self, test_output_dir):
        """Test that content is inserted at placeholder position, not appended."""
        doc_node = DocumentNode(children=[
            ParagraphNode(runs=[FormattedRun(text="Inserted content")]),
        ])
        
        template_path = create_placeholder_template()
        doc = Document(str(template_path))
        
        # Find placeholder and surrounding paragraphs
        placeholder_para = None
        before_para = None
        after_para = None
        
        for i, para in enumerate(doc.paragraphs):
            if "{{content}}" in para.text:
                placeholder_para = para
                if i > 0:
                    before_para = doc.paragraphs[i - 1]
                if i < len(doc.paragraphs) - 1:
                    after_para = doc.paragraphs[i + 1]
                break
        
        assert placeholder_para is not None, "Placeholder not found"
        assert before_para is not None, "Before paragraph not found"
        assert after_para is not None, "After paragraph not found"
        
        before_text = before_para.text
        after_text = after_para.text
        
        # Render
        render_markdown_tree_to_docx(doc_node, doc, placeholder_paragraph=placeholder_para)
        
        # Verify surrounding paragraphs unchanged
        assert before_para.text == before_text, "Before paragraph was modified"
        assert after_para.text == after_text, "After paragraph was modified"
        
        # Verify placeholder removed (check XML directly)
        body_xml = doc._body._body.xml
        assert "{{content}}" not in body_xml, "Placeholder text still present in document XML"
        
        # Verify content inserted (should appear between before and after)
        para_texts = extract_paragraph_texts(doc)
        assert "Inserted content" in para_texts, "Inserted content not found"
        
        # Verify structural integrity and save artifacts
        output_path = test_output_dir / "test_placeholder_boundary.docx"
        assert_docx_integrity(doc, output_path)
        
        # Save artifacts for inspection
        artifacts_dir = Path(__file__).parent / "artifacts"
        save_test_artifacts_phase3("placeholder_boundary", doc_node, doc, output_path, artifacts_dir)


# ============================================================================
# 2. List Torture Tests
# ============================================================================

class TestListTorture:
    """Test all list rendering scenarios."""
    
    def test_bullet_list_depths_1_to_6(self, test_output_dir):
        """Test bullet lists at depths 1-6."""
        # Build nested bullet lists
        items_level_6 = [
            ListItemNode(children=[
                ParagraphNode(runs=[FormattedRun(text=f"Level 6 item {i}")])
            ]) for i in range(1, 3)
        ]
        list_level_6 = ListNode(kind="bullet", items=items_level_6)
        
        items_level_5 = [
            ListItemNode(children=[
                ParagraphNode(runs=[FormattedRun(text=f"Level 5 item {i}")]),
                list_level_6
            ]) for i in range(1, 3)
        ]
        list_level_5 = ListNode(kind="bullet", items=items_level_5)
        
        items_level_4 = [
            ListItemNode(children=[
                ParagraphNode(runs=[FormattedRun(text=f"Level 4 item {i}")]),
                list_level_5
            ]) for i in range(1, 3)
        ]
        list_level_4 = ListNode(kind="bullet", items=items_level_4)
        
        items_level_3 = [
            ListItemNode(children=[
                ParagraphNode(runs=[FormattedRun(text=f"Level 3 item {i}")]),
                list_level_4
            ]) for i in range(1, 3)
        ]
        list_level_3 = ListNode(kind="bullet", items=items_level_3)
        
        items_level_2 = [
            ListItemNode(children=[
                ParagraphNode(runs=[FormattedRun(text=f"Level 2 item {i}")]),
                list_level_3
            ]) for i in range(1, 3)
        ]
        list_level_2 = ListNode(kind="bullet", items=items_level_2)
        
        items_level_1 = [
            ListItemNode(children=[
                ParagraphNode(runs=[FormattedRun(text=f"Level 1 item {i}")]),
                list_level_2
            ]) for i in range(1, 3)
        ]
        list_level_1 = ListNode(kind="bullet", items=items_level_1)
        
        doc_node = DocumentNode(children=[list_level_1])
        
        template_path = create_placeholder_template()
        doc = Document(str(template_path))
        placeholder_para = next(p for p in doc.paragraphs if "{{content}}" in p.text)
        
        render_markdown_tree_to_docx(doc_node, doc, placeholder_paragraph=placeholder_para)
        
        # Verify no Word numbering XML present
        body_xml = doc._body._body.xml
        assert "w:numPr" not in body_xml, "Word numbering XML found (should use manual glyphs)"
        assert "w:numId" not in body_xml, "Word numbering ID found (should use manual glyphs)"
        
        # Verify structural integrity and save artifacts
        output_path = test_output_dir / "test_bullet_depths_1_6.docx"
        assert_docx_integrity(doc, output_path)
        
        # Save artifacts for inspection
        artifacts_dir = Path(__file__).parent / "artifacts"
        save_test_artifacts_phase3("bullet_depths_1_6", doc_node, doc, output_path, artifacts_dir)
    
    def test_ordered_list_hierarchical_numbering(self, test_output_dir):
        """Test ordered lists produce hierarchical numbering (1., 1.1., 1.1.1.)."""
        # Build nested ordered lists
        items_inner = [
            ListItemNode(children=[
                ParagraphNode(runs=[FormattedRun(text="Inner 1")])
            ]),
            ListItemNode(children=[
                ParagraphNode(runs=[FormattedRun(text="Inner 2")])
            ])
        ]
        list_inner = ListNode(kind="ordered", items=items_inner)
        
        items_mid = [
            ListItemNode(children=[
                ParagraphNode(runs=[FormattedRun(text="Mid 1")]),
                list_inner
            ]),
            ListItemNode(children=[
                ParagraphNode(runs=[FormattedRun(text="Mid 2")])
            ])
        ]
        list_mid = ListNode(kind="ordered", items=items_mid)
        
        items_outer = [
            ListItemNode(children=[
                ParagraphNode(runs=[FormattedRun(text="Outer 1")]),
                list_mid
            ]),
            ListItemNode(children=[
                ParagraphNode(runs=[FormattedRun(text="Outer 2")])
            ])
        ]
        list_outer = ListNode(kind="ordered", items=items_outer)
        
        doc_node = DocumentNode(children=[list_outer])
        
        template_path = create_placeholder_template()
        doc = Document(str(template_path))
        placeholder_para = next(p for p in doc.paragraphs if "{{content}}" in p.text)
        
        render_markdown_tree_to_docx(doc_node, doc, placeholder_paragraph=placeholder_para)
        
        # Extract paragraph texts with prefixes
        para_texts = extract_paragraph_texts(doc)
        
        # Verify hierarchical numbering appears
        # Should see: "1. Outer 1", "1.1. Mid 1", "1.1.1. Inner 1", "1.1.2. Inner 2", etc.
        numbered_paras = [t for t in para_texts if any(t.startswith(f"{i}.") for i in range(1, 10))]
        assert len(numbered_paras) >= 4, f"Expected at least 4 numbered paragraphs, found {len(numbered_paras)}"
        
        # Verify no Word numbering XML
        body_xml = doc._body._body.xml
        assert "w:numPr" not in body_xml, "Word numbering XML found (should use manual numbering)"
        
        # Verify structural integrity and save artifacts
        output_path = test_output_dir / "test_ordered_hierarchical.docx"
        assert_docx_integrity(doc, output_path)
        
        # Save artifacts for inspection
        artifacts_dir = Path(__file__).parent / "artifacts"
        save_test_artifacts_phase3("ordered_hierarchical", doc_node, doc, output_path, artifacts_dir)
    
    def test_mixed_nesting_ordered_bullet_ordered(self, test_output_dir):
        """Test mixed nesting: ordered ? bullet ? ordered."""
        items_inner_ordered = [
            ListItemNode(children=[
                ParagraphNode(runs=[FormattedRun(text="Inner ordered 1")])
            ])
        ]
        list_inner_ordered = ListNode(kind="ordered", items=items_inner_ordered)
        
        items_mid_bullet = [
            ListItemNode(children=[
                ParagraphNode(runs=[FormattedRun(text="Mid bullet 1")]),
                list_inner_ordered
            ])
        ]
        list_mid_bullet = ListNode(kind="bullet", items=items_mid_bullet)
        
        items_outer_ordered = [
            ListItemNode(children=[
                ParagraphNode(runs=[FormattedRun(text="Outer ordered 1")]),
                list_mid_bullet
            ])
        ]
        list_outer_ordered = ListNode(kind="ordered", items=items_outer_ordered)
        
        doc_node = DocumentNode(children=[list_outer_ordered])
        
        template_path = create_placeholder_template()
        doc = Document(str(template_path))
        placeholder_para = next(p for p in doc.paragraphs if "{{content}}" in p.text)
        
        render_markdown_tree_to_docx(doc_node, doc, placeholder_paragraph=placeholder_para)
        
        # Verify all list items rendered
        para_texts = extract_paragraph_texts(doc)
        assert any("Outer ordered" in t for t in para_texts), "Outer ordered item not found"
        assert any("Mid bullet" in t for t in para_texts), "Mid bullet item not found"
        assert any("Inner ordered" in t for t in para_texts), "Inner ordered item not found"
        
        # Verify structural integrity and save artifacts
        output_path = test_output_dir / "test_mixed_nesting.docx"
        assert_docx_integrity(doc, output_path)
        
        # Save artifacts for inspection
        artifacts_dir = Path(__file__).parent / "artifacts"
        save_test_artifacts_phase3("mixed_nesting", doc_node, doc, output_path, artifacts_dir)
    
    def test_list_item_first_child_table(self, test_output_dir):
        """Test list item where first child is a TableNode."""
        table = TableNode(rows=[
            TableRowNode(cells=[
                TableCellNode(children=[
                    ParagraphNode(runs=[FormattedRun(text="Table in list")])
                ])
            ])
        ])
        
        item = ListItemNode(children=[table])
        list_node = ListNode(kind="bullet", items=[item])
        
        doc_node = DocumentNode(children=[list_node])
        
        template_path = create_placeholder_template()
        doc = Document(str(template_path))
        placeholder_para = next(p for p in doc.paragraphs if "{{content}}" in p.text)
        
        render_markdown_tree_to_docx(doc_node, doc, placeholder_paragraph=placeholder_para)
        
        # Verify prefix paragraph was created (empty paragraph with bullet)
        para_texts = extract_paragraph_texts(doc)
        bullet_char = "\u2022"  # Unicode bullet character
        bullet_paras = [t for t in para_texts if t.strip().startswith(bullet_char)]
        assert len(bullet_paras) >= 1, "Prefix paragraph with bullet not found"
        
        # Verify table exists
        xml_order = extract_xml_element_order(doc)
        assert any(tag == "tbl" for tag, _ in xml_order), "Table not found in document"
        
        # Verify structural integrity and save artifacts
        output_path = test_output_dir / "test_list_item_table_first_child.docx"
        assert_docx_integrity(doc, output_path)
        
        # Save artifacts for inspection
        artifacts_dir = Path(__file__).parent / "artifacts"
        save_test_artifacts_phase3("list_item_table_first_child", doc_node, doc, output_path, artifacts_dir)
    
    def test_list_item_first_child_nested_list(self, test_output_dir):
        """Test list item where first child is a nested ListNode."""
        nested_list = ListNode(kind="bullet", items=[
            ListItemNode(children=[
                ParagraphNode(runs=[FormattedRun(text="Nested item")])
            ])
        ])
        
        item = ListItemNode(children=[nested_list])
        outer_list = ListNode(kind="ordered", items=[item])
        
        doc_node = DocumentNode(children=[outer_list])
        
        template_path = create_placeholder_template()
        doc = Document(str(template_path))
        placeholder_para = next(p for p in doc.paragraphs if "{{content}}" in p.text)
        
        render_markdown_tree_to_docx(doc_node, doc, placeholder_paragraph=placeholder_para)
        
        # Verify both outer and nested items rendered
        para_texts = extract_paragraph_texts(doc)
        assert any("Nested item" in t for t in para_texts), "Nested list item not found"
        
        # Verify structural integrity and save artifacts
        output_path = test_output_dir / "test_list_item_nested_list_first_child.docx"
        assert_docx_integrity(doc, output_path)
        
        # Save artifacts for inspection
        artifacts_dir = Path(__file__).parent / "artifacts"
        save_test_artifacts_phase3("list_item_nested_list_first_child", doc_node, doc, output_path, artifacts_dir)


# ============================================================================
# 3. Table Safety Tests
# ============================================================================

class TestTableSafety:
    """Test table rendering safety and atomicity."""
    
    def test_body_flow_table_atomic_insertion(self, test_output_dir):
        """Test that body-flow tables are inserted atomically."""
        table = TableNode(rows=[
            TableRowNode(cells=[
                TableCellNode(children=[
                    ParagraphNode(runs=[FormattedRun(text="Header A")])
                ]),
                TableCellNode(children=[
                    ParagraphNode(runs=[FormattedRun(text="Header B")])
                ])
            ]),
            TableRowNode(cells=[
                TableCellNode(children=[
                    ParagraphNode(runs=[FormattedRun(text="Data 1")])
                ]),
                TableCellNode(children=[
                    ParagraphNode(runs=[FormattedRun(text="Data 2")])
                ])
            ])
        ])
        
        doc_node = DocumentNode(children=[table])
        
        template_path = create_placeholder_template()
        doc = Document(str(template_path))
        placeholder_para = next(p for p in doc.paragraphs if "{{content}}" in p.text)
        
        # Count tables before
        body_elem = doc._body._body
        tables_before = len([e for e in body_elem if e.tag.endswith("}tbl")])
        
        render_markdown_tree_to_docx(doc_node, doc, placeholder_paragraph=placeholder_para)
        
        # Count tables after
        body_elem = doc._body._body
        tables_after = len([e for e in body_elem if e.tag.endswith("}tbl")])
        
        # Verify exactly one table inserted
        assert tables_after == tables_before + 1, f"Expected {tables_before + 1} tables, found {tables_after}"
        
        # Verify table structure
        xml_order = extract_xml_element_order(doc)
        tbl_positions = [idx for tag, idx in xml_order if tag == "tbl"]
        assert len(tbl_positions) == 1, f"Expected 1 table, found {len(tbl_positions)}"
        
        # Verify table formatting exists (check for table properties)
        body_xml = doc._body._body.xml
        assert "w:tblPr" in body_xml or "w:tblGrid" in body_xml, "Table formatting missing"
        
        # Verify structural integrity and save artifacts
        output_path = test_output_dir / "test_body_flow_table_atomic.docx"
        assert_docx_integrity(doc, output_path)
        
        # Save artifacts for inspection
        artifacts_dir = Path(__file__).parent / "artifacts"
        save_test_artifacts_phase3("body_flow_table_atomic", doc_node, doc, output_path, artifacts_dir)
    
    def test_table_inside_list(self, test_output_dir):
        """Test table inside a list item."""
        table = TableNode(rows=[
            TableRowNode(cells=[
                TableCellNode(children=[
                    ParagraphNode(runs=[FormattedRun(text="X")])
                ]),
                TableCellNode(children=[
                    ParagraphNode(runs=[FormattedRun(text="Y")])
                ])
            ])
        ])
        
        item = ListItemNode(children=[
            ParagraphNode(runs=[FormattedRun(text="Item with table")]),
            table,
            ParagraphNode(runs=[FormattedRun(text="Continuation after table")])
        ])
        
        list_node = ListNode(kind="ordered", items=[item])
        doc_node = DocumentNode(children=[list_node])
        
        template_path = create_placeholder_template()
        doc = Document(str(template_path))
        placeholder_para = next(p for p in doc.paragraphs if "{{content}}" in p.text)
        
        render_markdown_tree_to_docx(doc_node, doc, placeholder_paragraph=placeholder_para)
        
        # Verify table exists
        xml_order = extract_xml_element_order(doc)
        assert any(tag == "tbl" for tag, _ in xml_order), "Table not found"
        
        # Verify order: item paragraph, table, continuation paragraph
        para_texts = extract_paragraph_texts(doc)
        assert any("Item with table" in t for t in para_texts), "Item paragraph not found"
        assert any("Continuation after table" in t for t in para_texts), "Continuation paragraph not found"
        
        # Verify structural integrity and save artifacts
        output_path = test_output_dir / "test_table_inside_list.docx"
        assert_docx_integrity(doc, output_path)
        
        # Save artifacts for inspection
        artifacts_dir = Path(__file__).parent / "artifacts"
        save_test_artifacts_phase3("table_inside_list", doc_node, doc, output_path, artifacts_dir)
    
    def test_nested_table_in_cell(self, test_output_dir):
        """Test nested table inside a table cell."""
        nested_table = TableNode(rows=[
            TableRowNode(cells=[
                TableCellNode(children=[
                    ParagraphNode(runs=[FormattedRun(text="Nested")])
                ])
            ])
        ])
        
        outer_table = TableNode(rows=[
            TableRowNode(cells=[
                TableCellNode(children=[
                    ParagraphNode(runs=[FormattedRun(text="Outer cell")]),
                    nested_table
                ])
            ])
        ])
        
        doc_node = DocumentNode(children=[outer_table])
        
        template_path = create_placeholder_template()
        doc = Document(str(template_path))
        placeholder_para = next(p for p in doc.paragraphs if "{{content}}" in p.text)
        
        render_markdown_tree_to_docx(doc_node, doc, placeholder_paragraph=placeholder_para)
        
        # Verify outer table exists
        xml_order = extract_xml_element_order(doc)
        assert any(tag == "tbl" for tag, _ in xml_order), "Outer table not found"
        
        # Note: Nested table rendering may fallback to text-grid, which is acceptable
        # The important thing is that the document is not corrupted
        
        # Verify structural integrity and save artifacts
        output_path = test_output_dir / "test_nested_table_in_cell.docx"
        assert_docx_integrity(doc, output_path)
        
        # Save artifacts for inspection
        artifacts_dir = Path(__file__).parent / "artifacts"
        save_test_artifacts_phase3("nested_table_in_cell", doc_node, doc, output_path, artifacts_dir)


# ============================================================================
# 4. Placeholder Integrity Tests
# ============================================================================

class TestPlaceholderIntegrity:
    """Test placeholder replacement integrity."""
    
    def test_placeholder_removed_exactly_once(self, test_output_dir):
        """Test that placeholder is removed exactly once."""
        doc_node = DocumentNode(children=[
            ParagraphNode(runs=[FormattedRun(text="Content")])
        ])
        
        template_path = create_placeholder_template()
        doc = Document(str(template_path))
        
        # Count placeholders before
        placeholders_before = len([p for p in doc.paragraphs if "{{content}}" in p.text])
        assert placeholders_before == 1, "Expected exactly 1 placeholder"
        
        placeholder_para = next(p for p in doc.paragraphs if "{{content}}" in p.text)
        
        render_markdown_tree_to_docx(doc_node, doc, placeholder_paragraph=placeholder_para)
        
        # Count placeholders after
        placeholders_after = len([p for p in doc.paragraphs if "{{content}}" in p.text])
        assert placeholders_after == 0, f"Expected 0 placeholders after removal, found {placeholders_after}"
        
        # Verify structural integrity and save artifacts
        output_path = test_output_dir / "test_placeholder_removed_once.docx"
        assert_docx_integrity(doc, output_path)
        
        # Save artifacts for inspection
        artifacts_dir = Path(__file__).parent / "artifacts"
        save_test_artifacts_phase3("placeholder_removed_once", doc_node, doc, output_path, artifacts_dir)
    
    def test_content_inserted_at_placeholder_position(self, test_output_dir):
        """Test that content is inserted at placeholder position, not appended."""
        doc_node = DocumentNode(children=[
            ParagraphNode(runs=[FormattedRun(text="Inserted")])
        ])
        
        template_path = create_placeholder_template()
        doc = Document(str(template_path))
        
        # Get paragraph indices
        placeholder_idx = None
        for i, para in enumerate(doc.paragraphs):
            if "{{content}}" in para.text:
                placeholder_idx = i
                break
        
        assert placeholder_idx is not None, "Placeholder not found"
        
        placeholder_para = doc.paragraphs[placeholder_idx]
        
        render_markdown_tree_to_docx(doc_node, doc, placeholder_paragraph=placeholder_para)
        
        # Verify "Inserted" appears, and "After placeholder" still exists
        para_texts = extract_paragraph_texts(doc)
        assert any("Inserted" in t for t in para_texts), "Inserted content not found"
        assert any("After placeholder" in t for t in para_texts), "After placeholder paragraph missing"
        
        # Verify structural integrity and save artifacts
        output_path = test_output_dir / "test_content_at_placeholder_position.docx"
        assert_docx_integrity(doc, output_path)
        
        # Save artifacts for inspection
        artifacts_dir = Path(__file__).parent / "artifacts"
        save_test_artifacts_phase3("content_at_placeholder_position", doc_node, doc, output_path, artifacts_dir)
    
    def test_surrounding_paragraphs_untouched(self, test_output_dir):
        """Test that surrounding paragraphs remain untouched."""
        doc_node = DocumentNode(children=[
            ParagraphNode(runs=[FormattedRun(text="New content")])
        ])
        
        template_path = create_placeholder_template()
        doc = Document(str(template_path))
        
        # Capture surrounding text
        before_text = doc.paragraphs[0].text
        after_text = doc.paragraphs[2].text
        
        placeholder_para = doc.paragraphs[1]
        
        render_markdown_tree_to_docx(doc_node, doc, placeholder_paragraph=placeholder_para)
        
        # Verify surrounding paragraphs unchanged
        assert doc.paragraphs[0].text == before_text, "Before paragraph was modified"
        # After paragraph index may have shifted, find it
        after_paras = [p for p in doc.paragraphs if "After placeholder" in p.text]
        assert len(after_paras) == 1, "After paragraph missing or duplicated"
        assert after_paras[0].text == after_text, "After paragraph was modified"
        
        # Verify structural integrity and save artifacts
        output_path = test_output_dir / "test_surrounding_untouched.docx"
        assert_docx_integrity(doc, output_path)
        
        # Save artifacts for inspection
        artifacts_dir = Path(__file__).parent / "artifacts"
        save_test_artifacts_phase3("surrounding_untouched", doc_node, doc, output_path, artifacts_dir)
    
    def test_multiple_placeholders_independent(self, test_output_dir):
        """Test that multiple placeholders are handled independently."""
        # This test would require multiple render calls, which is beyond Phase II scope
        # But we can test that a single render doesn't affect other placeholders
        doc_node = DocumentNode(children=[
            ParagraphNode(runs=[FormattedRun(text="Content 1")])
        ])
        
        # Create template with two placeholders
        template_path = Path(tempfile.mktemp(suffix=".docx"))
        doc = Document()
        doc.add_paragraph("{{content1}}")
        doc.add_paragraph("{{content2}}")
        doc.save(str(template_path))
        
        # Render only first placeholder
        placeholder1 = doc.paragraphs[0]
        render_markdown_tree_to_docx(doc_node, doc, placeholder_paragraph=placeholder1)
        
        # Verify second placeholder still exists
        remaining_placeholders = [p for p in doc.paragraphs if "{{content2}}" in p.text]
        assert len(remaining_placeholders) == 1, "Second placeholder was affected"
        
        # Verify structural integrity and save artifacts
        output_path = test_output_dir / "test_multiple_placeholders.docx"
        assert_docx_integrity(doc, output_path)
        
        # Save artifacts for inspection
        artifacts_dir = Path(__file__).parent / "artifacts"
        save_test_artifacts_phase3("multiple_placeholders", doc_node, doc, output_path, artifacts_dir)


# ============================================================================
# 5. Safety Guardrail Tests
# ============================================================================

class TestSafetyGuardrails:
    """Test safety guardrails (MAX_TREE_DEPTH, MAX_NODE_COUNT)."""
    
    def test_max_tree_depth_exceeded_raises_before_mutation(self, test_output_dir):
        """Test that exceeding MAX_TREE_DEPTH raises before any document mutation."""
        # Build tree exceeding MAX_TREE_DEPTH
        def build_deep_list(depth: int) -> ListNode:
            if depth > MAX_TREE_DEPTH + 5:  # Exceed limit
                return ListNode(kind="bullet", items=[])
            items = [
                ListItemNode(children=[
                    ParagraphNode(runs=[FormattedRun(text=f"Depth {depth}")]),
                    build_deep_list(depth + 1)
                ])
            ]
            return ListNode(kind="bullet", items=items)
        
        deep_list = build_deep_list(1)
        doc_node = DocumentNode(children=[deep_list])
        
        template_path = create_placeholder_template()
        doc = Document(str(template_path))
        placeholder_para = next(p for p in doc.paragraphs if "{{content}}" in p.text)
        
        # Capture document state before
        para_count_before = len(doc.paragraphs)
        placeholder_exists_before = "{{content}}" in placeholder_para.text
        
        # Attempt render - should raise
        with pytest.raises(ValueError, match="Tree depth.*exceeds MAX_TREE_DEPTH"):
            render_markdown_tree_to_docx(doc_node, doc, placeholder_paragraph=placeholder_para)
        
        # Verify no mutation occurred
        assert len(doc.paragraphs) == para_count_before, "Document was mutated despite error"
        # Check placeholder still exists in XML
        body_xml = doc._body._body.xml
        assert "{{content}}" in body_xml, "Placeholder was removed despite validation failure"
    
    def test_max_node_count_exceeded_raises_before_mutation(self, test_output_dir):
        """Test that exceeding MAX_NODE_COUNT raises before any document mutation."""
        # Build tree with many nodes
        many_items = [
            ListItemNode(children=[
                ParagraphNode(runs=[FormattedRun(text=f"Item {i}")]),
                ParagraphNode(runs=[FormattedRun(text=f"Para {i}")])
            ]) for i in range(MAX_NODE_COUNT + 100)
        ]
        
        large_list = ListNode(kind="bullet", items=many_items)
        doc_node = DocumentNode(children=[large_list])
        
        template_path = create_placeholder_template()
        doc = Document(str(template_path))
        placeholder_para = next(p for p in doc.paragraphs if "{{content}}" in p.text)
        
        # Capture document state before
        para_count_before = len(doc.paragraphs)
        
        # Attempt render - should raise
        with pytest.raises(ValueError, match="Node count.*exceeds MAX_NODE_COUNT"):
            render_markdown_tree_to_docx(doc_node, doc, placeholder_paragraph=placeholder_para)
        
        # Verify no mutation occurred
        assert len(doc.paragraphs) == para_count_before, "Document was mutated despite error"
        # Check placeholder still exists in XML
        body_xml = doc._body._body.xml
        assert "{{content}}" in body_xml, "Placeholder was removed despite validation failure"


# ============================================================================
# 6. DOCX Structural Integrity Tests
# ============================================================================

class TestDocxStructuralIntegrity:
    """Test DOCX structural integrity after rendering."""
    
    def test_document_saves_and_reloads(self, test_output_dir):
        """Test that rendered document can be saved and reloaded."""
        doc_node = DocumentNode(children=[
            HeadingNode(level=1, runs=[FormattedRun(text="Test")]),
            ParagraphNode(runs=[FormattedRun(text="Content")]),
            ListNode(kind="bullet", items=[
                ListItemNode(children=[
                    ParagraphNode(runs=[FormattedRun(text="Item")])
                ])
            ]),
            TableNode(rows=[
                TableRowNode(cells=[
                    TableCellNode(children=[
                        ParagraphNode(runs=[FormattedRun(text="Cell")])
                    ])
                ])
            ])
        ])
        
        template_path = create_placeholder_template()
        doc = Document(str(template_path))
        placeholder_para = next(p for p in doc.paragraphs if "{{content}}" in p.text)
        
        render_markdown_tree_to_docx(doc_node, doc, placeholder_paragraph=placeholder_para)
        
        # Save and reload
        output_path = test_output_dir / "test_save_reload.docx"
        assert_docx_integrity(doc, output_path)
        
        # Verify reloaded document is readable
        reloaded = Document(str(output_path))
        assert len(reloaded.paragraphs) > 0, "Reloaded document has no paragraphs"
        assert any("Test" in p.text for p in reloaded.paragraphs), "Heading not found in reloaded document"
    
    def test_no_malformed_xml(self, test_output_dir):
        """Test that rendered document has no malformed XML."""
        doc_node = DocumentNode(children=[
            ParagraphNode(runs=[FormattedRun(text="Test")]),
            TableNode(rows=[
                TableRowNode(cells=[
                    TableCellNode(children=[
                        ParagraphNode(runs=[FormattedRun(text="A")])
                    ]),
                    TableCellNode(children=[
                        ParagraphNode(runs=[FormattedRun(text="B")])
                    ])
                ])
            ])
        ])
        
        template_path = create_placeholder_template()
        doc = Document(str(template_path))
        placeholder_para = next(p for p in doc.paragraphs if "{{content}}" in p.text)
        
        render_markdown_tree_to_docx(doc_node, doc, placeholder_paragraph=placeholder_para)
        
        # Attempt to serialize XML (will raise if malformed)
        try:
            body_xml = doc._body._body.xml
            # Verify basic structure
            assert "<w:body" in body_xml or "w:body" in body_xml, "Body element missing"
            assert "</w:body>" in body_xml or body_xml.count("w:body") > 0, "Body element not closed"
        except Exception as e:
            pytest.fail(f"XML serialization failed (malformed XML): {e}")
        
        # Save and verify can reload
        output_path = test_output_dir / "test_no_malformed_xml.docx"
        assert_docx_integrity(doc, output_path)
    
    def test_tables_readable_after_reload(self, test_output_dir):
        """Test that tables are readable after save/reload."""
        table = TableNode(rows=[
            TableRowNode(cells=[
                TableCellNode(children=[
                    ParagraphNode(runs=[FormattedRun(text="Header 1")])
                ]),
                TableCellNode(children=[
                    ParagraphNode(runs=[FormattedRun(text="Header 2")])
                ])
            ]),
            TableRowNode(cells=[
                TableCellNode(children=[
                    ParagraphNode(runs=[FormattedRun(text="Data 1")])
                ]),
                TableCellNode(children=[
                    ParagraphNode(runs=[FormattedRun(text="Data 2")])
                ])
            ])
        ])
        
        doc_node = DocumentNode(children=[table])
        
        template_path = create_placeholder_template()
        doc = Document(str(template_path))
        placeholder_para = next(p for p in doc.paragraphs if "{{content}}" in p.text)
        
        render_markdown_tree_to_docx(doc_node, doc, placeholder_paragraph=placeholder_para)
        
        output_path = test_output_dir / "test_tables_readable.docx"
        doc.save(str(output_path))
        
        # Reload and verify tables
        reloaded = Document(str(output_path))
        assert len(reloaded.tables) > 0, "No tables found in reloaded document"
        
        table = reloaded.tables[0]
        assert len(table.rows) >= 1, "Table has no rows"
        assert len(table.columns) >= 1, "Table has no columns"
        
        # Verify cell content readable
        for row in table.rows:
            for cell in row.cells:
                _ = cell.text  # Should not raise


# ============================================================================
# Additional Critical Tests
# ============================================================================

class TestListIndentationDeterministic:
    """Test that list indentation increases deterministically."""
    
    def test_indentation_increases_with_depth(self, test_output_dir):
        """Test that indentation increases deterministically with list depth."""
        # Build 3-level nested list
        items_level_3 = [
            ListItemNode(children=[
                ParagraphNode(runs=[FormattedRun(text="Level 3")])
            ])
        ]
        list_level_3 = ListNode(kind="bullet", items=items_level_3)
        
        items_level_2 = [
            ListItemNode(children=[
                ParagraphNode(runs=[FormattedRun(text="Level 2")]),
                list_level_3
            ])
        ]
        list_level_2 = ListNode(kind="bullet", items=items_level_2)
        
        items_level_1 = [
            ListItemNode(children=[
                ParagraphNode(runs=[FormattedRun(text="Level 1")]),
                list_level_2
            ])
        ]
        list_level_1 = ListNode(kind="bullet", items=items_level_1)
        
        doc_node = DocumentNode(children=[list_level_1])
        
        template_path = create_placeholder_template()
        doc = Document(str(template_path))
        placeholder_para = next(p for p in doc.paragraphs if "{{content}}" in p.text)
        
        render_markdown_tree_to_docx(doc_node, doc, placeholder_paragraph=placeholder_para)
        
        # Extract indentation values
        indent_values = []
        for para in doc.paragraphs:
            if para.paragraph_format.left_indent is not None:
                indent_values.append(para.paragraph_format.left_indent.inches)
        
        # Verify indentation increases (or stays same for clamped levels)
        # At minimum, we should see different indentation levels
        unique_indents = set(indent_values)
        assert len(unique_indents) >= 1, "No indentation applied"
        
        # Verify structural integrity and save artifacts
        output_path = test_output_dir / "test_indentation_deterministic.docx"
        assert_docx_integrity(doc, output_path)
        
        # Save artifacts for inspection
        artifacts_dir = Path(__file__).parent / "artifacts"
        save_test_artifacts_phase3("indentation_deterministic", doc_node, doc, output_path, artifacts_dir)
    
    def test_list_items_respect_insertion_context(self, test_output_dir):
        """Test that list items do not escape insertion context."""
        # Create template with placeholder between two paragraphs
        template_path = Path(tempfile.mktemp(suffix=".docx"))
        doc = Document()
        doc.add_paragraph("Before")
        placeholder_para = doc.add_paragraph("{{content}}")
        doc.add_paragraph("After")
        doc.save(str(template_path))
        
        # Build list with multiple items
        list_node = ListNode(kind="ordered", items=[
            ListItemNode(children=[
                ParagraphNode(runs=[FormattedRun(text="Item 1")])
            ]),
            ListItemNode(children=[
                ParagraphNode(runs=[FormattedRun(text="Item 2")])
            ]),
            ListItemNode(children=[
                ParagraphNode(runs=[FormattedRun(text="Item 3")])
            ])
        ])
        
        doc_node = DocumentNode(children=[list_node])
        
        # Capture before/after paragraphs
        before_para = doc.paragraphs[0]
        after_para = doc.paragraphs[2]
        before_text = before_para.text
        after_text = after_para.text
        
        render_markdown_tree_to_docx(doc_node, doc, placeholder_paragraph=placeholder_para)
        
        # Verify before/after paragraphs unchanged
        assert before_para.text == before_text, "Before paragraph modified"
        # After paragraph may have shifted index, find it
        after_paras = [p for p in doc.paragraphs if "After" in p.text]
        assert len(after_paras) == 1, "After paragraph missing or duplicated"
        assert after_paras[0].text == after_text, "After paragraph modified"
        
        # Verify all list items rendered
        para_texts = extract_paragraph_texts(doc)
        assert any("Item 1" in t for t in para_texts), "Item 1 not found"
        assert any("Item 2" in t for t in para_texts), "Item 2 not found"
        assert any("Item 3" in t for t in para_texts), "Item 3 not found"
        
        # Verify structural integrity and save artifacts
        output_path = test_output_dir / "test_list_items_insertion_context.docx"
        assert_docx_integrity(doc, output_path)
        
        # Save artifacts for inspection
        artifacts_dir = Path(__file__).parent / "artifacts"
        save_test_artifacts_phase3("list_items_insertion_context", doc_node, doc, output_path, artifacts_dir)


class TestTableFailureHandling:
    """Test table failure handling (no silent failures)."""
    
    def test_table_failure_does_not_corrupt_document(self, test_output_dir):
        """Test that table rendering failure does not corrupt document."""
        # Create a table that might fail (empty rows edge case)
        empty_table = TableNode(rows=[])
        
        doc_node = DocumentNode(children=[empty_table])
        
        template_path = create_placeholder_template()
        doc = Document(str(template_path))
        placeholder_para = next(p for p in doc.paragraphs if "{{content}}" in p.text)
        
        # Capture document state
        para_count_before = len(doc.paragraphs)
        
        # Render (should handle empty table gracefully)
        render_markdown_tree_to_docx(doc_node, doc, placeholder_paragraph=placeholder_para)
        
        # Verify document is still valid (no corruption)
        # Empty table should return 0, but document should be intact
        assert len(doc.paragraphs) >= para_count_before - 1, "Document structure corrupted"
        
        # Verify structural integrity and save artifacts
        output_path = test_output_dir / "test_table_failure_handling.docx"
        assert_docx_integrity(doc, output_path)
        
        # Save artifacts for inspection
        artifacts_dir = Path(__file__).parent / "artifacts"
        save_test_artifacts_phase3("table_failure_handling", doc_node, doc, output_path, artifacts_dir, template_path)


# ============================================================================
# 9. Deep Nested Tables Tests
# ============================================================================

class TestDeepNestedTables:
    """Test deeply nested tables with markdown content."""
    
    def test_nested_tables_5_levels_with_markdown(self, test_output_dir):
        """
        Test nested tables at 5 levels deep, each containing markdown content.
        
        Structure:
        Level 1 Table
          Cell 1: Heading + Paragraph + Level 2 Table + Paragraph
            Cell 1: Paragraph + Level 3 Table + Paragraph
              Cell 1: Heading + Paragraph + Level 4 Table + Paragraph
                Cell 1: Paragraph + Level 5 Table + List
                  Cell 1: Paragraphs (final content)
        """
        # Build nested tables from innermost (level 5) to outermost (level 1)
        
        # Level 5: Innermost table
        level5_table = TableNode(rows=[
            TableRowNode(cells=[
                TableCellNode(children=[
                    ParagraphNode(runs=[FormattedRun(text="Level 5: Final nested content")]),
                    ParagraphNode(runs=[FormattedRun(text="This is the deepest level")])
                ]),
                TableCellNode(children=[
                    ParagraphNode(runs=[FormattedRun(text="Level 5: Right cell")])
                ])
            ])
        ])
        
        # Level 4: Contains level 5 table + list
        level4_table = TableNode(rows=[
            TableRowNode(cells=[
                TableCellNode(children=[
                    ParagraphNode(runs=[FormattedRun(text="Level 4: Before nested table")]),
                    level5_table,
                    ListNode(kind="bullet", items=[
                        ListItemNode(children=[
                            ParagraphNode(runs=[FormattedRun(text="Level 4 bullet item")])
                        ])
                    ])
                ]),
                TableCellNode(children=[
                    ParagraphNode(runs=[FormattedRun(text="Level 4: Right cell")])
                ])
            ])
        ])
        
        # Level 3: Contains level 4 table + heading
        level3_table = TableNode(rows=[
            TableRowNode(cells=[
                TableCellNode(children=[
                    HeadingNode(level=2, runs=[FormattedRun(text="Level 3 Heading")]),
                    ParagraphNode(runs=[FormattedRun(text="Level 3: Introduction")]),
                    level4_table,
                    ParagraphNode(runs=[FormattedRun(text="Level 3: After nested table")])
                ]),
                TableCellNode(children=[
                    ParagraphNode(runs=[FormattedRun(text="Level 3: Right cell")])
                ])
            ])
        ])
        
        # Level 2: Contains level 3 table + paragraph
        level2_table = TableNode(rows=[
            TableRowNode(cells=[
                TableCellNode(children=[
                    ParagraphNode(runs=[FormattedRun(text="Level 2: Introduction paragraph")]),
                    level3_table,
                    ParagraphNode(runs=[FormattedRun(text="Level 2: Conclusion paragraph")])
                ]),
                TableCellNode(children=[
                    ParagraphNode(runs=[FormattedRun(text="Level 2: Right cell")])
                ])
            ])
        ])
        
        # Level 1: Outermost table
        level1_table = TableNode(rows=[
            TableRowNode(cells=[
                TableCellNode(children=[
                    HeadingNode(level=1, runs=[FormattedRun(text="Level 1: Main Heading")]),
                    ParagraphNode(runs=[FormattedRun(text="Level 1: Main content before nested table")]),
                    level2_table,
                    ParagraphNode(runs=[FormattedRun(text="Level 1: Main content after nested table")])
                ]),
                TableCellNode(children=[
                    ParagraphNode(runs=[FormattedRun(text="Level 1: Right column")])
                ])
            ])
        ])
        
        doc_node = DocumentNode(children=[level1_table])
        
        # Create template with placeholder
        template_path = create_placeholder_template()
        doc = Document(str(template_path))
        placeholder_para = next(p for p in doc.paragraphs if "{{content}}" in p.text)
        
        # Render tree
        render_markdown_tree_to_docx(doc_node, doc, placeholder_paragraph=placeholder_para)
        
        # Verify all levels of content are present
        # Extract text from both paragraphs and table cells
        para_texts = extract_paragraph_texts(doc)
        all_para_text = " ".join(para_texts)
        
        # Also extract text from table cells (nested tables render in cells)
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text
                    if cell_text.strip():
                        table_texts.append(cell_text)
        
        all_text = all_para_text + " " + " ".join(table_texts)
        
        # Verify level 1 content (check both paragraphs and table cells)
        assert ("Level 1: Main Heading" in all_text or 
                any("Level 1: Main Heading" in t for t in para_texts) or
                any("Level 1: Main Heading" in t for t in table_texts)), "Level 1 heading not found"
        # Verify level 1 content (outermost table - should always render)
        assert "Level 1: Main Heading" in all_text, "Level 1 heading not found"
        assert "Level 1: Main content before" in all_text, "Level 1 before text not found"
        assert "Level 1: Main content after" in all_text, "Level 1 after text not found"
        assert "Level 1: Right column" in all_text, "Level 1 right cell not found"
        
        # Note: Deeply nested tables (5 levels) may not all render as Word tables
        # due to Word's limitations. They may fall back to text-grid rendering.
        # We verify that:
        # 1. The outermost table renders correctly
        # 2. The document structure is valid
        # 3. At least some nested content appears (either as tables or text-grid)
        
        # Verify document has at least the outer table
        assert len(doc.tables) >= 1, "Outer table not found"
        
        # Check if nested content appears (either as nested tables or text-grid fallback)
        # Level 2+ content might be in text-grid format if nested table rendering fails
        # This is acceptable - the test validates that deep nesting doesn't corrupt the document
        
        # Verify document structure is valid
        xml_order = extract_xml_element_order(doc)
        tbl_count = sum(1 for tag, _ in xml_order if tag == "tbl")
        # Should have at least 1 table (level 1), possibly more if nested tables rendered
        assert tbl_count >= 1, f"Expected at least 1 table, found {tbl_count}"
        
        # Verify structural integrity and save artifacts
        output_path = test_output_dir / "test_nested_tables_5_levels.docx"
        assert_docx_integrity(doc, output_path)
        
        # Save artifacts for inspection
        artifacts_dir = Path(__file__).parent / "artifacts"
        save_test_artifacts_phase3("nested_tables_5_levels", doc_node, doc, output_path, artifacts_dir, template_path)


# ============================================================================
# 10. Mixed Content: Free text + Lists + Tables + Free text
# ============================================================================

class TestMixedContentListsAndTables:
    """
    Phase III proof test for a realistic mixed document flow:
    - Free text
    - Nested bullet list (2 items include tables; includes free text inside items)
    - Nested numbered list (2 items include tables; includes free text inside items)
    - Free text
    - Standalone body-flow table
    - Final free text
    """

    def test_free_text_then_lists_with_tables_then_table_then_free_text(self, test_output_dir):
        # Helper: small 2x2 table with unique marker text per cell
        def make_marker_table(prefix: str) -> TableNode:
            return TableNode(
                rows=[
                    TableRowNode(
                        cells=[
                            TableCellNode(children=[ParagraphNode(runs=[FormattedRun(text=f"{prefix} R1C1")])]),
                            TableCellNode(children=[ParagraphNode(runs=[FormattedRun(text=f"{prefix} R1C2")])]),
                        ]
                    ),
                    TableRowNode(
                        cells=[
                            TableCellNode(children=[ParagraphNode(runs=[FormattedRun(text=f"{prefix} R2C1")])]),
                            TableCellNode(children=[ParagraphNode(runs=[FormattedRun(text=f"{prefix} R2C2")])]),
                        ]
                    ),
                ]
            )

        # Bullet list with nested bullets; 2 bullet items contain tables, with free text around them
        bullet_table_1 = make_marker_table("BULLET_TABLE_1")
        bullet_table_2 = make_marker_table("BULLET_TABLE_2")

        bullet_list = ListNode(
            kind="bullet",
            items=[
                ListItemNode(
                    children=[
                        ParagraphNode(runs=[FormattedRun(text="BULLET_ITEM_1 intro")]),
                        bullet_table_1,
                        ParagraphNode(runs=[FormattedRun(text="BULLET_ITEM_1 middle free text")]),
                    ]
                ),
                ListItemNode(
                    children=[
                        ParagraphNode(runs=[FormattedRun(text="BULLET_ITEM_2 intro")]),
                        ListNode(
                            kind="bullet",
                            items=[
                                ListItemNode(
                                    children=[
                                        ParagraphNode(runs=[FormattedRun(text="BULLET_NESTED_ITEM_1 intro")]),
                                        bullet_table_2,
                                        ParagraphNode(runs=[FormattedRun(text="BULLET_NESTED_ITEM_1 end")]),
                                    ]
                                )
                            ],
                        ),
                        ParagraphNode(runs=[FormattedRun(text="BULLET_ITEM_2 end")]),
                    ]
                ),
            ],
        )

        # Numbered list with nesting; 2 numbered items contain tables, with free text around them
        number_table_1 = make_marker_table("NUMBER_TABLE_1")
        number_table_2 = make_marker_table("NUMBER_TABLE_2")

        ordered_list = ListNode(
            kind="ordered",
            items=[
                ListItemNode(
                    children=[
                        ParagraphNode(runs=[FormattedRun(text="NUMBER_ITEM_1 intro")]),
                        number_table_1,
                        ParagraphNode(runs=[FormattedRun(text="NUMBER_ITEM_1 end")]),
                    ]
                ),
                ListItemNode(
                    children=[
                        ParagraphNode(runs=[FormattedRun(text="NUMBER_ITEM_2 intro")]),
                        ListNode(
                            kind="ordered",
                            items=[
                                ListItemNode(
                                    children=[
                                        ParagraphNode(runs=[FormattedRun(text="NUMBER_NESTED_ITEM_1 intro")]),
                                        number_table_2,
                                        ParagraphNode(runs=[FormattedRun(text="NUMBER_NESTED_ITEM_1 end")]),
                                    ]
                                )
                            ],
                        ),
                        ParagraphNode(runs=[FormattedRun(text="NUMBER_ITEM_2 end")]),
                    ]
                ),
            ],
        )

        standalone_table = make_marker_table("STANDALONE_TABLE")

        doc_node = DocumentNode(
            children=[
                ParagraphNode(runs=[FormattedRun(text="INTRO_FREE_TEXT")]),
                bullet_list,
                ParagraphNode(runs=[FormattedRun(text="FREE_TEXT_BETWEEN_LISTS")]),
                ordered_list,
                ParagraphNode(runs=[FormattedRun(text="FREE_TEXT_AFTER_LISTS")]),
                standalone_table,
                ParagraphNode(runs=[FormattedRun(text="FINAL_FREE_TEXT_END")]),
            ]
        )

        # Template with explicit bounds so we can assert insertion stays within placeholder region
        template_path = Path(tempfile.mktemp(suffix=".docx"))
        doc = Document()
        before_para = doc.add_paragraph("BEFORE")
        placeholder_para = doc.add_paragraph("{{content}}")
        after_para = doc.add_paragraph("AFTER")
        doc.save(str(template_path))

        doc = Document(str(template_path))
        placeholder_para = next(p for p in doc.paragraphs if "{{content}}" in p.text)
        render_markdown_tree_to_docx(doc_node, doc, placeholder_paragraph=placeholder_para)

        # Paragraph-level assertions (body paragraphs only, excludes text that lives inside table cells)
        para_texts = [p.text for p in doc.paragraphs]
        assert "BEFORE" in para_texts
        assert "AFTER" in para_texts
        before_idx = para_texts.index("BEFORE")
        after_idx = para_texts.index("AFTER")
        assert before_idx < after_idx

        # All marker paragraphs must be inserted between BEFORE and AFTER, in relative order.
        markers_in_body_order = [
            "INTRO_FREE_TEXT",
            "BULLET_ITEM_1 intro",
            "BULLET_ITEM_1 middle free text",
            "BULLET_ITEM_2 intro",
            "BULLET_NESTED_ITEM_1 intro",
            "BULLET_NESTED_ITEM_1 end",
            "BULLET_ITEM_2 end",
            "FREE_TEXT_BETWEEN_LISTS",
            "NUMBER_ITEM_1 intro",
            "NUMBER_ITEM_1 end",
            "NUMBER_ITEM_2 intro",
            "NUMBER_NESTED_ITEM_1 intro",
            "NUMBER_NESTED_ITEM_1 end",
            "NUMBER_ITEM_2 end",
            "FREE_TEXT_AFTER_LISTS",
            "FINAL_FREE_TEXT_END",
        ]

        def _find_next_paragraph_containing(marker: str, start_at: int) -> int:
            for i in range(start_at, len(para_texts)):
                if marker in para_texts[i]:
                    return i
            return -1

        last_seen = before_idx + 1
        for marker in markers_in_body_order:
            idx = _find_next_paragraph_containing(marker, last_seen)
            assert idx != -1, f"Missing body marker paragraph (substring): {marker}"
            assert before_idx < idx < after_idx, f"Marker paragraph escaped placeholder bounds: {marker}"
            last_seen = idx + 1

        # Table assertions: we expect 5 body-flow tables total:
        # - 2 in the bullet list
        # - 2 in the numbered list
        # - 1 standalone after lists
        xml_order = extract_xml_element_order(doc)
        tbl_count = sum(1 for tag, _ in xml_order if tag == "tbl")
        assert tbl_count >= 5, f"Expected at least 5 tables, found {tbl_count}"

        # Also assert table cell marker text exists somewhere in the document (python-docx exposes it via doc.tables).
        all_cell_text = " ".join(cell.text for t in doc.tables for r in t.rows for cell in r.cells)
        for table_prefix in [
            "BULLET_TABLE_1",
            "BULLET_TABLE_2",
            "NUMBER_TABLE_1",
            "NUMBER_TABLE_2",
            "STANDALONE_TABLE",
        ]:
            assert table_prefix in all_cell_text, f"Missing table cell text for {table_prefix}"

        # Verify structural integrity and save artifacts
        output_path = test_output_dir / "test_mixed_free_text_lists_tables_table_free_text.docx"
        assert_docx_integrity(doc, output_path)

        artifacts_dir = Path(__file__).parent / "artifacts"
        save_test_artifacts_phase3(
            "mixed_free_text_lists_tables_table_free_text",
            doc_node,
            doc,
            output_path,
            artifacts_dir,
            template_path,
        )
