"""
Test suite for scalar placeholder replacement across all locations.

Enforces:
- Invariant 1: Text Fidelity - all scalar values appear in output
- Invariant 2: Determinism - same input produces same output
- Invariant 3: Observability - all artifacts written to disk
"""
import pytest
from pathlib import Path
from docx import Document

from docx_template_export.models.export_models import WordExportRequest
from tests.test_artifacts_helpers import (
    run_export_and_collect_artifacts,
    assert_text_fidelity,
    assert_no_placeholders_remain,
    extract_all_text_from_docx,
)


@pytest.fixture
def artifacts_dir(test_output_dir):
    """Return artifacts directory for this test suite."""
    artifacts_base = test_output_dir.parent / "artifacts"
    artifacts_base.mkdir(exist_ok=True)
    return artifacts_base


@pytest.fixture
def scalar_template_body(test_fixtures_dir):
    """Create template with scalar placeholders in BODY FLOW."""
    template_path = test_fixtures_dir / "scalar_template_body.docx"
    doc = Document()
    doc.add_paragraph("Document ID: {{document_id}}")
    doc.add_paragraph("Title: {{title}}")
    doc.add_paragraph("Author: {{author}}")
    doc.add_paragraph("Version: {{version}}")
    template_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(template_path))
    return template_path


@pytest.fixture
def scalar_template_table(test_fixtures_dir):
    """Create template with scalar placeholders in table cells."""
    template_path = test_fixtures_dir / "scalar_template_table.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].paragraphs[0].text = "ID"
    table.rows[0].cells[1].paragraphs[0].text = "{{document_id}}"
    table.rows[1].cells[0].paragraphs[0].text = "Title"
    table.rows[1].cells[1].paragraphs[0].text = "{{title}}"
    template_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(template_path))
    return template_path


@pytest.fixture
def scalar_template_header_footer(test_fixtures_dir):
    """Create template with scalar placeholders in header and footer."""
    template_path = test_fixtures_dir / "scalar_template_header_footer.docx"
    doc = Document()
    
    # Header
    header = doc.sections[0].header
    header.paragraphs[0].text = "{{document_id}} - {{title}}"
    
    # Body
    doc.add_paragraph("Body content here")
    
    # Footer
    footer = doc.sections[0].footer
    footer.paragraphs[0].text = "Page {{page_number}}"
    
    template_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(template_path))
    return template_path


def test_scalar_replacement_body_flow(artifacts_dir, scalar_template_body):
    """Test scalar replacement in BODY FLOW paragraphs."""
    request = WordExportRequest(
        scalar_fields={
            "document_id": "DOC-12345",
            "title": "Test Document Title",
            "author": "John Doe",
            "version": "1.0.0",
        },
        block_fields={},
    )
    
    artifact_dir, markdown, extracted_text = run_export_and_collect_artifacts(
        test_name="scalar_replacement_body_flow",
        template_path=scalar_template_body,
        request=request,
        artifacts_dir=artifacts_dir,
        subdirectory="scalar_replacement",
    )
    
    # Assert text fidelity
    assert_text_fidelity(
        input_texts=[
            "DOC-12345",
            "Test Document Title",
            "John Doe",
            "1.0.0",
        ],
        extracted_text=extracted_text,
        test_name="scalar_replacement_body_flow",
    )
    
    # Assert no placeholders remain
    assert_no_placeholders_remain(
        extracted_text=extracted_text,
        scalar_placeholders=["document_id", "title", "author", "version"],
        block_placeholders=[],
        test_name="scalar_replacement_body_flow",
    )
    
    # Assert artifacts exist
    assert (artifact_dir / "input.json").exists()
    assert (artifact_dir / "output.docx").exists()


def test_scalar_replacement_table_cells(artifacts_dir, scalar_template_table):
    """Test scalar replacement in table cells."""
    request = WordExportRequest(
        scalar_fields={
            "document_id": "TBL-67890",
            "title": "Table Test Document",
        },
        block_fields={},
    )
    
    artifact_dir, markdown, extracted_text = run_export_and_collect_artifacts(
        test_name="scalar_replacement_table_cells",
        template_path=scalar_template_table,
        request=request,
        artifacts_dir=artifacts_dir,
        subdirectory="scalar_replacement",
    )
    
    # Assert text fidelity
    assert_text_fidelity(
        input_texts=[
            "TBL-67890",
            "Table Test Document",
        ],
        extracted_text=extracted_text,
        test_name="scalar_replacement_table_cells",
    )
    
    # Assert no placeholders remain
    assert_no_placeholders_remain(
        extracted_text=extracted_text,
        scalar_placeholders=["document_id", "title"],
        block_placeholders=[],
        test_name="scalar_replacement_table_cells",
    )


def test_scalar_replacement_header_footer(artifacts_dir, scalar_template_header_footer):
    """Test scalar replacement in headers and footers."""
    request = WordExportRequest(
        scalar_fields={
            "document_id": "HDR-11111",
            "title": "Header Footer Test",
            "page_number": "1",
        },
        block_fields={},
    )
    
    artifact_dir, markdown, extracted_text = run_export_and_collect_artifacts(
        test_name="scalar_replacement_header_footer",
        template_path=scalar_template_header_footer,
        request=request,
        artifacts_dir=artifacts_dir,
        subdirectory="scalar_replacement",
    )
    
    # Assert text fidelity
    assert_text_fidelity(
        input_texts=[
            "HDR-11111",
            "Header Footer Test",
            "1",
        ],
        extracted_text=extracted_text,
        test_name="scalar_replacement_header_footer",
    )
    
    # Assert no placeholders remain
    assert_no_placeholders_remain(
        extracted_text=extracted_text,
        scalar_placeholders=["document_id", "title", "page_number"],
        block_placeholders=[],
        test_name="scalar_replacement_header_footer",
    )


def test_scalar_replacement_determinism(artifacts_dir, scalar_template_body):
    """Test that same input produces identical output (Invariant 2: Determinism)."""
    request = WordExportRequest(
        scalar_fields={
            "document_id": "DET-99999",
            "title": "Determinism Test",
        },
        block_fields={},
    )
    
    # Run export twice
    artifact_dir1, markdown1, extracted_text1 = run_export_and_collect_artifacts(
        test_name="scalar_replacement_determinism_run1",
        template_path=scalar_template_body,
        request=request,
        artifacts_dir=artifacts_dir,
        subdirectory="scalar_replacement",
    )
    
    artifact_dir2, markdown2, extracted_text2 = run_export_and_collect_artifacts(
        test_name="scalar_replacement_determinism_run2",
        template_path=scalar_template_body,
        request=request,
        artifacts_dir=artifacts_dir,
        subdirectory="scalar_replacement",
    )
    
    # Assert markdown is identical
    assert markdown1 == markdown2, "Markdown output must be identical for same input"
    
    # Assert extracted text is identical (normalized)
    normalized1 = " ".join(extracted_text1.split())
    normalized2 = " ".join(extracted_text2.split())
    assert normalized1 == normalized2, "Extracted text must be identical for same input"
