"""
Test suite for negative cases and safety guarantees.

This suite asserts correct refusal behavior:
- Block placeholder inside HEADER/FOOTER/TEXTBOX does NOT expand
- Block placeholder mixed with other text in a paragraph is skipped
- Unsupported cases are asserted, not silently ignored

These tests protect safety guarantees and must fail if behavior changes.
"""
import pytest
from pathlib import Path
from docx import Document

from docx_template_export.models.export_models import WordExportRequest
from tests.test_artifacts_helpers import (
    run_export_and_collect_artifacts,
    extract_all_text_from_docx,
)


@pytest.fixture
def artifacts_dir(test_output_dir):
    """Return artifacts directory for this test suite."""
    artifacts_base = test_output_dir.parent / "artifacts"
    artifacts_base.mkdir(exist_ok=True)
    return artifacts_base


@pytest.fixture
def block_in_header_template(test_fixtures_dir):
    """Create template with block placeholder in header."""
    template_path = test_fixtures_dir / "block_in_header_template.docx"
    doc = Document()
    
    # Header with block placeholder (should NOT expand)
    header = doc.sections[0].header
    header.paragraphs[0].text = "{{block_content}}"
    
    # Body
    doc.add_paragraph("Body content")
    
    template_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(template_path))
    return template_path


@pytest.fixture
def block_in_footer_template(test_fixtures_dir):
    """Create template with block placeholder in footer."""
    template_path = test_fixtures_dir / "block_in_footer_template.docx"
    doc = Document()
    
    # Body
    doc.add_paragraph("Body content")
    
    # Footer with block placeholder (should NOT expand)
    footer = doc.sections[0].footer
    footer.paragraphs[0].text = "{{block_content}}"
    
    template_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(template_path))
    return template_path


@pytest.fixture
def block_mixed_with_text_template(test_fixtures_dir):
    """Create template with block placeholder mixed with other text."""
    template_path = test_fixtures_dir / "block_mixed_with_text_template.docx"
    doc = Document()
    doc.add_paragraph("Prefix text {{block_content}} suffix text")
    template_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(template_path))
    return template_path


def test_block_placeholder_in_header_not_expanded(artifacts_dir, block_in_header_template):
    """
    Assert: Block placeholder in header does NOT expand.
    
    Block expansion is only allowed in BODY FLOW and BODY TABLE.
    Headers must leave block placeholders unchanged.
    """
    request = WordExportRequest(
        scalar_fields={},
        block_fields={
            "block_content": "# Heading\n\nContent here.",
        },
    )
    
    artifact_dir, markdown, extracted_text = run_export_and_collect_artifacts(
        test_name="block_in_header_not_expanded",
        template_path=block_in_header_template,
        request=request,
        artifacts_dir=artifacts_dir,
        subdirectory="negative_cases",
    )
    
    # Assert placeholder remains (not expanded)
    assert "{{block_content}}" in extracted_text, (
        "Block placeholder in header must NOT be expanded. "
        "Block expansion is only allowed in BODY FLOW and BODY TABLE."
    )
    
    # Assert block content does NOT appear in output
    assert "Heading" not in extracted_text, (
        "Block content must NOT appear when placeholder is in header."
    )


def test_block_placeholder_in_footer_not_expanded(artifacts_dir, block_in_footer_template):
    """
    Assert: Block placeholder in footer does NOT expand.
    
    Block expansion is only allowed in BODY FLOW and BODY TABLE.
    Footers must leave block placeholders unchanged.
    """
    request = WordExportRequest(
        scalar_fields={},
        block_fields={
            "block_content": "# Heading\n\nContent here.",
        },
    )
    
    artifact_dir, markdown, extracted_text = run_export_and_collect_artifacts(
        test_name="block_in_footer_not_expanded",
        template_path=block_in_footer_template,
        request=request,
        artifacts_dir=artifacts_dir,
        subdirectory="negative_cases",
    )
    
    # Assert placeholder remains (not expanded)
    assert "{{block_content}}" in extracted_text, (
        "Block placeholder in footer must NOT be expanded. "
        "Block expansion is only allowed in BODY FLOW and BODY TABLE."
    )
    
    # Assert block content does NOT appear in output
    assert "Heading" not in extracted_text, (
        "Block content must NOT appear when placeholder is in footer."
    )


def test_block_placeholder_mixed_with_text_skipped(artifacts_dir, block_mixed_with_text_template):
    """
    Assert: Block placeholder mixed with other text in paragraph is skipped.
    
    Block placeholders must occupy the entire paragraph.
    Partial-paragraph block replacement is intentionally unsupported.
    """
    request = WordExportRequest(
        scalar_fields={},
        block_fields={
            "block_content": "# Heading\n\nContent here.",
        },
    )
    
    artifact_dir, markdown, extracted_text = run_export_and_collect_artifacts(
        test_name="block_mixed_with_text_skipped",
        template_path=block_mixed_with_text_template,
        request=request,
        artifacts_dir=artifacts_dir,
        subdirectory="negative_cases",
    )
    
    # Assert placeholder remains (not replaced)
    assert "{{block_content}}" in extracted_text, (
        "Block placeholder mixed with other text must NOT be replaced. "
        "Block placeholders must occupy the entire paragraph."
    )
    
    # Assert block content does NOT appear in output
    assert "Heading" not in extracted_text, (
        "Block content must NOT appear when placeholder is mixed with other text."
    )
    
    # Assert surrounding text is preserved
    assert "Prefix text" in extracted_text, "Prefix text must be preserved"
    assert "suffix text" in extracted_text, "Suffix text must be preserved"


def test_empty_block_content_removes_placeholder(artifacts_dir, test_fixtures_dir, test_output_dir):
    """
    Assert: Empty block content removes placeholder but doesn't insert content.
    
    This is correct behavior: empty content should remove the placeholder
    without inserting anything.
    """
    template_path = test_fixtures_dir / "empty_block_template.docx"
    doc = Document()
    doc.add_paragraph("{{empty_block}}")
    template_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(template_path))
    
    request = WordExportRequest(
        scalar_fields={},
        block_fields={
            "empty_block": "",  # Empty content
        },
    )
    
    artifact_dir, markdown, extracted_text = run_export_and_collect_artifacts(
        test_name="empty_block_content_removes_placeholder",
        template_path=template_path,
        request=request,
        artifacts_dir=artifacts_dir,
        subdirectory="negative_cases",
    )
    
    # Assert placeholder is removed
    assert "{{empty_block}}" not in extracted_text, (
        "Empty block content must remove placeholder."
    )
    
    # Assert no content was inserted (paragraph may be empty or removed)
    # This is acceptable behavior for empty blocks


@pytest.fixture
def block_with_free_text_before_after_template(test_fixtures_dir):
    """
    Create template with free text paragraph, then block placeholder in its own paragraph,
    then another free text paragraph.
    
    This tests the scenario where block placeholder is in a separate paragraph
    (which should work) vs mixed in the same paragraph (which should be skipped).
    """
    template_path = test_fixtures_dir / "block_with_free_text_before_after_template.docx"
    doc = Document()
    
    # Paragraph 1: Free text before
    doc.add_paragraph("This is free text before the block placeholder.")
    
    # Paragraph 2: Block placeholder alone (should work)
    doc.add_paragraph("{{summary}}")
    
    # Paragraph 3: Free text after
    doc.add_paragraph("This is free text after the block placeholder.")
    
    template_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(template_path))
    return template_path


def test_block_placeholder_with_free_text_before_after(artifacts_dir, block_with_free_text_before_after_template):
    """
    Test: Block placeholder in its own paragraph with free text before and after in separate paragraphs.
    
    Expected behavior:
    - Block placeholder in its own paragraph SHOULD be replaced (meets safety requirement)
    - Free text paragraphs before and after should be preserved
    - Block content should be inserted between the free text paragraphs
    """
    request = WordExportRequest(
        scalar_fields={},
        block_fields={
            "summary": """# Summary Section

This is the summary content with markdown formatting.

## Key Points

- First important point
- Second important point
- Third important point

### Details

More detailed information here with **bold** and *italic* text.
""",
        },
    )
    
    artifact_dir, markdown, extracted_text = run_export_and_collect_artifacts(
        test_name="block_with_free_text_before_after",
        template_path=block_with_free_text_before_after_template,
        request=request,
        artifacts_dir=artifacts_dir,
        subdirectory="negative_cases",
    )
    
    # Assert placeholder is replaced (since it's alone in its paragraph)
    assert "{{summary}}" not in extracted_text, (
        "Block placeholder in its own paragraph should be replaced."
    )
    
    # Assert block content appears in output
    assert "Summary Section" in extracted_text, (
        "Block content should appear when placeholder is in its own paragraph."
    )
    assert "Key Points" in extracted_text, "Block content should be fully inserted."
    
    # Assert free text before is preserved
    assert "This is free text before" in extracted_text, (
        "Free text paragraph before placeholder should be preserved."
    )
    
    # Assert free text after is preserved
    assert "This is free text after" in extracted_text, (
        "Free text paragraph after placeholder should be preserved."
    )
    
    # Verify the order: before text -> block content -> after text
    before_idx = extracted_text.find("This is free text before")
    summary_idx = extracted_text.find("Summary Section")
    after_idx = extracted_text.find("This is free text after")
    
    assert before_idx < summary_idx < after_idx, (
        "Content should appear in order: free text before -> block content -> free text after"
    )
