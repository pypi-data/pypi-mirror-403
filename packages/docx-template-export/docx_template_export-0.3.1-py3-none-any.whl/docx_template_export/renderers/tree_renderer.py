# Copyright � 2024 Ahsan Saeed
# Licensed under the Apache License, Version 2.0
# See LICENSE and NOTICE files for details.

"""
Tree-based renderer for markdown to DOCX conversion.

This module provides an experimental tree-based rendering path that attempts to preserve
order and nesting of markdown elements where structurally safe. It is parallel to the
existing block-based renderer and does not replace it. The tree renderer is feature-flagged
and experimental.

SAFETY INVARIANTS:
- Reuses existing proven-safe patterns from replace_block_placeholder_with_content
- Uses atomic table insertion strategy (scratch → clone → insert) for consistency
- Attempts to preserve list rendering semantics (manual glyph/numbering, no Word numbering XML)
- Tracks insertion_index to attempt accurate placeholder replacement order
- Validates tree depth and node count before rendering to attempt to prevent stack overflow

PYTHON-DOCX INTERNALS USAGE:
This module uses python-docx internals (_p, _tc, _body) for precise XML control.
This is intentional and matches the existing hardened renderer's approach.
Direct XML manipulation is required for:
- Exact insertion position control (insertion_parent.insert(insertion_index, ...))
- Table cell context tracking (cell._tc for cell element access)
- Atomic table insertion (cloning XML before insertion)

These patterns are battle-tested in the existing block-based renderer and are
preserved here for safety and determinism.
"""
import logging
from copy import deepcopy
from dataclasses import dataclass, field
from typing import List, Optional, Literal, Any

from docx import Document  # type: ignore
from docx.document import Document as DocumentType  # type: ignore
from docx.text.paragraph import Paragraph  # type: ignore
from docx.shared import Inches, Pt  # type: ignore
from docx.oxml import parse_xml  # type: ignore
from docx.oxml.ns import nsdecls, qn  # type: ignore

from docx_template_export.models.markdown_tree import (
    BlockNode,
    DocumentNode,
    HeadingNode,
    ParagraphNode,
    ListNode,
    ListItemNode,
    TableNode,
    TableRowNode,
    TableCellNode,
)
from docx_template_export.models.export_config import ListRenderConfig

logger = logging.getLogger(__name__)

# Internal safety limits (non-configurable, internal-only)
MAX_TREE_DEPTH = 50  # Maximum nesting depth to prevent stack overflow
MAX_NODE_COUNT = 10000  # Maximum total nodes to prevent excessive processing


def _validate_tree_safety(document_node: DocumentNode) -> None:
    """
    Validate tree depth and node count before rendering.
    
    SAFETY INVARIANT: Aborts rendering if limits exceeded to prevent
    stack overflow or excessive processing. Does not partially mutate document.
    
    ASSUMPTION: Tree depth and node count are validated BEFORE any document mutation.
    This function is called at the start of render_markdown_tree_to_docx(), ensuring
    that if validation fails, zero changes are made to the document. This prevents
    partial corruption and makes failures explicit (exception raised, not silent).
    
    The validation is recursive and counts all nodes in the tree, including nested
    structures (lists, tables, table cells). This attempts to catch deep nesting
    or excessive content before attempting to render.
    """
    def count_nodes_and_depth(node: BlockNode, depth: int = 0) -> tuple[int, int]:
        """Recursively count nodes and track max depth."""
        if depth > MAX_TREE_DEPTH:
            raise ValueError(f"Tree depth {depth} exceeds MAX_TREE_DEPTH {MAX_TREE_DEPTH}")
        
        count = 1
        max_depth = depth
        
        if isinstance(node, DocumentNode):
            for child in node.children:
                child_count, child_depth = count_nodes_and_depth(child, depth + 1)
                count += child_count
                max_depth = max(max_depth, child_depth)
        elif isinstance(node, ListNode):
            for item in node.items:
                item_count, item_depth = count_nodes_and_depth(item, depth + 1)
                count += item_count
                max_depth = max(max_depth, item_depth)
        elif isinstance(node, ListItemNode):
            for child in node.children:
                child_count, child_depth = count_nodes_and_depth(child, depth + 1)
                count += child_count
                max_depth = max(max_depth, child_depth)
        elif isinstance(node, TableNode):
            for row in node.rows:
                for cell in row.cells:
                    for child in cell.children:
                        child_count, child_depth = count_nodes_and_depth(child, depth + 1)
                        count += child_count
                        max_depth = max(max_depth, child_depth)
        
        return count, max_depth
    
    try:
        total_count, max_depth = count_nodes_and_depth(document_node)
        if total_count > MAX_NODE_COUNT:
            raise ValueError(f"Node count {total_count} exceeds MAX_NODE_COUNT {MAX_NODE_COUNT}")
    except ValueError as e:
        logger.error(f"Tree safety validation failed: {e}. Aborting rendering to prevent document corruption.")
        raise


@dataclass
class RenderContext:
    """
    Rendering context that tracks list state and indentation.
    
    SAFETY INVARIANT: This context replaces _list_continuation_level and implicit
    numbering from the block-based renderer. It provides explicit control over
    list state, ensuring deterministic rendering without continuation heuristics.
    
    ORDERED LIST COUNTER SEMANTICS:
    - ordered_counters is a stack parallel to list_stack
    - ordered_counters[-1] corresponds to the current ListNode (list_stack[-1]) when stack is non-empty
    - When entering a ListNode(kind="ordered"), a counter is pushed (initialized to 0)
    - When rendering a ListItemNode inside that list, ordered_counters[-1] is incremented exactly once
    - When exiting the ListNode, the counter is popped
    - This attempts to ensure hierarchical numbering (1., 1.1., 1.1.1.) matches list nesting depth
    - The counter is never reset or modified outside of push/increment/pop operations
    """
    list_stack: List[Literal["bullet", "ordered"]] = field(default_factory=list)
    ordered_counters: List[int] = field(default_factory=list)
    indent_level: int = 0
    in_table_cell: bool = False
    current_cell: Optional[Any] = None  # Current table cell being rendered into (python-docx Cell object)
    config: ListRenderConfig = field(default_factory=ListRenderConfig)
    
    def copy(self) -> "RenderContext":
        """Create a deep copy of the context."""
        return RenderContext(
            list_stack=self.list_stack.copy(),
            ordered_counters=self.ordered_counters.copy(),
            indent_level=self.indent_level,
            in_table_cell=self.in_table_cell,
            current_cell=self.current_cell,
            config=self.config,
        )


def render_markdown_tree_to_docx(
    document_node: DocumentNode,
    doc: DocumentType,
    placeholder_paragraph: Optional[Paragraph] = None,
    config: Optional[ListRenderConfig] = None,
) -> None:
    """
    Render markdown tree nodes to a Word document, attempting to preserve order.
    
    This is the entry point for tree-based rendering. It initializes the
    rendering context and attempts to render all children of the document node
    where structurally safe.
    
    SAFETY INVARIANT: This function attempts to preserve placeholder replacement order
    by tracking insertion_index and updating it based on actual element counts.
    
    Args:
        document_node: Root document node containing all content
        doc: Word document to render into
        placeholder_paragraph: Optional paragraph to replace (if None, appends to document)
        config: List rendering configuration (defaults to ListRenderConfig() if None)
    """
    # Safety check: validate tree depth and node count
    _validate_tree_safety(document_node)
    
    # Use default config if not provided (avoids mutable default argument)
    if config is None:
        config = ListRenderConfig()
    
    # Determine insertion point
    insertion_parent = None
    insertion_index = None
    cell = None
    
    if placeholder_paragraph is not None:
        # Detect if placeholder paragraph is inside a table cell
        # Check if paragraph._parent is a cell (has _tc attribute)
        if hasattr(placeholder_paragraph._parent, "_tc"):
            cell = placeholder_paragraph._parent
            logger.info("Tree renderer: placeholder is inside table cell; enabling cell-context rendering")
        
        # Capture insertion point before removing paragraph
        # This follows the existing hardened pattern from replace_block_placeholder_with_content
        if cell is not None:
            # Inside table cell: use cell._tc as insertion_parent
            insertion_parent = cell._tc
            # Compute insertion_index from w:tc parent
            actual_parent_elem = cell._tc
            if actual_parent_elem is not None:
                insertion_index = actual_parent_elem.index(placeholder_paragraph._p)
                # Remove placeholder paragraph
                actual_parent_elem.remove(placeholder_paragraph._p)
        else:
            # Body flow: use paragraph's XML parent
            actual_parent_elem = placeholder_paragraph._p.getparent()
            if actual_parent_elem is not None:
                insertion_parent = actual_parent_elem
                insertion_index = actual_parent_elem.index(placeholder_paragraph._p)
                # Remove placeholder paragraph
                actual_parent_elem.remove(placeholder_paragraph._p)
    
    context = RenderContext(config=config)
    
    # Set cell context if placeholder was in a cell
    if cell is not None:
        context.in_table_cell = True
        context.current_cell = cell
    
    # Render all children, tracking insertion index accurately
    # INVARIANT: render_node() returns the count of top-level Word elements it inserted.
    # The caller (this function) owns insertion_index and updates it based on actual counts.
    # This attempts to ensure placeholder replacement order: each child's elements are inserted
    # at the captured position, and subsequent children are inserted after them, subject to DOCX constraints.
    for child in document_node.children:
        count = render_node(child, doc, context, insertion_parent, insertion_index)
        # Update insertion index based on actual count of inserted elements
        if insertion_index is not None:
            insertion_index += count


def render_node(
    node: BlockNode,
    doc: DocumentType,
    context: RenderContext,
    insertion_parent: Optional[Any] = None,
    insertion_index: Optional[int] = None,
) -> int:
    """
    Recursive dispatcher that renders a node based on its type.
    
    This function attempts to preserve order and nesting by recursively rendering
    all children in the order they appear in the tree, subject to DOCX constraints.
    
    SAFETY INVARIANT: Returns the count of top-level Word elements inserted.
    This count is used by callers to accurately track insertion_index.
    
    Args:
        node: The node to render
        doc: Word document to render into
        context: Current rendering context (modified in-place for lists)
        insertion_parent: Optional parent element for insertion (BODY FLOW)
        insertion_index: Optional insertion index (BODY FLOW)
    
    Returns:
        Number of top-level Word elements inserted (0, 1, or more for lists)
    """
    if isinstance(node, DocumentNode):
        # INVARIANT: DocumentNode aggregates counts from all children.
        # The parent (render_markdown_tree_to_docx) owns insertion_index and updates it.
        # We track total_count here for return value, but insertion_index updates
        # are handled by the parent to attempt to maintain ordering.
        total_count = 0
        for child in node.children:
            count = render_node(child, doc, context, insertion_parent, insertion_index)
            total_count += count
            if insertion_index is not None:
                insertion_index += count
        return total_count
    
    elif isinstance(node, HeadingNode):
        return render_heading_node(node, doc, context, insertion_parent, insertion_index)
    
    elif isinstance(node, ParagraphNode):
        return render_paragraph_node(node, doc, context, insertion_parent, insertion_index)
    
    elif isinstance(node, ListNode):
        return render_list_node(node, doc, context, insertion_parent, insertion_index)
    
    elif isinstance(node, ListItemNode):
        return render_list_item_node(node, doc, context, insertion_parent, insertion_index)
    
    elif isinstance(node, TableNode):
        return render_table_node(node, doc, context, insertion_parent, insertion_index)
    
    elif isinstance(node, TableRowNode):
        # TableRowNode should not be rendered directly (handled by TableNode)
        logger.warning("TableRowNode rendered directly (should be inside TableNode)")
        return 0
    
    elif isinstance(node, TableCellNode):
        render_table_cell_node(node, doc, context)
        return 0  # TableCellNode doesn't insert top-level elements
    
    else:
        logger.warning(f"Unknown node type: {type(node)}")
        return 0


def _create_paragraph_safe(
    doc: DocumentType,
    context: RenderContext,
    insertion_parent: Optional[Any] = None,
    insertion_index: Optional[int] = None,
) -> Paragraph:
    """
    Create a paragraph using existing safe patterns.
    
    SAFETY INVARIANT: Reuses proven-safe XML insertion patterns from
    replace_block_placeholder_with_content. Respects table-cell context
    and insertion position exactly as the existing renderer does.
    
    PYTHON-DOCX INTERNALS: Uses cell._tc (table cell XML element) and
    doc._body (document body) for precise XML control. This matches the
    existing hardened renderer's approach.
    """
    if context.in_table_cell and context.current_cell is not None:
        # Inside table cell: create paragraph in cell (reuses existing pattern)
        # Uses cell._tc to access table cell XML element directly
        cell = context.current_cell
        p_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
        cell_elem = cell._tc  # python-docx internal: table cell XML element
        
        # SAFETY: Ensure insertion_index accounts for <w:tcPr> which MUST come first in <w:tc>
        adjusted_index = None
        if insertion_index is not None:
            # Find position of <w:tcPr> if it exists
            tcpr_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            tcpr_tag = f"{{{tcpr_ns}}}tcPr"
            tcpr_position = None
            for idx, child in enumerate(cell_elem):
                if child.tag == tcpr_tag:
                    tcpr_position = idx
                    break
            
            adjusted_index = insertion_index
            # CRITICAL: If <w:tcPr> exists, ensure adjusted_index is AFTER it
            if tcpr_position is not None and adjusted_index <= tcpr_position:
                adjusted_index = tcpr_position + 1
            
            cell_elem.insert(adjusted_index, p_xml)
        else:
            cell_elem.append(p_xml)
        
        return Paragraph(p_xml, cell)  # type: ignore[arg-type]
    elif insertion_parent is not None and insertion_index is not None:
        # Body flow: insert at specific position (reuses existing pattern)
        # Uses insertion_parent.insert() for exact position control
        p_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
        insertion_parent.insert(insertion_index, p_xml)
        return Paragraph(p_xml, doc._body)  # python-docx internal: document body
    else:
        # Body flow: add to document (fallback when no insertion context)
        return doc.add_paragraph()


def render_heading_node(
    node: HeadingNode,
    doc: DocumentType,
    context: RenderContext,
    insertion_parent: Optional[Any] = None,
    insertion_index: Optional[int] = None,
) -> int:
    """
    Render a heading node (H1-H6) to a Word document.
    
    This function creates a paragraph with the appropriate heading style
    and applies formatted text (preserving bold/italic formatting).
    
    Safety Guarantees:
        - Attempts to preserve formatting from FormattedRun objects where supported
        - Uses safe paragraph creation (_create_paragraph_safe)
        - Respects table cell context if rendering inside a cell
        - Attempts to maintain insertion order via insertion_index
    
    Args:
        node: HeadingNode containing level and formatted runs.
        doc: Word document to render into.
        context: Current rendering context (for table cell detection).
        insertion_parent: Optional parent element for insertion (BODY FLOW).
        insertion_index: Optional insertion index (BODY FLOW).
    
    Returns:
        Number of top-level elements inserted (always 1 for heading).
        This count is used by callers to update insertion_index.
    """
    from docx_template_export.services.word_export_service import (
        _heading_style_for_level,
        _set_paragraph_formatted_text,
    )
    
    # Extract text from runs for fallback
    text = "".join(run.text for run in node.runs) if node.runs else ""
    
    # Create paragraph using safe helper
    para = _create_paragraph_safe(doc, context, insertion_parent, insertion_index)
    
    # Apply heading style
    style = _heading_style_for_level(node.level)
    para.style = style
    
    # Set formatted text
    _set_paragraph_formatted_text(para, node.runs, text)
    
    return 1  # One paragraph inserted


def render_paragraph_node(
    node: ParagraphNode,
    doc: DocumentType,
    context: RenderContext,
    insertion_parent: Optional[Any] = None,
    insertion_index: Optional[int] = None,
) -> int:
    """
    Render a paragraph node to a Word document.
    
    This function creates a paragraph with formatted text (preserving bold/italic).
    If the paragraph is inside a list context, appropriate indentation is applied
    to match list item continuation paragraph semantics.
    
    Safety Guarantees:
        - Attempts to preserve formatting from FormattedRun objects where supported
        - Applies list indentation when inside list context (continuation paragraphs) where applicable
        - Uses safe paragraph creation (_create_paragraph_safe)
        - Respects table cell context if rendering inside a cell
        - Skips empty paragraphs (returns 0, no element inserted)
    
    Args:
        node: ParagraphNode containing formatted runs.
        doc: Word document to render into.
        context: Current rendering context (for list indentation and table cell detection).
        insertion_parent: Optional parent element for insertion (BODY FLOW).
        insertion_index: Optional insertion index (BODY FLOW).
    
    Returns:
        Number of top-level elements inserted (0 if empty paragraph, 1 otherwise).
        This count is used by callers to update insertion_index.
    """
    from docx_template_export.services.word_export_service import (
        _set_paragraph_formatted_text,
    )
    
    # Extract text from runs for fallback
    text = "".join(run.text for run in node.runs) if node.runs else ""
    
    if not text.strip() and not node.runs:
        return 0  # Skip empty paragraphs
    
    # Create paragraph using safe helper
    para = _create_paragraph_safe(doc, context, insertion_parent, insertion_index)
    
    # Apply indentation if inside a list (attempts to preserve existing list rendering semantics)
    if context.list_stack:
        max_depth = context.config.max_visual_depth if context.config.max_visual_depth is not None else 3
        indent_per_level = context.config.indent_inches_per_level
        hanging_indent = -context.config.hanging_indent_inches
        visual_level = min(len(context.list_stack), max_depth)
        para.paragraph_format.left_indent = Inches(indent_per_level * visual_level)
        para.paragraph_format.first_line_indent = Inches(hanging_indent)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
    
    # Set formatted text
    _set_paragraph_formatted_text(para, node.runs, text)
    
    return 1  # One paragraph inserted


def render_list_node(
    node: ListNode,
    doc: DocumentType,
    context: RenderContext,
    insertion_parent: Optional[Any] = None,
    insertion_index: Optional[int] = None,
) -> int:
    """
    Render a list node (bullet or ordered).
    
    SAFETY INVARIANT: Attempts to preserve list rendering semantics from existing
    block-based renderer. Manual glyph/numbering logic attempts to match block renderer behavior.
    
    Returns:
        Total count of paragraphs inserted (one per list item)
    """
    # Push list type onto stack
    context.list_stack.append(node.kind)
    
    # If ordered, initialize counter
    # INVARIANT: ordered_counters stack is parallel to list_stack.
    # When we push "ordered" onto list_stack, we push 0 onto ordered_counters.
    # This counter will be incremented exactly once per ListItemNode rendered
    # inside this ListNode, ensuring correct hierarchical numbering.
    if node.kind == "ordered":
        context.ordered_counters.append(0)
    
    # Update indent level
    context.indent_level = len(context.list_stack)
    
    # Render all items in order, tracking total count
    # INVARIANT: ListNode owns insertion_index for its items. We use current_index
    # as a local variable to track position, updating it based on actual element
    # counts returned by render_node(). This attempts to ensure list items are inserted in
    # order without gaps or overlaps, subject to DOCX constraints.
    total_count = 0
    current_index = insertion_index
    for item in node.items:
        count = render_node(item, doc, context, insertion_parent, current_index)
        total_count += count
        if current_index is not None:
            current_index += count
    
    # Pop list type from stack
    context.list_stack.pop()
    
    # If ordered, pop counter
    if node.kind == "ordered":
        context.ordered_counters.pop()
    
    # Update indent level
    context.indent_level = len(context.list_stack)
    
    return total_count


def render_list_item_node(
    node: ListItemNode,
    doc: DocumentType,
    context: RenderContext,
    insertion_parent: Optional[Any] = None,
    insertion_index: Optional[int] = None,
) -> int:
    """
    Render a list item node.
    
    SAFETY INVARIANT: Attempts to preserve list item rendering semantics:
    - Manual glyph/numbering (no Word numbering XML)
    - Visual indentation behavior attempts to match block renderer
    - Formatted runs (bold/italic) preserved where supported
    
    Returns:
        Total count of paragraphs inserted (first item paragraph + continuation paragraphs)
    """
    if not context.list_stack:
        logger.warning("List item rendered outside list context")
        # Render children as regular paragraphs
        total_count = 0
        current_index = insertion_index
        for child in node.children:
            count = render_node(child, doc, context, insertion_parent, current_index)
            total_count += count
            if current_index is not None:
                current_index += count
        return total_count
    
    list_kind = context.list_stack[-1]
    max_depth = context.config.max_visual_depth if context.config.max_visual_depth is not None else 3
    indent_per_level = context.config.indent_inches_per_level
    hanging_indent = -context.config.hanging_indent_inches
    visual_level = min(len(context.list_stack), max_depth)
    
    # Increment counter for ordered lists
    # INVARIANT: Each ListItemNode increments the counter exactly once.
    # ordered_counters[-1] corresponds to the current ListNode (list_stack[-1]).
    # This increment happens before computing the prefix, so the first item gets "1.",
    # the second gets "2.", etc. For nested lists, the counter stack provides
    # hierarchical numbering (e.g., 1.1., 1.2. for items in a nested list).
    if list_kind == "ordered" and context.ordered_counters:
        context.ordered_counters[-1] += 1
    
    # Compute prefix
    if list_kind == "bullet":
        # Get glyph based on depth
        bullet_glyphs = context.config.bullet_glyphs if context.config.bullet_glyphs else ("\u2022",)
        if not bullet_glyphs:
            bullet_glyphs = ("\u2022",)
        
        logical_level = len(context.list_stack)
        if logical_level <= max_depth:
            glyph_index = min(logical_level - 1, len(bullet_glyphs) - 1)
            prefix = f"{bullet_glyphs[glyph_index]}  "
        else:
            # Handle deep nesting
            if context.config.deep_bullet_strategy == "clamp_last":
                glyph_index = min(max_depth - 1, len(bullet_glyphs) - 1)
                prefix = f"{bullet_glyphs[glyph_index]}  "
            elif context.config.deep_bullet_strategy == "cycle":
                glyph_index = ((logical_level - 1) % len(bullet_glyphs))
                prefix = f"{bullet_glyphs[glyph_index]}  "
            else:  # textual
                prefix = f"[{logical_level}]  "
    else:  # ordered
        # Build hierarchical number
        number_parts = []
        for i in range(len(context.ordered_counters)):
            if i < len(context.ordered_counters):
                number_parts.append(str(context.ordered_counters[i]))
        number_str = ".".join(number_parts) + "."
        prefix = f"{number_str}  "
    
    # Render first child (if it's a paragraph, add prefix; otherwise render as-is)
    # ASSUMPTION: ListItemNode.children[0] is typically a ParagraphNode containing
    # the list item text. However, the tree structure allows any block type as the
    # first child (e.g., table, nested list). We handle both cases to preserve
    # exact tree structure without flattening.
    first_child = node.children[0] if node.children else None
    
    if isinstance(first_child, ParagraphNode):
        # Create paragraph with prefix using safe helper
        para = _create_paragraph_safe(doc, context, insertion_parent, insertion_index)
        
        # Apply indentation (attempts to preserve existing list rendering semantics)
        para.paragraph_format.left_indent = Inches(indent_per_level * visual_level)
        para.paragraph_format.first_line_indent = Inches(hanging_indent)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        
        # Add prefix (manual glyph/numbering, no Word numbering XML)
        prefix_run = para.add_run(prefix)
        prefix_run.bold = False
        prefix_run.italic = False
        
        # Add formatted runs from first child (attempts to preserve bold/italic formatting)
        if first_child.runs:
            for run_data in first_child.runs:
                run = para.add_run(run_data.text)
                run.bold = run_data.bold
                run.italic = run_data.italic
        
        # Render remaining children, tracking count
        total_count = 1  # First paragraph already inserted
        current_index = insertion_index
        if current_index is not None:
            current_index += 1
        for child in node.children[1:]:
            count = render_node(child, doc, context, insertion_parent, current_index)
            total_count += count
            if current_index is not None:
                current_index += count
        return total_count
    else:
        # First child is not a paragraph (e.g., table, nested list)
        # INVARIANT: We still create an empty paragraph with prefix for list item semantics.
        # This attempts to ensure visual consistency: every list item has a paragraph with the
        # bullet/number glyph, even if the actual content is a table or nested list.
        # The prefix paragraph is inserted first, then the actual content follows.
        # This attempts to match the existing block-based renderer's behavior and preserve
        # visual order from the markdown tree, subject to DOCX constraints.
        para = _create_paragraph_safe(doc, context, insertion_parent, insertion_index)
        para.paragraph_format.left_indent = Inches(indent_per_level * visual_level)
        para.paragraph_format.first_line_indent = Inches(hanging_indent)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        prefix_run = para.add_run(prefix)
        prefix_run.bold = False
        prefix_run.italic = False
        
        # Render all children in order
        total_count = 1  # Prefix paragraph
        current_index = insertion_index
        if current_index is not None:
            current_index += 1
        for child in node.children:
            count = render_node(child, doc, context, insertion_parent, current_index)
            total_count += count
            if current_index is not None:
                current_index += count
        return total_count


def _create_table_atomic(
    node: TableNode,
    doc: DocumentType,
    context: RenderContext,
    insertion_parent: Optional[Any] = None,
    insertion_index: Optional[int] = None,
) -> Optional[Any]:
    """
    Create a table using the atomic strategy (scratch → clone → insert).
    
    SAFETY INVARIANT: Reuses the proven-safe atomic table insertion pattern
    from _attempt_render_nested_word_table. Builds table in scratch document,
    validates, clones XML, then inserts atomically. Any failure results in
    zero document mutation.
    
    ATOMIC STRATEGY ALIGNMENT:
    This function matches the nested-table strategy exactly:
    1. Build in scratch document (isolated, no mutation of target)
    2. Apply formatting and indentation (using scratch document context)
    3. Clone XML element (deepcopy attempts to ensure no shared references)
    4. Validate cloned element (tag check, structure validation)
    5. Insert atomically (single operation, all-or-nothing)
    
    This attempts to ensure body-flow tables have similar safety guarantees as nested tables.
    Both paths use the same atomic pattern, attempting to prevent divergence and ensure
    consistent failure behavior (abort without partial mutation).
    
    NO AUTOMATIC FALLBACK IN PHASE II:
    If table creation fails, this function returns None. The caller (render_table_node)
    handles the failure by returning 0 (no elements inserted). There is no automatic
    fallback to text-grid in Phase II - that is a Phase III concern. This preserves
    explicit control and makes failures visible.
    """
    # SAFETY: If inserting into a w:tc (table cell), return None to force nested-table path
    # This prevents inserting a cloned w:tbl into w:tc with wrong parent wrappers
    if insertion_parent is not None:
        tc_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        tc_tag = f"{{{tc_ns}}}tc"
        if insertion_parent.tag == tc_tag:
            return None  # Force nested-table path for cells
    
    if not node.rows:
        return None
    
    # ASSUMPTION: TableNode.rows[0] defines the column count for the entire table.
    # All rows must have the same number of cells (enforced by tree parser).
    # If rows[0] is missing or has no cells, we cannot determine column count safely.
    num_cols = len(node.rows[0].cells) if node.rows else 0
    if num_cols == 0:
        return None
    
    try:
        # Build table in isolation using scratch document (atomic strategy)
        scratch_doc = Document()
        scratch_table = scratch_doc.add_table(rows=0, cols=num_cols)
        
        # Render rows into scratch table - first row is treated as header
        for row_idx, row_node in enumerate(node.rows):
            row_cells = scratch_table.add_row().cells
            
            # Render each cell
            for cell_idx, cell_node in enumerate(row_node.cells[:num_cols]):
                cell = row_cells[cell_idx]
                
                # Create context for cell rendering
                cell_context = context.copy()
                cell_context.in_table_cell = True
                cell_context.current_cell = cell
                cell_context.indent_level = 0
                
                # Render cell children
                for child in cell_node.children:
                    render_node(child, scratch_doc, cell_context, None, None)
        
        # Apply formatting BEFORE cloning (using scratch document context)
        from docx_template_export.services.word_export_service import (
            _apply_default_table_formatting,
        )
        _apply_default_table_formatting(scratch_table._tbl, scratch_doc._body)
        
        # Apply indentation if inside a list (before cloning)
        if context.list_stack:
            max_depth = context.config.max_visual_depth if context.config.max_visual_depth is not None else 3
            indent_per_level = context.config.indent_inches_per_level
            visual_level = min(len(context.list_stack), max_depth)
            indent_inches = indent_per_level * visual_level
            
            tbl_elem = scratch_table._tbl
            tbl_pr = None
            for child in tbl_elem:
                if child.tag == qn('w:tblPr'):
                    tbl_pr = child
                    break
            
            if tbl_pr is None:
                tbl_pr = tbl_elem.makeelement(qn('w:tblPr'))
                tbl_elem.insert(0, tbl_pr)
            
            tbl_ind = None
            for child in tbl_pr:
                if child.tag == qn('w:tblInd'):
                    tbl_ind = child
                    break
            
            if tbl_ind is None:
                tbl_ind = tbl_pr.makeelement(qn('w:tblInd'))
                tbl_pr.append(tbl_ind)
            
            twips = int(indent_inches * 1440)
            tbl_ind.set(qn('w:w'), str(twips))
            tbl_ind.set(qn('w:type'), 'dxa')
        
        # Extract and clone the already-formatted w:tbl element
        tbl_elem = scratch_table._tbl
        cloned_tbl = deepcopy(tbl_elem)
        
        # Validate: root element must be a table
        if not cloned_tbl.tag.endswith("}tbl"):
            logger.warning("Table validation failed; table not inserted")
            return None
        
        # Insert cloned table atomically (reuses existing atomic insertion pattern)
        # INVARIANT: Atomic insertion attempts to ensure zero document mutation on failure.
        # If insertion_parent/insertion_index are provided, we insert at exact position.
        # Otherwise, we append to doc._body as fallback. This fallback is safe because:
        # - It only occurs when no insertion context is available (e.g., direct call)
        # - The table is already validated and cloned, so appending is deterministic
        # - This matches the existing hardened renderer's fallback behavior
        if insertion_parent is not None and insertion_index is not None:
            insertion_parent.insert(insertion_index, cloned_tbl)
        else:
            # Fallback: append to document body (python-docx internal: doc._body)
            # This is safe because the table is validated and cloned before this point.
            doc._body.append(cloned_tbl)
        
        # Wrap as Table object for compatibility
        from docx.table import Table
        return Table(cloned_tbl, doc._body)
        
    except Exception as e:
        # SAFETY INVARIANT: Any failure during atomic table creation results in
        # zero document mutation. The table was built in a scratch document,
        # so no changes were made to the target document. We log and return None,
        # allowing the caller to handle the failure (no automatic fallback in Phase II).
        logger.warning(f"Atomic table creation failed: {e}. Table not inserted.")
        return None


def render_table_node(
    node: TableNode,
    doc: DocumentType,
    context: RenderContext,
    insertion_parent: Optional[Any] = None,
    insertion_index: Optional[int] = None,
) -> int:
    """
    Render a table node to a Word document.
    
    This function handles table rendering in two contexts:
    1. Body flow tables: Uses atomic strategy (scratch → clone → insert)
    2. Nested tables (inside table cells): Reuses existing proven-safe nested table logic
    
    Safety Guarantees:
        - Atomic table creation for body-flow tables (zero mutation on failure)
        - Reuses existing nested table rendering for cells (proven-safe pattern)
        - Table formatting applied before cloning (attempts to be deterministic)
        - List indentation preserved when table is inside a list where applicable
        - Insertion position attempted via insertion_index, subject to DOCX constraints
    
    Args:
        node: TableNode containing rows and cells.
        doc: Word document to render into.
        context: Current rendering context (for table cell detection and list indentation).
        insertion_parent: Optional parent element for insertion (BODY FLOW).
        insertion_index: Optional insertion index (BODY FLOW).
    
    Returns:
        Number of top-level elements inserted (1 if table created successfully, 0 if failed).
        This count is used by callers to update insertion_index.
    
    Note:
        If table creation fails (e.g., invalid structure), the function returns 0
        without inserting anything. This attempts to ensure zero document mutation on failure.
    """
    if not node.rows:
        return 0
    
    num_cols = len(node.rows[0].cells) if node.rows else 0
    if num_cols == 0:
        return 0
    
    # Detect if table is being inserted into a cell (either from context OR insertion_parent is w:tc)
    cell = None
    if context.in_table_cell and context.current_cell is not None:
        cell = context.current_cell
    elif insertion_parent is not None:
        # Check if insertion_parent is a w:tc element
        tc_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        tc_tag = f"{{{tc_ns}}}tc"
        if insertion_parent.tag == tc_tag:
            # Find the cell object that owns this w:tc element
            # Search through document tables to find matching cell
            for table in doc.tables:
                for row in table.rows:
                    for table_cell in row.cells:
                        if table_cell._tc is insertion_parent:
                            cell = table_cell
                            break
                    if cell is not None:
                        break
                if cell is not None:
                    break
    
    if cell is not None:
        # Nested table: use existing proven-safe nested table rendering logic
        from docx_template_export.services.word_export_service import (
            _attempt_render_nested_word_table,
            MarkdownBlock,
        )
        
        # Convert tree to MarkdownBlock for compatibility with existing logic
        header = []
        rows = []
        
        if node.rows:
            first_row = node.rows[0]
            header = [_extract_cell_text(cell_node) for cell_node in first_row.cells]
            rows = [[_extract_cell_text(cell_node) for cell_node in row.cells] for row in node.rows[1:]]
        
        block = MarkdownBlock(type="table", header=header, rows=rows)
        
        cell_elem = cell._tc
        # SAFETY: Compute insertion_index as true XML child index, accounting for <w:tcPr>
        # Use len(cell._tc) as base (append-at-end), then clamp to be after tcPr if it exists
        tcpr_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        tcpr_tag = f"{{{tcpr_ns}}}tcPr"
        
        # Find position of <w:tcPr> if it exists
        tcpr_position = None
        for idx, child in enumerate(cell_elem):
            if child.tag == tcpr_tag:
                tcpr_position = idx
                break
        
        # Compute insertion_index: use provided index if available, otherwise append at end
        if insertion_index is not None:
            cell_insertion_index = insertion_index
        else:
            cell_insertion_index = len(cell_elem)
        
        # CRITICAL: If <w:tcPr> exists, ensure insertion_index is AFTER it
        if tcpr_position is not None and cell_insertion_index <= tcpr_position:
            cell_insertion_index = tcpr_position + 1
        
        success = _attempt_render_nested_word_table(
            block,
            parent_cell_elem=cell_elem,
            insertion_index=cell_insertion_index,
            render_context=None,  # Not used in tree renderer
        )
        
        if not success:
            logger.info("Nested table rendered as text-grid fallback")
        
        return 1 if success else 0
    else:
        # Body flow: use atomic strategy (scratch → clone → insert)
        table = _create_table_atomic(node, doc, context, insertion_parent, insertion_index)
        return 1 if table is not None else 0


def render_table_cell_node(
    node: TableCellNode,
    doc: DocumentType,
    context: RenderContext,
) -> None:
    """
    Render a table cell node by rendering its children in order.
    
    This function renders all children of a table cell node (paragraphs, headings,
    lists, nested tables) into the current cell context. The cell context is
    updated to indicate we're inside a table cell, and indent level is reset.
    
    Safety Guarantees:
        - Children are rendered in order from the tree where structurally safe
        - Cell context is properly isolated (copied context, not shared)
        - Indent level reset attempts to ensure proper formatting inside cells
        - Block types are supported where applicable (headings, paragraphs, lists, nested tables)
    
    Args:
        node: TableCellNode containing children to render.
        doc: Word document being rendered into.
        context: Current rendering context. A copy is made for cell rendering
            to isolate cell context from parent context.
    
    Returns:
        None. Children are rendered directly into the current cell (from context.current_cell).
    
    Note:
        This function does not return an element count because table cells are
        not top-level elements. The cell itself is part of the table structure,
        and we're just populating its content.
    """
    # Create a new context for the cell
    cell_context = context.copy()
    cell_context.in_table_cell = True
    cell_context.indent_level = 0  # Reset indent inside cells
    
    # Render all children in order
    for child in node.children:
        render_node(child, doc, cell_context)


def _extract_cell_text(cell_node: TableCellNode) -> str:
    """
    Extract plain text from a table cell node for compatibility with existing logic.
    
    This function is used only when converting tree structures to MarkdownBlock
    format for compatibility with the existing nested table rendering logic
    (_attempt_render_nested_word_table). The tree renderer itself renders cells
    with full structure (headings, paragraphs, lists), but the nested table
    renderer expects plain text strings.
    
    This is a compatibility bridge that allows the tree renderer to reuse
    proven-safe nested table rendering code without duplicating logic.
    
    Args:
        cell_node: TableCellNode to extract text from. Only ParagraphNode children
            are considered; other node types are ignored.
    
    Returns:
        Plain text string containing concatenated text from all paragraph children,
        joined with spaces. If no paragraph children exist, returns empty string.
    
    Note:
        This function only extracts text from ParagraphNode children. Headings,
        lists, and nested tables are ignored. This is intentional for compatibility
        with the existing nested table renderer which expects simple text cells.
    """
    text_parts = []
    for child in cell_node.children:
        if isinstance(child, ParagraphNode):
            text_parts.append("".join(run.text for run in child.runs))
    return " ".join(text_parts)
