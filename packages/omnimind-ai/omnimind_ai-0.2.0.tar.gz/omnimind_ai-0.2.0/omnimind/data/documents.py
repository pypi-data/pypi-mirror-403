"""
OMNIMIND Document Processing
Support for various document formats: .txt, .pdf, .doc, .docx, .html, .md, .json, .csv

Features:
1. Universal document loading
2. Text extraction
3. Chunking for training
4. Streaming for large files
5. Metadata extraction

Supported Formats:
- .txt - Plain text
- .pdf - PDF documents
- .doc/.docx - Microsoft Word
- .html - Web pages
- .md - Markdown
- .json - JSON files
- .csv - CSV data
- .xml - XML documents
- .rtf - Rich text format
"""
import os
import re
import json
import html
from typing import Optional, List, Dict, Any, Generator, Union
from dataclasses import dataclass
from pathlib import Path


@dataclass
class Document:
    """Represents a loaded document"""
    content: str
    metadata: Dict[str, Any]
    source: str
    format: str
    
    def __len__(self):
        return len(self.content)
    
    def chunks(self, chunk_size: int = 1000, overlap: int = 100) -> List[str]:
        """Split document into overlapping chunks"""
        chunks = []
        start = 0
        while start < len(self.content):
            end = start + chunk_size
            chunk = self.content[start:end]
            chunks.append(chunk)
            start = end - overlap
        return chunks


class TextLoader:
    """Load plain text files"""
    
    EXTENSIONS = ['.txt', '.text']
    
    @staticmethod
    def load(path: str, encoding: str = 'utf-8') -> Document:
        with open(path, 'r', encoding=encoding, errors='ignore') as f:
            content = f.read()
        
        return Document(
            content=content,
            metadata={'encoding': encoding, 'size': len(content)},
            source=path,
            format='txt'
        )


class MarkdownLoader:
    """Load Markdown files"""
    
    EXTENSIONS = ['.md', '.markdown']
    
    @staticmethod
    def load(path: str) -> Document:
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        # Extract headers for metadata
        headers = re.findall(r'^#+\s+(.+)$', content, re.MULTILINE)
        
        return Document(
            content=content,
            metadata={'headers': headers[:10]},
            source=path,
            format='markdown'
        )


class HTMLLoader:
    """Load HTML files and extract text"""
    
    EXTENSIONS = ['.html', '.htm']
    
    @staticmethod
    def load(path: str, keep_structure: bool = False) -> Document:
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            raw_html = f.read()
        
        # Extract title
        title_match = re.search(r'<title[^>]*>([^<]+)</title>', raw_html, re.IGNORECASE)
        title = title_match.group(1) if title_match else None
        
        # Remove script and style
        clean = re.sub(r'<script[^>]*>.*?</script>', '', raw_html, flags=re.DOTALL | re.IGNORECASE)
        clean = re.sub(r'<style[^>]*>.*?</style>', '', clean, flags=re.DOTALL | re.IGNORECASE)
        
        if keep_structure:
            # Keep basic structure
            clean = re.sub(r'<br\s*/?>', '\n', clean)
            clean = re.sub(r'<p[^>]*>', '\n\n', clean)
            clean = re.sub(r'<h[1-6][^>]*>', '\n## ', clean)
            clean = re.sub(r'</h[1-6]>', '\n', clean)
        
        # Remove all remaining tags
        clean = re.sub(r'<[^>]+>', '', clean)
        
        # Decode HTML entities
        clean = html.unescape(clean)
        
        # Clean whitespace
        clean = re.sub(r'\n\s*\n', '\n\n', clean)
        clean = clean.strip()
        
        return Document(
            content=clean,
            metadata={'title': title, 'original_size': len(raw_html)},
            source=path,
            format='html'
        )


class JSONLoader:
    """Load JSON files"""
    
    EXTENSIONS = ['.json', '.jsonl']
    
    @staticmethod
    def load(path: str, text_field: str = None) -> Document:
        with open(path, 'r', encoding='utf-8') as f:
            if path.endswith('.jsonl'):
                # JSON Lines format
                data = [json.loads(line) for line in f if line.strip()]
            else:
                data = json.load(f)
        
        # Extract text
        if text_field:
            if isinstance(data, list):
                content = '\n\n'.join(str(item.get(text_field, '')) for item in data)
            else:
                content = str(data.get(text_field, ''))
        else:
            content = json.dumps(data, ensure_ascii=False, indent=2)
        
        return Document(
            content=content,
            metadata={'type': type(data).__name__, 'items': len(data) if isinstance(data, list) else 1},
            source=path,
            format='json'
        )


class CSVLoader:
    """Load CSV files"""
    
    EXTENSIONS = ['.csv', '.tsv']
    
    @staticmethod
    def load(path: str, text_columns: List[str] = None) -> Document:
        import csv
        
        delimiter = '\t' if path.endswith('.tsv') else ','
        
        rows = []
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            reader = csv.DictReader(f, delimiter=delimiter)
            headers = reader.fieldnames or []
            
            for row in reader:
                if text_columns:
                    text = ' | '.join(str(row.get(col, '')) for col in text_columns if col in row)
                else:
                    text = ' | '.join(str(v) for v in row.values())
                rows.append(text)
        
        content = '\n'.join(rows)
        
        return Document(
            content=content,
            metadata={'columns': headers, 'rows': len(rows)},
            source=path,
            format='csv'
        )


class PDFLoader:
    """Load PDF files"""
    
    EXTENSIONS = ['.pdf']
    
    @staticmethod
    def load(path: str) -> Document:
        """Load PDF - requires PyPDF2 or pdfplumber"""
        try:
            # Try PyPDF2 first
            from PyPDF2 import PdfReader
            
            reader = PdfReader(path)
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            
            content = '\n\n'.join(pages)
            metadata = {
                'pages': len(reader.pages),
                'info': dict(reader.metadata) if reader.metadata else {}
            }
            
        except ImportError:
            try:
                # Try pdfplumber
                import pdfplumber
                
                pages = []
                with pdfplumber.open(path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            pages.append(text)
                    
                    content = '\n\n'.join(pages)
                    metadata = {'pages': len(pdf.pages)}
                    
            except ImportError:
                # Fallback: try basic extraction
                content = f"[PDF file: {path}]\n[Install PyPDF2 or pdfplumber for text extraction]"
                metadata = {'error': 'PDF library not installed'}
        
        return Document(
            content=content,
            metadata=metadata,
            source=path,
            format='pdf'
        )


class DocxLoader:
    """Load Word documents (.docx)"""
    
    EXTENSIONS = ['.docx', '.doc']
    
    @staticmethod
    def load(path: str) -> Document:
        """Load Word doc - requires python-docx"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = '\n\n'.join(paragraphs)
            
            metadata = {
                'paragraphs': len(paragraphs),
                'tables': len(doc.tables)
            }
            
            # Extract core properties if available
            if doc.core_properties:
                metadata['title'] = doc.core_properties.title
                metadata['author'] = doc.core_properties.author
            
        except ImportError:
            content = f"[Word file: {path}]\n[Install python-docx for text extraction]"
            metadata = {'error': 'python-docx not installed'}
        
        return Document(
            content=content,
            metadata=metadata,
            source=path,
            format='docx'
        )


class XMLLoader:
    """Load XML files"""
    
    EXTENSIONS = ['.xml']
    
    @staticmethod
    def load(path: str) -> Document:
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        # Extract text content from tags
        text = re.sub(r'<[^>]+>', ' ', content)
        text = ' '.join(text.split())
        
        return Document(
            content=text,
            metadata={'raw_size': len(content)},
            source=path,
            format='xml'
        )


class DocumentLoader:
    """
    Universal document loader
    
    Automatically detects format and loads document
    
    Usage:
        loader = DocumentLoader()
        doc = loader.load("file.pdf")
        print(doc.content)
    """
    
    LOADERS = {
        '.txt': TextLoader,
        '.text': TextLoader,
        '.md': MarkdownLoader,
        '.markdown': MarkdownLoader,
        '.html': HTMLLoader,
        '.htm': HTMLLoader,
        '.json': JSONLoader,
        '.jsonl': JSONLoader,
        '.csv': CSVLoader,
        '.tsv': CSVLoader,
        '.pdf': PDFLoader,
        '.docx': DocxLoader,
        '.doc': DocxLoader,
        '.xml': XMLLoader,
    }
    
    @classmethod
    def load(cls, path: str, **kwargs) -> Document:
        """Load document from path"""
        ext = Path(path).suffix.lower()
        
        if ext not in cls.LOADERS:
            # Fallback to text
            return TextLoader.load(path)
        
        loader = cls.LOADERS[ext]
        return loader.load(path, **kwargs)
    
    @classmethod
    def load_directory(
        cls, 
        directory: str, 
        extensions: List[str] = None,
        recursive: bool = True
    ) -> Generator[Document, None, None]:
        """Load all documents from directory"""
        extensions = extensions or list(cls.LOADERS.keys())
        
        path = Path(directory)
        pattern = '**/*' if recursive else '*'
        
        for file_path in path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in extensions:
                try:
                    yield cls.load(str(file_path))
                except Exception as e:
                    print(f"Warning: Failed to load {file_path}: {e}")
    
    @classmethod
    def supported_formats(cls) -> List[str]:
        """List supported formats"""
        return list(set(cls.LOADERS.keys()))


class DocumentDataset:
    """
    Dataset for training from documents
    
    Usage:
        dataset = DocumentDataset.from_directory("./docs", tokenizer)
        for batch in dataset:
            train(batch)
    """
    
    def __init__(
        self,
        documents: List[Document],
        tokenizer,
        max_length: int = 2048,
        chunk_size: int = 1000,
        chunk_overlap: int = 100
    ):
        self.documents = documents
        self.tokenizer = tokenizer
        self.max_length = max_length
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Create chunks
        self.chunks = []
        for doc in documents:
            doc_chunks = doc.chunks(chunk_size, chunk_overlap)
            self.chunks.extend(doc_chunks)
    
    @classmethod
    def from_directory(
        cls,
        directory: str,
        tokenizer,
        extensions: List[str] = None,
        **kwargs
    ) -> 'DocumentDataset':
        """Create dataset from directory of documents"""
        loader = DocumentLoader()
        documents = list(loader.load_directory(directory, extensions))
        return cls(documents, tokenizer, **kwargs)
    
    @classmethod
    def from_files(
        cls,
        files: List[str],
        tokenizer,
        **kwargs
    ) -> 'DocumentDataset':
        """Create dataset from list of files"""
        loader = DocumentLoader()
        documents = [loader.load(f) for f in files]
        return cls(documents, tokenizer, **kwargs)
    
    def __len__(self):
        return len(self.chunks)
    
    def __getitem__(self, idx):
        chunk = self.chunks[idx]
        
        # Tokenize
        if hasattr(self.tokenizer, 'encode'):
            input_ids = self.tokenizer.encode(chunk)
        else:
            input_ids = self.tokenizer(chunk)['input_ids']
        
        # Truncate/pad
        if len(input_ids) > self.max_length:
            input_ids = input_ids[:self.max_length]
        
        import torch
        return {
            'input_ids': torch.tensor(input_ids),
            'labels': torch.tensor(input_ids)
        }


def load_document(path: str, **kwargs) -> Document:
    """Quick function to load any document"""
    return DocumentLoader.load(path, **kwargs)


def load_documents(paths: List[str]) -> List[Document]:
    """Load multiple documents"""
    return [DocumentLoader.load(p) for p in paths]


def documents_to_text(paths: List[str], separator: str = '\n\n---\n\n') -> str:
    """Load documents and combine to text"""
    docs = load_documents(paths)
    return separator.join(doc.content for doc in docs)


# Convenience functions for specific formats
def load_pdf(path: str) -> Document:
    return PDFLoader.load(path)

def load_docx(path: str) -> Document:
    return DocxLoader.load(path)

def load_html(path: str) -> Document:
    return HTMLLoader.load(path)

def load_txt(path: str) -> Document:
    return TextLoader.load(path)

def load_json(path: str, text_field: str = None) -> Document:
    return JSONLoader.load(path, text_field)

def load_csv(path: str, text_columns: List[str] = None) -> Document:
    return CSVLoader.load(path, text_columns)


if __name__ == "__main__":
    print("OMNIMIND Document Processing")
    print("=" * 40)
    print()
    print("Supported formats:")
    for fmt in DocumentLoader.supported_formats():
        print(f"  - {fmt}")
    print()
    print("Usage:")
    print("  from omnimind.documents import load_document")
    print("  doc = load_document('file.pdf')")
    print("  print(doc.content)")
