from __future__ import annotations

import io
import os
import re
import subprocess
import tempfile
from dataclasses import dataclass, field


@dataclass(frozen=True)
class DocumentSection:
    text: str
    page_start: int | None = None
    page_end: int | None = None
    section_heading: str | None = None


@dataclass
class ExtractionResult:
    sections: list[DocumentSection]
    page_count: int | None = None
    warnings: list[str] = field(default_factory=list)
    title_hint: str | None = None


_MARKDOWN_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")


def chunk_text_with_overlap(text: str, chunk_size: int, overlap: int) -> list[str]:
    text = text.strip()
    if not text:
        return []
    if chunk_size <= 0 or len(text) <= chunk_size:
        return [text]
    if overlap < 0:
        overlap = 0
    if overlap >= chunk_size:
        overlap = max(0, chunk_size - 1)

    chunks: list[str] = []
    start = 0
    length = len(text)
    while start < length:
        end = min(start + chunk_size, length)
        split = end
        if end < length:
            space = text.rfind(" ", start, end)
            if space > start:
                split = space
        chunk = text[start:split].strip()
        if chunk:
            chunks.append(chunk)
        if split >= length:
            break
        next_start = split - overlap
        if next_start <= start:
            next_start = split
        start = next_start

    return chunks or [text[:chunk_size]]


def decode_bytes_to_text(data: bytes) -> tuple[str, list[str]]:
    warnings: list[str] = []
    try:
        return data.decode("utf-8"), warnings
    except UnicodeDecodeError:
        decoded = data.decode("utf-8", errors="replace")
        if "\ufffd" in decoded:
            warnings.append("Text contained invalid utf-8 bytes; replacements applied.")
        return decoded, warnings


def extract_plain_text(text: str) -> ExtractionResult:
    cleaned = text.strip()
    return ExtractionResult(
        sections=[DocumentSection(text=cleaned)] if cleaned else [],
    )


def extract_markdown_sections(text: str) -> ExtractionResult:
    sections: list[DocumentSection] = []
    buffer: list[str] = []
    current_heading: str | None = None
    title_hint: str | None = None

    def flush():
        nonlocal buffer
        content = "\n".join(buffer).strip()
        if content:
            sections.append(
                DocumentSection(text=content, section_heading=current_heading)
            )
        buffer = []

    for line in text.splitlines():
        match = _MARKDOWN_HEADING_RE.match(line.strip())
        if match:
            flush()
            current_heading = match.group(2).strip()
            if title_hint is None:
                title_hint = current_heading or title_hint
            buffer.append(line.strip())
        else:
            buffer.append(line)

    flush()

    if not sections and text.strip():
        sections.append(DocumentSection(text=text.strip(), section_heading=None))

    return ExtractionResult(sections=sections, title_hint=title_hint)


def _extract_docx_sections_sync(data: bytes) -> ExtractionResult:
    try:
        import docx  # type: ignore
    except ImportError as exc:  # pragma: no cover - dependency missing
        raise RuntimeError("python-docx is required for .docx files.") from exc

    document = docx.Document(io.BytesIO(data))
    sections: list[DocumentSection] = []
    buffer: list[str] = []
    current_heading: str | None = None
    title_hint: str | None = None

    def flush():
        nonlocal buffer
        content = "\n".join(buffer).strip()
        if content:
            sections.append(
                DocumentSection(text=content, section_heading=current_heading)
            )
        buffer = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = ""
        try:
            style_name = paragraph.style.name or ""
        except Exception:  # pragma: no cover - style access varies
            style_name = ""
        if style_name.lower().startswith("heading"):
            flush()
            current_heading = text
            if title_hint is None:
                title_hint = current_heading
            buffer.append(text)
        else:
            buffer.append(text)

    flush()

    if not sections and document.paragraphs:
        full_text = "\n".join(
            paragraph.text for paragraph in document.paragraphs if paragraph.text
        ).strip()
        if full_text:
            sections.append(DocumentSection(text=full_text, section_heading=None))

    return ExtractionResult(sections=sections, title_hint=title_hint)


def _extract_pdf_sections_sync(data: bytes, *, ocr: bool) -> ExtractionResult:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as exc:  # pragma: no cover - dependency missing
        raise RuntimeError("pypdf is required for .pdf files.") from exc

    reader = PdfReader(io.BytesIO(data))
    page_count = len(reader.pages)
    page_texts: list[str] = []
    empty_pages: list[int] = []
    warnings: list[str] = []

    for idx, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if not text.strip():
            empty_pages.append(idx)
        page_texts.append(text)

    if ocr and empty_pages:
        try:
            import pytesseract  # type: ignore
            from pdf2image import convert_from_bytes  # type: ignore
        except ImportError:  # pragma: no cover - optional dependency missing
            warnings.append("OCR requested but pdf2image/pytesseract not installed.")
        else:
            try:
                images = convert_from_bytes(data)
                for idx in empty_pages:
                    if idx < len(images):
                        ocr_text = pytesseract.image_to_string(images[idx])
                        if ocr_text.strip():
                            page_texts[idx] = ocr_text
                        else:
                            warnings.append(f"OCR produced no text for page {idx + 1}.")
                    else:
                        warnings.append(f"OCR image missing for page {idx + 1}.")
            except Exception as exc:  # pragma: no cover - OCR errors vary
                warnings.append(f"OCR failed: {exc}")

    sections: list[DocumentSection] = []
    for idx, text in enumerate(page_texts):
        if not text.strip():
            continue
        page_number = idx + 1
        sections.append(
            DocumentSection(
                text=text.strip(),
                page_start=page_number,
                page_end=page_number,
                section_heading=None,
            )
        )

    if not sections:
        warnings.append("No text extracted from PDF.")

    return ExtractionResult(
        sections=sections,
        page_count=page_count,
        warnings=warnings,
    )


def _extract_doc_sections_sync(data: bytes) -> ExtractionResult:
    warnings: list[str] = []
    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as temp_file:
            temp_file.write(data)
            temp_file.flush()
            temp_path = temp_file.name
        result = subprocess.run(
            ["antiword", temp_path],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            check=False,
        )
    except FileNotFoundError as exc:  # pragma: no cover - binary missing
        raise RuntimeError("antiword is required for .doc files.") from exc
    finally:
        if temp_path and os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
            except OSError:
                warnings.append("Failed to remove temporary .doc file.")

    if result.returncode != 0:
        warnings.append(
            f"antiword exited with code {result.returncode}: "
            f"{result.stderr.decode('utf-8', errors='replace').strip()}"
        )

    text, decode_warnings = decode_bytes_to_text(result.stdout)
    warnings.extend(decode_warnings)

    sections: list[DocumentSection] = []
    cleaned = text.strip()
    if cleaned:
        sections.append(DocumentSection(text=cleaned))
    else:
        warnings.append("No text extracted from .doc file.")

    return ExtractionResult(sections=sections, warnings=warnings)


def extract_document_sections(
    file_type: str,
    *,
    text: str | None = None,
    data: bytes | None = None,
    ocr: bool = False,
) -> ExtractionResult:
    if file_type in {"txt", "text"}:
        if text is None and data is not None:
            text, warnings = decode_bytes_to_text(data)
            result = extract_plain_text(text)
            result.warnings.extend(warnings)
            return result
        return extract_plain_text(text or "")

    if file_type in {"md", "markdown"}:
        if text is None and data is not None:
            text, warnings = decode_bytes_to_text(data)
            result = extract_markdown_sections(text)
            result.warnings.extend(warnings)
            return result
        return extract_markdown_sections(text or "")

    if file_type == "docx":
        if data is None:
            raise ValueError(".docx ingestion requires binary data.")
        return _extract_docx_sections_sync(data)

    if file_type == "doc":
        if data is None:
            raise ValueError(".doc ingestion requires binary data.")
        return _extract_doc_sections_sync(data)

    if file_type == "pdf":
        if data is None:
            raise ValueError(".pdf ingestion requires binary data.")
        return _extract_pdf_sections_sync(data, ocr=ocr)

    raise ValueError(f"Unsupported file_type: {file_type}")
