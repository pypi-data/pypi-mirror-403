#!/usr/bin/env python3
# SPDX-License-Identifier: Apache-2.0
"""
Indexing pipeline for building a RAG index from PDF, DOCX, DOC, Markdown, and TXT documents.
Supports incremental indexing using a local SQLite database to track file states.
"""
import argparse
import hashlib
import itertools
import json
import logging
import os
import sqlite3
import sys
import threading
import time
from datetime import datetime
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Dict, Optional, Iterator, Set, Any

from docx import Document

from llama_index.core import (
    VectorStoreIndex,
    StorageContext,
    Settings,
    SimpleDirectoryReader,
    Document as LlamaIndexDocument,
    load_index_from_storage,
)
from llama_index.core.node_parser import SentenceSplitter
from llama_index.embeddings.fastembed import FastEmbedEmbedding

# Load configuration from config.yaml
from . import cfgload
from .cfgload import load_config
_config = load_config()

# Configuration from config.yaml
STORAGE_DIR = Path(_config["storage"]["storage_dir"])
STATE_DB_PATH = STORAGE_DIR / "ingestion_state.db"

# Stage 1 (embedding/vector search) configuration
RETRIEVAL_EMBED_MODEL_NAME = _config["retrieval"]["embed_model_name"]

# Stage 2 (FlashRank reranking, CPU-only, ONNX-based) configuration
RETRIEVAL_RERANK_MODEL_NAME = _config["retrieval"]["rerank_model_name"]

# Shared cache directory for embedding and reranking models
RETRIEVAL_MODEL_CACHE_DIR = Path(_config["storage"]["model_cache_dir"])

# BM25 index directory for file name matching
BM25_INDEX_DIR = STORAGE_DIR / "bm25_index"

# Heading store for document headings (stored separately to avoid metadata size issues)
HEADING_STORE_PATH = STORAGE_DIR / "heading_store.json"

# Metadata exclusion configuration
# These keys are excluded from the embedding text to save tokens and avoid length errors
EXCLUDED_EMBED_METADATA_KEYS = [
    "line_offsets",      # Large integer array, primary cause of length errors
    "document_headings", # Heading hierarchy array with positions, excluded like line_offsets
    "heading_path",      # Pre-computed heading hierarchy, stored separately to save chunk space
    "file_path",         # redundant with file_name/source, strict path less useful for semantic similarity
    "source",            # often same as file_path
    "creation_date",     # temporal, not semantic
    "last_modified_date",# temporal, not semantic
    "doc_ids",           # internal tracking
    "hash",              # internal tracking
]

# These keys are excluded from the LLM context to save context window
EXCLUDED_LLM_METADATA_KEYS = [
    "line_offsets",      # LLM needs text content, not integer map
    "hash",              # internal tracking
    "doc_ids",           # internal tracking
    "file_path",         # usually redundant if file_name is present
    "source",            # usually redundant
]

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)


# Default file type patterns
DEFAULT_INCLUDE_PATTERNS = ["**/*.pdf", "**/*.md", "**/*.txt", "**/*.docx", "**/*.doc"]


@dataclass
class DirectoryConfig:
    """Configuration for a single source directory."""
    path: Path
    enabled: bool = True
    include: List[str] = field(default_factory=lambda: DEFAULT_INCLUDE_PATTERNS.copy())
    exclude: List[str] = field(default_factory=list)
    recursive: bool = True


@dataclass
class IndexConfig:
    """Complete indexing configuration."""
    directories: List[DirectoryConfig]
    chunk_size: int = 1600
    chunk_overlap: int = 200


def load_index_config() -> IndexConfig:
    """Load indexing configuration from config.yaml.

    Raises:
        ValueError: If config is invalid
    """
    indexing_config = _config.get("indexing", {})

    if not indexing_config.get("directories"):
        raise ValueError(
            "Config must have at least one directory in 'indexing.directories'.\n"
            "Please update config.yaml with your directory configuration.\n"
            "Example:\n"
            "indexing:\n"
            "  directories:\n"
            '    - "./data"\n'
            "  chunk_size: 1600\n"
            "  chunk_overlap: 200\n"
        )

    logger.info("Loading indexing config from config.yaml")
    return _parse_index_config(indexing_config)


def _parse_index_config(config_data: dict) -> IndexConfig:
    """Parse raw config dict into IndexConfig."""
    # Get defaults section
    defaults = config_data.get("defaults", {})
    default_include = defaults.get("include", DEFAULT_INCLUDE_PATTERNS.copy())
    default_exclude = defaults.get("exclude", [])
    default_recursive = defaults.get("recursive", True)

    # Parse directories
    directories: List[DirectoryConfig] = []
    raw_dirs = config_data.get("directories", [])

    if not raw_dirs:
        raise ValueError("Config must have at least one directory in 'directories' list")

    for entry in raw_dirs:
        if isinstance(entry, str):
            # Simple path string - use defaults
            dir_config = DirectoryConfig(
                path=Path(entry),
                include=default_include.copy(),
                exclude=default_exclude.copy(),
                recursive=default_recursive,
            )
        elif isinstance(entry, dict):
            # Full directory config object
            path_str = entry.get("path")
            if not path_str:
                raise ValueError(f"Directory config missing 'path': {entry}")

            dir_config = DirectoryConfig(
                path=Path(path_str),
                enabled=entry.get("enabled", True),
                include=entry.get("include", default_include.copy()),
                exclude=entry.get("exclude", default_exclude.copy()),
                recursive=entry.get("recursive", default_recursive),
            )
        else:
            raise ValueError(f"Invalid directory entry: {entry}")

        directories.append(dir_config)

    return IndexConfig(
        directories=directories,
        chunk_size=config_data.get("chunk_size", 1600),
        chunk_overlap=config_data.get("chunk_overlap", 200),
    )


class HeadingStore:
    """Stores document headings separately from chunk metadata.

    This avoids the LlamaIndex SentenceSplitter metadata size validation issue,
    which checks metadata length before applying exclusions. By storing headings
    in a separate file, we keep chunk metadata small while preserving heading
    data for retrieval.
    """

    def __init__(self, store_path: Path):
        self.store_path = store_path
        self._data: Dict[str, List[dict]] = {}
        self._load()

    def _load(self):
        """Load heading data from disk."""
        if self.store_path.exists():
            try:
                with open(self.store_path, "r", encoding="utf-8") as f:
                    self._data = json.load(f)
            except Exception as e:
                logger.warning(f"Failed to load heading store: {e}")
                self._data = {}

    def _save(self):
        """Save heading data to disk."""
        self.store_path.parent.mkdir(parents=True, exist_ok=True)
        with open(self.store_path, "w", encoding="utf-8") as f:
            json.dump(self._data, f)

    def set_headings(self, file_path: str, headings: List[dict]):
        """Store headings for a file."""
        self._data[file_path] = headings
        self._save()

    def get_headings(self, file_path: str) -> List[dict]:
        """Get headings for a file."""
        return self._data.get(file_path, [])

    def remove_headings(self, file_path: str):
        """Remove headings for a file."""
        if file_path in self._data:
            del self._data[file_path]
            self._save()


# Module-level heading store instance (lazy initialized)
_heading_store: Optional["HeadingStore"] = None


def get_heading_store() -> HeadingStore:
    """Get the singleton HeadingStore instance."""
    global _heading_store
    if _heading_store is None:
        _heading_store = HeadingStore(HEADING_STORE_PATH)
    return _heading_store


@dataclass
class FileInfo:
    """Metadata about a file in the data source."""
    path: str
    hash: str
    last_modified: float
    source_dir: str = ""  # Tracks which configured directory this file came from


class IngestionState:
    """Manages the state of ingested files using a SQLite database."""

    def __init__(self, db_path: Path):
        self.db_path = db_path
        self._init_db()

    def _init_db(self):
        """Initialize the SQLite database schema with migration support."""
        self.db_path.parent.mkdir(parents=True, exist_ok=True)
        with sqlite3.connect(self.db_path) as conn:
            # Check if table exists
            cursor = conn.execute(
                "SELECT name FROM sqlite_master WHERE type='table' AND name='files'"
            )
            table_exists = cursor.fetchone() is not None

            if not table_exists:
                # Create new table with source_dir column
                conn.execute(
                    """
                    CREATE TABLE files (
                        path TEXT PRIMARY KEY,
                        hash TEXT NOT NULL,
                        last_modified REAL NOT NULL,
                        doc_ids TEXT NOT NULL,
                        source_dir TEXT DEFAULT ''
                    )
                    """
                )
            else:
                # Migration: add source_dir column if missing
                cursor = conn.execute("PRAGMA table_info(files)")
                columns = {row[1] for row in cursor}
                if "source_dir" not in columns:
                    conn.execute("ALTER TABLE files ADD COLUMN source_dir TEXT DEFAULT ''")
                    logger.info("Migrated files table: added source_dir column")

    def get_all_files(self) -> Dict[str, dict]:
        """Retrieve all tracked files and their metadata."""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.execute(
                "SELECT path, hash, last_modified, doc_ids, source_dir FROM files"
            )
            return {
                row[0]: {
                    "hash": row[1],
                    "last_modified": row[2],
                    "doc_ids": row[3].split(",") if row[3] else [],
                    "source_dir": row[4] if row[4] else "",
                }
                for row in cursor
            }

    def update_file_state(self, file_info: FileInfo, doc_ids: List[str]):
        """Update or insert the state for a file."""
        with sqlite3.connect(self.db_path) as conn:
            conn.execute(
                """
                INSERT INTO files (path, hash, last_modified, doc_ids, source_dir)
                VALUES (?, ?, ?, ?, ?)
                ON CONFLICT(path) DO UPDATE SET
                    hash=excluded.hash,
                    last_modified=excluded.last_modified,
                    doc_ids=excluded.doc_ids,
                    source_dir=excluded.source_dir
                """,
                (
                    file_info.path,
                    file_info.hash,
                    file_info.last_modified,
                    ",".join(doc_ids),
                    file_info.source_dir,
                ),
            )

    def remove_file_state(self, path: str):
        """Remove a file from the state tracking."""
        with sqlite3.connect(self.db_path) as conn:
            conn.execute("DELETE FROM files WHERE path = ?", (path,))


class DataSource(ABC):
    """Abstract base class for data sources."""

    @abstractmethod
    def iter_files(self) -> Iterator[FileInfo]:
        """Yield FileInfo for each file in the source."""
        pass

    @abstractmethod
    def load_file(self, file_info: FileInfo) -> List[LlamaIndexDocument]:
        """Load and return documents for a given file."""
        pass


def _compute_line_offsets(text: str) -> List[int]:
    """Compute character offset positions for each line start.

    Returns a list where line_offsets[i] is the character position where line i+1 starts.
    Line 1 starts at position 0 (implicit).
    """
    offsets = [0]  # Line 1 starts at position 0
    for i, char in enumerate(text):
        if char == '\n':
            offsets.append(i + 1)  # Next line starts after the newline
    return offsets


def _extract_markdown_headings(text: str) -> List[dict]:
    """Extract heading hierarchy from Markdown text using ATX-style syntax.

    Parses # Heading syntax and returns list of dicts with text, position, level.
    Handles ATX-style headings (# Heading) but not Setext (underlined).

    Returns:
        List of dicts with keys: text (str), position (int), level (int)
    """
    import re

    headings = []
    # Match ATX-style headings: line start, 1-6 #s, space, text
    pattern = re.compile(r'^(#{1,6})\s+(.+?)$', re.MULTILINE)

    # Find all code block ranges to skip headings inside them
    code_blocks = []
    for match in re.finditer(r'```.*?```', text, flags=re.DOTALL):
        code_blocks.append((match.start(), match.end()))

    def is_in_code_block(pos):
        """Check if position is inside a code block."""
        return any(start <= pos < end for start, end in code_blocks)

    for match in pattern.finditer(text):
        # Skip headings inside code blocks
        if is_in_code_block(match.start()):
            continue

        level = len(match.group(1))
        heading_text = match.group(2).strip()
        position = match.start()

        if heading_text:
            headings.append({
                "text": heading_text,
                "position": position,
                "level": level
            })

    return headings


def _extract_pdf_headings_from_outline(pdf_path: Path) -> List[dict]:
    """Extract headings from PDF outline/bookmarks (TOC).

    Returns list of dicts with text, position (estimated), level.
    Position is approximate based on cumulative page character counts.
    Falls back to empty list if PDF has no outline or extraction fails.

    Returns:
        List of dicts with keys: text (str), position (int), level (int)
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.warning("pypdf not available, skipping PDF heading extraction")
        return []

    try:
        reader = PdfReader(pdf_path)
        outline = reader.outline

        if not outline:
            return []

        def flatten_outline(items, level=1):
            """Flatten nested outline into list of (title, page_num, level)."""
            results = []
            for item in items:
                if isinstance(item, list):
                    results.extend(flatten_outline(item, level + 1))
                else:
                    page_num = reader.get_destination_page_number(item)
                    results.append((item.title, page_num, level))
            return results

        flat = flatten_outline(outline)
        headings = []
        for title, page_num, level in flat:
            # Estimate position by accumulating text from previous pages
            position = 0
            for page_idx in range(page_num):
                if page_idx < len(reader.pages):
                    position += len(reader.pages[page_idx].extract_text() or "")

            headings.append({
                "text": title.strip(),
                "position": position,
                "level": level
            })

        return headings

    except Exception as e:
        logger.warning(f"Failed to extract PDF outline from {pdf_path}: {e}")
        return []


class LocalFileSystemSource(DataSource):
    """Implementation of DataSource for the local file system with filtering."""

    def __init__(self, config: DirectoryConfig):
        self.config = config
        self.base_dir = config.path

    def is_available(self) -> bool:
        """Check if the directory is available and accessible."""
        try:
            if not self.base_dir.exists():
                return False
            if not self.base_dir.is_dir():
                return False
            # Try to list directory to verify access (important for network mounts)
            next(self.base_dir.iterdir(), None)
            return True
        except (OSError, PermissionError):
            return False

    def _matches_patterns(self, file_path: Path) -> bool:
        """Check if file matches include patterns and doesn't match exclude patterns.

        Uses Path.match() which supports ** glob patterns for directory matching.
        For directory exclusion patterns like **/*venv*/**, checks each path component.
        """
        import fnmatch

        try:
            rel_path = file_path.relative_to(self.base_dir)
        except ValueError:
            rel_path = Path(file_path.name)

        # Check exclude patterns first
        for pattern in self.config.exclude:
            # Handle directory exclusion patterns (e.g., **/*venv*/**, **/node_modules/**)
            # by checking if any directory component matches
            if pattern.startswith('**/') and pattern.endswith('/**'):
                # Extract the directory pattern (e.g., *venv* or node_modules)
                dir_pattern = pattern[3:-3]  # Remove **/ prefix and /** suffix
                for part in rel_path.parts[:-1]:  # Check all directory components (not filename)
                    if fnmatch.fnmatch(part, dir_pattern):
                        return False
            else:
                # Standard pattern matching
                if rel_path.match(pattern) or file_path.name == pattern:
                    return False

        # Check include patterns
        if not self.config.include:
            return True

        for pattern in self.config.include:
            # Path.match() supports ** for recursive directory matching
            if rel_path.match(pattern) or file_path.match(pattern):
                return True

        return False

    def iter_files(self) -> Iterator[FileInfo]:
        """Yield FileInfo for each matching file in the source."""
        if self.config.recursive:
            walker = os.walk(self.base_dir)
        else:
            # Non-recursive: only top-level files
            try:
                top_files = [
                    f for f in self.base_dir.iterdir() if f.is_file()
                ]
                walker = [(str(self.base_dir), [], [f.name for f in top_files])]
            except OSError as e:
                logger.warning(f"Could not list directory {self.base_dir}: {e}")
                return

        for root, _, files in walker:
            for file in files:
                file_path = Path(root) / file

                # Check patterns
                if not self._matches_patterns(file_path):
                    continue

                try:
                    yield self._create_file_info(file_path)
                except (OSError, IOError) as e:
                    logger.warning(f"Could not access file {file_path}: {e}")
                    continue

    def _create_file_info(self, file_path: Path) -> FileInfo:
        """Create FileInfo with source directory context."""
        hasher = hashlib.md5()
        with open(file_path, "rb") as f:
            buf = f.read(65536)
            while len(buf) > 0:
                hasher.update(buf)
                buf = f.read(65536)

        return FileInfo(
            path=str(file_path.absolute()),
            hash=hasher.hexdigest(),
            last_modified=file_path.stat().st_mtime,
            source_dir=str(self.base_dir.absolute()),
        )

    def load_file(self, file_info: FileInfo) -> List[LlamaIndexDocument]:
        file_path = Path(file_info.path)
        if file_path.suffix.lower() == ".docx":
            return split_docx_into_heading_documents(file_path)
        elif file_path.suffix.lower() == ".doc":
            # Convert .doc to .docx using LibreOffice, then process
            docx_path = _convert_doc_to_docx(file_path)
            if docx_path is None:
                logger.warning(f"Skipping {file_path}: could not convert .doc to .docx")
                return []
            try:
                docs = split_docx_into_heading_documents(docx_path)
                # Update metadata to point to original .doc file
                for doc in docs:
                    doc.metadata["file_path"] = str(file_path)
                    doc.metadata["file_name"] = file_path.name
                    if "source" in doc.metadata:
                        doc.metadata["source"] = str(file_path)
                return docs
            finally:
                # Clean up temp file
                if docx_path.exists():
                    docx_path.unlink()
        else:
            reader = SimpleDirectoryReader(
                input_files=[str(file_path)],
            )
            docs = reader.load_data()
            # Ensure dates are visible to LLM (remove from exclusion list)
            for doc in docs:
                if hasattr(doc, 'excluded_llm_metadata_keys') and doc.excluded_llm_metadata_keys:
                    doc.excluded_llm_metadata_keys = [
                        k for k in doc.excluded_llm_metadata_keys
                        if k not in ('creation_date', 'last_modified_date')
                    ]

            # Add line offsets for text-based files (markdown, txt) to enable line number lookup
            if file_path.suffix.lower() in {".md", ".txt"}:
                for doc in docs:
                    text = doc.get_content()
                    line_offsets = _compute_line_offsets(text)
                    doc.metadata["line_offsets"] = line_offsets

                    # Extract headings for Markdown and store separately
                    # (not in metadata to avoid SentenceSplitter size validation)
                    if file_path.suffix.lower() == ".md":
                        headings = _extract_markdown_headings(text)
                        get_heading_store().set_headings(str(file_path), headings)

            # Extract headings for PDF files and store separately
            if file_path.suffix.lower() == ".pdf":
                headings = _extract_pdf_headings_from_outline(file_path)
                get_heading_store().set_headings(str(file_path), headings)

            # Apply metadata exclusions
            for doc in docs:
                doc.excluded_embed_metadata_keys = EXCLUDED_EMBED_METADATA_KEYS
                doc.excluded_llm_metadata_keys = EXCLUDED_LLM_METADATA_KEYS

            return docs


class MultiDirectoryDataSource(DataSource):
    """Aggregates multiple LocalFileSystemSource instances."""

    def __init__(self, config: IndexConfig):
        self.config = config
        self.sources: List[LocalFileSystemSource] = []
        self.unavailable_dirs: List[DirectoryConfig] = []

        for dir_config in config.directories:
            if not dir_config.enabled:
                logger.info(f"Skipping disabled directory: {dir_config.path}")
                continue

            source = LocalFileSystemSource(dir_config)

            if source.is_available():
                self.sources.append(source)
                logger.info(f"Added directory source: {dir_config.path}")
            else:
                self.unavailable_dirs.append(dir_config)
                logger.warning(f"Directory unavailable, skipping: {dir_config.path}")

    def iter_files(self) -> Iterator[FileInfo]:
        """Iterate over files from all available sources."""
        seen_paths: Set[str] = set()

        for source in self.sources:
            for file_info in source.iter_files():
                # Deduplicate in case of overlapping mounts
                if file_info.path not in seen_paths:
                    seen_paths.add(file_info.path)
                    yield file_info

    def load_file(self, file_info: FileInfo) -> List[LlamaIndexDocument]:
        """Load file using the appropriate source based on source_dir."""
        # Find the source that owns this file
        for source in self.sources:
            if file_info.source_dir == str(source.base_dir.absolute()):
                return source.load_file(file_info)

        # Fallback: use first source (shouldn't happen normally)
        if self.sources:
            return self.sources[0].load_file(file_info)

        raise ValueError(f"No source available for file: {file_info.path}")

    def get_summary(self) -> Dict[str, Any]:
        """Return summary of configured directories."""
        return {
            "available": [str(s.base_dir) for s in self.sources],
            "unavailable": [str(d.path) for d in self.unavailable_dirs],
            "total_sources": len(self.sources),
        }


class SimpleProgressBar:
    """Lightweight progress bar using only the standard library."""

    def __init__(self, total: int, desc: str, unit: str = "item", width: int = 30):
        self.total = max(total, 0)
        self.desc = desc
        self.unit = unit
        self.width = width
        self.current = 0
        if self.total > 0:
            self._render()

    def update(self, step: int = 1) -> None:
        if self.total <= 0:
            return
        self.current = min(self.total, self.current + step)
        self._render()
        if self.current >= self.total:
            sys.stdout.write("\n")
            sys.stdout.flush()

    def _render(self) -> None:
        progress = self.current / self.total if self.total else 0
        filled = int(self.width * progress)
        bar = "#" * filled + "-" * (self.width - filled)
        sys.stdout.write(
            f"\r{self.desc} [{bar}] {progress * 100:5.1f}% ({self.current}/{self.total} {self.unit}s)"
        )
        sys.stdout.flush()


class Spinner:
    """Simple console spinner to indicate long-running steps."""

    def __init__(self, desc: str, interval: float = 0.1):
        self.desc = desc
        self.interval = interval
        self._stop_event = threading.Event()
        self._thread: threading.Thread | None = None
        self._line = desc

    def __enter__(self):
        self._thread = threading.Thread(target=self._spin, daemon=True)
        self._thread.start()
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self._stop_event.set()
        if self._thread:
            self._thread.join()
        # Clear spinner line
        sys.stdout.write("\r" + " " * len(self._line) + "\r")
        sys.stdout.flush()

    def _spin(self) -> None:
        for char in itertools.cycle("|/-\\"):
            if self._stop_event.is_set():
                break
            self._line = f"{self.desc} {char}"
            sys.stdout.write("\r" + self._line)
            sys.stdout.flush()
            time.sleep(self.interval)


def _embedding_cache_path(model_name: str, cache_dir: Path) -> Path:
    """Return the expected cache directory for a FastEmbed model."""
    return cache_dir / f"models--{model_name.replace('/', '--')}"


def _verify_model_cache_exists(cache_dir: Path) -> bool:
    """
    Verify that the cached model directory exists and contains the expected model files.
    """
    from fastembed import TextEmbedding

    try:
        models = TextEmbedding.list_supported_models()
        model_info = [m for m in models if m.get("model") == RETRIEVAL_EMBED_MODEL_NAME]
        if not model_info:
            return False

        model_info = model_info[0]
        hf_source = model_info.get("sources", {}).get("hf")
        if not hf_source:
            return False

        expected_dir = cache_dir / f"models--{hf_source.replace('/', '--')}"
        if not expected_dir.exists():
            return False

        snapshots_dir = expected_dir / "snapshots"
        if not snapshots_dir.exists():
            return False

        model_file = model_info.get("model_file", "model_optimized.onnx")
        for snapshot in snapshots_dir.iterdir():
            if snapshot.is_dir():
                model_path = snapshot / model_file
                if model_path.exists() or model_path.is_symlink():
                    return True

        return False
    except Exception:
        return False


def _get_cached_model_path(cache_dir: Path, model_name: str) -> Path | None:
    """Get the cached model directory path."""
    try:
        from huggingface_hub import snapshot_download
        from fastembed import TextEmbedding
        models = TextEmbedding.list_supported_models()
        model_info = [m for m in models if m.get("model") == model_name]
        if model_info:
            hf_source = model_info[0].get("sources", {}).get("hf")
            if hf_source:
                cache_dir_abs = cache_dir.resolve()
                model_dir = snapshot_download(
                    repo_id=hf_source,
                    local_files_only=True,
                    cache_dir=str(cache_dir_abs)
                )
                return Path(model_dir).resolve()
    except (ImportError, Exception):
        pass
    return None


def _create_fastembed_embedding(cache_dir: Path, offline: bool = False):
    """Create a FastEmbedEmbedding instance."""
    if offline:
        cached_model_path = _get_cached_model_path(cache_dir, RETRIEVAL_EMBED_MODEL_NAME)
        if cached_model_path:
            logger.info(
                f"Using cached model path to bypass download: {cached_model_path}"
            )
            return FastEmbedEmbedding(
                model_name=RETRIEVAL_EMBED_MODEL_NAME,
                cache_dir=str(cache_dir),
                specific_model_path=str(cached_model_path)
            )
        else:
            logger.warning(
                "Could not find cached model path, falling back to normal initialization"
            )

    return FastEmbedEmbedding(
        model_name=RETRIEVAL_EMBED_MODEL_NAME, cache_dir=str(cache_dir)
    )


def ensure_embedding_model_cached(cache_dir: Path, offline: bool = False) -> None:
    """Ensure the embedding model is available in the local cache."""
    if offline:
        logger.info("Verifying embedding model cache...")
        if _verify_model_cache_exists(cache_dir):
            logger.info("Embedding model found in cache")
        else:
            logger.error(
                "Offline mode enabled, but embedding model cache not found in %s",
                cache_dir,
            )
            raise FileNotFoundError(
                f"Embedding model '{RETRIEVAL_EMBED_MODEL_NAME}' not found in cache directory '{cache_dir}'. "
            )

    try:
        logger.info("Initializing embedding model from cache...")
        cache_dir_abs = cache_dir.resolve()
        if offline:
            os.environ["HF_HUB_CACHE"] = str(cache_dir_abs)

        _create_fastembed_embedding(cache_dir, offline=offline)
        logger.info("Embedding model initialized successfully")
        return
    except (ValueError, Exception) as e:
        # Simplified error handling for brevity, similar logic as original
        if offline:
            raise FileNotFoundError(f"Failed to load model offline: {e}") from e
        else:
            raise RuntimeError(f"Failed to download/initialize model: {e}") from e


def ensure_rerank_model_cached(cache_dir: Path, offline: bool = False) -> Path:
    """Ensure the reranking model is cached locally."""
    try:
        from flashrank import Ranker
    except ImportError as exc:
        raise ImportError(
            "flashrank is required for reranking."
        ) from exc

    cache_dir_abs = cache_dir.resolve()
    logger.info("Ensuring rerank model is available in cache...")

    # Map cross-encoder model names to FlashRank equivalents if needed
    model_name = RETRIEVAL_RERANK_MODEL_NAME
    # Note: FlashRank doesn't have L-6 models, so we map to L-12 equivalents
    model_mapping = {
        "cross-encoder/ms-marco-MiniLM-L-6-v2": "ms-marco-MiniLM-L-12-v2",  # L-6 not available, use L-12
        "ms-marco-MiniLM-L-6-v2": "ms-marco-MiniLM-L-12-v2",  # Direct mapping for L-6
    }
    if model_name in model_mapping:
        model_name = model_mapping[model_name]
    elif model_name.startswith("cross-encoder/"):
        # Extract model name after cross-encoder/ prefix and try to map
        base_name = model_name.replace("cross-encoder/", "")
        # If it's an L-6 model, map to L-12
        if "L-6" in base_name:
            model_name = base_name.replace("L-6", "L-12")
        else:
            model_name = base_name

    try:
        reranker = Ranker(model_name=model_name, cache_dir=str(cache_dir_abs))
        logger.info(f"FlashRank model '{model_name}' initialized successfully")
        return cache_dir_abs
    except Exception as exc:
        if offline:
            raise FileNotFoundError(
                f"Rerank model '{model_name}' not found in cache."
            ) from exc
        raise


def _parse_heading_level(style_name: str | None) -> int:
    """Best-effort extraction of a numeric heading level from a DOCX style name."""
    if not style_name:
        return 1
    try:
        if "Heading" in style_name:
            level_str = style_name.replace("Heading", "").strip()
            if level_str:
                return int(level_str)
    except (ValueError, AttributeError):
        pass
    return 1


def _get_doc_temp_dir() -> Path:
    """Get the temporary directory for .doc conversion, creating it if needed."""
    storage_dir = Path(cfgload.get("storage.storage_dir", "./storage"))
    temp_dir = storage_dir / "doc_temp"
    temp_dir.mkdir(parents=True, exist_ok=True)
    return temp_dir


def _convert_doc_to_docx(doc_path: Path) -> Path | None:
    """Convert a .doc file to .docx using LibreOffice.

    Returns path to temporary .docx file, or None if conversion fails.
    Caller is responsible for cleaning up the temp file.
    """
    import subprocess
    import shutil

    # Find LibreOffice executable
    soffice_paths = [
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",  # macOS
        "/usr/bin/soffice",  # Linux
        "/usr/bin/libreoffice",  # Linux alternative
        "soffice",  # Windows (in PATH)
    ]

    soffice = None
    for path in soffice_paths:
        if shutil.which(path):
            soffice = path
            break

    if not soffice:
        logger.warning(f"LibreOffice not found. Cannot convert {doc_path}")
        return None

    # Use storage directory for temp files (more reliable space than /tmp)
    temp_dir = _get_doc_temp_dir()

    try:
        result = subprocess.run(
            [soffice, "--headless", "--convert-to", "docx",
             "--outdir", str(temp_dir), str(doc_path)],
            capture_output=True,
            timeout=60,
        )
        if result.returncode != 0:
            logger.warning(f"LibreOffice conversion failed for {doc_path}: {result.stderr}")
            return None

        # Find the converted file
        docx_name = doc_path.stem + ".docx"
        docx_path = temp_dir / docx_name
        if docx_path.exists():
            return docx_path

        logger.warning(f"Converted file not found: {docx_path}")
    except subprocess.TimeoutExpired:
        logger.warning(f"LibreOffice conversion timed out for {doc_path}")
    except Exception as e:
        logger.warning(f"Error converting {doc_path}: {e}")

    return None


def split_docx_into_heading_documents(docx_path: Path) -> List[LlamaIndexDocument]:
    """Split DOCX into documents by heading."""
    docs: List[LlamaIndexDocument] = []
    try:
        doc = Document(docx_path)
    except Exception as e:
        logger.warning(f"Failed to open DOCX {docx_path}: {e}")
        return docs

    # Extract file dates from filesystem
    stat = docx_path.stat()
    creation_date = datetime.fromtimestamp(stat.st_ctime).strftime("%Y-%m-%d")
    last_modified_date = datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d")

    # Try to extract dates from DOCX core properties (more accurate than filesystem)
    try:
        core_props = doc.core_properties
        if core_props.created:
            creation_date = core_props.created.strftime("%Y-%m-%d")
        if core_props.modified:
            last_modified_date = core_props.modified.strftime("%Y-%m-%d")
    except Exception:
        pass  # Fall back to filesystem dates

    # First pass: Extract all headings with positions for hierarchy metadata
    all_headings = []
    char_position = 0
    for para in doc.paragraphs:
        style_name = getattr(para.style, "name", "") or ""
        is_heading = (
            style_name.startswith("Heading")
            or style_name.startswith("heading")
            or "Heading" in style_name
        )

        if is_heading and para.text.strip():
            heading_level = _parse_heading_level(style_name)
            all_headings.append({
                "text": para.text.strip(),
                "position": char_position,
                "level": heading_level
            })

        char_position += len(para.text) + 1  # +1 for newline

    # Store headings separately to avoid metadata size issues during chunking
    get_heading_store().set_headings(str(docx_path), all_headings)

    # Second pass: Split by heading (existing logic)
    current_heading: str | None = None
    current_level: int | None = None
    current_body: list[str] = []

    def flush_current():
        if not current_heading:
            return
        text = "\n".join(line for line in current_body if line is not None).strip()
        if not text:
            return

        # Build hierarchical heading_path by finding parent headings based on level
        heading_path = []
        if all_headings:
            # Find the index of the current heading in all_headings
            current_idx = None
            for idx, h in enumerate(all_headings):
                if h["text"] == current_heading and h["level"] == current_level:
                    current_idx = idx
                    break

            if current_idx is not None:
                # Build path by including all parent headings (those with lower level numbers)
                # Walk backwards from current heading and include headings with level < current_level
                path_headings = [all_headings[current_idx]]  # Start with current
                for idx in range(current_idx - 1, -1, -1):
                    h = all_headings[idx]
                    if h["level"] < path_headings[0]["level"]:
                        path_headings.insert(0, h)
                heading_path = [h["text"] for h in path_headings]

        metadata = {
            "file_path": str(docx_path),
            "file_name": docx_path.name,
            "source": str(docx_path),
            "heading": current_heading,
            "heading_level": current_level,
            "creation_date": creation_date,
            "last_modified_date": last_modified_date,
            "heading_path": heading_path,  # Pre-computed hierarchical path
        }
        docs.append(LlamaIndexDocument(
            text=text,
            metadata=metadata,
            excluded_embed_metadata_keys=EXCLUDED_EMBED_METADATA_KEYS,
            excluded_llm_metadata_keys=EXCLUDED_LLM_METADATA_KEYS,
        ))

    for para in doc.paragraphs:
        style_name = getattr(para.style, "name", "") or ""
        is_heading = (
            style_name.startswith("Heading")
            or style_name.startswith("heading")
            or "Heading" in style_name
        )

        if is_heading and para.text.strip():
            flush_current()
            current_heading = para.text.strip()
            current_level = _parse_heading_level(style_name)
            current_body = []
        else:
            if current_heading is not None:
                current_body.append(para.text)

    flush_current()

    if not docs:
        try:
            full_text = "\n".join(p.text for p in doc.paragraphs).strip()
        except Exception:
            full_text = ""

        if full_text:
            metadata = {
                "file_path": str(docx_path),
                "file_name": docx_path.name,
                "source": str(docx_path),
                "heading": None,
                "heading_level": None,
                "creation_date": creation_date,
                "last_modified_date": last_modified_date,
            }
            docs.append(LlamaIndexDocument(
                text=full_text,
                metadata=metadata,
                excluded_embed_metadata_keys=EXCLUDED_EMBED_METADATA_KEYS,
                excluded_llm_metadata_keys=EXCLUDED_LLM_METADATA_KEYS,
            ))

    logger.info(
        f"Split DOCX {docx_path} into {len(docs)} heading-based document(s)"
    )
    return docs


def tokenize_filename(filename: str) -> List[str]:
    """
    Tokenize a filename for BM25 indexing.

    Splits on delimiters (underscore, hyphen, dot, space) and camelCase.

    Examples:
        'cpp_styleguide.md' -> ['cpp', 'styleguide', 'md']
        'API-Reference-v2.pdf' -> ['api', 'reference', 'v2', 'pdf']
        'CamelCaseDoc.docx' -> ['camel', 'case', 'doc', 'docx']
    """
    import re

    name_parts = filename.rsplit('.', 1)
    base_name = name_parts[0]
    extension = name_parts[1] if len(name_parts) > 1 else ""

    # Split on explicit delimiters
    parts = re.split(r'[_\-\.\s]+', base_name)

    # Split camelCase within each part
    tokens = []
    for part in parts:
        camel_split = re.sub(r'([a-z])([A-Z])', r'\1 \2', part).split()
        tokens.extend(t.lower() for t in camel_split if t)

    # Add extension as a token
    if extension:
        tokens.append(extension.lower())

    return tokens


def build_bm25_index(index, storage_dir: Path) -> None:
    """
    Build a BM25 index over file names from the docstore.

    This enables keyword matching for queries like 'cpp styleguide' to find
    files named 'cpp_styleguide.md'.
    """
    from llama_index.retrievers.bm25 import BM25Retriever
    from llama_index.core.schema import TextNode

    logger.info("Building BM25 index for file name matching...")

    # Create filename nodes - one per unique file
    filename_nodes = []
    seen_files: Set[str] = set()

    for doc_id, node in index.docstore.docs.items():
        metadata = node.metadata or {}
        file_name = metadata.get("file_name", "")
        file_path = metadata.get("file_path", "")

        if not file_name or file_path in seen_files:
            continue
        seen_files.add(file_path)

        tokens = tokenize_filename(file_name)
        filename_nodes.append(TextNode(
            text=" ".join(tokens),
            metadata={"file_name": file_name, "file_path": file_path},
            id_=f"bm25_{file_path}"
        ))

    if not filename_nodes:
        logger.warning("No documents found for BM25 indexing")
        return

    logger.info(f"Creating BM25 index with {len(filename_nodes)} file name entries")

    bm25_retriever = BM25Retriever.from_defaults(
        nodes=filename_nodes,
        similarity_top_k=10,
    )

    bm25_dir = storage_dir / "bm25_index"
    bm25_dir.mkdir(parents=True, exist_ok=True)
    bm25_retriever.persist(str(bm25_dir))

    logger.info(f"BM25 index persisted to {bm25_dir}")


def configure_offline_mode(offline: bool, cache_dir: Path) -> None:
    """Configure environment variables for offline mode."""
    if offline:
        os.environ["HF_HUB_OFFLINE"] = "1"
        os.environ["TRANSFORMERS_OFFLINE"] = "1"
        os.environ["HF_DATASETS_OFFLINE"] = "1"
        cache_dir_abs = cache_dir.resolve()
        os.environ["HF_HOME"] = str(cache_dir_abs)
        os.environ["HF_HUB_CACHE"] = str(cache_dir_abs)
        os.environ["HF_DATASETS_CACHE"] = str(cache_dir_abs)
        logger.info("Offline mode enabled.")
    else:
        # Clear offline mode environment variables to allow downloads
        for var in ["HF_HUB_OFFLINE", "TRANSFORMERS_OFFLINE", "HF_DATASETS_OFFLINE"]:
            os.environ.pop(var, None)

    # Update huggingface_hub's cached constant (it caches at import time)
    try:
        from huggingface_hub import constants
        constants.HF_HUB_OFFLINE = offline
    except ImportError:
        pass


def build_index(
    download_only: bool = False,
    config_path: Path | None = None,
    model_cache_dir: Path | None = None,
) -> None:
    """Build and persist the vector index incrementally."""
    global _config, STORAGE_DIR, STATE_DB_PATH, RETRIEVAL_MODEL_CACHE_DIR, BM25_INDEX_DIR, HEADING_STORE_PATH
    global RETRIEVAL_EMBED_MODEL_NAME, RETRIEVAL_RERANK_MODEL_NAME

    if config_path:
        cfg = load_config(config_path)
        _config = cfg
        STORAGE_DIR = Path(cfg["storage"]["storage_dir"])
        STATE_DB_PATH = STORAGE_DIR / "ingestion_state.db"
        RETRIEVAL_MODEL_CACHE_DIR = Path(cfg["storage"]["model_cache_dir"])
        BM25_INDEX_DIR = STORAGE_DIR / "bm25_index"
        HEADING_STORE_PATH = STORAGE_DIR / "heading_store.json"
        RETRIEVAL_EMBED_MODEL_NAME = cfg["retrieval"]["embed_model_name"]
        RETRIEVAL_RERANK_MODEL_NAME = cfg["retrieval"]["rerank_model_name"]

    # Override model cache dir if specified via CLI
    if model_cache_dir:
        RETRIEVAL_MODEL_CACHE_DIR = model_cache_dir

    # Read offline setting from config; force online when downloading models
    offline = False if download_only else _config["retrieval"].get("offline", False)
    cache_dir = RETRIEVAL_MODEL_CACHE_DIR
    configure_offline_mode(offline, cache_dir)

    # Load configuration
    index_config = load_index_config()
    logger.info(f"Indexing configured with {len(index_config.directories)} directories")

    ensure_embedding_model_cached(cache_dir, offline=offline)
    try:
        ensure_rerank_model_cached(cache_dir, offline=offline)
    except FileNotFoundError:
        if download_only or offline:
            raise
        logger.warning("Rerank model could not be cached yet; continuing without it.")

    if download_only:
        logger.info("Models downloaded; skipping index build.")
        return

    # Initialize State and Multi-Directory Data Source
    ingestion_state = IngestionState(STATE_DB_PATH)
    data_source = MultiDirectoryDataSource(index_config)

    # Log directory summary
    summary = data_source.get_summary()
    logger.info(f"Active directories: {summary['available']}")
    if summary['unavailable']:
        logger.warning(f"Unavailable directories (skipped): {summary['unavailable']}")

    if not data_source.sources:
        logger.error("No available directories to index. Check your config.yaml indexing.directories.")
        return

    # Initialize Embedding Model
    logger.info(f"Initializing embedding model: {RETRIEVAL_EMBED_MODEL_NAME}")
    with Spinner("Initializing embedding model"):
        embed_model = _create_fastembed_embedding(RETRIEVAL_MODEL_CACHE_DIR, offline=offline)
    Settings.embed_model = embed_model

    # Configure Text Splitter using config values
    text_splitter = SentenceSplitter(
        chunk_size=index_config.chunk_size,
        chunk_overlap=index_config.chunk_overlap,
        separator=" ",
    )
    Settings.text_splitter = text_splitter

    # Load existing index or create new
    if (STORAGE_DIR / "docstore.json").exists():
        logger.info("Loading existing index context...")
        storage_context = StorageContext.from_defaults(persist_dir=str(STORAGE_DIR))
        index = load_index_from_storage(storage_context, embed_model=embed_model)
    else:
        logger.info("Creating new index context...")
        storage_context = StorageContext.from_defaults()
        index = VectorStoreIndex([], storage_context=storage_context, embed_model=embed_model)

    # Change Detection
    tracked_files = ingestion_state.get_all_files()
    found_files: Set[str] = set()
    files_to_process: List[FileInfo] = []

    logger.info("Scanning for changes...")
    for file_info in data_source.iter_files():
        found_files.add(file_info.path)
        existing_state = tracked_files.get(file_info.path)

        if existing_state:
            # Check if modified
            if existing_state["hash"] != file_info.hash:
                logger.info(f"Modified file detected: {file_info.path}")
                files_to_process.append(file_info)
        else:
            # New file
            logger.info(f"New file detected: {file_info.path}")
            files_to_process.append(file_info)

    # Identify Deleted Files
    deleted_files = set(tracked_files.keys()) - found_files
    for deleted_path in deleted_files:
        logger.info(f"Deleted file detected: {deleted_path}")
        doc_ids = tracked_files[deleted_path]["doc_ids"]
        for doc_id in doc_ids:
            try:
                index.delete_ref_doc(doc_id, delete_from_docstore=True)
            except Exception as e:
                logger.warning(f"Failed to delete doc {doc_id} from index: {e}")
        # Clean up heading data for deleted file
        get_heading_store().remove_headings(deleted_path)
        ingestion_state.remove_file_state(deleted_path)

    if not files_to_process and not deleted_files:
        logger.info("No changes detected. Index is up to date.")
        return

    # Process New/Modified Files
    if files_to_process:
        progress = SimpleProgressBar(len(files_to_process), desc="Processing files", unit="file")
        for file_info in files_to_process:
            # Remove old versions if they exist
            existing_state = tracked_files.get(file_info.path)
            if existing_state:
                for doc_id in existing_state["doc_ids"]:
                    try:
                        index.delete_ref_doc(doc_id, delete_from_docstore=True)
                    except KeyError:
                        pass # Document might already be gone

            # Load and Index New Version
            docs = data_source.load_file(file_info)
            doc_ids = []
            for doc in docs:
                index.insert(doc)
                doc_ids.append(doc.doc_id)

            # Update State
            ingestion_state.update_file_state(file_info, doc_ids)
            progress.update()

    # Persist Index
    STORAGE_DIR.mkdir(parents=True, exist_ok=True)
    logger.info(f"Persisting index to {STORAGE_DIR}")
    index.storage_context.persist(persist_dir=str(STORAGE_DIR))

    # Build BM25 index for file name matching
    build_bm25_index(index, STORAGE_DIR)

    logger.info("Indexing complete.")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Build the document index")
    parser.add_argument(
        "--download-models",
        action="store_true",
        help="Download the retrieval models and exit",
    )
    parser.add_argument(
        "--config",
        type=str,
        help="Path to config.yaml (overrides auto-discovery)",
    )
    parser.add_argument(
        "--model-cache-dir",
        type=str,
        help="Directory to download/cache models (overrides config)",
    )
    args = parser.parse_args()

    try:
        build_index(
            download_only=args.download_models,
            config_path=Path(args.config) if args.config else None,
            model_cache_dir=Path(args.model_cache_dir) if args.model_cache_dir else None,
        )
    except Exception as e:
        logger.error(f"Indexing failed: {e}", exc_info=True)
        raise
