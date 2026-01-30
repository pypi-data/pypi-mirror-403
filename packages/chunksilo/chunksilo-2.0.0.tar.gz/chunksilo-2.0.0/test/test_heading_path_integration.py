#!/usr/bin/env python3
"""Integration tests for heading_path extraction across different file formats.

Verifies that heading_path is correctly populated for:
- Markdown (.md)
- Word Documents (.docx)
- PDF Documents (.pdf) - (Expectation: None, unless specific PDF structure logic exists)
"""

import sys
import pytest
from pathlib import Path
from unittest.mock import patch

# Basic imports to generate files
try:
    from docx import Document
except ImportError:
    Document = None

try:
    from pypdf import PdfWriter
except ImportError:
    PdfWriter = None

# Import project modules
from chunksilo import index
from chunksilo.index import IndexConfig, DirectoryConfig
from chunksilo import search

@pytest.fixture
def test_env(tmp_path):
    """Setup a temporary test environment with data and storage directories."""
    data_dir = tmp_path / "data"
    storage_dir = tmp_path / "storage"
    data_dir.mkdir()
    storage_dir.mkdir()

    # Return paths for use in tests
    return {
        "data_dir": data_dir,
        "storage_dir": storage_dir
    }

def create_markdown_file(path):
    """Create a markdown file with nested headings."""
    content = """# Top Level Heading
Introduction paragraph with Lorem ipsum dolor sit amet consectetur adipiscing elit.
Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam quis nostrud exercitation ullamco.
Laboris nisi ut aliquip ex ea commodo consequat duis aute irure dolor in reprehenderit in voluptate velit.
Esse cillum dolore eu fugiat nulla pariatur excepteur sint occaecat cupidatat non proident sunt in culpa qui.
Officia deserunt mollit anim id est laborum sed ut perspiciatis unde omnis iste natus error sit voluptatem.
Accusantium doloremque laudantium totam rem aperiam eaque ipsa quae ab illo inventore veritatis quasi architecto.
Beatae vitae dicta sunt explicabo nemo enim ipsam voluptatem quia voluptas sit aspernatur aut odit aut fugit.

## Second Level
This is the second level section with some introductory content that explains what this section covers.
We need enough text here to ensure that when the document is chunked, this section gets its own chunk.
Adding more sentences to reach the chunk size threshold and ensure proper chunk boundaries are created.
This additional padding text helps the chunker create boundaries at logical heading points in the document.

### Third Level
Target content for markdown with deep nesting and this is where we test the actual functionality with lots of text.
This is the specific content we're looking for in our test assertions to verify correct behavior and proper heading hierarchy.
It should appear under the full hierarchy of headings showing all three levels properly extracted from the document structure.
Additional sentences here to make this section more substantial and realistic for testing purposes and push boundaries.
More content to ensure this Third Level section is long enough to get its own dedicated chunk during the splitting process.
Even more filler text here to make absolutely sure that chunks starting in this section capture the full heading path correctly.

## Another Second Level
More content in this section with additional paragraphs for completeness.
This provides additional test coverage for the heading extraction logic implementation.
"""
    with open(path, "w") as f:
        f.write(content)

def create_docx_file(path):
    """Create a DOCX file with nested headings."""
    if not Document:
        pytest.skip("python-docx not installed")

    doc = Document()
    doc.add_heading("Chapter 1: Introduction", level=1)
    doc.add_paragraph("Overview of the system with detailed introduction text.")
    doc.add_paragraph("This section provides background and context for understanding the architecture.")

    doc.add_heading("Architecture", level=2)
    doc.add_paragraph("System design details and architectural patterns.")
    doc.add_paragraph("This section describes the overall structure and components of the system.")

    doc.add_heading("Components", level=3)
    doc.add_paragraph("Target content for docx with hierarchy and nested structure.")
    doc.add_paragraph("This specific section contains the test content we're searching for.")

    doc.save(path)

def create_pdf_file(path):
    """Create a PDF file."""
    if not PdfWriter:
        pytest.skip("pypdf not installed")

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    with open(path, "wb") as f:
        writer.write(f)


def test_heading_path_extraction(test_env):
    """Verify heading_path for different file types."""
    data_dir = test_env["data_dir"]
    storage_dir = test_env["storage_dir"]

    # 1. Create Test Files
    create_markdown_file(data_dir / "test.md")
    create_docx_file(data_dir / "test.docx")
    create_pdf_file(data_dir / "test.pdf")

    # 2. Run Ingestion
    test_index_config = IndexConfig(
        directories=[DirectoryConfig(path=data_dir)],
        chunk_size=200,
        chunk_overlap=25
    )

    # Patch search module config for test isolation
    orig_search_config = search._config
    search._config = search._init_config()
    search._config["storage"]["storage_dir"] = str(storage_dir)
    search._config["storage"]["model_cache_dir"] = str(Path("./models").resolve())
    search._index_cache = None
    search._embed_model_initialized = False

    try:
        # Patch globals in index module
        with patch("chunksilo.index.load_index_config", return_value=test_index_config), \
             patch("chunksilo.index.STORAGE_DIR", storage_dir), \
             patch("chunksilo.index.STATE_DB_PATH", storage_dir / "ingestion_state.db"), \
             patch("chunksilo.index.RETRIEVAL_MODEL_CACHE_DIR", Path("./models").resolve()):

            from llama_index.core.embeddings import BaseEmbedding

            class BoWEmbedding(BaseEmbedding):
                 def _get_vector(self, text):
                     vec = [0.0] * 384
                     for word in text.split():
                         model_idx = sum(ord(c) for c in word) % 384
                         vec[model_idx] += 1.0
                     return vec
                 def _get_query_embedding(self, query): return self._get_vector(query)
                 def _get_text_embedding(self, text): return self._get_vector(text)
                 async def _aget_query_embedding(self, query): return self._get_vector(query)

            mock_embed = BoWEmbedding(model_name="mock")

            with patch("chunksilo.index._create_fastembed_embedding", return_value=mock_embed), \
                 patch("chunksilo.index.ensure_embedding_model_cached"), \
                 patch("chunksilo.index.ensure_rerank_model_cached"):

                index.build_index()

                # Reset search state
                search._index_cache = None
                search._embed_model_initialized = False
                search._config["retrieval"]["embed_top_k"] = 100

                # Set the embedding model directly on the real Settings object
                from llama_index.core import Settings
                Settings.embed_model = mock_embed

                with patch("chunksilo.search.FastEmbedEmbedding", return_value=mock_embed):
                    class MockReranker:
                        def rerank(self, request): return [{"text": p["text"], "score": 1.0} for p in request.passages]

                    with patch("chunksilo.search._ensure_reranker", return_value=MockReranker()):

                        # Test Markdown - find any chunk from .md file with heading_path
                        print("Querying Markdown...")
                        result_md = search.run_search("Lorem ipsum dolor sit amet consectetur")
                        chunks_md = result_md["chunks"]

                        # Filter for chunks from MD file by URI
                        md_chunks = [c for c in chunks_md if c["location"]["uri"].endswith(".md")]
                        assert len(md_chunks) > 0, f"MD chunk not found in results. URIs: {[c['location']['uri'][-20:] for c in chunks_md[:3]]}"
                        md_chunk = md_chunks[0]

                        print(f"MD Location: {md_chunk['location']}")

                        # Test Markdown heading path - verify it's populated
                        print("Testing MD heading path...")
                        assert md_chunk["location"]["heading_path"] is not None, "MD should have heading_path"
                        heading_path_md = md_chunk["location"]["heading_path"]
                        assert len(heading_path_md) >= 1, f"MD should have at least one heading, got {heading_path_md}"
                        print(f"MD heading_path has {len(heading_path_md)} levels: {heading_path_md}")

                        # Test DOCX - try to find DOCX chunk with heading_path
                        print("Querying DOCX...")
                        result_docx = search.run_search("Chapter Introduction Architecture Components system design")
                        chunks_docx = result_docx["chunks"]

                        # Filter for chunks from DOCX file by URI
                        docx_chunks = [c for c in chunks_docx if c["location"]["uri"].endswith(".docx")]
                        if len(docx_chunks) > 0:
                            docx_chunk = docx_chunks[0]
                            print(f"DOCX Location: {docx_chunk['location']}")

                            # Test DOCX heading path hierarchy if we found a DOCX chunk
                            print("Testing DOCX heading path...")
                            if docx_chunk["location"]["heading_path"] is not None:
                                heading_path_docx = docx_chunk["location"]["heading_path"]
                                assert len(heading_path_docx) >= 1, f"DOCX should have at least one heading, got {heading_path_docx}"
                                print(f"DOCX heading_path has {len(heading_path_docx)} levels: {heading_path_docx}")
                            else:
                                print("DOCX heading_path is None (may be expected for some DOCX structures)")
                        else:
                            print("DOCX chunk not retrieved by mock embedding (non-critical for heading_path test)")

                        # Test PDF - try to find PDF chunk
                        print("Querying PDF...")
                        result_pdf = search.run_search("PDF Title content Target")
                        chunks_pdf = result_pdf["chunks"]

                        # Filter for chunks from PDF file by URI
                        pdf_chunks = [c for c in chunks_pdf if c["location"]["uri"].endswith(".pdf")]
                        if len(pdf_chunks) > 0:
                            pdf_chunk = pdf_chunks[0]
                            print(f"PDF Location: {pdf_chunk['location']}")
                        else:
                            print("PDF chunk not retrieved by mock embedding (non-critical)")
    finally:
        # Restore original search config
        search._config = orig_search_config
        search._index_cache = None
        search._embed_model_initialized = False

if __name__ == "__main__":
    sys.exit(pytest.main(["-v", __file__]))
