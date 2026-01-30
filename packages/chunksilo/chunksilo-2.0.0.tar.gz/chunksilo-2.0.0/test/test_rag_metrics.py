#!/usr/bin/env python3
"""
RAG metrics test suite for system evaluation.

This test suite:
1. Downloads a diverse corpus of documents from the web (PDF, DOCX, Markdown, TXT)
2. Ingests them into the RAG system
3. Tests retrieval accuracy with diverse queries
4. Evaluates using standard RAG metrics (Precision@k, Recall@k, MRR, NDCG)
5. Challenges the models with various query types and edge cases
"""
import json
import logging
import math
import os
import sys
import time
from collections import defaultdict
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple
from urllib.parse import urlparse

import requests
from dotenv import load_dotenv

load_dotenv()

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)

# Import after logging is set up
from chunksilo.index import STORAGE_DIR, build_index
from chunksilo.search import run_search

# Test corpus configuration
TEST_DATA_DIR = Path(os.getenv("TEST_DATA_DIR", "./test_data"))
TEST_STORAGE_DIR = Path(os.getenv("TEST_STORAGE_DIR", "./test_storage"))
TEST_RESULTS_DIR = Path(os.getenv("TEST_RESULTS_DIR", "./test_results"))

# Configuration: abort on download failures
ABORT_ON_DOWNLOAD_FAILURE = os.getenv("ABORT_ON_DOWNLOAD_FAILURE", "1").lower() not in ("0", "false", "no")

# Document sources - diverse corpus from public domains
DOCUMENT_SOURCES = {
    "pdf": [
        # Academic papers and technical documentation
        "https://arxiv.org/pdf/1706.03762.pdf",  # Attention Is All You Need (Transformer paper)
        "https://arxiv.org/pdf/2005.14165.pdf",  # GPT-3 paper
        "https://arxiv.org/pdf/1810.04805.pdf",  # BERT paper
        # Additional academic papers for diversity
        "https://arxiv.org/pdf/2010.11929.pdf",  # Vision Transformer (ViT)
        "https://arxiv.org/pdf/1910.10683.pdf",  # T5: Text-To-Text Transfer Transformer
    ],
    "markdown": [
        # GitHub README files and documentation
        "https://raw.githubusercontent.com/python/cpython/main/README.rst",
        "https://raw.githubusercontent.com/facebook/react/main/README.md",
        "https://raw.githubusercontent.com/microsoft/vscode/main/README.md",
        "https://raw.githubusercontent.com/nodejs/node/main/README.md",
        "https://raw.githubusercontent.com/tensorflow/tensorflow/master/README.md",
    ],
    "txt": [
        # Plain text documents
        "https://www.gutenberg.org/files/1342/1342-0.txt",  # Pride and Prejudice
        "https://www.gutenberg.org/files/11/11-0.txt",  # Alice in Wonderland
    ],
    "docx": [
        # Note: DOCX files are harder to find publicly, so we'll generate some
        # or use sample documents. For now, we'll create test DOCX files programmatically.
    ],
}

# Test queries with expected answers/contexts
# Format: (query, expected_keywords, expected_file_patterns, difficulty)
TEST_QUERIES = [
    # Simple factual queries
    (
        "What is attention mechanism in transformers?",
        ["attention", "transformer", "self-attention"],
        ["1706.03762"],  # Transformer paper
        "easy",
    ),
    (
        "How does BERT work?",
        ["bert", "bidirectional", "encoder"],
        ["1810.04805"],  # BERT paper
        "easy",
    ),
    (
        "What is GPT-3?",
        ["gpt-3", "language model", "generative"],
        ["2005.14165"],  # GPT-3 paper
        "easy",
    ),
    # Complex multi-part queries
    (
        "Compare and contrast transformer architecture with BERT architecture",
        ["transformer", "bert", "architecture", "encoder"],
        ["1706.03762", "1810.04805"],
        "hard",
    ),
    (
        "What are the key innovations in language models from GPT-3 to transformers?",
        ["gpt-3", "transformer", "language model", "innovation"],
        ["2005.14165", "1706.03762"],
        "hard",
    ),
    # Edge cases - misspellings and variations
    (
        "What is attension mechansim?",  # Intentional misspellings
        ["attention", "mechanism"],
        ["1706.03762"],
        "medium",
    ),
    (
        "How do transformers work?",
        ["transformer", "attention", "encoder", "decoder"],
        ["1706.03762"],
        "medium",
    ),
    # Specific technical queries
    (
        "What is the architecture of the transformer model?",
        ["transformer", "architecture", "encoder", "decoder", "attention"],
        ["1706.03762"],
        "medium",
    ),
    (
        "Explain self-attention mechanism",
        ["self-attention", "attention", "mechanism"],
        ["1706.03762"],
        "medium",
    ),
    # Broad queries that should retrieve multiple documents
    (
        "What are neural language models?",
        ["language model", "neural", "nlp"],
        ["2005.14165", "1810.04805", "1706.03762"],
        "medium",
    ),
    # Negative queries (should not retrieve certain documents)
    (
        "What is Python programming?",
        ["python", "programming"],
        ["python"],  # Should NOT retrieve transformer/BERT papers
        "easy",
    ),
]


def validate_file_content(file_path: Path, expected_type: str) -> bool:
    """Validate that a downloaded file is of the expected type."""
    try:
        if not file_path.exists() or file_path.stat().st_size == 0:
            logger.error(f"File validation failed: {file_path.name} does not exist or is empty")
            return False
        
        # Check file header/magic bytes
        with open(file_path, "rb") as f:
            header = f.read(1024)
        
        if expected_type == "pdf":
            # PDF files start with %PDF
            if not header.startswith(b"%PDF"):
                # Check if it's HTML (common for 404 pages)
                if header.startswith(b"<!DOCTYPE") or header.startswith(b"<html") or header.startswith(b"<HTML"):
                    logger.error(
                        f"Download failed: {file_path.name} is HTML (likely 404 page), not a PDF. "
                        f"Starts with: {header[:50].decode('utf-8', errors='ignore')[:50]}"
                    )
                else:
                    logger.error(
                        f"Download failed: {file_path.name} does not appear to be a valid PDF. "
                        f"Starts with: {header[:50]}"
                    )
                return False
        elif expected_type in ("md", "markdown", "txt"):
            # Text files should be readable as UTF-8 or ASCII
            try:
                header.decode("utf-8")
            except UnicodeDecodeError:
                # Might be binary, check if it's actually text
                if b"\x00" in header[:100]:  # Null bytes suggest binary
                    logger.error(f"Download failed: {file_path.name} appears to be binary, not text")
                    return False
        
        return True
    except Exception as e:
        logger.error(f"Error validating {file_path}: {e}")
        return False


def download_file(url: str, output_path: Path, timeout: int = 30, expected_type: str = "auto") -> bool:
    """Download a file from URL to output path and validate it."""
    try:
        logger.info(f"Downloading {url} to {output_path}")
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
        }
        response = requests.get(url, timeout=timeout, stream=True, headers=headers)
        response.raise_for_status()
        
        # Check content type from response
        content_type = response.headers.get("Content-Type", "").lower()
        
        # Detect expected type if not specified
        if expected_type == "auto":
            if output_path.suffix == ".pdf":
                expected_type = "pdf"
            elif output_path.suffix in (".md", ".rst", ".txt"):
                expected_type = "txt"
            else:
                expected_type = "unknown"
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Handle zip files specially
        if url.endswith(".zip") or "zip" in content_type:
            import zipfile
            import io
            try:
                with zipfile.ZipFile(io.BytesIO(response.content)) as z:
                    # Extract PDFs from zip
                    for member in z.namelist():
                        if member.endswith(".pdf"):
                            with z.open(member) as source, open(output_path, "wb") as target:
                                target.write(source.read())
                            logger.info(f"Extracted {member} from zip")
                            if validate_file_content(output_path, "pdf"):
                                return True
                            else:
                                output_path.unlink()  # Remove invalid file
                                return False
            except zipfile.BadZipFile:
                logger.error(f"Download failed: Downloaded file is not a valid ZIP: {url}")
                return False
            return False
        
        # Download file
        with open(output_path, "wb") as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
        
        file_size = output_path.stat().st_size
        
        # Validate file content - this is critical for detecting HTML 404 pages masquerading as PDFs
        if expected_type == "pdf":
            if not validate_file_content(output_path, "pdf"):
                logger.error(f"Download validation failed for {url} - file is not a valid PDF")
                if output_path.exists():
                    output_path.unlink()  # Remove invalid file
                return False
        elif expected_type == "txt":
            if not validate_file_content(output_path, "txt"):
                logger.error(f"Download validation failed for {url} - file is not valid text")
                if output_path.exists():
                    output_path.unlink()  # Remove invalid file
                return False
        
        logger.info(f"Downloaded {output_path.name} ({file_size} bytes)")
        return True
    except requests.exceptions.HTTPError as e:
        if e.response.status_code == 404:
            logger.error(f"File not found (404): {url}")
        else:
            logger.error(f"HTTP error downloading {url}: {e}")
        return False
    except Exception as e:
        logger.error(f"Failed to download {url}: {e}")
        # Clean up partial download
        if output_path.exists():
            try:
                output_path.unlink()
            except Exception:
                pass
        return False


def create_sample_docx(output_path: Path, content: str, title: str = "Test Document") -> bool:
    """Create a sample DOCX file with content."""
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading(title, 0)
        
        # Split content into paragraphs
        for para_text in content.split("\n\n"):
            if para_text.strip():
                doc.add_paragraph(para_text.strip())
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        logger.info(f"Created DOCX file: {output_path}")
        return True
    except Exception as e:
        logger.warning(f"Failed to create DOCX {output_path}: {e}")
        return False


def cleanup_invalid_files(data_dir: Path) -> None:
    """Remove invalid files from previous test runs."""
    logger.info("Cleaning up invalid files from previous runs...")
    removed_count = 0
    
    for pdf_file in data_dir.rglob("*.pdf"):
        if not validate_file_content(pdf_file, "pdf"):
            logger.info(f"Removing invalid PDF: {pdf_file}")
            try:
                pdf_file.unlink()
                removed_count += 1
            except Exception as e:
                logger.warning(f"Failed to remove {pdf_file}: {e}")
    
    if removed_count > 0:
        logger.info(f"Removed {removed_count} invalid file(s)")


def download_test_corpus() -> Dict[str, List[Path]]:
    """Download a diverse corpus of test documents."""
    logger.info("=" * 80)
    logger.info("Downloading Test Corpus")
    logger.info("=" * 80)
    
    TEST_DATA_DIR.mkdir(parents=True, exist_ok=True)
    
    # Clean up any invalid files from previous runs
    cleanup_invalid_files(TEST_DATA_DIR)
    
    downloaded_files = defaultdict(list)
    download_failures = []
    
    # Download PDFs
    for url in DOCUMENT_SOURCES["pdf"]:
        filename = Path(urlparse(url).path).name
        if not filename.endswith(".pdf"):
            filename = f"{Path(urlparse(url).path).stem}.pdf"
        output_path = TEST_DATA_DIR / "pdf" / filename
        if download_file(url, output_path, expected_type="pdf"):
            downloaded_files["pdf"].append(output_path)
        else:
            download_failures.append(("pdf", url, output_path))
    
    # Download Markdown files
    for url in DOCUMENT_SOURCES["markdown"]:
        filename = Path(urlparse(url).path).name
        if not filename.endswith((".md", ".rst", ".txt")):
            # Generate unique filename based on URL
            url_stem = Path(urlparse(url).path).stem
            if url.endswith(".rst"):
                filename = f"{url_stem}.rst"
            else:
                filename = f"{url_stem}.md"
        output_path = TEST_DATA_DIR / "md" / filename
        # Handle filename conflicts by adding source identifier
        if output_path.exists():
            # Add a prefix to avoid overwriting
            url_parts = urlparse(url).netloc.split(".")
            prefix = url_parts[-2] if len(url_parts) >= 2 else "doc"
            output_path = TEST_DATA_DIR / "md" / f"{prefix}_{filename}"
        if download_file(url, output_path, expected_type="txt"):
            downloaded_files["markdown"].append(output_path)
        else:
            download_failures.append(("markdown", url, output_path))
    
    # Download TXT files
    for url in DOCUMENT_SOURCES["txt"]:
        filename = Path(urlparse(url).path).name
        if not filename.endswith(".txt"):
            filename = f"{Path(urlparse(url).path).stem}.txt"
        output_path = TEST_DATA_DIR / "txt" / filename
        if download_file(url, output_path, expected_type="txt"):
            downloaded_files["txt"].append(output_path)
        else:
            download_failures.append(("txt", url, output_path))
    
    # Create sample DOCX files
    docx_content = {
        "transformer_overview.docx": {
            "title": "Transformer Architecture Overview",
            "content": """
# Transformer Architecture

## Introduction
The Transformer architecture, introduced in "Attention Is All You Need", revolutionized natural language processing.

## Key Components

### Self-Attention Mechanism
Self-attention allows the model to weigh the importance of different words in a sequence.

### Encoder-Decoder Structure
The transformer uses an encoder-decoder architecture with multiple layers.

## Applications
Transformers are used in machine translation, text generation, and many other NLP tasks.
            """.strip(),
        },
        "bert_explained.docx": {
            "title": "BERT Explained",
            "content": """
# BERT: Bidirectional Encoder Representations from Transformers

## Overview
BERT is a bidirectional transformer model that reads text in both directions.

## Architecture
BERT uses only the encoder part of the transformer architecture.

## Training
BERT is pre-trained on large text corpora using masked language modeling and next sentence prediction.

## Applications
BERT is widely used for question answering, sentiment analysis, and named entity recognition.
            """.strip(),
        },
    }
    
    for filename, doc_data in docx_content.items():
        output_path = TEST_DATA_DIR / "docx" / filename
        if create_sample_docx(output_path, doc_data["content"], doc_data["title"]):
            downloaded_files["docx"].append(output_path)
    
    total_files = sum(len(files) for files in downloaded_files.values())
    logger.info(f"\nDownloaded {total_files} files:")
    for file_type, files in downloaded_files.items():
        logger.info(f"  {file_type}: {len(files)} files")
    
    # Report download failures
    if download_failures:
        logger.error(f"\n{'=' * 80}")
        logger.error(f"Download Failures: {len(download_failures)} file(s) failed to download")
        logger.error(f"{'=' * 80}")
        for file_type, url, output_path in download_failures:
            logger.error(f"  {file_type}: {url}")
        
        if ABORT_ON_DOWNLOAD_FAILURE:
            logger.error("\nAborting test suite due to download failures.")
            logger.error("Set ABORT_ON_DOWNLOAD_FAILURE=0 to continue despite failures.")
            raise RuntimeError(
                f"Failed to download {len(download_failures)} required file(s). "
                "Test suite aborted. Check the URLs and network connectivity."
            )
        else:
            logger.warning("\nContinuing despite download failures (ABORT_ON_DOWNLOAD_FAILURE=0).")
    
    if total_files == 0:
        raise RuntimeError(
            "No files were downloaded. Cannot proceed with tests. "
            "Check network connectivity and URL availability."
        )
    
    return dict(downloaded_files)


def _get_file_identifier(chunk: Dict) -> str:
    """Extract file identifier from chunk for pattern matching."""
    # Try location.file first (full path), then location.uri, then metadata fields
    file_path = chunk.get("location", {}).get("file", "")
    if not file_path:
        # Extract path from file:// URI if present
        uri = chunk.get("location", {}).get("uri", "")
        if uri and uri.startswith("file://"):
            file_path = uri[7:]  # Remove "file://" prefix
    if not file_path:
        file_path = chunk.get("metadata", {}).get("file_name", "")
    if not file_path:
        file_path = chunk.get("metadata", {}).get("file_path", "")
    return str(file_path).lower()


def precision_at_k(retrieved: List[Dict], relevant_file_patterns: List[str], k: int) -> float:
    """Calculate Precision@k."""
    if k == 0:
        return 0.0
    
    top_k = retrieved[:k]
    relevant_count = 0
    
    for chunk in top_k:
        file_id = _get_file_identifier(chunk)
        
        # Check if any relevant pattern matches
        if any(pattern.lower() in file_id for pattern in relevant_file_patterns):
            relevant_count += 1
    
    return relevant_count / k


def recall_at_k(retrieved: List[Dict], relevant_file_patterns: List[str], k: int, total_relevant: int) -> float:
    """Calculate Recall@k."""
    if total_relevant == 0:
        return 0.0
    
    top_k = retrieved[:k]
    relevant_retrieved = set()
    
    for chunk in top_k:
        file_id = _get_file_identifier(chunk)
        
        # Find which relevant pattern matches
        for pattern in relevant_file_patterns:
            if pattern.lower() in file_id:
                relevant_retrieved.add(pattern)
                break
    
    return len(relevant_retrieved) / total_relevant if total_relevant > 0 else 0.0


def mean_reciprocal_rank(retrieved: List[Dict], relevant_file_patterns: List[str]) -> float:
    """Calculate Mean Reciprocal Rank (MRR)."""
    for rank, chunk in enumerate(retrieved, start=1):
        file_id = _get_file_identifier(chunk)
        
        if any(pattern.lower() in file_id for pattern in relevant_file_patterns):
            return 1.0 / rank
    
    return 0.0


def ndcg_at_k(retrieved: List[Dict], relevant_file_patterns: List[str], k: int) -> float:
    """Calculate Normalized Discounted Cumulative Gain (NDCG@k)."""
    if not retrieved or k == 0:
        return 0.0
    
    # Use actual number of retrieved chunks, not k (in case we have fewer than k)
    actual_k = min(k, len(retrieved))
    if actual_k == 0:
        return 0.0
    
    def relevance_score(chunk: Dict) -> float:
        file_id = _get_file_identifier(chunk)
        # Binary relevance: 1 if relevant, 0 otherwise
        return 1.0 if any(pattern.lower() in file_id for pattern in relevant_file_patterns) else 0.0
    
    # Calculate DCG@k (using actual_k, not k)
    dcg = 0.0
    relevant_count = 0
    for i, chunk in enumerate(retrieved[:actual_k], start=1):
        rel = relevance_score(chunk)
        if rel > 0:
            relevant_count += 1
        # DCG formula: sum of (relevance / log2(rank + 1))
        # Note: using i (1-indexed) so log2(i+1) = log2(2) for first item = 1.0
        dcg += rel / math.log2(i + 1)
    
    # If no relevant documents were retrieved, NDCG is 0
    if relevant_count == 0:
        return 0.0
    
    # Calculate IDCG@k (ideal DCG - all relevant documents at the top)
    # IDCG assumes all relevant documents are ranked at positions 1, 2, 3, ...
    # We need to know how many relevant documents exist, but we only have patterns.
    # For simplicity, we assume each pattern corresponds to at least one document.
    # The ideal case is having all relevant documents at the top.
    num_relevant_docs = len(relevant_file_patterns)
    if num_relevant_docs == 0:
        return 0.0
    
    # IDCG: DCG if all relevant documents were perfectly ranked at the top
    # Use actual_k (not k) to match the number of documents we're actually evaluating
    # We can't have more than actual_k documents, and we can't have more than num_relevant_docs
    ideal_count = min(actual_k, num_relevant_docs)
    idcg = sum(1.0 / math.log2(i + 1) for i in range(1, ideal_count + 1))
    
    if idcg == 0:
        return 0.0
    
    # NDCG = DCG / IDCG (normalized to [0, 1])
    ndcg = dcg / idcg
    # Clamp to [0, 1] in case of floating point issues
    return max(0.0, min(1.0, ndcg))


def evaluate_query_with_retriever(
    query: str,
    expected_keywords: List[str],
    expected_file_patterns: List[str],
    difficulty: str,
    retriever_func,
) -> Dict[str, Any]:
    """Evaluate a single query against the RAG system."""
    logger.info(f"\nEvaluating query: {query}")
    logger.info(f"Expected patterns: {expected_file_patterns}")

    start_time = time.time()
    result = retriever_func(query)
    elapsed = time.time() - start_time
    
    chunks = result.get("chunks", [])
    
    # Calculate metrics at different k values
    k_values = [1, 3, 5, 10]
    metrics = {}
    
    for k in k_values:
        metrics[f"precision@{k}"] = precision_at_k(chunks, expected_file_patterns, k)
        metrics[f"recall@{k}"] = recall_at_k(chunks, expected_file_patterns, k, len(expected_file_patterns))
        # Calculate NDCG@k even if we have fewer than k chunks (ndcg_at_k handles this)
        if len(chunks) > 0:
            ndcg_value = ndcg_at_k(chunks, expected_file_patterns, k)
            metrics[f"ndcg@{k}"] = ndcg_value
            # Log if NDCG is suspiciously low
            if ndcg_value == 0.0 and len(expected_file_patterns) > 0:
                # Check if any relevant docs were retrieved
                actual_k = min(k, len(chunks))
                relevant_found = any(
                    any(p.lower() in _get_file_identifier(chunk)
                        for p in expected_file_patterns)
                    for chunk in chunks[:actual_k]
                )
                if not relevant_found:
                    logger.debug(f"  NDCG@{k}=0.0: No relevant documents found in top {actual_k} results")
        else:
            metrics[f"ndcg@{k}"] = 0.0
    
    metrics["mrr"] = mean_reciprocal_rank(chunks, expected_file_patterns)
    
    # Check if expected keywords appear in retrieved chunks
    keyword_matches = {}
    for keyword in expected_keywords:
        found = False
        for chunk in chunks[:5]:  # Check top 5 chunks
            text = chunk.get("text", "").lower()
            if keyword.lower() in text:
                found = True
                break
        keyword_matches[keyword] = found
    
    # Get top retrieved files
    top_files = []
    for chunk in chunks[:5]:
        file_name = chunk.get("metadata", {}).get("file_name", "unknown")
        top_files.append(file_name)
    
    evaluation = {
        "query": query,
        "difficulty": difficulty,
        "num_chunks_retrieved": len(chunks),
        "retrieval_time": elapsed,
        "metrics": metrics,
        "keyword_matches": keyword_matches,
        "top_files": top_files,
        "expected_patterns": expected_file_patterns,
    }
    
    logger.info(f"  Precision@5: {metrics['precision@5']:.3f}")
    logger.info(f"  Recall@5: {metrics['recall@5']:.3f}")
    logger.info(f"  MRR: {metrics['mrr']:.3f}")
    ndcg5 = metrics.get('ndcg@5', 0.0)
    logger.info(f"  NDCG@5: {ndcg5:.3f}")
    if ndcg5 == 0.0 and len(expected_file_patterns) > 0:
        # Check if relevant docs were actually retrieved
        relevant_in_top5 = sum(
            1 for chunk in chunks[:5]
            if any(
                pattern.lower() in _get_file_identifier(chunk)
                for pattern in expected_file_patterns
            )
        )
        if relevant_in_top5 == 0:
            logger.warning(f"  ⚠ NDCG@5=0.0: No relevant documents found in top 5 results")
            logger.warning(f"     Expected patterns: {expected_file_patterns}")
            logger.warning(f"     Top 5 files: {top_files[:5]}")
    logger.info(f"  Top files: {', '.join(top_files[:3])}")
    
    return evaluation


def evaluate_query(
    query: str,
    expected_keywords: List[str],
    expected_file_patterns: List[str],
    difficulty: str,
) -> Dict[str, Any]:
    """Evaluate a single query against the RAG system (uses global run_search)."""
    return evaluate_query_with_retriever(
        query, expected_keywords, expected_file_patterns, difficulty, run_search
    )


def run_rag_metrics_tests() -> Dict[str, Any]:
    """Run the complete RAG metrics test suite."""
    logger.info("=" * 80)
    logger.info("RAG Metrics Test Suite")
    logger.info("=" * 80)
    
    # Step 1: Download test corpus
    downloaded_files = download_test_corpus()
    
    if not any(downloaded_files.values()):
        logger.error("No documents downloaded. Cannot proceed with tests.")
        return {"error": "No documents downloaded"}
    
    # Step 2: Patch module globals for test isolation
    TEST_STORAGE_DIR.mkdir(parents=True, exist_ok=True)

    from chunksilo import index as index_mod
    from chunksilo import search as search_mod

    # Save originals for restoration
    orig_index_storage = index_mod.STORAGE_DIR
    orig_index_state_db = index_mod.STATE_DB_PATH
    orig_index_bm25 = index_mod.BM25_INDEX_DIR
    orig_index_heading = index_mod.HEADING_STORE_PATH
    orig_index_config = index_mod._config
    orig_search_config = search_mod._config

    try:
        # Patch index module globals
        index_mod.STORAGE_DIR = TEST_STORAGE_DIR
        index_mod.STATE_DB_PATH = TEST_STORAGE_DIR / "ingestion_state.db"
        index_mod.BM25_INDEX_DIR = TEST_STORAGE_DIR / "bm25_index"
        index_mod.HEADING_STORE_PATH = TEST_STORAGE_DIR / "heading_store.json"
        index_mod._config = {
            **index_mod._config,
            "indexing": {
                "directories": [str(TEST_DATA_DIR)],
                "chunk_size": 1600,
                "chunk_overlap": 200,
            },
        }

        # Patch search module config
        test_search_config = search_mod._init_config()
        test_search_config["storage"]["storage_dir"] = str(TEST_STORAGE_DIR)
        search_mod._config = test_search_config
        search_mod._index_cache = None
        search_mod._bm25_retriever_cache = None

        from chunksilo.index import build_index as build_test_index
        from chunksilo.search import load_llamaindex_index, run_search as run_search_reloaded
        
        # Step 3: Build index
        logger.info("\n" + "=" * 80)
        logger.info("Building Index from Test Corpus")
        logger.info("=" * 80)
        
        build_test_index()
        
        # Step 4: Load index
        logger.info("\n" + "=" * 80)
        logger.info("Loading Index")
        logger.info("=" * 80)
        
        index = load_llamaindex_index()
        logger.info("Index loaded successfully")
        
        # Step 5: Run evaluation queries
        logger.info("\n" + "=" * 80)
        logger.info("Running Evaluation Queries")
        logger.info("=" * 80)
        
        evaluations = []
        for query, keywords, patterns, difficulty in TEST_QUERIES:
            try:
                # Use the reloaded run_search function
                eval_result = evaluate_query_with_retriever(
                    query, keywords, patterns, difficulty, run_search_reloaded
                )
                evaluations.append(eval_result)
            except Exception as e:
                logger.error(f"Error evaluating query '{query}': {e}")
                import traceback
                traceback.print_exc()
                evaluations.append({
                    "query": query,
                    "error": str(e),
                })
        
        # Step 6: Calculate aggregate metrics
        logger.info("\n" + "=" * 80)
        logger.info("Calculating Aggregate Metrics")
        logger.info("=" * 80)
        
        successful_evals = [e for e in evaluations if "error" not in e]
        
        if successful_evals:
            aggregate_metrics = {
                "precision@1": sum(e["metrics"]["precision@1"] for e in successful_evals) / len(successful_evals),
                "precision@5": sum(e["metrics"]["precision@5"] for e in successful_evals) / len(successful_evals),
                "recall@5": sum(e["metrics"]["recall@5"] for e in successful_evals) / len(successful_evals),
                "mrr": sum(e["metrics"]["mrr"] for e in successful_evals) / len(successful_evals),
                "ndcg@5": sum(e["metrics"].get("ndcg@5", 0) for e in successful_evals) / len(successful_evals),
            }
            
            # Metrics by difficulty
            difficulty_metrics = defaultdict(lambda: {"count": 0, "precision@5": [], "recall@5": [], "mrr": []})
            for eval_result in successful_evals:
                diff = eval_result.get("difficulty", "unknown")
                difficulty_metrics[diff]["count"] += 1
                difficulty_metrics[diff]["precision@5"].append(eval_result["metrics"]["precision@5"])
                difficulty_metrics[diff]["recall@5"].append(eval_result["metrics"]["recall@5"])
                difficulty_metrics[diff]["mrr"].append(eval_result["metrics"]["mrr"])
            
            difficulty_summary = {}
            for diff, data in difficulty_metrics.items():
                if data["count"] > 0:
                    difficulty_summary[diff] = {
                        "count": data["count"],
                        "avg_precision@5": sum(data["precision@5"]) / len(data["precision@5"]),
                        "avg_recall@5": sum(data["recall@5"]) / len(data["recall@5"]),
                        "avg_mrr": sum(data["mrr"]) / len(data["mrr"]),
                    }
            
            logger.info("\nAggregate Metrics:")
            logger.info(f"  Precision@1: {aggregate_metrics['precision@1']:.3f}")
            logger.info(f"  Precision@5: {aggregate_metrics['precision@5']:.3f}")
            logger.info(f"  Recall@5: {aggregate_metrics['recall@5']:.3f}")
            logger.info(f"  MRR: {aggregate_metrics['mrr']:.3f}")
            logger.info(f"  NDCG@5: {aggregate_metrics['ndcg@5']:.3f}")
            
            logger.info("\nMetrics by Difficulty:")
            for diff, metrics in difficulty_summary.items():
                logger.info(f"  {diff}:")
                logger.info(f"    Count: {metrics['count']}")
                logger.info(f"    Avg Precision@5: {metrics['avg_precision@5']:.3f}")
                logger.info(f"    Avg Recall@5: {metrics['avg_recall@5']:.3f}")
                logger.info(f"    Avg MRR: {metrics['avg_mrr']:.3f}")
        else:
            aggregate_metrics = {}
            difficulty_summary = {}
            logger.warning("No successful evaluations to aggregate")
        
        # Step 7: Save results
        TEST_RESULTS_DIR.mkdir(parents=True, exist_ok=True)
        results_file = TEST_RESULTS_DIR / f"test_results_{int(time.time())}.json"
        
        results = {
            "timestamp": time.time(),
            "corpus_stats": {
                file_type: len(files) for file_type, files in downloaded_files.items()
            },
            "num_queries": len(TEST_QUERIES),
            "num_successful_evals": len(successful_evals),
            "aggregate_metrics": aggregate_metrics,
            "difficulty_summary": difficulty_summary,
            "evaluations": evaluations,
        }
        
        with open(results_file, "w") as f:
            json.dump(results, f, indent=2)
        
        logger.info(f"\nResults saved to: {results_file}")
        
        return results
        
    finally:
        # Restore original module globals
        index_mod.STORAGE_DIR = orig_index_storage
        index_mod.STATE_DB_PATH = orig_index_state_db
        index_mod.BM25_INDEX_DIR = orig_index_bm25
        index_mod.HEADING_STORE_PATH = orig_index_heading
        index_mod._config = orig_index_config
        search_mod._config = orig_search_config
        search_mod._index_cache = None
        search_mod._bm25_retriever_cache = None


def main():
    """Main entry point."""
    results = run_rag_metrics_tests()
    
    if "error" in results:
        logger.error(f"Test suite failed: {results['error']}")
        sys.exit(1)
    
    logger.info("\n" + "=" * 80)
    logger.info("Test Suite Completed Successfully")
    logger.info("=" * 80)


if __name__ == "__main__":
    main()
