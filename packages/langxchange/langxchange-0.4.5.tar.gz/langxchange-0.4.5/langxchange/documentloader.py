# # langxchange/document_loader_helper.py

import os
import re
import time
import json
import base64
import pandas as pd
from typing import List, Dict, Any, Generator, Optional, Tuple
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
from enum import Enum
from pathlib import Path

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import docx
except ImportError:
    docx = None

try:
    import tiktoken
    TIKTOKEN_AVAILABLE = True
except ImportError:
    TIKTOKEN_AVAILABLE = False

try:
    from PIL import Image, ExifTags
    from PIL.ExifTags import TAGS
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False


class ChunkingStrategy(Enum):
    """Enumeration of available chunking strategies."""
    CHARACTER = "character"  # Simple character-based chunking
    SENTENCE = "sentence"    # Sentence-aware chunking
    PARAGRAPH = "paragraph"  # Paragraph-aware chunking
    SEMANTIC = "semantic"    # Semantic chunking with overlap
    TOKEN = "token"         # Token-based chunking (requires tiktoken)


class ImageProcessingStrategy(Enum):
    """Enumeration of available image processing strategies."""
    OCR_TEXT = "ocr_text"           # Extract text using OCR
    DESCRIPTION = "description"     # Generate image description
    METADATA = "metadata"           # Extract technical metadata
    COMBINED = "combined"           # OCR + description + metadata
    VISUAL_ANALYSIS = "visual"      # Advanced visual analysis


class ExcelChunkingStrategy(Enum):
    """Enumeration of available Excel chunking strategies."""
    HEADER_ROW = "header_row"           # Pair headers with each row (best for QA)
    ROW_STRINGIFIED = "row_stringified" # Convert rows to strings with context
    TABLE_AS_CHUNK = "table_as_chunk"   # Treat entire table as one chunk (best for small files)


@dataclass
class ChunkMetadata:
    """Metadata for each text chunk."""
    source_file: str
    chunk_index: int
    total_chunks: int
    file_type: str
    page_number: Optional[int] = None
    section_title: Optional[str] = None
    char_start: Optional[int] = None
    char_end: Optional[int] = None
    token_count: Optional[int] = None
    # Image-specific metadata
    image_dimensions: Optional[Tuple[int, int]] = None
    image_format: Optional[str] = None
    image_mode: Optional[str] = None
    image_size_bytes: Optional[int] = None
    processing_strategy: Optional[str] = None
    ocr_confidence: Optional[float] = None
    exif_data: Optional[Dict[str, Any]] = None


@dataclass
class TextChunk:
    """Container for text chunk with metadata."""
    content: str
    metadata: ChunkMetadata


class DocumentLoaderHelper:
    """
    Enhanced helper to load and extract textual content from various document types
    with intelligent chunking strategies optimized for LLM processing.
    """

    def __init__(
        self, 
        chunk_size: int = 1000,
        overlap_size: int = 100,
        chunking_strategy: ChunkingStrategy = ChunkingStrategy.SEMANTIC,
        csv_chunksize: int = 1000,
        max_workers: int = 4,
        encoding_model: str = "cl100k_base",
        preserve_formatting: bool = True,
        min_chunk_size: int = 50,
        # Image processing parameters
        image_processing_strategy: ImageProcessingStrategy = ImageProcessingStrategy.COMBINED,
        ocr_language: str = "eng",
        include_image_metadata: bool = True,
        max_image_size: Tuple[int, int] = (2048, 2048),
        image_quality_threshold: float = 0.5,
        # Excel processing parameters
        excel_chunking_strategy: ExcelChunkingStrategy = ExcelChunkingStrategy.HEADER_ROW,
        excel_max_rows_per_chunk: int = 100,
        excel_include_sheet_name: bool = True
    ):
        """
        Initialize the DocumentLoaderHelper with enhanced chunking and image processing capabilities.
        
        :param chunk_size: Target size for each chunk (characters or tokens)
        :param overlap_size: Number of characters/tokens to overlap between chunks
        :param chunking_strategy: Strategy to use for chunking text
        :param csv_chunksize: Number of rows per chunk when loading CSV
        :param max_workers: Threads to use for parallel processing
        :param encoding_model: Tokenizer model for token counting (requires tiktoken)
        :param preserve_formatting: Whether to preserve document formatting
        :param min_chunk_size: Minimum size for a chunk to be considered valid
        :param image_processing_strategy: Strategy for processing images
        :param ocr_language: Language for OCR processing (default: English)
        :param include_image_metadata: Whether to extract image metadata (EXIF, etc.)
        :param max_image_size: Maximum image dimensions for processing (width, height)
        :param image_quality_threshold: Minimum OCR confidence threshold (0.0-1.0)
        :param excel_chunking_strategy: Strategy for chunking Excel files (default: HEADER_ROW)
        :param excel_max_rows_per_chunk: Maximum rows per chunk for table-as-chunk strategy
        :param excel_include_sheet_name: Whether to include sheet name in metadata
        """
        self.chunk_size = chunk_size
        self.overlap_size = overlap_size
        self.chunking_strategy = chunking_strategy
        self.csv_chunksize = csv_chunksize
        self.max_workers = max_workers
        self.encoding_model = encoding_model
        self.preserve_formatting = preserve_formatting
        self.min_chunk_size = min_chunk_size
        
        # Image processing parameters
        self.image_processing_strategy = image_processing_strategy
        self.ocr_language = ocr_language
        self.include_image_metadata = include_image_metadata
        self.max_image_size = max_image_size
        self.image_quality_threshold = image_quality_threshold

        # Excel processing parameters
        self.excel_chunking_strategy = excel_chunking_strategy
        self.excel_max_rows_per_chunk = excel_max_rows_per_chunk
        self.excel_include_sheet_name = excel_include_sheet_name

        # Initialize tokenizer if available and strategy requires it
        self.tokenizer = None
        if TIKTOKEN_AVAILABLE and chunking_strategy == ChunkingStrategy.TOKEN:
            try:
                self.tokenizer = tiktoken.get_encoding(encoding_model)
            except Exception:
                # Fallback to character-based if tokenizer fails
                self.chunking_strategy = ChunkingStrategy.SEMANTIC

        # Compile regex patterns for text processing
        self._sentence_pattern = re.compile(r'(?<=[.!?])\s+')
        self._paragraph_pattern = re.compile(r'\n\s*\n')
        self._whitespace_pattern = re.compile(r'\s+')
        self._bullet_pattern = re.compile(r'^[\s]*[•\-\*\+]\s+', re.MULTILINE)

        # Stats tracking
        self.stats = {
            "total_units": 0,
            "processed_units": 0,
            "total_chunks": 0,
            "times": {
                "load": 0.0,
                "chunk": 0.0,
                "total": 0.0
            }
        }

    def load(self, file_path: str) -> Generator[TextChunk, None, None]:
        """
        Extract text chunks from the file with metadata.
        
        :param file_path: Path to the document file
        :yields: TextChunk objects with content and metadata
        """
        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        start_total = time.perf_counter()
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # Load raw text units from file
        t0 = time.perf_counter()
        raw_units = self._load_file_content(file_path, file_ext)
        t1 = time.perf_counter()
        self.stats["times"]["load"] = t1 - t0

        self.stats["total_units"] = len(raw_units)
        self.stats["processed_units"] = 0

        # Process units in parallel with intelligent chunking
        t2 = time.perf_counter()
        all_chunks = []
        
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            future_to_unit = {
                executor.submit(self._process_unit, unit, file_path, file_ext, idx): (unit, idx)
                for idx, unit in enumerate(raw_units)
            }
            
            for future in as_completed(future_to_unit):
                unit, unit_idx = future_to_unit[future]
                try:
                    chunks = future.result()
                    all_chunks.extend(chunks)
                    self.stats["processed_units"] += 1
                except Exception as e:
                    print(f"Error processing unit {unit_idx}: {e}")

        # Update chunk metadata with correct total count
        total_chunks = len(all_chunks)
        self.stats["total_chunks"] = total_chunks
        
        for chunk in all_chunks:
            chunk.metadata.total_chunks = total_chunks
            
        t3 = time.perf_counter()
        self.stats["times"]["chunk"] = t3 - t2
        self.stats["times"]["total"] = time.perf_counter() - start_total

        # Yield chunks
        for chunk in all_chunks:
            yield chunk

    def _load_file_content(self, file_path: str, file_ext: str) -> List[Dict[str, Any]]:
        """Load file content into structured units."""
        # Text and document formats
        if file_ext == ".txt":
            return self._load_txt_units(file_path)
        elif file_ext == ".csv":
            return self._load_csv_units(file_path)
        elif file_ext == ".json":
            return self._load_json_units(file_path)
        elif file_ext == ".pdf":
            return self._load_pdf_units(file_path)
        elif file_ext in (".xls", ".xlsx"):
            return self._load_excel_units(file_path)
        elif file_ext == ".docx":
            return self._load_docx_units(file_path)
        # Image formats
        elif file_ext.lower() in self._get_supported_image_formats():
            return self._load_image_units(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

    def _get_supported_image_formats(self) -> List[str]:
        """Get list of supported image formats."""
        return [
            '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.tif',
            '.webp', '.ico', '.psd', '.svg', '.raw', '.cr2', '.nef',
            '.arw', '.dng', '.orf', '.rw2', '.pef', '.srw', '.x3f'
        ]

    def _process_unit(
        self, 
        unit: Dict[str, Any], 
        file_path: str, 
        file_ext: str, 
        unit_idx: int
    ) -> List[TextChunk]:
        """Process a single unit into chunks with metadata."""
        text_content = unit.get("content", "")
        
        # Clean and normalize text
        if self.preserve_formatting:
            cleaned_text = self._clean_text_preserve_format(text_content)
        else:
            cleaned_text = self._clean_text_normalize(text_content)
        
        # Skip empty or too small content
        if len(cleaned_text.strip()) < self.min_chunk_size:
            return []

        # Apply chunking strategy
        text_chunks = self._apply_chunking_strategy(cleaned_text)
        
        # Create TextChunk objects with metadata
        chunks = []
        for idx, chunk_text in enumerate(text_chunks):
            if len(chunk_text.strip()) >= self.min_chunk_size:
                metadata = ChunkMetadata(
                    source_file=os.path.basename(file_path),
                    chunk_index=len(chunks),
                    total_chunks=0,  # Will be updated later
                    file_type=file_ext.lstrip('.'),
                    page_number=unit.get("page_number"),
                    section_title=unit.get("section_title"),
                    token_count=self._count_tokens(chunk_text) if self.tokenizer else None
                )
                chunks.append(TextChunk(content=chunk_text, metadata=metadata))
        
        return chunks

    def _apply_chunking_strategy(self, text: str) -> List[str]:
        """Apply the selected chunking strategy to text."""
        if self.chunking_strategy == ChunkingStrategy.CHARACTER:
            return self._chunk_by_character(text)
        elif self.chunking_strategy == ChunkingStrategy.SENTENCE:
            return self._chunk_by_sentence(text)
        elif self.chunking_strategy == ChunkingStrategy.PARAGRAPH:
            return self._chunk_by_paragraph(text)
        elif self.chunking_strategy == ChunkingStrategy.SEMANTIC:
            return self._chunk_semantic(text)
        elif self.chunking_strategy == ChunkingStrategy.TOKEN:
            return self._chunk_by_token(text)
        else:
            return self._chunk_semantic(text)  # Default fallback

    def _chunk_by_character(self, text: str) -> List[str]:
        """Simple character-based chunking with overlap."""
        chunks = []
        start = 0
        text_len = len(text)
        
        while start < text_len:
            end = min(start + self.chunk_size, text_len)
            chunk = text[start:end]
            
            # Try to break at word boundary
            if end < text_len:
                last_space = chunk.rfind(' ')
                if last_space > start + self.chunk_size * 0.7:  # Don't break too early
                    chunk = text[start:start + last_space]
                    end = start + last_space
            
            chunks.append(chunk)
            start = end - self.overlap_size if end < text_len else end
            
        return chunks

    def _chunk_by_sentence(self, text: str) -> List[str]:
        """Sentence-aware chunking with overlap."""
        sentences = self._sentence_pattern.split(text)
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            # Check if adding this sentence would exceed chunk size
            potential_chunk = current_chunk + " " + sentence if current_chunk else sentence
            
            if len(potential_chunk) <= self.chunk_size:
                current_chunk = potential_chunk
            else:
                if current_chunk:
                    chunks.append(current_chunk)
                    # Start new chunk with overlap
                    if self.overlap_size > 0:
                        overlap_sentences = self._get_last_sentences(current_chunk, self.overlap_size)
                        current_chunk = overlap_sentences + " " + sentence if overlap_sentences else sentence
                    else:
                        current_chunk = sentence
                else:
                    # Single sentence is too long, split it
                    chunks.extend(self._chunk_by_character(sentence))
                    current_chunk = ""
        
        if current_chunk:
            chunks.append(current_chunk)
            
        return chunks

    def _chunk_by_paragraph(self, text: str) -> List[str]:
        """Paragraph-aware chunking."""
        paragraphs = self._paragraph_pattern.split(text)
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
                
            potential_chunk = current_chunk + "\n\n" + paragraph if current_chunk else paragraph
            
            if len(potential_chunk) <= self.chunk_size:
                current_chunk = potential_chunk
            else:
                if current_chunk:
                    chunks.append(current_chunk)
                
                # Handle large paragraphs
                if len(paragraph) > self.chunk_size:
                    chunks.extend(self._chunk_by_sentence(paragraph))
                    current_chunk = ""
                else:
                    current_chunk = paragraph
        
        if current_chunk:
            chunks.append(current_chunk)
            
        return chunks

    def _chunk_semantic(self, text: str) -> List[str]:
        """Intelligent semantic chunking that preserves meaning."""
        # First, try paragraph-based chunking
        paragraphs = self._paragraph_pattern.split(text)
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            # Check if this paragraph fits in current chunk
            potential_chunk = current_chunk + "\n\n" + paragraph if current_chunk else paragraph
            
            if len(potential_chunk) <= self.chunk_size:
                current_chunk = potential_chunk
            else:
                # Save current chunk if it exists
                if current_chunk:
                    chunks.append(current_chunk)
                
                # Handle large paragraphs with sentence-based chunking
                if len(paragraph) > self.chunk_size:
                    para_chunks = self._chunk_by_sentence(paragraph)
                    if para_chunks:
                        # Add overlap from previous chunk if available
                        if chunks and self.overlap_size > 0:
                            overlap = self._get_text_tail(chunks[-1], self.overlap_size)
                            para_chunks[0] = overlap + "\n" + para_chunks[0]
                        chunks.extend(para_chunks)
                        current_chunk = ""
                else:
                    # Add overlap from previous chunk
                    if chunks and self.overlap_size > 0:
                        overlap = self._get_text_tail(chunks[-1], self.overlap_size)
                        current_chunk = overlap + "\n" + paragraph
                    else:
                        current_chunk = paragraph
        
        if current_chunk:
            chunks.append(current_chunk)
            
        return chunks

    def _chunk_by_token(self, text: str) -> List[str]:
        """Token-based chunking using tiktoken."""
        if not self.tokenizer:
            return self._chunk_semantic(text)  # Fallback
            
        tokens = self.tokenizer.encode(text)
        chunks = []
        start = 0
        
        while start < len(tokens):
            end = min(start + self.chunk_size, len(tokens))
            chunk_tokens = tokens[start:end]
            chunk_text = self.tokenizer.decode(chunk_tokens)
            chunks.append(chunk_text)
            
            start = end - self.overlap_size if end < len(tokens) else end
            
        return chunks

    def _clean_text_preserve_format(self, text: str) -> str:
        """Clean text while preserving important formatting."""
        # Remove excessive whitespace but preserve structure
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # Max 2 consecutive newlines
        text = re.sub(r'[ \t]+', ' ', text)  # Normalize spaces and tabs
        text = text.strip()
        return text

    def _clean_text_normalize(self, text: str) -> str:
        """Normalize text by removing extra whitespace."""
        text = self._whitespace_pattern.sub(' ', text)
        text = text.strip()
        return text

    def _get_last_sentences(self, text: str, max_chars: int) -> str:
        """Get the last few sentences up to max_chars."""
        sentences = self._sentence_pattern.split(text)
        result = ""
        
        for sentence in reversed(sentences):
            sentence = sentence.strip()
            if not sentence:
                continue
                
            potential = sentence + " " + result if result else sentence
            if len(potential) > max_chars:
                break
            result = potential
            
        return result

    def _get_text_tail(self, text: str, max_chars: int) -> str:
        """Get the tail of text up to max_chars, preferring sentence boundaries."""
        if len(text) <= max_chars:
            return text
            
        # Try to break at sentence boundary
        tail = text[-max_chars:]
        sentence_start = tail.find('. ')
        if sentence_start > 0:
            return tail[sentence_start + 2:]
        
        # Try to break at word boundary
        space_pos = tail.find(' ')
        if space_pos > 0:
            return tail[space_pos + 1:]
            
        return tail

    def _count_tokens(self, text: str) -> int:
        """Count tokens in text."""
        if self.tokenizer:
            return len(self.tokenizer.encode(text))
        # Rough approximation: 1 token ≈ 4 characters
        return len(text) // 4

    # File-specific loading methods (enhanced versions)
    
    def _load_txt_units(self, path: str) -> List[Dict[str, Any]]:
        """Load text file content."""
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        return [{"content": content, "section_title": "Text Document"}]

    def _load_csv_units(self, path: str) -> List[Dict[str, Any]]:
        """Load CSV file content in chunks."""
        units = []
        for chunk_idx, df_chunk in enumerate(pd.read_csv(path, chunksize=self.csv_chunksize, dtype=str)):
            # Convert DataFrame to readable text
            df_chunk = df_chunk.fillna("")
            
            # Create a header for this chunk
            headers = " | ".join(df_chunk.columns)
            separator = "-" * len(headers)
            
            # Convert rows to text
            rows = []
            for _, row in df_chunk.iterrows():
                row_text = " | ".join(str(val) for val in row.values)
                rows.append(row_text)
            
            content = f"CSV Headers:\n{headers}\n{separator}\n" + "\n".join(rows)
            units.append({
                "content": content,
                "section_title": f"CSV Chunk {chunk_idx + 1}"
            })
        return units

    def _load_json_units(self, path: str) -> List[Dict[str, Any]]:
        """Load JSON file content."""
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        
        # Convert to readable format
        if isinstance(data, dict):
            content = json.dumps(data, indent=2, ensure_ascii=False)
        elif isinstance(data, list):
            content = json.dumps(data, indent=2, ensure_ascii=False)
        else:
            content = str(data)
            
        return [{"content": content, "section_title": "JSON Document"}]

    def _load_pdf_units(self, path: str) -> List[Dict[str, Any]]:
        """Load PDF file content page by page."""
        if PyPDF2 is None:
            raise ImportError("PyPDF2 is required for PDF support")
            
        units = []
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text() or ""
                if text.strip():
                    units.append({
                        "content": text,
                        "page_number": page_num,
                        "section_title": f"Page {page_num}"
                    })
        return units

    def _load_excel_units(self, path: str) -> List[Dict[str, Any]]:
        """
        Load Excel file content with advanced chunking strategies.
        
        Supports three chunking strategies:
        - HEADER_ROW: Each row becomes a self-contained chunk with headers
        - ROW_STRINGIFIED: Rows converted to readable strings with context
        - TABLE_AS_CHUNK: Entire tables as Markdown (for small files)
        """
        excel_file = pd.ExcelFile(path)
        units = []
        
        for sheet_index, sheet_name in enumerate(excel_file.sheet_names):
            df = pd.read_excel(path, sheet_name=sheet_name, dtype=str)
            df = df.fillna("")
            
            # Skip empty sheets
            if df.empty:
                continue
            
            # Apply chunking strategy
            if self.excel_chunking_strategy == ExcelChunkingStrategy.HEADER_ROW:
                sheet_units = self._chunk_excel_header_row(df, sheet_name, sheet_index)
            elif self.excel_chunking_strategy == ExcelChunkingStrategy.ROW_STRINGIFIED:
                sheet_units = self._chunk_excel_row_stringified(df, sheet_name, sheet_index)
            elif self.excel_chunking_strategy == ExcelChunkingStrategy.TABLE_AS_CHUNK:
                sheet_units = self._chunk_excel_table(df, sheet_name, sheet_index)
            else:
                # Fallback to header+row if unknown strategy
                sheet_units = self._chunk_excel_header_row(df, sheet_name, sheet_index)
            
            units.extend(sheet_units)
        
        return units
    
    def _chunk_excel_header_row(self, df: pd.DataFrame, sheet_name: str, sheet_index: int) -> List[Dict[str, Any]]:
        """
        Chunk Excel data using header+row strategy (best for QA).
        Each row becomes a self-contained chunk with column headers.
        """
        units = []
        column_headers = list(df.columns)
        total_rows = len(df)
        
        for row_idx, (_, row) in enumerate(df.iterrows(), start=1):
            # Create key-value pairs for each column
            row_data = []
            for col in column_headers:
                value = str(row[col]).strip()
                if value:  # Only include non-empty values
                    row_data.append(f"{col}: {value}")
            
            # Skip completely empty rows
            if not row_data:
                continue
            
            # Build content string
            content_parts = []
            if self.excel_include_sheet_name:
                content_parts.append(f"[Sheet: {sheet_name}]")
            content_parts.append(" | ".join(row_data))
            content = " ".join(content_parts)
            
            units.append({
                "content": content,
                "section_title": f"Excel Sheet: {sheet_name}",
                "sheet_name": sheet_name,
                "sheet_index": sheet_index,
                "row_number": row_idx,
                "total_rows": total_rows,
                "column_headers": column_headers,
                "excel_strategy": "header_row"
            })
        
        return units
    
    def _chunk_excel_row_stringified(self, df: pd.DataFrame, sheet_name: str, sheet_index: int) -> List[Dict[str, Any]]:
        """
        Chunk Excel data using row-wise stringified strategy.
        Convert each row into a single string with context.
        """
        units = []
        column_headers = list(df.columns)
        total_rows = len(df)
        
        for row_idx, (_, row) in enumerate(df.iterrows(), start=1):
            # Create readable string for the row
            row_values = []
            for col in column_headers:
                value = str(row[col]).strip()
                if value:
                    row_values.append(f"{col}: {value}")
            
            # Skip completely empty rows
            if not row_values:
                continue
            
            # Build content with sheet context
            if self.excel_include_sheet_name:
                content = f"[Sheet: {sheet_name}] [Row {row_idx}] " + " | ".join(row_values)
            else:
                content = f"[Row {row_idx}] " + " | ".join(row_values)
            
            units.append({
                "content": content,
                "section_title": f"Excel Sheet: {sheet_name}",
                "sheet_name": sheet_name,
                "sheet_index": sheet_index,
                "row_number": row_idx,
                "total_rows": total_rows,
                "column_headers": column_headers,
                "excel_strategy": "row_stringified"
            })
        
        return units
    
    def _chunk_excel_table(self, df: pd.DataFrame, sheet_name: str, sheet_index: int) -> List[Dict[str, Any]]:
        """
        Chunk Excel data using table-as-chunk strategy (best for small files).
        Convert entire table or chunks of rows to Markdown format.
        """
        units = []
        column_headers = list(df.columns)
        total_rows = len(df)
        
        # Split into chunks if table is too large
        num_chunks = (total_rows + self.excel_max_rows_per_chunk - 1) // self.excel_max_rows_per_chunk
        
        for chunk_idx in range(num_chunks):
            start_row = chunk_idx * self.excel_max_rows_per_chunk
            end_row = min((chunk_idx + 1) * self.excel_max_rows_per_chunk, total_rows)
            df_chunk = df.iloc[start_row:end_row]
            
            # Convert to Markdown table
            markdown_table = self._excel_to_markdown_table(df_chunk, column_headers)
            
            # Build content with sheet context
            content_parts = []
            if self.excel_include_sheet_name:
                content_parts.append(f"# Excel Sheet: {sheet_name}")
            if num_chunks > 1:
                content_parts.append(f"## Rows {start_row + 1}-{end_row} (of {total_rows})")
            content_parts.append(markdown_table)
            content = "\n\n".join(content_parts)
            
            units.append({
                "content": content,
                "section_title": f"Excel Sheet: {sheet_name}" + (f" (Part {chunk_idx + 1}/{num_chunks})" if num_chunks > 1 else ""),
                "sheet_name": sheet_name,
                "sheet_index": sheet_index,
                "row_number": start_row + 1,
                "total_rows": total_rows,
                "column_headers": column_headers,
                "excel_strategy": "table_as_chunk",
                "chunk_index": chunk_idx,
                "total_chunks": num_chunks
            })
        
        return units
    
    def _excel_to_markdown_table(self, df: pd.DataFrame, column_headers: List[str]) -> str:
        """Convert DataFrame to Markdown table format."""
        lines = []
        
        # Header row
        lines.append("| " + " | ".join(column_headers) + " |")
        
        # Separator row
        lines.append("| " + " | ".join(["---"] * len(column_headers)) + " |")
        
        # Data rows
        for _, row in df.iterrows():
            values = [str(row[col]).strip() or "-" for col in column_headers]
            lines.append("| " + " | ".join(values) + " |")
        
        return "\n".join(lines)

    def _load_docx_units(self, path: str) -> List[Dict[str, Any]]:
        """Load DOCX file content with structure preservation."""
        if docx is None:
            raise ImportError("python-docx is required for DOCX support")
            
        document = docx.Document(path)
        units = []
        current_section = []
        current_title = "Document Content"
        
        for element in document.paragraphs:
            text = element.text.strip()
            if not text:
                continue
                
            # Check if this might be a heading (simple heuristic)
            if len(text) < 100 and (element.style.name.startswith('Heading') or 
                                   text.isupper() or 
                                   (len(text.split()) <= 8 and not text.endswith('.'))):
                # Save previous section
                if current_section:
                    content = "\n".join(current_section)
                    units.append({
                        "content": content,
                        "section_title": current_title
                    })
                    current_section = []
                
                current_title = text
            else:
                current_section.append(text)
        
        # Add final section
        if current_section:
            content = "\n".join(current_section)
            units.append({
                "content": content,
                "section_title": current_title
            })
            
        return units

    def _load_image_units(self, path: str) -> List[Dict[str, Any]]:
        """Load and process image files."""
        if not PIL_AVAILABLE:
            raise ImportError("Pillow (PIL) is required for image support")
            
        try:
            # Basic image info
            file_stats = os.stat(path)
            image_info = {
                "content": "",
                "section_title": f"Image: {os.path.basename(path)}",
                "image_path": path,
                "image_size_bytes": file_stats.st_size
            }
            
            # Try to open and process the image
            try:
                with Image.open(path) as img:
                    # Basic image metadata
                    image_info.update({
                        "image_dimensions": img.size,
                        "image_format": img.format,
                        "image_mode": img.mode
                    })
                    
                    # Process image based on strategy
                    content_parts = []
                    
                    if self.image_processing_strategy in [ImageProcessingStrategy.OCR_TEXT, ImageProcessingStrategy.COMBINED]:
                        ocr_text = self._extract_text_from_image(img)
                        if ocr_text.strip():
                            content_parts.append(f"Text content (OCR):\n{ocr_text}")
                    
                    if self.image_processing_strategy in [ImageProcessingStrategy.DESCRIPTION, ImageProcessingStrategy.COMBINED]:
                        description = self._generate_image_description(img, path)
                        if description.strip():
                            content_parts.append(f"Image description:\n{description}")
                    
                    if self.image_processing_strategy in [ImageProcessingStrategy.METADATA, ImageProcessingStrategy.COMBINED]:
                        metadata_text = self._extract_image_metadata(img)
                        if metadata_text.strip():
                            content_parts.append(f"Technical metadata:\n{metadata_text}")
                    
                    if self.image_processing_strategy == ImageProcessingStrategy.VISUAL_ANALYSIS:
                        visual_analysis = self._perform_visual_analysis(img, path)
                        if visual_analysis.strip():
                            content_parts.append(f"Visual analysis:\n{visual_analysis}")
                    
                    # Combine all content
                    if content_parts:
                        image_info["content"] = "\n\n".join(content_parts)
                    else:
                        # Fallback: basic image information
                        image_info["content"] = f"Image file: {os.path.basename(path)}\nDimensions: {img.size[0]}x{img.size[1]}\nFormat: {img.format}\nMode: {img.mode}"
                        
            except Exception as e:
                # If image processing fails, create basic info
                image_info["content"] = f"Image file: {os.path.basename(path)}\nNote: Could not process image content ({str(e)})"
                
            return [image_info]
            
        except Exception as e:
            # Return error info if file can't be read
            return [{
                "content": f"Image file: {os.path.basename(path)}\nError: Could not read image file ({str(e)})",
                "section_title": f"Image (Error): {os.path.basename(path)}"
            }]

    def _extract_text_from_image(self, img: Image.Image) -> str:
        """Extract text from image using OCR."""
        if not TESSERACT_AVAILABLE:
            return "[OCR not available - pytesseract not installed]"
        
        try:
            # Resize image if too large for better OCR performance
            if img.size[0] > self.max_image_size[0] or img.size[1] > self.max_image_size[1]:
                img.thumbnail(self.max_image_size, Image.Resampling.LANCZOS)
            
            # Convert to RGB if necessary
            if img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Perform OCR
            text = pytesseract.image_to_string(img, lang=self.ocr_language)
            
            # Get confidence data if available
            try:
                data = pytesseract.image_to_data(img, lang=self.ocr_language, output_type=pytesseract.Output.DICT)
                confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
                if confidences:
                    avg_confidence = sum(confidences) / len(confidences) / 100.0  # Convert to 0-1 scale
                    if avg_confidence < self.image_quality_threshold:
                        text = f"[Low confidence OCR - {avg_confidence:.2f}]\n{text}"
            except:
                pass  # If confidence calculation fails, continue with text
            
            return text.strip()
            
        except Exception as e:
            return f"[OCR Error: {str(e)}]"

    def _generate_image_description(self, img: Image.Image, path: str) -> str:
        """Generate a description of the image content."""
        try:
            # Basic image analysis without external APIs
            description_parts = []
            
            # Image properties
            width, height = img.size
            aspect_ratio = width / height
            
            if aspect_ratio > 1.5:
                orientation = "landscape"
            elif aspect_ratio < 0.67:
                orientation = "portrait"
            else:
                orientation = "square"
            
            description_parts.append(f"A {orientation} image ({width}x{height} pixels)")
            
            # Color analysis
            if img.mode == 'RGB':
                # Get dominant colors
                img_small = img.resize((150, 150))
                colors = img_small.getcolors(maxcolors=256*256*256)
                if colors:
                    dominant_color = max(colors, key=lambda x: x[0])[1]
                    description_parts.append(f"Dominant color: RGB{dominant_color}")
            
            # File format info
            if hasattr(img, 'format'):
                description_parts.append(f"Format: {img.format}")
            
            # Simple visual characteristics
            if img.mode == 'L':
                description_parts.append("Grayscale image")
            elif img.mode == 'RGBA':
                description_parts.append("Image with transparency")
            
            return ". ".join(description_parts) + "."
            
        except Exception as e:
            return f"Basic image file: {os.path.basename(path)}"

    def _extract_image_metadata(self, img: Image.Image) -> str:
        """Extract technical metadata from image."""
        metadata_parts = []
        
        try:
            # Basic image info
            metadata_parts.append(f"Dimensions: {img.size[0]}x{img.size[1]} pixels")
            metadata_parts.append(f"Color mode: {img.mode}")
            metadata_parts.append(f"Format: {img.format}")
            
            # EXIF data if available
            if hasattr(img, '_getexif') and img._getexif():
                exif_data = img._getexif()
                exif_readable = {}
                
                for tag_id, value in exif_data.items():
                    tag = TAGS.get(tag_id, tag_id)
                    if isinstance(tag, str) and tag in ['DateTime', 'Make', 'Model', 'Software', 'ImageWidth', 'ImageLength']:
                        exif_readable[tag] = value
                
                if exif_readable:
                    metadata_parts.append("EXIF Data:")
                    for key, value in exif_readable.items():
                        metadata_parts.append(f"  {key}: {value}")
            
            # Additional PIL info
            if hasattr(img, 'info') and img.info:
                interesting_info = {k: v for k, v in img.info.items() 
                                  if k in ['dpi', 'compression', 'quality'] and v is not None}
                if interesting_info:
                    metadata_parts.append("Additional Info:")
                    for key, value in interesting_info.items():
                        metadata_parts.append(f"  {key}: {value}")
            
            return "\n".join(metadata_parts)
            
        except Exception as e:
            return f"Could not extract metadata: {str(e)}"

    def _perform_visual_analysis(self, img: Image.Image, path: str) -> str:
        """Perform advanced visual analysis of the image."""
        analysis_parts = []
        
        try:
            # Statistical analysis of the image
            width, height = img.size
            total_pixels = width * height
            
            analysis_parts.append(f"Image statistics:")
            analysis_parts.append(f"  Total pixels: {total_pixels:,}")
            analysis_parts.append(f"  Aspect ratio: {width/height:.2f}")
            
            # Brightness and contrast analysis
            if img.mode in ['RGB', 'L']:
                # Convert to grayscale for analysis
                gray_img = img.convert('L')
                
                # Get histogram
                histogram = gray_img.histogram()
                
                # Calculate average brightness
                total = sum(i * v for i, v in enumerate(histogram))
                avg_brightness = total / total_pixels / 255
                
                analysis_parts.append(f"  Average brightness: {avg_brightness:.2f} (0=dark, 1=bright)")
                
                # Simple contrast estimation
                # Calculate standard deviation of pixel values
                pixels = list(gray_img.getdata())
                mean = sum(pixels) / len(pixels)
                variance = sum((p - mean) ** 2 for p in pixels) / len(pixels)
                contrast = (variance ** 0.5) / 255
                
                analysis_parts.append(f"  Contrast estimate: {contrast:.2f} (0=low, 1=high)")
            
            # Color analysis for RGB images
            if img.mode == 'RGB':
                # Sample colors from different regions
                regions = [
                    ("top-left", (0, 0, width//3, height//3)),
                    ("center", (width//3, height//3, 2*width//3, 2*height//3)),
                    ("bottom-right", (2*width//3, 2*height//3, width, height))
                ]
                
                analysis_parts.append("Color regions:")
                for region_name, box in regions:
                    try:
                        region_img = img.crop(box)
                        # Get average color
                        region_data = list(region_img.getdata())
                        if region_data:
                            avg_r = sum(p[0] for p in region_data) / len(region_data)
                            avg_g = sum(p[1] for p in region_data) / len(region_data)
                            avg_b = sum(p[2] for p in region_data) / len(region_data)
                            analysis_parts.append(f"  {region_name}: RGB({avg_r:.0f}, {avg_g:.0f}, {avg_b:.0f})")
                    except:
                        continue
            
            return "\n".join(analysis_parts)
            
        except Exception as e:
            return f"Visual analysis error: {str(e)}"

    def get_statistics(self) -> Dict[str, Any]:
        """Get processing statistics."""
        stats = {
            "processing_stats": self.stats.copy(),
            "chunking_strategy": self.chunking_strategy.value,
            "chunk_size": self.chunk_size,
            "overlap_size": self.overlap_size,
            "tokenizer_available": self.tokenizer is not None,
            # Image processing stats
            "image_processing_strategy": self.image_processing_strategy.value,
            "image_support_available": PIL_AVAILABLE,
            "ocr_support_available": TESSERACT_AVAILABLE,
            "supported_image_formats": self._get_supported_image_formats()
        }
        return stats

        

# import os
# import time
# import pandas as pd
# from concurrent.futures import ThreadPoolExecutor, as_completed

# try:
#     import PyPDF2
# except ImportError:
#     PyPDF2 = None

# try:
#     import docx
# except ImportError:
#     docx = None


# class DocumentLoaderHelper:
#     """
#     Helper to load and extract textual content from various document types
#     (txt, csv, json, pdf, excel, docx) in parallel, with progress tracking.
#     """

#     def __init__(self, chunk_size: int = None, csv_chunksize: int = 1000, max_workers: int = 4):
#         """
#         :param chunk_size: maximum characters per text chunk
#         :param csv_chunksize: number of rows per chunk when loading CSV
#         :param max_workers: threads to use for parallel extraction/chunking
#         """
#         self.chunk_size = chunk_size
#         self.csv_chunksize = csv_chunksize
#         self.max_workers = max_workers

#         # Stats
#         self.stats = {
#             "total_units": 0,
#             "processed_units": 0,
#             "times": {
#                 "load": 0.0,
#                 "chunk": 0.0,
#                 "total": 0.0
#             }
#         }

#     def load(self, file_path: str):
#         """
#         Extract text chunks from the file in parallel and yield them.
#         Updates self.stats with timing and progress.
#         """
#         if not os.path.isfile(file_path):
#             raise FileNotFoundError(f"File not found: {file_path}")
#         start_total = time.perf_counter()

#         ext = os.path.splitext(file_path)[1].lower()

#         # 1) gather raw units
#         t0 = time.perf_counter()
#         if ext == ".txt":
#             units = self._load_txt_units(file_path)
#         elif ext == ".csv":
#             units = self._load_csv_units(file_path)
#         elif ext == ".json":
#             units = self._load_json_units(file_path)
#         elif ext == ".pdf":
#             units = self._load_pdf_units(file_path)
#         elif ext in (".xls", ".xlsx"):
#             units = self._load_excel_units(file_path)
#         elif ext == ".docx":
#             units = self._load_docx_units(file_path)
#         else:
#             raise ValueError(f"Unsupported file type: {ext}")
#         t1 = time.perf_counter()
#         self.stats["times"]["load"] = t1 - t0

#         self.stats["total_units"] = len(units)
#         self.stats["processed_units"] = 0

#         # 2) process units in parallel: chunking if needed
#         t2 = time.perf_counter()
#         with ThreadPoolExecutor(max_workers=self.max_workers) as exe:
#             futures = {exe.submit(self._chunk_unit, u): u for u in units}
#             for future in as_completed(futures):
#                 chunks = future.result()
#                 self.stats["processed_units"] += 1
#                 for c in chunks:
#                     yield c
#         t3 = time.perf_counter()
#         self.stats["times"]["chunk"] = t3 - t2
#         self.stats["times"]["total"] = time.perf_counter() - start_total

#     def _chunk_unit(self, text: str):
#         """
#         Split a single text unit into chunk_size pieces.
#         """
#         if not self.chunk_size or len(text) <= self.chunk_size:
#             return [text]
#         chunks = []
#         start = 0
#         n = len(text)
#         while start < n:
#             end = min(start + self.chunk_size, n)
#             seg = text[start:end]
#             # break on newline or space
#             cut = max(seg.rfind("\n"), seg.rfind(" "))
#             if cut > 0:
#                 chunks.append(text[start:start + cut])
#                 start += cut
#             else:
#                 chunks.append(seg)
#                 start = end
#         return chunks

#     def _load_txt_units(self, path: str):
#         with open(path, "r", encoding="utf-8", errors="ignore") as f:
#             if self.chunk_size:
#                 units = []
#                 while True:
#                     buf = f.read(self.chunk_size * 4)
#                     if not buf:
#                         break
#                     units.append(buf)
#                 return units
#             return [f.read()]

#     def _load_csv_units(self, path: str):
#         units = []
#         for df_chunk in pd.read_csv(path, chunksize=self.csv_chunksize, dtype=str):
#             text = df_chunk.fillna("").astype(str).agg(" ".join, axis=1)
#             units.append("\n".join(text.tolist()))
#         return units

#     def _load_json_units(self, path: str):
#         """
#         Load JSON file and produce text units per record.
#         Supports both arrays-of-objects and object-of-arrays.
#         """
#         df = pd.read_json(path, dtype=str)
#         # Normalize nested structures into flat table
#         df = pd.json_normalize(df.to_dict(orient="records"))
#         rows = df.fillna("").astype(str).agg(" ".join, axis=1)
#         return ["\n".join(rows.tolist())]

#     def _load_pdf_units(self, path: str):
#         if PyPDF2 is None:
#             raise ImportError("PyPDF2 is required for PDF support")
#         units = []
#         with open(path, "rb") as f:
#             reader = PyPDF2.PdfReader(f)
#             for page in reader.pages:
#                 units.append(page.extract_text() or "")
#         return units

#     def _load_excel_units(self, path: str):
#         df = pd.read_excel(path, dtype=str, engine="openpyxl")
#         rows = df.fillna("").astype(str).agg(" ".join, axis=1)
#         return ["\n".join(rows.tolist())]

#     def _load_docx_units(self, path: str):
#         if docx is None:
#             raise ImportError("python-docx is required for DOCX support")
#         document = docx.Document(path)
#         return [p.text for p in document.paragraphs if p.text]
