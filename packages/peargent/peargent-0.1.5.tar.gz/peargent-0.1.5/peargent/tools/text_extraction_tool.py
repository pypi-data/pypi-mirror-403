"""
Text Extraction Tool for Peargent
Extracts plain text from HTML, PDF, and DOCX files with optional metadata.
"""

import os
from pathlib import Path
from typing import Dict, Any, Optional
import re
import ipaddress
from urllib.parse import urlparse

from peargent import Tool


def _validate_url(url: str) -> None:
    """
    Validate URL to prevent SSRF attacks.
    
    Args:
        url: URL to validate
        
    Raises:
        ValueError: If URL is invalid or potentially dangerous
    """
    try:
        parsed = urlparse(url)
        
        # Only allow http and https schemes
        if parsed.scheme not in ('http', 'https'):
            raise ValueError(f"Only HTTP and HTTPS URLs are allowed, got: {parsed.scheme}")
        
        # Check if hostname exists
        if not parsed.hostname:
            raise ValueError("URL must have a valid hostname")
        
        # Block localhost and loopback addresses
        hostname_lower = parsed.hostname.lower()
        if hostname_lower in ('localhost', '127.0.0.1', '::1', '0.0.0.0'):
            raise ValueError("Access to localhost is not allowed")
        
        # Try to resolve and check if it's a private IP
        try:
            # Check if hostname is an IP address
            ip = ipaddress.ip_address(parsed.hostname)
        except ValueError:
            # Not an IP address, it's a hostname - that's okay
            # We could add DNS resolution here, but it adds complexity
            # and might block legitimate internal hostnames
            pass
        else:
            # Block private, loopback, link-local, and multicast addresses
            if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_multicast or ip.is_reserved:
                raise ValueError(f"Access to private/internal IP addresses is not allowed: {ip}")
            
    except Exception as e:
        if isinstance(e, ValueError) and "not allowed" in str(e):
            raise
        raise ValueError(f"Invalid URL: {e}")


def extract_text(
    file_path: str,
    extract_metadata: bool = False,
    max_length: Optional[int] = None
) -> Dict[str, Any]:
    """
    Extract plain text from various document formats.
    
    Supports: HTML, PDF, DOCX, TXT, MD
    
    Args:
        file_path: Path to the file or URL to extract text from
        extract_metadata: Whether to extract metadata (title, author, etc.)
        max_length: Maximum text length to return (truncates if exceeded)
        
    Returns:
        Dictionary containing:
            - text: Extracted plain text
            - metadata: Dict with file info (if extract_metadata=True)
            - format: Detected file format
            - success: Boolean indicating success
            - error: Error message if any
            
    Example:
        >>> result = extract_text("document.pdf")
        >>> print(result["text"])
        >>> print(result["format"])
    """
    try:
        # Validate URL if it's a URL
        if file_path.startswith(("http://", "https://")):
            _validate_url(file_path)
        
        # Validate file exists (if not a URL)
        if not file_path.startswith(("http://", "https://")):
            if not os.path.exists(file_path):
                return {
                    "text": "",
                    "metadata": {},
                    "format": "unknown",
                    "success": False,
                    "error": f"File not found: {file_path}"
                }
        
        # Detect format
        file_format = _detect_format(file_path)
        
        # Extract based on format
        if file_format == "html":
            text, metadata = _extract_html(file_path, extract_metadata)
        elif file_format == "pdf":
            text, metadata = _extract_pdf(file_path, extract_metadata)
        elif file_format == "docx":
            text, metadata = _extract_docx(file_path, extract_metadata)
        elif file_format in ["txt", "md"]:
            text, metadata = _extract_text_file(file_path, extract_metadata)
        else:
            return {
                "text": "",
                "metadata": {},
                "format": file_format,
                "success": False,
                "error": f"Unsupported format: {file_format}"
            }
        
        # Apply max_length if specified
        if max_length and len(text) > max_length:
            text = text[:max_length] + "..."
        
        return {
            "text": text.strip(),
            "metadata": metadata if extract_metadata else {},
            "format": file_format,
            "success": True,
            "error": None
        }
        
    except Exception as e:
        return {
            "text": "",
            "metadata": {},
            "format": "unknown",
            "success": False,
            "error": str(e)
        }


def _detect_format(file_path: str) -> str:
    """Detect file format from extension."""
    if file_path.startswith(("http://", "https://")):
        # Assume HTML for URLs
        return "html"
    
    ext = Path(file_path).suffix.lower()
    format_map = {
        ".html": "html",
        ".htm": "html",
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".txt": "txt",
        ".md": "md"
    }
    return format_map.get(ext, "unknown")


def _extract_html(file_path: str, extract_metadata: bool) -> tuple[str, Dict[str, Any]]:
    """Extract text from HTML file or URL."""
    try:
        from bs4 import BeautifulSoup  # type: ignore
    except ImportError:
        raise ImportError(
            "beautifulsoup4 is required for HTML extraction. "
            "Install it with: pip install beautifulsoup4"
        )
    
    # Read HTML content
    if file_path.startswith(("http://", "https://")):
        import urllib.request
        with urllib.request.urlopen(file_path, timeout=30) as response:
            content = response.read().decode('utf-8', errors='ignore')
    else:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
    
    soup = BeautifulSoup(content, 'html.parser')
    
    # Remove script and style elements
    for script in soup(["script", "style"]):
        script.decompose()
    
    # Extract text
    text = soup.get_text(separator='\n', strip=True)
    
    # Clean up extra whitespace
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    # Extract metadata if requested
    metadata = {}
    if extract_metadata:
        title_tag = soup.find('title')
        metadata['title'] = title_tag.string if title_tag else ""
        
        # Meta tags
        meta_desc = soup.find('meta', attrs={'name': 'description'})
        if meta_desc and meta_desc.get('content'):
            metadata['description'] = meta_desc.get('content')
        
        meta_author = soup.find('meta', attrs={'name': 'author'})
        if meta_author and meta_author.get('content'):
            metadata['author'] = meta_author.get('content')
        
        metadata['word_count'] = len(text.split())
        metadata['char_count'] = len(text)
    
    return text, metadata


def _extract_pdf(file_path: str, extract_metadata: bool) -> tuple[str, Dict[str, Any]]:
    """Extract text from PDF file."""
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        raise ImportError(
            "pypdf is required for PDF extraction. "
            "Install it with: pip install pypdf>=6.0.0 or pip install peargent[text-extraction]"
        )
    
    text_parts = []
    metadata = {}
    
    with open(file_path, 'rb') as f:
        reader = PdfReader(f)
        
        # Extract metadata
        if extract_metadata:
            info = reader.metadata
            if info:
                metadata['title'] = info.get('/Title', '')
                metadata['author'] = info.get('/Author', '')
                metadata['subject'] = info.get('/Subject', '')
                metadata['creator'] = info.get('/Creator', '')
                metadata['producer'] = info.get('/Producer', '')
                metadata['creation_date'] = str(info.get('/CreationDate', ''))
            metadata['page_count'] = len(reader.pages)
        
        # Extract text from all pages
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
    
    full_text = '\n\n'.join(text_parts)
    
    if extract_metadata:
        metadata['word_count'] = len(full_text.split())
        metadata['char_count'] = len(full_text)
    
    return full_text, metadata


def _extract_docx(file_path: str, extract_metadata: bool) -> tuple[str, Dict[str, Any]]:
    """Extract text from DOCX file."""
    try:
        from docx import Document  # type: ignore
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX extraction. "
            "Install it with: pip install python-docx"
        )
    
    doc = Document(file_path)
    
    # Extract text from paragraphs
    text_parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    full_text = '\n\n'.join(text_parts)
    
    # Extract metadata
    metadata = {}
    if extract_metadata:
        core_props = doc.core_properties
        metadata['title'] = core_props.title or ""
        metadata['author'] = core_props.author or ""
        metadata['subject'] = core_props.subject or ""
        metadata['created'] = str(core_props.created) if core_props.created else ""
        metadata['modified'] = str(core_props.modified) if core_props.modified else ""
        metadata['word_count'] = len(full_text.split())
        metadata['char_count'] = len(full_text)
        metadata['paragraph_count'] = len(doc.paragraphs)
    
    return full_text, metadata


def _extract_text_file(file_path: str, extract_metadata: bool) -> tuple[str, Dict[str, Any]]:
    """Extract text from plain text or markdown file."""
    # Try different encodings
    encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
    
    text = ""
    encoding_used = "utf-8"
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                text = f.read()
            encoding_used = encoding
            break
        except (UnicodeDecodeError, LookupError):
            continue
    
    if not text:
        # Fallback: read as binary and decode with errors='ignore'
        with open(file_path, 'rb') as f:
            text = f.read().decode('utf-8', errors='ignore')
    
    metadata = {}
    if extract_metadata:
        metadata['encoding'] = encoding_used
        metadata['word_count'] = len(text.split())
        metadata['char_count'] = len(text)
        metadata['line_count'] = len(text.splitlines())
        
        # Try to extract title from first line (for markdown)
        first_line = text.split('\n')[0].strip()
        if first_line.startswith('#'):
            metadata['title'] = first_line.lstrip('#').strip()
    
    return text, metadata


class TextExtractionTool(Tool):
    """
    Tool for extracting plain text from various document formats.
    
    Supports: HTML, PDF, DOCX, TXT, MD files and URLs
    
    Example:
        >>> from peargent.tools import TextExtractionTool
        >>> tool = TextExtractionTool()
        >>> result = tool.run({"file_path": "document.pdf", "extract_metadata": True})
        >>> print(result["text"])
    """
    
    def __init__(self):
        super().__init__(
            name="extract_text",
            description=(
                "Extract plain text from HTML, PDF, DOCX, TXT, and Markdown files. "
                "Can also extract from URLs. Optionally extracts metadata like title, "
                "author, page count, etc. Optional parameters: extract_metadata (bool, default: False), "
                "max_length (int, optional): maximum text length."
            ),
            input_parameters={
                "file_path": str
            },
            call_function=extract_text
        )


# Create default instance for easy import
text_extractor = TextExtractionTool()
