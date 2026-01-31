#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import argparse
import json
import os
import subprocess
import sys
import time
import fnmatch

from pkgmgr import config

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except Exception:
    Document = None
    Pt = None
    RGBColor = None


def _load_pkg_yaml(pkg_dir, pkg_yaml):
    if pkg_yaml:
        return pkg_yaml
    if not pkg_dir:
        return None
    return os.path.join(pkg_dir, "pkg.yaml")


def _read_update_json(path):
    with open(path, "rb") as f:
        raw = f.read()
    for encoding in ("utf-8", "euc-kr", "cp949"):
        try:
            return json.loads(raw.decode(encoding))
        except Exception:
            continue
    return json.loads(raw.decode("utf-8", errors="replace"))


def _find_latest_update(pkg_id):
    updates_dir = os.path.join(config.DEFAULT_STATE_DIR, "pkg", str(pkg_id), "updates")
    if not os.path.isdir(updates_dir):
        return None
    candidates = [name for name in os.listdir(updates_dir) if name.startswith("update-") and name.endswith(".json")]
    if not candidates:
        return None
    candidates.sort()
    return os.path.join(updates_dir, candidates[-1])


def _decode_output(raw):
    if isinstance(raw, str):
        return raw.strip()
    for encoding in ("utf-8", "euc-kr", "cp949"):
        try:
            return raw.decode(encoding).strip()
        except Exception:
            continue
    return raw.decode("utf-8", errors="replace").strip()


def _git_rev_parse(repo_root, args):
    try:
        out = subprocess.check_output(
            ["git", "rev-parse"] + args,
            cwd=repo_root,
            stderr=subprocess.STDOUT,
        )
        return _decode_output(out)
    except Exception:
        return None


def _git_log_first_commit(repo_root, keyword):
    cmd = [
        "git",
        "--no-pager",
        "log",
        "--reverse",
        "--format=%H",
        "--grep=%s" % keyword,
        "--regexp-ignore-case",
        "--all",
        "--",
    ]
    try:
        out = subprocess.check_output(cmd, cwd=repo_root, stderr=subprocess.STDOUT)
    except Exception:
        return None
    lines = _decode_output(out).splitlines()
    return lines[0].strip() if lines else None


def _resolve_repo_root(pkg_cfg, pkg_dir, update_data):
    repo_root = (pkg_cfg.get("git") or {}).get("repo_root") if pkg_cfg else None
    if repo_root:
        repo_root = os.path.expanduser(repo_root)
        if not os.path.isabs(repo_root):
            repo_root = os.path.abspath(os.path.join(pkg_dir or os.getcwd(), repo_root))
        if os.path.isdir(repo_root):
            return repo_root
    git_files = list(((update_data.get("checksums") or {}).get("git_files") or {}).keys())
    for path in git_files:
        base = os.path.dirname(path)
        root = _git_rev_parse(base, ["--show-toplevel"])
        if root:
            return root
    root = _git_rev_parse(os.getcwd(), ["--show-toplevel"])
    return root


def _resolve_pkg_output_dir(pkg_id, pkg_cfg, config_path=None):
    if pkg_cfg:
        pkg_block = pkg_cfg.get("pkg") or {}
        root = pkg_block.get("root")
        if root:
            return os.path.abspath(os.path.expanduser(str(root)))
    if config_path:
        try:
            main_cfg = config.load_main(path=config_path, allow_interactive=False)
        except Exception:
            main_cfg = None
    else:
        try:
            main_cfg = config.load_main(allow_interactive=False)
        except Exception:
            main_cfg = None
    if not main_cfg:
        return None
    release_root = main_cfg.get("pkg_release_root")
    if not release_root:
        return None
    return os.path.abspath(os.path.expanduser(os.path.join(release_root, str(pkg_id))))


def _collect_keywords(git_info):
    keywords = set()
    for kw in git_info.get("keywords") or []:
        if kw:
            keywords.add(str(kw))
    for commit in git_info.get("commits") or []:
        for kw in commit.get("keywords") or []:
            if kw:
                keywords.add(str(kw))
    return sorted(keywords)


def _select_keyword(git_info):
    keywords = _collect_keywords(git_info)
    if len(keywords) == 1:
        return keywords[0], False
    if keywords:
        return keywords[0], True
    return None, False


def _parse_commit_time(value):
    if not value:
        return None
    try:
        return time.strptime(value, "%a %b %d %H:%M:%S %Y %z")
    except Exception:
        return None


def _find_first_commit(commits, keyword):
    matches = [c for c in commits if keyword in (c.get("keywords") or [])]
    if not matches:
        return None
    dated = []
    for commit in matches:
        dt = _parse_commit_time(commit.get("authored_at") or commit.get("date"))
        dated.append((dt, commit))
    dated = [item for item in dated if item[0] is not None]
    if dated:
        return min(dated, key=lambda item: item[0])[1]
    return sorted(matches, key=lambda c: c.get("hash") or "")[0]


def _collect_files(commits, keyword):
    files = set()
    for commit in commits:
        if keyword not in (commit.get("keywords") or []):
            continue
        for path in commit.get("files") or []:
            if path:
                files.add(path)
    return sorted(files)


def _parse_ignore_patterns(values):
    patterns = []
    for raw in values or []:
        if raw is None:
            continue
        if isinstance(raw, (list, tuple)):
            items = raw
        else:
            items = str(raw).replace("\n", ",").replace(";", ",").split(",")
        for item in items:
            item = item.strip()
            if item:
                patterns.append(item)
    return patterns


def _is_ignored(path, repo_root, patterns):
    if not patterns:
        return False
    candidates = [path]
    basename = os.path.basename(path)
    if basename:
        candidates.append(basename)
    if repo_root and not os.path.isabs(path):
        candidates.append(os.path.join(repo_root, path))
    if repo_root and os.path.isabs(path):
        try:
            candidates.append(os.path.relpath(path, repo_root))
        except Exception:
            pass
    for candidate in candidates:
        for pattern in patterns:
            if fnmatch.fnmatch(candidate, pattern):
                return True
    return False


def _git_parent(repo_root, commit_hash):
    parent = _git_rev_parse(repo_root, ["%s^" % commit_hash])
    return parent or commit_hash


def _git_file_exists(repo_root, commit_hash, path):
    cmd = ["git", "cat-file", "-e", "%s:%s" % (commit_hash, path)]
    result = subprocess.run(cmd, cwd=repo_root, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    return result.returncode == 0


def _git_diff_no_index(repo_root, abs_path):
    cmd = ["git", "--no-pager", "diff", "--no-index", "/dev/null", abs_path]
    result = subprocess.run(cmd, cwd=repo_root, stdout=subprocess.PIPE, stderr=subprocess.STDOUT)
    return _decode_output(result.stdout)


def _git_log_name_status(repo_root, path):
    cmd = [
        "git",
        "--no-pager",
        "log",
        "--follow",
        "--name-status",
        "--format=%H",
        "--",
        path,
    ]
    result = subprocess.run(cmd, cwd=repo_root, stdout=subprocess.PIPE, stderr=subprocess.STDOUT)
    return _decode_output(result.stdout)


def _resolve_paths_with_history(repo_root, path):
    output = _git_log_name_status(repo_root, path)
    if not output:
        return [path], None
    current = None
    paths = [path]
    last_commit_with_file = None
    for line in output.splitlines():
        line = line.strip()
        if not line:
            continue
        if "\t" not in line and " " not in line:
            last_commit_with_file = line
            continue
        parts = line.split("\t")
        status = parts[0]
        if status.startswith("R") and len(parts) >= 3:
            old_path, new_path = parts[1], parts[2]
            current = new_path
            paths.append(new_path)
            continue
        if status.startswith("D") and len(parts) >= 2:
            current = parts[1]
            paths.append(current)
    if current:
        paths.append(current)
    # keep unique order
    seen = set()
    result = []
    for item in paths:
        if item not in seen:
            seen.add(item)
            result.append(item)
    return result, last_commit_with_file


def _pick_head_path(repo_root, path):
    if _git_file_exists(repo_root, "HEAD", path):
        return path
    candidates, _ = _resolve_paths_with_history(repo_root, path)
    for candidate in candidates:
        if _git_file_exists(repo_root, "HEAD", candidate):
            return candidate
    return None


def _git_diff(repo_root, start_commit, path):
    cmd = [
        "git",
        "--no-pager",
        "diff",
        "-U3",
        start_commit,
        "HEAD",
        "-M",
        "-C",
        "--",
        path,
    ]
    try:
        out = subprocess.check_output(cmd, cwd=repo_root, stderr=subprocess.STDOUT)
        diff_text = _decode_output(out)
        if diff_text.strip():
            return diff_text
        head_has = _git_file_exists(repo_root, "HEAD", path)
        start_has = _git_file_exists(repo_root, start_commit, path)
        if head_has and not start_has:
            abs_path = os.path.join(repo_root, path)
            if os.path.exists(abs_path):
                return _git_diff_no_index(repo_root, abs_path)
        if not head_has:
            candidates, last_commit = _resolve_paths_with_history(repo_root, path)
            for candidate in candidates:
                if candidate == path:
                    continue
                diff_retry = _git_diff(repo_root, start_commit, candidate)
                if diff_retry.strip():
                    return diff_retry
            if last_commit and _git_file_exists(repo_root, last_commit, path):
                cmd_deleted = [
                    "git",
                    "--no-pager",
                    "diff",
                    "-U3",
                    start_commit,
                    last_commit,
                    "-M",
                    "-C",
                    "--",
                    path,
                ]
                out_deleted = subprocess.check_output(
                    cmd_deleted, cwd=repo_root, stderr=subprocess.STDOUT
                )
                return _decode_output(out_deleted)
        return diff_text
    except Exception as exc:
        return "[export_source_review] git diff failed for %s: %s" % (path, str(exc))


def _add_diff_table(doc, file_path, diff_text):
    table = doc.add_table(rows=2, cols=1)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    header = table.cell(0, 0)
    header.text = file_path
    if header.paragraphs and header.paragraphs[0].runs:
        header.paragraphs[0].runs[0].bold = True
    body = table.cell(1, 0)
    body.text = ""
    paragraph = body.paragraphs[0]
    text = diff_text or "No changes in range."
    lines = text.splitlines() or [text]
    for idx, line in enumerate(lines):
        run = paragraph.add_run(line)
        if Pt is not None:
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        if RGBColor is not None:
            if line.startswith("+") and not line.startswith("+++"):
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            elif line.startswith("-") and not line.startswith("---"):
                run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        if idx < len(lines) - 1:
            run.add_break()
    doc.add_paragraph("")


def main(argv=None):
    argv = argv if argv is not None else sys.argv[1:]
    parser = argparse.ArgumentParser(
        description="Export per-file diff history for a keyword into a Word document."
    )
    parser.add_argument("--config", help="pkgmgr main config path")
    parser.add_argument("--pkg-id", required=True, help="pkg id (used to locate latest update JSON)")
    parser.add_argument("--docx", required=True, help="output docx path")
    parser.add_argument("--ignore", action="append", help="glob patterns to ignore")
    args = parser.parse_args(argv)

    if Document is None:
        print("[export_source_review] python-docx is required (pip install python-docx)")
        return 1

    pkg_yaml = None
    pkg_dir = None
    pkg_cfg = None

    update_path = _find_latest_update(args.pkg_id)
    if not update_path:
        print("[export_source_review] update json not found for pkg: %s" % args.pkg_id)
        return 1
    update_path = os.path.abspath(os.path.expanduser(update_path))
    if not os.path.exists(update_path):
        print("[export_source_review] update json not found: %s" % update_path)
        return 1

    data = _read_update_json(update_path)
    git_info = data.get("git") or {}
    commits = git_info.get("commits") or []
    keyword, multi_keywords = _select_keyword(git_info)
    if not keyword:
        print("[export_source_review] keyword not found in update json for pkg: %s" % args.pkg_id)
        return 1
    if multi_keywords:
        print("[export_source_review] multiple keywords found; using %s" % keyword)

    repo_root = _resolve_repo_root(pkg_cfg, pkg_dir, data)
    if not repo_root:
        print("[export_source_review] repo root not found for pkg: %s" % args.pkg_id)
        return 1

    commit_hash = _git_log_first_commit(repo_root, keyword)
    if not commit_hash:
        first_commit = _find_first_commit(commits, keyword)
        if not first_commit:
            print("[export_source_review] no commits found for keyword: %s" % keyword)
            return 1
        commit_hash = first_commit.get("hash") or first_commit.get("commit")
    if not commit_hash:
        print("[export_source_review] commit hash missing for keyword: %s" % keyword)
        return 1
    start_commit = _git_parent(repo_root, commit_hash)
    file_list = _collect_files(commits, keyword)
    if not file_list:
        git_files = list(((data.get("checksums") or {}).get("git_files") or {}).keys())
        for path in git_files:
            if path.startswith(repo_root):
                file_list.append(os.path.relpath(path, repo_root))
        file_list = sorted(set(file_list))

    if not file_list:
        print("[export_source_review] no files found for keyword: %s" % keyword)
        return 1

    ignore_patterns = _parse_ignore_patterns([args.ignore, os.environ.get("PKGMGR_REVIEW_IGNORE")])

    doc = Document()
    doc.add_paragraph("Source Review Export")
    doc.add_paragraph("Keyword: %s" % keyword)
    doc.add_paragraph("Range: %s..HEAD" % start_commit)
    doc.add_paragraph("Update JSON: %s" % update_path)
    doc.add_paragraph("")

    for path in file_list:
        if _is_ignored(path, repo_root, ignore_patterns):
            print("[export_source_review] skip (ignored): %s" % path)
            continue
        head_path = _pick_head_path(repo_root, path)
        if not head_path:
            print("[export_source_review] skip (file deleted or untracked in HEAD): %s" % path)
            continue
        if _is_ignored(head_path, repo_root, ignore_patterns):
            print("[export_source_review] skip (ignored): %s" % head_path)
            continue
        diff_text = _git_diff(repo_root, start_commit, head_path)
        _add_diff_table(doc, head_path, diff_text)

    out_path = args.docx
    if not out_path.lower().endswith(".docx"):
        out_path = out_path + ".docx"
    config_path = args.config or os.environ.get("PKGMGR_CONFIG")
    if os.sep not in out_path:
        base_dir = _resolve_pkg_output_dir(args.pkg_id, pkg_cfg, config_path=config_path)
        if not base_dir:
            base_dir = pkg_dir or os.path.join(config.DEFAULT_STATE_DIR, "pkg", str(args.pkg_id))
        export_dir = os.path.join(base_dir, "export")
        out_path = os.path.join(export_dir, out_path)
    out_dir = os.path.dirname(os.path.abspath(out_path))
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir)
    doc.save(out_path)
    print("[export_source_review] wrote %s" % out_path)
    return 0


if __name__ == "__main__":
    sys.exit(main())
