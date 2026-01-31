from cannect.schema.candb import CanMessage

from docx.document import Document
from docx.styles.styles import Styles
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from typing import List


class Message:
    COLORS = [
        'C4BC96',
        '8DB3E2',
        'B8CCE4',
        'E5B8B7',
        'D6E3BC',
        'CCC0D9',
        'FBD4B4',
        'B6DDE8'
    ]

    def __init__(self, doc:Document):
        self.doc:Document = doc
        return

    def _coloring(self, msg:CanMessage) -> List[str]:
        colors = ['A6A6A6'] * msg['DLC'] * 8
        for n in range(msg['DLC'] * 8):
            for m, sig in enumerate(msg):
                if sig['StartBit'] <= n < (sig['StartBit'] + sig['Length']):
                    colors[n] = self.COLORS[m % len(self.COLORS)]
                    break
        return colors

    @property
    def messageOverviewStyle(self) -> Styles:
        if not hasattr(self, "_overview_style"):
            style = self.doc.styles.add_style("_overview_style", WD_PARAGRAPH_ALIGNMENT.CENTER)
            style.font.name = "현대산스 text"
            style.font.size = Pt(11)
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.font.bold = True
            style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            self.__setattr__("_overview_style", style)
        return self.__getattribute__("_overview_style")

    @property
    def messageOverviewContentStyle(self) -> Styles:
        if not hasattr(self, "_overview_content_style"):
            style = self.doc.styles.add_style("_overview_content_style", WD_PARAGRAPH_ALIGNMENT.CENTER)
            style.font.name = "현대산스 text"
            style.font.size = Pt(11)
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.font.bold = False
            style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            self.__setattr__("_overview_content_style", style)
        return self.__getattribute__("_overview_content_style")

    @property
    def messageLayoutNumberStyle(self) -> Styles:
        if not hasattr(self, "_layout_number_style"):
            style = self.doc.styles.add_style("_layout_number_style", WD_PARAGRAPH_ALIGNMENT.CENTER)
            style.font.name = "현대산스 text"
            style.font.size = Pt(9)
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.font.bold = True
            style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            self.__setattr__("_layout_number_style", style)
        return self.__getattribute__("_layout_number_style")

    @property
    def messageLayoutStyle(self) -> Styles:
        if not hasattr(self, "_layout_style"):
            style = self.doc.styles.add_style("_layout_style", WD_PARAGRAPH_ALIGNMENT.CENTER)
            style.font.name = "현대산스 text"
            style.font.size = Pt(7)
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.font.bold = False
            style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            self.__setattr__("_layout_style", style)
        return self.__getattribute__("_layout_style")

    @property
    def signalListHeadingStyle(self) -> Styles:
        if not hasattr(self, "_signal_style"):
            style = self.doc.styles.add_style("_signal_style", WD_PARAGRAPH_ALIGNMENT.CENTER)
            style.font.name = "현대산스 text"
            style.font.size = Pt(10)
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.font.bold = True
            style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            self.__setattr__("_signal_style", style)
        return self.__getattribute__("_signal_style")

    @property
    def signalListStyle(self) -> Styles:
        if not hasattr(self, "_signal_content_style"):
            style = self.doc.styles.add_style("_signal_content_style", WD_PARAGRAPH_ALIGNMENT.CENTER)
            style.font.name = "현대산스 text"
            style.font.size = Pt(10)
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.font.bold = False
            style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            self.__setattr__("_signal_content_style", style)
        return self.__getattribute__("_signal_content_style")

    def addHeading(
        self,
        text:str,
        level:int=1,
        fontSize:int=20,
        bold:bool=True,
        alignCenter:bool=True
    ):
        heading = self.doc.add_heading(text, level=level)
        if alignCenter:
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        runner = heading.runs[0]
        runner.font.name = "현대산스 Text"
        runner.font.size = Pt(fontSize)
        runner.font.bold = bold
        runner.font.color.rgb = RGBColor(0, 0, 0)
        return

    def addMessageHeading(self, msg:CanMessage):
        self.addHeading(f"{msg['ID']}: {msg.name}", level=2, fontSize=14, alignCenter=False)
        return

    def addMessageSpec(self, msg:CanMessage):
        specs = {
            "ECU": "SENDER",
            "ID": "ID",
            "Message": "MESSAGE",
            "DLC": "DLC",
            "Send Type": "SEND TYPE",
            "Cycle Time": "CYCLE TIME [ms]",
            "ByteOrder": "BYTE ORDER",
            "SystemConstant": "SYSCON",
            "Codeword": "CODEWORD"
        }
        self.addHeading("Message Overview", level=3, fontSize=12, bold=True, alignCenter=False)
        table = self.doc.add_table(rows=len(specs), cols=2)
        table.style = 'Table Grid'
        for n, (key, value) in enumerate(specs.items()):
            left = table.rows[n].cells[0]
            left.width = Inches(2.0)
            name = left.paragraphs[0]
            name.text = value
            name.style = self.messageOverviewStyle

            right = table.rows[n].cells[1]
            right.width = self.doc.sections[0].page_width
            text = right.paragraphs[0]
            text.text = str(msg[key])
            text.style = self.messageOverviewContentStyle
        return

    def addMessageLayout(self, msg:CanMessage):
        self.addHeading("Message Layout", level=3, fontSize=12, bold=True, alignCenter=False)
        color = self._coloring(msg)
        table = self.doc.add_table(rows=msg['DLC'] * 2, cols=8)
        for n, row in enumerate(table.rows):
            byte = int(n / 2)
            for m, cell in enumerate(row.cells):
                bit = 8 * byte + 7 - m
                sig = msg.signals[msg.signals["StartBit"] == int(bit)]
                sig = '' if sig.empty else sig.iloc[0]["Signal"]

                text = cell.paragraphs[0]
                text.text = str(bit) if not n % 2 else sig
                text.style = self.messageLayoutNumberStyle if not n % 2 else self.messageLayoutStyle
                cell._tc.get_or_add_tcPr().append(
                    parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color[bit]))
                )
        return

    def addSignalList(self, msg:CanMessage):
        specs = {
            "Signal": "Signal",
            "InterfacedVariable": "@EMS",
            "Length": "Len.",
            "StartBit": "Addr.",
            "Value Type": "Type",
            "Factor": "Factor",
            "Offset": "Offset",
            "Unit": "Unit"
        }
        self.addHeading("Signal List", level=3, fontSize=12, bold=True, alignCenter=False)
        table = self.doc.add_table(rows=len(msg.signals) + 1, cols=8)
        table.style = "Table Grid"
        for n, row in enumerate(table.rows):
            for m, key in enumerate(specs.keys()):
                if not n:
                    text = row.cells[m].paragraphs[0]
                    text.text = specs[key]
                    text.style = self.signalListHeadingStyle
                else:
                    text = row.cells[m].paragraphs[0]
                    text.text = str(msg.signals.iloc[n - 1][key]).replace('"', '')
                    text.style = self.signalListStyle
        return

    def addSignalProperty(self, msg:CanMessage):
        specs = {
            "Signal": "Signal",
            "Definition": "Definition",
            "Value Table": "Value Table"
        }
        self.addHeading("Signal Property", level=3, fontSize=12, bold=True, alignCenter=False)

        for sig in msg:
            table = self.doc.add_table(rows=3, cols=2)
            table.style = 'Table Grid'
            for n, (key, value) in enumerate(specs.items()):
                left = table.rows[n].cells[0]
                left.width = Inches(2.0)
                name = left.paragraphs[0]
                name.text = value
                name.style = self.messageOverviewStyle

                right = table.rows[n].cells[1]
                right.width = self.doc.sections[0].page_width
                text = right.paragraphs[0]
                text.text = sig[key].replace("/0x", "\n0x") if key == "Value Table" else sig[key]
                text.style = self.messageOverviewContentStyle
            self.doc.add_paragraph()
        self.doc.add_paragraph("\n")
        return