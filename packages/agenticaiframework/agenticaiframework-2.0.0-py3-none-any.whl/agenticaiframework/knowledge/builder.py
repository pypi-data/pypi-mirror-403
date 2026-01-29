"""
Knowledge Builder - Create knowledge bases from various sources.

Supports:
- Web search and scraping
- API data ingestion
- Documents (PDF, DOCX, TXT, MD, etc.)
- Images (OCR and image analysis)
- Any file types
- Embedding generation for vector databases
"""

import os
import re
import json
import uuid
import logging
import mimetypes
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from typing import Any, Callable, Dict, Iterator, List, Optional, Tuple, Union
from enum import Enum
from datetime import datetime
from pathlib import Path

logger = logging.getLogger(__name__)


class SourceType(Enum):
    """Knowledge source types."""
    WEB_SEARCH = "web_search"
    WEB_PAGE = "web_page"
    API = "api"
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    MARKDOWN = "markdown"
    HTML = "html"
    IMAGE = "image"
    JSON = "json"
    CSV = "csv"
    EXCEL = "excel"
    CODE = "code"
    AUDIO = "audio"
    VIDEO = "video"
    CUSTOM = "custom"


@dataclass
class KnowledgeChunk:
    """A chunk of knowledge with metadata."""
    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    content: str = ""
    source: str = ""
    source_type: SourceType = SourceType.TXT
    metadata: Dict[str, Any] = field(default_factory=dict)
    embedding: Optional[List[float]] = None
    timestamp: str = field(default_factory=lambda: datetime.now().isoformat())
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for storage."""
        return {
            "id": self.id,
            "content": self.content,
            "source": self.source,
            "source_type": self.source_type.value,
            "metadata": self.metadata,
            "embedding": self.embedding,
            "timestamp": self.timestamp,
        }
    
    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> "KnowledgeChunk":
        return cls(
            id=data.get("id", str(uuid.uuid4())),
            content=data.get("content", ""),
            source=data.get("source", ""),
            source_type=SourceType(data.get("source_type", "txt")),
            metadata=data.get("metadata", {}),
            embedding=data.get("embedding"),
            timestamp=data.get("timestamp", datetime.now().isoformat()),
        )


@dataclass
class EmbeddingOutput:
    """Embedding output ready for vector database storage."""
    id: str
    embedding: List[float]
    content: str
    metadata: Dict[str, Any]
    
    def to_qdrant_point(self, collection_name: str = "default"):
        """Convert to Qdrant point format."""
        return {
            "id": self.id,
            "vector": self.embedding,
            "payload": {
                "content": self.content,
                **self.metadata,
            }
        }
    
    def to_pinecone_vector(self):
        """Convert to Pinecone vector format."""
        return {
            "id": self.id,
            "values": self.embedding,
            "metadata": {
                "content": self.content,
                **self.metadata,
            }
        }
    
    def to_weaviate_object(self, class_name: str = "Document"):
        """Convert to Weaviate object format."""
        return {
            "class": class_name,
            "id": self.id,
            "properties": {
                "content": self.content,
                **self.metadata,
            },
            "vector": self.embedding,
        }
    
    def to_chroma_document(self):
        """Convert to ChromaDB document format."""
        return {
            "ids": [self.id],
            "embeddings": [self.embedding],
            "documents": [self.content],
            "metadatas": [self.metadata],
        }
    
    def to_milvus_entity(self):
        """Convert to Milvus entity format."""
        return {
            "id": self.id,
            "vector": self.embedding,
            "content": self.content,
            **self.metadata,
        }
    
    def to_mongodb_document(self):
        """Convert to MongoDB Atlas Vector Search format."""
        return {
            "_id": self.id,
            "embedding": self.embedding,
            "content": self.content,
            "metadata": self.metadata,
        }
    
    def to_opensearch_document(self, index_name: str = "knowledge"):
        """Convert to OpenSearch format."""
        return {
            "_index": index_name,
            "_id": self.id,
            "_source": {
                "embedding": self.embedding,
                "content": self.content,
                **self.metadata,
            }
        }
    
    def to_pgvector_row(self):
        """Convert to pgvector format (PostgreSQL)."""
        return {
            "id": self.id,
            "embedding": self.embedding,  # Use pgvector extension
            "content": self.content,
            "metadata": json.dumps(self.metadata),
        }


class EmbeddingProvider(ABC):
    """Base class for embedding providers."""
    
    @abstractmethod
    def embed(self, text: str) -> List[float]:
        """Generate embedding for text."""
        pass
    
    @abstractmethod
    def embed_batch(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings for multiple texts."""
        pass
    
    @property
    @abstractmethod
    def dimension(self) -> int:
        """Get embedding dimension."""
        pass


class OpenAIEmbedding(EmbeddingProvider):
    """OpenAI embedding provider."""
    
    def __init__(
        self,
        model: str = "text-embedding-3-small",
        api_key: Optional[str] = None,
    ):
        self.model = model
        self.api_key = api_key or os.getenv("OPENAI_API_KEY")
        self._client = None
        
        # Model dimensions
        self._dimensions = {
            "text-embedding-3-small": 1536,
            "text-embedding-3-large": 3072,
            "text-embedding-ada-002": 1536,
        }
    
    def _get_client(self):
        if not self._client:
            try:
                from openai import OpenAI
                self._client = OpenAI(api_key=self.api_key)
            except ImportError:
                raise ImportError("OpenAI embeddings require: pip install openai")
        return self._client
    
    def embed(self, text: str) -> List[float]:
        client = self._get_client()
        response = client.embeddings.create(input=text, model=self.model)
        return response.data[0].embedding
    
    def embed_batch(self, texts: List[str]) -> List[List[float]]:
        client = self._get_client()
        response = client.embeddings.create(input=texts, model=self.model)
        return [item.embedding for item in response.data]
    
    @property
    def dimension(self) -> int:
        return self._dimensions.get(self.model, 1536)


class AzureOpenAIEmbedding(EmbeddingProvider):
    """Azure OpenAI embedding provider."""
    
    def __init__(
        self,
        deployment: str = "text-embedding-ada-002",
        api_key: Optional[str] = None,
        endpoint: Optional[str] = None,
        api_version: str = "2024-02-01",
    ):
        self.deployment = deployment
        self.api_key = api_key or os.getenv("AZURE_OPENAI_API_KEY")
        self.endpoint = endpoint or os.getenv("AZURE_OPENAI_ENDPOINT")
        self.api_version = api_version
        self._client = None
    
    def _get_client(self):
        if not self._client:
            try:
                from openai import AzureOpenAI
                self._client = AzureOpenAI(
                    api_key=self.api_key,
                    api_version=self.api_version,
                    azure_endpoint=self.endpoint,
                )
            except ImportError:
                raise ImportError("Azure OpenAI requires: pip install openai")
        return self._client
    
    def embed(self, text: str) -> List[float]:
        client = self._get_client()
        response = client.embeddings.create(input=text, model=self.deployment)
        return response.data[0].embedding
    
    def embed_batch(self, texts: List[str]) -> List[List[float]]:
        client = self._get_client()
        response = client.embeddings.create(input=texts, model=self.deployment)
        return [item.embedding for item in response.data]
    
    @property
    def dimension(self) -> int:
        return 1536


class HuggingFaceEmbedding(EmbeddingProvider):
    """HuggingFace local embedding provider."""
    
    def __init__(
        self,
        model_name: str = "sentence-transformers/all-MiniLM-L6-v2",
        device: str = "cpu",
    ):
        self.model_name = model_name
        self.device = device
        self._model = None
        self._tokenizer = None
    
    def _load_model(self):
        if self._model is None:
            try:
                from sentence_transformers import SentenceTransformer
                self._model = SentenceTransformer(self.model_name, device=self.device)
            except ImportError:
                raise ImportError("HuggingFace embeddings require: pip install sentence-transformers")
    
    def embed(self, text: str) -> List[float]:
        self._load_model()
        return self._model.encode(text).tolist()
    
    def embed_batch(self, texts: List[str]) -> List[List[float]]:
        self._load_model()
        return self._model.encode(texts).tolist()
    
    @property
    def dimension(self) -> int:
        self._load_model()
        return self._model.get_sentence_embedding_dimension()


class CohereEmbedding(EmbeddingProvider):
    """Cohere embedding provider."""
    
    def __init__(
        self,
        model: str = "embed-english-v3.0",
        api_key: Optional[str] = None,
        input_type: str = "search_document",
    ):
        self.model = model
        self.api_key = api_key or os.getenv("COHERE_API_KEY")
        self.input_type = input_type
        self._client = None
    
    def _get_client(self):
        if not self._client:
            try:
                import cohere
                self._client = cohere.Client(self.api_key)
            except ImportError:
                raise ImportError("Cohere embeddings require: pip install cohere")
        return self._client
    
    def embed(self, text: str) -> List[float]:
        client = self._get_client()
        response = client.embed(texts=[text], model=self.model, input_type=self.input_type)
        return response.embeddings[0]
    
    def embed_batch(self, texts: List[str]) -> List[List[float]]:
        client = self._get_client()
        response = client.embed(texts=texts, model=self.model, input_type=self.input_type)
        return response.embeddings
    
    @property
    def dimension(self) -> int:
        return 1024 if "v3" in self.model else 4096


class SourceLoader(ABC):
    """Base class for loading data from various sources."""
    
    @abstractmethod
    def load(self, source: str, **kwargs) -> List[KnowledgeChunk]:
        """Load and chunk content from source."""
        pass
    
    @abstractmethod
    def supports(self, source: str) -> bool:
        """Check if loader supports the source."""
        pass


class TextLoader(SourceLoader):
    """Load text files."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def supports(self, source: str) -> bool:
        ext = Path(source).suffix.lower()
        return ext in [".txt", ".text", ".log", ""]
    
    def load(self, source: str, **kwargs) -> List[KnowledgeChunk]:
        with open(source, "r", encoding="utf-8") as f:
            content = f.read()
        
        return self._chunk_text(content, source, SourceType.TXT)
    
    def _chunk_text(
        self,
        text: str,
        source: str,
        source_type: SourceType,
    ) -> List[KnowledgeChunk]:
        """Split text into chunks."""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            chunk_text = text[start:end]
            
            # Try to break at sentence or paragraph
            if end < len(text):
                for sep in ["\n\n", "\n", ". ", " "]:
                    last_sep = chunk_text.rfind(sep)
                    if last_sep > self.chunk_size // 2:
                        chunk_text = chunk_text[:last_sep + len(sep)]
                        break
            
            if chunk_text.strip():
                chunks.append(KnowledgeChunk(
                    content=chunk_text.strip(),
                    source=source,
                    source_type=source_type,
                    metadata={
                        "chunk_index": len(chunks),
                        "start_char": start,
                    }
                ))
            
            start += len(chunk_text) - self.chunk_overlap
            if start >= end:
                start = end
        
        return chunks


class MarkdownLoader(TextLoader):
    """Load Markdown files."""
    
    def supports(self, source: str) -> bool:
        ext = Path(source).suffix.lower()
        return ext in [".md", ".markdown"]
    
    def load(self, source: str, **kwargs) -> List[KnowledgeChunk]:
        with open(source, "r", encoding="utf-8") as f:
            content = f.read()
        
        return self._chunk_text(content, source, SourceType.MARKDOWN)


class PDFLoader(SourceLoader):
    """Load PDF documents."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def supports(self, source: str) -> bool:
        return Path(source).suffix.lower() == ".pdf"
    
    def load(self, source: str, **kwargs) -> List[KnowledgeChunk]:
        try:
            import pypdf
        except ImportError:
            try:
                import PyPDF2 as pypdf
            except ImportError:
                raise ImportError("PDF loading requires: pip install pypdf")
        
        chunks = []
        
        with open(source, "rb") as f:
            reader = pypdf.PdfReader(f)
            
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    chunk = KnowledgeChunk(
                        content=text,
                        source=source,
                        source_type=SourceType.PDF,
                        metadata={
                            "page": page_num + 1,
                            "total_pages": len(reader.pages),
                        }
                    )
                    chunks.append(chunk)
        
        return chunks


class DocxLoader(SourceLoader):
    """Load Word documents."""
    
    def supports(self, source: str) -> bool:
        return Path(source).suffix.lower() in [".docx", ".doc"]
    
    def load(self, source: str, **kwargs) -> List[KnowledgeChunk]:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("DOCX loading requires: pip install python-docx")
        
        doc = Document(source)
        content = "\n".join([para.text for para in doc.paragraphs if para.text])
        
        return [KnowledgeChunk(
            content=content,
            source=source,
            source_type=SourceType.DOCX,
            metadata={"paragraphs": len(doc.paragraphs)}
        )]


class ImageLoader(SourceLoader):
    """Load and extract text from images using OCR or vision models."""
    
    def __init__(
        self,
        ocr_provider: str = "pytesseract",  # or "openai_vision", "azure_vision"
        api_key: Optional[str] = None,
    ):
        self.ocr_provider = ocr_provider
        self.api_key = api_key
    
    def supports(self, source: str) -> bool:
        ext = Path(source).suffix.lower()
        return ext in [".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp"]
    
    def load(self, source: str, **kwargs) -> List[KnowledgeChunk]:
        if self.ocr_provider == "pytesseract":
            return self._load_with_tesseract(source)
        elif self.ocr_provider == "openai_vision":
            return self._load_with_openai_vision(source, **kwargs)
        else:
            return self._load_basic(source)
    
    def _load_basic(self, source: str) -> List[KnowledgeChunk]:
        """Return metadata only without OCR."""
        import base64
        
        with open(source, "rb") as f:
            data = f.read()
            b64 = base64.b64encode(data).decode()
        
        return [KnowledgeChunk(
            content=f"[Image: {Path(source).name}]",
            source=source,
            source_type=SourceType.IMAGE,
            metadata={
                "size_bytes": len(data),
                "base64_prefix": b64[:100] + "...",
            }
        )]
    
    def _load_with_tesseract(self, source: str) -> List[KnowledgeChunk]:
        """Extract text using Tesseract OCR."""
        try:
            import pytesseract
            from PIL import Image
        except ImportError:
            raise ImportError("OCR requires: pip install pytesseract pillow")
        
        image = Image.open(source)
        text = pytesseract.image_to_string(image)
        
        return [KnowledgeChunk(
            content=text,
            source=source,
            source_type=SourceType.IMAGE,
            metadata={
                "ocr_provider": "tesseract",
                "image_size": image.size,
            }
        )]
    
    def _load_with_openai_vision(
        self,
        source: str,
        prompt: str = "Extract all text and describe the content of this image.",
    ) -> List[KnowledgeChunk]:
        """Extract content using OpenAI Vision."""
        try:
            from openai import OpenAI
            import base64
        except ImportError:
            raise ImportError("OpenAI Vision requires: pip install openai")
        
        with open(source, "rb") as f:
            b64_image = base64.b64encode(f.read()).decode()
        
        client = OpenAI(api_key=self.api_key or os.getenv("OPENAI_API_KEY"))
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/jpeg;base64,{b64_image}"}
                    }
                ]
            }]
        )
        
        content = response.choices[0].message.content
        
        return [KnowledgeChunk(
            content=content,
            source=source,
            source_type=SourceType.IMAGE,
            metadata={
                "ocr_provider": "openai_vision",
                "model": "gpt-4o",
            }
        )]


class WebLoader(SourceLoader):
    """Load content from web pages."""
    
    def __init__(self, chunk_size: int = 1000):
        self.chunk_size = chunk_size
    
    def supports(self, source: str) -> bool:
        return source.startswith(("http://", "https://"))
    
    def load(self, source: str, **kwargs) -> List[KnowledgeChunk]:
        try:
            import requests
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("Web loading requires: pip install requests beautifulsoup4")
        
        response = requests.get(source, timeout=30)
        soup = BeautifulSoup(response.text, "html.parser")
        
        # Remove script and style elements
        for element in soup(["script", "style", "nav", "footer", "header"]):
            element.decompose()
        
        text = soup.get_text(separator="\n")
        # Clean up whitespace
        text = re.sub(r"\n\s*\n", "\n\n", text)
        text = text.strip()
        
        return [KnowledgeChunk(
            content=text,
            source=source,
            source_type=SourceType.WEB_PAGE,
            metadata={
                "title": soup.title.string if soup.title else "",
                "url": source,
            }
        )]


class WebSearchLoader(SourceLoader):
    """Load content from web search results."""
    
    def __init__(
        self,
        search_provider: str = "serper",  # or "google", "bing", "duckduckgo"
        api_key: Optional[str] = None,
        num_results: int = 5,
    ):
        self.search_provider = search_provider
        self.api_key = api_key
        self.num_results = num_results
        self._web_loader = WebLoader()
    
    def supports(self, source: str) -> bool:
        return source.startswith("search:")
    
    def load(self, source: str, **kwargs) -> List[KnowledgeChunk]:
        query = source.replace("search:", "").strip()
        
        if self.search_provider == "serper":
            return self._search_serper(query)
        elif self.search_provider == "duckduckgo":
            return self._search_duckduckgo(query)
        else:
            logger.warning(f"Unknown search provider: {self.search_provider}")
            return []
    
    def _search_serper(self, query: str) -> List[KnowledgeChunk]:
        """Search using Serper API."""
        try:
            import requests
        except ImportError:
            raise ImportError("Web search requires: pip install requests")
        
        api_key = self.api_key or os.getenv("SERPER_API_KEY")
        
        response = requests.post(
            "https://google.serper.dev/search",
            headers={"X-API-KEY": api_key, "Content-Type": "application/json"},
            json={"q": query, "num": self.num_results}
        )
        
        data = response.json()
        chunks = []
        
        for result in data.get("organic", []):
            chunk = KnowledgeChunk(
                content=f"{result.get('title', '')}\n{result.get('snippet', '')}",
                source=result.get("link", ""),
                source_type=SourceType.WEB_SEARCH,
                metadata={
                    "title": result.get("title"),
                    "url": result.get("link"),
                    "query": query,
                }
            )
            chunks.append(chunk)
        
        return chunks
    
    def _search_duckduckgo(self, query: str) -> List[KnowledgeChunk]:
        """Search using DuckDuckGo."""
        try:
            from duckduckgo_search import DDGS
        except ImportError:
            raise ImportError("DuckDuckGo search requires: pip install duckduckgo-search")
        
        chunks = []
        
        with DDGS() as ddgs:
            for result in ddgs.text(query, max_results=self.num_results):
                chunk = KnowledgeChunk(
                    content=f"{result.get('title', '')}\n{result.get('body', '')}",
                    source=result.get("href", ""),
                    source_type=SourceType.WEB_SEARCH,
                    metadata={
                        "title": result.get("title"),
                        "url": result.get("href"),
                        "query": query,
                    }
                )
                chunks.append(chunk)
        
        return chunks


class APILoader(SourceLoader):
    """Load content from REST APIs."""
    
    def supports(self, source: str) -> bool:
        return source.startswith("api:")
    
    def load(
        self,
        source: str,
        method: str = "GET",
        headers: Optional[Dict] = None,
        data: Optional[Dict] = None,
        json_path: Optional[str] = None,
        **kwargs,
    ) -> List[KnowledgeChunk]:
        try:
            import requests
        except ImportError:
            raise ImportError("API loading requires: pip install requests")
        
        url = source.replace("api:", "").strip()
        
        response = requests.request(
            method=method,
            url=url,
            headers=headers or {},
            json=data,
            timeout=30,
        )
        
        try:
            content = response.json()
            
            # Extract specific path if provided
            if json_path:
                for key in json_path.split("."):
                    content = content.get(key, content)
            
            content_str = json.dumps(content, indent=2) if isinstance(content, (dict, list)) else str(content)
            
        except json.JSONDecodeError:
            content_str = response.text
        
        return [KnowledgeChunk(
            content=content_str,
            source=url,
            source_type=SourceType.API,
            metadata={
                "status_code": response.status_code,
                "method": method,
            }
        )]


class JSONLoader(SourceLoader):
    """Load JSON files."""
    
    def supports(self, source: str) -> bool:
        return Path(source).suffix.lower() == ".json"
    
    def load(self, source: str, **kwargs) -> List[KnowledgeChunk]:
        with open(source, "r", encoding="utf-8") as f:
            data = json.load(f)
        
        return [KnowledgeChunk(
            content=json.dumps(data, indent=2),
            source=source,
            source_type=SourceType.JSON,
            metadata={"type": type(data).__name__}
        )]


class CSVLoader(SourceLoader):
    """Load CSV files."""
    
    def supports(self, source: str) -> bool:
        return Path(source).suffix.lower() == ".csv"
    
    def load(self, source: str, **kwargs) -> List[KnowledgeChunk]:
        try:
            import pandas as pd
        except ImportError:
            # Fallback to csv module
            import csv
            chunks = []
            with open(source, "r", encoding="utf-8") as f:
                reader = csv.DictReader(f)
                for i, row in enumerate(reader):
                    chunks.append(KnowledgeChunk(
                        content=json.dumps(row),
                        source=source,
                        source_type=SourceType.CSV,
                        metadata={"row_index": i}
                    ))
            return chunks
        
        df = pd.read_csv(source)
        
        chunks = []
        for i, row in df.iterrows():
            chunks.append(KnowledgeChunk(
                content=row.to_json(),
                source=source,
                source_type=SourceType.CSV,
                metadata={
                    "row_index": i,
                    "columns": list(df.columns),
                }
            ))
        
        return chunks


class KnowledgeBuilder:
    """
    Build knowledge bases from various sources with embedding generation.
    
    Example:
        >>> from agenticaiframework.knowledge import KnowledgeBuilder
        >>> 
        >>> # Initialize with embedding provider
        >>> builder = KnowledgeBuilder(
        ...     embedding_provider="openai",
        ...     embedding_model="text-embedding-3-small"
        ... )
        >>> 
        >>> # Add knowledge from various sources
        >>> builder.add_from_file("docs/manual.pdf")
        >>> builder.add_from_url("https://example.com/article")
        >>> builder.add_from_web_search("machine learning best practices")
        >>> builder.add_from_api("https://api.example.com/data")
        >>> builder.add_from_image("diagram.png")
        >>> 
        >>> # Get embeddings for vector database
        >>> embeddings = builder.get_embeddings()
        >>> 
        >>> # Store in Qdrant
        >>> qdrant_points = [e.to_qdrant_point() for e in embeddings]
        >>> 
        >>> # Or store in Pinecone
        >>> pinecone_vectors = [e.to_pinecone_vector() for e in embeddings]
    """
    
    def __init__(
        self,
        embedding_provider: str = "openai",
        embedding_model: Optional[str] = None,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.chunks: List[KnowledgeChunk] = []
        
        # Initialize embedding provider
        self.embedding = self._create_embedding_provider(
            embedding_provider,
            embedding_model,
        )
        
        # Initialize loaders
        self.loaders: List[SourceLoader] = [
            TextLoader(chunk_size, chunk_overlap),
            MarkdownLoader(chunk_size, chunk_overlap),
            PDFLoader(chunk_size, chunk_overlap),
            DocxLoader(),
            ImageLoader(),
            WebLoader(chunk_size),
            WebSearchLoader(),
            APILoader(),
            JSONLoader(),
            CSVLoader(),
        ]
    
    def _create_embedding_provider(
        self,
        provider: str,
        model: Optional[str],
    ) -> EmbeddingProvider:
        """Create embedding provider instance."""
        if provider == "openai":
            return OpenAIEmbedding(model=model or "text-embedding-3-small")
        elif provider == "azure":
            return AzureOpenAIEmbedding(deployment=model or "text-embedding-ada-002")
        elif provider == "huggingface":
            return HuggingFaceEmbedding(model_name=model or "sentence-transformers/all-MiniLM-L6-v2")
        elif provider == "cohere":
            return CohereEmbedding(model=model or "embed-english-v3.0")
        else:
            raise ValueError(f"Unknown embedding provider: {provider}")
    
    def add_loader(self, loader: SourceLoader) -> None:
        """Add a custom source loader."""
        self.loaders.insert(0, loader)
    
    def _find_loader(self, source: str) -> Optional[SourceLoader]:
        """Find appropriate loader for source."""
        for loader in self.loaders:
            if loader.supports(source):
                return loader
        return None
    
    def add(self, source: str, **kwargs) -> List[KnowledgeChunk]:
        """
        Add knowledge from any supported source.
        
        Args:
            source: File path, URL, or search query
            **kwargs: Additional arguments for the loader
            
        Returns:
            List of created knowledge chunks
        """
        loader = self._find_loader(source)
        if not loader:
            raise ValueError(f"No loader found for source: {source}")
        
        new_chunks = loader.load(source, **kwargs)
        self.chunks.extend(new_chunks)
        logger.info(f"Added {len(new_chunks)} chunks from: {source}")
        return new_chunks
    
    def add_from_file(self, file_path: str, **kwargs) -> List[KnowledgeChunk]:
        """Add knowledge from a file."""
        return self.add(file_path, **kwargs)
    
    def add_from_files(self, file_paths: List[str], **kwargs) -> List[KnowledgeChunk]:
        """Add knowledge from multiple files."""
        all_chunks = []
        for path in file_paths:
            all_chunks.extend(self.add_from_file(path, **kwargs))
        return all_chunks
    
    def add_from_directory(
        self,
        directory: str,
        extensions: Optional[List[str]] = None,
        recursive: bool = True,
    ) -> List[KnowledgeChunk]:
        """Add knowledge from all files in a directory."""
        extensions = extensions or [".txt", ".md", ".pdf", ".docx", ".json"]
        all_chunks = []
        
        path = Path(directory)
        pattern = "**/*" if recursive else "*"
        
        for file_path in path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in extensions:
                try:
                    chunks = self.add_from_file(str(file_path))
                    all_chunks.extend(chunks)
                except Exception as e:
                    logger.warning(f"Failed to load {file_path}: {e}")
        
        return all_chunks
    
    def add_from_url(self, url: str, **kwargs) -> List[KnowledgeChunk]:
        """Add knowledge from a web page."""
        return self.add(url, **kwargs)
    
    def add_from_urls(self, urls: List[str], **kwargs) -> List[KnowledgeChunk]:
        """Add knowledge from multiple URLs."""
        all_chunks = []
        for url in urls:
            try:
                chunks = self.add_from_url(url, **kwargs)
                all_chunks.extend(chunks)
            except Exception as e:
                logger.warning(f"Failed to load URL {url}: {e}")
        return all_chunks
    
    def add_from_web_search(
        self,
        query: str,
        num_results: int = 5,
        fetch_content: bool = False,
    ) -> List[KnowledgeChunk]:
        """
        Add knowledge from web search results.
        
        Args:
            query: Search query
            num_results: Number of results to fetch
            fetch_content: Whether to fetch full page content
        """
        # Update search loader results count
        for loader in self.loaders:
            if isinstance(loader, WebSearchLoader):
                loader.num_results = num_results
                break
        
        chunks = self.add(f"search:{query}")
        
        if fetch_content:
            # Fetch full content from search result URLs
            urls = [c.metadata.get("url") for c in chunks if c.metadata.get("url")]
            for url in urls:
                try:
                    self.add_from_url(url)
                except Exception as e:
                    logger.warning(f"Failed to fetch content from {url}: {e}")
        
        return chunks
    
    def add_from_api(
        self,
        url: str,
        method: str = "GET",
        headers: Optional[Dict] = None,
        data: Optional[Dict] = None,
        json_path: Optional[str] = None,
    ) -> List[KnowledgeChunk]:
        """Add knowledge from an API endpoint."""
        return self.add(
            f"api:{url}",
            method=method,
            headers=headers,
            data=data,
            json_path=json_path,
        )
    
    def add_from_image(
        self,
        image_path: str,
        ocr_provider: str = "pytesseract",
        vision_prompt: Optional[str] = None,
    ) -> List[KnowledgeChunk]:
        """
        Add knowledge from an image (OCR or vision model).
        
        Args:
            image_path: Path to image file
            ocr_provider: 'pytesseract' or 'openai_vision'
            vision_prompt: Custom prompt for vision model
        """
        # Update image loader
        for loader in self.loaders:
            if isinstance(loader, ImageLoader):
                loader.ocr_provider = ocr_provider
                break
        
        return self.add(image_path, prompt=vision_prompt)
    
    def add_text(self, text: str, source: str = "direct", metadata: Optional[Dict] = None) -> KnowledgeChunk:
        """Add text directly as knowledge."""
        chunk = KnowledgeChunk(
            content=text,
            source=source,
            source_type=SourceType.TXT,
            metadata=metadata or {},
        )
        self.chunks.append(chunk)
        return chunk
    
    def generate_embeddings(self, batch_size: int = 100) -> None:
        """Generate embeddings for all chunks."""
        chunks_without_embeddings = [c for c in self.chunks if c.embedding is None]
        
        for i in range(0, len(chunks_without_embeddings), batch_size):
            batch = chunks_without_embeddings[i:i + batch_size]
            texts = [c.content for c in batch]
            
            embeddings = self.embedding.embed_batch(texts)
            
            for chunk, embedding in zip(batch, embeddings):
                chunk.embedding = embedding
        
        logger.info(f"Generated embeddings for {len(chunks_without_embeddings)} chunks")
    
    def get_embeddings(self) -> List[EmbeddingOutput]:
        """
        Get all embeddings ready for vector database storage.
        
        Returns:
            List of EmbeddingOutput objects with vector DB conversion methods
        """
        # Generate embeddings if not already done
        self.generate_embeddings()
        
        outputs = []
        for chunk in self.chunks:
            if chunk.embedding:
                outputs.append(EmbeddingOutput(
                    id=chunk.id,
                    embedding=chunk.embedding,
                    content=chunk.content,
                    metadata={
                        "source": chunk.source,
                        "source_type": chunk.source_type.value,
                        "timestamp": chunk.timestamp,
                        **chunk.metadata,
                    }
                ))
        
        return outputs
    
    def get_chunks(self) -> List[KnowledgeChunk]:
        """Get all knowledge chunks."""
        return self.chunks
    
    def clear(self) -> None:
        """Clear all chunks."""
        self.chunks = []
    
    def export_to_json(self, file_path: str) -> None:
        """Export knowledge base to JSON file."""
        data = [chunk.to_dict() for chunk in self.chunks]
        with open(file_path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2)
    
    def import_from_json(self, file_path: str) -> None:
        """Import knowledge base from JSON file."""
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        
        for item in data:
            self.chunks.append(KnowledgeChunk.from_dict(item))


__all__ = [
    "SourceType",
    "KnowledgeChunk",
    "EmbeddingOutput",
    "EmbeddingProvider",
    "OpenAIEmbedding",
    "AzureOpenAIEmbedding",
    "HuggingFaceEmbedding",
    "CohereEmbedding",
    "SourceLoader",
    "TextLoader",
    "MarkdownLoader",
    "PDFLoader",
    "DocxLoader",
    "ImageLoader",
    "WebLoader",
    "WebSearchLoader",
    "APILoader",
    "JSONLoader",
    "CSVLoader",
    "KnowledgeBuilder",
]
