"""
Document ingestion pipeline with contextual chunking.

Collects documents, chunks them with context, generates embeddings via Modal,
and stores in pgvector.

Usage:
    python ingest.py ~/notes ~/projects/docs
    python ingest.py ~/notes --metadata '{"project": "personal"}'
"""

from __future__ import annotations

import argparse
import fnmatch
import hashlib
import json
import os
import re
import sys
from collections.abc import Generator
from dataclasses import dataclass, field
from datetime import UTC, datetime
from pathlib import Path

import psycopg
import yaml
from pgvector.psycopg import register_vector
from psycopg.rows import dict_row

from .config import config


def read_text_with_fallback(
    path: Path, encodings: tuple[str, ...] = ("utf-8", "windows-1252", "latin-1")
) -> str:
    """Read text file trying multiple encodings in order."""
    for encoding in encodings:
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    # Last resort: read with errors replaced
    return path.read_text(encoding="utf-8", errors="replace")


def matches_pattern(filename: str, patterns: list[str]) -> str | None:
    """Check if filename matches any pattern. Returns matched pattern or None."""
    for pattern in patterns:
        if fnmatch.fnmatch(filename, pattern) or fnmatch.fnmatch(filename.lower(), pattern.lower()):
            return pattern
    return None


# Patterns for detecting secrets in content
SECRET_PATTERNS = [
    (re.compile(r"-----BEGIN [A-Z ]* PRIVATE KEY-----"), "private key"),
    (re.compile(r"AKIA[0-9A-Z]{16}"), "AWS access key"),
    (re.compile(r"ghp_[a-zA-Z0-9]{36}"), "GitHub personal access token"),
    (re.compile(r"gho_[a-zA-Z0-9]{36}"), "GitHub OAuth token"),
    (re.compile(r"sk-[a-zA-Z0-9]{48}"), "OpenAI API key"),
    (re.compile(r"sk-ant-api[a-zA-Z0-9-]{80,}"), "Anthropic API key"),
]


def scan_content_for_secrets(content: str) -> str | None:
    """Scan content for potential secrets. Returns description if found, None otherwise."""
    # Only check first 10KB to avoid slow scans on large files
    sample = content[:10240]
    for pattern, description in SECRET_PATTERNS:
        if pattern.search(sample):
            return description
    return None


def is_minified(content: str, max_line_length: int = 1000) -> bool:
    """Detect if content appears to be minified JS/CSS."""
    lines = content.split("\n", 10)  # Only check first few lines
    if not lines:
        return False
    # Check if any of the first lines is extremely long
    for line in lines[:5]:
        if len(line) > max_line_length:
            # Also check it's not just a long string/comment - minified has lots of punctuation
            if line.count(";") > 20 or line.count(",") > 50 or line.count("{") > 20:
                return True
    return False


class FileSkipReason:
    """Result of file skip check."""

    def __init__(self, should_skip: bool, reason: str = "", is_security: bool = False):
        self.should_skip = should_skip
        self.reason = reason
        self.is_security = is_security  # True for blocked (security), False for skipped (low-value)


def check_file_skip(path: Path, content: str | None = None) -> FileSkipReason:
    """
    Check if a file should be skipped or blocked.

    Returns FileSkipReason with details.
    """
    filename = path.name

    # Check block patterns (security)
    if matched := matches_pattern(filename, config.block_patterns):
        return FileSkipReason(True, f"matches block pattern '{matched}'", is_security=True)

    # Check skip patterns (low-value)
    if matched := matches_pattern(filename, config.skip_patterns):
        return FileSkipReason(True, f"matches skip pattern '{matched}'", is_security=False)

    # Content-based checks (if content provided and scanning enabled)
    if content is not None and config.scan_content:
        # Check for secrets
        if secret_type := scan_content_for_secrets(content):
            return FileSkipReason(True, f"contains {secret_type}", is_security=True)

        # Check for minified JS/CSS
        if path.suffix in (".js", ".css") and is_minified(
            content, config.max_line_length_for_minified
        ):
            return FileSkipReason(True, "appears to be minified", is_security=False)

    return FileSkipReason(False)


@dataclass
class DocumentMetadata:
    """Metadata extracted from document or provided externally."""

    tags: list[str] = field(default_factory=list)
    project: str | None = None
    category: str | None = None
    status: str | None = None
    extra: dict = field(default_factory=dict)

    @classmethod
    def from_frontmatter(cls, frontmatter: dict) -> DocumentMetadata:
        """Create from YAML frontmatter."""
        extra = {
            k: v
            for k, v in frontmatter.items()
            if k not in {"tags", "project", "category", "status"}
        }
        if doc_date := extract_document_date(frontmatter):
            extra["document_date"] = doc_date
        return cls(
            tags=frontmatter.get("tags", []),
            project=frontmatter.get("project"),
            category=frontmatter.get("category"),
            status=frontmatter.get("status"),
            extra=extra,
        )

    def to_dict(self) -> dict:
        """Convert to JSON-serializable dict."""
        result = {}
        if self.tags:
            result["tags"] = self.tags
        if self.project:
            result["project"] = self.project
        if self.category:
            result["category"] = self.category
        if self.status:
            result["status"] = self.status
        if self.extra:
            result.update(self.extra)
        return result


@dataclass
class Document:
    """A document to be indexed."""

    source_path: str
    source_type: str
    title: str
    content: str
    metadata: DocumentMetadata = field(default_factory=DocumentMetadata)
    sections: list[tuple[str, str]] = field(default_factory=list)  # (header, content)

    # Structured fields for actionable items (tasks, events, emails)
    due_date: datetime | None = None  # Task deadlines
    event_start: datetime | None = None  # Calendar event start
    event_end: datetime | None = None  # Calendar event end
    status: str | None = None  # 'pending', 'completed', 'cancelled', etc.
    priority: int | None = None  # 1-5 scale (1=highest)


@dataclass
class Chunk:
    """A chunk ready for embedding."""

    content: str  # Original text (for display)
    embedding_text: str  # Contextualized text (for embedding)
    chunk_index: int
    token_count: int
    metadata: dict = field(default_factory=dict)


def content_hash(content: str) -> str:
    """Generate hash for deduplication/change detection."""
    return hashlib.sha256(content.encode()).hexdigest()[:16]


def extract_document_date(metadata: dict) -> str | None:
    """Extract document date from frontmatter/metadata, trying common field names."""
    date_fields = ["date", "created", "modified", "updated", "last_modified", "pubdate"]
    for field_name in date_fields:
        if value := metadata.get(field_name):
            if hasattr(value, "isoformat"):
                return value.isoformat()
            if isinstance(value, str):
                return value
    return None


def extract_frontmatter(content: str) -> tuple[dict, str]:
    """
    Extract YAML frontmatter from markdown content.

    Returns (frontmatter_dict, remaining_content).
    """
    if not content.startswith("---"):
        return {}, content

    # Find closing ---
    end_match = re.search(r"\n---\s*\n", content[3:])
    if not end_match:
        return {}, content

    frontmatter_text = content[3 : end_match.start() + 3]
    remaining = content[end_match.end() + 3 :]

    try:
        frontmatter = yaml.safe_load(frontmatter_text) or {}
        return frontmatter, remaining
    except yaml.YAMLError:
        return {}, content


def extract_sections_markdown(content: str) -> list[tuple[str, str]]:
    """
    Extract sections from markdown content.

    Returns list of (header, section_content) tuples.
    """
    # Split by headers (any level)
    parts = re.split(r"(^#{1,6}\s+.+$)", content, flags=re.MULTILINE)

    sections = []
    current_header = None

    for part in parts:
        if re.match(r"^#{1,6}\s+", part):
            current_header = part.strip().lstrip("#").strip()
        elif part.strip():
            sections.append((current_header, part.strip()))

    return sections


def extract_org_metadata(content: str) -> tuple[dict, str]:
    """
    Extract org-mode metadata from file header.

    Parses #+KEY: value lines at the start of the file.
    Returns (metadata_dict, remaining_content).
    """
    metadata = {}
    lines = content.split("\n")
    body_start = 0

    for i, line in enumerate(lines):
        match = re.match(r"^#\+(\w+):\s*(.*)$", line, re.IGNORECASE)
        if match:
            key = match.group(1).lower()
            value = match.group(2).strip()
            if key in metadata:
                # Handle multiple values (e.g., multiple #+TAGS lines)
                if isinstance(metadata[key], list):
                    metadata[key].append(value)
                else:
                    metadata[key] = [metadata[key], value]
            else:
                metadata[key] = value
            body_start = i + 1
        elif line.strip() and not line.startswith("#"):
            # Stop at first non-metadata, non-comment line
            break

    remaining = "\n".join(lines[body_start:])
    return metadata, remaining


def extract_org_tags(header: str) -> tuple[str, list[str]]:
    """
    Extract tags from an org header line.

    Org tags appear at end of header like: * Header text  :tag1:tag2:
    Returns (header_without_tags, list_of_tags).
    """
    match = re.search(r"\s+(:[:\w]+:)\s*$", header)
    if match:
        tag_str = match.group(1)
        tags = [t for t in tag_str.split(":") if t]
        header_clean = header[: match.start()].strip()
        return header_clean, tags
    return header, []


def extract_sections_org(content: str) -> list[tuple[str, str]]:
    """
    Extract sections from org-mode content.

    Org headers use * (one or more) at start of line.
    Returns list of (header, section_content) tuples.
    """
    # Split by org headers (any level)
    parts = re.split(r"(^\*+\s+.+$)", content, flags=re.MULTILINE)

    sections = []
    current_header = None

    for part in parts:
        if re.match(r"^\*+\s+", part):
            # Remove leading stars and any TODO keywords
            header = re.sub(r"^\*+\s+", "", part)
            # Remove common TODO keywords
            header = re.sub(r"^(TODO|DONE|WAITING|CANCELLED|NEXT|SOMEDAY)\s+", "", header)
            # Extract and remove tags
            header, _ = extract_org_tags(header)
            current_header = header.strip()
        elif part.strip():
            # Skip property drawers
            clean_part = re.sub(r":PROPERTIES:.*?:END:", "", part, flags=re.DOTALL)
            if clean_part.strip():
                sections.append((current_header, clean_part.strip()))

    return sections


# Org-mode TODO keywords (common defaults)
ORG_TODO_KEYWORDS = {"TODO", "DONE", "WAITING", "CANCELLED", "NEXT", "SOMEDAY"}
ORG_DONE_KEYWORDS = {"DONE", "CANCELLED"}


@dataclass
class OrgTodoItem:
    """Represents a parsed org-mode TODO item."""

    heading: str  # The heading text (without stars, keyword, priority, tags)
    raw_heading: str  # Original heading line for source_path anchor
    level: int  # Number of stars
    keyword: str | None  # TODO, DONE, etc.
    priority: str | None  # A, B, C
    tags: list[str]
    deadline: datetime | None
    scheduled: datetime | None
    closed: datetime | None
    content: str  # Body text under this heading


def parse_org_timestamp(ts: str) -> datetime | None:
    """Parse org-mode timestamp like <2024-01-15 Mon> or [2024-01-15 Mon 10:30]."""
    # Strip brackets
    ts = ts.strip("<>[]")
    # Try various formats
    formats = [
        "%Y-%m-%d %a %H:%M",  # <2024-01-15 Mon 10:30>
        "%Y-%m-%d %a",  # <2024-01-15 Mon>
        "%Y-%m-%d %H:%M",  # <2024-01-15 10:30>
        "%Y-%m-%d",  # <2024-01-15>
    ]
    for fmt in formats:
        try:
            dt = datetime.strptime(ts, fmt)
            return dt.replace(tzinfo=UTC)
        except ValueError:
            continue
    # Try just the date part
    match = re.match(r"(\d{4}-\d{2}-\d{2})", ts)
    if match:
        try:
            return datetime.strptime(match.group(1), "%Y-%m-%d").replace(tzinfo=UTC)
        except ValueError:
            pass
    return None


def extract_org_todo_items(content: str) -> list[OrgTodoItem]:
    """
    Extract TODO items from org-mode content.

    Parses headings with TODO keywords and extracts:
    - Status (TODO/DONE/etc.)
    - Priority ([#A]/[#B]/[#C])
    - Tags (:tag1:tag2:)
    - DEADLINE/SCHEDULED/CLOSED timestamps
    - Body content
    """
    items = []
    lines = content.split("\n")

    i = 0
    while i < len(lines):
        line = lines[i]

        # Match org heading with optional TODO keyword
        # Pattern: *+ [KEYWORD] [#PRIORITY] Title :tags:
        heading_match = re.match(
            r"^(\*+)\s+"  # Stars
            r"(?:(TODO|DONE|WAITING|CANCELLED|NEXT|SOMEDAY)\s+)?"  # Optional keyword
            r"(?:\[#([ABC])\]\s+)?"  # Optional priority
            r"(.+)$",  # Rest of heading
            line,
        )

        if heading_match:
            level = len(heading_match.group(1))
            keyword = heading_match.group(2)
            priority = heading_match.group(3)
            rest = heading_match.group(4)

            # Only process items with TODO keywords
            if keyword:
                # Extract tags from end of heading
                heading_text, tags = extract_org_tags(rest)

                # Collect body content until next heading of same or higher level
                body_lines = []
                deadline = None
                scheduled = None
                closed = None
                i += 1

                while i < len(lines):
                    next_line = lines[i]
                    # Check for next heading of same or higher level
                    next_heading = re.match(r"^(\*+)\s+", next_line)
                    if next_heading and len(next_heading.group(1)) <= level:
                        break

                    # Check for planning line (DEADLINE, SCHEDULED, CLOSED)
                    if re.match(r"^\s*(DEADLINE|SCHEDULED|CLOSED):", next_line):
                        if dl := re.search(r"DEADLINE:\s*(<[^>]+>)", next_line):
                            deadline = parse_org_timestamp(dl.group(1))
                        if sc := re.search(r"SCHEDULED:\s*(<[^>]+>)", next_line):
                            scheduled = parse_org_timestamp(sc.group(1))
                        if cl := re.search(r"CLOSED:\s*(\[[^\]]+\])", next_line):
                            closed = parse_org_timestamp(cl.group(1))
                    # Skip property drawers
                    elif next_line.strip() == ":PROPERTIES:":
                        while i < len(lines) and lines[i].strip() != ":END:":
                            i += 1
                    elif next_line.strip() and not next_line.strip().startswith(":"):
                        body_lines.append(next_line)

                    i += 1

                items.append(
                    OrgTodoItem(
                        heading=heading_text.strip(),
                        raw_heading=line,
                        level=level,
                        keyword=keyword,
                        priority=priority,
                        tags=tags,
                        deadline=deadline,
                        scheduled=scheduled,
                        closed=closed,
                        content="\n".join(body_lines).strip(),
                    )
                )
                continue

        i += 1

    return items


def org_todo_to_document(
    item: OrgTodoItem,
    file_path: Path,
    file_metadata: DocumentMetadata,
) -> Document:
    """Convert an OrgTodoItem to a Document with structured fields."""
    # Build org-mode link-style source path: file.org::*Heading
    # Use the heading text (not raw) for cleaner anchors
    anchor = f"*{item.keyword} {item.heading}" if item.keyword else f"*{item.heading}"
    source_path = f"{file_path.resolve()}::{anchor}"

    # Map org priority to numeric (A=1, B=2, C=3)
    priority_map = {"A": 1, "B": 2, "C": 3}
    priority = priority_map.get(item.priority) if item.priority else None
    # SOMEDAY items get lowest priority
    if item.keyword == "SOMEDAY":
        priority = 5

    # Map org keyword to status
    status = "completed" if item.keyword in ORG_DONE_KEYWORDS else "pending"

    # Use deadline or scheduled as due_date
    due_date = item.deadline or item.scheduled

    # Merge file tags with item tags
    tags = list(file_metadata.tags) + item.tags

    metadata = DocumentMetadata(
        tags=tags,
        project=file_metadata.project,
        category=file_metadata.category,
    )

    # Build content with context
    content_parts = [item.heading]
    if item.content:
        content_parts.append(item.content)
    content = "\n\n".join(content_parts)

    return Document(
        source_path=source_path,
        source_type="org-todo",
        title=item.heading,
        content=content,
        metadata=metadata,
        due_date=due_date,
        status=status,
        priority=priority,
    )


def extract_code_context(content: str, file_ext: str) -> dict:
    """
    Extract structural context from code files.

    Returns dict with classes, functions, imports found.
    """
    context = {
        "classes": [],
        "functions": [],
        "imports": [],
    }

    if file_ext == ".py":
        # Python classes and functions
        context["classes"] = re.findall(r"^class\s+(\w+)", content, re.MULTILINE)
        context["functions"] = re.findall(r"^def\s+(\w+)", content, re.MULTILINE)
        # Top-level imports
        imports = re.findall(r"^(?:from\s+(\S+)|import\s+(\S+))", content, re.MULTILINE)
        context["imports"] = [i[0] or i[1] for i in imports][:10]  # Limit

    elif file_ext in {".js", ".ts", ".jsx", ".tsx"}:
        # JavaScript/TypeScript
        context["classes"] = re.findall(r"class\s+(\w+)", content)
        context["functions"] = re.findall(
            r"(?:function\s+(\w+)|(?:const|let|var)\s+(\w+)\s*=\s*(?:async\s*)?\()",
            content,
        )
        context["functions"] = [f[0] or f[1] for f in context["functions"]]
        context["imports"] = re.findall(r"from\s+['\"]([^'\"]+)['\"]", content)[:10]

    return {k: v for k, v in context.items() if v}


def infer_project_from_path(path: Path) -> str | None:
    """
    Infer project name from file path.

    Looks for common patterns like:
    - ~/projects/{project}/...
    - ~/code/{project}/...
    - ~/notes/projects/{project}/...
    """
    parts = path.parts
    project_indicators = {"projects", "code", "repos", "src"}

    for i, part in enumerate(parts):
        if part.lower() in project_indicators and i + 1 < len(parts):
            return parts[i + 1]

    return None


def build_embedding_context(
    chunk_text: str,
    doc_title: str,
    source_path: str,
    source_type: str,
    section_header: str | None = None,
    metadata: DocumentMetadata | None = None,
    code_context: dict | None = None,
) -> str:
    """
    Build contextualized text for embedding.

    This is what the embedding model sees. The original chunk_text
    is stored separately for display.
    """
    parts = []

    # Document identity
    parts.append(f"Document: {doc_title}")

    # Source type context
    if source_type == "code":
        path = Path(source_path)
        parts.append(f"File: {path.name}")
        if code_context:
            if classes := code_context.get("classes"):
                parts.append(f"Classes: {', '.join(classes[:5])}")
            if functions := code_context.get("functions"):
                parts.append(f"Functions: {', '.join(functions[:5])}")

    # Project from metadata or path
    project = None
    if metadata and metadata.project:
        project = metadata.project
    else:
        project = infer_project_from_path(Path(source_path))

    if project:
        parts.append(f"Project: {project}")

    # Section context for long documents
    if section_header:
        parts.append(f"Section: {section_header}")

    # Tags/topics from metadata
    if metadata and metadata.tags:
        parts.append(f"Topics: {', '.join(metadata.tags[:5])}")

    if metadata and metadata.category:
        parts.append(f"Category: {metadata.category}")

    # The actual content
    parts.append(f"Content: {chunk_text}")

    return "\n".join(parts)


def chunk_text(
    text: str,
    chunk_size: int = config.chunk_size,
    chunk_overlap: int = config.chunk_overlap,
) -> Generator[tuple[int, str], None, None]:
    """
    Split text into overlapping chunks.

    Tries to break at paragraph/sentence boundaries.
    Uses approximate token count (4 chars ≈ 1 token).
    """
    char_size = chunk_size * config.chars_per_token
    char_overlap = chunk_overlap * config.chars_per_token

    if len(text) <= char_size:
        yield 0, text
        return

    # Split into paragraphs
    paragraphs = re.split(r"\n\n+", text)

    current_chunk = ""
    chunk_index = 0

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        if len(current_chunk) + len(para) + 2 <= char_size:
            current_chunk += para + "\n\n"
        else:
            if current_chunk.strip():
                yield chunk_index, current_chunk.strip()
                chunk_index += 1
                # Keep overlap
                overlap = current_chunk[-char_overlap:] if len(current_chunk) > char_overlap else ""
                current_chunk = overlap + para + "\n\n"
            else:
                # Single paragraph too large - split by sentences
                sentences = re.split(r"(?<=[.!?])\s+", para)
                for sentence in sentences:
                    if len(current_chunk) + len(sentence) + 1 <= char_size:
                        current_chunk += sentence + " "
                    else:
                        if current_chunk.strip():
                            yield chunk_index, current_chunk.strip()
                            chunk_index += 1
                            overlap = (
                                current_chunk[-char_overlap:]
                                if len(current_chunk) > char_overlap
                                else ""
                            )
                            current_chunk = overlap + sentence + " "
                        else:
                            # Single sentence too large - hard split
                            yield chunk_index, sentence[:char_size]
                            chunk_index += 1
                            current_chunk = (
                                sentence[-char_overlap:] if len(sentence) > char_overlap else ""
                            )

    if current_chunk.strip():
        yield chunk_index, current_chunk.strip()


def parse_markdown(path: Path, extra_metadata: dict | None = None) -> Document:
    """Parse a markdown (.md) file into a Document."""
    content = read_text_with_fallback(path)

    # Extract frontmatter
    frontmatter, body = extract_frontmatter(content)
    metadata = DocumentMetadata.from_frontmatter(frontmatter)

    # Merge extra metadata
    if extra_metadata:
        if "tags" in extra_metadata:
            metadata.tags.extend(extra_metadata["tags"])
        if "project" in extra_metadata:
            metadata.project = extra_metadata["project"]
        if "category" in extra_metadata:
            metadata.category = extra_metadata["category"]

    # Extract title
    title_match = re.search(r"^#\s+(.+)$", body, re.MULTILINE)
    title = title_match.group(1) if title_match else path.stem

    # Extract sections
    sections = extract_sections_markdown(body)

    return Document(
        source_path=str(path.resolve()),
        source_type="markdown",
        title=title,
        content=content,
        metadata=metadata,
        sections=sections,
    )


def parse_org(path: Path, extra_metadata: dict | None = None) -> Document:
    """Parse an org-mode (.org) file into a Document (file only, no TODO extraction)."""
    content = read_text_with_fallback(path)

    # Extract org metadata (#+KEY: value lines)
    org_meta, body = extract_org_metadata(content)

    # Build DocumentMetadata from org metadata
    tags = []
    # #+FILETAGS: :tag1:tag2:
    if filetags := org_meta.get("filetags"):
        tags.extend([t for t in filetags.split(":") if t])
    # #+TAGS: tag1 tag2
    if tag_str := org_meta.get("tags"):
        if isinstance(tag_str, list):
            for t in tag_str:
                tags.extend(t.split())
        else:
            tags.extend(tag_str.split())

    metadata = DocumentMetadata(
        tags=tags,
        project=org_meta.get("project"),
        category=org_meta.get("category"),
    )

    # Merge extra metadata
    if extra_metadata:
        if "tags" in extra_metadata:
            metadata.tags.extend(extra_metadata["tags"])
        if "project" in extra_metadata:
            metadata.project = extra_metadata["project"]
        if "category" in extra_metadata:
            metadata.category = extra_metadata["category"]

    # Extract title from #+TITLE or first header
    title = org_meta.get("title")
    if not title:
        title_match = re.search(r"^\*+\s+(.+)$", body, re.MULTILINE)
        if title_match:
            title, _ = extract_org_tags(title_match.group(1))
            # Remove TODO keywords from title
            title = re.sub(r"^(TODO|DONE|WAITING|CANCELLED|NEXT|SOMEDAY)\s+", "", title)
        else:
            title = path.stem

    # Extract sections
    sections = extract_sections_org(body)

    return Document(
        source_path=str(path.resolve()),
        source_type="org",
        title=title,
        content=content,
        metadata=metadata,
        sections=sections,
    )


def parse_org_with_todos(path: Path, extra_metadata: dict | None = None) -> list[Document]:
    """
    Parse an org-mode file into multiple Documents.

    Returns:
        - The file itself as one Document (source_type='org')
        - Each TODO item as a separate Document (source_type='org-todo')
    """
    # Parse the file document
    file_doc = parse_org(path, extra_metadata)

    # Extract TODO items
    content = read_text_with_fallback(path)
    todo_items = extract_org_todo_items(content)

    # Convert TODO items to Documents
    todo_docs = [org_todo_to_document(item, path, file_doc.metadata) for item in todo_items]

    # File document first, then TODO documents
    return [file_doc] + todo_docs


def parse_text(path: Path, extra_metadata: dict | None = None) -> Document:
    """Parse a plain text file into a Document (no special parsing)."""
    content = read_text_with_fallback(path)

    metadata = DocumentMetadata()
    if extra_metadata:
        metadata = DocumentMetadata(
            tags=extra_metadata.get("tags", []),
            project=extra_metadata.get("project"),
            category=extra_metadata.get("category"),
        )

    return Document(
        source_path=str(path.resolve()),
        source_type="text",
        title=path.stem,
        content=content,
        metadata=metadata,
        sections=[],  # No section parsing for raw text
    )


def parse_code(path: Path, extra_metadata: dict | None = None) -> Document:
    """Parse a code file into a Document."""
    content = read_text_with_fallback(path)

    metadata = DocumentMetadata()
    if extra_metadata:
        metadata = DocumentMetadata(
            tags=extra_metadata.get("tags", []),
            project=extra_metadata.get("project"),
            category=extra_metadata.get("category"),
        )

    # Auto-tag by language
    lang_tags = {
        ".py": "python",
        ".js": "javascript",
        ".ts": "typescript",
        ".sql": "sql",
        ".sh": "bash",
        ".yaml": "yaml",
        ".yml": "yaml",
    }
    if lang := lang_tags.get(path.suffix):
        if lang not in metadata.tags:
            metadata.tags.append(lang)

    return Document(
        source_path=str(path.resolve()),
        source_type="code",
        title=path.name,
        content=content,
        metadata=metadata,
    )


def is_url(s: str) -> bool:
    """Check if a string looks like a URL."""
    return s.startswith(("http://", "https://"))


def parse_url(url: str, extra_metadata: dict | None = None) -> Document:
    """Fetch and parse content from a URL using trafilatura."""
    try:
        import trafilatura
    except ImportError:
        raise ImportError(
            "trafilatura is required for URL ingestion. Install with: pip install local-kb[web]"
        )

    # Fetch and extract content
    downloaded = trafilatura.fetch_url(url)
    if downloaded is None:
        raise ValueError(f"Failed to fetch URL: {url}")

    # Extract text content and metadata
    result = trafilatura.extract(
        downloaded,
        include_comments=False,
        include_tables=True,
        output_format="txt",
    )
    if result is None:
        raise ValueError(f"Failed to extract content from URL: {url}")

    # Get metadata separately
    meta = trafilatura.extract_metadata(downloaded)

    # Build document metadata
    metadata = DocumentMetadata()
    if extra_metadata:
        metadata = DocumentMetadata(
            tags=extra_metadata.get("tags", []),
            project=extra_metadata.get("project"),
            category=extra_metadata.get("category"),
        )

    # Add URL-specific metadata
    if meta:
        if meta.title:
            metadata.extra["original_title"] = meta.title
        if meta.author:
            metadata.extra["author"] = meta.author
        if meta.date:
            metadata.extra["document_date"] = meta.date
        if meta.sitename:
            metadata.extra["site"] = meta.sitename
        if meta.description:
            metadata.extra["description"] = meta.description

    # Use fetched timestamp
    metadata.extra["fetched_at"] = datetime.now(UTC).isoformat()

    # Determine title
    title = meta.title if meta and meta.title else url

    return Document(
        source_path=url,
        source_type="web",
        title=title,
        content=result,
        metadata=metadata,
        sections=extract_sections_markdown(result),  # trafilatura output has markdown-like headers
    )


def parse_pdf_date(pdf_date: str | None) -> str | None:
    """Parse PDF date format (D:YYYYMMDDHHmmSS+TZ) to ISO format."""
    if not pdf_date:
        return None
    # Strip optional 'D:' prefix
    if pdf_date.startswith("D:"):
        pdf_date = pdf_date[2:]
    try:
        # Basic format: YYYYMMDDHHMMSS
        if len(pdf_date) >= 14:
            dt = datetime.strptime(pdf_date[:14], "%Y%m%d%H%M%S")
            return dt.isoformat()
        elif len(pdf_date) >= 8:
            dt = datetime.strptime(pdf_date[:8], "%Y%m%d")
            return dt.isoformat()
    except ValueError:
        pass
    return None


def parse_pdf(path: Path, extra_metadata: dict | None = None) -> Document:
    """Parse a PDF file into a Document using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError(
            "pymupdf is required for PDF ingestion. Install with: pip install local-kb[pdf]"
        )

    doc = fitz.open(path)
    metadata = DocumentMetadata()

    # Extract PDF metadata
    pdf_meta = doc.metadata
    if pdf_meta:
        if pdf_meta.get("title"):
            metadata.extra["original_title"] = pdf_meta["title"]
        if pdf_meta.get("author"):
            metadata.extra["author"] = pdf_meta["author"]
        if pdf_meta.get("subject"):
            metadata.extra["subject"] = pdf_meta["subject"]
        if pdf_meta.get("keywords"):
            # Keywords often comma-separated
            keywords = [k.strip() for k in pdf_meta["keywords"].split(",") if k.strip()]
            metadata.tags.extend(keywords)
        # Parse creation date
        if doc_date := parse_pdf_date(pdf_meta.get("creationDate")):
            metadata.extra["document_date"] = doc_date

    # Merge extra metadata
    if extra_metadata:
        if "tags" in extra_metadata:
            metadata.tags.extend(extra_metadata["tags"])
        if "project" in extra_metadata:
            metadata.project = extra_metadata["project"]
        if "category" in extra_metadata:
            metadata.category = extra_metadata["category"]

    # Extract text page by page as sections
    sections = []
    full_text_parts = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text().strip()
        if text:
            sections.append((f"Page {page_num + 1}", text))
            full_text_parts.append(text)

    doc.close()

    # Skip if no text extracted (likely scanned image)
    if not full_text_parts:
        raise ValueError(f"No text extracted from {path.name} - may be a scanned image (needs OCR)")

    # Determine title
    title = pdf_meta.get("title") if pdf_meta else None
    if not title:
        title = path.stem

    return Document(
        source_path=str(path.resolve()),
        source_type="pdf",
        title=title,
        content="\n\n".join(full_text_parts),
        metadata=metadata,
        sections=sections,
    )


def parse_docx(path: Path, extra_metadata: dict | None = None) -> Document:
    """Parse a DOCX file into a Document using python-docx."""
    try:
        import docx
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX ingestion. Install with: pip install local-kb[docx]"
        )

    doc = docx.Document(path)
    metadata = DocumentMetadata()

    # Extract core properties
    core = doc.core_properties
    if core:
        if core.title:
            metadata.extra["original_title"] = core.title
        if core.author:
            metadata.extra["author"] = core.author
        if core.keywords:
            # Keywords often comma or semicolon separated
            for sep in [",", ";"]:
                if sep in core.keywords:
                    keywords = [k.strip() for k in core.keywords.split(sep) if k.strip()]
                    metadata.tags.extend(keywords)
                    break
            else:
                # Single keyword or space-separated
                metadata.tags.extend(core.keywords.split())
        # Get document date (prefer created, fall back to modified)
        if core.created:
            metadata.extra["document_date"] = core.created.isoformat()
        elif core.modified:
            metadata.extra["document_date"] = core.modified.isoformat()

    # Merge extra metadata
    if extra_metadata:
        if "tags" in extra_metadata:
            metadata.tags.extend(extra_metadata["tags"])
        if "project" in extra_metadata:
            metadata.project = extra_metadata["project"]
        if "category" in extra_metadata:
            metadata.category = extra_metadata["category"]

    # Extract paragraphs with heading detection
    sections = []
    current_heading = None
    current_content = []
    full_text_parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        full_text_parts.append(text)

        # Check if paragraph is a heading
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            # Save previous section
            if current_content:
                sections.append((current_heading, "\n\n".join(current_content)))
                current_content = []
            current_heading = text
        else:
            current_content.append(text)

    # Don't forget the last section
    if current_content:
        sections.append((current_heading, "\n\n".join(current_content)))

    # Determine title
    title = core.title if core and core.title else path.stem

    return Document(
        source_path=str(path.resolve()),
        source_type="docx",
        title=title,
        content="\n\n".join(full_text_parts),
        metadata=metadata,
        sections=sections,
    )


def is_text_file(path: Path) -> bool:
    """Check if a file appears to be text (not binary)."""
    try:
        with open(path, "rb") as f:
            chunk = f.read(8192)
        # Check for null bytes (binary indicator)
        if b"\x00" in chunk:
            return False
        # Try to decode as UTF-8
        try:
            chunk.decode("utf-8")
            return True
        except UnicodeDecodeError:
            # Try other common encodings
            for encoding in ("windows-1252", "latin-1"):
                try:
                    chunk.decode(encoding)
                    return True
                except UnicodeDecodeError:
                    continue
            return False
    except OSError:
        return False


def parse_document(
    path: Path, extra_metadata: dict | None = None, force: bool = False
) -> list[Document]:
    """Parse a file into one or more Documents.

    Some file types (e.g., org-mode) produce multiple documents:
    - The file itself (for semantic search)
    - Individual actionable items like TODOs (for structured queries)

    Checks plugin registry first, then falls back to built-in parsers.
    If force=True, parse unknown extensions as text/code (for explicitly provided files).
    """
    # Check plugin registry first
    from .plugins.registry import PluginRegistry

    if parser := PluginRegistry.get_parser_for_file(path):
        return [parser.parse(path, extra_metadata)]

    # Fall back to built-in parsers
    if path.suffix == ".md":
        return [parse_markdown(path, extra_metadata)]
    elif path.suffix == ".org":
        # Org files produce multiple documents: file + TODO items
        return parse_org_with_todos(path, extra_metadata)
    elif path.suffix == ".pdf":
        return [parse_pdf(path, extra_metadata)]
    elif path.suffix == ".docx":
        return [parse_docx(path, extra_metadata)]
    elif path.suffix in config.document_extensions:
        return [parse_text(path, extra_metadata)]
    elif path.suffix in config.code_extensions:
        return [parse_code(path, extra_metadata)]
    elif force:
        return [parse_code(path, extra_metadata)]
    else:
        raise ValueError(f"Unsupported file type: {path.suffix}")


def collect_documents(
    root: Path,
    extra_metadata: dict | None = None,
) -> Generator[Document, None, None]:
    """Recursively collect documents from a directory, pruning ignored directories."""
    print(f"Scanning {root}...", file=sys.stderr, flush=True)
    scanned = 0
    collected = 0
    skipped_ext = 0

    for dirpath, dirnames, filenames in os.walk(root, topdown=True):
        # Prune ignored directories in-place (modifying dirnames affects traversal)
        dirnames[:] = [
            d for d in dirnames if not d.startswith(".") and d not in config.skip_directories
        ]

        for filename in filenames:
            path = Path(dirpath) / filename

            scanned += 1
            if scanned % 500 == 0:
                print(
                    f"  {scanned} files scanned, {collected} documents found...",
                    file=sys.stderr,
                    flush=True,
                )

            if path.suffix not in config.all_extensions:
                skipped_ext += 1
                continue

            # Check filename-based skip/block patterns first (before reading content)
            skip_check = check_file_skip(path)
            if skip_check.should_skip:
                prefix = "BLOCKED" if skip_check.is_security else "Skipping"
                print(f"{prefix}: {path} ({skip_check.reason})", file=sys.stderr)
                continue

            try:
                docs = parse_document(path, extra_metadata)
                if not docs:
                    continue

                # Content-based checks on the primary (file) document
                primary_doc = docs[0]
                if config.scan_content:
                    skip_check = check_file_skip(path, primary_doc.content)
                    if skip_check.should_skip:
                        prefix = "BLOCKED" if skip_check.is_security else "Skipping"
                        print(f"{prefix}: {path} ({skip_check.reason})", file=sys.stderr)
                        continue

                # Capture file mtime for staleness tracking
                mtime = datetime.fromtimestamp(path.stat().st_mtime, tz=UTC)
                mtime_iso = mtime.isoformat()

                # Yield all documents from this file
                for doc in docs:
                    doc.metadata.extra["file_modified_at"] = mtime_iso
                    collected += 1
                    yield doc
            except Exception as e:
                print(f"Error parsing {path}: {e}", file=sys.stderr)

    if scanned >= 1000:
        print(
            f"Scan complete: {scanned} files, {skipped_ext} wrong extension",
            file=sys.stderr,
            flush=True,
        )


def create_chunks(doc: Document) -> list[Chunk]:
    """
    Create contextual chunks from a document.

    Each chunk includes:
    - content: original text (for display)
    - embedding_text: contextualized text (for embedding)
    """
    chunks = []

    # For code files, extract structural context once
    code_context = None
    if doc.source_type == "code":
        ext = Path(doc.source_path).suffix
        code_context = extract_code_context(doc.content, ext)

    # Chunk by sections if available, otherwise whole document
    if doc.sections:
        chunk_index = 0
        for section_header, section_content in doc.sections:
            for _, section_chunk in chunk_text_generator(section_content):
                if not section_chunk.strip():
                    continue  # Skip empty chunks
                embedding_text = build_embedding_context(
                    chunk_text=section_chunk,
                    doc_title=doc.title,
                    source_path=doc.source_path,
                    source_type=doc.source_type,
                    section_header=section_header,
                    metadata=doc.metadata,
                    code_context=code_context,
                )

                chunks.append(
                    Chunk(
                        content=section_chunk,
                        embedding_text=embedding_text,
                        chunk_index=chunk_index,
                        token_count=len(section_chunk) // config.chars_per_token,
                        metadata={"section": section_header} if section_header else {},
                    )
                )
                chunk_index += 1
    else:
        for chunk_index, chunk_content in chunk_text(doc.content):
            if not chunk_content.strip():
                continue  # Skip empty chunks
            embedding_text = build_embedding_context(
                chunk_text=chunk_content,
                doc_title=doc.title,
                source_path=doc.source_path,
                source_type=doc.source_type,
                metadata=doc.metadata,
                code_context=code_context,
            )

            chunks.append(
                Chunk(
                    content=chunk_content,
                    embedding_text=embedding_text,
                    chunk_index=chunk_index,
                    token_count=len(chunk_content) // config.chars_per_token,
                )
            )

    return chunks


# Alias for the generator to avoid name collision
chunk_text_generator = chunk_text


class Ingester:
    """Handles document ingestion into pgvector."""

    def __init__(self, db_url: str, use_modal: bool = True):
        self.db_url = db_url
        self.use_modal = use_modal
        self._embedder = None

    @property
    def embedder(self):
        """Lazy-load embedder, falling back to local if Modal unavailable."""
        if self._embedder is None:
            if self.use_modal:
                try:
                    import modal

                    self._embedder = modal.Cls.from_name("knowledge-embedder", "Embedder")()
                except Exception as e:
                    print(f"Modal unavailable ({e}), using local CPU embedding", file=sys.stderr)
                    self.use_modal = False  # Update flag for embed_batch call path

            if not self.use_modal:
                from .local_embedder import embed_document

                class LocalEmbedder:
                    def embed_batch(self, texts):
                        return [embed_document(t) for t in texts]

                self._embedder = LocalEmbedder()
        return self._embedder

    def ingest_documents(self, documents: list[Document], batch_size: int = 50):
        """
        Ingest documents into the database.

        1. Check for existing documents (by hash)
        2. Create contextual chunks
        3. Generate embeddings via Modal (or local)
        4. Store in pgvector

        For files that produce multiple documents (e.g., org with TODOs):
        - Primary document: file.org (source_type='org')
        - Derived documents: file.org::*TODO ... (source_type='org-todo')
        When a primary document changes, all derived documents are deleted first.
        """
        with psycopg.connect(self.db_url, row_factory=dict_row) as conn:
            register_vector(conn)

            # Track which primary files we've already cleaned up derived docs for
            cleaned_derived = set()

            for doc in documents:
                doc_hash = content_hash(doc.content)

                # Determine if this is a derived document (has :: in path)
                is_derived = "::" in doc.source_path
                if is_derived:
                    base_path = doc.source_path.split("::")[0]
                else:
                    base_path = doc.source_path

                # Check if document exists and unchanged (FOR UPDATE to prevent race)
                existing = conn.execute(
                    "SELECT id FROM documents WHERE content_hash = %s FOR UPDATE",
                    (doc_hash,),
                ).fetchone()

                if existing:
                    # Content unchanged - but update file_modified_at if present
                    new_mtime = doc.metadata.extra.get("file_modified_at")
                    if new_mtime:
                        conn.execute(
                            """UPDATE documents
                               SET metadata = jsonb_set(metadata, '{file_modified_at}', to_jsonb(%s::text))
                               WHERE id = %s""",
                            (new_mtime, existing["id"]),
                        )
                        conn.commit()
                    print(f"Skipping (unchanged): {doc.source_path}")
                    continue

                # For primary documents: also delete any derived documents
                if not is_derived and base_path not in cleaned_derived:
                    deleted = conn.execute(
                        "DELETE FROM documents WHERE source_path LIKE %s RETURNING id",
                        (base_path + "::%",),
                    ).fetchall()
                    if deleted:
                        print(f"  Deleted {len(deleted)} derived documents from {base_path}")
                    cleaned_derived.add(base_path)

                # Check if same path exists with different hash (FOR UPDATE to prevent race)
                old_doc = conn.execute(
                    "SELECT id FROM documents WHERE source_path = %s FOR UPDATE",
                    (doc.source_path,),
                ).fetchone()

                if old_doc:
                    print(f"Updating: {doc.source_path}")
                    conn.execute("DELETE FROM documents WHERE id = %s", (old_doc["id"],))
                else:
                    print(f"Ingesting: {doc.source_path}")

                # Insert document (ON CONFLICT handles duplicate content from different paths)
                result = conn.execute(
                    """
                    INSERT INTO documents (
                        source_path, source_type, title, content, metadata, content_hash,
                        due_date, event_start, event_end, status, priority
                    )
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                    ON CONFLICT (content_hash) DO NOTHING
                    RETURNING id
                    """,
                    (
                        doc.source_path,
                        doc.source_type,
                        doc.title,
                        doc.content,
                        psycopg.types.json.Json(doc.metadata.to_dict()),
                        doc_hash,
                        doc.due_date,
                        doc.event_start,
                        doc.event_end,
                        doc.status,
                        doc.priority,
                    ),
                ).fetchone()

                if result is None:
                    print(f"  Skipping (duplicate content): {doc.source_path}")
                    continue

                doc_id = result["id"]

                # Create chunks
                chunks = create_chunks(doc)

                if not chunks:
                    conn.commit()
                    continue

                # Generate embeddings (batch to avoid OOM on GPU)
                embedding_texts = [c.embedding_text for c in chunks]
                embed_batch_size = 100  # Max texts per GPU call

                print(f"  Generating embeddings for {len(chunks)} chunks...")
                if self.use_modal:
                    embeddings = []
                    for i in range(0, len(embedding_texts), embed_batch_size):
                        batch = embedding_texts[i : i + embed_batch_size]
                        embeddings.extend(self.embedder.embed_batch.remote(batch))
                else:
                    embeddings = self.embedder.embed_batch(embedding_texts)

                # Insert chunks with embeddings
                for chunk, embedding in zip(chunks, embeddings):
                    conn.execute(
                        """
                        INSERT INTO chunks 
                            (document_id, chunk_index, content, embedding_text, embedding, token_count, metadata)
                        VALUES (%s, %s, %s, %s, %s, %s, %s)
                        """,
                        (
                            doc_id,
                            chunk.chunk_index,
                            chunk.content,
                            chunk.embedding_text,
                            embedding,
                            chunk.token_count,
                            psycopg.types.json.Json(chunk.metadata),
                        ),
                    )

                conn.commit()
                print(f"  → {len(chunks)} chunks indexed")

    def delete_document(self, source_path: str):
        """Remove a document and its chunks."""
        with psycopg.connect(self.db_url) as conn:
            result = conn.execute(
                "DELETE FROM documents WHERE source_path = %s RETURNING id",
                (source_path,),
            ).fetchone()
            conn.commit()
            return result is not None


def main():
    """CLI entry point."""
    parser = argparse.ArgumentParser(
        description="Ingest documents into knowledge base",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
    python ingest.py ~/notes
    python ingest.py ~/projects/myapp --metadata '{"project": "myapp"}'
    python ingest.py document.md --local  # Use CPU embedding
        """,
    )
    parser.add_argument("paths", nargs="+", type=Path, help="Files or directories to ingest")
    parser.add_argument(
        "--metadata",
        type=json.loads,
        default={},
        help='JSON metadata to attach (e.g., \'{"project": "myapp"}\')',
    )
    parser.add_argument(
        "--db-url",
        default=config.db_url,
        help="Database URL",
    )
    parser.add_argument(
        "--local",
        action="store_true",
        help="Use local CPU embedding instead of Modal",
    )

    args = parser.parse_args()

    ingester = Ingester(args.db_url, use_modal=not args.local)

    # Collect documents
    documents = []
    for path in args.paths:
        path = path.resolve()
        if path.is_dir():
            documents.extend(collect_documents(path, args.metadata))
        elif path.is_file():
            # Check security patterns first
            skip_check = check_file_skip(path)
            if skip_check.should_skip:
                prefix = "BLOCKED" if skip_check.is_security else "Skipping"
                print(f"{prefix}: {path} ({skip_check.reason})", file=sys.stderr)
                continue

            # For explicitly provided files, try to parse even with unknown extension
            # Always allow .pdf and .docx even if not in config (user may have old config)
            if path.suffix in config.all_extensions or path.suffix in (".pdf", ".docx"):
                documents.extend(parse_document(path, args.metadata))
            elif is_text_file(path):
                print(f"Parsing as text: {path}", file=sys.stderr)
                documents.extend(parse_document(path, args.metadata, force=True))
            else:
                print(f"Skipping binary file: {path}", file=sys.stderr)

    if not documents:
        print("No documents found to ingest")
        return

    print(f"Found {len(documents)} documents to process")
    ingester.ingest_documents(documents)
    print("Done!")


if __name__ == "__main__":
    main()
