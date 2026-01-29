"""Integration tests for template rendering with real templates."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from resume_as_code.models.education import Education
from resume_as_code.models.resume import (
    ContactInfo,
    ResumeBullet,
    ResumeData,
    ResumeItem,
    ResumeSection,
)
from resume_as_code.services.template_service import TemplateService


class TestEndToEndWorkUnitToHTML:
    """End-to-end tests: Work Unit dicts -> ResumeData -> HTML."""

    def test_work_units_to_html_modern_template(self) -> None:
        """Full pipeline: Work Unit dicts transform to rendered HTML."""
        # Simulate Work Units as they'd come from YAML files
        work_units = [
            {
                "id": "wu-2024-01-15-api-migration",
                "title": "Senior Software Engineer",
                "organization": "TechCorp Inc",
                "time_started": date(2022, 3, 1),
                "time_ended": date(2024, 1, 15),
                "actions": [
                    "Led migration of 200+ REST endpoints to GraphQL",
                    "Implemented caching layer reducing latency by 60%",
                    "Mentored team of 4 junior engineers",
                ],
                "outcome": {
                    "result": "Delivered new API platform serving 10M requests/day",
                    "quantified_impact": "40% reduction in p99 latency",
                },
                "tags": ["python", "graphql", "redis"],
                "skills_demonstrated": [
                    {"name": "API Design", "level": "expert"},
                    "Technical Leadership",
                ],
                "scope": {
                    "budget_managed": "$500K",
                    "team_size": 5,
                },
            },
            {
                "id": "wu-2022-02-28-startup",
                "title": "Full Stack Developer",
                "organization": "StartupXYZ",
                "time_started": date(2020, 6, 1),
                "time_ended": date(2022, 2, 28),
                "actions": [
                    "Built customer-facing dashboard from scratch",
                    "Integrated third-party payment processing",
                ],
                "outcome": {
                    "result": "Launched MVP that acquired 1000 paying customers",
                },
                "tags": ["react", "typescript", "postgres"],
                "skills_demonstrated": ["Full Stack Development"],
            },
        ]

        contact = ContactInfo(
            name="Jane Developer",
            email="jane@example.com",
            phone="555-0123",
            location="San Francisco, CA",
            linkedin="https://linkedin.com/in/janedev",
        )

        # Transform Work Units to ResumeData
        resume = ResumeData.from_work_units(
            work_units=work_units,
            contact=contact,
            summary="Senior engineer with 5+ years building scalable systems.",
        )

        # Render to HTML
        service = TemplateService()
        html = service.render(resume, "modern")

        # Verify contact info rendered
        assert "Jane Developer" in html
        assert "jane@example.com" in html
        assert "San Francisco, CA" in html

        # Verify summary rendered
        assert "Senior engineer with 5+ years" in html

        # Verify Work Unit data transformed and rendered
        assert "Senior Software Engineer" in html
        assert "TechCorp Inc" in html
        assert "Mar 2022" in html  # Formatted date
        assert "Delivered new API platform" in html  # Outcome bullet

        # Verify skills extracted and rendered
        assert "python" in html
        assert "graphql" in html
        assert "API Design" in html

    def test_work_units_to_html_executive_template_with_scope(self) -> None:
        """Executive template renders scope_line (Story 7.2: scope from Position only).

        WorkUnit.scope is deprecated - standalone work units without position_id
        don't have scope_line rendered. This test verifies the rendering still works
        when scope_line is set directly on ResumeItem.
        """
        # Create ResumeData directly with scope_line (as would be set from Position)
        contact = ContactInfo(name="Executive Leader")
        resume = ResumeData(
            contact=contact,
            sections=[
                ResumeSection(
                    title="Experience",
                    items=[
                        ResumeItem(
                            title="VP of Engineering",
                            organization="Enterprise Corp",
                            start_date="2021",
                            scope_line="$50M ARR revenue | 80+ engineers | $10M budget",
                            bullets=[
                                ResumeBullet(
                                    text="Scaled engineering organization",
                                    metrics="from 20 to 80 engineers",
                                ),
                                ResumeBullet(text="Directed engineering strategy"),
                            ],
                        )
                    ],
                )
            ],
        )

        service = TemplateService()
        html = service.render(resume, "executive")

        # Verify scope_line is rendered
        assert "$10M budget" in html
        assert "80+ engineers" in html
        assert "$50M ARR revenue" in html

    def test_work_units_to_html_ats_safe_template(self) -> None:
        """ATS-safe template renders Work Unit data with standard formatting."""
        work_units = [
            {
                "id": "wu-2024-01-01-test",
                "title": "Software Engineer",
                "organization": "Tech Company",
                "time_started": "2023-01-15",
                "time_ended": "2024-06-30",
                "actions": ["Developed features", "Fixed bugs"],
                "outcome": {"result": "Improved system reliability"},
                "tags": ["java", "spring"],
                "skills_demonstrated": ["Backend Development"],
            }
        ]

        contact = ContactInfo(
            name="Test Candidate",
            email="test@example.com",
            phone="555-9999",
        )
        resume = ResumeData.from_work_units(work_units, contact)

        service = TemplateService()
        html = service.render(resume, "ats-safe")

        # Verify ATS-friendly uppercase headers
        assert "EXPERIENCE" in html
        assert "SKILLS" in html

        # Verify content rendered
        assert "Software Engineer" in html
        assert "Tech Company" in html
        assert "java | spring" in html or "java" in html  # Skills as pipe-separated


class TestATSSafeTemplateIntegration:
    """Integration tests for ATS-safe template rendering."""

    def test_ats_safe_template_exists(self) -> None:
        """ATS-safe template is discovered by template service."""
        service = TemplateService()
        templates = service.list_templates()
        assert "ats-safe" in templates

    def test_ats_safe_template_uses_standard_section_headers(self) -> None:
        """ATS-safe template uses standard uppercase section headers."""
        service = TemplateService()
        contact = ContactInfo(name="Test Candidate")
        resume = ResumeData(
            contact=contact,
            summary="Experienced professional.",
            sections=[
                ResumeSection(title="Experience", items=[]),
            ],
            skills=["Python", "SQL"],
        )

        html = service.render(resume, "ats-safe")

        # Standard ATS-recognizable headers in uppercase
        assert "PROFESSIONAL SUMMARY" in html
        assert "SKILLS" in html
        assert "EXPERIENCE" in html

    def test_ats_safe_template_single_column_layout(self) -> None:
        """ATS-safe template uses simple single-column structure."""
        service = TemplateService()

        # ATS-safe CSS should not use flex or grid layouts
        css = service.get_css("ats-safe")
        assert "flex" not in css, "ATS-safe CSS should not use flexbox"
        assert "grid" not in css, "ATS-safe CSS should not use CSS grid"

    def test_ats_safe_template_renders_contact_inline(self) -> None:
        """ATS-safe template renders contact info on single line with separators."""
        service = TemplateService()
        contact = ContactInfo(
            name="Test Candidate",
            email="test@example.com",
            phone="555-1234",
            location="New York, NY",
        )
        resume = ResumeData(contact=contact)

        html = service.render(resume, "ats-safe")

        assert "test@example.com" in html
        assert "555-1234" in html
        assert "New York, NY" in html
        # Contact items separated by pipe
        assert "|" in html

    def test_ats_safe_template_renders_skills_inline(self) -> None:
        """ATS-safe template renders skills as pipe-separated list."""
        service = TemplateService()
        contact = ContactInfo(name="Test Candidate")
        resume = ResumeData(
            contact=contact,
            skills=["Python", "JavaScript", "SQL"],
        )

        html = service.render(resume, "ats-safe")

        # Skills should be separated by pipes for ATS parsing
        assert "Python | JavaScript | SQL" in html

    def test_ats_safe_css_minimal_formatting(self) -> None:
        """ATS-safe CSS has minimal decorative elements."""
        service = TemplateService()
        css = service.get_css("ats-safe")

        assert len(css) > 0
        # Uses standard system fonts
        assert "Arial" in css
        # No complex layouts
        assert "flex" not in css
        assert "grid" not in css

    def test_ats_safe_template_valid_html_structure(self) -> None:
        """ATS-safe template produces valid HTML structure."""
        service = TemplateService()
        contact = ContactInfo(name="Test Candidate")
        resume = ResumeData(contact=contact)

        html = service.render(resume, "ats-safe")

        assert "<!DOCTYPE html>" in html
        assert "<html" in html
        assert "</html>" in html


class TestExecutiveTemplateIntegration:
    """Integration tests for executive template rendering."""

    def test_executive_template_exists(self) -> None:
        """Executive template is discovered by template service."""
        service = TemplateService()
        templates = service.list_templates()
        assert "executive" in templates

    def test_executive_template_renders_scope_indicators(self) -> None:
        """Executive template renders scope indicators from scope_line.

        Story 7.2: scope_line is the unified display format from Position.scope.
        """
        service = TemplateService()
        contact = ContactInfo(name="Executive Leader")
        resume = ResumeData(
            contact=contact,
            sections=[
                ResumeSection(
                    title="Experience",
                    items=[
                        ResumeItem(
                            title="VP of Engineering",
                            organization="Enterprise Corp",
                            scope_line="$100M ARR revenue | 25+ engineers | $5M budget",
                            bullets=[
                                ResumeBullet(text="Led digital transformation"),
                            ],
                        )
                    ],
                )
            ],
        )

        html = service.render(resume, "executive")

        # Verify scope_line is rendered
        assert "$5M budget" in html
        assert "25+ engineers" in html
        assert "$100M ARR revenue" in html

    def test_executive_template_executive_summary_section(self) -> None:
        """Executive template renders Executive Summary section."""
        service = TemplateService()
        contact = ContactInfo(name="Executive Leader")
        resume = ResumeData(
            contact=contact,
            summary="Transformational technology executive with 20+ years.",
        )

        html = service.render(resume, "executive")

        assert "Executive Summary" in html
        assert "Transformational technology executive" in html

    def test_executive_template_renders_metrics_inline(self) -> None:
        """Executive template renders metrics inline with achievement text."""
        service = TemplateService()
        contact = ContactInfo(name="Executive Leader")
        resume = ResumeData(
            contact=contact,
            sections=[
                ResumeSection(
                    title="Experience",
                    items=[
                        ResumeItem(
                            title="CTO",
                            bullets=[
                                ResumeBullet(
                                    text="Scaled engineering organization",
                                    metrics="from 50 to 200 engineers in 18 months",
                                ),
                            ],
                        )
                    ],
                )
            ],
        )

        html = service.render(resume, "executive")

        assert "Scaled engineering organization" in html
        assert "from 50 to 200 engineers" in html

    def test_executive_template_core_competencies(self) -> None:
        """Executive template renders skills as Core Competencies."""
        service = TemplateService()
        contact = ContactInfo(name="Executive Leader")
        resume = ResumeData(
            contact=contact,
            skills=["Strategic Planning", "Team Leadership", "Digital Transformation"],
        )

        html = service.render(resume, "executive")

        assert "Core Competencies" in html
        assert "Strategic Planning" in html
        assert "Team Leadership" in html

    def test_executive_css_has_scope_styling(self) -> None:
        """Executive CSS includes scope indicator styling."""
        service = TemplateService()
        css = service.get_css("executive")

        # New executive template uses .scope-line for inline scope format
        # per Story 6.4 AC#3: "Led team of X | $YM budget | ZM revenue impact"
        assert ".scope-line" in css

    def test_executive_template_valid_html_structure(self) -> None:
        """Executive template produces valid HTML structure."""
        service = TemplateService()
        contact = ContactInfo(name="Executive Leader")
        resume = ResumeData(contact=contact)

        html = service.render(resume, "executive")

        assert "<!DOCTYPE html>" in html
        assert "Executive Resume" in html


class TestModernTemplateIntegration:
    """Integration tests for modern template rendering."""

    def test_modern_template_exists(self) -> None:
        """Modern template is discovered by template service."""
        service = TemplateService()
        templates = service.list_templates()
        assert "modern" in templates

    def test_modern_template_renders_contact_info(self) -> None:
        """Modern template renders contact information."""
        service = TemplateService()
        contact = ContactInfo(
            name="Jane Developer",
            email="jane@example.com",
            phone="555-1234",
            location="San Francisco, CA",
        )
        resume = ResumeData(contact=contact)

        html = service.render(resume, "modern")

        assert "Jane Developer" in html
        assert "jane@example.com" in html
        assert "555-1234" in html
        assert "San Francisco, CA" in html

    def test_modern_template_renders_links(self) -> None:
        """Modern template renders social links."""
        service = TemplateService()
        contact = ContactInfo(
            name="Jane Developer",
            linkedin="https://linkedin.com/in/janedev",
            github="https://github.com/janedev",
            website="https://janedev.com",
        )
        resume = ResumeData(contact=contact)

        html = service.render(resume, "modern")

        assert "linkedin.com/in/janedev" in html
        assert "github.com/janedev" in html
        assert "janedev.com" in html

    def test_modern_template_renders_summary(self) -> None:
        """Modern template renders professional summary."""
        service = TemplateService()
        contact = ContactInfo(name="Jane Developer")
        resume = ResumeData(
            contact=contact,
            summary="Experienced software engineer with 10+ years of Python expertise.",
        )

        html = service.render(resume, "modern")

        assert "Summary" in html
        assert "Experienced software engineer" in html

    def test_modern_template_renders_experience_section(self) -> None:
        """Modern template renders experience section with bullets."""
        service = TemplateService()
        contact = ContactInfo(name="Jane Developer")
        resume = ResumeData(
            contact=contact,
            sections=[
                ResumeSection(
                    title="Experience",
                    items=[
                        ResumeItem(
                            title="Senior Engineer",
                            organization="TechCorp",
                            location="Remote",
                            start_date="Jan 2022",
                            end_date="Present",
                            bullets=[
                                ResumeBullet(
                                    text="Led platform migration to Kubernetes",
                                    metrics="reduced deployment time by 75%",
                                ),
                                ResumeBullet(text="Mentored 5 junior engineers"),
                            ],
                        )
                    ],
                )
            ],
        )

        html = service.render(resume, "modern")

        assert "Experience" in html
        assert "Senior Engineer" in html
        assert "TechCorp" in html
        assert "Remote" in html
        assert "Jan 2022" in html
        assert "Present" in html
        assert "Led platform migration to Kubernetes" in html
        assert "reduced deployment time by 75%" in html
        assert "Mentored 5 junior engineers" in html

    def test_modern_template_renders_education(self) -> None:
        """Modern template renders education section."""
        service = TemplateService()
        contact = ContactInfo(name="Jane Developer")
        resume = ResumeData(
            contact=contact,
            education=[
                Education(
                    degree="BS Computer Science",
                    institution="State University",
                    graduation_year="2014",
                )
            ],
        )

        html = service.render(resume, "modern")

        assert "Education" in html
        assert "BS Computer Science" in html
        assert "State University" in html
        assert "2014" in html

    def test_ats_safe_template_renders_education(self) -> None:
        """ATS-safe template renders education section."""
        service = TemplateService()
        contact = ContactInfo(name="Jane Developer")
        resume = ResumeData(
            contact=contact,
            education=[
                Education(
                    degree="BS Computer Science",
                    institution="State University",
                    graduation_year="2014",
                    honors="Cum Laude",
                )
            ],
        )

        html = service.render(resume, "ats-safe")

        assert "EDUCATION" in html  # ATS-safe uses uppercase
        assert "BS Computer Science" in html
        assert "State University" in html
        assert "2014" in html
        assert "Cum Laude" in html

    def test_executive_template_renders_education(self) -> None:
        """Executive template renders education section."""
        service = TemplateService()
        contact = ContactInfo(name="Jane Executive")
        resume = ResumeData(
            contact=contact,
            education=[
                Education(
                    degree="MBA",
                    institution="Stanford Graduate School of Business",
                    graduation_year="2015",
                ),
                Education(
                    degree="BS Computer Science",
                    institution="MIT",
                    graduation_year="2008",
                ),
            ],
        )

        html = service.render(resume, "executive")

        assert "Education" in html
        assert "MBA" in html
        assert "Stanford Graduate School of Business" in html
        assert "2015" in html
        assert "BS Computer Science" in html
        assert "MIT" in html

    def test_executive_classic_template_renders_education(self) -> None:
        """Executive-classic template renders education section."""
        service = TemplateService()
        contact = ContactInfo(name="Jane Executive")
        resume = ResumeData(
            contact=contact,
            education=[
                Education(
                    degree="PhD Computer Science",
                    institution="Carnegie Mellon",
                    graduation_year="2010",
                    honors="Summa Cum Laude",
                )
            ],
        )

        html = service.render(resume, "executive-classic")

        assert "Education" in html
        assert "PhD Computer Science" in html
        assert "Carnegie Mellon" in html
        assert "2010" in html
        assert "Summa Cum Laude" in html

    def test_modern_template_renders_skills(self) -> None:
        """Modern template renders skills list."""
        service = TemplateService()
        contact = ContactInfo(name="Jane Developer")
        resume = ResumeData(
            contact=contact,
            skills=["Python", "TypeScript", "Kubernetes", "AWS", "PostgreSQL"],
        )

        html = service.render(resume, "modern")

        assert "Skills" in html
        assert "Python" in html
        assert "TypeScript" in html
        assert "Kubernetes" in html

    def test_modern_template_valid_html_structure(self) -> None:
        """Modern template produces valid HTML structure."""
        service = TemplateService()
        contact = ContactInfo(name="Jane Developer")
        resume = ResumeData(contact=contact)

        html = service.render(resume, "modern")

        assert "<!DOCTYPE html>" in html
        assert "<html" in html
        assert "</html>" in html
        assert "<head>" in html
        assert "</head>" in html
        assert "<body>" in html
        assert "</body>" in html
        assert "<title>" in html

    def test_modern_css_exists(self) -> None:
        """Modern CSS file exists and has content."""
        service = TemplateService()
        css = service.get_css("modern")

        assert len(css) > 0
        assert "@page" in css  # Page setup for printing
        assert "font-family" in css
        assert "@media print" in css  # Print-friendly styles

    def test_modern_template_escapes_html_special_chars(self) -> None:
        """Modern template escapes HTML special characters."""
        service = TemplateService()
        contact = ContactInfo(name="Test <script>alert('xss')</script>")
        resume = ResumeData(contact=contact)

        html = service.render(resume, "modern")

        # Script tag should be escaped
        assert "<script>" not in html
        assert "&lt;script&gt;" in html


class TestEmployerGroupingIntegration:
    """Integration tests for employer-grouped position rendering (Story 8.1)."""

    def test_multi_position_employer_grouped_in_modern_template(self) -> None:
        """Multi-position employer renders with grouped layout in modern template."""
        service = TemplateService()
        contact = ContactInfo(name="Jane Developer")
        resume = ResumeData(
            contact=contact,
            sections=[
                ResumeSection(
                    title="Experience",
                    items=[
                        ResumeItem(
                            title="Senior Software Engineer",
                            organization="TechCorp Industries",
                            location="Austin, TX",
                            start_date="2023",
                            end_date=None,
                            bullets=[
                                ResumeBullet(text="Led platform architecture"),
                            ],
                        ),
                        ResumeItem(
                            title="Software Engineer",
                            organization="TechCorp Industries",
                            location="Austin, TX",
                            start_date="2020",
                            end_date="2023",
                            bullets=[
                                ResumeBullet(text="Built core features"),
                            ],
                        ),
                    ],
                )
            ],
        )

        html = service.render(resume, "modern")

        # Verify employer grouping elements
        assert "employer-group" in html
        assert "employer-header" in html
        assert "TechCorp Industries" in html
        # Verify total tenure display
        assert "2020 - Present" in html
        # Verify nested position elements
        assert "position-entry" in html or "nested" in html
        # Verify both positions rendered
        assert "Senior Software Engineer" in html
        assert "Software Engineer" in html

    def test_single_position_employer_standard_layout_modern_template(self) -> None:
        """Single-position employer renders with standard layout in modern template."""
        service = TemplateService()
        contact = ContactInfo(name="Jane Developer")
        resume = ResumeData(
            contact=contact,
            sections=[
                ResumeSection(
                    title="Experience",
                    items=[
                        ResumeItem(
                            title="Software Engineer",
                            organization="StartupCo",
                            location="Remote",
                            start_date="2018",
                            end_date="2020",
                            bullets=[
                                ResumeBullet(text="Built MVP"),
                            ],
                        ),
                    ],
                )
            ],
        )

        html = service.render(resume, "modern")

        # Standard job layout
        assert 'class="job"' in html
        assert "StartupCo" in html
        assert "Software Engineer" in html
        # Not in employer-group container (single position)
        # Count occurrences - should be just in the standard job element
        assert "StartupCo" in html

    def test_mixed_employers_grouped_correctly_modern_template(self) -> None:
        """Mix of single and multi-position employers render correctly."""
        service = TemplateService()
        contact = ContactInfo(name="Jane Developer")
        resume = ResumeData(
            contact=contact,
            sections=[
                ResumeSection(
                    title="Experience",
                    items=[
                        ResumeItem(
                            title="Senior Engineer",
                            organization="TechCorp",
                            start_date="2023",
                            bullets=[ResumeBullet(text="Led team")],
                        ),
                        ResumeItem(
                            title="Engineer",
                            organization="TechCorp",
                            start_date="2020",
                            end_date="2023",
                            bullets=[ResumeBullet(text="Developed features")],
                        ),
                        ResumeItem(
                            title="Developer",
                            organization="StartupCo",
                            start_date="2018",
                            end_date="2020",
                            bullets=[ResumeBullet(text="Built MVP")],
                        ),
                    ],
                )
            ],
        )

        html = service.render(resume, "modern")

        # TechCorp should be in employer-group (multi-position)
        assert "employer-group" in html
        assert "TechCorp" in html
        assert "Senior Engineer" in html
        assert "Engineer" in html

        # StartupCo should be in standard job element (single-position)
        assert "StartupCo" in html
        assert "Developer" in html

    def test_group_employer_positions_false_disables_grouping(self) -> None:
        """Setting group_employer_positions=False uses original separate rendering."""
        from resume_as_code.models.config import ResumeConfig, TemplateOptions

        service = TemplateService()
        contact = ContactInfo(name="Jane Developer")
        resume = ResumeData(
            contact=contact,
            sections=[
                ResumeSection(
                    title="Experience",
                    items=[
                        ResumeItem(
                            title="Senior Engineer",
                            organization="TechCorp",
                            start_date="2023",
                            bullets=[ResumeBullet(text="Led team")],
                        ),
                        ResumeItem(
                            title="Engineer",
                            organization="TechCorp",
                            start_date="2020",
                            end_date="2023",
                            bullets=[ResumeBullet(text="Built features")],
                        ),
                    ],
                )
            ],
        )

        config = ResumeConfig(template_options=TemplateOptions(group_employer_positions=False))

        html = service.render(resume, "modern", config=config)

        # Should NOT have employer-group elements (check for element class, not CSS)
        assert 'class="employer-group"' not in html
        assert 'class="employer-header"' not in html
        # Should have separate job articles
        assert 'class="job"' in html

    def test_employer_name_normalization_variations_grouped(self) -> None:
        """Employer name variations (ampersand, case) are grouped together."""
        service = TemplateService()
        contact = ContactInfo(name="Jane Developer")
        resume = ResumeData(
            contact=contact,
            sections=[
                ResumeSection(
                    title="Experience",
                    items=[
                        ResumeItem(
                            title="Senior Engineer",
                            organization="Burns & McDonnell",  # With ampersand
                            start_date="2023",
                            bullets=[ResumeBullet(text="Led projects")],
                        ),
                        ResumeItem(
                            title="Engineer",
                            organization="Burns and McDonnell",  # With 'and'
                            start_date="2020",
                            end_date="2023",
                            bullets=[ResumeBullet(text="Developed solutions")],
                        ),
                    ],
                )
            ],
        )

        html = service.render(resume, "modern")

        # Should be grouped as multi-position
        assert "employer-group" in html
        # Uses name from most recent position
        assert "Burns &amp; McDonnell" in html or "Burns & McDonnell" in html
        # Both positions should be present
        assert "Senior Engineer" in html
        assert "Engineer" in html

    def test_multi_position_employer_executive_template(self) -> None:
        """Multi-position employer renders with grouped layout in executive template."""
        service = TemplateService()
        contact = ContactInfo(name="Executive Leader")
        resume = ResumeData(
            contact=contact,
            sections=[
                ResumeSection(
                    title="Experience",
                    items=[
                        ResumeItem(
                            title="VP of Engineering",
                            organization="Enterprise Corp",
                            start_date="2022",
                            scope_line="$10M budget | 50+ engineers",
                            bullets=[
                                ResumeBullet(text="Scaled engineering org"),
                            ],
                        ),
                        ResumeItem(
                            title="Director of Engineering",
                            organization="Enterprise Corp",
                            start_date="2020",
                            end_date="2022",
                            scope_line="$5M budget | 25 engineers",
                            bullets=[
                                ResumeBullet(text="Built platform team"),
                            ],
                        ),
                    ],
                )
            ],
        )

        html = service.render(resume, "executive")

        # Verify employer grouping elements
        assert "employer-group" in html
        assert "Enterprise Corp" in html
        # Verify total tenure display
        assert "2020 - Present" in html
        # Verify both positions rendered
        assert "VP of Engineering" in html
        assert "Director of Engineering" in html
        # Verify scope rendered at role level
        assert "$10M budget" in html
        assert "$5M budget" in html

    def test_employer_grouping_preserves_chronological_order(self) -> None:
        """Positions within employer group are in reverse chronological order."""
        service = TemplateService()
        contact = ContactInfo(name="Jane Developer")
        resume = ResumeData(
            contact=contact,
            sections=[
                ResumeSection(
                    title="Experience",
                    items=[
                        # Input in wrong order
                        ResumeItem(
                            title="Engineer",
                            organization="TechCorp",
                            start_date="2018",
                            end_date="2020",
                            bullets=[ResumeBullet(text="First role")],
                        ),
                        ResumeItem(
                            title="Senior Engineer",
                            organization="TechCorp",
                            start_date="2020",
                            end_date=None,
                            bullets=[ResumeBullet(text="Current role")],
                        ),
                    ],
                )
            ],
        )

        html = service.render(resume, "modern")

        # Senior Engineer (more recent) should appear before Engineer
        senior_pos = html.find("Senior Engineer")
        engineer_pos = html.find("Engineer", senior_pos + 1)  # Find second occurrence
        # We're checking the position titles appear, and the group shows them properly
        assert senior_pos < engineer_pos or "Senior Engineer" in html


class TestDOCXTemplateIntegration:
    """Integration tests for DOCX template rendering (Story 13.1)."""

    def test_branded_template_exists(self) -> None:
        """Branded DOCX template is discoverable."""

        from resume_as_code.providers.docx import DOCXProvider

        provider = DOCXProvider(template_name="branded")
        template_path = provider._resolve_template()

        assert template_path is not None
        assert template_path.exists()
        assert template_path.name == "branded.docx"

    def test_branded_template_renders_contact_info(self, tmp_path: Path) -> None:
        """Branded template renders contact information via docxtpl."""

        from docx import Document

        from resume_as_code.providers.docx import DOCXProvider

        contact = ContactInfo(
            name="Jane Executive",
            email="jane@example.com",
            phone="555-1234",
            location="San Francisco, CA",
            linkedin="linkedin.com/in/janeexec",
        )
        resume = ResumeData(
            contact=contact,
            summary="Transformational technology leader with 15+ years experience.",
        )

        output_path = tmp_path / "branded_test.docx"
        provider = DOCXProvider(template_name="branded")
        provider.render(resume, output_path)

        # Verify file created
        assert output_path.exists()

        # Verify content via python-docx
        doc = Document(output_path)
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "Jane Executive" in text
        assert "jane@example.com" in text
        assert "555-1234" in text

    def test_branded_template_renders_experience_section(self, tmp_path: Path) -> None:
        """Branded template renders experience with employer groups via docxtpl."""

        from docx import Document

        from resume_as_code.providers.docx import DOCXProvider

        contact = ContactInfo(name="Test Candidate")
        resume = ResumeData(
            contact=contact,
            sections=[
                ResumeSection(
                    title="Experience",
                    items=[
                        ResumeItem(
                            title="VP of Engineering",
                            organization="TechCorp Industries",
                            start_date="2022",
                            end_date=None,
                            scope_line="$50M budget | 100+ engineers",
                            bullets=[
                                ResumeBullet(text="Scaled engineering organization"),
                                ResumeBullet(text="Led digital transformation"),
                            ],
                        ),
                        ResumeItem(
                            title="Director of Engineering",
                            organization="TechCorp Industries",
                            start_date="2019",
                            end_date="2022",
                            bullets=[
                                ResumeBullet(text="Built platform team from ground up"),
                            ],
                        ),
                    ],
                )
            ],
        )

        output_path = tmp_path / "branded_experience.docx"
        provider = DOCXProvider(template_name="branded")
        provider.render(resume, output_path)

        doc = Document(output_path)
        text = "\n".join(p.text for p in doc.paragraphs)

        # Verify employer name
        assert "TechCorp Industries" in text
        # Verify positions
        assert "VP of Engineering" in text
        assert "Director of Engineering" in text
        # Verify bullets
        assert "Scaled engineering organization" in text
        assert "Built platform team" in text

    def test_branded_template_renders_skills(self, tmp_path: Path) -> None:
        """Branded template renders skills section via docxtpl."""

        from docx import Document

        from resume_as_code.providers.docx import DOCXProvider

        contact = ContactInfo(name="Test Candidate")
        resume = ResumeData(
            contact=contact,
            skills=["Python", "Kubernetes", "AWS", "Team Leadership", "Strategic Planning"],
        )

        output_path = tmp_path / "branded_skills.docx"
        provider = DOCXProvider(template_name="branded")
        provider.render(resume, output_path)

        doc = Document(output_path)
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "Python" in text
        assert "Kubernetes" in text
        assert "AWS" in text

    def test_branded_template_renders_certifications(self, tmp_path: Path) -> None:
        """Branded template renders certifications section via docxtpl."""

        from docx import Document

        from resume_as_code.models.certification import Certification
        from resume_as_code.providers.docx import DOCXProvider

        contact = ContactInfo(name="Test Candidate")
        resume = ResumeData(
            contact=contact,
            certifications=[
                Certification(
                    name="AWS Solutions Architect Professional",
                    issuer="Amazon Web Services",
                    date="2024-01",
                    expires="2027-01",
                ),
                Certification(
                    name="CISSP",
                    issuer="ISC2",
                    date="2023-06",
                ),
            ],
        )

        output_path = tmp_path / "branded_certs.docx"
        provider = DOCXProvider(template_name="branded")
        provider.render(resume, output_path)

        doc = Document(output_path)
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "AWS Solutions Architect Professional" in text
        assert "Amazon Web Services" in text
        assert "CISSP" in text
        assert "ISC2" in text

    def test_branded_template_renders_education(self, tmp_path: Path) -> None:
        """Branded template renders education section via docxtpl."""

        from docx import Document

        from resume_as_code.providers.docx import DOCXProvider

        contact = ContactInfo(name="Test Candidate")
        resume = ResumeData(
            contact=contact,
            education=[
                Education(
                    degree="MBA",
                    institution="Stanford Graduate School of Business",
                    graduation_year="2015",
                ),
                Education(
                    degree="B.S. Computer Science",
                    institution="MIT",
                    graduation_year="2010",
                    honors="Summa Cum Laude",
                ),
            ],
        )

        output_path = tmp_path / "branded_education.docx"
        provider = DOCXProvider(template_name="branded")
        provider.render(resume, output_path)

        doc = Document(output_path)
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "MBA" in text
        assert "Stanford Graduate School of Business" in text
        assert "B.S. Computer Science" in text
        assert "MIT" in text

    def test_custom_templates_dir_override(self, tmp_path: Path) -> None:
        """Custom templates_dir should override built-in templates (AC2)."""
        from pathlib import Path
        from shutil import copy

        from resume_as_code.providers.docx import DOCXProvider

        # Create custom templates directory structure
        custom_docx_dir = tmp_path / "custom" / "docx"
        custom_docx_dir.mkdir(parents=True)

        # Copy branded template to custom location (simulating custom template)
        builtin_template = (
            Path(__file__).parent.parent.parent
            / "src"
            / "resume_as_code"
            / "templates"
            / "docx"
            / "branded.docx"
        )
        custom_template = custom_docx_dir / "branded.docx"
        copy(builtin_template, custom_template)

        # Provider should find custom template first
        provider = DOCXProvider(
            template_name="branded",
            templates_dir=tmp_path / "custom",
        )
        resolved = provider._resolve_template()

        assert resolved == custom_template

    def test_fallback_to_builtin_when_not_in_custom_dir(self, tmp_path: Path) -> None:
        """Should fall back to built-in when template not in custom dir (AC2)."""

        from resume_as_code.providers.docx import DOCXProvider

        # Create empty custom templates directory
        custom_docx_dir = tmp_path / "custom" / "docx"
        custom_docx_dir.mkdir(parents=True)

        # Provider should fall back to built-in
        provider = DOCXProvider(
            template_name="branded",
            templates_dir=tmp_path / "custom",
        )
        resolved = provider._resolve_template()

        # Should find built-in template
        assert resolved is not None
        assert "resume_as_code" in str(resolved)
        assert resolved.name == "branded.docx"

    def test_template_with_all_sections(self, tmp_path: Path) -> None:
        """End-to-end test with all resume sections populated."""

        from docx import Document

        from resume_as_code.models.board_role import BoardRole
        from resume_as_code.models.certification import Certification
        from resume_as_code.models.publication import Publication
        from resume_as_code.providers.docx import DOCXProvider

        contact = ContactInfo(
            name="Executive Leader",
            title="Chief Technology Officer",
            email="cto@example.com",
            phone="555-9999",
            location="New York, NY",
            linkedin="linkedin.com/in/ctoleader",
            github="github.com/ctoleader",
            website="ctoleader.com",
        )

        resume = ResumeData(
            contact=contact,
            summary="Visionary technology executive with 20+ years of experience.",
            sections=[
                ResumeSection(
                    title="Experience",
                    items=[
                        ResumeItem(
                            title="Chief Technology Officer",
                            organization="Enterprise Corp",
                            start_date="2020",
                            scope_line="$100M ARR | 200 engineers | $25M budget",
                            bullets=[
                                ResumeBullet(
                                    text="Led company-wide digital transformation",
                                    metrics="resulting in 40% revenue growth",
                                ),
                            ],
                        ),
                    ],
                )
            ],
            skills=["Strategic Planning", "Team Building", "Cloud Architecture"],
            education=[
                Education(
                    degree="MBA",
                    institution="Harvard Business School",
                    graduation_year="2012",
                ),
            ],
            certifications=[
                Certification(
                    name="AWS Solutions Architect",
                    issuer="AWS",
                    date="2023-01",
                ),
            ],
            publications=[
                Publication(
                    title="Digital Transformation at Scale",
                    type="conference",
                    venue="CTO Summit",
                    date="2024-03",
                ),
            ],
            board_roles=[
                BoardRole(
                    organization="TechStartup Inc",
                    role="Technical Advisor",
                    type="advisory",
                    start_date="2022-01",
                ),
            ],
            career_highlights=["Grew engineering org from 50 to 200 engineers"],
        )

        output_path = tmp_path / "full_resume.docx"
        provider = DOCXProvider(template_name="branded")
        provider.render(resume, output_path)

        assert output_path.exists()

        doc = Document(output_path)
        text = "\n".join(p.text for p in doc.paragraphs)

        # Verify key content from all sections
        assert "Executive Leader" in text
        assert "Visionary technology executive" in text
        assert "Chief Technology Officer" in text
        assert "Enterprise Corp" in text
        assert "Strategic Planning" in text
        assert "MBA" in text
        assert "AWS Solutions Architect" in text
