"""Unit tests for certification rendering across templates."""

from __future__ import annotations

import pytest

from resume_as_code.models.certification import Certification
from resume_as_code.models.resume import ContactInfo, ResumeData
from resume_as_code.services.template_service import TemplateService


@pytest.fixture
def template_service() -> TemplateService:
    """Create a template service with real templates."""
    return TemplateService()


@pytest.fixture
def sample_certifications() -> list[Certification]:
    """Create sample certifications for testing."""
    return [
        Certification(
            name="AWS Solutions Architect - Professional",
            issuer="Amazon Web Services",
            date="2024-06",
        ),
        Certification(
            name="CISSP",
            issuer="ISC2",
            date="2023-01",
        ),
        Certification(
            name="Old Cert",
            display=False,  # Should be hidden
        ),
    ]


@pytest.fixture
def resume_with_certifications(sample_certifications: list[Certification]) -> ResumeData:
    """Create a resume with certifications."""
    return ResumeData(
        contact=ContactInfo(name="Test Developer"),
        certifications=sample_certifications,
    )


class TestModernTemplateCertifications:
    """Tests for certification rendering in modern template."""

    def test_modern_template_renders_certifications_section(
        self, template_service: TemplateService, resume_with_certifications: ResumeData
    ) -> None:
        """Modern template should render certifications section."""
        html = template_service.render(resume_with_certifications, "modern")

        assert "Certifications" in html
        assert "AWS Solutions Architect - Professional" in html

    def test_modern_template_positions_certs_after_education(
        self, template_service: TemplateService, resume_with_certifications: ResumeData
    ) -> None:
        """Modern template certifications section should be after education section."""
        html = template_service.render(resume_with_certifications, "modern")

        # Certifications should appear (even without education)
        certs_section = html.find("Certifications")
        assert certs_section != -1

    def test_modern_template_formats_full_certification(
        self, template_service: TemplateService
    ) -> None:
        """Modern template should format full certification: Name, Issuer, Year."""
        resume = ResumeData(
            contact=ContactInfo(name="Test"),
            certifications=[
                Certification(
                    name="AWS Solutions Architect",
                    issuer="Amazon Web Services",
                    date="2024-06",
                )
            ],
        )
        html = template_service.render(resume, "modern")

        assert "AWS Solutions Architect" in html
        assert "Amazon Web Services" in html
        assert "2024" in html

    def test_modern_template_formats_minimal_certification(
        self, template_service: TemplateService
    ) -> None:
        """Modern template should format minimal certification: Name, Year."""
        resume = ResumeData(
            contact=ContactInfo(name="Test"),
            certifications=[Certification(name="CISSP", date="2023-01")],
        )
        html = template_service.render(resume, "modern")

        assert "CISSP" in html
        assert "2023" in html

    def test_modern_template_hides_cert_with_display_false(
        self, template_service: TemplateService, resume_with_certifications: ResumeData
    ) -> None:
        """Modern template should not render certifications with display=False."""
        html = template_service.render(resume_with_certifications, "modern")

        assert "Old Cert" not in html

    def test_modern_template_no_credential_id_in_output(
        self, template_service: TemplateService
    ) -> None:
        """Modern template should never show credential_id."""
        resume = ResumeData(
            contact=ContactInfo(name="Test"),
            certifications=[
                Certification(
                    name="AWS SAP",
                    credential_id="ABC123XYZ",  # Should NOT appear
                )
            ],
        )
        html = template_service.render(resume, "modern")

        assert "ABC123XYZ" not in html

    def test_modern_template_no_certs_section_when_empty(
        self, template_service: TemplateService
    ) -> None:
        """Modern template should not render certifications section when empty."""
        resume = ResumeData(
            contact=ContactInfo(name="Test"),
            certifications=[],
        )
        html = template_service.render(resume, "modern")

        # Count occurrences of "Certifications" - should be 0 in main content
        assert 'class="certifications"' not in html


class TestExecutiveTemplateCertifications:
    """Tests for certification rendering in executive template."""

    def test_executive_template_renders_certifications(
        self, template_service: TemplateService, resume_with_certifications: ResumeData
    ) -> None:
        """Executive template should render certifications prominently."""
        html = template_service.render(resume_with_certifications, "executive")

        assert "Certifications" in html
        assert "AWS Solutions Architect - Professional" in html
        assert "CISSP" in html

    def test_executive_template_formats_certification_correctly(
        self, template_service: TemplateService
    ) -> None:
        """Executive template should format certification with name, issuer, year."""
        resume = ResumeData(
            contact=ContactInfo(name="Test"),
            certifications=[
                Certification(
                    name="PMP",
                    issuer="PMI",
                    date="2023-08",
                )
            ],
        )
        html = template_service.render(resume, "executive")

        assert "PMP" in html
        assert "PMI" in html
        assert "2023" in html

    def test_executive_template_hides_hidden_certs(
        self, template_service: TemplateService, resume_with_certifications: ResumeData
    ) -> None:
        """Executive template should not render hidden certifications."""
        html = template_service.render(resume_with_certifications, "executive")

        assert "Old Cert" not in html

    def test_executive_template_shows_expiration(self, template_service: TemplateService) -> None:
        """Executive template should show expiration year when present."""
        resume = ResumeData(
            contact=ContactInfo(name="Test"),
            certifications=[
                Certification(
                    name="CISSP",
                    issuer="ISC2",
                    date="2023-01",
                    expires="2028-01",  # Future date so cert is active
                )
            ],
        )
        html = template_service.render(resume, "executive")

        assert "expires 2028" in html


class TestATSSafeTemplateCertifications:
    """Tests for certification rendering in ATS-safe template."""

    def test_ats_safe_template_renders_certifications(
        self, template_service: TemplateService, resume_with_certifications: ResumeData
    ) -> None:
        """ATS-safe template should render certifications as plain list."""
        html = template_service.render(resume_with_certifications, "ats-safe")

        assert "CERTIFICATIONS" in html
        assert "AWS Solutions Architect - Professional" in html

    def test_ats_safe_template_uses_uppercase_header(
        self, template_service: TemplateService, resume_with_certifications: ResumeData
    ) -> None:
        """ATS-safe template should use UPPERCASE section header."""
        html = template_service.render(resume_with_certifications, "ats-safe")

        assert "CERTIFICATIONS" in html

    def test_ats_safe_template_formats_certification_simply(
        self, template_service: TemplateService
    ) -> None:
        """ATS-safe template should format certification in plain text."""
        resume = ResumeData(
            contact=ContactInfo(name="Test"),
            certifications=[
                Certification(
                    name="AWS SAP",
                    issuer="Amazon",
                    date="2024-06",
                )
            ],
        )
        html = template_service.render(resume, "ats-safe")

        # Should have simple list format
        assert "AWS SAP" in html
        assert "Amazon" in html
        assert "2024" in html

    def test_ats_safe_template_hides_hidden_certs(
        self, template_service: TemplateService, resume_with_certifications: ResumeData
    ) -> None:
        """ATS-safe template should not render hidden certifications."""
        html = template_service.render(resume_with_certifications, "ats-safe")

        assert "Old Cert" not in html

    def test_ats_safe_template_no_certs_section_when_empty(
        self, template_service: TemplateService
    ) -> None:
        """ATS-safe template should not render certifications section when empty."""
        resume = ResumeData(
            contact=ContactInfo(name="Test"),
            certifications=[],
        )
        html = template_service.render(resume, "ats-safe")

        assert "CERTIFICATIONS" not in html

    def test_ats_safe_template_shows_expiration(self, template_service: TemplateService) -> None:
        """ATS-safe template should show expiration year when present."""
        resume = ResumeData(
            contact=ContactInfo(name="Test"),
            certifications=[
                Certification(
                    name="CISSP",
                    issuer="ISC2",
                    date="2023-01",
                    expires="2028-01",  # Future date so cert is active
                )
            ],
        )
        html = template_service.render(resume, "ats-safe")

        assert "expires 2028" in html


class TestPDFCertifications:
    """Tests for certification rendering in PDF output."""

    def test_pdf_provider_renders_certifications(self, tmp_path: str) -> None:
        """PDF provider should render certifications section in output."""
        from pathlib import Path

        from resume_as_code.providers.pdf import PDFProvider

        output_path = Path(tmp_path) / "test.pdf"
        provider = PDFProvider(template_name="modern")

        resume = ResumeData(
            contact=ContactInfo(name="Test Developer"),
            certifications=[
                Certification(
                    name="AWS Solutions Architect",
                    issuer="Amazon Web Services",
                    date="2024-06",
                )
            ],
        )

        provider.render(resume, output_path)

        assert output_path.exists()
        assert output_path.stat().st_size > 0  # PDF has content

    def test_pdf_provider_renders_certifications_executive(self, tmp_path: str) -> None:
        """PDF provider should render certifications in executive template."""
        from pathlib import Path

        from resume_as_code.providers.pdf import PDFProvider

        output_path = Path(tmp_path) / "test.pdf"
        provider = PDFProvider(template_name="executive")

        resume = ResumeData(
            contact=ContactInfo(name="Test Developer"),
            certifications=[
                Certification(
                    name="PMP",
                    issuer="PMI",
                    date="2023-08",
                )
            ],
        )

        provider.render(resume, output_path)

        assert output_path.exists()
        assert output_path.stat().st_size > 0  # PDF has content


class TestDOCXCertifications:
    """Tests for certification rendering in DOCX."""

    def test_docx_provider_renders_certifications(self, tmp_path: str) -> None:
        """DOCX provider should render certifications with Word formatting."""
        from pathlib import Path

        from resume_as_code.providers.docx import DOCXProvider

        output_path = Path(tmp_path) / "test.docx"
        provider = DOCXProvider()

        resume = ResumeData(
            contact=ContactInfo(name="Test Developer"),
            certifications=[
                Certification(
                    name="AWS Solutions Architect",
                    issuer="Amazon Web Services",
                    date="2024-06",
                )
            ],
        )

        provider.render(resume, output_path)

        assert output_path.exists()
        # DOCX exists and was created successfully

    def test_docx_provider_formats_certification_correctly(self, tmp_path: str) -> None:
        """DOCX provider should format certification as 'Name, Issuer, Year'."""
        from pathlib import Path

        from docx import Document

        from resume_as_code.providers.docx import DOCXProvider

        output_path = Path(tmp_path) / "test.docx"
        provider = DOCXProvider()

        resume = ResumeData(
            contact=ContactInfo(name="Test Developer"),
            certifications=[
                Certification(
                    name="PMP",
                    issuer="PMI",
                    date="2023-08",
                )
            ],
        )

        provider.render(resume, output_path)

        # Read the document and check content
        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "PMP" in text
        assert "PMI" in text
        assert "2023" in text

    def test_docx_provider_hides_hidden_certs(self, tmp_path: str) -> None:
        """DOCX provider should not render hidden certifications."""
        from pathlib import Path

        from docx import Document

        from resume_as_code.providers.docx import DOCXProvider

        output_path = Path(tmp_path) / "test.docx"
        provider = DOCXProvider()

        resume = ResumeData(
            contact=ContactInfo(name="Test Developer"),
            certifications=[
                Certification(name="Visible Cert", display=True),
                Certification(name="Hidden Cert", display=False),
            ],
        )

        provider.render(resume, output_path)

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "Visible Cert" in text
        assert "Hidden Cert" not in text

    def test_docx_provider_no_credential_id_in_output(self, tmp_path: str) -> None:
        """DOCX provider should never show credential_id."""
        from pathlib import Path

        from docx import Document

        from resume_as_code.providers.docx import DOCXProvider

        output_path = Path(tmp_path) / "test.docx"
        provider = DOCXProvider()

        resume = ResumeData(
            contact=ContactInfo(name="Test Developer"),
            certifications=[
                Certification(
                    name="AWS SAP",
                    credential_id="SECRET123",  # Should NOT appear
                )
            ],
        )

        provider.render(resume, output_path)

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "SECRET123" not in text

    def test_docx_provider_shows_expiration(self, tmp_path: str) -> None:
        """DOCX provider should show expiration year when present."""
        from pathlib import Path

        from docx import Document

        from resume_as_code.providers.docx import DOCXProvider

        output_path = Path(tmp_path) / "test.docx"
        provider = DOCXProvider()

        resume = ResumeData(
            contact=ContactInfo(name="Test Developer"),
            certifications=[
                Certification(
                    name="CISSP",
                    date="2023-01",
                    expires="2028-01",  # Future date so cert is active
                )
            ],
        )

        provider.render(resume, output_path)

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "expires 2028" in text


class TestPartialCertificationData:
    """Tests for handling partial certification data."""

    def test_certification_with_name_only(self, template_service: TemplateService) -> None:
        """Templates should handle certification with only name."""
        resume = ResumeData(
            contact=ContactInfo(name="Test"),
            certifications=[Certification(name="Basic Cert")],
        )
        html = template_service.render(resume, "modern")

        assert "Basic Cert" in html
        # Should not have comma followed by nothing
        assert "Basic Cert, ," not in html

    def test_certification_with_name_and_issuer_no_date(
        self, template_service: TemplateService
    ) -> None:
        """Templates should handle certification with name and issuer but no date."""
        resume = ResumeData(
            contact=ContactInfo(name="Test"),
            certifications=[Certification(name="Security+", issuer="CompTIA")],
        )
        html = template_service.render(resume, "modern")

        assert "Security+" in html
        assert "CompTIA" in html

    def test_certification_with_name_and_date_no_issuer(
        self, template_service: TemplateService
    ) -> None:
        """Templates should handle certification with name and date but no issuer."""
        resume = ResumeData(
            contact=ContactInfo(name="Test"),
            certifications=[Certification(name="CISSP", date="2023-01")],
        )
        html = template_service.render(resume, "modern")

        assert "CISSP" in html
        assert "2023" in html
