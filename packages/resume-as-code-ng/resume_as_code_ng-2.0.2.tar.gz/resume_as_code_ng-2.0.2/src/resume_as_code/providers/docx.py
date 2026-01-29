"""DOCX provider using python-docx."""

from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from resume_as_code.models.certification import Certification
from resume_as_code.models.education import Education
from resume_as_code.models.errors import RenderError
from resume_as_code.models.resume import ResumeData, ResumeItem

logger = logging.getLogger(__name__)


class DOCXProvider:
    """Provider for generating DOCX resumes using python-docx or docxtpl templates."""

    def __init__(
        self,
        template_name: str | None = None,
        templates_dir: Path | None = None,
    ) -> None:
        """Initialize DOCXProvider with optional template support.

        Args:
            template_name: Name of DOCX template (without .docx extension).
            templates_dir: Custom directory to search for templates first.
        """
        self.template_name = template_name
        self.templates_dir = templates_dir

    def _resolve_template(self) -> Path | None:
        """Find DOCX template file, checking custom dir first.

        Template resolution order (AC2):
        1. Custom templates_dir/docx/{template_name}.docx
        2. Built-in templates/docx/{template_name}.docx

        Returns:
            Path to template if found, None otherwise.
        """
        if not self.template_name:
            return None

        template_filename = f"{self.template_name}.docx"

        # Check custom templates dir first
        if self.templates_dir:
            custom_path = self.templates_dir / "docx" / template_filename
            if custom_path.exists():
                return custom_path

        # Fall back to built-in templates
        builtin_path = Path(__file__).parent.parent / "templates" / "docx" / template_filename
        if builtin_path.exists():
            return builtin_path

        return None

    def render(self, resume: ResumeData, output_path: Path) -> Path:
        """Render resume to DOCX file.

        Args:
            resume: ResumeData to render.
            output_path: Path for output DOCX file.

        Returns:
            Path to generated DOCX.

        Raises:
            RenderError: If DOCX generation fails.
        """
        try:
            doc = self._build_document(resume)

            # Ensure output directory exists (idempotent for build command's mkdir)
            output_path.parent.mkdir(parents=True, exist_ok=True)

            # Save document
            doc.save(str(output_path))

            return output_path
        except PermissionError as e:
            raise RenderError(
                message=f"Permission denied writing to {output_path}: {e}",
                path=str(output_path),
                suggestion="Check file permissions or close the file if open elsewhere",
            ) from e
        except OSError as e:
            raise RenderError(
                message=f"DOCX generation failed: {e}",
                path=str(output_path),
                suggestion="Ensure the output directory is writable",
            ) from e
        except Exception as e:
            raise RenderError(
                message=f"Failed to render DOCX: {e}",
                path=str(output_path),
                suggestion="Check the resume data for issues",
            ) from e

    def render_to_bytes(self, resume: ResumeData) -> bytes:
        """Render resume to DOCX bytes.

        Useful for streaming or in-memory processing.

        Args:
            resume: ResumeData to render.

        Returns:
            DOCX content as bytes.

        Raises:
            RenderError: If DOCX generation fails.
        """
        try:
            doc = self._build_document(resume)

            # Write to in-memory buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            return buffer.read()
        except Exception as e:
            raise RenderError(
                message=f"Failed to render DOCX to bytes: {e}",
                suggestion="Check the resume data for issues",
            ) from e

    def _build_document(self, resume: ResumeData) -> Any:
        """Build the Word document from resume data.

        Uses docxtpl template if found, otherwise falls back to programmatic generation.

        Args:
            resume: ResumeData to render.

        Returns:
            python-docx Document object.
        """
        template_path = self._resolve_template()

        if template_path:
            return self._render_from_template(resume, template_path)

        # Fallback to programmatic generation (AC3)
        if self.template_name:
            logger.warning(
                f"DOCX template '{self.template_name}' not found, using programmatic generation"
            )

        return self._build_programmatic(resume)

    def _render_from_template(self, resume: ResumeData, template_path: Path) -> Any:
        """Render resume using docxtpl template.

        Args:
            resume: ResumeData to render.
            template_path: Path to .docx template file.

        Returns:
            python-docx Document object.
        """
        from docxtpl import DocxTemplate  # type: ignore[import-untyped]

        doc = DocxTemplate(str(template_path))
        context = self._build_template_context(resume)
        doc.render(context)
        return doc.docx  # Returns underlying python-docx Document

    def _build_template_context(self, resume: ResumeData) -> dict[str, Any]:
        """Build Jinja2 context dictionary for docxtpl template.

        Args:
            resume: ResumeData to render.

        Returns:
            Dictionary with all template variables.
        """
        # Build employer groups for grouped position rendering
        employer_groups = self._build_employer_groups(resume)

        return {
            # Contact info
            "contact": {
                "name": resume.contact.name,
                "title": resume.contact.title,
                "email": resume.contact.email,
                "phone": resume.contact.phone,
                "location": resume.contact.location,
                "linkedin": resume.contact.linkedin,
                "github": resume.contact.github,
                "website": resume.contact.website,
            },
            # Summary
            "summary": resume.summary,
            # Sections (experience, skills, etc.)
            "sections": [
                {
                    "title": section.title,
                    "items": [
                        {
                            "title": item.title,
                            "organization": item.organization,
                            "location": item.location,
                            "start_date": item.start_date,
                            "end_date": item.end_date or "Present",
                            "scope_line": item.scope_line,
                            "bullets": [b.text for b in item.bullets],
                        }
                        for item in section.items
                    ],
                }
                for section in resume.sections
            ],
            # Employer groups (for grouped position rendering)
            "employer_groups": employer_groups,
            # Skills
            "skills": resume.skills,
            # Optional sections - always return lists (empty if no data)
            # to avoid 'NoneType' is not iterable errors in templates
            "certifications": [
                {
                    "name": c.name,
                    "issuer": c.issuer,
                    "date": c.date,
                    "expires": c.expires,
                    "year": c.date[:4] if c.date else None,
                }
                for c in resume.get_active_certifications()
            ],
            "education": [
                {
                    "degree": e.degree,
                    "institution": e.institution,
                    "graduation_year": e.graduation_year,
                    "honors": e.honors,
                    "gpa": e.gpa,
                }
                for e in resume.education
                if e.display
            ],
            "publications": [
                {
                    "title": p.title,
                    "type": p.type,
                    "venue": p.venue,
                    "date": p.date,
                    "url": p.url,
                }
                for p in resume.get_sorted_publications()
            ],
            "board_roles": [
                {
                    "organization": b.organization,
                    "role": b.role,
                    "type": b.type,
                    "start_date": b.start_date,
                    "end_date": b.end_date,
                    "focus": b.focus,
                }
                for b in resume.get_sorted_board_roles()  # Sorted by type priority and date
            ],
            "highlights": resume.career_highlights or [],
            # Story 7.19: Tailored notice for customized resumes
            "tailored_notice_text": resume.tailored_notice_text,
        }

    def _build_employer_groups(self, resume: ResumeData) -> list[dict[str, Any]]:
        """Build employer-grouped positions for template rendering.

        Groups multiple positions at the same employer together.

        Args:
            resume: ResumeData to extract positions from.

        Returns:
            List of employer groups with nested positions.
        """
        # Find Experience section
        experience_section = next((s for s in resume.sections if s.title == "Experience"), None)
        if not experience_section:
            return []

        # Group positions by employer (organization)
        # Track both positions and location (from first/most recent position)
        groups: dict[str, dict[str, Any]] = {}
        for item in experience_section.items:
            employer = item.organization or "Unknown"
            if employer not in groups:
                groups[employer] = {
                    "positions": [],
                    "location": item.location,  # Use first position's location
                }
            groups[employer]["positions"].append(
                {
                    "title": item.title,
                    "start_date": item.start_date,
                    "end_date": item.end_date or "Present",
                    "scope_line": item.scope_line,
                    "bullets": [b.text for b in item.bullets],
                }
            )

        # Convert to list maintaining order
        return [
            {
                "employer": employer,
                "location": data["location"],
                "positions": data["positions"],
                "date_range": self._compute_employer_date_range(data["positions"]),
                "is_multi_position": len(data["positions"]) > 1,
            }
            for employer, data in groups.items()
        ]

    def _compute_employer_date_range(self, positions: list[dict[str, Any]]) -> str:
        """Compute overall date range for an employer.

        Args:
            positions: List of position dicts with start_date and end_date.

        Returns:
            Date range string like "2020-01 - Present".
        """
        start_dates = [p["start_date"] for p in positions if p.get("start_date")]
        end_dates = [p["end_date"] for p in positions if p.get("end_date")]

        earliest = min(start_dates) if start_dates else ""
        latest = "Present" if "Present" in end_dates else (max(end_dates) if end_dates else "")

        return f"{earliest} - {latest}" if earliest else ""

    def _build_programmatic(self, resume: ResumeData) -> Any:
        """Build Word document programmatically (fallback method).

        Args:
            resume: ResumeData to render.

        Returns:
            python-docx Document object.
        """
        doc: Any = Document()

        # Set up page margins
        for doc_section in doc.sections:
            doc_section.top_margin = Inches(0.75)
            doc_section.bottom_margin = Inches(0.75)
            doc_section.left_margin = Inches(0.75)
            doc_section.right_margin = Inches(0.75)

        # Header with name and contact info
        self._add_header(doc, resume)

        # Summary section
        if resume.summary:
            self._add_section_heading(doc, "Summary")
            p = doc.add_paragraph(resume.summary)
            p.paragraph_format.space_after = Pt(12)

        # Experience sections
        for resume_section in resume.sections:
            self._add_section_heading(doc, resume_section.title)
            for idx, item in enumerate(resume_section.items):
                is_last = idx == len(resume_section.items) - 1
                self._add_experience_item(doc, item, is_last=is_last)

        # Education section
        displayable_education = [edu for edu in resume.education if edu.display]
        if displayable_education:
            self._add_section_heading(doc, "Education")
            for idx, edu in enumerate(displayable_education):
                is_last = idx == len(displayable_education) - 1
                self._add_education_item(doc, edu, is_last=is_last)

        # Certifications section
        active_certs = resume.get_active_certifications()
        if active_certs:
            self._add_certifications_section(doc, active_certs)

        # Skills section
        if resume.skills:
            self._add_section_heading(doc, "Skills")
            p = doc.add_paragraph(", ".join(resume.skills))
            p.paragraph_format.space_after = Pt(12)

        return doc

    def _add_header(self, doc: Any, resume: ResumeData) -> None:
        """Add header with contact information.

        Args:
            doc: Word document to add header to.
            resume: Resume data containing contact info.
        """
        # Name (centered, bold, large font)
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(resume.contact.name)
        name_run.bold = True
        name_run.font.size = Pt(24)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.paragraph_format.space_after = Pt(6)

        # Contact line (email | phone | location)
        contact_parts: list[str] = []
        if resume.contact.email:
            contact_parts.append(resume.contact.email)
        if resume.contact.phone:
            contact_parts.append(resume.contact.phone)
        if resume.contact.location:
            contact_parts.append(resume.contact.location)

        if contact_parts:
            contact_para = doc.add_paragraph(" | ".join(contact_parts))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.paragraph_format.space_after = Pt(3)

        # Links line (linkedin | github | website)
        link_parts: list[str] = []
        if resume.contact.linkedin:
            link_parts.append(resume.contact.linkedin)
        if resume.contact.github:
            link_parts.append(resume.contact.github)
        if resume.contact.website:
            link_parts.append(resume.contact.website)

        if link_parts:
            links_para = doc.add_paragraph(" | ".join(link_parts))
            links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            links_para.paragraph_format.space_after = Pt(12)

    def _add_section_heading(self, doc: Any, title: str) -> None:
        """Add a section heading (Heading 2 style).

        Args:
            doc: Word document to add heading to.
            title: Section title text.
        """
        heading = doc.add_heading(title, level=2)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

    def _add_experience_item(self, doc: Any, item: ResumeItem, *, is_last: bool = False) -> None:
        """Add an experience item with title, org, dates, and bullets.

        Args:
            doc: Word document to add item to.
            item: Experience item to render.
            is_last: Whether this is the last item in the section.
        """
        # Title line with org and dates
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(item.title)
        title_run.bold = True

        if item.organization:
            title_para.add_run(f" | {item.organization}")

        if item.start_date:
            dates = f"{item.start_date} - {item.end_date or 'Present'}"
            title_para.add_run(f"  ({dates})")

        title_para.paragraph_format.space_after = Pt(3)

        # Scope indicators for executive roles (unified scope_line from Position)
        if item.scope_line:
            scope_para = doc.add_paragraph(item.scope_line)
            scope_para.paragraph_format.left_indent = Inches(0.25)
            for run in scope_para.runs:
                run.italic = True
            scope_para.paragraph_format.space_after = Pt(3)

        # Bullets using Word list style
        for bullet in item.bullets:
            bullet_para = doc.add_paragraph(bullet.text, style="List Bullet")
            bullet_para.paragraph_format.space_after = Pt(3)

        # Add spacer only if not the last item (avoid trailing whitespace)
        if not is_last:
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(6)

    def _add_education_item(self, doc: Any, edu: Education, *, is_last: bool = False) -> None:
        """Add an education item.

        Args:
            doc: Word document to add item to.
            edu: Education item to render.
            is_last: Whether this is the last item in the section.
        """
        # Format: "Degree, Institution, Year - Honors" or "(GPA: X)"
        parts = [edu.degree, edu.institution]
        if edu.graduation_year:
            parts.append(edu.graduation_year)

        text = ", ".join(parts)
        if edu.honors:
            text += f" - {edu.honors}"
        elif edu.gpa:
            text += f" (GPA: {edu.gpa})"

        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(3)

        # Add spacer only if not the last item
        if not is_last:
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(6)

    def _add_certifications_section(self, doc: Any, certifications: list[Certification]) -> None:
        """Add certifications section with Word bullet list formatting.

        Args:
            doc: Word document to add section to.
            certifications: List of active certifications to render.
        """
        self._add_section_heading(doc, "Certifications")

        for cert in certifications:
            # Build certification display text: "Name, Issuer, Year"
            parts: list[str] = [cert.name]
            if cert.issuer:
                parts.append(cert.issuer)
            if cert.date:
                parts.append(cert.date[:4])  # Year only
            if cert.expires:
                parts.append(f"expires {cert.expires[:4]}")

            cert_text = ", ".join(parts)
            bullet_para = doc.add_paragraph(cert_text, style="List Bullet")
            bullet_para.paragraph_format.space_after = Pt(3)
