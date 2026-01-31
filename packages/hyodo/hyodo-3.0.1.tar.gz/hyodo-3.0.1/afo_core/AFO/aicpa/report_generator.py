# Trinity Score: 90.0 (Established by Chancellor)
"""AICPA Report Generator - 문서 생성기

眞 (Truth): 정확한 세금 데이터 기반
美 (Beauty): 전문적이고 아름다운 보고서
孝 (Serenity): 버튼 하나로 자동 생성

의존성: python-docx (pip install python-docx)
"""

import csv
import io
import logging
import tempfile
from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING

logger = logging.getLogger(__name__)

# python-docx는 선택적 의존성
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor  # noqa: F401

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None  # type: ignore[misc, assignment]
    WD_ALIGN_PARAGRAPH = None  # type: ignore[misc, assignment]
    logger.warning("[ReportGenerator] python-docx not installed. Word generation disabled.")

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument


def _add_report_header(doc: "DocxDocument", client_name: str) -> None:
    """Step 1: 보고서 헤더 추가"""
    header = doc.add_heading("AFO AICPA Tax Strategy Report", 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph(f"Client: {client_name}")
    doc.add_paragraph("Prepared by: Julie Union[CPA, AFO] Kingdom")
    doc.add_paragraph("-" * 50)


def _add_executive_summary(doc: "DocxDocument", tax_result: dict) -> None:
    """Step 2: 요약 섹션 추가"""
    doc.add_heading("Executive Summary", level=1)
    summary_text = f"""
Based on the 2025 OBBBA tax regulations, we have analyzed your financial situation
and prepared personalized recommendations to optimize your tax position.

Filing Status: {tax_result.get("filing_status", "N/A").upper()}
Gross Income: ${tax_result.get("gross_income", 0):,}
Total Tax Liability: ${tax_result.get("total_tax", 0):,}
Effective Tax Rate: {tax_result.get("effective_federal_rate", 0):.2f}%
"""
    doc.add_paragraph(summary_text)


def _add_tax_analysis_table(doc: "DocxDocument", tax_result: dict) -> None:
    """Step 3: 세금 분석 테이블 추가"""
    doc.add_heading("Tax Analysis", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Item"
    hdr_cells[1].text = "Amount"

    tax_items = [
        ("Gross Income", f"${tax_result.get('gross_income', 0):,}"),
        ("Taxable Income", f"${tax_result.get('taxable_income', 0):,}"),
        ("Federal Tax", f"${tax_result.get('federal_tax', 0):,}"),
        ("CA State Tax", f"${tax_result.get('state_tax', 0):,}"),
        ("Total Tax", f"${tax_result.get('total_tax', 0):,}"),
        ("After-Tax Income", f"${tax_result.get('after_tax_income', 0):,}"),
    ]

    for item, amount in tax_items:
        row = table.add_row().cells
        row[0].text = item
        row[1].text = amount


def _add_sweet_spot_analysis(doc: "DocxDocument", tax_result: dict) -> None:
    """Step 4: OBBBA Sweet Spot 분석 추가"""
    doc.add_heading("OBBBA Sweet Spot Analysis", level=1)
    headroom = tax_result.get("sweet_spot_headroom", 0)
    roth_rec = tax_result.get("roth_conversion_recommendation", 0)

    if headroom > 0:
        doc.add_paragraph(
            f"""
✅ OPPORTUNITY IDENTIFIED

Your current taxable income is within the 12% bracket with ${headroom:,} of headroom.

RECOMMENDATION: Consider converting ${roth_rec:,} from Traditional IRA to Roth IRA
this year to take advantage of the lower tax rate.

This strategy allows you to pay taxes now at 12% instead of potentially higher rates
in the future when Required Minimum Distributions (RMDs) begin.
"""
        )
    else:
        doc.add_paragraph("No additional Sweet Spot opportunity available at current income level.")


def _add_roth_simulation(doc: "DocxDocument", roth_simulation: dict) -> None:
    """Step 5: Roth Ladder 시뮬레이션 테이블 추가"""
    doc.add_heading("Roth Ladder Strategy (Multi-Year)", level=1)
    doc.add_paragraph(
        f"""
Based on the OBBBA provisions (2025-2028), we recommend a phased Roth Conversion strategy:

Estimated Total Savings: ${roth_simulation.get("estimated_savings", 0):,}
"""
    )
    ladder_table = doc.add_table(rows=1, cols=4)
    ladder_table.style = "Table Grid"
    headers = ["Year", "Conversion", "Tax Paid", "Remaining IRA"]
    for i, h in enumerate(headers):
        ladder_table.rows[0].cells[i].text = h

    for year_data in roth_simulation.get("years", []):
        row = ladder_table.add_row().cells
        row[0].text = str(year_data["year"])
        row[1].text = f"${year_data['conversion_amount']:,}"
        row[2].text = f"${year_data['tax_paid']:,}"
        row[3].text = f"${year_data['remaining_ira']:,}"


def _add_risk_alerts(doc: "DocxDocument", tax_result: dict) -> None:
    """Step 6: 리스크 경고 추가"""
    if tax_result.get("irmaa_warning"):
        doc.add_heading("⚠️ Risk Alert: IRMAA", level=1)
        doc.add_paragraph(
            """
WARNING: Your Modified Adjusted Gross Income (MAGI) may trigger
Income-Related Monthly Adjustment Amount (IRMAA) surcharges for Medicare Part B and D.

Consider income smoothing strategies to avoid crossing IRMAA thresholds.
"""
        )


def _add_advice_and_closing(doc: "DocxDocument", tax_result: dict) -> None:
    """Step 7: 조언 및 마무리 추가"""
    doc.add_heading("Julie's Advice", level=1)
    advice_p = doc.add_paragraph()
    advice_p.add_run(tax_result.get("advice", "No additional advice at this time."))
    doc.add_paragraph("-" * 50)
    doc.add_paragraph("Generated by AFO Union[Kingdom, AICPA] Agent Army")
    doc.add_paragraph("眞善美孝永 - Truth, Goodness, Beauty, Serenity, Forever")


def generate_strategy_report(
    client_name: str,
    tax_result: dict,
    roth_simulation: dict | None = None,
    output_path: str | None = None,
) -> str:
    """세금 전략 보고서 생성 (Word .docx) - Refactored"""
    if not DOCX_AVAILABLE:
        return "Error: python-docx not installed. Run: pip install python-docx"

    try:
        doc = Document()
        _add_report_header(doc, client_name)
        _add_executive_summary(doc, tax_result)
        _add_tax_analysis_table(doc, tax_result)
        _add_sweet_spot_analysis(doc, tax_result)

        if roth_simulation:
            _add_roth_simulation(doc, roth_simulation)

        _add_risk_alerts(doc, tax_result)
        _add_advice_and_closing(doc, tax_result)

        if not output_path:
            safe_name = client_name.replace(" ", "_")
            # Use tempfile for secure temporary file creation
            temp_dir = Path(tempfile.gettempdir())
            output_path = str(temp_dir / f"{safe_name}_Tax_Strategy_Report.docx")

        doc.save(output_path)
        logger.info(f"[ReportGenerator] Word 보고서 생성 완료: {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"[ReportGenerator] 보고서 생성 실패: {e}")
        return f"Error: {e!s}"


def generate_turbotax_csv(client_name: str, tax_data: dict, output_path: str | None = None) -> str:
    """TurboTax 입력용 CSV 생성

    孝 (Serenity): 수동 입력 제거
    """
    output = io.StringIO()
    writer = csv.writer(output)

    # 헤더
    writer.writerow(["Field Name", "Value", "Source", "Note"])

    # 데이터 매핑
    rows = [
        ["Taxpayer Name", client_name, "System", ""],
        [
            "Filing Status",
            tax_data.get("filing_status", "Single").upper(),
            "Client Input",
            "",
        ],
        [
            "Gross Income (AGI)",
            tax_data.get("gross_income", 0),
            "W-2/1099",
            "Verify with documents",
        ],
        [
            "Taxable Income",
            tax_data.get("taxable_income", 0),
            "AFO Calc",
            "After deductions",
        ],
        ["Federal Tax", tax_data.get("federal_tax", 0), "AFO Calc", "2025 OBBBA rates"],
        ["State Tax (CA)", tax_data.get("state_tax", 0), "AFO Calc", ""],
        [
            "Roth Conversion",
            tax_data.get("roth_conversion_recommendation", 0),
            "AFO Strategy",
            "Recommended",
        ],
        ["Tax Year", "2025", "System", "OBBBA in effect"],
    ]

    for row in rows:
        writer.writerow(row)

    content = output.getvalue()

    if output_path:
        with open(output_path, "w") as f:
            f.write(content)
        logger.info(f"[ReportGenerator] TurboTax CSV 생성: {output_path}")
        return output_path

    return content


def generate_quickbooks_csv(
    client_name: str, transaction_data: dict, output_path: str | None = None
) -> str:
    """QuickBooks Online 입력용 CSV 생성

    포맷: Batch Enter Transactions 호환
    """
    output = io.StringIO()
    writer = csv.writer(output)

    # QuickBooks 표준 헤더
    writer.writerow(["Date", "Description", "Debit Account", "Credit Account", "Amount", "Class"])

    # 세금 납부 트랜잭션
    date = datetime.now().strftime("%m/%d/%Y")
    tax_amount = transaction_data.get("total_tax", 0)

    writer.writerow(
        [
            date,
            f"Tax Payment - {client_name}",
            "Tax Expense",
            "Bank Account",
            tax_amount,
            "Personal",
        ]
    )

    # Roth Conversion 트랜잭션 (있는 경우)
    roth_amount = transaction_data.get("roth_conversion_recommendation", 0)
    if roth_amount > 0:
        writer.writerow(
            [
                date,
                f"Roth Conversion - {client_name}",
                "Roth IRA",
                "Traditional IRA",
                roth_amount,
                "Retirement",
            ]
        )

    content = output.getvalue()

    if output_path:
        with open(output_path, "w") as f:
            f.write(content)
        logger.info(f"[ReportGenerator] QuickBooks CSV 생성: {output_path}")
        return output_path

    return content


def generate_email_draft(
    client_name: str,
    tax_result: dict,
    next_steps: str = "Please review and let me know if you have questions.",
) -> str:
    """고객 이메일 초안 생성

    美 (Beauty): 전문적이면서 친근한 톤
    """
    subject = f"Tax Strategy Update for {client_name} - 2025 OBBBA Analysis"

    tax_result.get("roth_conversion_recommendation", 0)
    sweet_spot = tax_result.get("sweet_spot_headroom", 0)

    body = f"""Subject: {subject}

Dear {client_name},

I hope this email finds you well.

Based on my analysis of the 2025 OBBBA tax regulations, I have prepared a personalized tax strategy report for you.

[KEY FINDINGS]
- Your current effective tax rate: {tax_result.get("effective_federal_rate", 0):.2f}%
- Total estimated tax liability: ${tax_result.get("total_tax", 0):,}
- OBBBA Sweet Spot opportunity: ${sweet_spot:,}

[RECOMMENDED ACTION]
{tax_result.get("advice", "No immediate action required.")}

I have attached the detailed report to this email. {next_steps}

Best regards,

Julie Kim, CPA
AFO AICPA Group
---
Powered by AFO Kingdom | 眞善美孝永
"""

    logger.info(f"[ReportGenerator] Email draft generated for {client_name}")
    return body
