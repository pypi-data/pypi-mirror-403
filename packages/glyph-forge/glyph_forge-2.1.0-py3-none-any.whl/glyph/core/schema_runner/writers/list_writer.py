from __future__ import annotations
from typing import Any, Dict, List, Optional

from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from .font_utils import convert_underline_value
from .base_writer import BaseWriter
from ..utils.list_numbering import get_format_config, get_numid_for_type


class ListWriter(BaseWriter):
    """
    Writes list content to the document with full support for all 10 list types.

    List Configuration Schema
    =========================

    The `style.list` object supports the following configuration:

    {
      "format": str,      # "bullet", "decimal", "upperRoman", "lowerRoman", "upperLetter", "lowerLetter"
      "numId": str,       # DOCX numbering ID (1-12 for standard types, or custom)
      "ilvl": str,        # Indentation level ("0", "1", "2", ...)
      "lvlText": str,     # Level text/bullet character (e.g., "•", "%1.", "%1)")
      "numFmt": str       # Number format (same as "format" typically)
    }

    All fields are optional. If not provided, defaults are inferred from descriptor type.

    Standard List Types and NumId Mappings
    =======================================

    Bullet Types:
      - L-BULLET-SOLID (•, ●, -, *) → numId: 1
      - L-BULLET-HOLLOW (◦, o) → numId: 4
      - L-BULLET-SQUARE (▪, ■) → numId: 5

    Ordered Types:
      - L-ORDERED-DOTTED (1., 2., 3.) → numId: 6
      - L-ORDERED-PARA-NUM (1), 2), 3)) → numId: 7
      - L-ORDERED-ROMAN-UPPER (I., II., III.) → numId: 8
      - L-ORDERED-ALPHA-UPPER (A., B., C.) → numId: 9
      - L-ORDERED-ALPHA-LOWER-PAREN (a), b), c)) → numId: 10
      - L-ORDERED-ALPHA-LOWER-DOT (a., b., c.) → numId: 11
      - L-ORDERED-ROMAN-LOWER (i., ii., iii.) → numId: 12

    Usage Examples
    ==============

    Example 1: Automatic format from type
      {
        "type": "L-BULLET-SOLID",
        "features": {"text": "First item"},
        "style": {}  # Format auto-inferred from type
      }

    Example 2: Explicit format configuration
      {
        "type": "L-ORDERED-DOTTED",
        "features": {"text": "First item"},
        "style": {
          "list": {
            "format": "decimal",
            "numId": "6",
            "ilvl": "0",
            "lvlText": "%1."
          }
        }
      }

    Example 3: Custom numId override
      {
        "type": "L-BULLET-SOLID",
        "features": {"text": "First item"},
        "style": {
          "list": {
            "numId": "99"  # Use custom numbering from template
          }
        }
      }

    Entry Points
    ============
      - write(descriptor, style, *, plaintext=None)    # Single item
      - write_block(block, *, plaintext=None)          # Plaintext/DOCX grouped block
    """

    def __init__(self, document, inject_numbering: bool = True, default_bullet_glyph: str = "•"):
        super().__init__(document)
        self.inject_numbering = inject_numbering
        self.default_bullet_glyph = default_bullet_glyph

        # List counters for manual rendering (when numbering.xml not available)
        # Format: {(list_type, ilvl): counter}
        self._list_counters: Dict[tuple, int] = {}
        self._last_list_type: Optional[str] = None

    def _resolve_list_metadata(self, descriptor: Dict[str, Any], effective_style: Dict[str, Any]) -> Dict[str, Any]:
        """
        Resolve list metadata (numId, format, lvlText) from descriptor type and style.

        Priority:
        1. Explicit style.list configuration (user override)
        2. Auto-inferred from descriptor type (using standard mappings)
        3. Fallback to defaults

        Args:
            descriptor: Pattern descriptor with type field
            effective_style: Merged style dict

        Returns:
            Dict with numId, ilvl, format, lvlText keys
        """
        # Start with explicit list metadata from style
        list_meta = dict(effective_style.get("list") or {})

        # If no explicit metadata, try to infer from type
        if not list_meta.get("numId") or not list_meta.get("format"):
            desc_type = descriptor.get("type", "")

            # Get format config from standard mappings
            format_config = get_format_config(desc_type)

            if format_config:
                # Apply defaults for missing fields
                if not list_meta.get("numId"):
                    list_meta["numId"] = format_config["numId"]
                if not list_meta.get("format"):
                    list_meta["format"] = format_config["format"]
                if not list_meta.get("lvlText"):
                    list_meta["lvlText"] = format_config["lvlText"]
                if not list_meta.get("numFmt"):
                    list_meta["numFmt"] = format_config.get("numFmt", format_config["format"])

        # Ensure ilvl has a default
        if "ilvl" not in list_meta:
            list_meta["ilvl"] = "0"

        return list_meta

    # ---------------------------------------------------------------------
    # Numbering.xml detection
    # ---------------------------------------------------------------------
    def _has_numbering_definition(self, numId: str) -> bool:
        """
        Check if the document has a numbering definition for the given numId.

        Returns:
            True if numbering.xml has the numId definition, False otherwise
        """
        # For now, always return False to use manual prefix generation
        # This ensures all list types render correctly without requiring numbering.xml
        #
        # TODO: Implement actual numbering.xml parsing to check if numId exists
        # This would enable true DOCX numbering for templates that have definitions
        return False

    # ---------------------------------------------------------------------
    # Low-level XML helper to attach w:numPr (numId + ilvl) to a paragraph
    # ---------------------------------------------------------------------
    def _set_num_pr(self, paragraph, numId: str, ilvl: str = "0") -> None:
        """
        Inject:
          <w:pPr>
            <w:numPr>
              <w:ilvl w:val="0"/>
              <w:numId w:val="1"/>
            </w:numPr>
          </w:pPr>
        """
        p_elm = paragraph._p
        pPr = p_elm.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p_elm.insert(0, pPr)

        numPr = OxmlElement("w:numPr")
        ilvl_el = OxmlElement("w:ilvl")
        ilvl_el.set(qn("w:val"), str(ilvl))
        numId_el = OxmlElement("w:numId")
        numId_el.set(qn("w:val"), str(numId))
        numPr.append(ilvl_el)
        numPr.append(numId_el)

        existing = pPr.find(qn("w:numPr"))
        if existing is not None:
            pPr.remove(existing)
        pPr.append(numPr)

    # ---------------------------------------------------------------------
    # Helpers
    # ---------------------------------------------------------------------
    def _choose_list_style(self, format_val: Optional[str]) -> str:
        """
        Choose a paragraph style that roughly matches the list type.
        Returns a style *name* that exists in the document, else a safe fallback.
        """
        fmt = (format_val or "").lower()
        candidates = []

        # Try format-specific styles first
        if fmt == "bullet":
            candidates = ["List Bullet", "ListBullet", "List Paragraph"]
        elif fmt in {"decimal", "upperroman", "lowerroman", "upperletter", "lowerletter"}:
            candidates = ["List Number", "ListNumber", "List Paragraph"]
        else:
            candidates = ["List Paragraph"]

        # Try each candidate, return first that exists
        for candidate in candidates:
            try:
                _ = self.doc.styles[candidate]
                return candidate
            except KeyError:
                continue

        # Ultimate fallback to Normal (always exists)
        return "Normal"

    def _apply_font_overrides(self, paragraph, font_dict: Dict[str, Any]) -> None:
        if not font_dict:
            return
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        if "name" in font_dict:
            run.font.name = font_dict["name"]
        if "size" in font_dict and font_dict["size"]:
            run.font.size = Pt(font_dict["size"])
        if "bold" in font_dict:
            run.font.bold = bool(font_dict["bold"])
        if "italic" in font_dict:
            run.font.italic = bool(font_dict["italic"])
        if "underline" in font_dict:
            run.font.underline = convert_underline_value(font_dict["underline"])
        if "strike" in font_dict:
            run.font.strike = bool(font_dict["strike"])
        if "highlight" in font_dict and font_dict["highlight"]:
            highlight_map = {
                "yellow": WD_COLOR_INDEX.YELLOW,
                "brightGreen": WD_COLOR_INDEX.BRIGHT_GREEN,
                "turquoise": WD_COLOR_INDEX.TURQUOISE,
                "pink": WD_COLOR_INDEX.PINK,
                "blue": WD_COLOR_INDEX.BLUE,
                "red": WD_COLOR_INDEX.RED,
                "darkBlue": WD_COLOR_INDEX.DARK_BLUE,
                "teal": WD_COLOR_INDEX.TEAL,
                "green": WD_COLOR_INDEX.GREEN,
                "violet": WD_COLOR_INDEX.VIOLET,
                "darkRed": WD_COLOR_INDEX.DARK_RED,
                "darkYellow": WD_COLOR_INDEX.DARK_YELLOW,
                "gray50": WD_COLOR_INDEX.GRAY_50,
                "gray25": WD_COLOR_INDEX.GRAY_25,
            }
            highlight_val = font_dict["highlight"]
            if highlight_val in highlight_map:
                run.font.highlight_color = highlight_map[highlight_val]
        if "color" in font_dict and font_dict["color"]:
            color_val = font_dict["color"]
            if isinstance(color_val, str):
                # accept "FF0000" or "#FF0000"
                hex_str = color_val.lstrip("#")
                run.font.color.rgb = RGBColor.from_string(hex_str)
            elif isinstance(color_val, RGBColor):
                run.font.color.rgb = color_val

    def _apply_paragraph_overrides(self, paragraph, para_dict: Dict[str, Any], indent_level: Optional[int] = None) -> None:
        # Default indent based on level (if no explicit indent provided)
        if indent_level is not None and "indent_left" not in para_dict:
            paragraph.paragraph_format.left_indent = Inches(0.25 * max(0, int(indent_level)))

        if not para_dict:
            return

        # Alignment
        align = (para_dict.get("alignment") or "").lower()
        if align == "left":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif align == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == "justify":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Spacing before paragraph (in twips, convert to points)
        if "spacing_before" in para_dict:
            paragraph.paragraph_format.space_before = Pt(para_dict["spacing_before"] / 20)

        # Spacing after paragraph (in twips, convert to points)
        if "spacing_after" in para_dict:
            paragraph.paragraph_format.space_after = Pt(para_dict["spacing_after"] / 20)

        # Line spacing
        if "line_spacing" in para_dict:
            paragraph.paragraph_format.line_spacing = para_dict["line_spacing"]

        # Indentation (from style dict, overrides indent_level)
        if "indent_left" in para_dict:
            paragraph.paragraph_format.left_indent = Pt(para_dict["indent_left"] / 20)
        if "indent_right" in para_dict or "right_indent" in para_dict:
            indent_val = para_dict.get("indent_right") or para_dict.get("right_indent")
            paragraph.paragraph_format.right_indent = Pt(indent_val / 20)
        if "indent_hanging" in para_dict:
            paragraph.paragraph_format.first_line_indent = Pt(-para_dict["indent_hanging"] / 20)
        elif "first_line_indent" in para_dict:
            paragraph.paragraph_format.first_line_indent = Pt(para_dict["first_line_indent"] / 20)

        # Paragraph shading (background color)
        # Handle both dict {"fill": "color"} and string "color" formats
        if "shading" in para_dict and para_dict["shading"]:
            shading_value = para_dict["shading"]
            if isinstance(shading_value, dict) and "fill" in shading_value:
                self._apply_paragraph_shading(paragraph, shading_value["fill"])
            elif isinstance(shading_value, str):
                self._apply_paragraph_shading(paragraph, shading_value)

        # Paragraph borders
        if "borders" in para_dict and para_dict["borders"]:
            self._apply_paragraph_borders(paragraph, para_dict["borders"])
    def _generate_list_prefix(self, list_type: str, counter: int, lvlText: str) -> str:
        """
        Generate the appropriate list prefix for manual rendering.

        Args:
            list_type: List type (L-BULLET-SOLID, L-ORDERED-DOTTED, etc.)
            counter: Current counter value (1, 2, 3, ...)
            lvlText: Level text template (e.g., "%1.", "%1)", "•")

        Returns:
            Formatted prefix string
        """
        # Bullet types - use the bullet character directly
        if list_type.startswith("L-BULLET"):
            if lvlText and "%1" not in lvlText:
                return lvlText  # Direct bullet character
            # Fallback based on type
            if "HOLLOW" in list_type:
                return "◦"
            elif "SQUARE" in list_type:
                return "▪"
            else:  # SOLID
                return "•"

        # Ordered types - generate number/letter/roman
        elif list_type.startswith("L-ORDERED"):
            # Determine the number format
            if "ROMAN-UPPER" in list_type:
                number_str = self._int_to_roman(counter).upper()
            elif "ROMAN-LOWER" in list_type:
                number_str = self._int_to_roman(counter).lower()
            elif "ALPHA-UPPER" in list_type:
                number_str = self._int_to_letter(counter).upper()
            elif "ALPHA-LOWER" in list_type:
                number_str = self._int_to_letter(counter).lower()
            else:  # Decimal (DOTTED, PARA-NUM, or generic ORDERED)
                number_str = str(counter)

            # Determine the suffix
            if "PAREN" in list_type or (lvlText and lvlText.endswith(")")):
                return f"{number_str})"
            else:  # DOT or default
                return f"{number_str}."

        # Fallback
        return "•"

    def _int_to_roman(self, num: int) -> str:
        """Convert integer to Roman numeral."""
        val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
        syms = ['m', 'cm', 'd', 'cd', 'c', 'xc', 'l', 'xl', 'x', 'ix', 'v', 'iv', 'i']
        roman_num = ''
        for i in range(len(val)):
            count = num // val[i]
            roman_num += syms[i] * count
            num -= val[i] * count
        return roman_num

    def _int_to_letter(self, num: int) -> str:
        """Convert integer to letter (1=a, 2=b, ..., 26=z, 27=aa)."""
        result = ""
        while num > 0:
            num -= 1  # Make it 0-indexed
            result = chr(ord('a') + (num % 26)) + result
            num //= 26
        return result or "a"

    def _get_or_increment_counter(self, list_type: str, ilvl: str = "0") -> int:
        """
        Get and increment the counter for a list type.

        Resets counter when list type changes (new list detected).

        Args:
            list_type: List type (L-BULLET-SOLID, L-ORDERED-DOTTED, etc.)
            ilvl: Indentation level

        Returns:
            Current counter value
        """
        key = (list_type, ilvl)

        # Reset counters if we switched to a different list type
        if list_type != self._last_list_type:
            self._list_counters.clear()
            self._last_list_type = list_type

        # Get current counter and increment
        current = self._list_counters.get(key, 0) + 1
        self._list_counters[key] = current

        return current

    def _get_glyph_for_list_type(self, list_type: str) -> str:
        """
        Get the appropriate glyph/lvlText for a list type.

        Args:
            list_type: List type (L-BULLET-SOLID, L-ORDERED-DOTTED, etc.)

        Returns:
            The lvlText/glyph for the list type, or default bullet if not found
        """
        from glyph.core.schema_runner.utils.list_numbering import get_format_config

        config = get_format_config(list_type)
        if config:
            return config.get("lvlText", self.default_bullet_glyph)
        return self.default_bullet_glyph

    def _prepend_glyph_if_needed(self, paragraph, glyph: Optional[str], list_type: Optional[str] = None, ilvl: str = "0") -> None:
        """
        For plaintext lists without numbering injection, prepend the appropriate prefix.

        Handles both bullet types (static glyphs) and ordered types (generated numbers).

        Args:
            paragraph: The paragraph to modify
            glyph: The glyph/lvlText template (e.g., "•", "%1.", "%1)")
            list_type: Optional list type for counter management
            ilvl: Indentation level
        """
        if not glyph and not list_type:
            return

        # Generate the prefix
        if list_type and list_type.startswith("L-ORDERED"):
            # Ordered list - generate number/letter/roman
            counter = self._get_or_increment_counter(list_type, ilvl)
            prefix = self._generate_list_prefix(list_type, counter, glyph or "")
        elif "%1" in str(glyph):
            # Has placeholder but no list_type - can't generate, skip
            return
        else:
            # Bullet or static glyph
            if list_type:
                counter = self._get_or_increment_counter(list_type, ilvl)
                prefix = self._generate_list_prefix(list_type, 1, glyph or "")
            else:
                prefix = glyph or self.default_bullet_glyph

        # Insert at the start of the first run
        if paragraph.runs:
            paragraph.runs[0].text = f"{prefix} {paragraph.runs[0].text}"
        else:
            paragraph.add_run(f"{prefix} ")

    # ---------------------------------------------------------------------
    # Public API
    # ---------------------------------------------------------------------
    def write(self, descriptor: Dict[str, Any], style: Dict[str, Any] = None, *, plaintext: Optional[str] = None):
        """
        Write a single list item (or a small set if descriptor['items'] is present).
        The `plaintext` param is the authoritative line text if provided.
        """
        # Effective style: router already merges, but merge again defensively.
        effective_style: Dict[str, Any] = {**(descriptor.get("style") or {}), **(style or {})}

        # Handle section changes (columns, orientation, etc.)
        if "section" in effective_style:
            self._handle_section_change(effective_style["section"])

        # Resolve list metadata (with auto-inference from type)
        list_meta = self._resolve_list_metadata(descriptor, effective_style)
        numId = list_meta.get("numId")
        ilvl = str(list_meta.get("ilvl", "0"))
        fmt = list_meta.get("format")
        lvlText = list_meta.get("lvlText")  # may be a private-use bullet char

        # items: either explicit array or synthesize from features/ plaintext
        items: List[Dict[str, Any]] = descriptor.get("items") or []
        if not items:
            text = (plaintext or descriptor.get("features", {}).get("text") or "")
            items = [{"text": text, "style": effective_style}]

        # STEP A: Try to use style_id from schema first
        style_id_from_schema = effective_style.get("style_id")
        if style_id_from_schema:
            try:
                _ = self.doc.styles[style_id_from_schema]
                fallback_style_name = style_id_from_schema
            except KeyError:
                # style_id doesn't exist, fall back to format-based selection
                fallback_style_name = self._choose_list_style(fmt)
        else:
            # No style_id provided, use format-based selection
            fallback_style_name = self._choose_list_style(fmt)

        for it in items:
            text = it.get("text", "") or ""
            item_style = it.get("style") or {}

            # CRITICAL: Strip bullet/number markers from plaintext to avoid duplication
            # The DOCX numbering will add its own markers, so we need to remove:
            # - "- item" → "item"
            # - "• item" → "item"
            # - "1. item" → "item"
            # - "a) item" → "item"
            if text:
                from glyph.core.analysis.plaintext.parsers.list_parser import normalize_bullet_line
                try:
                    _, clean_text, _, _ = normalize_bullet_line(text)
                    text = clean_text
                except Exception:
                    # If parsing fails, try regex fallback to strip common markers
                    import re
                    text = re.sub(r'^\s*([\-–—•▪◦●·\*■o]|\d+\.?|\d+\)|[a-zA-Z][\.\)]|[ivxlcdmIVXLCDM]+[\.\)])\s*', '', text)

            # Merge effective_style font and paragraph properties with item_style
            # Item-specific properties take precedence
            merged_font = {**(effective_style.get("font") or {}), **(item_style.get("font") or {})}
            merged_paragraph = {**(effective_style.get("paragraph") or {}), **(item_style.get("paragraph") or {})}

            # prefer item's own indent_level; else map from ilvl if present
            indent_level = item_style.get("indent_level")
            if indent_level is None:
                # try to derive from ilvl if available
                try:
                    indent_level = int(ilvl)
                except Exception:
                    indent_level = 0

            # STEP 1: Create EMPTY paragraph
            p = self.doc.add_paragraph()

            # STEP 2: Apply style (if it exists)
            if fallback_style_name:
                try:
                    p.style = fallback_style_name
                except KeyError:
                    # Style doesn't exist in template, continue without it
                    pass

            # STEP 3: Add text in NEW run and apply font overrides
            if text:
                run = p.add_run(text)

                # STEP 4: Apply font overrides to THIS run using merged properties
                if merged_font:
                    if "name" in merged_font:
                        run.font.name = merged_font["name"]
                    if "size" in merged_font and merged_font["size"]:
                        run.font.size = Pt(merged_font["size"])
                    if "bold" in merged_font:
                        run.font.bold = bool(merged_font["bold"])
                    if "italic" in merged_font:
                        run.font.italic = bool(merged_font["italic"])
                    if "underline" in merged_font:
                        run.font.underline = convert_underline_value(merged_font["underline"])
                    if "strike" in merged_font:
                        run.font.strike = bool(merged_font["strike"])
                    if "highlight" in merged_font and merged_font["highlight"]:
                        highlight_map = {
                            "yellow": WD_COLOR_INDEX.YELLOW,
                            "brightGreen": WD_COLOR_INDEX.BRIGHT_GREEN,
                            "turquoise": WD_COLOR_INDEX.TURQUOISE,
                            "pink": WD_COLOR_INDEX.PINK,
                            "blue": WD_COLOR_INDEX.BLUE,
                            "red": WD_COLOR_INDEX.RED,
                            "darkBlue": WD_COLOR_INDEX.DARK_BLUE,
                            "teal": WD_COLOR_INDEX.TEAL,
                            "green": WD_COLOR_INDEX.GREEN,
                            "violet": WD_COLOR_INDEX.VIOLET,
                            "darkRed": WD_COLOR_INDEX.DARK_RED,
                            "darkYellow": WD_COLOR_INDEX.DARK_YELLOW,
                            "gray50": WD_COLOR_INDEX.GRAY_50,
                            "gray25": WD_COLOR_INDEX.GRAY_25,
                        }
                        highlight_val = merged_font["highlight"]
                        if highlight_val in highlight_map:
                            run.font.highlight_color = highlight_map[highlight_val]
                    if "color" in merged_font and merged_font["color"]:
                        color_val = merged_font["color"]
                        if isinstance(color_val, str):
                            hex_str = color_val.lstrip("#")
                            run.font.color.rgb = RGBColor.from_string(hex_str)
                        elif isinstance(color_val, RGBColor):
                            run.font.color.rgb = color_val
                    if "all_caps" in merged_font:
                        run.font.all_caps = merged_font["all_caps"]
                    if "small_caps" in merged_font:
                        run.font.small_caps = merged_font["small_caps"]

            # Numbering injection
            # NOTE: Only inject if template has numbering.xml definitions
            # Otherwise, fall back to manual prefix generation
            if self.inject_numbering and numId and self._has_numbering_definition(numId):
                try:
                    self._set_num_pr(p, numId, ilvl)
                except Exception:
                    # Injection failed - fall back to manual prepending
                    list_type = descriptor.get("type", "")
                    glyph = lvlText or self.default_bullet_glyph
                    self._prepend_glyph_if_needed(p, glyph, list_type=list_type, ilvl=ilvl)
            else:
                # No numbering injection → use manual prefix generation
                list_type = descriptor.get("type", "")
                glyph = lvlText or self.default_bullet_glyph
                self._prepend_glyph_if_needed(p, glyph, list_type=list_type, ilvl=ilvl)

            # Apply paragraph overrides using merged properties
            self._apply_paragraph_overrides(p, merged_paragraph, indent_level=indent_level)

        return self.doc

    def write_block(self, block: Dict[str, Any], *, plaintext: Optional[str] = None):
        """
        Write a whole list block (typically from plaintext handlers).
        Expects block["payload"]["items"] = [{text, ilvl}, ...]
        Optional: block["payload"]["numId"], block["payload"]["ilvl"] (top-level)
        """
        # Handle section changes (columns, orientation, etc.)
        style = block.get("style", {})
        if "section" in style:
            self._handle_section_change(style["section"])

        payload = block.get("payload") or {}
        items: List[Dict[str, Any]] = payload.get("items") or []
        numId = payload.get("numId")  # usually None for plaintext
        top_ilvl = str(payload.get("ilvl", "0"))

        # Extract list types (for counter management)
        list_types = payload.get("list_types", [])
        primary_list_type = list_types[0] if list_types else None

        # Try to get style_id from block or payload
        style_id_from_block = (block.get("style") or {}).get("style_id") or (payload.get("style") or {}).get("style_id")
        if style_id_from_block:
            try:
                _ = self.doc.styles[style_id_from_block]
                fallback_style_name = style_id_from_block
            except KeyError:
                # style_id doesn't exist, fall back to default
                fallback_style_name = self._choose_list_style("bullet")
        else:
            # Choose a reasonable fallback paragraph style
            fallback_style_name = self._choose_list_style("bullet")

        for it in items:
            text = (it.get("text") or "").strip()
            ilvl = str(it.get("ilvl", top_ilvl))
            item_font = it.get("font") or {}
            # Use matched type if available, otherwise fall back to primary_list_type
            item_list_type = it.get("matched_type") or primary_list_type

            # STEP 1: Create EMPTY paragraph
            p = self.doc.add_paragraph()

            # STEP 2: Apply style (if it exists)
            if fallback_style_name:
                try:
                    p.style = fallback_style_name
                except KeyError:
                    # Style doesn't exist in template, continue without it
                    pass

            # STEP 3: Add text in NEW run
            if text:
                run = p.add_run(text)

                # STEP 4: Apply font overrides to THIS run
                if item_font:
                    if "name" in item_font:
                        run.font.name = item_font["name"]
                    if "size" in item_font and item_font["size"]:
                        run.font.size = Pt(item_font["size"])
                    if "bold" in item_font:
                        run.font.bold = bool(item_font["bold"])
                    if "italic" in item_font:
                        run.font.italic = bool(item_font["italic"])
                    if "underline" in item_font:
                        run.font.underline = convert_underline_value(item_font["underline"])
                    if "strike" in item_font:
                        run.font.strike = bool(item_font["strike"])
                    if "highlight" in item_font and item_font["highlight"]:
                        highlight_map = {
                            "yellow": WD_COLOR_INDEX.YELLOW,
                            "brightGreen": WD_COLOR_INDEX.BRIGHT_GREEN,
                            "turquoise": WD_COLOR_INDEX.TURQUOISE,
                            "pink": WD_COLOR_INDEX.PINK,
                            "blue": WD_COLOR_INDEX.BLUE,
                            "red": WD_COLOR_INDEX.RED,
                            "darkBlue": WD_COLOR_INDEX.DARK_BLUE,
                            "teal": WD_COLOR_INDEX.TEAL,
                            "green": WD_COLOR_INDEX.GREEN,
                            "violet": WD_COLOR_INDEX.VIOLET,
                            "darkRed": WD_COLOR_INDEX.DARK_RED,
                            "darkYellow": WD_COLOR_INDEX.DARK_YELLOW,
                            "gray50": WD_COLOR_INDEX.GRAY_50,
                            "gray25": WD_COLOR_INDEX.GRAY_25,
                        }
                        highlight_val = item_font["highlight"]
                        if highlight_val in highlight_map:
                            run.font.highlight_color = highlight_map[highlight_val]
                    if "color" in item_font and item_font["color"]:
                        color_val = item_font["color"]
                        if isinstance(color_val, str):
                            hex_str = color_val.lstrip("#")
                            run.font.color.rgb = RGBColor.from_string(hex_str)
                        elif isinstance(color_val, RGBColor):
                            run.font.color.rgb = color_val
                    if "all_caps" in item_font:
                        run.font.all_caps = item_font["all_caps"]
                    if "small_caps" in item_font:
                        run.font.small_caps = item_font["small_caps"]

            # Try numbering injection only if numId is present (DOCX-derived blocks)
            if self.inject_numbering and numId and self._has_numbering_definition(numId):
                try:
                    self._set_num_pr(p, numId, ilvl)
                except Exception:
                    # Injection failed - use manual prefix
                    # Resolve correct glyph from list type
                    glyph = self._get_glyph_for_list_type(item_list_type) if item_list_type else self.default_bullet_glyph
                    self._prepend_glyph_if_needed(p, glyph, list_type=item_list_type, ilvl=ilvl)
            else:
                # Plaintext or no numbering: use manual prefix generation
                # Resolve correct glyph from list type
                glyph = self._get_glyph_for_list_type(item_list_type) if item_list_type else self.default_bullet_glyph
                self._prepend_glyph_if_needed(p, glyph, list_type=item_list_type, ilvl=ilvl)

            # Simple indent for nested bullets
            try:
                indent_level = int(ilvl)
            except Exception:
                indent_level = 0
            self._apply_paragraph_overrides(p, {}, indent_level=indent_level)

        return self.doc

    def _apply_paragraph_shading(self, paragraph, color: str):
        """
        Apply background shading to a paragraph.

        :param paragraph: The paragraph to shade
        :param color: Hex color string like "FFFF00" or "#FFFF00"
        """
        # Remove # prefix if present
        hex_color = color.lstrip("#")

        # Get paragraph properties element
        pPr = paragraph._element.get_or_add_pPr()

        # Remove existing shading if present
        existing_shd = pPr.find(qn('w:shd'))
        if existing_shd is not None:
            pPr.remove(existing_shd)

        # Create new shading element
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')  # Pattern type
        shd.set(qn('w:color'), 'auto')  # Foreground color
        shd.set(qn('w:fill'), hex_color)  # Background color

        pPr.append(shd)

    def _apply_paragraph_borders(self, paragraph, borders: dict):
        """
        Apply borders to a paragraph.

        :param paragraph: The paragraph to add borders to
        :param borders: Dict with border definitions for top, bottom, left, right
                       Each border should have: color, size, style (optional)
                       Example: {"bottom": {"color": "FF0000", "size": 6, "style": "single"}}
        """
        # Get paragraph properties element
        pPr = paragraph._element.get_or_add_pPr()

        # Remove existing borders if present
        existing_pBdr = pPr.find(qn('w:pBdr'))
        if existing_pBdr is not None:
            pPr.remove(existing_pBdr)

        # Create new paragraph borders element
        pBdr = OxmlElement('w:pBdr')

        # Define border sides
        border_sides = ['top', 'bottom', 'left', 'right']

        for side in border_sides:
            if side in borders:
                border_spec = borders[side]

                # Create border element for this side
                border_elem = OxmlElement(f'w:{side}')

                # Set border style (default: single line)
                border_style = border_spec.get('style', 'single')
                border_elem.set(qn('w:val'), border_style)

                # Set border color (required)
                if 'color' in border_spec:
                    hex_color = border_spec['color'].lstrip('#')
                    border_elem.set(qn('w:color'), hex_color)
                else:
                    border_elem.set(qn('w:color'), '000000')  # Default to black

                # Set border size in eighths of a point (default: 6 = 0.75pt)
                border_size = border_spec.get('size', 6)
                border_elem.set(qn('w:sz'), str(border_size))

                # Set space between border and text (in points)
                border_space = border_spec.get('space', 1)
                border_elem.set(qn('w:space'), str(border_space))

                pBdr.append(border_elem)

        # Only add pBdr if we have at least one border
        if len(pBdr) > 0:
            pPr.append(pBdr)
