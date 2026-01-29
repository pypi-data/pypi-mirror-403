from __future__ import annotations
from typing import Any, Dict, List, Optional

from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from .font_utils import convert_underline_value
from .base_writer import BaseWriter


class TableWriter(BaseWriter):
    """
    Handles writing tables into the document.

    Descriptor-oriented usage (DOCX path):
      write(descriptor, style=None, *, plaintext=None, plaintext_rows=None)

      descriptor example:
      {
        "id": "tbl_1",
        "type": "T",
        "features": {
          "table": { "columns": 3, "header_rows": 1 },
        },
        "style": { "table_style": "Table Grid", "alignment": "left", "borders": "all" }
      }

    Block-oriented usage (plaintext path):
      write_block(block, *, plaintext=None)

      block example:
      {
        "kind": "table",
        "payload": { "rows": [["Project","Role","Year"], ["glyph","Dev","2025"]] }
      }
    """

    def __init__(self, document, default_table_style: str = "Table Grid"):
        super().__init__(document)
        self.default_table_style = default_table_style

    # ---------------------------------------------------------------------
    # Public API
    # ---------------------------------------------------------------------
    def write(
        self,
        descriptor: Dict[str, Any],
        style: Optional[Dict[str, Any]] = None,
        *,
        plaintext: Optional[str] = None,
        plaintext_rows: Optional[List[List[str]]] = None,
    ):
        """
        Build a table using the schema descriptor and optional plaintext rows.
        The `plaintext` param is accepted (for parity with router) but not required.
        """
        # Gather style config first to check for section changes
        features = descriptor.get("features", {})
        style_cfg: Dict[str, Any] = {}
        for candidate in (features.get("style"), descriptor.get("style"), style):
            if isinstance(candidate, dict):
                style_cfg.update(candidate)

        # Handle section changes (columns, orientation, etc.)
        if "section" in style_cfg:
            self._handle_section_change(style_cfg["section"])
        # Gather table config from multiple possible locations
        table_cfg = {}
        for candidate in (descriptor.get("table"), features.get("table")):
            if isinstance(candidate, dict):
                table_cfg.update(candidate)

        # Determine rows
        rows_data = self._resolve_rows(plaintext_rows, descriptor)
        if not rows_data:
            # If columns are defined but no rows, create one empty row
            ncols = int(table_cfg.get("columns", 0)) or 1
            rows_data = [[""] * ncols]

        return self._render_table(rows_data, table_cfg, style_cfg)

    def write_block(self, block: Dict[str, Any], *, plaintext: Optional[str] = None):
        """
        Build a table from a plaintext block (payload.rows is required).
        """
        # Handle section changes (columns, orientation, etc.)
        block_style = block.get("style", {})
        if "section" in block_style:
            self._handle_section_change(block_style["section"])

        payload = block.get("payload") or {}
        rows_data: List[List[str]] = payload.get("rows") or []
        # Basic defaults for style/cfg in plaintext path
        table_cfg: Dict[str, Any] = {
            "header_rows": 1 if rows_data else 0,
        }
        style_cfg: Dict[str, Any] = {
            "table_style": self.default_table_style,
            "alignment": "left",
            "borders": "all",
            "autofit": True,
        }
        return self._render_table(rows_data, table_cfg, style_cfg)

    # ---------------------------------------------------------------------
    # Core rendering
    # ---------------------------------------------------------------------
    def _render_table(self, rows_data: List[List[str]], table_cfg: Dict[str, Any], style_cfg: Dict[str, Any]):
        nrows = len(rows_data)
        ncols = self._infer_ncols(rows_data, table_cfg)

        table = self.doc.add_table(rows=nrows, cols=ncols)

        # Apply table style
        style_name = style_cfg.get("table_style") or self.default_table_style
        self._apply_table_style(table, style_name)

        # Table alignment & layout
        self._apply_table_alignment(table, style_cfg.get("alignment"))
        self._apply_table_autofit(table, style_cfg.get("autofit", True))
        self._apply_column_widths(table, style_cfg.get("col_widths"))

        # Borders (simple)
        if style_cfg.get("borders"):
            self._set_table_borders(table, style_cfg["borders"])

        # Fill cells - apply global font defaults
        font_style = style_cfg.get("font", {})
        for i, row in enumerate(rows_data):
            for j in range(ncols):
                cell = table.cell(i, j)
                text = str(row[j]) if j < len(row) else ""
                # Apply font overrides to cell text
                self._set_cell_text_with_font(cell, text, font_style)

        # Header formatting
        header_rows = int(table_cfg.get("header_rows", 0) or 0)
        header_shading = style_cfg.get("header_shading")  # hex like "FCFCFC" or "#FCFCFC"
        if header_rows > 0:
            for i in range(min(header_rows, nrows)):
                for j in range(ncols):
                    self._bold_cell(table.cell(i, j), True)
                    if header_shading:
                        self._shade_cell(table.cell(i, j), header_shading)

        # Cell paragraph alignment (optional global default)
        cell_align = style_cfg.get("cell_alignment")
        if cell_align:
            for i in range(nrows):
                for j in range(ncols):
                    self._apply_cell_paragraph_alignment(table.cell(i, j), cell_align)

        # Apply font overrides to all cells
        font_overrides = style_cfg.get("font")
        if font_overrides:
            for i in range(nrows):
                for j in range(ncols):
                    self._apply_cell_font_overrides(table.cell(i, j), font_overrides)

        # Apply paragraph overrides to all cells
        paragraph_overrides = style_cfg.get("paragraph")
        if paragraph_overrides:
            for i in range(nrows):
                for j in range(ncols):
                    self._apply_cell_paragraph_overrides(table.cell(i, j), paragraph_overrides)

        return table

    # ---------------------------------------------------------------------
    # Helpers
    # ---------------------------------------------------------------------
    def _resolve_rows(self, plaintext_rows: Optional[List[List[str]]], descriptor: Dict[str, Any]) -> List[List[str]]:
        if plaintext_rows:
            return plaintext_rows
        # Look for rows in descriptor (rare, but keep for DOCX-driven flows)
        features = descriptor.get("features", {})
        for key in ("rows", "table_rows"):
            val = features.get(key)
            if isinstance(val, list) and all(isinstance(r, list) for r in val):
                return val
        return []

    def _infer_ncols(self, rows: List[List[str]], table_cfg: Dict[str, Any]) -> int:
        if rows:
            return max(len(r) for r in rows)
        cols = int(table_cfg.get("columns", 0)) or 1
        return max(1, cols)

    def _apply_table_style(self, table, style_name: str):
        if not style_name:
            return
        # Try the provided name; fallback to default if unavailable
        try:
            table.style = style_name
        except Exception:
            try:
                table.style = self.default_table_style
            except Exception:
                pass  # leave Word default

    def _apply_table_alignment(self, table, align: Optional[str]):
        if not align:
            return
        a = align.lower()
        if a == "left":
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
        elif a == "center":
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
        elif a == "right":
            table.alignment = WD_TABLE_ALIGNMENT.RIGHT

    def _apply_table_autofit(self, table, autofit: bool):
        try:
            table.autofit = bool(autofit)
        except Exception:
            pass

    def _apply_column_widths(self, table, widths: Optional[List[float]]):
        """
        widths: list of inches floats, e.g. [1.5, 2.0, 1.0]
        """
        if not widths:
            return
        ncols = len(table.columns)
        for j in range(min(ncols, len(widths))):
            w = widths[j]
            if w is None:
                continue
            for cell in table.columns[j].cells:
                try:
                    cell.width = Inches(float(w))
                except Exception:
                    # python-docx column width quirks; set each paragraph's left/right indents instead if needed
                    pass

    def _set_cell_text(self, cell, text: str):
        # replace content safely (DEPRECATED - use _set_cell_text_with_font for font support)
        cell.text = str(text)

    def _set_cell_text_with_font(self, cell, text: str, font_style: dict = None):
        """Set cell text and apply font overrides."""
        # Clear existing content
        cell.text = ""

        # Add paragraph with text
        p = cell.paragraphs[0]  # Cell always has at least one paragraph
        if text:
            run = p.add_run(str(text))

            # Apply font overrides if provided
            if font_style:
                if "name" in font_style:
                    run.font.name = font_style["name"]
                if "size" in font_style and font_style["size"]:
                    run.font.size = Pt(font_style["size"])
                if "bold" in font_style:
                    run.font.bold = font_style["bold"]
                if "italic" in font_style:
                    run.font.italic = font_style.get("italic", False)
                if "underline" in font_style:
                    run.font.underline = convert_underline_value(font_style["underline"])
                if "strike" in font_style:
                    run.font.strike = font_style["strike"]
                if "highlight" in font_style and font_style["highlight"]:
                    highlight_map = {
                        "yellow": WD_COLOR_INDEX.YELLOW,
                        "brightGreen": WD_COLOR_INDEX.BRIGHT_GREEN,
                        "turquoise": WD_COLOR_INDEX.TURQUOISE,
                        "pink": WD_COLOR_INDEX.PINK,
                        "blue": WD_COLOR_INDEX.BLUE,
                        "red": WD_COLOR_INDEX.RED,
                        "darkBlue": WD_COLOR_INDEX.DARK_BLUE,
                        "teal": WD_COLOR_INDEX.TEAL,
                        "green": WD_COLOR_INDEX.GREEN,
                        "violet": WD_COLOR_INDEX.VIOLET,
                        "darkRed": WD_COLOR_INDEX.DARK_RED,
                        "darkYellow": WD_COLOR_INDEX.DARK_YELLOW,
                        "gray50": WD_COLOR_INDEX.GRAY_50,
                        "gray25": WD_COLOR_INDEX.GRAY_25,
                    }
                    highlight_val = font_style["highlight"]
                    if highlight_val in highlight_map:
                        run.font.highlight_color = highlight_map[highlight_val]
                if "color" in font_style and font_style["color"]:
                    color_val = font_style["color"]
                    if isinstance(color_val, str):
                        # accept "FF0000" or "#FF0000"
                        hex_str = color_val.lstrip("#")
                        run.font.color.rgb = RGBColor.from_string(hex_str)
                    elif isinstance(color_val, RGBColor):
                        run.font.color.rgb = color_val
                if "all_caps" in font_style:
                    run.font.all_caps = font_style["all_caps"]
                if "small_caps" in font_style:
                    run.font.small_caps = font_style["small_caps"]

    def _bold_cell(self, cell, bold: bool = True):
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = bold

    def _apply_cell_paragraph_alignment(self, cell, align: str):
        a = (align or "").lower()
        for p in cell.paragraphs:
            if a == "left":
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif a == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif a == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif a == "justify":
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ---- XML helpers for borders & shading ----
    def _set_table_borders(self, table, kind: str = "all"):
        """
        kind: "all" | "outer" | "none"
        """
        kind = (kind or "all").lower()
        tbl = table._tbl
        tblPr = tbl.tblPr

        # ensure tblBorders exists
        tblBorders = getattr(tblPr, "tblBorders", None)
        if tblBorders is None:
            tblBorders = OxmlElement("w:tblBorders")
            tblPr.append(tblBorders)

        if kind == "none":
            tblPr.remove(tblBorders)
            return

        # border types
        border_types = ["top", "left", "bottom", "right"]
        if kind == "all":
            border_types += ["insideH", "insideV"]

        for bt in border_types:
            el = tblBorders.find(f"w:{bt}", namespaces=tblPr.nsmap)
            if el is None:
                el = OxmlElement(f"w:{bt}")
                tblBorders.append(el)
            # use qn() for namespaced attributes
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "auto")


    def _shade_cell(self, cell, color: str):
        """
        color: hex like "FCFCFC" or "#FCFCFC"
        """
        if not color:
            return
        hexval = color[1:] if color.startswith("#") else color
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = tcPr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tcPr.append(shd)
        shd.set(qn("w:fill"), hexval)
        # You can also set w:color / w:val if you want patterns

    def _apply_cell_font_overrides(self, cell, font_dict: Dict[str, Any]):
        """
        Apply font overrides to all paragraphs and runs in a cell.
        Similar to paragraph_writer's font override logic.
        """
        if not font_dict:
            return

        for p in cell.paragraphs:
            # Get or create a run
            run = p.runs[0] if p.runs else p.add_run()

            if "name" in font_dict:
                run.font.name = font_dict["name"]
            if "size" in font_dict and font_dict["size"]:
                run.font.size = Pt(font_dict["size"])
            if "bold" in font_dict:
                run.font.bold = bool(font_dict["bold"])
            if "italic" in font_dict:
                run.font.italic = bool(font_dict["italic"])
            if "underline" in font_dict:
                run.font.underline = bool(font_dict["underline"])
            if "color" in font_dict and font_dict["color"]:
                color_val = font_dict["color"]
                if isinstance(color_val, str):
                    # accept "FF0000" or "#FF0000"
                    hex_str = color_val.lstrip("#")
                    run.font.color.rgb = RGBColor.from_string(hex_str)
                elif isinstance(color_val, RGBColor):
                    run.font.color.rgb = color_val

    def _apply_cell_paragraph_overrides(self, cell, para_dict: Dict[str, Any]):
        """
        Apply paragraph overrides to all paragraphs in a cell.
        Similar to paragraph_writer and list_writer paragraph override logic.
        """
        if not para_dict:
            return

        for p in cell.paragraphs:
            # Spacing overrides
            if "spacing_after" in para_dict:
                p.paragraph_format.space_after = Pt(para_dict["spacing_after"])
            if "spacing_before" in para_dict:
                p.paragraph_format.space_before = Pt(para_dict["spacing_before"])
            if "line_spacing" in para_dict:
                p.paragraph_format.line_spacing = para_dict["line_spacing"]

            # Alignment is already handled by _apply_cell_paragraph_alignment
            # but we can also handle it here for consistency
            if "alignment" in para_dict:
                align = para_dict["alignment"].lower()
                if align == "left":
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif align == "center":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == "right":
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif align == "justify":
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
