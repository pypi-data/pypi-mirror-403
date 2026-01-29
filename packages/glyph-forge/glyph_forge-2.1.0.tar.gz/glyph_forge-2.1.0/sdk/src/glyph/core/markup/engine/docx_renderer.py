
"""
Glyph Markup DOCX Renderer
===========================

Renders AST to DOCX using python-docx.

This is a direct python-docx implementation, independent of schema_runner.
Future versions may integrate with schema_runner for consistency.
"""

from typing import Optional, Union, Dict, Any, TYPE_CHECKING
from pathlib import Path
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_COLOR_INDEX,
    WD_UNDERLINE,
    WD_BREAK,
    WD_LINE_SPACING,
)
from docx.enum.section import WD_ORIENT, WD_SECTION

from ..parser.ast import DocumentNode, BlockNode, ParagraphNode, RunNode
from .layout_resolver import resolve_classes, LayoutBundle

if TYPE_CHECKING:
    from .integration import ImageRegistry


class DocxRenderer:
    """
    Renders a Glyph markup AST to a DOCX document.

    Uses python-docx directly for document generation.
    """

    def __init__(
        self,
        template_path: Optional[Union[str, Path]] = None,
        image_registry: Optional["ImageRegistry"] = None,
    ):
        """
        Initialize renderer.

        Args:
            template_path: Optional path to template DOCX file
            image_registry: Optional ImageRegistry mapping image IDs to file paths
        """
        if template_path:
            self.doc = Document(str(template_path))
        else:
            self.doc = Document()

        # Current section context (for layout properties)
        self.current_section_props = {}

        # Image registry for resolving image-id-{key} references
        self.image_registry = image_registry

    def render(self, ast: DocumentNode) -> Document:
        """
        Render AST to DOCX document.

        Args:
            ast: The parsed DocumentNode

        Returns:
            python-docx Document object
        """
        for block in ast.blocks:
            self.render_block(block)

        return self.doc

    def render_block(self, block: BlockNode):
        """
        Render a single block.

        Applies section-level properties and renders paragraphs.
        Detects image blocks and renders them with render_image().

        Args:
            block: BlockNode to render
        """
        # Resolve block classes
        layout = resolve_classes(block.classes)

        # Apply section properties if any
        had_section_props = bool(layout.section_props)
        if had_section_props:
            self.apply_section_properties(layout.section_props)

        # Check if this is an image block
        if layout.image_props and "image_id" in layout.image_props:
            self.render_image(layout.image_props, block.paragraphs)
            return

        # Check for direct image path
        if layout.image_props and "image_path" in layout.image_props:
            self.render_image(layout.image_props, block.paragraphs)
            return

        # Render paragraphs within the block
        for para in block.paragraphs:
            # Merge block-level paragraph props with paragraph-specific props
            para_layout = resolve_classes(para.classes)
            merged_para_props = {**layout.paragraph_props, **para_layout.paragraph_props}
            merged_run_props = {**layout.run_props, **para_layout.run_props}

            self.render_paragraph(
                para,
                paragraph_props=merged_para_props,
                run_props=merged_run_props,
            )

        # CRITICAL: Close section if this block had section properties
        # This ensures columns don't bleed into subsequent blocks
        if had_section_props and layout.section_props.get("columns", 1) > 1:
            # Next block should start with default properties (1 column)
            # Create section break to close this columned section
            if self.doc.paragraphs:
                # Mark that next section should revert to single column
                # This will be handled by apply_section_properties when next block is rendered
                # or we can proactively create a closing section
                default_props = {
                    "columns": 1  # Revert to single column
                }
                # Only create closing section if properties actually changed
                if self._section_props_changed(default_props):
                    self.apply_section_properties(default_props)

    def render_paragraph(
        self,
        para: ParagraphNode,
        paragraph_props: Optional[Dict[str, Any]] = None,
        run_props: Optional[Dict[str, Any]] = None,
    ):
        """
        Render a paragraph with its runs.

        Args:
            para: ParagraphNode to render
            paragraph_props: Paragraph-level properties
            run_props: Default run-level properties for runs in this paragraph
        """
        paragraph_props = paragraph_props or {}
        run_props = run_props or {}

        # Create paragraph
        p = self.doc.add_paragraph()

        # Apply paragraph style if specified
        if "style_id" in paragraph_props:
            style_id = paragraph_props["style_id"]
            # Map slug to actual style name
            style_name = self.map_style_slug(style_id)
            valid_styles = {s.name for s in self.doc.styles}
            if style_name in valid_styles:
                p.style = style_name

        # Render runs
        for run_node in para.runs:
            # Merge run-specific props with inherited props
            run_layout = resolve_classes(run_node.classes)
            merged_run_props = {**run_props, **run_layout.run_props}

            run = p.add_run(run_node.text)
            self.apply_run_properties(run, merged_run_props)

        # Apply paragraph formatting
        self.apply_paragraph_properties(p, paragraph_props)

    def apply_run_properties(self, run, props: Dict[str, Any]):
        """
        Apply run-level properties to a docx run.

        Args:
            run: python-docx Run object
            props: Dictionary of run properties
        """
        # Font name
        if "font_name" in props:
            run.font.name = props["font_name"]

        # Font size
        if "size" in props:
            run.font.size = Pt(props["size"])

        # Bold
        if "bold" in props:
            run.font.bold = props["bold"]

        # Italic
        if "italic" in props:
            run.font.italic = props["italic"]

        # Underline
        if "underline" in props:
            underline_val = props["underline"]
            if underline_val is True:
                run.font.underline = WD_UNDERLINE.SINGLE
            elif underline_val is False:
                run.font.underline = False
            elif isinstance(underline_val, str):
                # Map string to WD_UNDERLINE enum
                underline_map = {
                    "single": WD_UNDERLINE.SINGLE,
                    "double": WD_UNDERLINE.DOUBLE,
                    "dotted": WD_UNDERLINE.DOTTED,
                    "wave": WD_UNDERLINE.WAVY,
                    "thick": WD_UNDERLINE.THICK,
                }
                run.font.underline = underline_map.get(underline_val, WD_UNDERLINE.SINGLE)

        # Strike
        if "strike" in props:
            run.font.strike = props["strike"]

        if "double_strike" in props:
            run.font.double_strike = props["double_strike"]

        # Color
        if "color" in props:
            hex_color = props["color"].lstrip("#")
            run.font.color.rgb = RGBColor.from_string(hex_color)

        # Highlight
        if "highlight" in props:
            highlight_map = {
                "yellow": WD_COLOR_INDEX.YELLOW,
                "green": WD_COLOR_INDEX.GREEN,
                "cyan": WD_COLOR_INDEX.TURQUOISE,
                "magenta": WD_COLOR_INDEX.PINK,
                "blue": WD_COLOR_INDEX.BLUE,
                "red": WD_COLOR_INDEX.RED,
                "dark-blue": WD_COLOR_INDEX.DARK_BLUE,
                "dark-cyan": WD_COLOR_INDEX.TEAL,
                "dark-green": WD_COLOR_INDEX.GREEN,
                "dark-magenta": WD_COLOR_INDEX.VIOLET,
                "dark-red": WD_COLOR_INDEX.DARK_RED,
                "dark-yellow": WD_COLOR_INDEX.DARK_YELLOW,
                "dark-gray": WD_COLOR_INDEX.GRAY_50,
                "light-gray": WD_COLOR_INDEX.GRAY_25,
                "black": WD_COLOR_INDEX.BLACK,
                "none": None,
            }
            highlight_val = props["highlight"]
            if highlight_val in highlight_map:
                hl_color = highlight_map[highlight_val]
                if hl_color:
                    run.font.highlight_color = hl_color

        # All caps
        if "all_caps" in props:
            run.font.all_caps = props["all_caps"]

        # Small caps
        if "small_caps" in props:
            run.font.small_caps = props["small_caps"]

        # Superscript/subscript
        if "superscript" in props:
            run.font.superscript = props["superscript"]

        if "subscript" in props:
            run.font.subscript = props["subscript"]

        # Advanced properties
        for prop_name in ["hidden", "outline", "shadow", "emboss", "imprint"]:
            if prop_name in props:
                setattr(run.font, prop_name, props[prop_name])

    def apply_paragraph_properties(self, paragraph, props: Dict[str, Any]):
        """
        Apply paragraph-level properties.

        Args:
            paragraph: python-docx Paragraph object
            props: Dictionary of paragraph properties
        """
        # Alignment
        if "alignment" in props:
            align_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
                "distribute": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
            }
            paragraph.alignment = align_map.get(props["alignment"], WD_ALIGN_PARAGRAPH.LEFT)

        # Indentation (convert from points to EMU)
        if "left_indent" in props:
            paragraph.paragraph_format.left_indent = Pt(props["left_indent"])

        if "right_indent" in props:
            paragraph.paragraph_format.right_indent = Pt(props["right_indent"])

        if "first_line_indent" in props:
            paragraph.paragraph_format.first_line_indent = Pt(props["first_line_indent"])

        # Spacing
        if "space_before" in props:
            paragraph.paragraph_format.space_before = Pt(props["space_before"])

        if "space_after" in props:
            paragraph.paragraph_format.space_after = Pt(props["space_after"])

        # Line spacing
        if "line_spacing" in props:
            line_spacing = props["line_spacing"]
            if isinstance(line_spacing, float):
                # Multiple (1.0, 1.5, 2.0)
                paragraph.paragraph_format.line_spacing = line_spacing
            elif line_spacing == 1.0:
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            elif line_spacing == 1.5:
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            elif line_spacing == 2.0:
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

        if "line_spacing_pt" in props:
            # Exact spacing in points
            paragraph.paragraph_format.line_spacing = Pt(props["line_spacing_pt"])

        # Pagination
        if "keep_together" in props:
            paragraph.paragraph_format.keep_together = props["keep_together"]

        if "keep_with_next" in props:
            paragraph.paragraph_format.keep_with_next = props["keep_with_next"]

        if "page_break_before" in props:
            paragraph.paragraph_format.page_break_before = props["page_break_before"]

        if "widow_control" in props:
            paragraph.paragraph_format.widow_control = props["widow_control"]

    def apply_section_properties(self, props: Dict[str, Any]):
        """
        Apply section-level properties.

        Creates a new section if needed.

        Args:
            props: Dictionary of section properties
        """
        # Check if section properties have changed from current section
        if self._section_props_changed(props):
            # Create new section with continuous break (for column changes)
            section = self._create_new_section()
        else:
            # Use existing section
            section = self.doc.sections[-1] if len(self.doc.sections) > 0 else self.doc.sections[0]

        # Orientation
        if "orientation" in props:
            if props["orientation"] == "landscape":
                section.orientation = WD_ORIENT.LANDSCAPE
                # Swap width/height for landscape
                new_width = section.page_height
                new_height = section.page_width
                section.page_width = new_width
                section.page_height = new_height
            else:
                section.orientation = WD_ORIENT.PORTRAIT

        # Page size
        if "page_width" in props:
            section.page_width = Inches(props["page_width"])

        if "page_height" in props:
            section.page_height = Inches(props["page_height"])

        # Margins
        if "margin_left" in props:
            section.left_margin = Inches(props["margin_left"])

        if "margin_right" in props:
            section.right_margin = Inches(props["margin_right"])

        if "margin_top" in props:
            section.top_margin = Inches(props["margin_top"])

        if "margin_bottom" in props:
            section.bottom_margin = Inches(props["margin_bottom"])

        # Columns (requires XML manipulation)
        if "columns" in props:
            self.set_section_columns(section, props["columns"])

        # Update current section properties tracker
        self.current_section_props = props.copy()

    def render_image(
        self,
        image_props: Dict[str, Any],
        paragraphs: Optional[list] = None,
    ):
        """
        Render an inline image with optional caption.

        Args:
            image_props: Dictionary with image properties:
                - image_id: ID to look up in image_registry
                - image_path: Direct path to image file (alternative to image_id)
                - image_width: Width in inches
                - image_height: Height in inches
                - image_alignment: 'left', 'center', or 'right'
                - image_caption: If True, render following text as caption
            paragraphs: Optional list of ParagraphNode for caption text
        """
        # Resolve image path
        image_path = None

        if "image_path" in image_props:
            image_path = image_props["image_path"]
        elif "image_id" in image_props and self.image_registry:
            image_id = image_props["image_id"]
            image_path = self.image_registry.get_path(image_id)

        if not image_path or not os.path.exists(image_path):
            # Image not found - skip silently or add placeholder text
            if paragraphs:
                # Still render caption/text content
                for para in paragraphs:
                    self.render_paragraph(para)
            return

        # Create paragraph for image
        p = self.doc.add_paragraph()
        run = p.add_run()

        # Calculate dimensions
        width = None
        height = None

        if "image_width" in image_props:
            width = Inches(image_props["image_width"])

        if "image_height" in image_props:
            height = Inches(image_props["image_height"])

        # Add image to run
        try:
            picture = run.add_picture(image_path, width=width, height=height)

            # Set alt text via XML if available
            if "alt_text" in image_props:
                try:
                    picture._inline.docPr.set("descr", image_props["alt_text"])
                except Exception:
                    pass  # Alt text is optional
        except Exception as e:
            # If image fails to load, add error text
            run.text = f"[Image not found: {image_path}]"
            return

        # Apply alignment
        alignment = image_props.get("image_alignment", "left")
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }
        p.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)

        # Render caption if present
        if image_props.get("image_caption") and paragraphs:
            for para in paragraphs:
                # Render caption text with Caption style
                cap_p = self.doc.add_paragraph()

                # Try to apply Caption style
                valid_styles = {s.name for s in self.doc.styles}
                if "Caption" in valid_styles:
                    cap_p.style = "Caption"

                # Center caption by default
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Add caption text
                for run_node in para.runs:
                    cap_p.add_run(run_node.text)

    def _get_current_section_props(self) -> Dict[str, Any]:
        """
        Get properties of the current section.

        Returns:
            Dictionary of current section properties
        """
        if len(self.doc.sections) == 0:
            return {}

        section = self.doc.sections[-1]
        props = {}

        # Extract orientation
        if section.orientation == WD_ORIENT.LANDSCAPE:
            props["orientation"] = "landscape"
        else:
            props["orientation"] = "portrait"

        # Extract page dimensions (convert from EMU to inches)
        props["page_width"] = section.page_width.inches
        props["page_height"] = section.page_height.inches

        # Extract margins (convert from EMU to inches)
        props["margin_left"] = section.left_margin.inches
        props["margin_right"] = section.right_margin.inches
        props["margin_top"] = section.top_margin.inches
        props["margin_bottom"] = section.bottom_margin.inches

        # Extract column count from XML
        from docx.oxml.ns import qn
        sectPr = section._sectPr
        cols_elem = sectPr.find(qn("w:cols"))
        if cols_elem is not None:
            num_attr = cols_elem.get(qn("w:num"))
            props["columns"] = int(num_attr) if num_attr else 1
        else:
            props["columns"] = 1

        return props

    def _section_props_changed(self, new_props: Dict[str, Any]) -> bool:
        """
        Check if section properties have changed from current section.

        Args:
            new_props: New section properties to apply

        Returns:
            True if properties differ and new section is needed
        """
        # If no sections exist yet, no change needed (will use default section)
        if len(self.doc.sections) == 0:
            return False

        # If no paragraphs added yet, modify the first section directly
        if len(self.doc.paragraphs) == 0:
            return False

        # Use cached current properties if available
        current_props = self.current_section_props if self.current_section_props else self._get_current_section_props()

        # Check critical properties that require new section
        critical_props = ["columns", "orientation", "page_width", "page_height"]

        for prop in critical_props:
            if prop in new_props:
                current_val = current_props.get(prop)
                new_val = new_props[prop]

                # Special handling for numeric comparisons (floating point tolerance)
                if isinstance(new_val, (int, float)) and isinstance(current_val, (int, float)):
                    if abs(new_val - current_val) > 0.01:
                        return True
                elif new_val != current_val:
                    return True

        return False

    def _create_new_section(self):
        """
        Create a new section with a continuous section break.

        Attaches section break to the LAST EXISTING paragraph rather than
        creating an empty paragraph, avoiding unwanted blank lines.

        Returns:
            The newly created section
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # If no paragraphs exist yet, cannot create a section break
        # Caller should modify the default section instead
        if not self.doc.paragraphs:
            # Return the default section
            return self.doc.sections[0]

        # Get the LAST existing paragraph (real content)
        last_para = self.doc.paragraphs[-1]

        # Get the paragraph's XML element
        pPr = last_para._element.get_or_add_pPr()

        # CRITICAL: Get CURRENT section's properties to preserve them
        # The sectPr we attach to the paragraph defines the END of the current section
        # So we need to copy current section's properties (columns, margins, etc.)
        current_section = self.doc.sections[-1]
        current_sectPr = current_section._sectPr

        # Create NEW section properties element for this paragraph
        sectPr = OxmlElement("w:sectPr")

        # Copy important properties from current section
        # This preserves columns, orientation, page size, margins
        from copy import deepcopy

        # Copy page size
        pgSz = current_sectPr.find(qn("w:pgSz"))
        if pgSz is not None:
            sectPr.append(deepcopy(pgSz))

        # Copy margins
        pgMar = current_sectPr.find(qn("w:pgMar"))
        if pgMar is not None:
            sectPr.append(deepcopy(pgMar))

        # Copy columns (CRITICAL - this was being lost!)
        cols = current_sectPr.find(qn("w:cols"))
        if cols is not None:
            sectPr.append(deepcopy(cols))

        # Add section type (continuous break) - this goes AFTER other properties
        sectType = OxmlElement("w:type")
        sectType.set(qn("w:val"), "continuous")
        sectPr.append(sectType)

        # Attach section properties to the LAST REAL paragraph
        # This makes the section break occur AFTER this paragraph
        # The properties in sectPr apply to the CURRENT section (the one ending)
        pPr.append(sectPr)

        # Return the new section (it's now the last one)
        # The new section starts AFTER the paragraph we just modified
        return self.doc.sections[-1]

    def set_section_columns(self, section, num_columns: int):
        """
        Set number of columns in a section.

        python-docx doesn't expose columns API, so we use XML.

        Args:
            section: Section object
            num_columns: Number of columns (1-3)
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        sectPr = section._sectPr
        # Remove existing cols element if present
        cols_elem = sectPr.find(qn("w:cols"))
        if cols_elem is not None:
            sectPr.remove(cols_elem)

        # Create new cols element
        if num_columns > 1:
            cols = OxmlElement("w:cols")
            cols.set(qn("w:num"), str(num_columns))
            sectPr.append(cols)

    def map_style_slug(self, slug: str) -> str:
        """
        Map style slug to actual Word style name.

        Args:
            slug: Style slug (e.g., "heading-1", "body")

        Returns:
            Actual style name
        """
        style_map = {
            "heading-1": "Heading 1",
            "heading-2": "Heading 2",
            "heading-3": "Heading 3",
            "body": "Body Text",
            "normal": "Normal",
            "caption": "Caption",
            "quote": "Quote",
            "intense-quote": "Intense Quote",
            "list-paragraph": "List Paragraph",
        }
        return style_map.get(slug, slug)


def render_ast_to_docx(
    ast: DocumentNode,
    template_path: Optional[Union[str, Path]] = None,
    image_registry: Optional["ImageRegistry"] = None,
) -> Document:
    """
    Render an AST to a python-docx Document.

    Args:
        ast: The DocumentNode to render
        template_path: Optional template DOCX path
        image_registry: Optional ImageRegistry for resolving image IDs

    Returns:
        python-docx Document object

    Examples:
        >>> from glyph.core.markup.parser import parse_markup
        >>> ast = parse_markup("$glyph-bold\\nHello\\n$glyph")
        >>> doc = render_ast_to_docx(ast)
        >>> doc.save("output.docx")
    """
    renderer = DocxRenderer(template_path, image_registry=image_registry)
    return renderer.render(ast)
