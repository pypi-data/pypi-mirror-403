from __future__ import annotations
from typing import Any, Dict, List, Optional

from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from .font_utils import convert_underline_value
from .base_writer import BaseWriter


class RowWriter(BaseWriter):
    """
    Handles writing row layout blocks into the document.

    Row layouts are single-row "layout tables" used for aligned content (e.g., two-column resume lines).
    They are NOT data tables - they are layout primitives optimized for:
    - Single-row or few-row alignment
    - No borders (by default)
    - Predictable width control
    - Cell-level alignment overrides

    Descriptor-oriented usage (Markup path):
      write(descriptor, style=None, *, plaintext=None)

      descriptor example:
      {
        "id": "row_1",
        "type": "R",
        "features": {
          "row": {
            "cols": 2,
            "widths": [70, 30],  # percentages
            "width_mode": "percent",
            "alignment": "left",
            "indent": 0.0,
          },
          "cells": [
            {
              "content": [{"text": "Eli Lilly"}],
              "align": "left",
              "valign": "top",
            },
            {
              "content": [{"text": "December 2023 – December 2025"}],
              "align": "right",
              "valign": "top",
            }
          ]
        },
        "style": {
          "row": {"space_before": 0, "space_after": 0},
          "cell": {"padding": {"top": 0.05, "bottom": 0.05, "left": 0.05, "right": 0.05}, "borders": "none"},
          "font": {...},
          "paragraph": {...}
        }
      }

    Block-oriented usage (plaintext path):
      write_block(block, *, plaintext=None)

      block example:
      {
        "kind": "row",
        "payload": {
          "cells": [
            {"text": "Cell 1", "align": "left"},
            {"text": "Cell 2", "align": "right"}
          ],
          "widths": [70, 30]
        }
      }
    """

    def __init__(self, document):
        super().__init__(document)

    # ---------------------------------------------------------------------
    # Public API
    # ---------------------------------------------------------------------
    def write(
        self,
        descriptor: Dict[str, Any],
        style: Optional[Dict[str, Any]] = None,
        *,
        plaintext: Optional[str] = None,
    ):
        """
        Build a row layout using the schema descriptor.
        """
        # Gather style config first to check for section changes
        features = descriptor.get("features", {})
        style_cfg: Dict[str, Any] = {}
        for candidate in (features.get("style"), descriptor.get("style"), style):
            if isinstance(candidate, dict):
                style_cfg.update(candidate)

        # Handle section changes (columns, orientation, etc.)
        if "section" in style_cfg:
            self._handle_section_change(style_cfg["section"])

        # Gather row config
        row_cfg = features.get("row", {})
        cells_cfg = features.get("cells", [])

        if not cells_cfg:
            # No cells defined - create empty row with one cell
            ncols = int(row_cfg.get("cols", 1))
            cells_cfg = [{"content": [{"text": ""}]} for _ in range(ncols)]

        return self._render_row(cells_cfg, row_cfg, style_cfg)

    def write_block(self, block: Dict[str, Any], *, plaintext: Optional[str] = None):
        """
        Build a row layout from a plaintext block (payload.cells is required).
        """
        # Handle section changes
        block_style = block.get("style", {})
        if "section" in block_style:
            self._handle_section_change(block_style["section"])

        payload = block.get("payload") or {}
        cells_data: List[Dict[str, Any]] = payload.get("cells") or []

        # Convert simple cell data to content format
        cells_cfg = []
        for cell_data in cells_data:
            cell_cfg = {
                "content": [{"text": cell_data.get("text", "")}],
                "align": cell_data.get("align", "left"),
                "valign": cell_data.get("valign", "top"),
            }
            cells_cfg.append(cell_cfg)

        # Basic defaults for row config
        row_cfg: Dict[str, Any] = {
            "cols": len(cells_cfg),
            "widths": payload.get("widths", []),
            "width_mode": payload.get("width_mode", "auto"),
            "alignment": "left",
        }

        # Basic defaults for style
        style_cfg: Dict[str, Any] = {
            "cell": {
                "padding": {"top": 0.05, "bottom": 0.05, "left": 0.05, "right": 0.05},
                "borders": "none",
            },
        }

        return self._render_row(cells_cfg, row_cfg, style_cfg)

    # ---------------------------------------------------------------------
    # Core rendering
    # ---------------------------------------------------------------------
    def _render_row(
        self,
        cells_cfg: List[Dict[str, Any]],
        row_cfg: Dict[str, Any],
        style_cfg: Dict[str, Any],
    ):
        """
        Render a single-row layout table.

        Args:
            cells_cfg: List of cell configurations (content, align, valign, etc.)
            row_cfg: Row configuration (cols, widths, alignment, etc.)
            style_cfg: Style configuration (cell defaults, font, paragraph)
        """
        ncols = int(row_cfg.get("cols", len(cells_cfg)))
        if ncols == 0:
            ncols = 1

        # Create 1-row table
        table = self.doc.add_table(rows=1, cols=ncols)

        # Apply table-level properties
        self._apply_table_properties(table, row_cfg, style_cfg)

        # Apply column widths
        self._apply_column_widths(table, row_cfg, style_cfg)

        # Fill cells with content
        row = table.rows[0]
        for j in range(ncols):
            cell = row.cells[j]
            cell_cfg = cells_cfg[j] if j < len(cells_cfg) else {"content": [{"text": ""}]}
            self._fill_cell(cell, cell_cfg, style_cfg)

        # Apply borders (none by default)
        self._apply_row_borders(table, style_cfg)

        return table

    # ---------------------------------------------------------------------
    # Table-level properties
    # ---------------------------------------------------------------------
    def _apply_table_properties(
        self,
        table,
        row_cfg: Dict[str, Any],
        style_cfg: Dict[str, Any],
    ):
        """
        Apply table-level properties: alignment, indent, spacing.
        """
        # Table alignment
        alignment = row_cfg.get("alignment", "left")
        if alignment == "left":
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
        elif alignment == "center":
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
        elif alignment == "right":
            table.alignment = WD_TABLE_ALIGNMENT.RIGHT

        # Table indent (from left margin)
        indent = row_cfg.get("indent", 0.0)
        if indent > 0:
            self._apply_table_indent(table, indent)

        # Disable autofit (we want fixed widths)
        try:
            table.autofit = False
        except Exception:
            pass

        # Row spacing (space before/after)
        row_style = style_cfg.get("row", {})
        space_before = row_style.get("space_before", 0)
        space_after = row_style.get("space_after", 0)

        # Apply spacing to first paragraph in first cell
        # (this is a workaround for row-level spacing)
        if space_before or space_after:
            first_cell = table.rows[0].cells[0]
            if first_cell.paragraphs:
                p = first_cell.paragraphs[0]
                if space_before:
                    p.paragraph_format.space_before = Pt(space_before)
                if space_after:
                    p.paragraph_format.space_after = Pt(space_after)

        # Keep properties (keep_with_next, keep_together)
        # These are best-effort and applied to first paragraph
        keep_with_next = row_style.get("keep_with_next", False)
        keep_together = row_style.get("keep_together", False)

        if keep_with_next or keep_together:
            first_cell = table.rows[0].cells[0]
            if first_cell.paragraphs:
                p = first_cell.paragraphs[0]
                if keep_with_next:
                    p.paragraph_format.keep_with_next = True
                if keep_together:
                    p.paragraph_format.keep_together = True

    def _apply_table_indent(self, table, indent_inches: float):
        """
        Apply left indent to table using OXML.
        """
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        # Remove existing indent
        tblInd = tblPr.find(qn("w:tblInd"))
        if tblInd is not None:
            tblPr.remove(tblInd)

        # Add new indent
        tblInd = OxmlElement("w:tblInd")
        tblInd.set(qn("w:w"), str(int(indent_inches * 1440)))  # Convert to twips
        tblInd.set(qn("w:type"), "dxa")
        tblPr.append(tblInd)

    # ---------------------------------------------------------------------
    # Column widths
    # ---------------------------------------------------------------------
    def _apply_column_widths(
        self,
        table,
        row_cfg: Dict[str, Any],
        style_cfg: Dict[str, Any],
    ):
        """
        Apply column widths based on widths array and width_mode.

        Supports:
        - Percentages: [70, 30] with sum ~100
        - Ratios: [3, 1] with sum != 100
        - Auto: no widths specified
        """
        widths = row_cfg.get("widths", [])
        width_mode = row_cfg.get("width_mode", "auto")
        total_width_pct = row_cfg.get("total_width", 100)  # Percent of available width

        if not widths or width_mode == "auto":
            # Auto mode - let Word handle it
            return

        ncols = len(table.columns)

        # Parse widths into percentages
        percentages = self._parse_widths(widths, ncols, width_mode)

        # Get available width (page width - margins)
        # We'll use a heuristic: 6.5 inches for standard 8.5x11 with 1" margins
        # This should be computed from section properties, but for now use heuristic
        available_width = 6.5  # inches

        # Apply total_width_pct
        total_width = available_width * (total_width_pct / 100.0)

        # Set column widths
        for j in range(min(ncols, len(percentages))):
            pct = percentages[j]
            width_inches = total_width * (pct / 100.0)

            # Set width for all cells in column
            for cell in table.columns[j].cells:
                try:
                    cell.width = Inches(width_inches)
                except Exception:
                    # python-docx width quirks - some versions don't support this
                    pass

    def _parse_widths(
        self,
        widths: List[float],
        num_cols: int,
        width_mode: str,
    ) -> List[float]:
        """
        Parse widths into percentages.

        Args:
            widths: List of width values
            num_cols: Expected number of columns
            width_mode: "percent", "ratio", or "auto"

        Returns:
            List of percentages (sum to 100)
        """
        if not widths:
            # Equal distribution
            return [100.0 / num_cols] * num_cols

        if len(widths) != num_cols:
            # Mismatch - use equal distribution
            return [100.0 / num_cols] * num_cols

        if width_mode == "percent":
            # Already percentages
            return list(widths)

        elif width_mode == "ratio":
            # Convert ratios to percentages
            total = sum(widths)
            if total == 0:
                return [100.0 / num_cols] * num_cols
            return [(w / total * 100.0) for w in widths]

        else:
            # Auto-detect: if sum is close to 100, assume percentages
            total = sum(widths)
            if abs(total - 100) < 1:
                return list(widths)
            else:
                # Assume ratios
                return [(w / total * 100.0) for w in widths]

    # ---------------------------------------------------------------------
    # Cell filling
    # ---------------------------------------------------------------------
    def _fill_cell(
        self,
        cell,
        cell_cfg: Dict[str, Any],
        style_cfg: Dict[str, Any],
    ):
        """
        Fill cell with rich content (paragraphs).

        Args:
            cell: python-docx cell object
            cell_cfg: Cell configuration (content, align, valign, padding, etc.)
            style_cfg: Global style configuration (font, paragraph defaults)
        """
        # Apply cell vertical alignment
        valign = cell_cfg.get("valign", "top")
        self._apply_cell_vertical_alignment(cell, valign)

        # Apply cell padding
        self._apply_cell_padding(cell, cell_cfg, style_cfg)

        # Get content (list of paragraph descriptors)
        content_list = cell_cfg.get("content", [])
        if not content_list:
            content_list = [{"text": ""}]

        # Get default font and paragraph styles
        font_style = style_cfg.get("font", {})
        para_style = style_cfg.get("paragraph", {})

        # Get cell-level alignment override
        cell_align = cell_cfg.get("align", "left")

        # Clear default paragraph
        cell.text = ""

        # Add paragraphs
        for i, content in enumerate(content_list):
            if i == 0:
                # Use existing first paragraph
                p = cell.paragraphs[0]
            else:
                # Add new paragraph
                p = cell.add_paragraph()

            # Apply paragraph content
            self._apply_paragraph_content(p, content, font_style, para_style, cell_align)

    def _apply_cell_vertical_alignment(self, cell, valign: str):
        """
        Apply vertical alignment to cell.
        """
        v = valign.lower()
        if v == "top":
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        elif v in ("middle", "center"):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        elif v == "bottom":
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM

    def _apply_cell_padding(
        self,
        cell,
        cell_cfg: Dict[str, Any],
        style_cfg: Dict[str, Any],
    ):
        """
        Apply cell padding (margins).

        Cell padding can be:
        - Specified in cell_cfg["padding"]
        - Or use defaults from style_cfg["cell"]["padding"]
        """
        # Get default padding
        default_padding = style_cfg.get("cell", {}).get("padding", {})
        if isinstance(default_padding, (int, float)):
            # Single value for all sides
            default_padding = {
                "top": default_padding,
                "bottom": default_padding,
                "left": default_padding,
                "right": default_padding,
            }

        # Get cell-specific padding (overrides defaults)
        cell_padding = cell_cfg.get("padding", {})
        if isinstance(cell_padding, (int, float)):
            cell_padding = {
                "top": cell_padding,
                "bottom": cell_padding,
                "left": cell_padding,
                "right": cell_padding,
            }

        # Merge with defaults
        padding = {**default_padding, **cell_padding}

        # Default padding if nothing specified
        if not padding:
            padding = {"top": 0.05, "bottom": 0.05, "left": 0.05, "right": 0.05}

        # Apply padding using OXML
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()

        # Remove existing padding
        tcMar = tcPr.find(qn("w:tcMar"))
        if tcMar is not None:
            tcPr.remove(tcMar)

        # Add new padding
        tcMar = OxmlElement("w:tcMar")

        for side in ["top", "left", "bottom", "right"]:
            if side in padding:
                val_inches = padding[side]
                val_twips = int(val_inches * 1440)  # Convert inches to twips

                side_elem = OxmlElement(f"w:{side}")
                side_elem.set(qn("w:w"), str(val_twips))
                side_elem.set(qn("w:type"), "dxa")
                tcMar.append(side_elem)

        tcPr.append(tcMar)

    def _apply_paragraph_content(
        self,
        paragraph,
        content: Dict[str, Any],
        font_style: Dict[str, Any],
        para_style: Dict[str, Any],
        cell_align: str,
    ):
        """
        Apply content to a paragraph.

        Args:
            paragraph: python-docx paragraph object
            content: Content descriptor (text, runs, style overrides)
            font_style: Default font style
            para_style: Default paragraph style
            cell_align: Cell alignment override
        """
        # Merge paragraph styles
        merged_para_style = {**para_style, **(content.get("style", {}).get("paragraph", {}))}

        # Apply paragraph alignment
        align = content.get("align", cell_align).lower()
        if align == "left":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif align == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == "justify":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Apply paragraph spacing
        if "spacing_before" in merged_para_style:
            paragraph.paragraph_format.space_before = Pt(merged_para_style["spacing_before"])
        if "spacing_after" in merged_para_style:
            paragraph.paragraph_format.space_after = Pt(merged_para_style["spacing_after"])
        if "line_spacing" in merged_para_style:
            paragraph.paragraph_format.line_spacing = merged_para_style["line_spacing"]

        # Get text content
        text = content.get("text", "")

        # Check if content has pre-split runs
        runs = content.get("runs", [])

        if runs:
            # Use pre-split runs
            for run_data in runs:
                run_text = run_data.get("text", "")
                run = paragraph.add_run(run_text)

                # Merge font styles
                merged_font_style = {**font_style, **(run_data.get("style", {}).get("font", {}))}
                self._apply_font_to_run(run, merged_font_style)
        else:
            # Single run with text
            if text:
                run = paragraph.add_run(text)

                # Merge font styles
                merged_font_style = {**font_style, **(content.get("style", {}).get("font", {}))}
                self._apply_font_to_run(run, merged_font_style)

    def _apply_font_to_run(self, run, font_style: Dict[str, Any]):
        """
        Apply font style to a run.
        """
        if not font_style:
            return

        if "name" in font_style:
            run.font.name = font_style["name"]

        if "size" in font_style and font_style["size"]:
            run.font.size = Pt(font_style["size"])

        if "bold" in font_style:
            run.font.bold = bool(font_style["bold"])

        if "italic" in font_style:
            run.font.italic = bool(font_style["italic"])

        if "underline" in font_style:
            run.font.underline = convert_underline_value(font_style["underline"])

        if "strike" in font_style:
            run.font.strike = bool(font_style["strike"])

        if "highlight" in font_style and font_style["highlight"]:
            highlight_map = {
                "yellow": WD_COLOR_INDEX.YELLOW,
                "brightGreen": WD_COLOR_INDEX.BRIGHT_GREEN,
                "turquoise": WD_COLOR_INDEX.TURQUOISE,
                "pink": WD_COLOR_INDEX.PINK,
                "blue": WD_COLOR_INDEX.BLUE,
                "red": WD_COLOR_INDEX.RED,
                "darkBlue": WD_COLOR_INDEX.DARK_BLUE,
                "teal": WD_COLOR_INDEX.TEAL,
                "green": WD_COLOR_INDEX.GREEN,
                "violet": WD_COLOR_INDEX.VIOLET,
                "darkRed": WD_COLOR_INDEX.DARK_RED,
                "darkYellow": WD_COLOR_INDEX.DARK_YELLOW,
                "gray50": WD_COLOR_INDEX.GRAY_50,
                "gray25": WD_COLOR_INDEX.GRAY_25,
            }
            highlight_val = font_style["highlight"]
            if highlight_val in highlight_map:
                run.font.highlight_color = highlight_map[highlight_val]

        if "color" in font_style and font_style["color"]:
            color_val = font_style["color"]
            if isinstance(color_val, str):
                # accept "FF0000" or "#FF0000"
                hex_str = color_val.lstrip("#")
                run.font.color.rgb = RGBColor.from_string(hex_str)
            elif isinstance(color_val, RGBColor):
                run.font.color.rgb = color_val

        if "all_caps" in font_style:
            run.font.all_caps = bool(font_style["all_caps"])

        if "small_caps" in font_style:
            run.font.small_caps = bool(font_style["small_caps"])

    # ---------------------------------------------------------------------
    # Borders
    # ---------------------------------------------------------------------
    def _apply_row_borders(self, table, style_cfg: Dict[str, Any]):
        """
        Apply borders to row table.

        For row layouts, borders are "none" by default (invisible table).
        Support "debug" mode for development.
        """
        borders_mode = style_cfg.get("cell", {}).get("borders", "none")

        if borders_mode == "none":
            # Remove all borders (make invisible table)
            self._remove_all_borders(table)

        elif borders_mode == "debug":
            # Thin borders for debugging
            self._apply_debug_borders(table)

    def _remove_all_borders(self, table):
        """
        Remove all table and cell borders to create invisible table.
        """
        # Remove table-level borders
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is not None:
            tblBorders = tblPr.find(qn("w:tblBorders"))
            if tblBorders is not None:
                tblPr.remove(tblBorders)

        # Remove cell-level borders
        for row in table.rows:
            for cell in row.cells:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()

                # Remove existing borders
                tcBorders = tcPr.find(qn("w:tcBorders"))
                if tcBorders is not None:
                    tcPr.remove(tcBorders)

                # Add nil borders
                tcBorders = OxmlElement("w:tcBorders")
                for border_name in ["top", "left", "bottom", "right"]:
                    border = OxmlElement(f"w:{border_name}")
                    border.set(qn("w:val"), "nil")
                    tcBorders.append(border)

                tcPr.append(tcBorders)

    def _apply_debug_borders(self, table):
        """
        Apply thin borders for debugging row layouts.
        """
        tbl = table._tbl
        tblPr = tbl.tblPr

        # Ensure tblBorders exists
        tblBorders = tblPr.find(qn("w:tblBorders"))
        if tblBorders is None:
            tblBorders = OxmlElement("w:tblBorders")
            tblPr.append(tblBorders)

        # Add thin borders
        border_types = ["top", "left", "bottom", "right", "insideH", "insideV"]
        for bt in border_types:
            el = tblBorders.find(qn(f"w:{bt}"))
            if el is None:
                el = OxmlElement(f"w:{bt}")
                tblBorders.append(el)

            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")  # Thin border
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "CCCCCC")  # Light gray
