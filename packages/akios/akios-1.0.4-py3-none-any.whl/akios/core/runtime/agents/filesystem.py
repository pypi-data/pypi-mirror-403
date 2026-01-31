# Copyright (C) 2025-2026 AKIOUD AI, SAS <contact@akioud.ai>
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

"""
Filesystem Agent - Filesystem access agent with path whitelisting + read-only default

Provides secure file system operations with strict access controls.
"""

import os
import stat
from pathlib import Path
from typing import Dict, Any, Optional, List

from .base import BaseAgent, AgentError
from akios.config import get_settings  # For test patching

# Document processing libraries
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    from pdfminer.high_level import extract_text as pdfminer_extract
    HAS_PDFMINER = True
except ImportError:
    HAS_PDFMINER = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Delay security imports to avoid validation during package import
from akios.core.audit import append_audit_event

# Pre-import PII redaction to avoid import hangs during agent execution
try:
    from akios.security.pii import apply_pii_redaction as _pii_redaction_func
except Exception:
    # Fallback if PII import fails
    _pii_redaction_func = lambda x: x


class FilesystemAgent(BaseAgent):
    """
    Filesystem agent for secure file operations.

    Enforces path whitelisting and read-only defaults for security.
    """

    def __init__(self, allowed_paths: Optional[List[str]] = None, read_only: bool = True, **kwargs):
        super().__init__(**kwargs)
        default_paths = self._get_default_allowed_paths()
        if allowed_paths:
            default_paths.extend(allowed_paths)
        self.allowed_paths = default_paths
        self.read_only = read_only

    def _get_default_allowed_paths(self) -> List[str]:
        """Get default allowed paths"""
        # In a real implementation, this would be configurable
        # For security, only allow specific directories
        return [
            "./templates",
            "./data/input",
            "./data/output"
        ]

    def execute(self, action: str, parameters: Dict[str, Any]) -> Dict[str, Any]:
        """
        Execute filesystem action.

        Args:
            action: Action to perform ('read', 'write', 'list', 'exists', 'stat')
            parameters: Action parameters

        Returns:
            Action result
        """
        self.validate_parameters(action, parameters)

        # Apply security before file operations (delayed import)
        from akios.security import enforce_sandbox
        apply_pii_redaction = _pii_redaction_func
        enforce_sandbox()

        # Validate path security first to get the path variable
        path = parameters.get('path', '')
        self._validate_path_security(path, action)

        # Apply PII redaction to inputs
        if 'content' in parameters and isinstance(parameters['content'], str):
            original_content = parameters['content']
            
            # Apply PII redaction with timeout to prevent hangs
            try:
                import signal
                
                def timeout_handler(signum, frame):
                    raise TimeoutError("PII redaction timed out")
                
                signal.signal(signal.SIGALRM, timeout_handler)
                signal.alarm(5)  # 5 second timeout for PII redaction
                
                try:
                    parameters['content'] = apply_pii_redaction(parameters['content'])
                finally:
                    signal.alarm(0)  # Cancel the alarm
                    
            except TimeoutError:
                # If PII redaction times out, keep original content
                import sys
                print(f"Warning: PII redaction timed out for input, keeping original content", file=sys.stderr)
                parameters['content'] = original_content
            
            if parameters['content'] != original_content:
                # Log PII redaction event
                append_audit_event({
                    'workflow_id': parameters.get('workflow_id', 'unknown'),
                    'step': parameters.get('step', 0),
                    'agent': 'filesystem',
                    'action': 'pii_redaction_input',
                    'result': 'success',
                    'metadata': {
                        'path': self._sanitize_path(path),
                        'redacted_length': len(parameters['content'])
                    }
                })

        # Execute the action
        result = self._execute_action(action, parameters)

        # Audit the file operation
        append_audit_event({
            'workflow_id': parameters.get('workflow_id', 'unknown'),
            'step': parameters.get('step', 0),
            'agent': 'filesystem',
            'action': action,
            'result': 'success',
            'metadata': {
                'path': self._sanitize_path(path),
                'operation': action,
                'read_only_mode': self.read_only,
                'path_allowed': True
            }
        })

        # Apply PII redaction to outputs if they contain sensitive content
        if 'content' in result and isinstance(result['content'], str):
            original_content = result['content']
            
            # Apply PII redaction with timeout to prevent hangs
            try:
                import signal
                
                def timeout_handler(signum, frame):
                    raise TimeoutError("PII redaction timed out")
                
                signal.signal(signal.SIGALRM, timeout_handler)
                signal.alarm(5)  # 5 second timeout for PII redaction
                
                try:
                    result['content'] = apply_pii_redaction(result['content'])
                finally:
                    signal.alarm(0)  # Cancel the alarm
                    
            except TimeoutError:
                # If PII redaction times out, keep original content
                import sys
                print(f"Warning: PII redaction timed out for output, keeping original content", file=sys.stderr)
                result['content'] = original_content
            
            if result['content'] != original_content:
                # Log PII redaction event
                append_audit_event({
                    'workflow_id': parameters.get('workflow_id', 'unknown'),
                    'step': parameters.get('step', 0),
                    'agent': 'filesystem',
                    'action': 'pii_output_redaction',
                    'result': 'success',
                    'metadata': {
                        'path': self._sanitize_path(path),
                        'redacted_length': len(result['content'])
                    }
                })

        return result

    def _execute_action(self, action: str, parameters: Dict[str, Any]) -> Dict[str, Any]:
        """Execute the specific filesystem action"""
        path = Path(parameters['path'])

        if action == 'read':
            return self._read_file(path)
        elif action == 'write':
            return self._write_file(path, parameters)
        elif action == 'list':
            return self._list_directory(path)
        elif action == 'exists':
            return {'exists': path.exists(), 'path': str(path)}
        elif action == 'stat':
            return self._get_file_stats(path)
        elif action == 'analyze':
            return self._analyze_file(path)
        else:
            raise AgentError(f"Unsupported filesystem action: {action}")

    def _read_file(self, path: Path) -> Dict[str, Any]:
        """Read file contents with multi-format support"""
        try:
            if not path.exists():
                raise AgentError(f"File does not exist: {path}")

            if not path.is_file():
                raise AgentError(f"Path is not a file: {path}")

            # Detect file type and read accordingly
            file_type = self._detect_file_type(path)

            if file_type == 'pdf':
                return self._read_pdf_file(path)
            elif file_type == 'docx':
                return self._read_docx_file(path)
            else:
                return self._read_text_file(path)

        except Exception as e:
            raise AgentError(f"Failed to read file {path}: {e}")

    def _detect_file_type(self, path: Path) -> str:
        """Detect file type based on extension and content"""
        suffix = path.suffix.lower()

        # Check by extension first
        if suffix == '.pdf':
            return 'pdf'
        elif suffix == '.docx':
            return 'docx'
        elif suffix == '.txt' or not suffix:
            return 'text'
        else:
            # Unknown extension, try to detect by content
            try:
                with open(path, 'rb') as f:
                    header = f.read(8)

                # PDF files start with %PDF-
                if header.startswith(b'%PDF-'):
                    return 'pdf'
                # DOCX files are ZIP archives with specific structure
                elif header.startswith(b'PK\x03\x04'):
                    return 'docx'
                else:
                    return 'text'  # Default to text
            except:
                return 'text'  # Default to text on error

    def _read_text_file(self, path: Path) -> Dict[str, Any]:
        """Read plain text file"""
        try:
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()

            # PII redaction is now handled at the agent level in execute()
            # to avoid double processing and timeout issues

            return {
                'content': content,
                'size': len(content),
                'path': str(path),
                'encoding': 'utf-8',
                'file_type': 'text'
            }
        except UnicodeDecodeError:
            raise AgentError(f"File is not a valid text file: {path}")

    def _read_pdf_file(self, path: Path) -> Dict[str, Any]:
        """Extract text from PDF file using robust multi-library approach"""
        # Check if PDF libraries are available
        if not HAS_PDFMINER and not HAS_PYPDF2:
            # Graceful fallback: return placeholder indicating PDF support not available
            return {
                'content': f'[PDF FILE - Text extraction not available. Install pdfminer or PyPDF2 for PDF text extraction. File: {path}]',
                'size': len('[PDF FILE - Text extraction not available]'),
                'path': str(path),
                'encoding': 'utf-8',
                'file_type': 'pdf',
                'extraction_status': 'library_not_available',
                'supported_libraries': ['pdfminer', 'PyPDF2']
            }

        extraction_errors = []

        # Try pdfminer first (more robust for text extraction)
        if HAS_PDFMINER:
            try:
                content = pdfminer_extract(str(path))
                if content and content.strip():
                    return {
                        'content': content,
                        'size': len(content),
                        'path': str(path),
                        'encoding': 'utf-8',
                        'file_type': 'pdf',
                        'extraction_method': 'pdfminer'
                    }
            except Exception as e:
                extraction_errors.append(f"pdfminer: {e}")

        # Fallback to PyPDF2
        if HAS_PYPDF2:
            try:
                with open(path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text_content = []

                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text and text.strip():
                            text_content.append(text)

                    content = '\n\n'.join(text_content)

                    if content and content.strip():
                        return {
                            'content': content,
                            'size': len(content),
                            'path': str(path),
                            'encoding': 'utf-8',
                            'file_type': 'pdf',
                            'pages': len(pdf_reader.pages),
                            'extraction_method': 'PyPDF2'
                        }
            except Exception as e:
                extraction_errors.append(f"PyPDF2: {e}")

        # If extraction failed but libraries are available, return informative placeholder
        # This ensures workflows don't fail just because PDF text extraction failed
        error_info = f"PDF text extraction failed. Errors: {'; '.join(extraction_errors) if extraction_errors else 'Unknown error'}"
        return {
            'content': f'[PDF FILE - {error_info}. File: {path}]',
            'size': len(f'[PDF FILE - {error_info}]'),
            'path': str(path),
            'encoding': 'utf-8',
            'file_type': 'pdf',
            'extraction_status': 'failed',
            'extraction_errors': extraction_errors
        }

    def _read_docx_file(self, path: Path) -> Dict[str, Any]:
        """Extract text from DOCX file with graceful fallback"""
        if not HAS_DOCX:
            # Graceful fallback: return placeholder indicating DOCX support not available
            return {
                'content': f'[DOCX FILE - Text extraction not available. Install python-docx for DOCX text extraction. File: {path}]',
                'size': len('[DOCX FILE - Text extraction not available]'),
                'path': str(path),
                'encoding': 'utf-8',
                'file_type': 'docx',
                'extraction_status': 'library_not_available',
                'supported_libraries': ['python-docx']
            }

        try:
            doc = Document(path)
            text_content = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)

            content = '\n\n'.join(text_content)

            return {
                'content': content,
                'size': len(content),
                'path': str(path),
                'encoding': 'utf-8',
                'file_type': 'docx'
            }
        except Exception as e:
            # Graceful fallback on extraction failure
            error_info = f"DOCX text extraction failed: {str(e)}"
            return {
                'content': f'[DOCX FILE - {error_info}. File: {path}]',
                'size': len(f'[DOCX FILE - {error_info}]'),
                'path': str(path),
                'encoding': 'utf-8',
                'file_type': 'docx',
                'extraction_status': 'failed',
                'extraction_error': str(e)
            }

    def _write_file(self, path: Path, parameters: Dict[str, Any]) -> Dict[str, Any]:
        """Write content to file"""
        if self.read_only:
            raise AgentError("Filesystem agent is in read-only mode")

        content = parameters.get('content', '')
        append = parameters.get('append', False)
        mode = 'a' if append else parameters.get('mode', 'w')  # 'a' for append, 'w' for write

        # Validate file size limit
        content_size_mb = len(content.encode('utf-8')) / (1024 * 1024)  # Size in MB
        if content_size_mb > self.settings.max_file_size_mb:
            raise AgentError(
                f"File size {content_size_mb:.1f}MB exceeds limit of {self.settings.max_file_size_mb}MB"
            )

        # Check available disk space for write operations (optional safety check)
        if not append:  # Only check for new writes, not appends to existing files
            try:
                import psutil
                # Use the base directory for disk space check, not the potentially non-existent parent
                base_dir = path.parent
                while not base_dir.exists() and base_dir.parent != base_dir:
                    base_dir = base_dir.parent
                disk_usage = psutil.disk_usage(base_dir)
                available_mb = disk_usage.free / (1024 * 1024 * 1024)  # Convert to GB for readability
                required_mb = content_size_mb
                if required_mb > available_mb * 1024:  # Convert GB to MB for comparison
                    raise AgentError(
                        f"Insufficient disk space. Required: {required_mb:.1f}MB, "
                        f"Available: {available_mb:.1f}GB"
                    )
            except ImportError:
                # psutil not available, skip disk space check
                pass
            except Exception:
                # If disk space check fails for any reason, skip it
                pass

        try:
            # Ensure parent directory exists
            path.parent.mkdir(parents=True, exist_ok=True)

            with open(path, mode, encoding='utf-8') as f:
                f.write(content)

            return {
                'written': True,
                'path': str(path),
                'size': len(content),
                'mode': mode
            }

        except (OSError, IOError) as e:
            raise AgentError(f"Failed to write file {path}: {e}") from e

    def _list_directory(self, path: Path) -> Dict[str, Any]:
        """List directory contents"""
        try:
            if not path.exists():
                raise AgentError(f"Directory does not exist: {path}")

            if not path.is_dir():
                raise AgentError(f"Path is not a directory: {path}")

            items = []
            for item in path.iterdir():
                items.append({
                    'name': item.name,
                    'path': str(item),
                    'type': 'directory' if item.is_dir() else 'file',
                    'size': item.stat().st_size if item.is_file() else 0
                })

            return {
                'items': items,
                'count': len(items),
                'path': str(path)
            }

        except (OSError, IOError) as e:
            raise AgentError(f"Failed to list directory {path}: {e}") from e

    def _get_file_stats(self, path: Path) -> Dict[str, Any]:
        """Get file/directory statistics"""
        try:
            if not path.exists():
                raise AgentError(f"Path does not exist: {path}")

            stat_info = path.stat()

            return {
                'path': str(path),
                'exists': True,
                'type': 'directory' if path.is_dir() else 'file',
                'size': stat_info.st_size,
                'modified': stat_info.st_mtime,
                'permissions': oct(stat_info.st_mode)[-3:],
                'readable': os.access(path, os.R_OK),
                'writable': os.access(path, os.W_OK)
            }

        except (OSError, IOError) as e:
            raise AgentError(f"Failed to get stats for {path}: {e}") from e

    def _analyze_file(self, path: Path) -> Dict[str, Any]:
        """Comprehensive file analysis including patterns, entropy, and hashes"""
        try:
            if not path.exists():
                raise AgentError(f"File does not exist: {path}")

            if not path.is_file():
                raise AgentError(f"Path is not a file: {path}")

            # Read file content for analysis
            content = self._read_file_for_analysis(path)

            # Perform all analysis types
            patterns = self._analyze_patterns(content)
            entropy = self._calculate_entropy(content)
            hashes = self._calculate_hash(path)

            # Get basic file stats
            stats = self._get_file_stats(path)

            return {
                'path': str(path),
                'patterns': patterns,
                'entropy': entropy,
                'hashes': hashes,
                'stats': stats,
                'file_type': self._detect_file_type(path)
            }

        except Exception as e:
            raise AgentError(f"Failed to analyze file {path}: {e}")

    def _read_file_for_analysis(self, path: Path) -> str:
        """Read file content specifically for analysis (text extraction)"""
        file_type = self._detect_file_type(path)

        if file_type == 'pdf':
            pdf_result = self._read_pdf_file(path)
            return pdf_result.get('content', '')
        elif file_type == 'docx':
            docx_result = self._read_docx_file(path)
            return docx_result.get('content', '')
        else:
            # For text and unknown files, try to read as text
            text_result = self._read_text_file(path)
            return text_result.get('content', '')

    def _analyze_patterns(self, content: str) -> Dict[str, Any]:
        """Extract security-relevant patterns from content"""
        import re

        patterns = {
            'emails': re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', content),
            'ips': re.findall(r'\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b', content),
            'urls': re.findall(r'https?://[^\s<>"{}|\\^`\[\]]+', content),
            'credit_cards': re.findall(r'\b(?:\d{4}[ -]?){3}\d{4}\b', content),
            'phone_numbers': re.findall(r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b', content),
            'numbers': re.findall(r'\b\d+\.?\d*\b', content),
            'dates': re.findall(r'\b\d{4}-\d{2}-\d{2}\b', content)
        }

        # Count occurrences
        pattern_counts = {key: len(value) for key, value in patterns.items()}

        return {
            'patterns': patterns,
            'counts': pattern_counts
        }

    def _calculate_entropy(self, content: str) -> float:
        """Calculate Shannon entropy for file content"""
        import math

        if not content:
            return 0.0

        entropy = 0
        content_length = len(content)

        # Count character frequencies
        char_counts = {}
        for char in content:
            char_counts[char] = char_counts.get(char, 0) + 1

        # Calculate entropy
        for count in char_counts.values():
            probability = count / content_length
            entropy -= probability * math.log2(probability)

        return round(entropy, 2)

    def _calculate_hash(self, path: Path) -> Dict[str, str]:
        """Calculate file hashes"""
        import hashlib

        hash_md5 = hashlib.md5()
        hash_sha256 = hashlib.sha256()

        with open(path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
                hash_sha256.update(chunk)

        return {
            'md5': hash_md5.hexdigest(),
            'sha256': hash_sha256.hexdigest()
        }

    def _validate_path_security(self, path_str: str, action: str) -> None:
        """Validate that the path is allowed and secure"""
        # Convert Path objects to strings
        if isinstance(path_str, Path):
            path_str = str(path_str)

        # Security checks that apply to all paths - check before path resolution
        if '..' in path_str:
            raise AgentError("Path traversal not allowed")

        # Special handling for timestamped output paths
        if action == 'write' and 'run_' in path_str and 'data/output' in path_str:
            # For timestamped paths like "data/output/run_20260107_132409/file.txt",
            # validate against the base "data/output" directory only
            base_output_path = "data/output"
            resolved_path = self._safe_absolute_path(base_output_path)

            # Check if base output directory is allowed
            allowed = False
            for allowed_path in self.allowed_paths:
                allowed_full_path = self._safe_absolute_path(allowed_path)
                if resolved_path.startswith(allowed_full_path):
                    allowed = True
                    break

            if not allowed:
                raise AgentError(f"Base output path not in allowed list: {resolved_path}")
        else:
            # Standard validation for non-timestamped paths
            path = Path(path_str)
            
            # Use safe absolute path instead of resolve() to avoid hangs in cgroups
            try:
                resolved_path = self._safe_absolute_path(str(path))
            except FileNotFoundError:
                # For read operations, path should exist
                if action in ['read', 'list', 'stat', 'analyze']:
                    raise AgentError(f"Path does not exist: {path_str}")
                # For write operations, parent directory should exist
                try:
                    resolved_path = self._safe_absolute_path(str(path.parent)) + "/" + path.name
                except FileNotFoundError:
                    raise AgentError(f"Parent directory does not exist: {path_str}")

            # Check against allowed paths
            allowed = False
            for allowed_path in self.allowed_paths:
                allowed_full_path = self._safe_absolute_path(allowed_path)
                if resolved_path.startswith(allowed_full_path):
                    allowed = True
                    break

            if not allowed:
                raise AgentError(f"Path not in allowed list: {resolved_path}")

        if action in ['write'] and self.read_only:
            raise AgentError("Write operations not allowed in read-only mode")

        # Prevent access to sensitive system files
        sensitive_paths = ['/etc/passwd', '/etc/shadow', '/proc', '/sys', '/dev']
        for sensitive in sensitive_paths:
            if resolved_path.startswith(sensitive):
                raise AgentError(f"Access to sensitive system path blocked: {sensitive}")

    def _safe_absolute_path(self, path_str: str) -> str:
        """Get absolute path without resolving symlinks to avoid hangs in cgroups"""
        path = Path(path_str)
        
        # Use absolute() instead of resolve() to avoid symlink resolution
        # This prevents hangs in restricted filesystem environments
        try:
            abs_path = str(path.absolute())
        except Exception:
            # Fallback: just return the original path
            abs_path = str(path)
        
        return abs_path

    def _sanitize_path(self, path_str: str) -> str:
        """Sanitize path for audit logging"""
        # Remove sensitive information from paths
        # For now, just return as-is since filesystem paths are generally safe to log
        return path_str

    def validate_parameters(self, action: str, parameters: Dict[str, Any]) -> None:
        """Validate action parameters"""
        if 'path' not in parameters:
            raise AgentError("Filesystem operations require 'path' parameter")

        path = parameters['path']
        if not isinstance(path, str):
            raise AgentError("'path' must be a string")

        # Check read-only mode for write operations
        if action == 'write' and self.read_only:
            raise AgentError("Write operations not allowed in read-only mode")

        # Additional validation based on action
        if action == 'write' and 'content' not in parameters:
            raise AgentError("Write action requires 'content' parameter")

        # Analyze action validation
        if action == 'analyze':
            # Analyze requires the same path validation as read
            pass

    def get_supported_actions(self) -> List[str]:
        """Get supported filesystem actions"""
        actions = ['read', 'list', 'exists', 'stat', 'analyze']
        if not self.read_only:
            actions.append('write')
        return actions

    def add_allowed_path(self, path: str) -> None:
        """Add an allowed path (admin operation)"""
        self.allowed_paths.append(path)

    def set_read_only(self, read_only: bool) -> None:
        """Set read-only mode"""
        self.read_only = read_only
