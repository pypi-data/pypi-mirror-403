import os
from typing import List

import xlrd
import xlwt
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Pt, RGBColor
from docxtpl import DocxTemplate, RichText
from openpyxl import load_workbook
from openpyxl.styles import Side, Border
from xlstpl import writer as xls_writer
from xlsxtpl import writerx as xlsx_writerx
from xlutils.copy import copy

from ctools import cjson
from ctools.office.word_fill_entity import WordTable, WordCell, Style, Font, Alignment

EXCEL_WORD_ALIGNMENT = {"left": WD_TABLE_ALIGNMENT.LEFT,
                        "center": WD_TABLE_ALIGNMENT.CENTER,
                        "right": WD_TABLE_ALIGNMENT.RIGHT}
EXCEL_WORD_VERTICAL = {"top": WD_ALIGN_VERTICAL.TOP,
                       "center": WD_ALIGN_VERTICAL.CENTER,
                       "bottom": WD_ALIGN_VERTICAL.BOTTOM}


class DocxTools:

  @staticmethod
  def get_docx(path=None):
    return Document(path)

  @staticmethod
  def add_rich_text(wildcard: dict):
    try:
      new_wildcard = {}
      if wildcard:
        for k, v in wildcard.items():
          if "_color" in k:
            n_k = k[:k.index('_color')]
            new_wildcard[n_k] = RichText(wildcard[n_k], color=v)
          elif k not in new_wildcard:
            new_wildcard[k] = v
      return new_wildcard
    except Exception as e:
      print("使用word_fill.add_rich_text函数异常: %s" % e)
    return wildcard

  def tables_format(self, input_path: str, table_data: List[dict] = None):
    if table_data:
      doc = Document(input_path)
      for index, table in enumerate(doc.tables):
        if index <= len(table_data) - 1:
          td = table_data[index]
          start_position = td.get('start_position')
          data = td.get('data')

          for r in range(0, len(data)):
            for c in range(0, len(data[r])):
              row = r + start_position[0]
              col = c + start_position[1]

              for i in range((row - len(table.rows)) + 1):
                table.add_row()
              try:
                table.cell(row, col).text = str(data[r][c])
              except IndexError:
                pass

          self.set_cell_border(table)
      doc.save(input_path)

  def set_cell_border(self, table):
    return self.set_table_boarder(
      table,
      top={"sz": 4, "val": "single", "color": "#000000"},
      bottom={"sz": 4, "val": "single", "color": "#000000"},
      left={"sz": 4, "val": "single", "color": "#000000"},
      right={"sz": 4, "val": "single", "color": "#000000"},
      insideV={"sz": 4, "val": "single", "color": "#000000"},
      insideH={"sz": 4, "val": "single", "color": "#000000"}
    )

  @staticmethod
  def set_table_boarder(table, **kwargs):
    """
    Set table`s border
    Usage:
    set_table_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    borders = OxmlElement('w:tblBorders')
    for tag in ('bottom', 'top', 'left', 'right', 'insideV', 'insideH'):
      edge_data = kwargs.get(tag)
      if edge_data:
        any_border = OxmlElement(f'w:{tag}')
        for key in ["sz", "val", "color", "space", "shadow"]:
          if key in edge_data:
            any_border.set(qn(f'w:{key}'), str(edge_data[key]))
        borders.append(any_border)
        table._tbl.tblPr.append(borders)

  @staticmethod
  def merge_word(document: Document, input_file):
    """
    合并docx文件内容
    :param document:
    :param input_file:
    :return:
    """
    doc = Document(input_file)
    for element in doc.element.body:
      document.element.body.append(element)

  @staticmethod
  def merge_excel(document: Document, excel_path):
    """
    合并excel文件内容
    :param document:
    :param excel_path: excel文件路径
    :return:
    """
    wt = ExcelTools().get_excel_cells(excel_path)
    max_rows = wt.max_rows
    max_cols = wt.max_cols
    cells = wt.cells
    merged_cells = wt.merged_cells

    doc_body = document.element.body
    # 创建Word表格
    table = document.add_table(rows=max_rows, cols=max_cols, style="Table Grid")
    if len(doc_body) > 0:
      doc_body[-1].addnext(table._tbl)
    table.autofit = False

    for cell in cells:
      position = cell.position
      value = cell.value
      style = cell.style

      t_cell = table.cell(position[0], position[1])
      t_cell.text = str(value)
      paragraph = t_cell.paragraphs[0]
      run = paragraph.runs[0]

      if style:
        if style.font:
          run.font.name = style.font.name
          run.font.size = Pt(style.font.size)
          run.font.bold = style.font.bold
          run.font.italic = style.font.italic
          run.font.underline = style.font.underline
          run._element.rPr.rFonts.set(qn('w:eastAsia'), style.font.name)
          if style.font.color:
            run.font.color.rgb = RGBColor.from_string(style.font.color)

        if style.bg_color:
          t_cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=style.bg_color)))

        paragraph.paragraph_format.alignment = EXCEL_WORD_ALIGNMENT.get(style.alignment.horizontal)
        t_cell.vertical_alignment = EXCEL_WORD_VERTICAL.get(style.alignment.vertical)

    # 合并单元格
    for mc in merged_cells:
      first_cell = None
      first_group_cell = []
      for p in mc:
        for i, r in enumerate(p):
          if i == 0:
            first_cell = table.cell(r[0], r[1])
            first_group_cell.append(first_cell)
          else:
            first_cell.merge(table.cell(r[0], r[1]))

      if len(first_group_cell) > 1:
        for i, cell in enumerate(first_group_cell):
          if i == 0:
            first_cell = cell
          else:
            first_cell.merge(cell)


class ExcelTools:

  @staticmethod
  def tables_format_xls(input_path: str, table_data: List[dict] = None):
    if table_data:
      wb = xlrd.open_workbook(input_path, formatting_info=True)
      wb = copy(wb)

      borders = xlwt.Borders()
      borders.left = 1
      borders.right = 1
      borders.top = 1
      borders.bottom = 1

      style = xlwt.XFStyle()
      style.borders = borders

      sheet = wb.get_sheet(0)
      for td in table_data:
        start_position = td.get('start_position')
        data = td.get('data')

        for r in range(0, len(data)):
          for c in range(0, len(data[r])):
            row = r + start_position[0]
            col = c + start_position[1]
            sheet.write(row, col, data[r][c], style=style)

      wb.save(input_path)

  @staticmethod
  def tables_format_xlsx(input_path: str, table_data: List[dict] = None):
    if table_data:
      wb = load_workbook(input_path)
      sheet = wb[wb.sheetnames[0]]

      side = Side(style='thin', color='FF000000')
      border = Border(left=side, right=side, top=side, bottom=side)

      for td in table_data:
        start_position = td.get('start_position')
        data = td.get('data')

        for r in range(0, len(data)):
          for c in range(0, len(data[r])):
            row = r + start_position[0] + 1
            col = c + start_position[1] + 1
            cell = sheet.cell(row=row, column=col)
            cell.value = data[r][c]
            cell.border = border

      wb.save(filename=input_path)

  def get_excel_cells(self, excel_path) -> WordTable:
    """
    获取excel所有有效单元格信息
    :param excel_path:
    :return:
    """
    cells = []
    merged_cells = []
    eff_max_rows, eff_max_cols, real_max_rows, real_max_cols = 0, 0, 0, 0
    suffix = os.path.splitext(excel_path)[-1].lower()
    if suffix == ".xls":
      wb = xlrd.open_workbook(excel_path, formatting_info=True)
      sheet = wb.sheets()[0]

      eff_max_rows, eff_min_rows, eff_max_cols, eff_min_cols, real_max_rows, real_max_cols = self.get_excel_rows_cols(sheet, is_xlsx=False)

      for r in range(0, real_max_rows):
        for c in range(0, real_max_cols):
          val = sheet.cell_value(r, c)
          if val and len(str(val).strip()) > 0:
            position = [r - eff_min_rows, c - eff_min_cols]
            xfx = sheet.cell_xf_index(position[0], position[1])
            xf = wb.xf_list[xfx]

            horizontal = None
            vertical = None
            hor_align = xf.alignment.hor_align
            if hor_align == 1:
              horizontal = "left"
            elif hor_align == 2:
              horizontal = "center"
            elif hor_align == 3:
              horizontal = "right"

            vert_align = xf.alignment.vert_align
            if vert_align == 0:
              vertical = "top"
            elif vert_align == 1:
              vertical = "center"
            elif vert_align == 2:
              vertical = "bottom"

            font = wb.font_list[xf.font_index]

            style = Style()
            font_size = round(font.height / 20)
            if font_size:
              if eff_max_cols < 20:
                font_size = 5 if round(font_size * 0.8) <= 5 else round(font_size * 0.8)
              else:
                font_size = 5 if round(font_size * 0.7) <= 5 else round(font_size * 0.7)
            style.font = Font(font.name, font_size, bool(font.bold), bool(font.italic), bool(font.underlined))
            # style.font.set_color(cell.font.color)
            style.alignment = Alignment(horizontal, vertical)
            # style.set_bg_color(cell.fill.start_color)

            word_cell = WordCell()
            word_cell.position = position
            word_cell.value = str(val)
            word_cell.style = style
            cells.append(word_cell)

      merged_cells = self.get_excel_merged_cells(sheet, eff_min_rows, eff_min_cols, is_xlsx=False)

    elif suffix == ".xlsx":
      wb = load_workbook(excel_path)
      sheet = wb[wb.sheetnames[0]]
      eff_max_rows, eff_min_rows, eff_max_cols, eff_min_cols, real_max_rows, real_max_cols = self.get_excel_rows_cols(sheet)

      for r in range(1, real_max_rows + 1):
        for c in range(1, real_max_cols + 1):
          cell = sheet.cell(r, c)
          val = cell.value
          if val and len(str(val).strip()) > 0:
            position = [r - eff_min_rows, c - eff_min_cols]

            style = Style()
            font_size = cell.font.size
            if font_size:
              if eff_max_cols < 20:
                font_size = 5 if round(font_size * 0.8) <= 5 else round(font_size * 0.8)
              else:
                font_size = 5 if round(font_size * 0.7) <= 5 else round(font_size * 0.7)

            style.font = Font(cell.font.name, font_size, cell.font.bold, cell.font.italic, cell.font.underline)
            style.font.set_color(cell.font.color)
            vertical = cell.alignment.vertical
            if cell.alignment.vertical is None:
              vertical = "bottom"
            style.alignment = Alignment(cell.alignment.horizontal, vertical)
            style.set_bg_color(cell.fill.start_color)

            word_cell = WordCell()
            word_cell.position = position
            word_cell.value = str(val)
            word_cell.style = style
            cells.append(word_cell)

      merged_cells = self.get_excel_merged_cells(sheet, eff_min_rows, eff_min_cols)

    word_table = WordTable()
    word_table.max_rows = eff_max_rows
    word_table.max_cols = eff_max_cols
    word_table.cells = cells
    word_table.merged_cells = merged_cells
    return word_table

  @staticmethod
  def get_excel_rows_cols(sheet, is_xlsx=True):
    """
    获取excel有效的最大行数和列数
    :param sheet:
    :param is_xlsx:
    :return:
    """
    row_list, col_list = [], []
    if is_xlsx:
      start_index = 1
      real_max_rows, real_max_cols = sheet.max_row, sheet.max_column
      real_max_cols = sheet.max_column
    else:
      start_index = 0
      real_max_rows, real_max_cols = sheet.nrows, sheet.ncols

    index = 0
    eff_max_rows, eff_max_cols = real_max_rows, real_max_cols
    eff_min_rows, eff_min_cols = start_index, start_index
    for r in range(start_index, real_max_rows):
      for c in range(start_index, real_max_cols):
        val = sheet.cell(r, c).value if is_xlsx else sheet.cell_value(r, c)
        if val and len(str(val).strip()) > 0:
          row_list.append(r)
          col_list.append(c)
          if index == 0:
            eff_min_rows = r
            eff_min_cols = c
            index = 1
    if len(row_list) > 0:
      eff_max_rows = max(row_list) - min(row_list) + start_index + 1
    if len(col_list) > 0:
      eff_max_cols = max(col_list) - min(col_list) + start_index + 1
    return eff_max_rows, eff_min_rows, eff_max_cols, eff_min_cols, real_max_rows, real_max_cols

  @staticmethod
  def get_excel_merged_cells(sheet, d_row, d_col, is_xlsx=True):
    """
    获取excel所有合并的单元格信息
    :param sheet:
    :param d_row: 有效起始行和实际起始行的差值
    :param d_col:
    :param is_xlsx:
    :return:
    """
    merged_cells = []
    if is_xlsx:
      merged_ranges = sheet.merged_cells.ranges
      for merged_range in merged_ranges:
        cells = []
        for r in range(merged_range.min_row, merged_range.max_row + 1):
          rows = []
          for c in range(merged_range.min_col, merged_range.max_col + 1):
            rows.append([r - d_row, c - d_col])
          cells.append(rows)
        merged_cells.append(cells)
    else:
      mergeds = sheet.merged_cells
      for merged in mergeds:
        cells = []
        for r in range(merged[0], merged[1]):
          rows = []
          for c in range(merged[2], merged[3]):
            rows.append([r - d_row, c - d_col])
          cells.append(rows)
        merged_cells.append(cells)

    return merged_cells


def template_format(input_path: str, output_path: str, wildcard: dict = None, table_data=None):
  """
  填充模板信息
  :param input_path: 模板文件路径
  :param output_path: 填充后输出文件路径
  :param wildcard: 模板通配符数据
  :param table_data: 模板中表格数据
  :return:
  """
  if os.path.exists(input_path):
    suffix = os.path.splitext(input_path)[-1].lower()
    if table_data and not isinstance(table_data, list):
      table_data = cjson.loads(table_data)

    if suffix == ".xls":
      tpl = xls_writer.BookWriter(input_path)
      tpl.render_book([wildcard])
      tpl.save(output_path)
      ExcelTools().tables_format_xls(output_path, table_data)
    elif suffix == ".xlsx":
      tpl = xlsx_writerx.BookWriter(input_path)
      tpl.render_book([wildcard])
      tpl.save(output_path)
      ExcelTools().tables_format_xlsx(output_path, table_data)
    elif suffix == ".docx":
      tpl = DocxTemplate(input_path)
      wildcard = DocxTools.add_rich_text(wildcard)
      tpl.render(wildcard)
      tpl.save(output_path)
      DocxTools().tables_format(output_path, table_data)


def check_format(input_path: str):
  """
  校验文件后缀与文件格式是否匹配
  :param input_path:
  :return:
  """
  res = True
  try:
    suffix = os.path.splitext(input_path)[-1].lower()
    if suffix == ".xls":
      xls_writer.BookWriter(input_path)
    elif suffix == ".xlsx":
      xlsx_writerx.BookWriter(input_path)
    elif suffix == ".docx":
      DocxTemplate(input_path)
  except Exception:
    res = False
  return res


def excel_optimize(input_path: str):
  """
  优化excel内容清除无法使用的隐患
  :param input_path:
  :return:
  """
  try:
    suffix = os.path.splitext(input_path)[-1].lower()
    if suffix == ".xlsx":
      wb = load_workbook(input_path)
      sheet = wb[wb.sheetnames[0]]
      real_max_rows, real_max_cols = sheet.max_row, sheet.max_column
      is_optimize = False

      if real_max_cols >= 100:
        sheet.delete_cols(100, real_max_cols)
        is_optimize = True

      if is_optimize:
        wb.save(input_path)
  except Exception as e:
    print("优化模板内容清除无法使用的隐患异常: %s" % e)

# 示例用法：从第1行到第5行、从第1列到第3列的区域复制到Word文档中
# excel_file = r'C:\Users\DELL\xxx/xxx-rpa/document\test-2024-04-01_09-05-26.xlsx'       # Excel文件名
# excel_file_1 = r'E:\test\c.xlsx'       # Excel文件名
# excel_file_2 = r'E:\test\b.xlsx'       # Excel文件名
# word_file = r'E:\test\test.docx'     # 输出Word文件名
# # xls_insert_docx(excel_file, word_file, start_row=1, end_row=11, start_col=1, end_col=28)
#
# doc = Document()
# input_files = r"E:\test\样式问题2_11-31-28.docx"
# DocxTools.merge_word(doc, input_files)
# DocxTools.merge_excel(doc, excel_file_1)
# DocxTools.merge_word(doc, input_files)
# DocxTools.merge_excel(doc, excel_file_1)
# DocxTools.merge_excel(doc, excel_file_2)
# DocxTools.merge_excel(doc, excel_file_1)
# DocxTools.merge_word(doc, input_files)
#
# doc.save(word_file)


# input_path = [r"E:\test\a.xlsx", r"E:\test\哈尔滨地铁数据文书模板.docx"]
#
#
# wildcard = {"title": "模板标题", "department": "技术平台部"}
# table_1 = {
#   'start_position': (13, 0),
#   'data': [['111', '112', '113'], ['121', '122', '123']]
# }
#
# table_2 = {
#   'start_position': (20, 0),
#   'data': [['211', '212', '213'], ['221', '222', '223']]
# }
#
# table_3 = {
#   'start_position': (1, 0),
#   'data': [['311', '312', '313'], ['321', '322', '323']]
# }
#
# table = {
#   'start_position': (13, 0),
#   'data': [['111', '112', '113'], ['121', '122', '123']]
# }
# doc_1 = [table]
# doc_2 = [table]
#
# doc_table = [doc_1, doc_2]
#
#
# doc_table_data = cjson.dumps(doc_table)
# rest_data = cjson.loads(doc_table_data)
#
# # print(type(doc_table_data), doc_table_data)
#
# output_dir = r"E:\test\out"
# for i, path in enumerate(input_path):
#   output_path = os.path.join(output_dir, os.path.split(path)[-1])
#   template_format(path, output_path, wildcard=wildcard, doc_table_data=rest_data[i])


# ext = os.path.splitext(input_path)[-1]
# print(check_format(input_path))

# from docxtpl import DocxTemplate, RichText
