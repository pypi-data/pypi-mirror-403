# from docxtpl import DocxTemplate
#
# # tpl = DocxTemplate('/Users/haoyang/Desktop/xxx.docx')
# tpl = DocxTemplate('/Users/haoyang/Desktop/123.doc')
#
# # 设置好各标签需要填写的内容
# context = {'xxxx': '计算机科学与技术', 'cccc': '2022050513'}
# # 将标签内容填入模板中
# tpl.render(context)
# # 保存
# tpl.save('/Users/haoyang/Desktop/new_test2.docx')

from docx import Document


def merge_word_files(input_files: [], output_file: str):
  merged_doc = Document()
  for file in input_files:
    doc = Document(file)
    for element in doc.element.body:
      merged_doc.element.body.append(element)
  merged_doc.save(output_file)


def read_word_file(input_file: str):
  doc = Document(input_file)
  text = []
  for paragraph in doc.paragraphs:
    text.append(paragraph.text)
  return "\n".join(text)
