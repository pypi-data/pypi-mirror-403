# %% #@ Libraries imports

import pandas as pd
import numpy as np
import requests
import json
from io import BytesIO
import copy
import zipfile as zp
import xlsxwriter
import xml.etree.ElementTree as ET
import uuid
from .classes import Form
import docx as dcx
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from re import findall
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from chardet import detect

# %% #@ Functions

def save_to_excel(data = {}, filename="output.xlsx", column_width=25,  row_colours={0: "#D8E4BC", 1: "#C5D9F1"}, row_bold=[0], row_wrap=[1], autofilter=True, freeze_panes=True, to_bytes=False):

    if to_bytes == True:
        workbook = xlsxwriter.Workbook(filename, {'in_memory': True})
    else:
        workbook = xlsxwriter.Workbook(filename)

    workbook.use_zip64()

    for k,v in data.items():

        worksheet = workbook.add_worksheet(name=k)

        for i in range(len(v.columns)):
            worksheet.write(0, i, v.columns[i])

        for i in range(len(v)):
            for j in range(len(v.columns)):
                if pd.isna(v.iloc[i, j]):
                    pass
                else:
                    worksheet.write(i+1, j, v.iloc[i, j])

        worksheet.set_column(0, len(v.columns), width=column_width)

        for i in range(len(v)):
            a = {}
            if i in list(row_colours.keys()):
                a["bg_color"] = row_colours[i]
            if i in row_bold:
                a["bold"] = True
            if i in row_wrap:
                a["text_wrap"] = True
            if len(a) != 0:
                worksheet.set_row(i, cell_format=workbook.add_format(a))

        if autofilter:
            worksheet.autofilter(1, 0, len(v), len(v.columns)-1)

        if freeze_panes:
            worksheet.freeze_panes(2, 0)

    workbook.close()

    if to_bytes == True:
        return filename

# %% #@ Classes

class Process_questionnaire():
    def __init__(self):
        self.attachments = None
        self.survey = None
        self.choices = None
        self.settings = None
        self.form_title = None
        self.form_version = None
        self.languages = None


    @classmethod
    def strip_double_column(cls, df):
        df.columns = [i.replace("::", ":") for i in df.columns]
        return df

    def get_data_from_files(self, form_filename, attachements_list_filenames):

        survey = pd.read_excel(form_filename, na_values=[
                               ' ', ''], keep_default_na=False, sheet_name="survey").dropna(how='all')
        self.survey = Process_questionnaire.strip_double_column(survey)
        choices = pd.read_excel(form_filename, sheet_name="choices", na_values=[
                                ' ', ''], keep_default_na=False).dropna(how='all')
        self.choices = Process_questionnaire.strip_double_column(choices)
        settings = pd.read_excel(form_filename, sheet_name="settings", na_values=[
            ' ', ''], keep_default_na=False).dropna(how='all')
        self.settings = settings
        self.form_title = self.settings['form_title'].iloc[0]
        self.form_version = self.settings['version'].iloc[0]
        attachments = {}
        for j in attachements_list_filenames:
            attachments[j] = pd.read_csv(j)
        self.attachments = attachments

    def get_data_from_odk_object(self, odk_object):
        self.survey = Process_questionnaire.strip_double_column(
            odk_object.survey)
        self.choices = Process_questionnaire.strip_double_column(
            odk_object.choices)
        self.settings = odk_object.settings
        self.form_title = self.settings['form_title'].iloc[0]
        self.form_version = self.settings['version'].iloc[0] if odk_object.form_is_published(
        ) else "FORM_NOT_PUBLISHED"
        self.attachments = odk_object.attachments
        self.form_is_published = odk_object.form_is_published()


    def get_languages(self):
        language = []
        for column in self.survey.columns:
            if column[:5] == "label":
                if len(column) == 5:
                    language.append("")
                else:
                    language.append(column.split(":")[1])
        self.languages = sorted(list(set(language)))

    def process(self, highlight_color={"begin_group": "4F81BD", "end_group": "B8CCE4", "begin_repeat": "9BBB59", "end_repeat": "D6E3BC", "calculate": "D9D9D9", "header_row": "919191"}, language=None, paragraph_spacing_points=3, compress_long_choices=True, to_memory_filename=False):

        document = dcx.Document()
        section = document.sections[-1]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = dcx.enum.section.WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
        section.top_margin = dcx.shared.Cm(0.5)
        section.bottom_margin = dcx.shared.Cm(0.5)
        section.left_margin = dcx.shared.Cm(1)
        section.right_margin = dcx.shared.Cm(1)
        heading = document.add_heading(
            f"Form title = {self.form_title}\nForm version = {str(self.form_version)}{"\n" if (language == None or language == "") else "\nForm language = "+language.split(" ")[0]}\n\n")
        heading.alignment = 1

        p = document.add_heading('Headings explained', level=2)

        p = document.add_paragraph()
        run = p.add_run('Question code')
        run.bold = True
        run = p.add_run(': The code used to identify the question')
        run.italic = True
        run = p.add_run('\n')
        run.italic = True

        run = p.add_run('Question type')
        run.bold = True
        run = p.add_run(': The question type. Common types are "text", "integer", "select_one", "select_multiple", "date", "time", "image, Note". Questions of type "calculate" are internal variables used in ODK to process calculations and should be ignored for the purpose of reviewing the questionnaire.')
        run.italic = True
        run = p.add_run('\n')
        run.italic = True

        run = p.add_run('Question')
        run.bold = True
        run = p.add_run(
            ': The question label in the specified language (default language English)')
        run.italic = True
        run = p.add_run('\n')
        run.italic = True

        run = p.add_run('Hint')
        run.bold = True
        run = p.add_run(
            ': The hint in the specified language (default language English)')
        run.italic = True
        run = p.add_run('\n')
        run.italic = True

        run = p.add_run('Select options')
        run.bold = True
        run = p.add_run(
            ': For questions of type select_one or select_multiple, these are the options that can be selected.')
        run.italic = True
        run = p.add_run('\n')
        run.italic = True

        run = p.add_run('Logic')
        run.bold = True
        run = p.add_run(
            ': These are the logics defined for the question. Different logic can be defined. ')
        run.italic = True
        run = p.add_run('Relevant')
        run.font.color.rgb = RGBColor(50, 0, 255)
        run.italic = True
        run = p.add_run(' defines the logic for showing the question or not. ')
        run.italic = True
        run = p.add_run('Default value')
        run.font.color.rgb = RGBColor(50, 0, 255)
        run.italic = True
        run = p.add_run(' is the value that the question assumes by default. ')
        run.italic = True
        run = p.add_run('Constrain value')
        run.font.color.rgb = RGBColor(50, 0, 255)
        run.italic = True
        run = p.add_run(
            ' specifies the limits imposed on the values that can be typed in. ')
        run.italic = True
        run = p.add_run('Calculation')
        run.font.color.rgb = RGBColor(50, 0, 255)
        run.italic = True
        run = p.add_run(
            ' specifies the calculations for questions of type calculate. ')
        run.italic = True
        run = p.add_run('Choice filter')
        run.font.color.rgb = RGBColor(50, 0, 255)
        run.italic = True
        run = p.add_run(
            ' specifies if the selectable options for questions of type select_one or select_multiple should be shown according to some logic. ')
        run.italic = True
        run = p.add_run('Repeat count')
        run.font.color.rgb = RGBColor(50, 0, 255)
        run.italic = True
        run = p.add_run(
            ' specifies the number of times that a repeat block is repeated. ')
        run.italic = True
        run = p.add_run('\n')
        run.italic = True

        p = document.add_heading('Rows highlighting explained', level=2)

        legend = document.add_table(rows=0, cols=2)
        for k, v in {"A block of questions begins": "4F81BD",
                     "A block of questions ends": "B8CCE4",
                     "A repeated block of questions begins": "9BBB59",
                     "A repeated block fo questions ends": "D6E3BC",
                     "A questions of type \"calculate\"": "D9D9D9",
                     "The table headers": "919191",
                     "Cells modified after review": "FFFF00"}.items():
            row_cells = legend.add_row().cells
            paragraph = row_cells[1].paragraphs[0]
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(k)
            row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = row_cells[0].paragraphs[0]
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(0)
            shading_elm = parse_xml(f'<w:shd {nsdecls('w')} w:fill="{v}"/>')
            legend.rows[-1].cells[0]._tc.get_or_add_tcPr().append(shading_elm)

        for cell in legend.columns[0].cells:
            cell.width = dcx.shared.Cm(2)
        for row in legend.rows:
            row.height = dcx.shared.Cm(1)

        p = document.add_paragraph()
        run = p.add_run('\n')
        run = p.add_run('\n')
        run.italic = True

        table = document.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        title_cells = table.rows[0].cells
        headings = ["Question code", "Question type",
                    "Question", "Hint", "Select Options", "Logic"]
        for i in range(len(headings)):
            paragraph = title_cells[i].paragraphs[0]
            run = paragraph.add_run(headings[i])
            run.bold = True

        def set_repeat_table_header(row):
            """ set repeat table row on every new page
            """
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            tblHeader = OxmlElement('w:tblHeader')
            tblHeader.set(qn('w:val'), "true")
            trPr.append(tblHeader)
            return row

        def get_choices(choice):
            choice_names = list(
                self.choices[f"label{"" if (language == None or language == "") else ":"+language}"].loc[self.choices["list_name"] == choice].map(str))
            choice_labels = list(
                self.choices["name"].loc[self.choices["list_name"] == choice].map(str))
            zipped = list(zip(choice_labels, choice_names))
            zipped = [" = ".join(x) for x in zipped]
            if compress_long_choices:
                if len(zipped) > 50:
                    zipped1 = zipped[0:25]
                    zipped2 = zipped[-25:-1]
                    zipped = zipped1 + \
                        ["...The list is longer than 50, some elements are omitted..."]+zipped2
            zipped = "\n".join(zipped)
            return zipped

        def get_from_file(file,label="label"):
            choice_list = list(self.attachments[file][label].map(str))
            if compress_long_choices:
                if len(choice_list) > 50:
                    choice_list1 = choice_list[0:25]
                    choice_list2 = choice_list[-25:-1]
                    choice_list = choice_list1 + \
                        ["...The list is longer than 50, some elements are omitted..."]+choice_list2
            return "\n".join(choice_list)

        def get_alternative_label_from_file(x):
            if type(x) is float:
                return "label"
            y = x.split(",")
            for i in y:
                if "label" in i:
                    return i.split("=")[1]
            return "label"

        def process_enclosing_variables(data):
            x = findall(r"\${.+?}", data)
            for i in x:
                data = data.replace(i, i[2:-1])
            return data

        def process_current_input(data):
            x = findall(r"\.", data)
            for i in x:
                data = data.replace(i, "input ")
            return data

        def process_string_only(column, index, cell_index):
            if not pd.isna(s[column].iloc[index]):

                paragraph = row_cells[cell_index].paragraphs[0]
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_before = Pt(paragraph_spacing_points)
                paragraph_format.space_after = Pt(paragraph_spacing_points)
                run = paragraph.add_run(process_enclosing_variables(
                    str(s[column].iloc[index])))
                if self.survey["type"].iloc[index].split(" ")[0] == "calculate":
                    run.italics = True
                    run.font.size = Pt(8)

        def process_string_only_combined(cells, index):
            columns_names = ["Relevant: ", "Default value: ", "Constrain value: ",
                             "Calculation: ", "Choice filter: ", "Repeat count: "]
            column_labels = ["relevant", "default", "constraint",
                             "calculation", "choice_filter", "repeat_count"]
            paragraph = cells[5].paragraphs[0]
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(paragraph_spacing_points)
            paragraph_format.space_after = Pt(paragraph_spacing_points)
            count = 0
            for j in range(len(column_labels)):
                if column_labels[j] in self.survey.columns:
                    if (not pd.isna(self.survey[column_labels[j]].iloc[index])):
                        count += 1
            for j in range(len(column_labels)):
                if column_labels[j] in self.survey.columns:
                    if not pd.isna(self.survey[column_labels[j]].iloc[index]):
                        run = paragraph.add_run(columns_names[j])
                        run.italics = True
                        run.font.color.rgb = RGBColor(50, 0, 255)
                        if self.survey["type"].iloc[index].split(" ")[0] == "calculate":
                            run.font.size = Pt(8)
                        if column_labels[j] == "constraint":
                            run = paragraph.add_run(
                                f"{process_current_input(process_enclosing_variables(str(self.survey[column_labels[j]].iloc[index])))}")
                            if self.survey["type"].iloc[index].split(" ")[0] == "calculate":
                                run.font.size = Pt(8)
                        else:
                            run = paragraph.add_run(
                                f"{process_enclosing_variables(str(self.survey[column_labels[j]].iloc[index]))}")
                            if self.survey["type"].iloc[index].split(" ")[0] == "calculate":
                                run.font.size = Pt(8)
                        count -= 1
                        if count != 0:
                            run = paragraph.add_run("\n")

        def cell_shading(index=None, index_counter=None, header_row=False):
            if index != None:
                sss = self.survey["type"].iloc[index].split(" ")[0]
                if sss in highlight_color.keys():
                    for j in range(6):
                        shading_elm = parse_xml(
                            f'<w:shd {nsdecls('w')} w:fill="{highlight_color[sss]}"/>')
                        table.rows[index+1 - index_counter].cells[j]._tc.get_or_add_tcPr().append(
                            shading_elm)
            if header_row:
                for j in range(6):
                    shading_elm = parse_xml(
                        f'<w:shd {nsdecls('w')} w:fill="{highlight_color['header_row']}"/>')
                    table.rows[0].cells[j]._tc.get_or_add_tcPr().append(
                        shading_elm)

        cell_shading(header_row=True)
        set_repeat_table_header(table.rows[0])
        index_counter = 0
        for i in range(len(self.survey)):
            s = self.survey
            c = self.choices
            if self.survey["relevant"].iloc[i] == "false()":
                index_counter += 1
            elif self.survey["type"].iloc[i].split(" ")[0] == "start":
                index_counter += 1
            elif self.survey["type"].iloc[i].split(" ")[0] == "end":
                index_counter += 1
            elif self.survey["type"].iloc[i].split(" ")[0] == "deviceid":
                index_counter += 1
            elif self.survey["type"].iloc[i].split(" ")[0] == "phonenumber":
                index_counter += 1
            else:
                row_cells = table.add_row().cells
                cell_shading(index=i, index_counter=index_counter)
                paragraph = row_cells[0].paragraphs[0]
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_before = Pt(paragraph_spacing_points)
                paragraph_format.space_after = Pt(paragraph_spacing_points)
                run = paragraph.add_run(s["name"].iloc[i] if not pd.isna(
                    s["name"].iloc[i]) else "")
                if self.survey["type"].iloc[i].split(" ")[0] == "calculate":
                    run.italics = True
                    run.font.size = Pt(8)

                paragraph = row_cells[1].paragraphs[0]
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_before = Pt(paragraph_spacing_points)
                paragraph_format.space_after = Pt(paragraph_spacing_points)
                run = paragraph.add_run(s["type"].iloc[i].split(" ")[0])
                if self.survey["type"].iloc[i].split(" ")[0] == "calculate":
                    run.italics = True
                    run.font.size = Pt(8)

                process_string_only(
                    f"label{"" if (language == None or language == "") else ":"+language}", i, 2)
                process_string_only(
                    f"hint{"" if (language == None or language == "") else ":"+language}", i, 3)

                if (s["type"].iloc[i].split(" ")[0] == "select_one") or (s["type"].iloc[i].split(" ")[0] == "select_multiple"):
                    row_cells[4].text = get_choices(
                        s["type"].iloc[i].split(" ")[1])
                elif (s["type"].iloc[i].split(" ")[0] == "select_one_from_file") or (s["type"].iloc[i].split(" ")[0] == "select_multiple_from_file"):
                    if type(s["parameters"].iloc[i]) is str:
                        label = get_alternative_label_from_file(
                            s["parameters"].iloc[i])
                    else:
                        label="label"
                    row_cells[4].text = get_from_file(
                        s["type"].iloc[i].split(" ")[1],label=label)
                else:
                    row_cells[4].text = ""
                process_string_only_combined(row_cells, i)
        if to_memory_filename != False:
            out = copy.deepcopy(to_memory_filename)
            document.save(out)
            return out
        else:
            document.save(
                f"{self.form_title}{"" if (self.languages == None) else ("" if (language == None or language == "") else "-" + language.split(" ")[0])}-{"Version_"+str(self.form_version)}.docx")

class ODK():

    def __init__(self, url):
        self.url = url
        self.form = None
        self.project = None
        self.survey = None
        self.choices = None
        self.settings = None
        self.attachments = None
        self.version = None

    def connect(self, email, password):

        self.email = email
        self.password = password

        req = requests.post(self.url+'/v1/sessions', data=json.dumps(
            {"email": self.email, "password": self.password}), headers={'Content-Type': 'application/json'})

        self.token = req.json()["token"]
        self.headers = {'Authorization': 'Bearer '+self.token}
        return self

    def set_target(self, project_name, form_name):
        self.project_name = project_name
        self.form_name = form_name
        self.project = self.get_project()
        self.form = self.get_form()
        if len(self.published_form_versions()[0]) == 0:
            draft = True
        else:
            draft = False
        self.survey = self.get_survey(draft=draft)
        self.choiches = self.get_choices(draft=draft)
        self.settings = self.get_settings(draft=draft)
        self.attachments = self.get_attachments(draft=draft)
        return self

    def list_projects(self, archived=False):
        req = requests.get(self.url+'/v1/projects', headers=self.headers)
        if archived == False:
            projects = [req.json()[i]["name"] for i in range(
                len(req.json())) if req.json()[i]["archived"] != True]
        else:
            projects = [req.json()[i]["name"] for i in range(
                len(req.json())) if req.json()[i]["archived"] == True]
        return projects

    def get_project(self):
        req = requests.get(self.url+'/v1/projects', headers=self.headers)
        project = [req.json()[i]["id"] for i in range(len(req.json()))
                   if req.json()[i]["name"] == self.project_name][0]

        return project

    def list_forms(self, project=None):
        req = requests.get(self.url+'/v1/projects', headers=self.headers)
        if project != None:
            project = [req.json()[i]["id"] for i in range(
                len(req.json())) if req.json()[i]["name"] == project][0]
        else:
            project = [req.json()[i]["id"] for i in range(
                len(req.json())) if req.json()[i]["name"] == self.project_name][0]
        req = requests.get(self.url+'/v1/projects/' +
                           str(project)+"/forms", headers=self.headers)
        forms = [req.json()[i]["name"] for i in range(len(req.json()))]
        return forms

    def get_form(self):

        req = requests.get(self.url+'/v1/projects/' +
                           str(self.get_project())+"/forms", headers=self.headers)
        form = [req.json()[i]["xmlFormId"] for i in range(len(req.json()))
                if req.json()[i]["name"] == self.form_name][0]

        return form

    def form_is_published(self, project_name=None, form_name=None):
        if project_name == None:
            project_name = self.project_name
        if form_name == None:
            form_name = self.form_name

        req = requests.get(self.url+'/v1/projects', headers=self.headers)
        projectid = [req.json()[i]["id"] for i in range(len(req.json()))
                     if req.json()[i]["name"] == project_name][0]
        req = requests.get(self.url+'/v1/projects/' +
                           str(projectid)+"/forms", headers=self.headers)
        form = [req.json()[i]for i in range(len(req.json()))
                if req.json()[i]["name"] == form_name][0]
        if form["publishedAt"] == None:
            return False
        else:
            return True

    def published_form_versions(self):
        req = requests.get(
            f"{self.url}/v1/projects/{str(self.get_project())}/forms/{self.get_form()}/versions", headers=self.headers)
        versions = [req.json()[i]["version"] for i in range(len(req.json()))]
        created_at = [req.json()[i]["publishedAt"]
                      for i in range(len(req.json()))]
        return versions, created_at

    def set_form_version(self, version=None):
        if version == None:
            version = self.published_form_versions()[0][0]
        self.survey = self.get_survey(version)
        self.choiches = self.get_choices(version)
        self.settings = self.get_settings(version)
        self.attachments = self.get_attachments(version)
        self.version = version

    def save_form(self, path="", save_file=True, xml=False, version=None):

        if version == None:
            version = self.published_form_versions()[0][0]

        if xml:
            extension = '.xml'
        else:
            extension = '.xlsx'

        req = requests.get(
            f"{self.url}/v1/projects/{str(self.project)}/forms/{self.form}/versions/{version}{extension}", headers=self.headers).content

        if save_file:
            file = open(path+"form_v"+version+extension, "wb")
            file.write(req)
            file.close()
        else:
            return BytesIO(req)

    def get_submissions(self):

        req = (requests.get(self.url+'/v1/projects/' +
                            str(self.project)+"/forms/" +
                            self.form+"/submissions.csv?",
                            headers=self.headers))
        df = pd.read_csv(BytesIO(req.content))
        return df

    def get_survey(self, version=None, draft=False):
        if version == None:
            if draft == False:
                version = self.published_form_versions()[0][0]

        if draft == False:
            req = requests.get(
                f"{self.url}/v1/projects/{str(self.project)}/forms/{self.form}/versions/{version}.xlsx", headers=self.headers)
        else:
            req = requests.get(
                f"{self.url}/v1/projects/{str(self.project)}/forms/{self.form}/draft.xlsx", headers=self.headers)

        survey = pd.read_excel(BytesIO(req.content), na_values=[
            ' ', ''], keep_default_na=False).dropna(how='all')
        self.survey = survey
        return survey

    def get_choices(self, version=None, draft=False):

        if version == None:
            if draft == False:
                version = self.published_form_versions()[0][0]

        if draft == False:
            req = requests.get(
                f"{self.url}/v1/projects/{str(self.project)}/forms/{self.form}/versions/{version}.xlsx", headers=self.headers)
        else:
            req = requests.get(
                f"{self.url}/v1/projects/{str(self.project)}/forms/{self.form}/draft.xlsx", headers=self.headers)
        choices = pd.read_excel(BytesIO(req.content), sheet_name="choices", na_values=[
                                ' ', ''], keep_default_na=False).dropna(how='all')
        self.choices = choices
        return choices

    def get_settings(self, version=None, draft=False):

        if version == None:
            if draft == False:
                version = self.published_form_versions()[0][0]

        if draft == False:
            req = requests.get(
                f"{self.url}/v1/projects/{str(self.project)}/forms/{self.form}/versions/{version}.xlsx", headers=self.headers)
        else:
            req = requests.get(
                f"{self.url}/v1/projects/{str(self.project)}/forms/{self.form}/draft.xlsx", headers=self.headers)

        settings = pd.read_excel(BytesIO(req.content), sheet_name="settings", na_values=[
                                 ' ', ''], keep_default_na=False).dropna(how='all')
        self.settings = settings
        return settings

    def get_repeats(self):

        req = (requests.get(self.url+'/v1/projects/' +
                            str(self.project)+"/forms/"+self.form +
                            "/submissions.csv.zip?attachments=false",
                            headers=self.headers))
        zipfile = zp.ZipFile(BytesIO(req.content))

        repeats = {}

        form_id = str(pd.read_excel(BytesIO(requests.get(self.url+'/v1/projects/'+str(self.project)+"/forms/"+self.form+".xlsx", headers=self.headers).content),
                                    sheet_name="settings")["form_id"].iloc[0])

        for j in self.survey["name"].loc[self.survey["type"] == "begin_repeat"]:
            repeats[j] = pd.read_csv(zipfile.open(
                form_id+"-" + j+".csv"), na_values=[' ', ''], keep_default_na=False).dropna(how='all')

        return repeats

    def get_attachments(self, version=None, draft=False):

        if version == None:
            if draft == False:
                version = self.published_form_versions()[0][0]

        if draft == False:
            req = requests.get(
                f"{self.url}/v1/projects/{str(self.project)}/forms/{self.form}/versions/{version}/attachments", headers=self.headers)
        else:
            req = requests.get(
                f"{self.url}/v1/projects/{str(self.project)}/forms/{self.form}/draft/attachments", headers=self.headers)

        attachments = {}

        for j in req.json():
            if draft == False:
                try:
                    attachments[j["name"]] = pd.read_csv(BytesIO((requests.get(f"{self.url}/v1/projects/{str(self.project)}/forms/{self.form}/versions/{version}/attachments/{j["name"]}", headers=self.headers)).content)) if j["name"].split(
                        ".")[-1] == "csv" else BytesIO((requests.get(f"{self.url}/v1/projects/{str(self.project)}/forms/{self.form}/versions/{version}/attachments{j["name"]}", headers=self.headers)).content)
                except UnicodeDecodeError:
                    encoding = detect((requests.get(
                        f"{self.url}/v1/projects/{str(self.project)}/forms/{self.form}/versions/{version}/attachments/{j["name"]}", headers=self.headers)).content)["encoding"]
                    attachments[j["name"]] = pd.read_csv(BytesIO((requests.get(f"{self.url}/v1/projects/{str(self.project)}/forms/{self.form}/versions/{version}/attachments/{j["name"]}", headers=self.headers)).content), encoding=encoding) if j["name"].split(
                        ".")[-1] == "csv" else BytesIO((requests.get(f"{self.url}/v1/projects/{str(self.project)}/forms/{self.form}/versions/{version}/attachments{j["name"]}", headers=self.headers)).content)

            else:
                try:
                    attachments[j["name"]] = pd.read_csv(BytesIO((requests.get(f"{self.url}/v1/projects/{str(self.project)}/forms/{self.form}/draft/attachments/{j["name"]}", headers=self.headers)).content)) if j["name"].split(
                        ".")[-1] == "csv" else BytesIO((requests.get(f"{self.url}/v1/projects/{str(self.project)}/forms/{self.form}/draft/attachments{j["name"]}", headers=self.headers)).content)
                except UnicodeDecodeError:
                    encoding = detect((requests.get(
                        f"{self.url}/v1/projects/{str(self.project)}/forms/{self.form}/draft/attachments/{j["name"]}", headers=self.headers)).content)["encoding"]
                    pd.read_csv(BytesIO((requests.get(f"{self.url}/v1/projects/{str(self.project)}/forms/{self.form}/draft/attachments/{j["name"]}", headers=self.headers)).content),encoding=encoding) if j["name"].split(
                        ".")[-1] == "csv" else BytesIO((requests.get(f"{self.url}/v1/projects/{str(self.project)}/forms/{self.form}/draft/attachments{j["name"]}", headers=self.headers)).content)
                    
        return attachments

    def get_media(self):

        req = requests.get(self.url+'/v1/projects/'+str(self.project) +
                           "/forms/"+self.form+".xlsx", headers=self.headers).content

        req = (requests.post(self.url+'/v1/projects/' +
                             str(self.project)+"/forms/" +
                             self.form+"/submissions.csv.zip?",
                             headers=self.headers))
        zipfile = zp.ZipFile(BytesIO(req.content))
        media = {}
        for name in zipfile.namelist():
            if (name.split('/')[0] == 'media') & (len(name) > 6):
                media[name.split('/')[-1]] = zipfile.read(name)
        return media

    def get_group_repeat_names(self):
        out = []
        for j in range(len(self.survey)):
            if self.survey['type'].iloc[j] in ['begin_group', 'begin_repeat']:
                out.append(self.survey["name"].iloc[j].strip())
        return out

    @classmethod
    def removing_group_repeat_names(cls, column_names, group_repeat_names):
        for j in group_repeat_names:
            for k in range(len(column_names)):
                s = column_names[k]
                if s.startswith(j+"-") and len(s) > len(j+"-"):
                    column_names[k] = column_names[k][len(j+"-"):]
        return column_names

    def processing_submission(self, process_datetimes=False):

        df = self.get_submissions()

        choices_dict = {}
        for key in set(self.choices['list_name'].map(lambda x: x.strip())):
            choices_dict[key] = dict(zip(self.choices["name"].loc[self.choices['list_name'].map(lambda x: x.strip()) == key].map(
                str), self.choices["label::English (en)"].loc[self.choices['list_name'].map(lambda x: x.strip()) == key]))

        def remove_tail(list_in):
            a = []
            for j in list_in:
                if j[-2:] == ".0":
                    a.append(j[:-2])
                else:
                    a.append(j)
            return a

        def select_one(select, value):
            y = choices_dict[select][str(value).replace(".0", "")]
            return y

        def select_multiple(select, value):
            z = []
            for s in remove_tail(list(str(value).split(" "))):
                z.append(choices_dict[select][str(s)])
            return " \n".join(z)

        def select_one_from_file(select, value):
            y = self.attachments[select]
            z = y["label"].loc[y["name"].map(str) == str(value)].iloc[0]
            return z

        def select_multiple_from_file(select, value):
            y = self.attachments[select]
            z = []
            for i in range(len(y)):
                if str(y["name"].iloc[i]) in remove_tail(list(str(value).split(" "))):
                    z.append(y["label"].iloc[i])
            return " \n".join(z)

        def rank(select, value):
            z = []
            for s in remove_tail(list(str(value).split(" "))):
                z.append(choices_dict[select][str(s)])
            return " \n".join(z)

        func = {"select_one_from_file": select_one_from_file,
                "select_one": select_one, "select_multiple": select_multiple, "select_multiple_from_file": select_multiple_from_file,
                "rank": rank}

        df.columns = self.removing_group_repeat_names(
            list(df.columns), self.get_group_repeat_names())

        for i in df.columns:
            b = self.survey["type"].loc[self.survey["name"] == i]
            if len(b) == 0:
                pass
            else:
                c = b.iloc[0].split(" ")[0]
                if c in list(func.keys()):
                    choice_list_name = b.iloc[0].split(" ")[1]
                    for j in range(len(df)):
                        if pd.isna(df[i].iloc[j]):
                            pass
                        else:
                            try:
                                df[i].iat[j] = func[c](
                                    choice_list_name, df[i].iat[j])
                            except:
                                pass

        df = df.loc[df["ReviewState"] != "rejected"]

        if process_datetimes:
            if "SubmissionDate" in df.columns:
                df["SubmissionDate"] = pd.to_datetime(df["SubmissionDate"], format="%Y-%m-%dT%H:%M:%S.%fZ")
            if 'start' in df.columns:
                df["start"] = pd.to_datetime(df["start"], format="%Y-%m-%dT%H:%M:%S.%f%z")

            for j in self.survey["name"].loc[self.survey["type"] == "datetime"]:
                try:
                    df[j] = pd.to_datetime(
                        df[j], format="%Y-%m-%dT%H:%M:%S.%f%z")
                except:
                    df[j] = pd.to_datetime(df[j], format="mixed")

            for j in self.survey["name"].loc[self.survey["type"] == "date"]:
                try:
                    df[j] = pd.to_datetime(df[j], format="%Y-%m-%d").dt.date
                except:
                    pass

            for j in self.survey["name"].loc[self.survey["type"] == "time"]:
                try:
                    df[j] = pd.to_datetime(
                        df[j], format="%H:%M:%S.%f%z").dt.time
                except:
                    pass

        return df

    def processing_repeats(self, data=None, process_datetimes=False):

        repeats = self.get_repeats()
        df = self.processing_submission() if type(data) == type(None) else data
        set_not_rejected = list(df["KEY"])

        choices_dict = {}
        for key in set(self.choices['list_name'].map(lambda x: x.strip())):
            choices_dict[key] = dict(zip(self.choices["name"].loc[self.choices['list_name'].map(lambda x: x.strip()) == key].map(
                str), self.choices["label::English (en)"].loc[self.choices['list_name'].map(lambda x: x.strip()) == key]))

        def remove_tail(list_in):
            a = []
            for j in list_in:
                if j[-2:] == ".0":
                    a.append(j[:-2])
                else:
                    a.append(j)
            return a

        def select_one(select, value):
            y = choices_dict[select][str(value).replace(".0", "")]
            return y

        def select_multiple(select, value):
            z = []
            for s in remove_tail(list(str(value).split(" "))):
                z.append(choices_dict[select][str(s)])
            return " \n".join(z)

        def select_one_from_file(select, value):
            y = pd.read_csv(select)
            z = y["label"].loc[y["name"] == str(value)].iloc[0]
            return z

        def select_multiple_from_file(select, value):
            y = pd.read_csv(select)
            z = []
            for i in range(len(y)):
                if str(y["name"].iloc[i]) in remove_tail(list(str(value).split(" "))):
                    z.append(y["label"].iloc[i])
            return " \n".join(z)

        def rank(select, value):
            z = []
            for s in remove_tail(list(str(value).split(" "))):
                z.append(choices_dict[select][str(s)])
            return " \n".join(z)

        func = {"select_one_from_file": select_one_from_file,
                "select_one": select_one, "select_multiple": select_multiple, "select_multiple_from_file": select_multiple_from_file,
                "rank":rank}

        group_names = self.get_group_repeat_names()

        for k in repeats.keys():
            repeats[k].columns = self.removing_group_repeat_names(
                list(repeats[k].columns), group_names)

            for i in repeats[k].columns:

                b = self.survey["type"].loc[self.survey["name"] == i]
                if len(b) == 0:
                    pass
                else:
                    c = b.iloc[0].split(" ")[0]
                    if c in list(func.keys()):
                        choice_list_name = b.iloc[0].split(" ")[1]
                        for j in range(len(repeats[k])):
                            if pd.isna(repeats[k][i].iloc[j]):
                                pass
                            else:
                                try:
                                    repeats[k][i].iat[j] = func[c](
                                        choice_list_name, repeats[k][i].iat[j])
                                except:
                                    pass

        for j in repeats.keys():

            repeats[j] = repeats[j].loc[[True if repeats[j]["PARENT_KEY"].iloc[i].split(
                "/")[0] in set_not_rejected else False for i in range(len(repeats[j]))]]

            if process_datetimes:

                for i in self.survey["name"].loc[self.survey["type"] == "datetime"]:
                    if i in repeats[j].columns:
                        try:
                            repeats[j][i] = pd.to_datetime(
                                repeats[j][i], format="%Y-%m-%dT%H:%M:%S.%f%z")
                        except:
                            repeats[j][i] = pd.to_datetime(
                                repeats[j][i], format="mixed")

                for i in self.survey["name"].loc[self.survey["type"] == "date"]:
                    if i in repeats[j].columns:
                        try:
                            repeats[j][i] = pd.to_datetime(
                                repeats[j][i], format="%Y-%m-%d").dt.date
                        except:
                            repeats[j][i] = pd.to_datetime(
                                repeats[j][i], format="mixed").dt.date

                for i in self.survey["name"].loc[self.survey["type"] == "time"]:
                    if i in repeats[j].columns:
                        try:
                            repeats[j][i] = pd.to_datetime(
                                repeats[j][i], format="%H:%M:%S.%f%z").dt.time
                        except:
                            repeats[j][i] = pd.to_datetime(
                                repeats[j][i], format="mixed").dt.time

        return repeats

    def process_all(self, variable='', time_variable='start', process_datetimes=False, process_media=True):

        submissions = self.processing_submission(
            process_datetimes=process_datetimes)
        survey = self.survey.dropna(how='all')
        choices = self.choices
        settings = self.settings
        repeats = self.processing_repeats(process_datetimes=process_datetimes)
        survey_name = self.form_name
        form = self.form
        variable = variable
        time_variable = time_variable
        if process_media:
            media = self.get_media()
        else:
            media = None
        attachments = self.attachments

        return Form(submissions, survey, choices, settings, repeats, survey_name, form, variable, time_variable, media, attachments)

    def add_questions(self,data):

        df = copy.deepcopy(data)

        for j in df.select_dtypes(include=['datetime64', 'datetimetz']).columns:
            df[j] = df[j].astype(str)
        if 'start' in df.columns:
            df['start'] = df['start'].astype(str)
        a = []
        for j in data.columns:
            if j in list(self.survey["name"]):
                x = self.survey["label::English (en)"].loc[self.survey["name"]
                                                           == j].iloc[0]
                a.append(x)
            else:
                a.append(np.nan)

        df.loc[-1] = a
        df.sort_index(inplace=True)

        return df

    def save_data(self, path=""):

        req = requests.get(self.url+'/v1/projects/'+str(self.project) +
                           "/forms/"+self.form+".xlsx", headers=self.headers).content

        version = str(pd.read_excel(BytesIO(req),
                                    sheet_name="settings")["version"].iloc[0])
        req = (requests.post(self.url+'/v1/projects/' +
                             str(self.project)+"/forms/" +
                             self.form+"/submissions.csv.zip?",
                             headers=self.headers))

        file = open(path+self.form_name+"_v"+version+".zip", "wb")
        file.write(req.content)
        file.close()

    def listing_submissions(self,draft=False):

        req = (requests.get(f"{self.url}/v1/projects/{str(self.project)}/forms/{self.form}{"/draft" if draft==True else ""}/submissions",
                            headers=self.headers))
        return req.json()

    def get_submission_metadata(self, instance, draft=False):

        req = (requests.get(f"{self.url}/v1/projects/{str(self.project)}/forms/{self.form}{"/draft" if draft == True else ""}/submissions/{instance}", headers=self.headers))
        return req.json()

    def get_submission_xml(self, instance,draft=False):

        req = (requests.get(f"{self.url}/v1/projects/{str(self.project)}/forms/{self.form}{"/draft" if draft == True else ""}/submissions/{instance}.xml", headers=self.headers))
        return req.content

    def put_submission(self, instance, data,draft=False):

        req = (requests.put(url=f"{self.url}/v1/projects/{str(self.project)}/forms/{self.form}{"/draft" if draft == True else ""}/submissions/{instance}", data=data, headers=self.headers))
        return req

    def create_submission(self, data, draft=False):

        req = requests.post(url=f"{self.url}/v1/projects/{str(self.project)}/forms/{self.form}{"/draft" if draft == True else ""}/submissions", data=data, headers=self.headers)
        return req

    def get_parent_tag(self, tag):

        n = self.survey.loc[self.survey['name'] == tag].index[0]
        begin_group = len(
            self.survey.iloc[:n].loc[self.survey['type'] == 'begin_group'])
        end_group = len(
            self.survey.iloc[:n].loc[self.survey['type'] == 'end_group'])
        begin_repeat = len(
            self.survey.iloc[:n].loc[self.survey['type'] == 'begin_repeat'])
        end_repeat = len(
            self.survey.iloc[:n].loc[self.survey['type'] == 'end_repeat'])

        if end_repeat < begin_repeat:
            return (self.survey['name'].iloc[:n].loc[self.survey['type'] == 'begin_repeat']).iloc[-1]
        elif end_group < begin_group:
            return (self.survey['name'].iloc[:n].loc[self.survey['type'] == 'begin_group']).iloc[-1]
        else:
            return None

    def return_element(self, tree, data: str):
        out = []
        for elem in tree.iter():
            if elem.tag == data:
                out.append(elem)
            else:
                pass
        if len(out) == 0:
            return None
        else:
            return out

    def modify_variable_xml(self, xml, variable: str, function, mask=None):
        """
        The argument "mask" is used to change values of a variable belonging to a repeat block.
        "mask" is by default equal to "None". If left as default, all entries for such variable in the repeated block will be changed.
        "mask" can also be set to a list of booleans (i.e. [True,False,True,True]), equal in length to the number of entries under the repeat group for the variable and the submission being edited. Entries that in order correspond to True are edited, entries that correspond to False are not.
        """
        tree = ET.parse(BytesIO(xml))
        elements = self.return_element(tree, variable)
        if mask == None:
            mask = [True]*len(elements)
        else:
            mask = mask
        if elements == None:
            print(f"{variable} is not in the xml")
            return xml
        else:
            try:
                for d in range(len(elements)):
                    if mask[d] == True:
                        k = elements[d].text
                        elements[d].text = function(k)
                xml_out = BytesIO()
                tree.write(xml_out, encoding='utf-8')
                return xml_out.getvalue()
            except:
                print('an error occurred while processing for variable ', variable)
                return xml

    def update_xml(self, xml):

        tree = ET.parse(BytesIO(xml))
        root = tree.getroot()

        if tree.find('meta').find('deprecatedID') == None:
            old = tree.find('meta').find('instanceID').text
            tree.find('meta').find(
                'instanceID').text = 'uuid:'+str(uuid.uuid4())
            deprecated = ET.Element("deprecatedID")
            deprecated.text = old
            root.find('meta').append(deprecated)

        else:
            if len(tree.find('meta').find('deprecatedID').text) > 0:
                old = tree.find('meta').find('instanceID').text
                tree.find('meta').find(
                    'instanceID').text = 'uuid:'+str(uuid.uuid4())
                root.find('meta').find('deprecatedID').text = old
        xml_out = BytesIO()
        tree.write(xml_out, encoding='utf-8')
        return xml_out.getvalue()

    def change_submission(self, xml, id):
        self.put_submission(id, self.update_xml(xml))

    def drop_variable_xml(self, xml, variable: str, parent_tag=None):

        tree = ET.parse(BytesIO(xml))
        root = tree.getroot()
        for elem in tree.iter():
            if elem.tag == variable:
                if parent_tag == None:
                    root.remove(elem)
                else:
                    self.return_element(tree, parent_tag).remove(elem)
        xml_out = BytesIO()
        tree.write(xml_out, encoding='utf-8')
        return xml_out.getvalue()

    def add_variable_xml(self, xml, variable: str, parent_tag=None):

        tree = ET.parse(BytesIO(xml))
        root = tree.getroot()
        if type(self.return_element(tree, variable)) == type(None):
            if parent_tag == None:
                child = ET.SubElement(root, variable)
            else:
                child = ET.SubElement(
                    self.return_element(tree, parent_tag)[0], variable)
            xml_out = BytesIO()
            tree.write(xml_out, encoding='utf-8')
            return xml_out.getvalue()
