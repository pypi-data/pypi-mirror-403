#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
DOC和PDF格式转换器测试文件
"""

import os
from kit import TranslationHelper


def test_doc_conversion():
    """测试DOC格式转换功能"""
    print("=== DOC格式转换测试 ===")
    
    # 检查是否安装了python-docx库
    try:
        import docx
    except ImportError:
        print("警告: 未安装python-docx库，跳过DOC测试")
        print("安装命令: pip install python-docx")
        return
    
    # 创建测试用的Word文档
    doc_path = create_test_word_document()
    
    try:
        helper = TranslationHelper()
        
        # 提取文本
        extracted_texts = helper.extract_text(doc_path, 'docx')
        print(f"提取的文本数量: {len(extracted_texts)}")
        print("提取的文本:")
        for i, text in enumerate(extracted_texts):
            print(f"  {i+1}. {text}")
        
        # 模拟翻译
        translated_texts = [f"[翻译]{text}" for text in extracted_texts]
        print("\n模拟翻译结果:")
        for i, text in enumerate(translated_texts):
            print(f"  {i+1}. {text}")
        
        # 应用翻译
        output_path = helper.apply_translation(doc_path, translated_texts, 'docx')
        print(f"\n翻译后的文档已保存至: {output_path}")
        
        # 验证输出文件是否存在
        if os.path.exists(output_path):
            print("✓ DOC格式转换测试成功")
        else:
            print("✗ DOC格式转换测试失败: 输出文件未创建")
            
    except Exception as e:
        print(f"✗ DOC格式转换测试出现异常: {e}")
    finally:
        # 清理测试文件
        if os.path.exists(doc_path):
            os.remove(doc_path)


def test_pdf_conversion():
    """测试PDF格式转换功能"""
    print("\n=== PDF格式转换测试 ===")
    
    # 检查是否安装了PyMuPDF库
    try:
        import fitz
    except ImportError:
        print("警告: 未安装PyMuPDF库，跳过PDF测试")
        print("安装命令: pip install PyMuPDF")
        return
    
    # 创建测试用的PDF文档
    pdf_path = create_test_pdf_document()
    
    try:
        helper = TranslationHelper()
        
        # 提取文本
        extracted_texts = helper.extract_text(pdf_path, 'pdf')
        print(f"提取的文本数量: {len(extracted_texts)}")
        print("提取的文本:")
        for i, text in enumerate(extracted_texts):
            print(f"  {i+1}. {text}")
        
        # 模拟翻译
        translated_texts = [f"[翻译]{text}" for text in extracted_texts]
        print("\n模拟翻译结果:")
        for i, text in enumerate(translated_texts):
            print(f"  {i+1}. {text}")
        
        # 应用翻译
        output_path = helper.apply_translation(pdf_path, translated_texts, 'pdf')
        print(f"\n翻译后的文档已保存至: {output_path}")
        
        # 验证输出文件是否存在
        if os.path.exists(output_path):
            print("✓ PDF格式转换测试成功")
        else:
            print("✗ PDF格式转换测试失败: 输出文件未创建")
            
    except Exception as e:
        print(f"✗ PDF格式转换测试出现异常: {e}")
    finally:
        # 清理测试文件
        if os.path.exists(pdf_path):
            os.remove(pdf_path)


def create_test_word_document():
    """创建测试用的Word文档"""
    import docx
    
    doc = docx.Document()
    doc.add_heading('测试文档', 0)
    doc.add_paragraph('这是第一个段落。')
    doc.add_paragraph('这是第二个段落，包含一些文本内容。')
    
    # 添加表格
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = '表格单元格1'
    table.cell(0, 1).text = '表格单元格2'
    table.cell(1, 0).text = '表格单元格3'
    table.cell(1, 1).text = '表格单元格4'
    
    doc_path = 'test_document.docx'
    doc.save(doc_path)
    return doc_path


def create_test_pdf_document():
    """创建测试用的PDF文档"""
    import fitz  # PyMuPDF
    
    # 创建一个新的PDF文档
    pdf = fitz.open()
    page = pdf.new_page()
    
    # 添加文本内容
    page.insert_text((50, 72), "测试PDF文档", fontsize=14)
    page.insert_text((50, 108), "这是第一行文本内容。", fontsize=12)
    page.insert_text((50, 132), "这是第二行文本内容。", fontsize=12)
    page.insert_text((50, 156), "这是第三行文本内容。", fontsize=12)
    
    pdf_path = 'test_document.pdf'
    pdf.save(pdf_path)
    pdf.close()
    return pdf_path


def test_format_detection():
    """测试格式检测功能"""
    print("\n=== 格式检测测试 ===")
    
    helper = TranslationHelper()
    supported_formats = helper.get_supported_formats()
    print(f"支持的格式: {supported_formats}")
    
    # 测试获取转换器
    for format_type in ['docx', 'pdf']:
        try:
            converter = helper.get_converter(format_type)
            print(f"✓ 成功获取 {format_type} 转换器: {type(converter).__name__}")
        except Exception as e:
            print(f"✗ 获取 {format_type} 转换器失败: {e}")


if __name__ == "__main__":
    # 测试格式检测
    test_format_detection()
    
    # 测试DOC转换
    test_doc_conversion()
    
    # 测试PDF转换
    test_pdf_conversion()
    
    print("\n=== 所有测试完成 ===")