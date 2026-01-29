from docx import Document
from docx.document import Document as DocObject
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table
from docx.section import Section
from typing import Optional

def translate_word_document(
    input_path: str,
    output_path: str,
    translate_func  # 已有的翻译函数（输入字符串，输出翻译后的字符串）
) -> None:
    """
    翻译Word文档并保留源格式（.docx格式）
    
    Args:
        input_path: 输入Word文档路径（必须是.docx格式）
        output_path: 输出翻译后文档的路径
        translate_func: 翻译函数，要求：输入单个字符串，输出翻译后的字符串
    """
    # 1. 打开源文档和创建目标文档（复制源文档格式）
    source_doc = Document(input_path)
    target_doc = Document()
    
    # 2. 复制文档的核心样式（确保格式一致）
    copy_document_styles(source_doc, target_doc)
    
    # 3. 处理文档主体内容（段落、表格等）
    for element in source_doc.element.body:
        if element.tag.endswith('p'):  # 处理段落
            handle_paragraph(element, target_doc, translate_func)
        elif element.tag.endswith('tbl'):  # 处理表格
            handle_table(element, target_doc, translate_func)
    
    # 4. 处理页眉页脚（所有节的页眉页脚）
    for section_idx, section in enumerate(source_doc.sections):
        target_section = target_doc.sections[section_idx]
        # 处理页眉
        handle_header_footer(section.header, target_section.header, translate_func)
        # 处理页脚
        handle_header_footer(section.footer, target_section.footer, translate_func)
    
    # 5. 保存翻译后的文档
    target_doc.save(output_path)
    print(f"翻译完成！文件已保存至：{output_path}")

def copy_document_styles(source_doc: DocObject, target_doc: DocObject) -> None:
    """复制源文档的样式到目标文档（确保格式一致）"""
    # 复制段落样式
    for style in source_doc.styles:
        style_name = style.name
        if style_name.startswith('Heading') or style_name == 'Normal':
            try:
                source_style = source_doc.styles[style_name]
                target_style = target_doc.styles[style_name]
                # 复制字体格式
                target_style.font.name = source_style.font.name
                target_style.font.size = source_style.font.size
                target_style.font.bold = source_style.font.bold
                target_style.font.italic = source_style.font.italic
                target_style.font.color.rgb = source_style.font.color.rgb
                # 复制段落格式
                target_style.paragraph_format.alignment = source_style.paragraph_format.alignment
                target_style.paragraph_format.line_spacing = source_style.paragraph_format.line_spacing
            except KeyError:
                continue

def handle_paragraph(source_element, target_doc: DocObject, translate_func) -> None:
    """处理单个段落：翻译文本并保留格式"""
    # 创建新段落，复制原段落样式
    target_para = target_doc.add_paragraph()
    source_para = Paragraph(source_element, target_doc)
    
    # 复制段落格式
    target_para.alignment = source_para.alignment
    target_para.paragraph_format.line_spacing = source_para.paragraph_format.line_spacing
    target_para.paragraph_format.space_before = source_para.paragraph_format.space_before
    target_para.paragraph_format.space_after = source_para.paragraph_format.space_after
    
    # 按格式相似性对runs进行分组
    run_groups = group_runs_by_format(source_para.runs)
    
    # 对每个组进行翻译处理
    for group in run_groups:
        # 如果组内都是空文本，则直接复制
        if all(not run.text.strip() for run in group):
            for source_run in group:
                target_run = target_para.add_run(source_run.text)
                copy_run_format(source_run, target_run)
            continue
            
        # 收集组内所有文本
        group_text = "".join(run.text for run in group)
        
        # 翻译组文本
        if group_text.strip():
            translated_text = translate_func(group_text)
        else:
            translated_text = group_text
            
        # 直接创建一个run包含所有翻译后的文本（不再分割）
        target_run = target_para.add_run(translated_text)
        # 复制组内第一个run的格式（因为组内所有run格式相同）
        copy_run_format(group[0], target_run)
def group_runs_by_format(runs):
    """根据格式相似性将runs分组"""
    if not runs:
        return []
        
    groups = []
    current_group = [runs[0]]
    
    for i in range(1, len(runs)):
        if is_format_similar(runs[i-1], runs[i]):
            current_group.append(runs[i])
        else:
            groups.append(current_group)
            current_group = [runs[i]]
    
    groups.append(current_group)
    return groups

def is_format_similar(run1, run2):
    """判断两个run的格式是否相似"""
    # 检查主要格式属性是否相同
    formats_match = (
        run1.font.name == run2.font.name and
        run1.font.size == run2.font.size and
        run1.font.bold == run2.font.bold and
        run1.font.italic == run2.font.italic and
        run1.font.underline == run2.font.underline and
        run1.font.color.rgb == run2.font.color.rgb and
        run1.font.highlight_color == run2.font.highlight_color
    )
    
    return formats_match

def copy_run_format(source_run, target_run):
    """复制run格式"""
    target_run.font.name = source_run.font.name
    target_run.font.size = source_run.font.size
    target_run.font.bold = source_run.font.bold
    target_run.font.italic = source_run.font.italic
    target_run.font.underline = source_run.font.underline
    target_run.font.color.rgb = source_run.font.color.rgb
    target_run.font.highlight_color = source_run.font.highlight_color

def handle_table(source_element, target_doc: DocObject, translate_func) -> None:
    """处理表格：翻译单元格文本并保留表格格式"""
    # 创建新表格，复制原表格行数和列数
    source_table = Table(source_element, target_doc)
    target_table = target_doc.add_table(
        rows=len(source_table.rows),
        cols=len(source_table.columns)
    )
    
    # 复制表格样式（边框、背景色等）
    target_table.style = source_table.style
    
    # 遍历表格的每个单元格
    for row_idx, source_row in enumerate(source_table.rows):
        target_row = target_table.rows[row_idx]
        for col_idx, source_cell in enumerate(source_row.cells):
            target_cell = target_row.cells[col_idx]
            
            # 复制单元格格式（宽度、对齐方式）
            target_cell.width = source_cell.width
            target_cell.vertical_alignment = source_cell.vertical_alignment
            
            # 处理单元格内的段落（一个单元格可能有多个段落）
            for source_para in source_cell.paragraphs:
                # 创建新段落并复制格式
                target_para = target_cell.add_paragraph()
                target_para.alignment = source_para.alignment
                
                # 遍历段落中的Run并翻译
                for source_run in source_para.runs:
                    if source_run.text.strip():
                        translated_text = translate_func(source_run.text)
                    else:
                        translated_text = source_run.text
                    
                    # 复制Run格式
                    target_run = target_para.add_run(translated_text)
                    target_run.font = source_run.font

def handle_header_footer(source_header_footer, target_header_footer, translate_func) -> None:
    """处理页眉/页脚：翻译文本并保留格式"""
    # 清空目标页眉/页脚的默认内容
    for para in target_header_footer.paragraphs:
        para.clear()
    
    # 复制源页眉/页脚的内容和格式
    for source_para in source_header_footer.paragraphs:
        target_para = target_header_footer.add_paragraph()
        target_para.alignment = source_para.alignment
        
        for source_run in source_para.runs:
            if source_run.text.strip():
                translated_text = translate_func(source_run.text)
            else:
                translated_text = source_run.text
            
            target_run = target_para.add_run(translated_text)
            target_run.font = source_run.font

# ------------------------------
# 测试示例（需替换为你的translate函数）
# ------------------------------
if __name__ == "__main__":
    from test import *
    translator = BaiduTranslator(config=config)
    # 调用翻译函数
    translate_word_document(
        input_path="test.docx",    # 你的输入Word文档路径
        output_path="translated.docx",# 翻译后的输出路径
        translate_func=translator.translate, # 你的真实翻译函数
    )
    print(translator.get_performance_metrics())