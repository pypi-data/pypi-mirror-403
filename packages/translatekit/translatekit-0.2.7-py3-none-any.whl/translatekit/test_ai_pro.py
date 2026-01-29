from docx import Document
from docx.document import Document as DocObject
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table
from docx.section import Section
from typing import List, Tuple, Any

def extract_text_from_word_document(input_path: str) -> List[Tuple[str, Any]]:
    """
    从Word文档中提取所有需要翻译的文本
    
    Args:
        input_path: 输入Word文档路径（必须是.docx格式）
        
    Returns:
        包含文本和位置信息的列表，每个元素为(文本, 位置信息)
    """
    source_doc = Document(input_path)
    text_elements = []
    
    # 1. 提取文档主体内容（段落）
    for para_idx, paragraph in enumerate(source_doc.paragraphs):
        if paragraph.text.strip():  # 只处理非空段落
            run_groups = group_runs_by_format(paragraph.runs)
            for group_idx, group in enumerate(run_groups):
                group_text = "".join(run.text for run in group)
                if group_text.strip():
                    position_info = {
                        'type': 'paragraph',
                        'para_idx': para_idx,
                        'group_idx': group_idx,
                        'run_lengths': [len(run.text) for run in group]
                    }
                    text_elements.append((group_text, position_info))
    
    # 2. 提取表格内容
    for table_idx, table in enumerate(source_doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    if paragraph.text.strip():
                        for run_idx, run in enumerate(paragraph.runs):
                            if run.text.strip():
                                position_info = {
                                    'type': 'table',
                                    'table_idx': table_idx,
                                    'row_idx': row_idx,
                                    'cell_idx': cell_idx,
                                    'para_idx': para_idx,
                                    'run_idx': run_idx
                                }
                                text_elements.append((run.text, position_info))
    
    # 3. 提取页眉页脚内容
    for section_idx, section in enumerate(source_doc.sections):
        # 页眉
        for para_idx, paragraph in enumerate(section.header.paragraphs):
            if paragraph.text.strip():
                for run_idx, run in enumerate(paragraph.runs):
                    if run.text.strip():
                        position_info = {
                            'type': 'header',
                            'section_idx': section_idx,
                            'para_idx': para_idx,
                            'run_idx': run_idx
                        }
                        text_elements.append((run.text, position_info))
        
        # 页脚
        for para_idx, paragraph in enumerate(section.footer.paragraphs):
            if paragraph.text.strip():
                for run_idx, run in enumerate(paragraph.runs):
                    if run.text.strip():
                        position_info = {
                            'type': 'footer',
                            'section_idx': section_idx,
                            'para_idx': para_idx,
                            'run_idx': run_idx
                        }
                        text_elements.append((run.text, position_info))
    
    return text_elements

def apply_translations_to_word_document(
    input_path: str,
    output_path: str,
    translated_elements: List[Tuple[str, Any]]
) -> None:
    """
    将翻译后的文本应用到Word文档中
    
    Args:
        input_path: 输入Word文档路径
        output_path: 输出翻译后文档的路径
        translated_elements: 包含翻译后文本和位置信息的列表
    """
    # 打开源文档和创建目标文档
    source_doc = Document(input_path)
    target_doc = Document()
    
    # 复制文档样式
    copy_document_styles(source_doc, target_doc)
    
    # 创建空的目标文档结构
    initialize_document_structure(source_doc, target_doc)
    
    # 应用翻译到各个部分
    apply_translations_to_paragraphs(translated_elements, target_doc)
    apply_translations_to_tables(translated_elements, target_doc)
    apply_translations_to_headers_footers(translated_elements, target_doc)
    
    # 保存文档
    target_doc.save(output_path)
    print(f"翻译完成！文件已保存至：{output_path}")

def initialize_document_structure(source_doc: DocObject, target_doc: DocObject) -> None:
    """初始化目标文档结构，复制源文档的段落和表格"""
    # 复制段落结构
    for _ in source_doc.paragraphs:
        target_doc.add_paragraph()
    
    # 复制表格结构
    for table in source_doc.tables:
        target_doc.add_table(
            rows=len(table.rows),
            cols=len(table.columns)
        )

def apply_translations_to_paragraphs(translated_elements: List[Tuple[str, Any]], target_doc: DocObject) -> None:
    """将翻译应用到段落"""
    paragraph_elements = [elem for elem in translated_elements if elem[1]['type'] == 'paragraph']
    
    for text, position_info in paragraph_elements:
        para_idx = position_info['para_idx']
        group_idx = position_info['group_idx']
        run_lengths = position_info['run_lengths']
        
        if para_idx < len(target_doc.paragraphs):
            paragraph = target_doc.paragraphs[para_idx]
            
            # 确保段落有足够的runs
            while len(paragraph.runs) <= group_idx:
                paragraph.add_run("")
            
            # 按原始run长度分布翻译文本
            translated_chars = 0
            for i, run_length in enumerate(run_lengths):
                run_text = text[translated_chars:translated_chars + run_length]
                translated_chars += run_length
                
                if i < len(paragraph.runs):
                    paragraph.runs[i].text = run_text

def apply_translations_to_tables(translated_elements: List[Tuple[str, Any]], target_doc: DocObject) -> None:
    """将翻译应用到表格"""
    table_elements = [elem for elem in translated_elements if elem[1]['type'] == 'table']
    
    for text, position_info in table_elements:
        table_idx = position_info['table_idx']
        row_idx = position_info['row_idx']
        cell_idx = position_info['cell_idx']
        para_idx = position_info['para_idx']
        run_idx = position_info['run_idx']
        
        if (table_idx < len(target_doc.tables) and 
            row_idx < len(target_doc.tables[table_idx].rows) and
            cell_idx < len(target_doc.tables[table_idx].rows[row_idx].cells)):
            
            cell = target_doc.tables[table_idx].rows[row_idx].cells[cell_idx]
            
            if para_idx < len(cell.paragraphs):
                paragraph = cell.paragraphs[para_idx]
                
                # 确保段落有足够的runs
                while len(paragraph.runs) <= run_idx:
                    paragraph.add_run("")
                
                if run_idx < len(paragraph.runs):
                    paragraph.runs[run_idx].text = text

def apply_translations_to_headers_footers(translated_elements: List[Tuple[str, Any]], target_doc: DocObject) -> None:
    """将翻译应用到页眉页脚"""
    header_elements = [elem for elem in translated_elements if elem[1]['type'] == 'header']
    footer_elements = [elem for elem in translated_elements if elem[1]['type'] == 'footer']
    
    # 处理页眉
    for text, position_info in header_elements:
        section_idx = position_info['section_idx']
        para_idx = position_info['para_idx']
        run_idx = position_info['run_idx']
        
        if section_idx < len(target_doc.sections):
            header = target_doc.sections[section_idx].header
            
            if para_idx < len(header.paragraphs):
                paragraph = header.paragraphs[para_idx]
                
                # 确保段落有足够的runs
                while len(paragraph.runs) <= run_idx:
                    paragraph.add_run("")
                
                if run_idx < len(paragraph.runs):
                    paragraph.runs[run_idx].text = text
    
    # 处理页脚
    for text, position_info in footer_elements:
        section_idx = position_info['section_idx']
        para_idx = position_info['para_idx']
        run_idx = position_info['run_idx']
        
        if section_idx < len(target_doc.sections):
            footer = target_doc.sections[section_idx].footer
            
            if para_idx < len(footer.paragraphs):
                paragraph = footer.paragraphs[para_idx]
                
                # 确保段落有足够的runs
                while len(paragraph.runs) <= run_idx:
                    paragraph.add_run("")
                
                if run_idx < len(paragraph.runs):
                    paragraph.runs[run_idx].text = text

# 以下辅助函数保持不变
def copy_document_styles(source_doc: DocObject, target_doc: DocObject) -> None:
    """复制源文档的样式到目标文档（确保格式一致）"""
    for style in source_doc.styles:
        style_name = style.name
        if style_name.startswith('Heading') or style_name == 'Normal':
            try:
                source_style = source_doc.styles[style_name]
                target_style = target_doc.styles[style_name]
                # 复制字体格式
                target_style.font.name = source_style.font.name
                target_style.font.size = source_style.font.size
                target_style.font.bold = source_style.font.bold
                target_style.font.italic = source_style.font.italic
                target_style.font.color.rgb = source_style.font.color.rgb
                # 复制段落格式
                target_style.paragraph_format.alignment = source_style.paragraph_format.alignment
                target_style.paragraph_format.line_spacing = source_style.paragraph_format.line_spacing
            except KeyError:
                continue

def group_runs_by_format(runs):
    """根据格式相似性将runs分组"""
    if not runs:
        return []
        
    groups = []
    current_group = [runs[0]]
    
    for i in range(1, len(runs)):
        if is_format_similar(runs[i-1], runs[i]):
            current_group.append(runs[i])
        else:
            groups.append(current_group)
            current_group = [runs[i]]
    
    groups.append(current_group)
    return groups

def is_format_similar(run1, run2):
    """判断两个run的格式是否相似"""
    formats_match = (
        run1.font.name == run2.font.name and
        run1.font.size == run2.font.size and
        run1.font.bold == run2.font.bold and
        run1.font.italic == run2.font.italic and
        run1.font.underline == run2.font.underline and
        run1.font.color.rgb == run2.font.color.rgb and
        run1.font.highlight_color == run2.font.highlight_color
    )
    
    return formats_match

# 使用示例
if __name__ == "__main__":
    from test import *
    translator = BaiduTranslator(config=config)
    
    # 1. 提取需要翻译的文本
    text_elements = extract_text_from_word_document("24.docx")
    texts_to_translate = [text for text, _ in text_elements]
    position_infos = [info for _, info in text_elements]
    
    # 2. 批量翻译所有文本
    translated_texts = translator.translate(texts_to_translate)
    
    # 3. 将翻译后的文本应用回文档
    translated_elements = list(zip(translated_texts, position_infos))
    apply_translations_to_word_document(
        input_path="24.docx",
        output_path="translated.docx",
        translated_elements=translated_elements
    )
    
    print(translator.get_performance_metrics())