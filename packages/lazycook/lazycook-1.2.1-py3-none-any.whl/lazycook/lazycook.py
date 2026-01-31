#lets you download the conversations
#takes docs
import asyncio
from PyPDF2 import PdfReader
import hashlib
import json
import logging
import mimetypes
import os
import threading
import time
from dataclasses import dataclass, asdict
from datetime import datetime, timedelta
from enum import Enum
from functools import wraps
from pathlib import Path
from threading import Lock
from typing import Dict, Any, List, Optional, Callable
from rich.markdown import Markdown
from rich import box
import textwrap
from rich.text import Text
from rich.align import Align
from docx import Document as DocxDocument
from rich.panel import Panel
from rich.filesize import decimal
from rich.prompt import Confirm
import google.generativeai as genai
from rich.box import ROUNDED
from rich.columns import Columns
from rich.console import Console
from rich.live import Live
from rich.progress import (
    Progress,
    TextColumn,
    BarColumn,
    TimeRemainingColumn,
    TimeElapsedColumn,
    MofNCompleteColumn
)
import matplotlib
matplotlib.use('Qt5Agg')
import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
import numpy as np
import io
import base64
from rich.progress_bar import ProgressBar
from rich.prompt import Prompt, Confirm
from rich.rule import Rule
from rich.status import Status
from rich.style import Style
from rich.syntax import Syntax
from rich.table import Table
from rich.text import Text

"""# --- Logging Configuration ---
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[logging.FileHandler('multi_agent_assistant.log'), logging.StreamHandler()]
)
logger = logging.getLogger(__name__)
"""
# --- Logging Configuration ---
import sys
import io

# Force UTF-8 encoding BEFORE any logging setup
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace', line_buffering=True)
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace', line_buffering=True)

# Simple logging with UTF-8
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('multi_agent_assistant.log', encoding='utf-8', errors='replace'),
        logging.StreamHandler(sys.stdout)
    ],
    force=True
)
logger = logging.getLogger(__name__)


# --- Decorators ---
def log_errors(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        try:
            return func(*args, **kwargs)
        except Exception as e:
            import traceback
            logger.error(f"Error in {func.__name__}: {e}")
            logger.error(f"Full traceback: {traceback.format_exc()}")
            raise  # Re-raise the exception instead of swallowing it

    return wrapper


# --- Data Classes ---
class AgentRole(Enum):
    GENERATOR = "generator"
    ANALYZER = "analyzer"
    OPTIMIZER = "optimizer"
    VALIDATOR = "validator"


@dataclass
class AgentResponse:
    agent_role: AgentRole
    content: str
    confidence: float
    suggestions: List[str]
    errors_found: List[str]
    improvements: List[str]
    metadata: Dict[str, Any]

    def to_dict(self):
        return {
            'agent_role': self.agent_role.value,
            'content': self.content,
            'confidence': self.confidence,
            'suggestions': self.suggestions,
            'errors_found': self.errors_found,
            'improvements': self.improvements,
            'metadata': self.metadata
        }


@dataclass
class MultiAgentSession:
    session_id: str
    user_query: str
    iterations: List[Dict[str, Any]]
    final_response: str
    quality_score: float
    total_iterations: int
    timestamp: datetime
    context_used: str
    metadata: Dict[str, Any] = None

    def to_dict(self):
        serializable_iterations = []
        for iteration in self.iterations:
            serializable_iteration = iteration.copy()
            for key in ['generator_response', 'analyzer_response', 'optimizer_response', 'validator_response']:
                if key in serializable_iteration and hasattr(serializable_iteration[key], 'to_dict'):
                    serializable_iteration[key] = serializable_iteration[key].to_dict()
                elif key in serializable_iteration and isinstance(serializable_iteration[key], dict):
                    if 'agent_role' in serializable_iteration[key] and hasattr(
                            serializable_iteration[key]['agent_role'], 'value'):
                        serializable_iteration[key]['agent_role'] = serializable_iteration[key]['agent_role'].value
            serializable_iterations.append(serializable_iteration)
        return {
            'session_id': self.session_id,
            'user_query': self.user_query,
            'iterations': serializable_iterations,
            'final_response': self.final_response,
            'quality_score': self.quality_score,
            'total_iterations': self.total_iterations,
            'timestamp': self.timestamp.isoformat(),
            'context_used': self.context_used,
            'metadata': self.metadata or {}
        }


@dataclass
class Conversation:
    id: str
    user_id: str
    timestamp: datetime
    user_message: str
    ai_response: str
    multi_agent_session: Optional['MultiAgentSession']
    context: str
    sentiment: str
    topics: List[str]
    potential_followups: List[str]

    def to_dict(self):
        return {
            'id': self.id,
            'user_id': self.user_id,
            'timestamp': self.timestamp.isoformat(),
            'user_message': self.user_message,
            'ai_response': self.ai_response,
            'multi_agent_session': self.multi_agent_session.to_dict() if self.multi_agent_session else None,
            'sentiment': self.sentiment,
            'topics': self.topics,
            'potential_followups': self.potential_followups
        }

    @classmethod
    def from_dict(cls, data: Dict):
        multi_agent_data = data.get('multi_agent_session')
        multi_agent_session = None
        if multi_agent_data:
            multi_agent_session = MultiAgentSession(
                session_id=multi_agent_data['session_id'],
                user_query=multi_agent_data['user_query'],
                iterations=multi_agent_data['iterations'],
                final_response=multi_agent_data['final_response'],
                quality_score=multi_agent_data['quality_score'],
                total_iterations=multi_agent_data['total_iterations'],
                timestamp=datetime.fromisoformat(multi_agent_data['timestamp']),
                context_used=multi_agent_data.get('context_used', '')
            )
        return cls(
            id=data['id'],
            user_id=data['user_id'],
            timestamp=datetime.fromisoformat(data['timestamp']),
            user_message=data['user_message'],
            ai_response=data['ai_response'],
            multi_agent_session=multi_agent_session,
            context="", # Don't load/store full context in history records
            sentiment=data.get('sentiment', 'neutral'),
            topics=data['topics'],
            potential_followups=data['potential_followups']
        )


@dataclass
class Task:
    id: str
    conversation_id: str
    task_type: str
    description: str
    priority: int
    status: str
    created_at: datetime
    scheduled_for: datetime
    result: Optional[str] = None
    metadata: Optional[Dict] = None

    def to_dict(self):
        return {
            'id': self.id,
            'conversation_id': self.conversation_id,
            'task_type': self.task_type,
            'description': self.description,
            'priority': self.priority,
            'status': self.status,
            'created_at': self.created_at.isoformat(),
            'scheduled_for': self.scheduled_for.isoformat(),
            'result': self.result,
            'metadata': self.metadata
        }

    @classmethod
    def from_dict(cls, data: Dict):
        return cls(
            id=data['id'],
            conversation_id=data['conversation_id'],
            task_type=data['task_type'],
            description=data['description'],
            priority=data['priority'],
            status=data['status'],
            created_at=datetime.fromisoformat(data['created_at']),
            scheduled_for=datetime.fromisoformat(data['scheduled_for']),
            result=data.get('result'),
            metadata=data.get('metadata')
        )


@dataclass
class Document:
    id: str
    filename: str
    content: str
    file_type: str
    file_size: int
    upload_time: datetime
    user_id: str
    hash_value: str
    metadata: Dict[str, Any]
    processing_status: str = "success"  # NEW: Tracks if content extraction was successful
    processing_message: Optional[str] = None # NEW: Stores messages for failed/partial processing

    def to_dict(self):
        return {
            'id': self.id,
            'filename': self.filename,
            'content': self.content,
            'file_type': self.file_type,
            'file_size': self.file_size,
            'upload_time': self.upload_time.isoformat(),
            'user_id': self.user_id,
            'hash_value': self.hash_value,
            'metadata': self.metadata,
            'processing_status': self.processing_status, # Include in dict
            'processing_message': self.processing_message # Include in dict
        }

    @classmethod
    def from_dict(cls, data: Dict):
        return cls(
            id=data['id'],
            filename=data['filename'],
            content=data['content'],
            file_type=data['file_type'],
            file_size=data['file_size'],
            upload_time=datetime.fromisoformat(data['upload_time']),
            user_id=data['user_id'],
            hash_value=data['hash_value'],
            metadata=data.get('metadata', {}),
            processing_status=data.get('processing_status', 'success'), # Handle existing data
            processing_message=data.get('processing_message') # Handle existing data
        )


# --- Custom Progress Bar ---
class AnimatedProgressBar(ProgressBar):
    """Custom animated progress bar with different styles for each stage"""

    def __init__(self, style="green", pulse_style="grey50"):
        super().__init__()
        self.style = style
        self.pulse_style = pulse_style
        self._animation_frames = ["⠋", "⠙", "⠹", "⠸", "⠼", "⠴", "⠦", "⠧", "⠇", "⠏"]
        self._frame_index = 0

    def get_animation_frame(self):
        self._frame_index = (self._frame_index + 1) % len(self._animation_frames)
        return self._animation_frames[self._frame_index]

    def __get_ranges__(self, task):
        total = task.total or 100
        completed = task.completed
        if total == 0:
            return [(0, 1, self.pulse_style)]

        width = self.width or 20
        full = completed / total

        # Add animation to the pulse style
        animated_pulse = f"{self.get_animation_frame()} {self.pulse_style}"

        ranges = [(0, full, self.style)]
        if full < 1.0:
            ranges.append((full, 1.0, animated_pulse))
        return ranges


# --- File Manager ---
class TextFileManager:
    def __init__(self, data_dir: str = "multi_agent_data", conversation_limit: int = 70, document_limit: int = 50, api_key: Optional[str] = None):
        self.conversation_limit = conversation_limit
        self.document_limit = document_limit
        self.data_dir = Path(data_dir)
        self.file_lock = Lock()
        self.api_key = api_key

        # FIX 1: Initialize missing cached_context
        self._cached_context = {}
        self._context_cache_time = {}
        self._cache_ttl = timedelta(minutes=5)  # Cache expires after 5 min

        self.data_dir.mkdir(exist_ok=True)
        self.conversations_file = self.data_dir / "conversations.json"
        self.tasks_file = self.data_dir / "tasks.json"
        self.documents_file = self.data_dir / "documents.json"
        self.session_conversations_file = self.data_dir / "new_convo.json"
        self._ensure_files_exist()
        self._initialize_session_file()
        self.max_documents_per_user = document_limit
        self.max_storage_per_user = 100 * 1024 * 1024 #100MB

    def _get_effective_limit(self, provided_limit: Optional[int]) -> int:
        """
        Resolve the effective limit to use.
        Priority: provided_limit > instance limit > default (70)
        """
        if provided_limit is not None and provided_limit > 0:
            return provided_limit
        return getattr(self, 'conversation_limit', 70)

    def get_conversation_context(self, user_id: str, limit: int = None, user_query: str = "") -> str:
        """Get context for a user with relevance-based document filtering"""
        limit = self._get_effective_limit(limit)

        # Build fresh context - Always get fresh for the current query to ensure relevance
        session_conversations = self.get_session_conversations(user_id, limit // 2)
        historical_conversations = self.get_recent_conversations(user_id, limit // 2)

        # Remove duplicates
        historical_ids = {conv.id for conv in historical_conversations}
        unique_session_convs = [conv for conv in session_conversations
                                if conv.id not in historical_ids]

        # Combine and sort
        all_conversations = unique_session_convs + historical_conversations
        all_conversations.sort(key=lambda x: x.timestamp, reverse=True)
        conversations = all_conversations[:limit]  # Apply final limit

        if not conversations:
            return "No previous conversation history available."

        context_parts = ["=== CONVERSATION CONTEXT (Session + History) ==="]
        
        # Track topics for follow-up detection
        recent_topics = []
        if conversations:
            for conv in conversations[:3]: # Last 3 convs
                if conv.topics:
                    recent_topics.extend(conv.topics)

        for i, conv in enumerate(conversations):
            source = "Current Session" if conv in unique_session_convs else "Previous Session"
            context_parts.append(
                f"\n--- Conversation {i + 1} ({conv.timestamp.strftime('%Y-%m-%d %H:%M')}) [{source}] ---"
            )
            context_parts.append(f"USER: {conv.user_message}")
            context_parts.append(f"ASSISTANT: {conv.ai_response}")

            if conv.multi_agent_session:
                session = conv.multi_agent_session
                context_parts.append(
                    f"[Quality: {session.quality_score:.2f} | Iterations: {session.total_iterations}]"
                )

            if conv.topics:
                context_parts.append(f"[Topics: {', '.join(conv.topics)}]")

        # Add document context
        docs_context = self.get_documents_context(user_id, 50, full_content=True, user_query=user_query, recent_topics=recent_topics)
        if docs_context:
            context_parts.append(f"\n--- 📄 RELEVANT DOCUMENTS (Reference Only) ---")
            context_parts.append("Note: These documents are for background information. Do NOT incorporate them if they are unrelated to the current query.")
            context_parts.append(docs_context)

        context_parts.append("\n=== END OF CONTEXT ===")
        context = "\n".join(context_parts)

        return context



    def _ensure_files_exist(self):
        for file_path in [self.conversations_file, self.tasks_file, self.documents_file]:
            if not file_path.exists():
                file_path.write_text("[]")

    def cleanup_session_file(self):
        """Clean up the session conversation file"""
        try:
            if self.session_conversations_file.exists():
                self.session_conversations_file.unlink()
                logger.info("Session conversation file cleaned up")
        except Exception as e:
            logger.error(f"Failed to cleanup session file: {e}")

    def _read_json_file(self, file_path: Path) -> List[Dict]:
        try:
            with self.file_lock:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return json.load(f)
        except json.JSONDecodeError as e:
            logger.error(f"JSON decode error in {file_path}: {e}")
            backup_path = file_path.with_suffix('.backup')
            file_path.rename(backup_path)
            return []
        except FileNotFoundError:
            return []

    def debug_context_flow(self, user_id: str):
        """Debug method to check context flow"""
        session_convs = self.get_session_conversations(user_id, 5)
        all_convs = self.get_recent_conversations(user_id, 5)

        print(f"Session conversations: {len(session_convs)}")
        print(f"All conversations: {len(all_convs)}")

        context = self.get_conversation_context(user_id)
        print(f"Context length: {len(context)} chars")
        print(f"Context preview: {context[:200]}...")

    def _write_json_file(self, file_path: Path, data: List[Dict]):
        with self.file_lock:
            try:
                json_content = json.dumps(data, indent=2, ensure_ascii=False)
                temp_path = file_path.with_suffix('.tmp')
                with open(temp_path, 'w', encoding='utf-8') as temp_file:
                    temp_file.write(json_content)
                temp_path.replace(file_path)
            except Exception as e:
                logger.error(f"Failed to write {file_path}: {e}")
                raise

    def _initialize_session_file(self):
        """Initialize or clear the session conversation file"""
        try:
            self.session_conversations_file.write_text("[]")
            logger.info("Session conversation file initialized")
        except Exception as e:
            logger.error(f"Failed to initialize session file: {e}")

    @log_errors
    def save_conversation(self, conversation: Conversation):
        # Invalidate cache for this user
        user_id = conversation.user_id
        keys_to_remove = [k for k in self._cached_context.keys() if k.startswith(f"{user_id}_")]
        for key in keys_to_remove:
            self._cached_context.pop(key, None)
            self._context_cache_time.pop(key, None)

        # Original save logic
        conversations = self._read_json_file(self.conversations_file)
        conversations = [c for c in conversations if c.get('id') != conversation.id]
        conversations.append(conversation.to_dict())
        conversations.sort(key=lambda x: x.get('timestamp', ''), reverse=True)
        
        # Enforce history limit
        if len(conversations) > self.conversation_limit:
            conversations = conversations[:self.conversation_limit]
            
        self._write_json_file(self.conversations_file, conversations)

        # Save to session file
        session_conversations = self._read_json_file(self.session_conversations_file)
        session_conversations = [c for c in session_conversations if c.get('id') != conversation.id]
        session_conversations.append(conversation.to_dict())
        session_conversations.sort(key=lambda x: x.get('timestamp', ''), reverse=True)
        
        # Limit session entries too
        if len(session_conversations) > 20:
            session_conversations = session_conversations[:20]
            
        self._write_json_file(self.session_conversations_file, session_conversations)

    @log_errors
    def clear_cached_context(self, user_id: str):
        """Clear cached context for a user - now actually works"""
        keys_to_remove = [k for k in self._cached_context.keys() if k.startswith(f"{user_id}_")]
        for key in keys_to_remove:
            self._cached_context.pop(key, None)
            self._context_cache_time.pop(key, None)
        logger.info(f"Cleared {len(keys_to_remove)} cached contexts for {user_id}")

    @log_errors
    def save_conversations_batch(self, conversations: List[Conversation]):
        conversations_data = [conv.to_dict() for conv in conversations]
        self._write_json_file(self.conversations_file, conversations_data)

    @log_errors
    def get_session_conversations(self, user_id: str, limit: int = None) -> List[Conversation]:
        limit = self._get_effective_limit(limit)
        session_data = self._read_json_file(self.session_conversations_file)
        user_conversations = [c for c in session_data if c.get('user_id') == user_id]
        user_conversations.sort(key=lambda x: x.get('timestamp', ''), reverse=True)
        return [Conversation.from_dict(conv_data) for conv_data in user_conversations[:limit]]

    @log_errors
    def get_recent_conversations(self, user_id: str, limit: int = None) -> List[Conversation]:
        limit = self._get_effective_limit(limit)
        conversations_data = self._read_json_file(self.conversations_file)
        user_conversations = [c for c in conversations_data if c.get('user_id') == user_id]
        user_conversations.sort(key=lambda x: x.get('timestamp', ''), reverse=True)
        return [Conversation.from_dict(conv_data) for conv_data in user_conversations[:limit]]


    @log_errors
    def save_task(self, task: Task):
        tasks = self._read_json_file(self.tasks_file)
        tasks = [t for t in tasks if t.get('id') != task.id]
        tasks.append(task.to_dict())
        tasks.sort(key=lambda x: x.get('created_at', ''), reverse=True)
        self._write_json_file(self.tasks_file, tasks)

    @log_errors
    def get_recent_conversations(self, user_id: str, limit: int = None) -> List[Conversation]:
        if limit is None:
            limit = getattr(self, 'conversation_limit', 70)
        conversations_data = self._read_json_file(self.conversations_file)
        user_conversations = [c for c in conversations_data if c.get('user_id') == user_id]
        user_conversations.sort(key=lambda x: x.get('timestamp', ''), reverse=True)
        return [Conversation.from_dict(conv_data) for conv_data in user_conversations[:limit]]




    @log_errors
    def get_pending_tasks(self) -> List[Task]:
        tasks_data = self._read_json_file(self.tasks_file)
        now = datetime.now()
        pending_tasks = []
        for task_data in tasks_data:
            try:
                if task_data.get('status') == 'pending' and datetime.fromisoformat(
                        task_data.get('scheduled_for', '')) <= now:
                    pending_tasks.append(Task.from_dict(task_data))
            except Exception as e:
                logger.warning(f"Skipping malformed task data: {e}")
        pending_tasks.sort(key=lambda x: (-x.priority, x.scheduled_for))
        return pending_tasks

    @log_errors
    def get_storage_stats(self) -> Dict[str, Any]:
        conversations_data = self._read_json_file(self.conversations_file)
        tasks_data = self._read_json_file(self.tasks_file)
        documents_data = self._read_json_file(self.documents_file)
        user_stats = {}
        for conv in conversations_data:
            user_id = conv.get('user_id', 'unknown')
            user_stats[user_id] = user_stats.get(user_id, 0) + 1
        return {
            'total_conversations': len(conversations_data),
            'total_tasks': len(tasks_data),
            'total_documents': len(documents_data),
            'users': user_stats,
            'oldest_conversation': min([c.get('timestamp', '') for c in conversations_data], default='none'),
            'newest_conversation': max([c.get('timestamp', '') for c in conversations_data], default='none'),
            'files_exist': {
                'conversations': self.conversations_file.exists(),
                'tasks': self.tasks_file.exists(),
                'documents': self.documents_file.exists(),
                'session_conversations': len(self._read_json_file(self.session_conversations_file)),
                'session_file_exists': self.session_conversations_file.exists(),
            }
        }

    @log_errors
    def save_document(self, document: Document):
        documents = self._read_json_file(self.documents_file)
        documents = [d for d in documents if d.get('id') != document.id]
        documents.append(document.to_dict())
        documents.sort(key=lambda x: x.get('upload_time', ''), reverse=True)
        self._write_json_file(self.documents_file, documents)
        logger.info(f"Document saved: {document.filename}")

    @log_errors
    def get_user_documents(self, user_id: str, limit: int = 20) -> List[Document]:
        documents_data = self._read_json_file(self.documents_file)
        user_docs = [d for d in documents_data if d.get('user_id') == user_id]
        user_docs.sort(key=lambda x: x.get('upload_time', ''), reverse=True)
        return [Document.from_dict(doc_data) for doc_data in user_docs[:limit]]

    @log_errors
    def delete_document(self, document_id: str, user_id: str) -> bool:
        try:
            documents = self._read_json_file(self.documents_file)
            original_count = len(documents)
            documents = [d for d in documents if not (d.get('id') == document_id and d.get('user_id') == user_id)]
            if len(documents) < original_count:
                self._write_json_file(self.documents_file, documents)
                return True
            return False
        except Exception as e:
            logger.error(f"Failed to delete document: {e}")
            return False

    @log_errors
    def get_documents_context(self, user_id: str, limit: int = 50, full_content: bool = True, user_query: str = "", recent_topics: List[str] = None) -> str:
        """Get document context with relevance filtering based on query and topics"""
        documents = self.get_user_documents(user_id, limit)
        if not documents:
            return ""

        # Simple relevance filtering
        query_words = set(user_query.lower().split()) if user_query else set()
        topics = set([t.lower() for t in recent_topics]) if recent_topics else set()
        
        # Stop words to ignore
        stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'if', 'then', 'else', 'when', 'at', 'from', 'by', 'for', 'with', 'about', 'to', 'in', 'on', 'is', 'are', 'was', 'were', 'tell', 'me', 'more', 'about', 'this', 'that'}
        query_keywords = query_words - stop_words

        # If query is very vague (e.g. "tell me more"), use recent topics as priority
        is_vague_query = len(query_keywords) <= 1
        
        context_parts = []
        for i, doc in enumerate(documents):
            # Relevance Check
            content_lower = doc.content.lower()
            filename_lower = doc.filename.lower()
            
            # 1. Check if query keywords appear in document
            keyword_match = any(word in content_lower or word in filename_lower for word in query_keywords) if query_keywords else False
            
            # 2. Check if recent topics appear in document
            topic_match = any(topic in content_lower or topic in filename_lower for topic in topics) if topics else False
            
            # DECISION: Include if it matches the current query, 
            # OR if the query is vague AND it matches recent conversation topics.
            is_relevant = keyword_match or (is_vague_query and topic_match)
            
            # EMERGENCY: If we have no context and only 1-2 documents, include them as a guess
            if not is_relevant and len(documents) <= 2 and not is_vague_query:
                is_relevant = True

            if is_relevant:
                context_parts.append(f"\n--- Document {i + 1}: {doc.filename} (Relevant) ---")
                if full_content:
                    context_parts.append(doc.content)
                else:
                    content_preview = doc.content[:500] + "..." if len(doc.content) > 500 else doc.content
                    context_parts.append(content_preview)

        return "\n".join(context_parts)

    @log_errors
    def process_uploaded_file(self, file_path: str, user_id: str) -> Optional[Document]:
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                return None

            file_type = mimetypes.guess_type(file_path)[0] or 'text/plain'
            content = ""
            processing_status = "success"
            processing_message = None

            if file_type.startswith('text/'):
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            elif file_type == 'application/pdf':
                try:
                    reader = PdfReader(file_path)
                    number_of_pages = len(reader.pages)
                    content_parts = []

                    for i in range(number_of_pages):
                        page = reader.pages[i]
                        text = page.extract_text()
                        if text and text.strip():
                            content_parts.append(text)

                    content = "\n\n".join(content_parts) if content_parts else ""

                    # FALLBACK: If no text was extracted, use Gemini for OCR
                    if not content.strip() and self.api_key:
                        logger.info(f"PDF extraction yielded no text for {file_path.name}. Falling back to Gemini OCR...")
                        try:
                            genai.configure(api_key=self.api_key)
                            model = genai.GenerativeModel('gemini-2.5-flash')
                            
                            # Upload to Gemini for processing
                            sample_file = genai.upload_file(path=str(file_path), display_name=file_path.name)
                            
                            # Wait for the file to be processed
                            while sample_file.state.name == "PROCESSING":
                                time.sleep(2)
                                sample_file = genai.get_file(sample_file.name)
                            
                            if sample_file.state.name == "FAILED":
                                raise Exception(f"Gemini file processing failed: {sample_file.state.name}")
                                
                            response = model.generate_content([
                                sample_file,
                                "Extract all text from this PDF document. Provide only the extracted text without any commentary."
                            ])
                            content = response.text
                            logger.info(f"Gemini OCR extracted {len(content)} characters from PDF: {file_path.name}")
                            
                            # Clean up file from Gemini
                            genai.delete_file(sample_file.name)
                            
                        except Exception as ocr_err:
                            logger.error(f"Gemini OCR fallback failed: {ocr_err}")
                            content = f"[PDF - No text content found and Gemini OCR failed: {str(ocr_err)}]"
                            processing_status = "ocr_failed"
                            processing_message = str(ocr_err)
                    elif not content.strip():
                        content = "[PDF - No text content extracted and no API key for OCR fallback]"
                        processing_status = "no_text_detected"
                        processing_message = "No text found in PDF and API key not configured for OCR."
                    else:
                        logger.info(f"PDF extracted normally: {number_of_pages} pages, {len(content)} characters")

                except Exception as e:
                    logger.error(f"Error extracting PDF content: {e}")
                    content = f"[PDF - Error extracting content: {str(e)}]"
                    processing_status = "error"
                    processing_message = str(e)

            elif file_type.startswith('image/'):
                logger.info(f"Processing image with Gemini Vision: {file_path.name}")
                try:
                    if self.api_key:
                        genai.configure(api_key=self.api_key)
                        model = genai.GenerativeModel('gemini-2.5-flash')
                        
                        # Upload to Gemini for processing
                        sample_file = genai.upload_file(path=str(file_path), display_name=file_path.name)
                        
                        # Wait for the file to be processed
                        while sample_file.state.name == "PROCESSING":
                            time.sleep(2)
                            sample_file = genai.get_file(sample_file.name)
                        
                        if sample_file.state.name == "FAILED":
                            raise Exception(f"Gemini file processing failed: {sample_file.state.name}")
                            
                        # Request both OCR and visual description
                        response = model.generate_content([
                            sample_file,
                            "Analyze this image. First, extract all text found in the image. Second, provide a detailed visual description of the image content (objects, scene, colors). Combine both into a clear report."
                        ])
                        content = response.text
                        logger.info(f"Gemini Vision processed image: {file_path.name}")
                        
                        # Clean up file from Gemini
                        genai.delete_file(sample_file.name)
                    else:
                        content = f"[Image file: {file_path.name} - No API key configured for Gemini Vision.]"
                        processing_status = "error"
                        processing_message = "API key missing for Gemini Vision processing."

                except Exception as e:
                    logger.error(f"Error with Gemini Vision for {file_path.name}: {e}")
                    content = f"[Image file: {file_path.name} - Gemini Vision error: {str(e)}]"
                    processing_status = "error"
                    processing_message = str(e)

            elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                logger.info(f"Processing DOCX file: {file_path.name}")
                try:
                    doc = DocxDocument(file_path)
                    content = "\n".join([para.text for para in doc.paragraphs])
                    logger.info(f"DOCX extracted: {len(content)} characters")
                except Exception as e:
                    logger.error(f"Error extracting DOCX content: {e}")
                    content = f"[DOCX - Error extracting content: {str(e)}]"
                    processing_status = "error"
                    processing_message = str(e)

            elif file_type == 'text/markdown':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            elif file_type == 'text/csv':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            else:
                content = f"[Binary file: {file_path.name} - content not extractable]"
                processing_status = "unsupported_type"
                processing_message = "File type not supported for content extraction."

            with open(file_path, 'rb') as f:
                file_hash = hashlib.md5(f.read()).hexdigest()

            file_size = file_path.stat().st_size
            document = Document(
                id=f"{user_id}_{int(time.time())}_{file_hash[:8]}",
                filename=file_path.name,
                content=content,
                file_type=file_type,
                file_size=file_size,
                upload_time=datetime.now(),
                user_id=user_id,
                hash_value=file_hash,
                metadata={
                    'original_path': str(file_path),
                    'processed_at': datetime.now().isoformat()
                },
                processing_status=processing_status,
                processing_message=processing_message
            )
            self.save_document(document)
            return document
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            return None

    @log_errors
    def clear_all_conversations(self, user_id: str) -> bool:
        """Clear all conversations for a specific user"""
        try:
            # Clear from main conversations file
            conversations = self._read_json_file(self.conversations_file)
            original_count = len([c for c in conversations if c.get('user_id') == user_id])
            conversations = [c for c in conversations if c.get('user_id') != user_id]
            self._write_json_file(self.conversations_file, conversations)

            # Clear from session conversations file
            session_conversations = self._read_json_file(self.session_conversations_file)
            session_conversations = [c for c in session_conversations if c.get('user_id') != user_id]
            self._write_json_file(self.session_conversations_file, session_conversations)

            # Clear cached context
            self._cached_context.pop(user_id, None)

            logger.info(f"Cleared {original_count} conversations for user {user_id}")
            return original_count > 0
        except Exception as e:
            logger.error(f"Failed to clear conversations: {e}")
            return False

    @log_errors
    def clear_all_documents(self, user_id: str) -> int:
        """Clear all documents for a specific user and return count deleted"""
        try:
            documents = self._read_json_file(self.documents_file)
            user_docs_count = len([d for d in documents if d.get('user_id') == user_id])
            documents = [d for d in documents if d.get('user_id') != user_id]
            self._write_json_file(self.documents_file, documents)
            logger.info(f"Cleared {user_docs_count} documents for user {user_id}")
            return user_docs_count
        except Exception as e:
            logger.error(f"Failed to clear documents: {e}")
            return 0

    @log_errors
    def delete_document_by_filename(self, filename: str, user_id: str) -> bool:
        """Delete a document by filename for specific user"""
        try:
            documents = self._read_json_file(self.documents_file)
            original_count = len(documents)
            documents = [d for d in documents if not (d.get('filename') == filename and d.get('user_id') == user_id)]
            if len(documents) < original_count:
                self._write_json_file(self.documents_file, documents)
                logger.info(f"Deleted document: {filename} for user {user_id}")
                return True
            return False
        except Exception as e:
            logger.error(f"Failed to delete document {filename}: {e}")
            return False

    @log_errors
    def get_documents_capacity_info(self, user_id: str) -> Dict[str, Any]:
        """Get document capacity information for a user"""
        try:
            documents = self.get_user_documents(user_id, 1000)  # Get all docs
            total_size = sum(doc.file_size for doc in documents)
            total_count = len(documents)

            # Define limits (configurable)
            max_documents = getattr(self, 'max_documents_per_user', 50)
            max_total_size = getattr(self, 'max_storage_per_user', 50 * 1024 * 1024)  # 50MB default

            return {
                'total_documents': total_count,
                'total_size_bytes': total_size,
                'total_size_mb': round(total_size / (1024 * 1024), 2),
                'max_documents': max_documents,
                'max_size_mb': round(max_total_size / (1024 * 1024), 2),
                'documents_remaining': max(0, max_documents - total_count),
                'size_remaining_mb': max(0, round((max_total_size - total_size) / (1024 * 1024), 2)),
                'at_document_limit': total_count >= max_documents,
                'at_size_limit': total_size >= max_total_size,
                'usage_percentage': min(100, round((total_size / max_total_size) * 100, 1))
            }
        except Exception as e:
            logger.error(f"Failed to get capacity info: {e}")
            return {
                'total_documents': 0,
                'total_size_bytes': 0,
                'total_size_mb': 0,
                'max_documents': 50,
                'max_size_mb': 50,
                'documents_remaining': 50,
                'size_remaining_mb': 50,
                'at_document_limit': False,
                'at_size_limit': False,
                'usage_percentage': 0
            }

    @log_errors
    def clear_old_conversations(self, user_id: str, keep_recent: int = 20) -> int:
        """Clear old conversations, keeping only the most recent ones"""
        try:
            conversations = self._read_json_file(self.conversations_file)
            user_conversations = [c for c in conversations if c.get('user_id') == user_id]
            other_conversations = [c for c in conversations if c.get('user_id') != user_id]

            # Sort by timestamp and keep only recent ones
            user_conversations.sort(key=lambda x: x.get('timestamp', ''), reverse=True)
            kept_conversations = user_conversations[:keep_recent]
            deleted_count = len(user_conversations) - len(kept_conversations)

            # Combine with other users' conversations
            all_conversations = other_conversations + kept_conversations
            self._write_json_file(self.conversations_file, all_conversations)

            # Also clean session conversations
            session_conversations = self._read_json_file(self.session_conversations_file)
            session_conversations = [c for c in session_conversations if c.get('user_id') != user_id]
            self._write_json_file(self.session_conversations_file, session_conversations)

            logger.info(f"Cleared {deleted_count} old conversations for user {user_id}, kept {len(kept_conversations)}")
            return deleted_count
        except Exception as e:
            logger.error(f"Failed to clear old conversations: {e}")
            return 0

    @log_errors
    def export_conversations_to_file(self, user_id: str, format: str = "txt", limit: int = None) -> Optional[str]:
        """Export conversations to a downloadable file"""
        try:
            limit = self._get_effective_limit(limit)
            conversations = self.get_recent_conversations(user_id, limit)

            if not conversations:
                return None

            # Sort oldest to newest for reading flow
            conversations.sort(key=lambda x: x.timestamp)

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"chat_export_{user_id}_{timestamp}.{format}"
            filepath = self.data_dir / filename

            if format == "txt":
                content = self._format_txt_export(conversations)
            elif format == "md":
                content = self._format_markdown_export(conversations)
            elif format == "json":
                content = self._format_json_export(conversations)
            else:
                return None

            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)

            logger.info(f"Exported {len(conversations)} conversations to {filename}")
            return str(filepath)

        except Exception as e:
            logger.error(f"Export failed: {e}")
            return None

    def _format_txt_export(self, conversations: List[Conversation]) -> str:
        """Format conversations as plain text"""
        lines = []
        lines.append("=" * 80)
        lines.append("LAZYCOOK CHAT EXPORT")
        lines.append(f"Exported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append(f"Total Conversations: {len(conversations)}")
        lines.append("=" * 80)
        lines.append("")

        for i, conv in enumerate(conversations, 1):
            lines.append(f"\n{'=' * 80}")
            lines.append(f"CONVERSATION #{i}")
            lines.append(f"Date: {conv.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
            lines.append(f"{'=' * 80}")
            lines.append(f"\nUSER:\n{conv.user_message}")
            lines.append(f"\nASSISTANT:\n{conv.ai_response}")

            if conv.multi_agent_session:
                lines.append(f"\n[Quality Score: {conv.multi_agent_session.quality_score:.2f}]")
                lines.append(f"[Iterations: {conv.multi_agent_session.total_iterations}]")

            lines.append("")

        return "\n".join(lines)

    def _format_markdown_export(self, conversations: List[Conversation]) -> str:
        """Format conversations as Markdown"""
        lines = []
        lines.append("# 🔥 LAZYCOOK Chat Export")
        lines.append(f"\n**Exported:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append(f"**Total Conversations:** {len(conversations)}")
        lines.append("\n---\n")

        for i, conv in enumerate(conversations, 1):
            lines.append(f"\n## 💬 Conversation #{i}")
            lines.append(f"**Date:** {conv.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
            lines.append(f"\n### 👤 User:\n{conv.user_message}")
            lines.append(f"\n### 🤖 Assistant:\n{conv.ai_response}")

            if conv.multi_agent_session:
                lines.append(
                    f"\n*Quality: {conv.multi_agent_session.quality_score:.2f} | Iterations: {conv.multi_agent_session.total_iterations}*")

            lines.append("\n---\n")

        return "\n".join(lines)

    def _format_json_export(self, conversations: List[Conversation]) -> str:
        """Format conversations as JSON"""
        export_data = {
            "export_info": {
                "exported_at": datetime.now().isoformat(),
                "total_conversations": len(conversations),
                "format_version": "1.0"
            },
            "conversations": [conv.to_dict() for conv in conversations]
        }
        return json.dumps(export_data, indent=2, ensure_ascii=False)

# --- AI Agent ---
class AIAgent:
    def __init__(self, api_key: str, role: AgentRole, temperature: float = 0.7):
        genai.configure(api_key=api_key)
        self.role = role
        self.model = genai.GenerativeModel(
            model_name='models/gemini-2.5-flash',  # Most stable vision/plotting model
            generation_config={
                "temperature": temperature,
                "top_p": 0.8,
                "top_k": 40,
                "max_output_tokens": 15000,
            }
        )
        self.prompts_dir = Path("prompts")
        self.prompts_dir.mkdir(exist_ok=True)
        self._create_default_instructions()

    def _create_default_instructions(self):
        """Create default instruction files if they don't exist"""
        defaults = {
                'generator_instructions.txt': """LAZYCOOK MISSION: Minimize human intervention by providing COMPREHENSIVE yet CONCISE responses that require NO follow-up questions or editing.
        
        1. PRIMARY FOCUS: The latest user query is your immediate task.
        2. TOPICAL RELEVANCE: Review the conversation history and uploaded documents in CONTEXT.
        3. STRICT SEPARATION: Only use historical data or document content IF it is directly relevant to the current query.
        4. PIVOT DETECTION: If the user suddenly changes topic (e.g., from Engines to Cooking), IGNORE the previous documents entirely.
        5. LATEST INTENT: If the user query is a pivot or a new topic, treat it as a fresh start, using history only for continuity or style.
        6. NO POLLUTION: Do not mix ideas or terminology from previous files if the current query has shifted scope.
        
        7. COMPREHENSIVE COVERAGE (CRITICAL):
            - Address EVERY part of the user's question - if they ask about A, B, and C, cover ALL THREE thoroughly
            - For multi-part questions, create clear sections for each part
            - Include examples, use cases, and practical applications
            - Anticipate related questions and address them proactively
            - Self-check: "Have I addressed every aspect? Would the user need to ask a follow-up?"
        
        8. RELEVANCE FILTER (ANTI-VERBOSITY):
            - ONLY include content explicitly requested or directly relevant to the query
            - DON'T add unrequested extras (e.g., if they ask about A, B, C - don't add D, E, F)
            - DON'T include tangential information "for completeness" unless asked
            - Self-check: "Did the user actually ask for this detail, or am I adding bloat?"
            - COMPREHENSIVE ≠ VERBOSE: Be thorough on what's asked, skip what's not
        
        9. STRUCTURE & CLARITY:
            - Use clear headings, bullet points, and numbered lists
            - Provide a detailed, well-structured response that is self-contained and accurate
            - Include practical examples where relevant and consider multiple approaches
            - Make it easy to scan and understand
            - Keep explanations concise but complete
        
        10. DATA VISUALIZATION ASSESSMENT:
            - Review the user query for data complexity. If it involves comparisons, trends, proportions, or multi-faceted data that would be easier to understand visually, you MUST explicitly state at the end of your response: "DATA_VISUALIZATION_RECOMMENDED: [Reasoning]".
            - Identify and extract the specific data points that should be charted.
            - Do NOT generate any JSON or <plot_config> blocks yourself; the Optimizer will handle the actual chart creation.
            - Just prepare the ground by formatting the numerical data in a clear markdown table.
        
        11. COMPLETENESS CHECK:
            - Before finishing, verify: Does this response fully satisfy the user's query?
            - Would they need to ask "tell me more" or "what about X"?
            - If yes, expand your response NOW.
        
        12. CONCISENESS CHECK:
            - Before finishing, verify: Is every paragraph necessary?
            - Would removing this section make the answer incomplete?
            - If no, it's bloat - remove it.
            - Goal: Complete but NOT verbose
        
        13. Rate your confidence in this solution (0-1).""",

                'analyzer_instructions.txt': """LAZYCOOK MISSION: Ensure responses are COMPREHENSIVE yet CONCISE - users need NO follow-ups or editing.

            1. COMPLETENESS AUDIT (HIGHEST PRIORITY):
                - Did the generator address EVERY part of the user's query?
                - For multi-part questions (e.g., "What is X and how does Y work?"), verify EACH part is covered
                - Check for missing aspects, topics, or perspectives
                - Flag if the response would require follow-up questions like "tell me more" or "what about Z?"
            
            2. VERBOSITY AUDIT (NEW - CRITICAL):
                - Did the generator add UNREQUESTED content not in the original query?
                - Flag any tangential information, extra examples, or topics NOT explicitly asked for
                - Check: "Would a user need to trim/edit this response?"
                - Penalize bloat: If user asked about A, B, C and got A, B, C, D, E, F - flag D, E, F as bloat
                - COMPREHENSIVE ≠ VERBOSE: Thorough on requested topics, skip unrequested ones
            
            3. CONTEXT POLLUTION CHECK: Verify if the solution accidentally mixed content from unrelated previous files or chat history into the current topic.
            
            4. RELEVANCE AUDIT: If the user query is new, ensure the assistant didn't over-rely on background documents that are no longer relevant.
            
            5. DEPTH & BREADTH (for requested topics only):
                - Is the explanation detailed enough for what was ASKED?
                - Are examples provided where REQUESTED?
                - Are edge cases mentioned if RELEVANT to the query?
            
            6. Identify factual errors, logical inconsistencies, or gaps.
            
            7. Check if the solution addresses the LATEST user query as the priority.
            
            8. Spot "hallucinated connections" between the current topic and old uploaded files.
            
            9. Check if the solution maintains conversational continuity.
            
            10. Suggest improvements: Add missing requested content, REMOVE unrequested bloat.
            
            11. Rate the overall quality (0-1) - be STRICT on completeness AND conciseness.
            
            12. Be thorough but constructive.

    Format your response as JSON:
    {
        "analysis": "Your detailed analysis",
        "errors_found": ["error1", "error2"],
        "gaps_identified": ["gap1", "gap2", "missing_aspect1"],
        "improvements_needed": ["improvement1", "Add coverage of aspect X", "Expand on topic Y"],
        "bloat_identified": ["Unrequested topic D", "Tangential example E", "Extra table F not asked for"],
        "quality_score": 0.75,
        "strengths": ["strength1", "strength2"],
        "recommendations": ["rec1", "rec2"],
        "data_visualization_needed": true, // Set to true if the query involves data that is inherently complex to grasp textually (comparisons, timelines, distributions) and charts would aid "better understanding", regardless of it being a thesis or not.
        "context_adherence": 0.8,
        "continuity_score": 0.7,
        "completeness_score": 0.8, // Rate how completely all REQUESTED aspects were covered (0-1)
        "conciseness_score": 0.85, // NEW: Rate relevance - penalize unrequested content (0-1, lower = more bloat)
        "note": "IMPORTANT: If data_visualization_needed is true, you MUST specifically instruct the Optimizer to 'Create 2-3 informative charts using <plot_config> based on the provided data tables' in the 'improvements_needed' list."
    }""",

                'optimizer_instructions.txt': """LAZYCOOK MISSION: Create responses so COMPLETE yet CONCISE that users need NO follow-ups or editing.

    1. COMPREHENSIVE EXPANSION (TOP PRIORITY):
       - Review the analyzer's "gaps_identified" and "improvements_needed" lists
       - For EVERY missing REQUESTED aspect or topic, add a dedicated section
       - If the query has multiple parts (A, B, C), ensure ALL are covered in depth
       - Add examples, use cases, comparisons ONLY where requested or directly relevant
    
    2. BLOAT REMOVAL (NEW - CRITICAL):
       - Review the analyzer's "bloat_identified" list
       - REMOVE any unrequested content, tangential examples, or extra topics
       - If user asked about A, B, C - don't include D, E, F even if "related"
       - Check "conciseness_score" - if low (<0.80), aggressively trim bloat
       - COMPREHENSIVE ≠ VERBOSE: Thorough on requested topics, skip unrequested ones
    
    3. ERROR CORRECTION:
       - Fix all identified errors and "context pollution" (mixing unrelated topics)
       - Ensure the solution fulfills the current user intent
       - If the history contains unrelated files, DISCARD their influence
    
    4. STRUCTURE & CLARITY:
       - Organize with clear headings for each major REQUESTED topic/aspect
       - Use bullet points, numbered lists, and tables for readability
       - Make the response scannable and well-structured
       - Keep explanations concise but complete
    
    5. DEPTH ENHANCEMENT (for requested topics only):
       - Don't just answer "what" - explain "why" and "how" WHERE ASKED
       - Include edge cases, alternatives, and best practices IF RELEVANT
       - Provide context and background where HELPFUL, not exhaustive
    
    6. COMPLETENESS & CONCISENESS VERIFICATION:
       - Self-check 1: "Would the user need to ask 'tell me more' or 'what about X'?"
       - If yes, expand NOW before finalizing
       - Self-check 2: "Would the user need to trim/edit this response?"
       - If yes, remove bloat NOW before finalizing
       - Goal: Complete on what's asked, concise on delivery
    
    7. DATA VISUALIZATION (PRINCIPAL RESPONSIBILITY):
       - You are the ONLY agent responsible for generating the final <plot_config> blocks.
       - If the Analyzer flagged 'data_visualization_needed': true, OR if you independently perceive that the data (comparisons, trends, ratios) is complex enough that a user would benefit from "better understanding" through visuals, you MUST generate 2-3 relevant charts.
       - Do not just rely on keywords; use your judgment on whether a visual aid makes the information more accessible.
       - Provide a clear, formatted markdown table of the data first.
       - THEN, provide a JSON block wrapped EXACTLY in <plot_config> tags containing a "charts" list.
       - DO NOT use markdown code blocks (```json) inside the <plot_config> tags.
       
       Correct Format (Strictly follow this JSON):
       <plot_config>
       {
         "charts": [
           {
             "type": "bar|line|pie|scatter",
             "title": "Title of Chart",
             "xaxis_title": "Label for X",
             "yaxis_title": "Label for Y",
             "labels": ["Item 1", "Item 2"], 
             "datasets": [
               {
                 "label": "Series Name",
                 "data": [10, 20]
               }
             ]
           }
         ]
       }
       </plot_config>
    
    Format your response as JSON:
{
    "optimized_solution": "Your improved, COMPREHENSIVE yet CONCISE solution here with the <plot_config> block included if needed",
    "changes_made": ["change1", "Added coverage of aspect X", "Expanded section on Y", "Removed unrequested topic Z"],
    "errors_fixed": ["fix1", "fix2"],
    "enhancements": ["enhancement1", "Added examples for requested topic A", "Included comparison of B vs C as asked"],
    "bloat_removed": ["Removed unrequested section D", "Trimmed tangential example E"],
    "confidence": 0.95,
    "context_integration": "How previous conversations were integrated",
    "completeness_improvements": ["Addressed missing aspect X", "Expanded on topic Y"],
    "conciseness_improvements": ["Removed bloat Z", "Trimmed verbose explanation of W"]
}
""",

                'validator_instructions.txt': """1. Check for factual accuracy
    2. Ensure the solution fully addresses the user query
    3. Verify all claims can be supported by the context
    4. Check for logical consistency
    5. Rate confidence (0-1)
    6. Provide specific validation feedback"""
            }


        for filename, content in defaults.items():
                filepath = self.prompts_dir / filename
                filepath.write_text(content.strip(), encoding='utf-8')

    def _load_instructions(self, instruction_file: str) -> str:
        """Load instructions from file, with fallback"""
        try:
            filepath = self.prompts_dir / instruction_file
            if filepath.exists():
                return filepath.read_text(encoding='utf-8').strip()
        except Exception as e:
            logger.error(f"Error loading {instruction_file}: {e}")

        # Return empty string as fallback (prompt will still work)
        return ""
    @log_errors
    async def process(self, user_query: str, context: str = "", previous_iteration: Dict = None, is_new_session: bool = False) -> AgentResponse:
        if self.role == AgentRole.GENERATOR:
            return await self._generate_solution(user_query, context, is_new_session=is_new_session)
        elif self.role == AgentRole.ANALYZER:
            return await self._analyze_solution(user_query, context, previous_iteration)
        elif self.role == AgentRole.OPTIMIZER:
            return await self._optimize_solution(user_query, context, previous_iteration)
        elif self.role == AgentRole.VALIDATOR:
            return await self._validate_solution(user_query, context, previous_iteration)

    @log_errors
    async def _generate_solution(self, user_query: str, context: str, is_new_session: bool = False) -> AgentResponse:
        instruct = self._load_instructions('generator_instructions.txt')
        
        # Determine if we should greet based on session history
        greeting_instruction = ""
        if is_new_session:
            greeting_instruction = "GREETING RULE: This is the very first message in a new session. Start your response with a brief, friendly greeting (e.g., 'Hello Hitarth!')."
        else:
            greeting_instruction = "GREETING RULE: This is a continuation of a conversation. DO NOT greet the user again. Start directly with the answer or solution."

        prompt = f"""
        Role: Solution Generator Agent
        Task: Provide a comprehensive initial solution to the user's query using conversation history.

        {greeting_instruction}

        IMPORTANT: Read the context carefully and refer to previous conversations to understand what the user is asking about.
                   Your name is LAZYCOOK an AI that is specially designed to minimize user interaction by performing 4 tasks alltogether(genrating->analyzing->optimizing->validating) so that user has to do the least work.

        📜 CONTEXT:
        {context}

        👤 USER QUERY: {user_query}

        {"NOTE: This is a professional research/thesis request. You MUST include numerical data tables and corresponding <plot_config> charts to compare models, show history, or illustrate trends as part of your comprehensive response." if any(word in user_query.lower() for word in ['detailed', 'thesis', 'research', 'report', 'analyze', 'paper', 'comparison', 'compare']) or "Uploaded Document" in context else ""}

        Instructions: {instruct}
        
        """
        try:
            response = await self.model.generate_content_async(prompt)
            return AgentResponse(
                agent_role=self.role,
                content=response.text,
                confidence=0.8,
                suggestions=[],
                errors_found=[],
                improvements=[],
                metadata={"context_length": len(context)}
            )
        except Exception as e:
            logger.error(f"Error in _generate_solution: {e}")
            return AgentResponse(
                agent_role=self.role,
                content=f"Error generating solution: {e}",
                confidence=0.5,
                suggestions=[],
                errors_found=[str(e)],
                improvements=[],
                metadata={"error": str(e)}
            )

    # REPLACE the _analyze_solution method in AIAgent class

    @log_errors
    async def _analyze_solution(self, user_query: str, context: str, previous_iteration: Dict) -> AgentResponse:
        generator_response = previous_iteration.get("generator_response", {})
        solution = generator_response.get("content", "")
        instruct = self._load_instructions('analyzer_instructions.txt')
        # Simple heuristic checks instead of LLM judgment
        word_count = len(solution.split())
        has_structure = '\n' in solution

        # Start optimistic
        confidence = 0.90
        errors_found = []
        improvements = []

        # Only downgrade for obvious issues
        if word_count < 50:
            confidence = 0.82
            improvements.append("Response could be more detailed")

        if word_count > 200 and not has_structure:
            confidence -= 0.05
            improvements.append("Could benefit from formatting")



        prompt = f"""
                Role: Critical Analyzer Agent
                Task: Analyze the provided solution for errors, gaps, and improvements, considering conversation history.

                {context}

                Original User Query: {user_query}

                Solution to Analyze:
                {solution}

                Instructions:{instruct}
               
                """

        try:
            response = await self.model.generate_content_async(prompt)
            # AFTER - More robust extraction
            response_text = response.text.strip()

            # Handle code blocks more reliably
            if '```json' in response_text:
                start_idx = response_text.find('```json') + 7
                end_idx = response_text.rfind('```')
                if end_idx > start_idx:
                    response_text = response_text[start_idx:end_idx].strip()
            elif '```' in response_text:
                # Handle generic code blocks
                start_idx = response_text.find('```') + 3
                end_idx = response_text.rfind('```')
                if end_idx > start_idx:
                    response_text = response_text[start_idx:end_idx].strip()

            # Attempt Robust JSON Parsing
            try:
                # Fix common LLM JSON errors: unescaped backslashes in strings
                import re
                fixed_json = re.sub(r'\\(?![\\"/bfnrt] | u[0-9a-fA-F]{4})', r'\\\\', response_text)
                data = json.loads(fixed_json)
                content = data.get("analysis", response_text)
                confidence = data.get("quality_score", 0.7)
                suggestions = data.get("recommendations", [])
                errors_found = data.get("errors_found", [])
                improvements = data.get("improvements_needed", [])
            except Exception as json_err:
                logger.warning(f"JSON parse failed in Analyzer, attempting emergency recovery: {json_err}")
                content = response_text
                confidence = 0.6
                suggestions = []
                errors_found = []
                improvements = []
                data = {"content_was_recovered": True}

            return AgentResponse(
                agent_role=self.role,
                content=content,
                confidence=confidence,
                suggestions=suggestions,
                errors_found=errors_found,
                improvements=improvements,
                metadata=data
            )
        except Exception as e:
            logger.error(f"Analyzer agent error: {e}")
            return AgentResponse(
                agent_role=self.role,
                content="Analysis completed with some limitations.",
                confidence=0.6,
                suggestions=[],
                errors_found=[],
                improvements=[],
                metadata={"error": str(e)}
            )

    @log_errors
    async def _optimize_solution(self, user_query: str, context: str, previous_iteration: Dict) -> AgentResponse:
        generator_response = previous_iteration.get("generator_response", {})
        analyzer_response = previous_iteration.get("analyzer_response", {})
        original_solution = generator_response.get("content", "")
        analysis = analyzer_response.get("content", "")
        errors = analyzer_response.get("errors_found", [])
        improvements = analyzer_response.get("improvements", [])
        instruct = self._load_instructions('optimizer_instructions.txt')
        prompt = f"""
        Role: Solution Optimizer Agent
        Task: Create an improved solution based on analysis feedback and conversation history.

        📜 CONTEXT:
        {context}

        👤 ORIGINAL USER QUERY: {user_query}

        🤖 ORIGINAL SOLUTION:
        {original_solution}

        🔍 ANALYSIS FEEDBACK:
        {analysis}

        ❌ IDENTIFIED ERRORS:
        {errors}

        💡 SUGGESTED IMPROVEMENTS:
        {improvements}

        Instructions:{instruct}
        

        
        """
        try:
            response = await self.model.generate_content_async(prompt)
            response_text = response.text.strip()
            
            # Use more robust extraction logic
            json_text = response_text
            if '```json' in response_text:
                start_idx = response_text.find('```json') + 7
                end_idx = response_text.rfind('```')
                if end_idx > start_idx:
                    json_text = response_text[start_idx:end_idx].strip()
            elif '```' in response_text:
                start_idx = response_text.find('```') + 3
                end_idx = response_text.rfind('```')
                if end_idx > start_idx:
                    json_text = response_text[start_idx:end_idx].strip()
            
            # Attempt Robust JSON Parsing
            try:
                # Fix common LLM JSON errors: unescaped backslashes in strings (often from LaTeX)
                # But don't break valid escapes like \n, \t, etc.
                import re
                
                # Try to fix unescaped backslashes
                # This regex looks for backslashes that are NOT followed by valid JSON escape chars
                fixed_json = re.sub(r'\\(?![\\"/bfnrt] | u[0-9a-fA-F]{4})', r'\\\\', json_text)
                
                data = json.loads(fixed_json)
                content = data.get("optimized_solution", response_text)
                confidence = data.get("confidence", 0.9)
                enhancements = data.get("enhancements", [])
                changes = data.get("changes_made", [])
            except Exception as json_err:
                logger.warning(f"JSON parse failed in Optimizer, attempting emergency recovery: {json_err}")
                
                # SOPHISTICATED RECOVERY:
                # LLMs often fail to escape internal quotes in long strings like "optimized_solution".
                # We'll identify known metadata keys at the END of the response and work backwards.
                content = response_text
                
                # Check for known metadata keys that usually follow optimized_solution
                meta_keys = ["changes_made", "errors_fixed", "enhancements", "confidence", "context_integration"]
                earliest_meta_pos = len(response_text)
                
                for key in meta_keys:
                    # Look for the last occurrence of "<key>":
                    key_pattern = f'"{key}"\s*:'
                    matches = list(re.finditer(key_pattern, response_text))
                    if matches:
                        last_match = matches[-1]
                        if last_match.start() < earliest_meta_pos:
                            earliest_meta_pos = last_match.start()
                
                # If we found metadata keys, the content is everything between "optimized_solution": and the first meta key
                if '"optimized_solution":' in response_text:
                    start_pattern = '"optimized_solution"\s*:\s*"'
                    start_match = re.search(start_pattern, response_text)
                    if start_match:
                        start_pos = start_match.end()
                        # The content ends before the metadata keys start
                        # Need to back up before the quote and comma
                        end_pos = earliest_meta_pos
                        # Walk back from end_pos to find the quote and comma
                        temp_content = response_text[start_pos:end_pos].rstrip()
                        # Strip trailing comma and quote
                        temp_content = re.sub(r',?\s*"$', '', temp_content)
                        content = temp_content
                    else:
                        # Fallback: Just take everything after the header
                        content = response_text.split('"optimized_solution":', 1)[-1]
                
                # Final cleanups for recovered content
                if content.startswith('"'): content = content[1:]
                content = re.sub(r'["\} \n\r,]+$', '', content)
                # Unescape common characters
                content = content.replace('\\n', '\n').replace('\\"', '"').replace('\\\\', '\\')

                confidence = 0.85
                enhancements = ["Recovered from malformed JSON response (Anti-Truncation Logic Applied)"]
                changes = ["Manual content recovery"]
                data = {"content_was_recovered": True, "original_error": str(json_err)}

            return AgentResponse(
                agent_role=self.role,
                content=content,
                confidence=confidence,
                suggestions=enhancements,
                errors_found=[],
                improvements=changes,
                metadata=data
            )
        except Exception as e:
            logger.error(f"Optimizer agent error: {e}")
            return AgentResponse(
                agent_role=self.role,
                content=original_solution,
                confidence=0.7,
                suggestions=[],
                errors_found=[],
                improvements=[],
                metadata={"error": str(e)}
            )

    @log_errors
    async def _validate_solution(self, user_query: str, context: str, previous_iteration: Dict) -> AgentResponse:
        optimizer_response = previous_iteration.get("optimizer_response", {})
        solution = optimizer_response.get("content", "")
        instruct = self._load_instructions('validator_instructions.txt')
        prompt = f"""
        Role: Validator Agent
        Task: Validate the final solution for accuracy and completeness.

        📜 CONTEXT:
        {context}

        👤 USER QUERY: {user_query}

        🤖 SOLUTION TO VALIDATE:
        {solution}

        Instructions:
        1. Check for factual accuracy
        2. Ensure the solution fully addresses the user query
        3. Verify all claims can be supported by the context
        4. Check for logical consistency
        5. Rate confidence (0-1)
        6. Provide specific validation feedback
        """
        try:
            response = await self.model.generate_content_async(prompt)
            return AgentResponse(
                agent_role=self.role,
                content=response.text,
                confidence=0.9,
                suggestions=[],
                errors_found=[],
                improvements=[],
                metadata={"validation": "passed"}
            )
        except Exception as e:
            logger.error(f"Validator agent error: {e}")
            return AgentResponse(
                agent_role=self.role,
                content=f"Validation error: {e}",
                confidence=0.5,
                suggestions=[],
                errors_found=[str(e)],
                improvements=[],
                metadata={"error": str(e)}
            )


# Add this class RIGHT AFTER AIAgent class (around line 800)

class QueryComplexityAnalyzer:
    """Analyzes query complexity to determine which agents are needed"""

    def __init__(self):
        self.simple_patterns = [
            'hello', 'hi', 'hey', 'thanks', 'thank you', 'ok', 'okay',
            'yes', 'no', 'bye', 'goodbye'
        ]

        # ENHANCED: Expanded keywords to catch more comprehensive requests
        self.complex_keywords = [
            'analyze', 'compare', 'evaluate', 'detailed', 'comprehensive',
            'explain why', 'how does', 'difference between', 'pros and cons',
            'step by step', 'in detail', 'thoroughly', 'thorough', 'complete',
            'full', 'all', 'everything', 'research', 'report', 'thesis', 'paper',
            'document', 'create', 'build', 'develop', 'implement', 'design',
            'guide', 'tutorial', 'explain', 'describe', 'discuss', 'overview',
            'breakdown', 'deep dive', 'elaborate', 'extensive', 'in-depth'
        ]

        self.code_keywords = [
            'code', 'python', 'function', 'class', 'debug', 'error',
            'implementation', 'algorithm', 'script', 'program', 'software'
        ]
        
        # NEW: Multi-part indicators
        self.multi_part_indicators = [
            'and', 'also', 'additionally', 'furthermore', 'moreover',
            'as well as', 'along with', 'including', 'plus'
        ]

    def analyze_complexity(self, query: str) -> str:
        query_lower = query.lower().strip()
        word_count = len(query.split())

        # Check for exact matches at word boundaries
        query_words = set(query_lower.split())

        # Simple queries (greetings, short responses) - ONLY for very basic interactions
        if word_count <= 3 and query_words & set(self.simple_patterns):
            return 'simple'

        # ENHANCED: Complex queries detection (prioritize comprehensive responses)
        
        # 1. Explicit complex keywords
        if any(keyword in query_lower for keyword in self.complex_keywords):
            return 'complex'

        # 2. Code-related queries
        if any(keyword in query_lower for keyword in self.code_keywords):
            return 'complex'

        # 3. Multi-part questions (e.g., "What is X and how does Y work?")
        multi_part_count = sum(1 for indicator in self.multi_part_indicators if indicator in query_lower)
        if multi_part_count >= 2 or (multi_part_count >= 1 and word_count > 10):
            return 'complex'

        # 4. Questions with multiple question marks or semicolons
        if query.count('?') > 1 or query.count(';') > 0:
            return 'complex'

        # 5. LOWERED threshold: Longer queries (15+ words instead of 30+)
        if word_count >= 15:
            return 'complex'

        # 6. Very short queries (4-8 words) that aren't greetings
        if 4 <= word_count <= 8:
            return 'medium'

        # DEFAULT: Favor complex over medium for ambiguous cases
        return 'complex' if word_count > 8 else 'medium'

    def get_agent_pipeline(self, complexity: str) -> dict:
        """Return which agents to use based on complexity"""
        pipelines = {
            'simple': {
                'use_generator': True,
                'use_analyzer': False,
                'use_optimizer': False,
                'use_validator': False,
                'max_iterations': 1,
                'description': 'Quick response (1 API call)'
            },
            'medium': {
                'use_generator': True,
                'use_analyzer': True,
                'use_optimizer': True,
                'use_validator': True,  # CHANGED: Added validator for medium queries
                'max_iterations': 1,
                'description': 'Balanced response (4 API calls)'
            },
            'complex': {
                'use_generator': True,
                'use_analyzer': True,
                'use_optimizer': True,
                'use_validator': True,
                'max_iterations': 3,  # CHANGED: Increased from 2 to 3 iterations
                'description': 'Comprehensive response (4-12 API calls, typically 4-8)'
            }
        }
        return pipelines.get(complexity, pipelines['complex'])  # CHANGED: Default to complex instead of medium


# Add this new class BEFORE the MultiAgentSystem class (around line 800)

class QualityMetrics:
    @staticmethod
    def calculate_objective_quality(
            user_query: str,
            response: str,
            analyzer_feedback: AgentResponse,
            context: str = ""
    ) -> Dict[str, Any]:
        """Calculate objective quality metrics with query-type adaptation"""

        try:
            scores = {}

            # STEP 1: Detect query type
            query_lower = user_query.lower()
            query_words = query_lower.split()
            query_length = len(query_words)

            # Factual query indicators
            factual_starters = ['what is', 'who is', 'when did', 'where is', 'what\'s', 'who\'s']
            is_factual = (
                    any(query_lower.startswith(starter) for starter in factual_starters) or
                    query_length <= 8  # Short questions are usually factual
            )

            # Definition query indicators
            is_definition = any(word in query_lower for word in ['define', 'meaning of', 'what does', 'what is'])

            # Complex query indicators
            is_complex = (
                    query_length > 15 or
                    any(word in query_lower for word in
                        ['explain', 'describe', 'analyze', 'compare', 'discuss', 'detail']) or
                    '?' in user_query and user_query.count('?') > 1  # Multi-part
            )

            # Determine query type
            if is_factual and not is_complex:
                query_type = 'factual'
            elif is_definition and not is_complex:
                query_type = 'definition'
            else:
                query_type = 'complex'

            logger.info(f"Query classified as: {query_type}")

            # STEP 2: Calculate metrics based on query type
            response_length = len(response.split())

            # === COMPLETENESS (Type-Aware) ===
            if query_type == 'factual':
                # For factual queries: Did it give a direct answer?
                # Check if response is concise and on-topic
                if 5 <= response_length <= 50:
                    scores['completeness'] = 0.98  # Perfect for factual
                elif response_length < 5:
                    scores['completeness'] = 0.75  # Too short
                else:
                    scores['completeness'] = 0.88  # Verbose but complete

            elif query_type == 'definition':
                # For definitions: 1-3 sentences is ideal
                if 15 <= response_length <= 100:
                    scores['completeness'] = 0.95
                elif response_length < 15:
                    scores['completeness'] = 0.80
                else:
                    scores['completeness'] = 0.85

            else:  # complex
                # Use your existing keyword-based logic
                stopwords = {
                    'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
                    'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had',
                    'do', 'does', 'did', 'will', 'would', 'should', 'could', 'may', 'might',
                    'can', 'i', 'you', 'he', 'she', 'it', 'we', 'they', 'what', 'when',
                    'where', 'who', 'which', 'how', 'why'
                }

                query_words_set = set(user_query.lower().split())
                response_words_set = set(response.lower().split())
                query_keywords = query_words_set - stopwords
                response_keywords = response_words_set - stopwords

                if query_keywords:
                    direct_match = len(query_keywords & response_keywords) / len(query_keywords)
                    keyword_coverage = direct_match

                    # Length appropriate?
                    length_appropriate = response_length >= 50
                    if length_appropriate:
                        keyword_coverage *= 1.15

                    # Has explanation?
                    explanatory_words = {'because', 'since', 'therefore', 'thus', 'means'}
                    has_explanation = bool(response_keywords & explanatory_words)
                    if has_explanation:
                        keyword_coverage *= 1.12

                    scores['completeness'] = max(0.70, min(0.98, keyword_coverage * 1.2))
                else:
                    scores['completeness'] = 0.88

            # === ACCURACY ===
            error_count = len(analyzer_feedback.errors_found)
            if error_count == 0:
                scores['accuracy'] = 0.98
            elif error_count == 1:
                scores['accuracy'] = 0.88
            else:
                scores['accuracy'] = max(0.75, 0.98 - (error_count * 0.08))

            # === LENGTH (Type-Aware) ===
            if query_type == 'factual':
                # Factual: shorter is better
                if 3 <= response_length <= 30:
                    scores['length'] = 0.98
                elif response_length < 3:
                    scores['length'] = 0.70
                else:
                    scores['length'] = max(0.80, 0.98 - ((response_length - 30) / 100))

            elif query_type == 'definition':
                # Definitions: 15-100 words ideal
                if 15 <= response_length <= 100:
                    scores['length'] = 0.98
                elif response_length < 15:
                    scores['length'] = 0.85
                else:
                    scores['length'] = 0.90

            else:  # complex
                # Complex: longer is better (your existing logic)
                if query_length <= 5:
                    ideal_min, ideal_max = 20, 300
                elif query_length <= 15:
                    ideal_min, ideal_max = 50, 600
                else:
                    ideal_min, ideal_max = 100, 1000

                if response_length < ideal_min:
                    scores['length'] = max(0.80, response_length / ideal_min)
                elif response_length > ideal_max:
                    scores['length'] = 0.90
                else:
                    scores['length'] = 0.98

            # === STRUCTURE ===
            structure_indicators = {
                'has_paragraphs': '\n\n' in response,
                'has_lists': any(marker in response for marker in ['1.', '2.', '•', '-', '*']),
                'has_sections': response.count('\n') > 5,
                'has_examples': any(word in response.lower() for word in ['example', 'for instance', 'such as']),
            }
            structure_count = sum(structure_indicators.values())

            # Factual queries don't need structure
            if query_type == 'factual':
                scores['structure'] = 0.95  # Structure irrelevant for facts
            else:
                scores['structure'] = min(0.98, 0.75 + (structure_count * 0.06))

            # === CONTEXT USAGE (Type-Aware) ===
            if query_type == 'factual':
                # Factual queries rarely need context
                scores['context_usage'] = 0.95  # Don't penalize
            else:
                # Complex queries benefit from context
                stopwords = {
                    'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
                    'is', 'are', 'was', 'were', 'be', 'been', 'being'
                }

                if context and len(context) > 100:
                    context_words = set(context.lower().split()[:200])
                    response_words_set = set(response.lower().split())
                    context_keywords = context_words - stopwords
                    response_keywords = response_words_set - stopwords

                    if context_keywords:
                        context_usage = len(context_keywords & response_keywords) / min(len(context_keywords), 50)
                        scores['context_usage'] = min(0.98, context_usage * 1.1)
                    else:
                        scores['context_usage'] = 0.88
                else:
                    scores['context_usage'] = 0.92

            # === POLISH ===
            improvement_count = len(analyzer_feedback.improvements)
            if improvement_count == 0:
                scores['polish'] = 0.98
            elif improvement_count <= 2:
                scores['polish'] = 0.92
            else:
                scores['polish'] = max(0.82, 0.98 - (improvement_count * 0.04))

            # === WEIGHTED SCORE (Type-Aware) ===
            if query_type == 'factual':
                weights = {
                    'completeness': 0.35,  # Most important for facts
                    'accuracy': 0.40,  # Accuracy critical
                    'length': 0.15,  # Brevity matters
                    'structure': 0.0,  # Irrelevant
                    'context_usage': 0.0,  # Irrelevant
                    'polish': 0.10
                }
            elif query_type == 'definition':
                weights = {
                    'completeness': 0.30,
                    'accuracy': 0.35,
                    'length': 0.15,
                    'structure': 0.10,
                    'context_usage': 0.0,
                    'polish': 0.10
                }
            else:  # complex
                weights = {
                    'completeness': 0.20,
                    'accuracy': 0.30,
                    'length': 0.10,
                    'structure': 0.15,
                    'context_usage': 0.10,
                    'polish': 0.15
                }

            overall_score = sum(scores[key] * weights[key] for key in scores.keys())

            # DATA VISUALIZATION PENALTY (MANDATORY FOR THESIS/RESEARCH)
            viz_needed = analyzer_feedback.metadata.get("data_visualization_needed", False)
            viz_present = "<plot_config>" in response or "```json" in response
            
            if viz_needed and not viz_present:
                logger.warning("Applying penalty for missing data visualization mandated by analyzer")
                overall_score *= 0.85 # Significant penalty to force an iteration with graphs
            
            return {
                'overall': overall_score,
                'breakdown': scores,
                'weights': weights,
                'query_type': query_type  # Include for debugging
            }

        except Exception as e:
            logger.error(f"Error in calculate_objective_quality: {e}")
            import traceback
            logger.error(traceback.format_exc())

            return {
                'overall': 0.85,
                'breakdown': {
                    'completeness': 0.85,
                    'accuracy': 0.85,
                    'length': 0.85,
                    'structure': 0.85,
                    'context_usage': 0.85,
                    'polish': 0.85
                },
                'weights': {
                    'completeness': 0.20,
                    'accuracy': 0.30,
                    'length': 0.10,
                    'structure': 0.15,
                    'context_usage': 0.10,
                    'polish': 0.15
                }
            }

    @staticmethod
    def get_quality_tier(score: float) -> str:
        if score >= 0.95:
            return "Excellent"
        elif score >= 0.90:
            return "Very Good"
        elif score >= 0.85:
            return "Good"
        elif score >= 0.75:
            return "Acceptable"
        else:
            return "Needs Improvement"

class DataVisualizer:
    """Handles parsing data patterns and generating plots from agent responses"""
    
    def __init__(self, output_dir: str = "multi_agent_data/plots"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        try:
            sns.set_theme(style="darkgrid")
        except:
            pass # Fallback if seaborn fails

    def extract_plot_configs(self, text: str) -> List[Dict]:
        """Extract multi-plot configurations from response text"""
        configs = []
        try:
            import re
            
            # 1. Extract anything between plot_config tags
            pattern = r'<plot_config>(.*?)</plot_config>'
            matches = re.finditer(pattern, text, re.DOTALL)
            for match in matches:
                try:
                    content = match.group(1).strip()
                    # Strip markdown code blocks if the agent included them anyway
                    if content.startswith('```'):
                        start = content.find('{')
                        end = content.rfind('}')
                        if start != -1 and end != -1:
                            content = content[start:end+1]
                    
                    data = json.loads(content)
                    if isinstance(data, dict):
                        if 'charts' in data:
                            configs.extend(data['charts'])
                        else:
                            configs.append(data)
                    elif isinstance(data, list):
                        configs.extend(data)
                except Exception as e:
                    logger.error(f"Failed to parse plot_config: {e}")

            # 2. Extract JSON blocks that match the 'charts' structure (terminal_plotly style)
            json_pattern = r'```json\s*(\{[\s\S]*?\})\s*```'
            json_matches = re.finditer(json_pattern, text)
            for j_match in json_matches:
                try:
                    data = json.loads(j_match.group(1))
                    if 'charts' in data:
                        configs.extend(data['charts'])
                except Exception:
                    pass # Not a chart JSON, ignore
                    
        except Exception as e:
            logger.error(f"Failed to extract plot data: {e}")
        return configs

    def create_plots(self, configs: List[Dict], session_id: str) -> List[str]:
        """Create multiple plots based on configurations, showing all in one window first if multiple"""
        if not configs:
            return []
            
        saved_paths = []
        
        # Display all plots together in one window for interactive review (like terminal_plotly)
        if len(configs) > 0:
            try:
                import matplotlib
                if sys.platform == 'win32':
                    matplotlib.use('Qt5Agg') # Ensure interactive backend
                import matplotlib.pyplot as plt
                
                num_plots = len(configs)
                fig_total, axes = plt.subplots(num_plots, 1, figsize=(10, 6 * num_plots))
                if num_plots == 1:
                    axes = [axes]
                
                for i, config in enumerate(configs):
                    self._draw_plot_on_ax(config, axes[i])
                
                plt.tight_layout()
                plt.show() # This opens the combined window
            except Exception as e:
                logger.warning(f"Failed to show combined interactive plot: {e}")

        # Still save individual plots and open them as before for records
        for i, config in enumerate(configs):
            path = self._create_single_plot(config, session_id, i)
            if path:
                saved_paths.append(path)
        return saved_paths

    def _draw_plot_on_ax(self, config: Dict, ax):
        """Helper to draw a specific plot configuration on an existing axis"""
        try:
            import seaborn as sns
            import numpy as np
            
            plot_type = config.get('type', 'bar').lower()
            title = config.get('title', 'Data Visualization')
            
            # Robust Data Extraction
            labels = []
            datasets = []
            data_content = config.get('data', config) # Use config as fallback for flat data
            
            # Target fields from config or data_content
            labels = config.get('labels', config.get('x', []))
            datasets = config.get('datasets', [])
            
            if isinstance(data_content, dict):
                if not labels: labels = data_content.get('labels', data_content.get('x', []))
                if not datasets: datasets = data_content.get('datasets', [])
            
            # CASE 1: Tabular data (list of dicts)
            if isinstance(data_content, list) and len(data_content) > 0:
                x_field = config.get('x_field')
                y_field = config.get('y_field')
                if not x_field or not y_field:
                    keys = list(data_content[0].keys())
                    if len(keys) >= 2:
                        x_field = x_field or keys[0]
                        y_field = y_field or keys[1]
                if x_field and y_field:
                    labels = [item.get(x_field, '') for item in data_content]
                    datasets = [{"label": config.get('y_label', y_field), "data": [item.get(y_field, 0) for item in data_content]}]
            
            # Normalize datasets (handle flat data or alternate keys)
            if not datasets and 'data' in config:
                d = config.get('data')
                if isinstance(d, list):
                    datasets = [{"label": config.get('title', 'Data'), "data": d}]
            
            # Ensure datasets have standard 'data' key and extract values
            for ds in datasets:
                if 'values' in ds and 'data' not in ds:
                    ds['data'] = ds['values']
                elif 'y' in ds and 'data' not in ds:
                    ds['data'] = ds['y']

            # FALLBACK: If labels are missing but we have data, generate numeric labels
            if not labels and datasets:
                first_ds_data = datasets[0].get('data', [])
                if first_ds_data:
                    labels = [str(i+1) for i in range(len(first_ds_data))]
            
            # Numeric conversion for labels if they look like numbers (helps with scaling)
            try:
                if labels and all(isinstance(l, (int, float)) or (isinstance(l, str) and l.replace('.','',1).replace('-','',1).isdigit()) for l in labels):
                    labels = [float(l) for l in labels]
            except:
                pass

            x_label = config.get('x_label', config.get('xaxis_title', ''))
            y_label = config.get('y_label', config.get('yaxis_title', ''))
            palette = config.get('palette', 'viridis')
            colors = sns.color_palette(palette, n_colors=max(len(datasets), 1))
            
            if plot_type == 'bar':
                width = 0.8 / len(datasets) if datasets else 0.8
                x = np.arange(len(labels))
                for i, ds in enumerate(datasets):
                    offset = (i - (len(datasets)-1)/2) * width
                    d_values = ds.get('data', [])
                    ax.bar(x + offset, d_values, width, label=ds.get('label', f'Series {i+1}'), color=colors[i])
                ax.set_xticks(x)
                ax.set_xticklabels(labels)
            elif plot_type == 'line':
                for i, ds in enumerate(datasets):
                    d_values = ds.get('data', [])
                    ax.plot(labels, d_values, marker='o', linewidth=2, label=ds.get('label', f'Series {i+1}'), color=colors[i])
            elif plot_type == 'pie':
                if datasets:
                    d_values = datasets[0].get('data', [])
                    ax.pie(d_values, labels=labels, autopct='%1.1f%%', shadow=True, startangle=140, colors=sns.color_palette(palette, len(labels)))
            elif plot_type == 'scatter':
                for i, ds in enumerate(datasets):
                    d_values = ds.get('data', [])
                    ax.scatter(labels, d_values, label=ds.get('label', f'Series {i+1}'), color=colors[i], s=100, alpha=0.7)
            
            ax.set_title(title, fontsize=14, fontweight='bold')
            if x_label: ax.set_xlabel(x_label)
            if y_label: ax.set_ylabel(y_label)
            if plot_type != 'pie': ax.legend(frameon=True, fontsize='small')
            plt.setp(ax.get_xticklabels(), rotation=45 if len(labels) > 5 else 0)
        except Exception as e:
            logger.error(f"Error drawing on axis: {e}")

    def _create_single_plot(self, config: Dict, session_id: str, index: int) -> Optional[str]:
        """Creates and saves a single plot as an image file in the background"""
        try:
            import matplotlib
            matplotlib.use('Agg') # Use non-interactive backend for background saving
            import matplotlib.pyplot as plt
            import seaborn as sns
            
            theme = config.get('theme', 'darkgrid')
            try:
                sns.set_theme(style=theme)
            except:
                plt.style.use('ggplot')
                
            fig, ax = plt.subplots(figsize=(10, 6))
            self._draw_plot_on_ax(config, ax)
            plt.tight_layout()
            
            file_name = f"plot_{session_id}_{index}_{int(time.time())}.png"
            file_path = self.output_dir / file_name
            plt.savefig(file_path, dpi=150, bbox_inches='tight')
            plt.close(fig)
            
            # Automatically open the plot record if on a desktop environment
            try:
                import sys
                import subprocess
                if sys.platform == 'win32':
                    os.startfile(file_path)
                elif sys.platform == 'darwin':
                    subprocess.Popen(['open', str(file_path)])
                else:
                    subprocess.Popen(['xdg-open', str(file_path)])
            except Exception as e:
                logger.warning(f"Could not automatically open plot record: {e}")
                
            return str(file_path)
        except Exception as e:
            logger.error(f"Error creating plot record: {e}")
            try: plt.close() # Ensure resources are freed
            except: pass
            return None


# --- Multi-Agent System ---
class MultiAgentSystem:
    # REPLACE __init__ method in MultiAgentSystem class

    def __init__(self, api_key: str):
        # Lower temperatures for more precise, consistent outputs
        self.generator = AIAgent(api_key, AgentRole.GENERATOR, temperature=0.7)  # Creative but focused
        self.analyzer = AIAgent(api_key, AgentRole.ANALYZER, temperature=0.2)  # Very strict
        self.optimizer = AIAgent(api_key, AgentRole.OPTIMIZER, temperature=0.4)  # Precise refinement
        self.validator = AIAgent(api_key, AgentRole.VALIDATOR, temperature=0.1)  # Extremely strict

        # Add complexity analyzer
        self.complexity_analyzer = QueryComplexityAnalyzer()
        
        # Add Data Visualizer
        self.visualizer = DataVisualizer()

        # NEW: Higher quality settings with more lenient thresholds
        self.max_iterations = 5  # CHANGED: Increased from 4 to 5 for more refinement
        self.quality_threshold = 0.92  # CHANGED: Lowered from 0.95 to 0.92 to allow more iterations
        self.use_validator = True  # Always validate for high quality

        # Statistics tracking
        self.stats = {
            'simple_queries': 0,
            'medium_queries': 0,
            'complex_queries': 0,
            'total_api_calls': 0,
            'quality_scores': [],  # NEW: Track all quality scores
            'iterations_per_query': []  # NEW: Track iteration counts
        }

    # REPLACE process_query method in MultiAgentSystem class

    @log_errors
    async def process_query(self, user_query: str, context: str = "",
                            progress_callback: Optional[Callable] = None,
                            is_new_session: bool = False) -> MultiAgentSession:
        session_id = f"session_{int(time.time())}"

        # Analyze query complexity
        complexity = self.complexity_analyzer.analyze_complexity(user_query)
        pipeline = self.complexity_analyzer.get_agent_pipeline(complexity)

        # Update stats
        self.stats[f'{complexity}_queries'] += 1

        # Log routing decision
        logger.info(f"Query complexity: {complexity} - {pipeline['description']}")

        if progress_callback:
            progress_callback("context", 0,
                              f"📊 Routing: {complexity.upper()} - {pipeline['description']}")

        iterations = []
        current_iteration = 0
        final_response = ""
        quality_score = 0.0
        api_calls_used = 0

        # NEW: Initialize quality metrics
        quality_metrics_calculator = QualityMetrics()

        # Show context being used
        if progress_callback:
            progress_callback("context", 5, f"📜 Using context ({len(context.split()) if context else 0} words)")

        # NEW: Iteration-based quality thresholds (more lenient to allow comprehensive coverage)
        iteration_thresholds = {
            0: 0.92,  # CHANGED: First attempt - good (was 0.95)
            1: 0.89,  # CHANGED: Second attempt - acceptable (was 0.92)
            2: 0.86,  # CHANGED: Third attempt - acceptable (was 0.88)
            3: 0.84,  # CHANGED: Fourth attempt - acceptable (was 0.85)
            4: 0.82   # NEW: Fifth attempt - minimum acceptable
        }

        while current_iteration < pipeline['max_iterations']:
            iteration_data = {"iteration": current_iteration + 1, "timestamp": datetime.now().isoformat()}

            # Generator Agent (ALWAYS RUN)
            if progress_callback:
                progress_callback("generator", 20 + (current_iteration * 25),
                                  "🔧 Generator Agent creating solution...")
            generator_response = await self.generator.process(user_query, context,
                                                               iterations[-1] if iterations else None,
                                                               is_new_session=is_new_session)
            iteration_data["generator_response"] = asdict(generator_response)
            api_calls_used += 1

            # Analyzer Agent (CONDITIONAL)
            if pipeline['use_analyzer']:
                if progress_callback:
                    progress_callback("analyzer", 35 + (current_iteration * 25),
                                      "🔍 Analyzer Agent reviewing...")
                analyzer_response = await self.analyzer.process(user_query, context, iteration_data)
                iteration_data["analyzer_response"] = asdict(analyzer_response)
                api_calls_used += 1
            else:
                analyzer_response = AgentResponse(
                    agent_role=AgentRole.ANALYZER,
                    content="Skipped for simple query",
                    confidence=0.9,
                    suggestions=[],
                    errors_found=[],
                    improvements=[],
                    metadata={"skipped": True}
                )
                iteration_data["analyzer_response"] = asdict(analyzer_response)

            # Optimizer Agent (CONDITIONAL)
            if pipeline['use_optimizer']:
                if progress_callback:
                    progress_callback("optimizer", 50 + (current_iteration * 25),
                                      "⚡ Optimizer Agent refining...")
                optimizer_response = await self.optimizer.process(user_query, context, iteration_data)
                iteration_data["optimizer_response"] = asdict(optimizer_response)
                api_calls_used += 1
            else:
                optimizer_response = generator_response
                iteration_data["optimizer_response"] = asdict(generator_response)

            # Validator Agent (CONDITIONAL)
            if pipeline['use_validator']:
                if progress_callback:
                    progress_callback("validator", 75 + (current_iteration * 25),
                                      "✅ Validator Agent validating...")
                validator_response = await self.validator.process(user_query, context, iteration_data)
                iteration_data["validator_response"] = asdict(validator_response)
                api_calls_used += 1
            else:
                validator_response = AgentResponse(
                    agent_role=AgentRole.VALIDATOR,
                    content="Skipped for efficiency",
                    confidence=0.9,
                    suggestions=[],
                    errors_found=[],
                    improvements=[],
                    metadata={"skipped": True}
                )
                iteration_data["validator_response"] = asdict(validator_response)

            # NEW: Calculate objective quality metrics
            objective_metrics = quality_metrics_calculator.calculate_objective_quality(
                user_query=user_query,
                response=optimizer_response.content,
                analyzer_feedback=analyzer_response,
                context=context
            )

            # ADD THIS SAFETY CHECK
            if objective_metrics is None or 'overall' not in objective_metrics:
                logger.error("Failed to calculate objective metrics, using fallback")
                objective_metrics = {
                    'overall': 0.85,
                    'breakdown': {
                        'completeness': 0.85,
                        'accuracy': 0.85,
                        'length': 0.85,
                        'structure': 0.85,
                        'context_usage': 0.85,
                        'polish': 0.85
                    },
                    'weights': {}
                }

            # NEW: Combine subjective (agent confidence) with objective metrics
            subjective_score = (
                    optimizer_response.confidence * 0.40 +
                    validator_response.confidence * 0.30 +
                    analyzer_response.confidence * 0.30
            )

            # NEW: visualization check - Penalty if analyzer says it's needed but it's missing
            viz_needed = analyzer_response.metadata.get("data_visualization_needed", False)
            viz_present = "<plot_config>" in optimizer_response.content
            
            quality_score = (objective_metrics['overall'] * 0.50 + subjective_score * 0.50)
            
            if viz_needed and not viz_present:
                quality_score *= 0.85 # Significant penalty to force another iteration
                logger.info("⚠ Data visualization needed but missing in optimizer output, downgrading quality score.")

            # Add boost for excellent responses
            word_count = len(optimizer_response.content.split())
            if word_count > 100 and not analyzer_response.errors_found:
                quality_score = min(0.98, quality_score * 1.03)

            # Set reasonable floor
            quality_score = max(0.86, quality_score)

            # Store detailed metrics
            iteration_data['quality_metrics'] = {
                'objective': objective_metrics,
                'subjective': subjective_score,
                'combined': quality_score,
                'tier': quality_metrics_calculator.get_quality_tier(quality_score)
            }

            # Log quality details
            logger.info(
                f"Iteration {current_iteration + 1}: "
                f"Objective={objective_metrics['overall']:.3f}, "
                f"Subjective={subjective_score:.3f}, "
                f"Combined={quality_score:.3f}, "
                f"Tier={quality_metrics_calculator.get_quality_tier(quality_score)}"
            )

            iterations.append(iteration_data)
            final_response = optimizer_response.content

            # Get threshold for current iteration
            current_threshold = iteration_thresholds.get(current_iteration, self.quality_threshold)

            # NEW: More sophisticated exit condition
            can_exit = (
                    quality_score >= current_threshold and  # Meets threshold
                    not analyzer_response.errors_found and  # No errors
                    objective_metrics['breakdown']['accuracy'] >= 0.90  # Good accuracy
            )

            if can_exit:
                logger.info(
                    f"✓ Quality threshold met: {quality_score:.3f} >= {current_threshold:.3f} "
                    f"(iteration {current_iteration + 1})"
                )
                break
            else:
                # FIXED: Only one log statement
                logger.info(
                    f"⚠ Quality threshold not met: {quality_score:.3f} < {current_threshold:.3f} "
                    f"(iteration {current_iteration + 1}, continuing...)"
                )

            current_iteration += 1

        # NEW: Process Data Visualizations
        plot_paths = []
        try:
            configs = self.visualizer.extract_plot_configs(final_response)
            if configs:
                logger.info(f"Found {len(configs)} plot configurations. Generating charts...")
                plot_paths = self.visualizer.create_plots(configs, session_id)
        except Exception as e:
            logger.error(f"Failed to generate plots: {e}")

        # Update total API calls
        self.stats['total_api_calls'] += api_calls_used

        logger.info(
            f"Query completed: {api_calls_used} API calls, "
            f"{len(iterations)} iterations, "
            f"final quality: {quality_score:.3f}, "
            f"plots generated: {len(plot_paths)}"
        )

        session = MultiAgentSession(
            session_id=session_id,
            user_query=user_query,
            iterations=iterations,
            final_response=final_response,
            quality_score=quality_score,
            total_iterations=len(iterations),
            timestamp=datetime.now(),
            context_used=context[:1000] + "..." if len(context) > 1000 else context,
            metadata={"plot_paths": plot_paths, "api_calls": api_calls_used}
        )
        
        # Add plot paths to metadata
        if plot_paths:
            # MultiAgentSession doesn't have a metadata field by default, we'll add it
            if not hasattr(session, 'metadata'):
                session.metadata = {}
            session.metadata['plot_paths'] = plot_paths
            
        return session

    # Add this method to your MultiAgentSystem class (after process_query method, around line 1450)

    @log_errors
    def get_routing_stats(self) -> Dict[str, Any]:
        """Get intelligent routing statistics"""
        total_queries = (
                self.stats['simple_queries'] +
                self.stats['medium_queries'] +
                self.stats['complex_queries']
        )

        if total_queries == 0:
            return {
                "message": "No queries processed yet",
                "total_queries": 0,
                "total_api_calls": 0
            }

        # Calculate API calls that would have been used without routing
        # (assuming all queries would use 4 agents = 4 calls per query)
        baseline_api_calls = total_queries * 4

        # Calculate average API calls per query
        avg_calls = self.stats['total_api_calls'] / total_queries if total_queries > 0 else 0

        # Calculate API calls saved
        api_calls_saved = baseline_api_calls - self.stats['total_api_calls']

        return {
            'total_queries': total_queries,
            'simple_queries': self.stats['simple_queries'],
            'medium_queries': self.stats['medium_queries'],
            'complex_queries': self.stats['complex_queries'],
            'total_api_calls': self.stats['total_api_calls'],
            'avg_calls_per_query': round(avg_calls, 2),
            'api_calls_saved': api_calls_saved,
            'efficiency_percentage': round((1 - avg_calls / 4) * 100, 1) if avg_calls > 0 else 0
        }

# --- Autonomous Assistant ---
class AutonomousMultiAgentAssistant:
    def __init__(self, gemini_api_key: str):
        self.console = Console()
        self.file_manager = TextFileManager()
        self.multi_agent_system = MultiAgentSystem(gemini_api_key)
        self.running = False
        self.task_executor_thread = None
        self._cached_context = {}

    def start(self):
        self.running = True
        self.task_executor_thread = threading.Thread(target=self._task_executor_loop, daemon=True)
        self.task_executor_thread.start()

    def stop(self):
        self.file_manager.cleanup_session_file()
        self.running = False
        if self.task_executor_thread:
            self.task_executor_thread.join(timeout=5)

    def get_cached_context(self, user_id: str) -> str:
        # Always get fresh context to include latest session conversations
        return self.file_manager.get_conversation_context(user_id, 70)

    def clear_cached_context(self, user_id: str):
        """Clear cached context for a user"""
        self._cached_context.pop(user_id, None)

    @log_errors
    async def process_user_message(self, user_id: str, message: str, reset_context: bool = False,
                                   progress_callback: Optional[Callable] = None) -> str:
        if reset_context:
            self.clear_cached_context(user_id)

        # Always get fresh context instead of cached for session conversations
        context = self.file_manager.get_conversation_context(user_id, 70, user_query=message)

        # Debug: Print context being used (remove in production)
        print(f"DEBUG: Using context with {len(context.split())} words")

        # Detect if this is the first message in a new session
        session_convs = self.file_manager.get_session_conversations(user_id)
        is_new_session = len(session_convs) == 0

        multi_agent_session = await self.multi_agent_system.process_query(
            message,
            context,
            progress_callback=progress_callback,
            is_new_session=is_new_session
        )

        conversation_id = f"{user_id}_{int(time.time())}"
        conversation = Conversation(
            id=conversation_id,
            user_id=user_id,
            timestamp=datetime.now(),
            user_message=message,
            ai_response=multi_agent_session.final_response,
            multi_agent_session=multi_agent_session,
            context=context[:2000] + "..." if len(context) > 2000 else context,
            sentiment="neutral",
            topics=[],
            potential_followups=[]
        )
        self.file_manager.save_conversation(conversation)
        await self._analyze_and_create_tasks(conversation)
        return multi_agent_session.final_response, multi_agent_session

    @log_errors
    async def _analyze_and_create_tasks(self, conversation: Conversation):
        suggested_tasks = []
        user_message = conversation.user_message.lower()

        if "smart home" in user_message:
            suggested_tasks = [
                "Research latest smart thermostat models and energy savings",
                "Compare smart lighting systems within 500-800 budget range"
            ]
        elif "sensors" in user_message:
            suggested_tasks = [
                "Research motion sensors for smart lighting automation",
                "Compare ambient light sensors for energy efficiency"
            ]
        elif "python" in user_message or "code" in user_message:
            suggested_tasks = [
                f"Find Python code examples related to: {conversation.user_message[:50]}",
                f"Explain best practices for: {conversation.user_message[:50]}"
            ]
        elif "machine learning" in user_message:
            suggested_tasks = [
                f"Research latest developments in: {conversation.user_message[:50]}",
                f"Find practical applications of: {conversation.user_message[:50]}"
            ]

        for i, task_desc in enumerate(suggested_tasks):
            task_id = f"{conversation.id}_task_{i}"
            schedule_delay = timedelta(minutes=2 + (i * 3))
            task = Task(
                id=task_id,
                conversation_id=conversation.id,
                task_type="research",
                description=task_desc,
                priority=2,
                status="pending",
                created_at=datetime.now(),
                scheduled_for=datetime.now() + schedule_delay,
                metadata={"user_id": conversation.user_id}
            )
            self.file_manager.save_task(task)

    @log_errors
    def _task_executor_loop(self):
        while self.running:
            try:
                pending_tasks = self.file_manager.get_pending_tasks()
                for task in pending_tasks[:2]:
                    asyncio.run(self._execute_task_async(task))
                time.sleep(45)
            except Exception as e:
                logger.error(f"Task executor error: {e}")
                time.sleep(60)

    @log_errors
    async def _execute_task_async(self, task: Task):
        try:
            task.status = "in_progress"
            self.file_manager.save_task(task)

            context = ""
            if task.metadata and 'user_id' in task.metadata:
                context = self.file_manager.get_conversation_context(task.metadata['user_id'], 3)

            session = await self.multi_agent_system.process_query(
                f"Research and provide information about: {task.description}", context
            )
            task.result = session.final_response
            task.status = "completed"
            task.metadata["quality_score"] = session.quality_score
            task.metadata["iterations"] = session.total_iterations
            self.file_manager.save_task(task)
        except Exception as e:
            logger.error(f"Task execution error: {e}")
            task.status = "failed"
            task.result = f"Error: {str(e)}"
            self.file_manager.save_task(task)

    @log_errors
    def get_user_insights(self, user_id: str) -> Dict[str, Any]:
        conversations = self.file_manager.get_recent_conversations(user_id, 70)
        if not conversations:
            return {"message": "No conversation history found"}

        quality_scores = []
        total_iterations = []
        context_usage = []
        for conv in conversations:
            if conv.multi_agent_session:
                quality_scores.append(conv.multi_agent_session.quality_score)
                total_iterations.append(conv.multi_agent_session.total_iterations)
                context_usage.append(len(conv.context) > 100)

        avg_quality = sum(quality_scores) / len(quality_scores) if quality_scores else 0
        avg_iterations = sum(total_iterations) / len(total_iterations) if total_iterations else 0
        context_usage_rate = sum(context_usage) / len(context_usage) if context_usage else 0

        return {
            "total_conversations": len(conversations),
            "multi_agent_sessions": len([c for c in conversations if c.multi_agent_session]),
            "average_quality_score": round(avg_quality, 2),
            "average_iterations": round(avg_iterations, 1),
            "highest_quality": max(quality_scores) if quality_scores else 0,
            "context_usage_rate": round(context_usage_rate * 100, 1),
            "topics": self._extract_topics(conversations),
            "last_interaction": conversations[0].timestamp.isoformat() if conversations else None,
            "conversation_file_status": "Active" if self.file_manager.conversations_file.exists() else "Missing"
        }

    def _extract_topics(self, conversations: List[Conversation]) -> List[str]:
        all_topics = []
        for conv in conversations:
            all_topics.extend(conv.topics)
        topic_counts = {}
        for topic in all_topics:
            topic_counts[topic] = topic_counts.get(topic, 0) + 1
        return sorted(topic_counts.items(), key=lambda x: x[1], reverse=True)[:5]

# --- Rich CLI ---
class RichMultiAgentCLI:
    def __init__(self, api_key: str):
        self.console = Console()  # This line was missing!
        self.assistant = AutonomousMultiAgentAssistant(api_key)
        self.user_id = "user_001"
        self.session_start = datetime.now()

        # Command aliases
        # ALSO: Update the command_aliases dictionary in RichMultiAgentCLI.__init__ (around line 1550)
        # Find this section and ADD the 'routing' aliases:

        self.command_aliases = {
            'q': 'quit', 'exit': 'quit',
            'c': 'chat',
            'i': 'insights',
            't': 'tasks',
            'd': 'docs',
            'dl': 'download',  # ADD THIS
            'save': 'download',  # ADD THIS
            'export': 'download',  # ADD THIS
            'h': 'help',
            'b': 'back',
            'm': 'maintenance',
            'clean': 'maintenance',
            'del': 'maintenance',
            'clear': 'cls',
            'r': 'routing',  # ADD THIS LINE
        }

        # Custom progress bar styles
        self.progress_styles = {
            "context": Style(color="yellow", blink=False, bold=True),
            "generator": Style(color="green", blink=False, bold=True),
            "analyzer": Style(color="yellow", blink=False, bold=True),
            "optimizer": Style(color="blue", blink=False, bold=True),
            "validator": Style(color="cyan", blink=False, bold=True),
            "complete": Style(color="green", blink=True, bold=True),
            "error": Style(color="red", blink=True, bold=True)
        }

    def display_routing_stats(self):
        """Display intelligent routing statistics"""
        progress = self.create_progress()
        task = progress.add_task("[bold cyan]📊 Loading routing statistics...[/bold cyan]", total=100)

        with Live(progress, auto_refresh=True, console=self.console, refresh_per_second=20):
            # Animate to 40%
            current_percent = 0
            while current_percent < 40:
                current_percent += 2
                progress.update(task, completed=current_percent)
                time.sleep(0.03)

            # Get routing stats
            stats = self.assistant.multi_agent_system.get_routing_stats()

            progress.update(task, completed=70,
                            description="[bold cyan]📊 Analyzing routing efficiency...[/bold cyan]")

            current_percent = 70
            while current_percent < 90:
                current_percent += 1
                progress.update(task, completed=current_percent)
                time.sleep(0.02)

            # Complete animation
            while current_percent < 100:
                current_percent += 1
                progress.update(task, completed=current_percent)
                time.sleep(0.02)

            progress.update(task, completed=100, description="[bold green]✅ Statistics loaded![/bold green]")

        if "message" in stats:
            self.console.print(Panel(
                f"[yellow]{stats['message']}[/yellow]",
                title="[bold]📊 Routing Statistics[/bold]",
                border_style="yellow"
            ))
            return

        # Create routing statistics table
        routing_table = Table(
            show_header=True,
            header_style="bold cyan",
            box=ROUNDED,
            title="Intelligent Query Routing Statistics"
        )
        routing_table.add_column("Complexity", style="bold white", width=15)
        routing_table.add_column("Count", justify="right", style="bold magenta", width=10)
        routing_table.add_column("Percentage", justify="right", style="bold green", width=12)
        routing_table.add_column("API Calls", justify="right", style="bold yellow", width=12)

        total = stats['total_queries']

        # Add rows with color coding
        complexities = [
            ('Simple', stats['simple_queries'], 'green', '1 per query'),
            ('Medium', stats['medium_queries'], 'yellow', '3 per query'),
            ('Complex', stats['complex_queries'], 'red', '4-8 per query')
        ]

        for label, count, color, calls in complexities:
            percentage = round((count / total * 100), 1) if total > 0 else 0
            routing_table.add_row(
                f"[{color}]{label}[/{color}]",
                str(count),
                f"{percentage}%",
                calls
            )

        self.console.print(Panel(
            routing_table,
            title="[bold cyan]📊 Query Routing Breakdown[/bold cyan]",
            border_style="cyan",
            padding=(1, 1)
        ))

        # Efficiency summary
        efficiency_text = f"""
               [bold]🎯 Routing Efficiency:[/bold]
               • Total Queries Processed: [bold]{stats['total_queries']}[/bold]
               • Total API Calls Used: [bold]{stats['total_api_calls']}[/bold]
               • Average Calls per Query: [bold]{stats['avg_calls_per_query']}[/bold]
               • API Calls Saved: [bold green]{stats['api_calls_saved']}[/bold green]

               [bold yellow]💡 Cost Optimization:[/bold yellow]
               • Without routing: [dim]{stats['total_queries'] * 4} API calls[/dim]
               • With routing: [bold green]{stats['total_api_calls']} API calls[/bold green]
               • Efficiency gain: [bold green]{round((1 - stats['avg_calls_per_query'] / 4) * 100, 1)}%[/bold green]

               [bold cyan]📈 Distribution:[/bold cyan]
               • Simple queries (fast): {stats['simple_queries']} ({round(stats['simple_queries'] / total * 100, 1) if total > 0 else 0}%)
               • Medium queries (balanced): {stats['medium_queries']} ({round(stats['medium_queries'] / total * 100, 1) if total > 0 else 0}%)
               • Complex queries (thorough): {stats['complex_queries']} ({round(stats['complex_queries'] / total * 100, 1) if total > 0 else 0}%)
               """

        self.console.print(Panel(efficiency_text, border_style="green", padding=(1, 2)))

        # Visual representation
        if total > 0:
            simple_bar = "█" * int(stats['simple_queries'] / total * 20)
            medium_bar = "█" * int(stats['medium_queries'] / total * 20)
            complex_bar = "█" * int(stats['complex_queries'] / total * 20)

            visual_text = f"""
                   [bold]Query Distribution:[/bold]
                   Simple:  [green]{simple_bar}[/green] {stats['simple_queries']}
                   Medium:  [yellow]{medium_bar}[/yellow] {stats['medium_queries']}
                   Complex: [red]{complex_bar}[/red] {stats['complex_queries']}
                   """
            self.console.print(Panel(visual_text, border_style="blue", padding=(1, 2)))



    def display_banner(self):
        """Display the exact ASCII art banner design for LAZYCOOK."""
        # Exact raw ASCII art from the design (multi-line string)
        raw_banner = """
    ╔═══════════════════════════════════════════════════════════════════════════╗
    ║                                                                           ║
    ║  ██╗     ██╗  ███████╗██╗   ██╗ ██████╗ ██████╗  ██████╗ ██╗  ██╗         ║
    ║  ██║    ████║    ███╔╝╚██╗ ██╔╝██╔════╝██╔═══██╗██╔═══██╗██║ ██╔╝         ║
    ║  ██║   ██╔██║   ███╔╝  ╚████╔╝ ██║     ██║   ██║██║   ██║█████╔╝          ║
    ║  ██║  ██╔╝██║  ███╔╝    ╚██╔╝  ██║     ██║   ██║██║   ██║██╔═██╗          ║
    ║  ██████╔╝ ██║ ███████╗   ██║   ╚██████╗╚██████╔╝╚██████╔╝██║  ██╗         ║
    ║  ╚═════╝  ╚═╝ ╚══════╝   ╚═╝    ╚═════╝ ╚═════╝  ╚═════╝ ╚═╝  ╚═╝         ║
    ║                                                                           ║
    ║                         === Let it cook! ===                              ║
    ║                                                                           ║
    ╚═══════════════════════════════════════════════════════════════════════════╝
        """

        # Dedent to remove Python indentation and clean (ensure exact 79-char width)
        banner_lines = textwrap.dedent(raw_banner).strip().split("\n")
        # Verify/enforce alignment: Pad any short lines to 79 chars (your design width)
        design_width = 79
        banner_lines = [line.ljust(design_width) if len(line) < design_width else line for line in banner_lines]

        # Build styled Text object: Bold cyan for art lines, preserve dim markup for subtitle
        banner_text = Text()
        for i, line in enumerate(banner_lines):
            if "[dim]" in line:  # Subtitle line: Keep Rich markup as-is (dim with emojis)
                banner_text.append(line + "\n", style="")  # No extra style; markup handles dim
            else:  # Art/border lines: Bold cyan
                # Replace any existing markup if needed, but apply bold cyan
                clean_line = line.replace("[dim]", "").replace("[/dim]", "")  # Clean if any
                banner_text.append(clean_line + "\n", style="bold cyan")

        # Create a simple panel (no border, as art has its own; just for containment)
        panel = Panel(
            banner_text,
            title="",  # No title to match exact design
            border_style="",  # No border (use art's box)
            padding=(0, 0),  # Zero padding to preserve exact spacing/indentation
            box=box.MINIMAL,  # Invisible/minimal box (no extra lines)
            width=design_width  # Fixed width to match design exactly
        )

        # Center the entire panel (horizontal alignment; vertical middle for balance)
        self.console.print(Align.center(panel, vertical="middle"))

    def display_help(self):
        help_text = """
        [bold magenta] Available Commands:[/bold magenta]
        • [bold green]chat[/bold green]         Start a conversation
        • [bold green]download[/bold green]     Download chat history (NEW!)
        • [bold green]insights[/bold green]     View user interaction statistics
        • [bold green]chat[/bold green]         Start a conversation (type [bold red]back[/bold red] to exit)
        • [bold green]insights[/bold green]     View user interaction statistics
        • [bold green]docs[/bold green]         Manage uploaded documents (view docs)
        • [bold green]maintenance[/bold green]  Cleanup & Upload docs
        • [bold green]tasks[/bold green]        Show autonomous tasks and status
        • [bold green]quality[/bold green]      Display recent quality scores
        • [bold green]routing[/bold green]      View intelligent routing statistics (NEW!)
        • [bold green]agents[/bold green]       Show multi-agent architecture
        • [bold green]context[/bold green]      Preview conversation context
        • [bold green]files[/bold green]        Check data file status
        • [bold green]stats[/bold green]        System performance statistics
        • [bold green]cls[/bold green]          Clear the console screen
        • [bold green]help[/bold green]         Show this help menu
        • [bold red]quit[/bold red]            Exit the application

        [bold yellow] Maintenance Features:[/bold yellow]
        • Delete specific documents or clear all documents
        • Clear conversations with various options (all/recent)
        • Document capacity management (100 docs / 100MB limit)
        • Full system cleanup options

        [bold cyan]Shortcuts:[/bold cyan]
        • Type 'r' for routing stats
        • Type 'c' for chat mode
        • Type 'i' for insights
        • Type 'q' to quit

        [dim]Type 'back' to return to main menu from any command.[/dim]
        [bold yellow] Note:[/bold yellow] The system uses conversation history and uploaded documents
        to provide more accurate and context-aware responses.
        """
        self.console.print(Panel(help_text, border_style="magenta", padding=(1, 2)))

    def create_progress(self, description: str = "Processing...") -> Progress:
        """Create a progress bar with custom animated styling"""
        return Progress(
            TextColumn("[bold cyan]{task.description}", justify="right"),
            BarColumn(
                bar_width=30,
                pulse_style=self.progress_styles["context"],
                complete_style="green",
                finished_style="green"
                # Removed animation_time parameter as it's not supported
            ),
            MofNCompleteColumn(),
            TimeRemainingColumn(),
            TimeElapsedColumn(),
            console=self.console,
            auto_refresh=True,
            refresh_per_second=20  # Smoother animation
        )

    def display_agents(self):
        """Display multi-agent architecture info"""
        agents_info = """
        [bold cyan]🤖 Multi-Agent Architecture:[/bold cyan]
        • [bold green]Generator Agent[/bold green]: Creates initial solutions (temp: 0.8)
        • [bold yellow]Analyzer Agent[/bold yellow]: Reviews for errors (temp: 0.3)  
        • [bold blue]Optimizer Agent[/bold blue]: Refines solutions (temp: 0.5)
        • [bold cyan]Validator Agent[/bold cyan]: Validates accuracy (temp: 0.2)
        """
        self.console.print(Panel(agents_info, border_style="cyan"))

    def display_tasks(self):
        """Display autonomous tasks and their status"""
        progress = self.create_progress()
        task = progress.add_task("[bold cyan]⚙️ Loading tasks...[/bold cyan]", total=100)

        with Live(progress, auto_refresh=True, console=self.console, refresh_per_second=20):
            # Animate to 30%
            current_percent = 0
            while current_percent < 30:
                current_percent += 2
                progress.update(task, completed=current_percent)
                time.sleep(0.03)

            # Get pending and recent tasks
            pending_tasks = self.assistant.file_manager.get_pending_tasks()
            all_tasks_data = self.assistant.file_manager._read_json_file(self.assistant.file_manager.tasks_file)
            progress.update(task, completed=60, description="[bold cyan]⚙️ Analyzing task queue...[/bold cyan]")

            current_percent = 60
            while current_percent < 80:
                current_percent += 1
                progress.update(task, completed=current_percent)
                time.sleep(0.02)

            # Filter recent tasks (last 20)
            recent_tasks = sorted(all_tasks_data, key=lambda x: x.get('created_at', ''), reverse=True)[:20]

            # Create tasks table
            tasks_table = Table(
                show_header=True,
                header_style="bold cyan",
                box=ROUNDED,
                title="Autonomous Tasks Status"
            )
            tasks_table.add_column("ID", style="dim", width=8)
            tasks_table.add_column("Status", style="bold", width=12)
            tasks_table.add_column("Priority", justify="center", style="bold magenta", width=8)
            tasks_table.add_column("Type", style="white", width=12)
            tasks_table.add_column("Description", style="green", width=40)
            tasks_table.add_column("Scheduled", style="dim", width=12)

            # Add tasks to table
            for task_data in recent_tasks:
                status_color = {
                    'pending': '[yellow]⏳ Pending[/yellow]',
                    'in_progress': '[blue]🔄 Running[/blue]',
                    'completed': '[green]✅ Done[/green]',
                    'failed': '[red]❌ Failed[/red]'
                }.get(task_data.get('status', 'unknown'), '[dim]❓ Unknown[/dim]')

                priority = task_data.get('priority', 0)
                priority_display = f"🔥 {priority}" if priority >= 3 else f"📌 {priority}" if priority >= 2 else f"📋 {priority}"

                scheduled_time = task_data.get('scheduled_for', '')
                if scheduled_time:
                    try:
                        scheduled_dt = datetime.fromisoformat(scheduled_time)
                        scheduled_display = scheduled_dt.strftime("%m-%d %H:%M")
                    except:
                        scheduled_display = "Invalid"
                else:
                    scheduled_display = "Not set"

                tasks_table.add_row(
                    task_data.get('id', 'N/A').split('_')[-1][:8],
                    status_color,
                    priority_display,
                    task_data.get('task_type', 'unknown'),
                    (task_data.get('description', 'No description')[:37] + "..." if len(
                        task_data.get('description', '')) > 40 else task_data.get('description', 'No description')),
                    scheduled_display
                )

            # Complete animation
            while current_percent < 100:
                current_percent += 1
                progress.update(task, completed=current_percent)
                time.sleep(0.02)

            progress.update(task, completed=100, description="[bold green]✓ Tasks loaded![/bold green]")

        # Display results
        self.console.print(Panel(
            tasks_table,
            title="[bold cyan]⚙️ Autonomous Tasks[/bold cyan]",
            border_style="cyan",
            padding=(1, 1)
        ))

        # Task summary
        status_counts = {}
        for task_data in all_tasks_data:
            status = task_data.get('status', 'unknown')
            status_counts[status] = status_counts.get(status, 0) + 1

        summary_text = f"""
        [bold]📊 Task Summary:[/bold]
        • Total Tasks: [bold]{len(all_tasks_data)}[/bold]
        • Pending: [bold yellow]{status_counts.get('pending', 0)}[/bold yellow]
        • Completed: [bold green]{status_counts.get('completed', 0)}[/bold green]
        • Failed: [bold red]{status_counts.get('failed', 0)}[/bold red]
        • Active Queue: [bold]{len(pending_tasks)} ready to execute[/bold]
        """
        self.console.print(Panel(summary_text, border_style="blue", padding=(1, 2)))

    def display_quality(self):
        """Display recent quality scores and metrics"""
        progress = self.create_progress()
        task = progress.add_task("[bold cyan]📈 Analyzing quality...[/bold cyan]", total=100)

        with Live(progress, auto_refresh=True, console=self.console, refresh_per_second=20):
            # Animate to 40%
            current_percent = 0
            while current_percent < 40:
                current_percent += 2
                progress.update(task, completed=current_percent)
                time.sleep(0.03)

            # Get recent conversations with multi-agent sessions
            conversations = self.assistant.file_manager.get_recent_conversations(self.user_id, 10)
            progress.update(task, completed=70, description="[bold cyan]📈 Computing quality metrics...[/bold cyan]")

            current_percent = 70
            while current_percent < 90:
                current_percent += 1
                progress.update(task, completed=current_percent)
                time.sleep(0.02)

            # Filter conversations with multi-agent sessions
            quality_sessions = []
            for conv in conversations:
                if conv.multi_agent_session:
                    quality_sessions.append({
                        'timestamp': conv.timestamp,
                        'quality_score': conv.multi_agent_session.quality_score,
                        'iterations': conv.multi_agent_session.total_iterations,
                        'query': conv.user_message[:50] + "..." if len(conv.user_message) > 50 else conv.user_message,
                        'session_id': conv.multi_agent_session.session_id
                    })

            # Complete animation
            while current_percent < 100:
                current_percent += 1
                progress.update(task, completed=current_percent)
                time.sleep(0.02)

            progress.update(task, completed=100, description="[bold green]✓ Quality analysis complete![/bold green]")

        if not quality_sessions:
            self.console.print(Panel(
                "[yellow]No multi-agent sessions found for quality analysis[/yellow]",
                title="[bold]📈 Quality Metrics[/bold]",
                border_style="yellow"
            ))
            return

        # Create quality table
        quality_table = Table(
            show_header=True,
            header_style="bold cyan",
            box=ROUNDED,
            title="Recent Quality Scores"
        )
        quality_table.add_column("Session", style="dim", width=12)
        quality_table.add_column("Quality", justify="center", style="bold", width=10)
        quality_table.add_column("Iterations", justify="center", style="bold magenta", width=10)
        quality_table.add_column("Query", style="green", width=35)
        quality_table.add_column("Timestamp", style="dim", width=12)

        # Add quality sessions to table
        for session in quality_sessions:
            # Color code quality scores
            score = session['quality_score']
            if score >= 0.85:
                quality_display = f"[bold green]{score:.2f}[/bold green] 🔥"
            elif score >= 0.70:
                quality_display = f"[bold yellow]{score:.2f}[/bold yellow] 📈"
            else:
                quality_display = f"[bold red]{score:.2f}[/bold red] ⚠️"

            # Color code iterations
            iterations = session['iterations']
            iter_display = f"[green]{iterations}[/green]" if iterations <= 2 else f"[yellow]{iterations}[/yellow]" if iterations <= 3 else f"[red]{iterations}[/red]"

            quality_table.add_row(
                session['session_id'].split('_')[-1][:8],
                quality_display,
                iter_display,
                session['query'],
                session['timestamp'].strftime("%m-%d %H:%M")
            )

        # Display results
        self.console.print(Panel(
            quality_table,
            title="[bold cyan]📈 Quality Metrics[/bold cyan]",
            border_style="cyan",
            padding=(1, 1)
        ))

        # Quality statistics
        if quality_sessions:
            scores = [s['quality_score'] for s in quality_sessions]
            iterations = [s['iterations'] for s in quality_sessions]

            avg_quality = sum(scores) / len(scores)
            max_quality = max(scores)
            min_quality = min(scores)
            avg_iterations = sum(iterations) / len(iterations)

            quality_trend = "📈 Improving" if len(scores) >= 3 and scores[0] > scores[-1] else "📊 Stable" if len(
                scores) >= 3 else "🔍 Analyzing"

            stats_text = f"""
            [bold]📊 Quality Statistics:[/bold]
            • Average Quality: [bold]{avg_quality:.2f}[/bold] ({quality_trend})
            • Highest Score: [bold green]{max_quality:.2f}[/bold green]
            • Lowest Score: [bold red]{min_quality:.2f}[/bold red]
            • Avg Iterations: [bold]{avg_iterations:.1f}[/bold]
            • Sessions Analyzed: [bold]{len(quality_sessions)}[/bold]
            • Quality Threshold: [bold]0.85[/bold] (system target)
            """
            self.console.print(Panel(stats_text, border_style="green", padding=(1, 2)))

    def display_files(self):
        """Display data file status and information"""
        progress = self.create_progress()
        task = progress.add_task("[bold cyan]📁 Checking file status...[/bold cyan]", total=100)

        with Live(progress, auto_refresh=True, console=self.console, refresh_per_second=20):
            # Animate to 40%
            current_percent = 0
            while current_percent < 40:
                current_percent += 2
                progress.update(task, completed=current_percent)
                time.sleep(0.03)

            # Get file stats
            stats = self.assistant.file_manager.get_storage_stats()
            progress.update(task, completed=70, description="[bold cyan]📁 Analyzing file system...[/bold cyan]")

            current_percent = 70
            while current_percent < 90:
                current_percent += 1
                progress.update(task, completed=current_percent)
                time.sleep(0.02)

            # Create file status table
            files_table = Table(
                show_header=True,
                header_style="bold cyan",
                box=ROUNDED,
                title="Data File Status"
            )
            files_table.add_column("File", style="bold white", width=20)
            files_table.add_column("Status", style="bold", width=10)
            files_table.add_column("Records", justify="right", style="bold magenta", width=10)
            files_table.add_column("Location", style="dim", width=30)

            # File status with color coding
            for file_name, exists in stats['files_exist'].items():
                status = "[green]✓ Active[/green]" if exists else "[red]✗ Missing[/red]"
                record_count = "0"

                if file_name == 'conversations':
                    record_count = str(stats['total_conversations'])
                elif file_name == 'tasks':
                    record_count = str(stats['total_tasks'])
                elif file_name == 'documents':
                    record_count = str(stats['total_documents'])

                files_table.add_row(
                    f"{file_name}.json",
                    status,
                    record_count,
                    f"./multi_agent_data/{file_name}.json"
                )

            # Complete animation
            while current_percent < 100:
                current_percent += 1
                progress.update(task, completed=current_percent)
                time.sleep(0.02)

            progress.update(task, completed=100, description="[bold green]✓ File status loaded![/bold green]")

        # Display results
        self.console.print(Panel(
            files_table,
            title="[bold cyan]📁 Data File Status[/bold cyan]",
            border_style="cyan",
            padding=(1, 1)
        ))

        # Additional file info
        info_text = f"""
        [bold]📊 Storage Summary:[/bold]
        • Data Directory: [bold]./multi_agent_data/[/bold]
        • Total Users: [bold]{len(stats['users'])}[/bold]
        • Oldest Record: [bold]{stats['oldest_conversation'][:10] if stats['oldest_conversation'] != 'none' else 'None'}[/bold]
        • Newest Record: [bold]{stats['newest_conversation'][:10] if stats['newest_conversation'] != 'none' else 'None'}[/bold]
        """
        self.console.print(Panel(info_text, border_style="blue", padding=(1, 2)))

    def display_stats(self):
        """Display system performance statistics"""
        progress = self.create_progress()
        task = progress.add_task("[bold cyan]📊 Computing statistics...[/bold cyan]", total=100)

        with Live(progress, auto_refresh=True, console=self.console, refresh_per_second=20):
            # Animate to 30%
            current_percent = 0
            while current_percent < 30:
                current_percent += 2
                progress.update(task, completed=current_percent)
                time.sleep(0.03)

            # Get various stats
            storage_stats = self.assistant.file_manager.get_storage_stats()
            user_insights = self.assistant.get_user_insights(self.user_id)
            progress.update(task, completed=60, description="[bold cyan]📊 Analyzing performance...[/bold cyan]")

            current_percent = 60
            while current_percent < 80:
                current_percent += 1
                progress.update(task, completed=current_percent)
                time.sleep(0.02)

            # System uptime
            uptime = datetime.now() - self.session_start

            # Create stats table
            stats_table = Table(
                show_header=True,
                header_style="bold cyan",
                box=ROUNDED,
                title="System Performance Statistics"
            )
            stats_table.add_column("Metric", style="bold white", width=25)
            stats_table.add_column("Value", style="bold green", width=20)
            stats_table.add_column("Details", style="dim", width=30)

            # Add rows
            stats_table.add_row("Session Uptime", str(uptime).split('.')[0], "Current session duration")
            stats_table.add_row("Total Conversations", str(storage_stats['total_conversations']), "All users combined")
            stats_table.add_row("User Conversations", str(user_insights.get('total_conversations', 0)),
                                "Your conversations")
            stats_table.add_row("Multi-Agent Sessions", str(user_insights.get('multi_agent_sessions', 0)),
                                "Advanced processing used")
            stats_table.add_row("Avg Quality Score", f"{user_insights.get('average_quality_score', 0):.2f}",
                                "Out of 1.00")
            stats_table.add_row("Avg Iterations", f"{user_insights.get('average_iterations', 0):.1f}",
                                "Per multi-agent session")
            stats_table.add_row("Context Usage", f"{user_insights.get('context_usage_rate', 0)}%",
                                "Sessions using context")
            stats_table.add_row("Total Documents", str(storage_stats['total_documents']), "All uploaded files")
            stats_table.add_row("Active Tasks", str(len(self.assistant.file_manager.get_pending_tasks())),
                                "Autonomous tasks pending")

            # Complete animation
            while current_percent < 100:
                current_percent += 1
                progress.update(task, completed=current_percent)
                time.sleep(0.02)

            progress.update(task, completed=100, description="[bold green]✓ Statistics computed![/bold green]")

        # Display results
        self.console.print(Panel(
            stats_table,
            title="[bold cyan]📊 System Performance Statistics[/bold cyan]",
            border_style="cyan",
            padding=(1, 1)
        ))

        # Performance indicators
        perf_text = f"""
        [bold]🎯 Performance Indicators:[/bold]
        • Quality Trend: [bold]{"🔥 Excellent" if user_insights.get('average_quality_score', 0) > 0.8 else "📈 Good" if user_insights.get('average_quality_score', 0) > 0.6 else "⚠️ Improving"}[/bold]
        • Processing Efficiency: [bold]{"⚡ Fast" if user_insights.get('average_iterations', 0) < 2 else "🔄 Standard" if user_insights.get('average_iterations', 0) < 3 else "🐌 Complex"}[/bold]
        • Memory Usage: [bold]{"💾 {:.1f}MB".format(sum([len(str(storage_stats)) for _ in range(storage_stats['total_conversations'])]) / 1024)}[/bold]
        • Context Utilization: [bold]{"🎯 High" if user_insights.get('context_usage_rate', 0) > 70 else "📊 Medium" if user_insights.get('context_usage_rate', 0) > 30 else "📉 Low"}[/bold]
        """
        self.console.print(Panel(perf_text, border_style="green", padding=(1, 2)))

    def display_response(self, user_message: str, ai_response: str, processing_time: float, context_used: str = ""):
        """Display user message, context, topics, and AI response with enhanced Markdown rendering and alignment."""
        # Clean response: strip technical tags for display
        import re
        ai_response_display = ai_response
        ai_response_display = re.sub(r'<plot_config>.*?</plot_config>', '', ai_response_display, flags=re.DOTALL)
        ai_response_display = re.sub(r'```json\s*\{[\s\S]*?"charts"[\s\S]*?\}\s*```', '', ai_response_display, flags=re.DOTALL)
        ai_response_display = ai_response_display.strip()
        
        # Fetch latest conversation for topics and quality info
        conversations = self.assistant.file_manager.get_recent_conversations(self.user_id, 1)
        quality_info = ""
        topics = []
        if conversations and conversations[0].multi_agent_session:
            session = conversations[0].multi_agent_session
            quality_info = f"Quality: {session.quality_score:.2f} | Iterations: {session.total_iterations} | Time: {processing_time:.1f}s"
            topics = conversations[0].topics  # Extract topics for display

        # 1. User Message Panel (simple blue style)
        user_panel = Panel(
            Text(f"You", style="bold blue") + Text(f": {user_message}", style="italic blue"),
            title="Your Message",
            border_style="blue",  # Simple solid color
            padding=(1, 2),
            box=box.SHADOW if hasattr(box, 'SHADOW') else box.ROUNDED,  # Shadow for depth (fallback to rounded)
            expand=True  # Fill width, wrap text to prevent horizontal scrolling
        )
        self.console.print(user_panel)

        # 2. Context Preview (only if meaningful, with bullets for multi-line)
        if context_used and len(
                context_used.strip()) > 50 and context_used != "No previous conversation history available.":
            # Clean and bullet-ify context preview (first 3-5 lines)
            context_lines = [line.strip() for line in context_used.split("\n") if line.strip()][:5]
            if len(context_lines) > 1:
                context_preview = "\n".join([f"• {line}" for line in context_lines]) + "\n..."
            else:
                context_preview = context_lines[0] + "..."

            context_panel = Panel(
                Text(context_preview, style="dim yellow"),
                title="📜 Context Preview",
                border_style="yellow",  # Simple solid color
                padding=(1, 2),
                expand=True
            )
            self.console.print(context_panel)

        # 3. Main Content: AI Response + Topics in Columns (no horizontal scroll)
        main_content = []

        # AI Response: Render as Markdown for beautiful handling of #, *, \n
        try:
            # Use Rich Markdown to render headers, bullets, etc. (no raw #, *, \n)
            markdown_content = Markdown(ai_response_display)
            # Clean any raw artifacts if Markdown parsing fails partially
            ai_response_clean = ai_response_display.replace("*\n", "• ").replace("# ", "🔹 ")  # Fallback cleaning
            if not markdown_content:  # If Markdown fails, use cleaned text
                markdown_content = Text(ai_response_clean, style="white")

            response_panel = Panel(
                markdown_content,
                title=Text("🤖 Assistant Response", style="bold green"),
                border_style="green",  # Simple solid color
                padding=(1, 2),
                box=box.SHADOW if hasattr(box, 'SHADOW') else box.ROUNDED,
                expand=True
            )
            main_content.append(response_panel)
        except Exception as e:
            # Fallback: Plain text with basic cleaning (remove raw #, convert * to bullets, handle \n)
            import logging  # Ensure logger is available (or use print for debug)
            logging.getLogger(__name__).warning(f"Markdown rendering failed: {e}")  # Log for debugging
            cleaned_response = ai_response_display.replace("\n\n", "\n").replace("# ",
                                                                         "🔹 ")  # Headers to icons, clean extra newlines
            cleaned_response = cleaned_response.replace("* ", "• ").replace("- ", "• ")  # Bullets without raw *
            cleaned_response = "\n".join(
                [line.strip() for line in cleaned_response.split("\n") if line.strip()])  # Remove empty lines
            response_panel = Panel(
                Text(cleaned_response, style="white", justify="left"),  # Justify for left-alignment and wrapping
                title=Text("🤖 Assistant Response", style="bold green"),
                border_style="green",  # Simple solid color
                padding=(1, 2),
                expand=True
            )
            main_content.append(response_panel)

        # Topics Sidebar: Bold bullets if topics exist (narrow to avoid scrolling)
        topics_content = ""
        if topics:
            topics_bullets = "\n".join(
                [f"• [bold cyan]{topic}[/bold cyan]" for topic in topics[:5]])  # Limit to 5, bold
            topics_content = Panel(
                Text(topics_bullets, style="white"),
                title=Text("🏷️ Key Topics", style="bold magenta"),
                border_style="magenta",  # Simple solid color
                padding=(1, 1),
                width=25  # Fixed narrow width for sidebar (prevents overflow)
            )

        # Layout: Columns for response (wide) + topics (narrow sidebar) - expands to fit terminal
        if topics_content:
            columns_layout = Columns([response_panel, topics_content], padding=(0, 1), expand=True)
        else:
            columns_layout = response_panel  # No sidebar if no topics

        self.console.print(columns_layout)

        # 4. Quality Footer (badge-style, simple cyan)
        if quality_info:
            # Get latest conversation for detailed metrics
            conversations = self.assistant.file_manager.get_recent_conversations(self.user_id, 1)
            if conversations and conversations[0].multi_agent_session:
                session = conversations[0].multi_agent_session

                # Check if quality metrics exist in iterations
                if session.iterations and 'quality_metrics' in session.iterations[-1]:
                    quality_metrics = session.iterations[-1]['quality_metrics']
                    self.display_quality_breakdown(quality_metrics)

            # Simple quality footer
            quality_panel = Panel(
                Text(quality_info, style="bold dim"),
                title="📊 Processing Metrics",
                border_style="cyan",
                padding=(0, 1),
                expand=True
            )
            self.console.print(quality_panel)

    def display_quality_breakdown(self, quality_metrics: Dict[str, Any]):
        """Display detailed quality metrics breakdown"""

        if not quality_metrics or 'objective' not in quality_metrics:
            return

        objective = quality_metrics['objective']
        breakdown = objective.get('breakdown', {})

        # Create quality breakdown table
        quality_table = Table(
            show_header=True,
            header_style="bold cyan",
            box=ROUNDED,
            title="Quality Metrics Breakdown"
        )
        quality_table.add_column("Metric", style="bold white", width=18)
        quality_table.add_column("Score", justify="center", style="bold", width=8)
        quality_table.add_column("Weight", justify="center", style="dim", width=8)
        quality_table.add_column("Status", style="bold", width=15)

        # Color coding for scores
        def get_score_color(score: float) -> str:
            if score >= 0.95:
                return "green"
            elif score >= 0.90:
                return "cyan"
            elif score >= 0.85:
                return "yellow"
            else:
                return "red"

        def get_status_icon(score: float) -> str:
            if score >= 0.95:
                return "🔥 Excellent"
            elif score >= 0.90:
                return "✅ Very Good"
            elif score >= 0.85:
                return "📈 Good"
            else:
                return "⚠️ Needs Work"

        weights = objective.get('weights', {})

        # Add rows for each metric
        for metric, score in breakdown.items():
            color = get_score_color(score)
            weight = weights.get(metric, 0)
            quality_table.add_row(
                metric.replace('_', ' ').title(),
                f"[{color}]{score:.2f}[/{color}]",
                f"{weight:.0%}",
                get_status_icon(score)
            )

        # Overall score row
        overall = objective.get('overall', 0)
        overall_color = get_score_color(overall)
        quality_table.add_row(
            "[bold]OVERALL[/bold]",
            f"[bold {overall_color}]{overall:.3f}[/bold {overall_color}]",
            "100%",
            f"[bold]{quality_metrics.get('tier', 'Unknown')}[/bold]"
        )

        self.console.print(Panel(
            quality_table,
            title="[bold cyan]📊 Quality Analysis[/bold cyan]",
            border_style="cyan",
            padding=(1, 1)
        ))

        # Combined score info
        combined_info = f"""
        [bold]Scoring Method:[/bold]
        • Objective Metrics: [bold cyan]{objective['overall']:.3f}[/bold cyan] (40% weight)
        • Agent Confidence: [bold yellow]{quality_metrics.get('subjective', 0):.3f}[/bold yellow] (60% weight)
        • Combined Score: [bold green]{quality_metrics.get('combined', 0):.3f}[/bold green]
        • Quality Tier: [bold]{quality_metrics.get('tier', 'Unknown')}[/bold]
        """

        self.console.print(Panel(combined_info, border_style="blue", padding=(1, 2)))

    def show_agent_progress(self, message: str):
        """Show detailed agent processing with animated progress bars"""
        with Live(auto_refresh=True, console=self.console, refresh_per_second=20) as live:
            progress = self.create_progress()

            # Main task
            main_task = progress.add_task("[bold cyan]🔄 Processing your request...[/bold cyan]", total=100)

            # Agent processing stages with animations
            stages = [
                ("context", 5, "[bold yellow]📜 Analyzing context and documents...[/bold yellow]", "yellow"),
                ("generator", 30, "[bold green]🔧 Generator Agent creating solution...[/bold green]", "green"),
                ("analyzer", 50, "[bold yellow]🔍 Analyzer Agent reviewing for errors...[/bold yellow]", "yellow"),
                ("optimizer", 70, "[bold blue]⚡ Optimizer Agent refining solution...[/bold blue]", "blue"),
                ("validator", 90, "[bold cyan]✅ Validator Agent validating accuracy...[/bold cyan]", "cyan"),
                ("complete", 100, "[bold green]✓ All agents completed processing![/bold green]", "green")
            ]

            for stage, percent, description, color in stages:
                # Update progress bar style
                progress.columns[1].pulse_style = self.progress_styles.get(stage, self.progress_styles["context"])
                progress.columns[1].complete_style = color

                # Animate progress
                current_percent = progress.tasks[main_task].completed or 0
                while current_percent < percent:
                    current_percent += 1
                    progress.update(main_task, completed=current_percent)
                    live.update(progress)
                    time.sleep(0.03)  # Smooth animation

                progress.update(main_task, completed=percent, description=description)
                live.update(progress)
                time.sleep(0.2)  # Pause at each stage

            live.update(progress)

    async def process_message_with_progress(self, message: str) -> str:
        """Process message with detailed animated agent progress tracking"""
        progress = self.create_progress()
        main_task = progress.add_task("[bold cyan]🔄 Processing your request...[/bold cyan]", total=100)

        with Live(progress, auto_refresh=True, console=self.console, refresh_per_second=20):
            # Agent processing stages
            agent_stages = {
                "context": (5, "[bold yellow]📜 Analyzing context...[/bold yellow]", "yellow"),
                "generator": (30, "[bold green]🔧 Generator Agent working...[/bold green]", "green"),
                "analyzer": (50, "[bold yellow]🔍 Analyzer Agent reviewing...[/bold yellow]", "yellow"),
                "optimizer": (70, "[bold blue]⚡ Optimizer Agent refining...[/bold blue]", "blue"),
                "validator": (90, "[bold cyan]✅ Validator Agent validating...[/bold cyan]", "cyan")
            }

            def progress_callback(stage: str, percent: int, description: str = ""):
                if stage in agent_stages:
                    stage_percent, stage_desc, color = agent_stages[stage]

                    # Animate to target percentage - start from current progress
                    current_task = progress.tasks[main_task]
                    current_percent = int(
                        current_task.completed) if current_task.completed else 0  # Get current progress as int
                    while current_percent < stage_percent:
                        current_percent += 1
                        progress.update(main_task, completed=current_percent)
                        time.sleep(0.02)

                    # Update style and description
                    progress.columns[1].pulse_style = self.progress_styles.get(stage, self.progress_styles["context"])
                    progress.columns[1].complete_style = color
                    progress.update(main_task, completed=stage_percent,
                                    description=stage_desc if not description else description)

            start_time = time.time()
            response, session = await self.assistant.process_user_message(
                self.user_id,
                message,
                reset_context=False,
                progress_callback=progress_callback
            )

            # Final animation to 100%
            current_percent = int(progress.tasks[main_task].completed)  # Get current progress as int
            while current_percent < 100:
                current_percent += 1
                progress.update(main_task, completed=current_percent)
                time.sleep(0.02)

            progress.update(main_task, completed=100, description="[bold green]✓ Processing complete![/bold green]")
            processing_time = time.time() - start_time

            # Add a newline after the Live context ends to separate from next output
        self.console.print()  # This adds the needed separation

        # Get context used for display
        context = self.assistant.get_cached_context(self.user_id)

        # Display the response with context
        self.display_response(message, response, processing_time, context)

        # Notify about generated plots
        if hasattr(session, 'metadata') and session.metadata and 'plot_paths' in session.metadata:
            plot_paths = session.metadata['plot_paths']
            if plot_paths:
                self.console.print()
                plot_info = "\n".join([f"• [bold cyan]{Path(p).name}[/bold cyan]" for p in plot_paths])
                self.console.print(Panel(
                    f"[bold green]📊 {len(plot_paths)} Data Visualization(s) generated![/bold green]\n\n"
                    f"Files saved to:\n{plot_info}\n\n"
                    f"[dim]You can find these in the 'multi_agent_data/plots' directory.[/dim]\n"
                    f"[bold yellow]Tip: You can open these files to view the high-quality charts.[/bold yellow]",
                    title="[bold green]📈 Visualization Ready[/bold green]",
                    border_style="green",
                    padding=(1, 2)
                ))
                
                # Proactive: ask to open on Windows if it's the first time or just tell them how
                if os.name == 'nt':
                     self.console.print(f"[dim]To view, run: [bold]start {plot_paths[0]}[/bold][/dim]")

        return response

    def display_insights(self):
        """Display insights with animated progress tracking"""
        progress = self.create_progress()
        task = progress.add_task("[bold cyan]📊 Loading insights...[/bold cyan]", total=100)

        with Live(progress, auto_refresh=True, console=self.console, refresh_per_second=20):
            # Animate to 30%
            current_percent = 0
            while current_percent < 30:
                current_percent += 2
                progress.update(task, completed=current_percent)
                time.sleep(0.03)

            insights = self.assistant.get_user_insights(self.user_id)
            progress.update(task, completed=60, description="[bold cyan]📊 Analyzing conversation data...[/bold cyan]")
            # Animate to 80%
            current_percent = 60
            while current_percent < 80:
                current_percent += 1
                progress.update(task, completed=current_percent)
                time.sleep(0.02)

            if "message" in insights:
                self.console.print(Panel(f"[yellow]{insights['message']}[/yellow]", title="[bold]📊 Insights[/bold]",
                                         border_style="yellow"))
            else:
                stats_text = f"""
                [bold]📈 Session Statistics:[/bold]
                • Total Conversations: [bold]{insights.get('total_conversations', 0)}[/bold]
                • Multi-Agent Sessions: [bold]{insights.get('multi_agent_sessions', 0)}[/bold]
                • Avg. Quality Score: [bold]{insights.get('average_quality_score', 0):.2f}[/bold]
                • Avg. Iterations: [bold]{insights.get('average_iterations', 0):.1f}[/bold]
                • Context Usage Rate: [bold]{insights.get('context_usage_rate', 0)}%[/bold]
                • Session Duration: [bold]{datetime.now() - self.session_start}[/bold]
                """
                self.console.print(
                    Panel(stats_text, title="[bold cyan]📊 Insights[/bold cyan]", border_style="cyan", padding=(1, 2)))

                if insights.get('topics'):
                    topics_text = "\n".join(
                        [f"• [bold]{topic}[/bold]: {count}" for topic, count in insights['topics'][:5]])
                    self.console.print(Panel(
                        f"[bold magenta]🏷️ Topics:[/bold magenta]\n{topics_text}",
                        title="[bold magenta]🏷️ Discussion Topics[/bold magenta]",
                        border_style="magenta",
                        padding=(1, 2)
                    ))

            # Complete the progress
            while current_percent < 100:
                current_percent += 1
                progress.update(task, completed=current_percent)
                time.sleep(0.02)

            progress.update(task, completed=100, description="[bold green]✓ Insights loaded![/bold green]")

    def display_context(self):
        """Display context preview with animated progress"""
        progress = self.create_progress()
        task = progress.add_task("[bold cyan]📜 Loading context...[/bold cyan]", total=100)

        with Live(progress, auto_refresh=True, console=self.console, refresh_per_second=20):
            # Animate to 40%
            current_percent = 0
            while current_percent < 40:
                current_percent += 2
                progress.update(task, completed=current_percent)
                time.sleep(0.03)

            context = self.assistant.file_manager.get_conversation_context(self.user_id, 5)
            progress.update(task, completed=70, description="[bold cyan]📜 Compiling context preview...[/bold cyan]")
            current_percent = 70  # Update current_percent to match the progress

            # Animate to 90%
            while current_percent < 90:
                current_percent += 1
                progress.update(task, completed=current_percent)
                time.sleep(0.02)

            if not context or context == "No previous conversation history available.":
                self.console.print(Panel(
                    "[yellow]No conversation context available[/yellow]",
                    title="[bold]📄 Context Preview[/bold]",
                    border_style="yellow"
                ))
            else:
                # Enhanced context visualization
                context_lines = context.split('\n')
                context_preview = []

                for line in context_lines:
                    if line.startswith('--- Conversation'):
                        context_preview.append(f"\n[bold magenta]{line}[/bold magenta]")
                    elif line.startswith('👤 USER:'):
                        context_preview.append(f"[bold blue]{line}[/bold blue]")
                    elif line.startswith('🤖 ASSISTANT:'):
                        context_preview.append(f"[bold green]{line}[/bold green]")
                    elif line.startswith('--- Document'):
                        context_preview.append(f"\n[bold cyan]{line}[/bold cyan]")
                    else:
                        context_preview.append(f"[dim]{line}[/dim]")

                preview_text = "".join(context_preview[:20]) + "\n..." if len(context_preview) > 20 else "".join(
                    context_preview)

                context_panel = Panel(
                    preview_text,
                    title=f"[bold cyan]📄 Context Preview ({len(context.split())} words)[/bold cyan]",
                    border_style="cyan",
                    padding=(1, 2)
                )
                self.console.print(context_panel)

                # Show statistics
                stats = f"""
                [bold]Context Statistics:[/bold]
                • Total words: {len(context.split())}
                • Conversations: {len([line for line in context_lines if line.startswith('--- Conversation')])}
                • Documents referenced: {len([line for line in context_lines if line.startswith('--- Document')])}
                """
                self.console.print(
                    Panel(stats, title="[bold]📊 Context Stats[/bold]", border_style="blue", padding=(1, 1)))

            # Complete the progress
            while current_percent < 100:
                current_percent += 1
                progress.update(task, completed=current_percent)
                time.sleep(0.02)

            progress.update(task, completed=100, description="[bold green]✓ Context loaded![/bold green]")

    def upload_document(self):
        """Document upload with capacity checking and animated progress tracking"""
        # Check capacity first
        capacity_info = self.assistant.file_manager.get_documents_capacity_info(self.user_id)

        if capacity_info['at_document_limit']:
            self.console.print(Panel(
                f"[red]❌ Document limit reached![/red]\n"
                f"You have {capacity_info['total_documents']}/{capacity_info['max_documents']} documents.\n"
                f"Please delete some documents first.",
                title="[bold red]Upload Failed[/bold red]",
                border_style="red"
            ))
            return None

        if capacity_info['at_size_limit']:
            self.console.print(Panel(
                f"[red]❌ Storage limit reached![/red]\n"
                f"You have used {capacity_info['total_size_mb']:.1f}MB/{capacity_info['max_size_mb']}MB.\n"
                f"Please delete some documents first.",
                title="[bold red]Upload Failed[/bold red]",
                border_style="red"
            ))
            return None

        file_path = Prompt.ask("[bold cyan]📁 Enter file path to upload[/bold cyan] (or 'back' to cancel)")
        if file_path.strip().lower() in ['back', 'exit', 'cancel']:
            return None

        file_path = Path(file_path.strip().strip('"\''))
        if not file_path.exists():
            self.console.print(f"[red]File not found: {file_path}[/red]")
            return None

        file_size = file_path.stat().st_size

        # Check if this file would exceed limits
        if capacity_info['total_documents'] + 1 > capacity_info['max_documents']:
            self.console.print("[red]This upload would exceed the document limit[/red]")
            return None

        if capacity_info['total_size_bytes'] + file_size > self.assistant.file_manager.max_storage_per_user:
            needed_mb = round((capacity_info['total_size_bytes'] + file_size) / (1024 * 1024), 2)
            self.console.print(
                f"[red]This upload would exceed storage limit ({needed_mb}MB > {capacity_info['max_size_mb']}MB)[/red]")
            return None

        if file_size > 5 * 1024 * 1024:
            self.console.print("[red]File too large. Maximum size is 5MB per file[/red]")
            return None

        progress = self.create_progress()
        task = progress.add_task("[bold cyan]📁 Preparing upload...[/bold cyan]", total=100)

        with Live(progress, auto_refresh=True, console=self.console, refresh_per_second=20):
            # Simulate file analysis
            current_percent = 0
            while current_percent < 20:
                current_percent += 2
                progress.update(task, completed=current_percent,
                                description=f"[bold cyan]📁 Analyzing file ({decimal(file_size)})...[/bold cyan]")
                time.sleep(0.05)

            # Simulate processing
            progress.update(task, completed=60, description="[bold cyan]💾 Uploading and indexing...[/bold cyan]")
            current_percent = 60
            try:
                document = self.assistant.file_manager.process_uploaded_file(str(file_path), self.user_id)

                # Simulate final processing
                while current_percent < 95:
                    current_percent += 1
                    progress.update(task, completed=current_percent)
                    time.sleep(0.02)

                if document:
                    progress.update(task, completed=100,
                                    description="[bold green]✅ Document uploaded successfully![/bold green]")

                    # Show updated capacity info
                    new_capacity = self.assistant.file_manager.get_documents_capacity_info(self.user_id)

                    info_table = Table(show_header=False, box=None, padding=(0, 2))
                    info_table.add_column(style="bold cyan", width=15)
                    info_table.add_column(style="white")
                    info_table.add_row("Filename:", document.filename)
                    info_table.add_row("Type:", document.file_type)
                    info_table.add_row("Size:", decimal(document.file_size))
                    info_table.add_row("Documents:",
                                       f"{new_capacity['total_documents']}/{new_capacity['max_documents']}")
                    info_table.add_row("Storage:",
                                       f"{new_capacity['total_size_mb']:.1f}MB/{new_capacity['max_size_mb']}MB")
                    info_table.add_row("Preview:", document.content[:100] + "..." if len(
                        document.content) > 100 else document.content)

                    self.console.print(
                        Panel(info_table, title="[bold green]📄 Document Info[/bold green]", border_style="green",
                              padding=(1, 2)))
                else:
                    progress.update(task, completed=100,
                                    description="[bold red]❌ Failed to process document[/bold red]")
                    self.console.print("[red]Failed to process document[/red]")

                return document
            except Exception as e:
                progress.update(task, completed=100, description=f"[bold red]❌ Error: {str(e)}[/bold red]")
                self.console.print(f"[red]Error processing document: {str(e)}[/red]")
                return None

    def display_documents(self):
        """Display documents with animated progress"""
        progress = self.create_progress()
        task = progress.add_task("[bold cyan]📁 Loading documents...[/bold cyan]", total=100)

        with Live(progress, auto_refresh=True, console=self.console, refresh_per_second=20):
            # Animate to 30%
            current_percent = 0
            while current_percent < 30:
                current_percent += 2
                progress.update(task, completed=current_percent)
                time.sleep(0.03)

            documents = self.assistant.file_manager.get_user_documents(self.user_id)
            progress.update(task, completed=60, description="[bold cyan]📁 Compiling document list...[/bold cyan]")

            # Animate to 80%
            while current_percent < 80:
                current_percent += 1
                progress.update(task, completed=current_percent)
                time.sleep(0.02)

            if not documents:
                self.console.print(
                    Panel("[yellow]No documents uploaded yet[/yellow]", title="[bold]📄 Your Documents[/bold]",
                          border_style="yellow"))
            else:
                docs_table = Table(
                    show_header=True,
                    header_style="bold cyan",
                    box=ROUNDED,
                    expand=True,
                    padding=(0, 1),
                    title="Your Documents"
                )
                docs_table.add_column("ID", style="dim", width=8)
                docs_table.add_column("Filename", style="bold cyan", width=25)
                docs_table.add_column("Type", style="white", width=12)
                docs_table.add_column("Size", justify="right", style="bold magenta", width=10)
                docs_table.add_column("Uploaded", style="dim", width=12)
                docs_table.add_column("Preview", style="green", width=30)

                for doc in documents:
                    preview = doc.content[:50] + "..." if len(doc.content) > 50 else doc.content
                    preview = preview.replace('\n', ' ').replace('\r', ' ')
                    docs_table.add_row(
                        doc.id.split('_')[-1],
                        doc.filename,
                        doc.file_type.split('/')[-1],
                        decimal(doc.file_size),
                        doc.upload_time.strftime("%m-%d %H:%M"),
                        preview
                    )
                self.console.print(Panel(
                    docs_table,
                    title=f"[bold cyan]📄 Your Documents ({len(documents)} total)[/bold cyan]",
                    border_style="cyan",
                    padding=(1, 1)
                ))

            # Complete the progress
            while current_percent < 100:
                current_percent += 1
                progress.update(task, completed=current_percent)
                time.sleep(0.02)

            progress.update(task, completed=100, description="[bold green]✓ Documents loaded![/bold green]")

        # Add these methods to the RichMultiAgentCLI class (around line 1200, after existing display methods)

    def delete_specific_document(self):
        """Delete a specific document with interactive selection"""
        progress = self.create_progress()
        task = progress.add_task("[bold cyan]📄 Loading documents...[/bold cyan]", total=100)

        with Live(progress, auto_refresh=True, console=self.console, refresh_per_second=20):
            # Animate to 40%
            current_percent = 0
            while current_percent < 40:
                current_percent += 2
                progress.update(task, completed=current_percent)
                time.sleep(0.03)

            documents = self.assistant.file_manager.get_user_documents(self.user_id)

            # Complete animation
            while current_percent < 100:
                current_percent += 2
                progress.update(task, completed=current_percent)
                time.sleep(0.02)

            progress.update(task, completed=100, description="[bold green]✅ Documents loaded![/bold green]")

        if not documents:
            self.console.print(Panel(
                "[yellow]No documents to delete[/yellow]",
                title="[bold]📄 Delete Document[/bold]",
                border_style="yellow"
            ))
            return

        # Display documents for selection
        docs_table = Table(
            show_header=True,
            header_style="bold cyan",
            box=ROUNDED,
            title="Select Document to Delete"
        )
        docs_table.add_column("#", style="bold magenta", width=3)
        docs_table.add_column("Filename", style="bold cyan", width=25)
        docs_table.add_column("Type", style="white", width=12)
        docs_table.add_column("Size", justify="right", style="bold magenta", width=10)
        docs_table.add_column("Uploaded", style="dim", width=12)

        for i, doc in enumerate(documents, 1):
            docs_table.add_row(
                str(i),
                doc.filename,
                doc.file_type.split('/')[-1],
                decimal(doc.file_size),
                doc.upload_time.strftime("%m-%d %H:%M")
            )

        self.console.print(Panel(
            docs_table,
            title="[bold red]🗑️ Delete Document[/bold red]",
            border_style="red",
            padding=(1, 1)
        ))

        try:
            choice = Prompt.ask(
                f"[bold cyan]Enter document number (1-{len(documents)}) or 'back' to cancel[/bold cyan]"
            ).strip()

            if choice.lower() in ['back', 'cancel', 'exit']:
                return

            doc_index = int(choice) - 1
            if 0 <= doc_index < len(documents):
                selected_doc = documents[doc_index]

                if Confirm.ask(f"[bold red]Are you sure you want to delete '{selected_doc.filename}'?[/bold red]"):
                    if self.assistant.file_manager.delete_document(selected_doc.id, self.user_id):
                        self.console.print(
                            f"[bold green]✅ Document '{selected_doc.filename}' deleted successfully![/bold green]")
                    else:
                        self.console.print(
                            f"[bold red]❌ Failed to delete document '{selected_doc.filename}'[/bold red]")
            else:
                self.console.print("[red]Invalid selection[/red]")

        except ValueError:
            self.console.print("[red]Invalid input. Please enter a number.[/red]")
        except Exception as e:
            self.console.print(f"[red]Error: {str(e)}[/red]")

    def clear_conversations_menu(self):
        """Interactive menu for clearing conversations"""
        options_text = """
        [bold magenta]🗑️ Clear Conversations Options:[/bold magenta]

        • [bold green]1[/bold green] - Clear ALL conversations (complete reset)
        • [bold yellow]2[/bold yellow] - Clear old conversations (keep recent 20)
        • [bold blue]3[/bold blue] - Clear old conversations (keep recent 10)
        • [bold red]4[/bold red] - Clear old conversations (keep recent 5)
        • [bold dim]back[/bold dim] - Return to main menu

        [bold yellow]⚠️ Warning:[/bold yellow] This action cannot be undone!
        """

        self.console.print(Panel(options_text, border_style="red", padding=(1, 2)))

        choice = Prompt.ask("[bold cyan]Select option[/bold cyan]").strip().lower()

        if choice in ['back', 'cancel', 'exit']:
            return

        # Show current conversation count
        conversations = self.assistant.file_manager.get_recent_conversations(self.user_id, 1000)
        current_count = len(conversations)

        self.console.print(f"[bold]Current conversations: {current_count}[/bold]")

        if choice == '1':
            # Clear all conversations
            if Confirm.ask(
                    f"[bold red]Are you sure you want to delete ALL {current_count} conversations?[/bold red]"):
                progress = self.create_progress()
                task = progress.add_task("[bold red]🗑️ Clearing all conversations...[/bold red]", total=100)

                with Live(progress, auto_refresh=True, console=self.console, refresh_per_second=20):
                    for i in range(100):
                        progress.update(task, completed=i + 1)
                        time.sleep(0.02)

                    success = self.assistant.file_manager.clear_all_conversations(self.user_id)

                if success:
                    self.console.print("[bold green]✅ All conversations cleared successfully![/bold green]")
                    # Clear cached context
                    self.assistant.clear_cached_context(self.user_id)
                else:
                    self.console.print("[bold red]❌ Failed to clear conversations[/bold red]")

        elif choice in ['2', '3', '4']:
            # Clear old conversations with different keep amounts
            keep_amounts = {'2': 20, '3': 10, '4': 5}
            keep_recent = keep_amounts[choice]

            if current_count <= keep_recent:
                self.console.print(
                    f"[yellow]You only have {current_count} conversations. No need to clear old ones.[/yellow]")
                return

            to_delete = current_count - keep_recent
            if Confirm.ask(
                    f"[bold yellow]Delete {to_delete} old conversations and keep the {keep_recent} most recent?[/bold yellow]"):
                progress = self.create_progress()
                task = progress.add_task(f"[bold yellow]🗑️ Clearing {to_delete} old conversations...[/bold yellow]",
                                         total=100)

                with Live(progress, auto_refresh=True, console=self.console, refresh_per_second=20):
                    for i in range(100):
                        progress.update(task, completed=i + 1)
                        time.sleep(0.02)

                    deleted_count = self.assistant.file_manager.clear_old_conversations(self.user_id, keep_recent)

                self.console.print(
                    f"[bold green]✅ Cleared {deleted_count} old conversations, kept {keep_recent} recent ones![/bold green]")
                # Clear cached context to refresh
                self.assistant.clear_cached_context(self.user_id)
        else:
            self.console.print("[red]Invalid option selected[/red]")

    def manage_documents_menu(self):
        """Enhanced document management with capacity info and delete options"""
        while True:
            # Get capacity info
            capacity_info = self.assistant.file_manager.get_documents_capacity_info(self.user_id)

            # Create capacity display
            usage_bar = "█" * int(capacity_info['usage_percentage'] / 5) + "░" * (
                    20 - int(capacity_info['usage_percentage'] / 5))
            capacity_color = "red" if capacity_info['usage_percentage'] > 80 else "yellow" if capacity_info[
                                                                                                  'usage_percentage'] > 60 else "green"

            menu_text = f"""
            [bold magenta]📄 Document Management:[/bold magenta]

            [bold]Storage Usage:[/bold]
            [{capacity_color}]{usage_bar}[/{capacity_color}] {capacity_info['usage_percentage']}%
            • Documents: {capacity_info['total_documents']}/{capacity_info['max_documents']} ({capacity_info['documents_remaining']} remaining)
            • Storage: {capacity_info['total_size_mb']}MB/{capacity_info['max_size_mb']}MB ({capacity_info['size_remaining_mb']}MB remaining)

            [bold]Options:[/bold]
            • [bold green]1[/bold green] - View all documents
            • [bold blue]2[/bold blue] - Upload new document
            • [bold yellow]3[/bold yellow] - Delete specific document
            • [bold red]4[/bold red] - Clear all documents
            • [bold dim]back[/bold dim] - Return to main menu
            """

            # Add warnings if at limits
            if capacity_info['at_document_limit']:
                menu_text += "\n[bold red]⚠️ Document limit reached! Delete some documents to upload new ones.[/bold red]"
            if capacity_info['at_size_limit']:
                menu_text += "\n[bold red]⚠️ Storage limit reached! Delete some documents to free space.[/bold red]"

            self.console.print(Panel(menu_text, border_style="cyan", padding=(1, 2)))

            choice = Prompt.ask("[bold cyan]Select option[/bold cyan]").strip().lower()

            if choice in ['back', 'cancel', 'exit']:
                break
            elif choice == '1':
                self.display_documents()
                Prompt.ask("[bold]Press Enter to continue...[/bold]")
            elif choice == '2':
                if capacity_info['at_document_limit']:
                    self.console.print("[red]❌ Document limit reached! Delete some documents first.[/red]")
                elif capacity_info['at_size_limit']:
                    self.console.print("[red]❌ Storage limit reached! Delete some documents first.[/red]")
                else:
                    self.upload_document()
                Prompt.ask("[bold]Press Enter to continue...[/bold]")
            elif choice == '3':
                self.delete_specific_document()
                Prompt.ask("[bold]Press Enter to continue...[/bold]")
            elif choice == '4':
                if Confirm.ask(
                        f"[bold red]Are you sure you want to delete ALL {capacity_info['total_documents']} documents?[/bold red]"):
                    progress = self.create_progress()
                    task = progress.add_task("[bold red]🗑️ Clearing all documents...[/bold red]", total=100)

                    with Live(progress, auto_refresh=True, console=self.console, refresh_per_second=20):
                        for i in range(100):
                            progress.update(task, completed=i + 1)
                            time.sleep(0.02)

                        deleted_count = self.assistant.file_manager.clear_all_documents(self.user_id)

                    self.console.print(
                        f"[bold green]✅ Cleared {deleted_count} documents successfully![/bold green]")
                Prompt.ask("[bold]Press Enter to continue...[/bold]")
            else:
                self.console.print("[red]Invalid option selected[/red]")

    def display_maintenance_menu(self):
        """Display maintenance and cleanup options"""
        maintenance_text = """
        [bold magenta]🔧 Maintenance & Cleanup:[/bold magenta]

        • [bold green]1[/bold green] - Clear conversations (with options)
        • [bold blue]2[/bold blue] - Manage documents (view/delete/capacity)
        • [bold yellow]3[/bold yellow] - View storage statistics
        • [bold red]4[/bold red] - Full system cleanup (conversations + documents)
        • [bold dim]back[/bold dim] - Return to main menu

        [bold yellow]⚠️ Warning:[/bold yellow] Cleanup operations cannot be undone!
        """

        self.console.print(Panel(maintenance_text, border_style="yellow", padding=(1, 2)))

        choice = Prompt.ask("[bold cyan]Select maintenance option[/bold cyan]").strip().lower()

        if choice in ['back', 'cancel', 'exit']:
            return
        elif choice == '1':
            self.clear_conversations_menu()
        elif choice == '2':
            self.manage_documents_menu()
        elif choice == '3':
            self.display_stats()
        elif choice == '4':
            # Full cleanup
            conversations = self.assistant.file_manager.get_recent_conversations(self.user_id, 1000)
            documents = self.assistant.file_manager.get_user_documents(self.user_id, 1000)

            warning_text = f"""
            [bold red]⚠️ FULL SYSTEM CLEANUP WARNING ⚠️[/bold red]

            This will delete:
            • All {len(conversations)} conversations
            • All {len(documents)} documents
            • All cached context data

            This action is [bold red]IRREVERSIBLE[/bold red]!
            """

            self.console.print(Panel(warning_text, border_style="red", padding=(1, 2)))

            if Confirm.ask("[bold red]Are you absolutely sure you want to perform a full cleanup?[/bold red]"):
                if Confirm.ask("[bold red]This is your final warning. Continue with full cleanup?[/bold red]"):
                    progress = self.create_progress()
                    task = progress.add_task("[bold red]🗑️ Performing full system cleanup...[/bold red]", total=100)

                    with Live(progress, auto_refresh=True, console=self.console, refresh_per_second=20):
                        # Clear conversations
                        progress.update(task, completed=25,
                                        description="[bold red]🗑️ Clearing conversations...[/bold red]")
                        conv_success = self.assistant.file_manager.clear_all_conversations(self.user_id)

                        # Clear documents
                        progress.update(task, completed=50,
                                        description="[bold red]🗑️ Clearing documents...[/bold red]")
                        doc_count = self.assistant.file_manager.clear_all_documents(self.user_id)

                        # Clear cache
                        progress.update(task, completed=75, description="[bold red]🗑️ Clearing cache...[/bold red]")
                        self.assistant.clear_cached_context(self.user_id)

                        # Complete
                        progress.update(task, completed=100,
                                        description="[bold green]✅ Full cleanup complete![/bold green]")
                        time.sleep(1)

                    self.console.print(f"[bold green]✅ Full cleanup completed successfully![/bold green]")
                    self.console.print(f"• Conversations cleared: {len(conversations)}")
                    self.console.print(f"• Documents cleared: {doc_count}")
                    self.console.print("• Cache cleared")
        else:
            self.console.print("[red]Invalid option selected[/red]")

    def download_chat_menu(self):
        """Main download menu with all options"""
        menu_text = """
        [bold magenta]💾 Download Chat History:[/bold magenta]

        • [bold green]1[/bold green] - Save current chat only
        • [bold blue]2[/bold blue] - Select specific chats to download
        • [bold yellow]3[/bold yellow] - Download all chats (bulk export)
        • [bold cyan]4[/bold cyan] - Download last N chats (10/20/50)
        • [bold dim]back[/bold dim] - Cancel
        """

        self.console.print(Panel(menu_text, border_style="cyan", padding=(1, 2)))

        choice = Prompt.ask("[bold cyan]Select option (1-4)[/bold cyan]").strip()

        if choice.lower() in ['back', 'cancel', 'exit']:
            return

        if choice == '1':
            self.download_current_chat()
        elif choice == '2':
            self.download_selected_chats()
        elif choice == '3':
            # Existing bulk download logic
            self._bulk_download_all()
        elif choice == '4':
            # Existing limited download logic
            self._bulk_download_limited()
        else:
            self.console.print("[red]Invalid option[/red]")

    def _bulk_download_all(self):
        """Download all conversations"""
        format_choice = Prompt.ask("[bold cyan]Format? (1=txt, 2=md, 3=json)[/bold cyan]").strip()
        format_map = {'1': 'txt', '2': 'md', '3': 'json'}
        export_format = format_map.get(format_choice, 'txt')

        progress = self.create_progress()
        task = progress.add_task("[bold cyan]💾 Exporting all chats...[/bold cyan]", total=100)

        with Live(progress, auto_refresh=True, console=self.console, refresh_per_second=20):
            for i in range(0, 100, 5):
                progress.update(task, completed=i)
                time.sleep(0.03)

            filepath = self.assistant.file_manager.export_conversations_to_file(
                self.user_id,
                format=export_format,
                limit=None  # All conversations
            )

            progress.update(task, completed=100, description="[bold green]✅ Complete![/bold green]")

        if filepath:
            file_size = Path(filepath).stat().st_size
            self.console.print(Panel(
                f"[bold green]✅ Exported all chats![/bold green]\n"
                f"📁 {Path(filepath).name}\n"
                f"📏 {decimal(file_size)}",
                border_style="green"
            ))

    def _bulk_download_limited(self):
        """Download limited number of conversations"""
        scope_choice = Prompt.ask("[bold cyan]How many? (a=10, b=20, c=50)[/bold cyan]").strip().lower()
        scope_map = {'a': 10, 'b': 20, 'c': 50}
        limit = scope_map.get(scope_choice, 20)

        format_choice = Prompt.ask("[bold cyan]Format? (1=txt, 2=md, 3=json)[/bold cyan]").strip()
        format_map = {'1': 'txt', '2': 'md', '3': 'json'}
        export_format = format_map.get(format_choice, 'txt')

        progress = self.create_progress()
        task = progress.add_task(f"[bold cyan]💾 Exporting last {limit} chats...[/bold cyan]", total=100)

        with Live(progress, auto_refresh=True, console=self.console, refresh_per_second=20):
            for i in range(0, 100, 5):
                progress.update(task, completed=i)
                time.sleep(0.03)

            filepath = self.assistant.file_manager.export_conversations_to_file(
                self.user_id,
                format=export_format,
                limit=limit
            )

            progress.update(task, completed=100, description="[bold green]✅ Complete![/bold green]")

        if filepath:
            file_size = Path(filepath).stat().st_size
            self.console.print(Panel(
                f"[bold green]✅ Exported last {limit} chats![/bold green]\n"
                f"📁 {Path(filepath).name}\n"
                f"📏 {decimal(file_size)}",
                border_style="green"
            ))

    def download_current_chat(self):
        """Download only the current/latest conversation"""
        menu_text = """
        [bold magenta]💾 Download Current Chat:[/bold magenta]

        Choose format:
        • [bold green]1[/bold green] - Plain Text (.txt)
        • [bold blue]2[/bold blue] - Markdown (.md)
        • [bold yellow]3[/bold yellow] - JSON (.json)
        • [bold dim]back[/bold dim] - Cancel
        """

        self.console.print(Panel(menu_text, border_style="cyan", padding=(1, 2)))

        format_choice = Prompt.ask("[bold cyan]Select format (1-3)[/bold cyan]").strip()

        if format_choice.lower() in ['back', 'cancel', 'exit']:
            return

        format_map = {'1': 'txt', '2': 'md', '3': 'json'}
        export_format = format_map.get(format_choice)

        if not export_format:
            self.console.print("[red]Invalid format selection[/red]")
            return

        # Export only the most recent conversation
        progress = self.create_progress()
        task = progress.add_task("[bold cyan]💾 Saving current chat...[/bold cyan]", total=100)

        with Live(progress, auto_refresh=True, console=self.console, refresh_per_second=20):
            for i in range(0, 100, 5):
                progress.update(task, completed=i)
                time.sleep(0.02)

            filepath = self.assistant.file_manager.export_conversations_to_file(
                self.user_id,
                format=export_format,
                limit=1  # Only export last conversation
            )

            progress.update(task, completed=100, description="[bold green]✅ Current chat saved![/bold green]")

        if filepath:
            file_size = Path(filepath).stat().st_size

            success_text = f"""
            [bold green]✅ Current chat saved successfully![/bold green]

            📁 File: [bold cyan]{Path(filepath).name}[/bold cyan]
            📏 Size: [bold]{decimal(file_size)}[/bold]
            📍 Location: [dim]{filepath}[/dim]
            """

            self.console.print(Panel(success_text, border_style="green", padding=(1, 2)))
        else:
            self.console.print("[red]❌ No conversation to save[/red]")

    def download_selected_chats(self):
        """Interactively select specific chats to download"""

        # Load conversations
        progress = self.create_progress()
        task = progress.add_task("[bold cyan]📋 Loading conversations...[/bold cyan]", total=100)

        with Live(progress, auto_refresh=True, console=self.console, refresh_per_second=20):
            for i in range(0, 100, 10):
                progress.update(task, completed=i)
                time.sleep(0.02)

            conversations = self.assistant.file_manager.get_recent_conversations(self.user_id, 50)
            progress.update(task, completed=100, description="[bold green]✅ Loaded![/bold green]")

        if not conversations:
            self.console.print("[yellow]No conversations found[/yellow]")
            return

        # Display conversations for selection
        convs_table = Table(
            show_header=True,
            header_style="bold cyan",
            box=ROUNDED,
            title=f"Select Conversations to Download (Total: {len(conversations)})"
        )
        convs_table.add_column("#", style="bold magenta", width=4)
        convs_table.add_column("Date", style="dim", width=16)
        convs_table.add_column("Your Message", style="cyan", width=40)
        convs_table.add_column("Quality", style="green", width=8)

        for i, conv in enumerate(conversations, 1):
            user_msg = conv.user_message[:37] + "..." if len(conv.user_message) > 40 else conv.user_message
            quality = f"{conv.multi_agent_session.quality_score:.2f}" if conv.multi_agent_session else "N/A"

            convs_table.add_row(
                str(i),
                conv.timestamp.strftime("%Y-%m-%d %H:%M"),
                user_msg,
                quality
            )

        self.console.print(Panel(convs_table, border_style="cyan", padding=(1, 1)))

        # Get user selection
        selection_text = """
        [bold cyan]Enter your selection:[/bold cyan]
        • Single: [bold]5[/bold] (conversation #5)
        • Range: [bold]1-10[/bold] (conversations 1 to 10)
        • Multiple: [bold]1,5,8[/bold] (conversations 1, 5, and 8)
        • All: [bold]all[/bold]
        • Cancel: [bold]back[/bold]
        """
        self.console.print(selection_text)

        selection = Prompt.ask("[bold cyan]Your selection[/bold cyan]").strip().lower()

        if selection in ['back', 'cancel', 'exit']:
            return

        # Parse selection
        selected_indices = set()

        try:
            if selection == 'all':
                selected_indices = set(range(len(conversations)))
            elif '-' in selection:
                # Range selection (e.g., "1-10")
                start, end = map(int, selection.split('-'))
                selected_indices = set(range(start - 1, min(end, len(conversations))))
            elif ',' in selection:
                # Multiple selection (e.g., "1,3,5")
                selected_indices = {int(x.strip()) - 1 for x in selection.split(',')}
            else:
                # Single selection (e.g., "5")
                selected_indices = {int(selection) - 1}

            # Validate indices
            selected_indices = {i for i in selected_indices if 0 <= i < len(conversations)}

            if not selected_indices:
                self.console.print("[red]No valid conversations selected[/red]")
                return

        except ValueError:
            self.console.print("[red]Invalid selection format[/red]")
            return

        # Show selection summary
        self.console.print(f"\n[bold green]Selected {len(selected_indices)} conversation(s)[/bold green]")

        # Choose format
        format_choice = Prompt.ask("[bold cyan]Format? (1=txt, 2=md, 3=json)[/bold cyan]").strip()
        format_map = {'1': 'txt', '2': 'md', '3': 'json'}
        export_format = format_map.get(format_choice, 'txt')

        # Export selected conversations
        selected_convs = [conversations[i] for i in sorted(selected_indices)]

        # Temporarily save selected conversations
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"selected_chats_{timestamp}.{export_format}"
        filepath = self.assistant.file_manager.data_dir / filename

        progress = self.create_progress()
        task = progress.add_task("[bold cyan]💾 Exporting selected chats...[/bold cyan]", total=100)

        with Live(progress, auto_refresh=True, console=self.console, refresh_per_second=20):
            for i in range(0, 60, 5):
                progress.update(task, completed=i)
                time.sleep(0.02)

            # Format content based on export format
            if export_format == 'txt':
                content = self.assistant.file_manager._format_txt_export(selected_convs)
            elif export_format == 'md':
                content = self.assistant.file_manager._format_markdown_export(selected_convs)
            else:
                content = self.assistant.file_manager._format_json_export(selected_convs)

            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)

            for i in range(60, 100, 5):
                progress.update(task, completed=i)
                time.sleep(0.02)

            progress.update(task, completed=100, description="[bold green]✅ Export complete![/bold green]")

        file_size = filepath.stat().st_size

        success_text = f"""
        [bold green]✅ Selected chats exported successfully![/bold green]

        📊 Exported: [bold]{len(selected_convs)} conversation(s)[/bold]
        📁 File: [bold cyan]{filename}[/bold cyan]
        📏 Size: [bold]{decimal(file_size)}[/bold]
        📍 Location: [dim]{filepath}[/dim]
        """

        self.console.print(Panel(success_text, border_style="green", padding=(1, 2)))

    async def handle_chat_mode(self):
        """Handle chat mode with persistent context and exit option"""
        self.console.print(Panel(
            "[bold cyan]Entering chat mode. Type [bold red]back[/bold red] or [bold red]exit[/bold red] to return to main menu.[/bold cyan]\n"
            "[bold yellow]Note:[/bold yellow] The system uses your conversation history and uploaded documents "
            "to provide more accurate and context-aware responses.",
            border_style="cyan",
            padding=(1, 2)
        ))

        while True:
            message = Prompt.ask("[bold cyan]Your Query[/bold cyan]")
            if message.strip().lower() in ['back', 'exit']:
                break

            if message.strip():
                # Debug: Check what context is being used
                context = self.assistant.file_manager.get_conversation_context(self.user_id, 70)
                self.console.print(
                    f"[dim]DEBUG: Context loaded - {len(context.split())} words, {len(context.splitlines())} lines[/dim]")

                # Show animated agent processing
                self.show_agent_progress(message)
                await self.process_message_with_progress(message)

    async def run_interactive(self):
        """Main interactive loop with command navigation"""
        self.display_banner()
        self.console.print("\n")
        with Status("[bold cyan]🔄 Starting multi-agent assistant...[/bold cyan]", console=self.console, spinner="dots"):
            self.assistant.start()
            await asyncio.sleep(1)
        self.console.print("[bold green]✓ LAZYCOOK is now running! [/bold green]\n")
        self.display_help()

        try:
            while True:
                self.console.print()
                command = Prompt.ask("[bold blue]Enter command[/bold blue] (or 'quit' to exit)").strip().lower()
                command = self.command_aliases.get(command, command)

                if command in ["quit", "exit"]:
                    if Confirm.ask("[bold red]Are you sure you want to quit?[/bold red]"):
                        break

                elif command == "chat":
                    await self.handle_chat_mode()

                elif command == "insights":
                    self.display_insights()
                    Prompt.ask("[bold]Press Enter to continue...[/bold]")

                elif command == "context":
                    self.display_context()
                    Prompt.ask("[bold]Press Enter to continue...[/bold]")

                elif command == "docs":
                    self.display_documents()
                    Prompt.ask("[bold]Press Enter to continue...[/bold]")

                elif command == "clear":
                    self.console.clear()
                    self.display_banner()
                elif command == "agents":
                    self.display_agents()
                    Prompt.ask("[bold]Press Enter to continue...[/bold]")

                elif command == "tasks":
                    self.display_tasks()
                    Prompt.ask("[bold]Press Enter to continue...[/bold]")

                elif command == "download":
                    self.download_chat_menu()
                    Prompt.ask("[bold]Press Enter to continue...[/bold]")

                elif command == "quality":
                    self.display_quality()
                    Prompt.ask("[bold]Press Enter to continue...[/bold]")

                elif command == "files":
                    self.display_files()
                    Prompt.ask("[bold]Press Enter to continue...[/bold]")

                elif command == "stats":
                    self.display_stats()
                    Prompt.ask("[bold]Press Enter to continue...[/bold]")

                elif command == "routing":
                    self.display_routing_stats()
                    Prompt.ask("[bold]Press Enter to continue...[/bold]")


                elif command == "help":
                    self.display_help()

                elif command == "maintenance":
                    self.display_maintenance_menu()
                    Prompt.ask("[bold]Press Enter to continue...[/bold]")

                elif command == "cls":
                    self.console.clear()
                    self.display_banner()

                else:
                    self.console.print(f"[red]Unknown command: {command}[/red]")
                    self.console.print("[dim]Type 'help' for available commands.[/dim]")

        finally:
            with Status("[bold cyan]🔄 Shutting down multi-agent assistant...[/bold cyan]", console=self.console,
                        spinner="dots"):
                self.assistant.stop()
                await asyncio.sleep(1)
            self.console.print("\n[bold green]✓ Multi-Agent Assistant stopped successfully![/bold green]")
            self.console.print("[dim]Thank you for using the Enhanced Multi-Agent Assistant![/dim]")


class MultiAgentAssistantConfig:
    """Configuration class for external usage of the Multi-Agent Assistant"""

    def __init__(self, api_key: str, conversation_limit: int = 70, document_limit: int = 50):
        self.api_key = api_key
        self.conversation_limit = conversation_limit
        self.document_limit = document_limit

    def create_assistant(self):
        """Create assistant with properly configured file manager"""
        assistant = AutonomousMultiAgentAssistant(self.api_key)
        # Replace the file manager with configured one
        assistant.file_manager = TextFileManager(
            conversation_limit=self.conversation_limit,
            document_limit=self.document_limit,
            api_key=self.api_key
        )
        return assistant

    def create_cli(self):
        """Create CLI with configured file manager"""
        cli = RichMultiAgentCLI(self.api_key)
        # Replace the file manager in the assistant
        cli.assistant.file_manager = TextFileManager(
            conversation_limit=self.conversation_limit,
            document_limit=self.document_limit,
            api_key=self.api_key
        )
        return cli

    async def run_cli(self):
        """Run the CLI interface with custom configuration"""
        cli = self.create_cli()
        try:
            await cli.run_interactive()
        except KeyboardInterrupt:
            Console().print("\n[yellow]⚠️ Interrupted by user[/yellow]")
        except Exception as e:
            import traceback
            console = Console()
            console.print(f"[bold red]❌ Error: {str(e)}[/bold red]")
            console.print(f"[red]Full traceback:[/red]")
            console.print(traceback.format_exc())
# --- Main ---
def main():
    console = Console()
    api_key = os.getenv("GEMINI_API_KEY")

    if not api_key:
        console.print("[bold red]❌ API key is required to run the assistant.[/bold red]")
        return

    with console.status("[bold cyan]🔄 Testing Gemini API connection...[/bold cyan]"):
        try:
            genai.configure(api_key=api_key)
            test_model = genai.GenerativeModel('gemini-2.5-flash')
            test_response = test_model.generate_content("Hello")
        except Exception as e:
            console.print(f"[bold red]❌ API connection failed: {e}[/bold red]")
            return

    console.print("[bold green]✓ API connection successful![/bold green]")
    cli = RichMultiAgentCLI(api_key)
    try:
        asyncio.run(cli.run_interactive())
    except KeyboardInterrupt:
        console.print("\n[yellow]⚠️ Interrupted by user[/yellow]")
    except Exception as e:
        import traceback
        console.print(f"[bold red]❌ Error: {str(e)}[/bold red]")
        console.print(f"[red]Full traceback:[/red]")
        console.print(traceback.format_exc())


def create_assistant(api_key: str, conversation_limit: int = 70, document_limit: int = 50):
    """Factory function to create a configured assistant for external use"""
    return MultiAgentAssistantConfig(api_key, conversation_limit, document_limit)

if __name__ == "__main__":
    main()