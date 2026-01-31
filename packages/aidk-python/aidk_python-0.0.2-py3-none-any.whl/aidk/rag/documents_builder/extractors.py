import os
import importlib.util
from typing import List, Dict, Tuple, Optional

from .base import BaseExtractor
from .splitters import Chunker
from aidk.tools.webscraping import scrape_web


class FileExtractor(BaseExtractor):
    """Extract text from plain text files."""
    
    def __init__(self, file_path: str, chunker: Chunker):
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        self.file_path = file_path
        self.chunker = chunker
    
    def extract(self) -> Tuple[List[str], List[Dict], List[str]]:
        """Extract and split text file."""
        with open(self.file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        chunks = self.chunker.split(text)
        
        base_metadata = {
            'file_path': self.file_path,
            'file_name': os.path.basename(self.file_path),
            'chunk_strategy': self.chunker.__class__.__name__
        }
        
        return self._generate_chunks_with_metadata(text, chunks, base_metadata)


class StringExtractor(BaseExtractor):
    """Extract text from string."""
    
    def __init__(self, text: str, source_name: str = "text_string", chunker: Optional[Chunker] = None):
        if not text or not text.strip():
            raise ValueError("Text cannot be empty")
        
        if chunker is None:
            raise ValueError("chunker parameter is required")
        
        self.text = text
        self.source_name = source_name
        self.chunker = chunker
    
    def extract(self) -> Tuple[List[str], List[Dict], List[str]]:
        """Extract and split string."""
        chunks = self.chunker.split(self.text)
        
        base_metadata = {
            'source_type': 'text_string',
            'source_name': self.source_name,
            'chunk_strategy': self.chunker.__class__.__name__
        }
        
        return self._generate_chunks_with_metadata(self.text, chunks, base_metadata)


class DocExtractor(BaseExtractor):
    """Extract text from Word documents (.doc and .docx)."""
    
    def __init__(self, file_path: str, extraction_method: str = "auto", chunker: Optional[Chunker] = None):
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if chunker is None:
            raise ValueError("chunker parameter is required")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext not in ['.doc', '.docx']:
            raise ValueError(f"Unsupported format: {file_ext}")
        
        self.file_path = file_path
        self.file_ext = file_ext
        self.extraction_method = extraction_method
        self.chunker = chunker
    
    def extract(self) -> Tuple[List[str], List[Dict], List[str]]:
        """Extract text from Word document."""
        text, doc_props = self._extract_text()
        chunks = self.chunker.split(text)
        
        base_metadata = {
            'file_path': self.file_path,
            'file_name': os.path.basename(self.file_path),
            'document_format': self.file_ext[1:],
            'extraction_method': self.extraction_method,
            'chunk_strategy': self.chunker.__class__.__name__,
            **doc_props
        }
        
        return self._generate_chunks_with_metadata(text, chunks, base_metadata)
    
    def _extract_text(self) -> Tuple[str, Dict]:
        """Extract text and metadata from Word document."""
        method = self._select_extraction_method()
        
        if method == "docx":
            return self._extract_with_docx()
        elif method == "docx2txt":
            return self._extract_with_docx2txt()
        else:
            raise ValueError(f"Unsupported extraction method: {method}")
    
    def _select_extraction_method(self) -> str:
        """Select extraction method based on availability."""
        if self.extraction_method != "auto":
            return self.extraction_method
        
        if self.file_ext == '.docx' and self._has_module("docx"):
            return "docx"
        elif self._has_module("docx2txt"):
            return "docx2txt"
        else:
            raise ImportError("docx2txt required. Install with: pip install docx2txt")
    
    @staticmethod
    def _has_module(name: str) -> bool:
        """Check if module is installed."""
        return importlib.util.find_spec(name) is not None
    
    def _extract_with_docx(self) -> Tuple[str, Dict]:
        """Extract using python-docx."""
        import docx
        doc = docx.Document(self.file_path)
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        text = "\n\n".join(text_parts)
        
        props = {}
        core = doc.core_properties
        if core.title:
            props['document_title'] = core.title
        if core.author:
            props['document_author'] = core.author
        if core.subject:
            props['document_subject'] = core.subject
        if core.created:
            props['document_created'] = str(core.created)
        if core.modified:
            props['document_modified'] = str(core.modified)
        
        return text, props
    
    def _extract_with_docx2txt(self) -> Tuple[str, Dict]:
        """Extract using docx2txt."""
        import docx2txt
        text = docx2txt.process(self.file_path)
        return text, {}


class PDFExtractor(BaseExtractor):
    """Extract text from PDF documents."""
    
    def __init__(self, file_path: str, page_range: Optional[Tuple[int, int]] = None, chunker: Optional[Chunker] = None):
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if chunker is None:
            raise ValueError("chunker parameter is required")
        
        if not file_path.lower().endswith('.pdf'):
            raise ValueError("Only PDF files are supported")
        
        if not self._has_module("PyPDF2"):
            raise ImportError("PyPDF2 required. Install with: pip install PyPDF2")
        
        self.file_path = file_path
        self.page_range = page_range
        self.chunker = chunker
    
    def extract(self) -> Tuple[List[str], List[Dict], List[str]]:
        """Extract text from PDF."""
        text, pdf_props, page_info = self._extract_text()
        chunks = self.chunker.split(text)
        
        base_metadata = {
            'file_path': self.file_path,
            'file_name': os.path.basename(self.file_path),
            'document_format': 'pdf',
            'chunk_strategy': self.chunker.__class__.__name__,
            **pdf_props,
            **page_info
        }
        
        return self._generate_chunks_with_metadata(text, chunks, base_metadata)
    
    def _extract_text(self) -> Tuple[str, Dict, Dict]:
        """Extract text and metadata from PDF."""
        import PyPDF2
        
        with open(self.file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            total = len(reader.pages)
            
            start, end = self._validate_page_range(total)
            
            text_parts = []
            for i in range(start - 1, end):
                page_text = reader.pages[i].extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text)
            
            text = "\n\n".join(text_parts)
            
            pdf_props = self._extract_pdf_metadata(reader)
            page_info = {
                'total_pages': total,
                'extracted_pages_start': start,
                'extracted_pages_end': end,
                'extracted_pages_count': end - start + 1
            }
        
        return text, pdf_props, page_info
    
    def _validate_page_range(self, total: int) -> Tuple[int, int]:
        """Validate and return page range."""
        if self.page_range is None:
            return 1, total
        
        start, end = self.page_range
        if start < 1 or end > total or start > end:
            raise ValueError(f"Invalid range: {self.page_range}. Must be 1-{total}")
        
        return start, end
    
    def _extract_pdf_metadata(self, reader) -> Dict:
        """Extract PDF metadata."""
        props = {}
        if not hasattr(reader, 'metadata') or reader.metadata is None:
            return props
        
        metadata = reader.metadata
        mapping = {
            '/Title': 'pdf_title',
            '/Author': 'pdf_author',
            '/Subject': 'pdf_subject',
            '/Creator': 'pdf_creator',
            '/Producer': 'pdf_producer',
            '/CreationDate': 'pdf_creation_date',
            '/ModDate': 'pdf_modification_date',
        }
        
        for key, dest_key in mapping.items():
            try:
                val = metadata.get(key) if hasattr(metadata, 'get') else getattr(metadata, key, None)
                if val:
                    props[dest_key] = str(val)
            except Exception:
                pass
        
        return props
    
    @staticmethod
    def _has_module(name: str) -> bool:
        """Check if module is installed."""
        return importlib.util.find_spec(name) is not None


class URLExtractor(BaseExtractor):
    """Extract text from web pages."""
    
    def __init__(self, url: str, engine: str = "requests", deep: bool = False, chunker: Optional[Chunker] = None):
        if chunker is None:
            raise ValueError("chunker parameter is required")
        
        self.url = url
        self.engine = engine
        self.deep = deep
        self.chunker = chunker
    
    def extract(self) -> Tuple[List[str], List[Dict], List[str]]:
        """Extract text from URL."""
        result = scrape_web(self.url, engine=self.engine, deep=self.deep)
        
        if not result or not result.get("text"):
            raise ValueError(f"Failed to extract from URL: {self.url}")
        
        text = result["text"]
        chunks = self.chunker.split(text)
        
        base_metadata = {
            'url': self.url,
            'source_type': 'web_page',
            'scraping_engine': self.engine,
            'deep_extraction': self.deep,
            'chunk_strategy': self.chunker.__class__.__name__
        }
        
        return self._generate_chunks_with_metadata(text, chunks, base_metadata)
