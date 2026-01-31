"""
Export research sessions to various formats.

Supports:
- Markdown (.md): Clean, readable format with full citations
- JSON (.json): Machine-readable with all metadata
- DOCX (.docx): Professional Word document via Marko + python-docx

DOCX generation uses:
- Marko: CommonMark-compliant Markdown parser with object-based AST
- python-docx: Native Python DOCX generation with full styling control

This approach provides:
- Zero external runtime dependencies (no Pandoc binary needed)
- Native page break control via paragraph_format.page_break_before
- Proper heading hierarchy and styling
- Full Markdown feature support
- Professional typography
"""

from __future__ import annotations

import json
import logging
import re
from dataclasses import dataclass
from datetime import UTC, datetime
from enum import Enum
from io import BytesIO
from pathlib import Path
from typing import TYPE_CHECKING, Any

from gemini_research_mcp.config import LOGGER_NAME

if TYPE_CHECKING:
    from gemini_research_mcp.storage import ResearchSession

logger = logging.getLogger(LOGGER_NAME)


class ExportFormat(str, Enum):
    """Supported export formats."""

    MARKDOWN = "markdown"
    JSON = "json"
    DOCX = "docx"


@dataclass
class ExportResult:
    """Result of an export operation."""

    format: ExportFormat
    filename: str
    content: bytes
    mime_type: str

    @property
    def size_human(self) -> str:
        """Human-readable file size."""
        size: float = len(self.content)
        for unit in ["B", "KB", "MB"]:
            if size < 1024:
                return f"{size:.1f} {unit}"
            size /= 1024
        return f"{size:.1f} GB"


# =============================================================================
# Markdown Export
# =============================================================================


def _format_markdown_export(session: ResearchSession) -> str:
    """Format a research session as a Markdown document."""
    lines: list[str] = []

    # Title
    title = session.title or session.query[:60]
    lines.append(f"# {title}")
    lines.append("")

    # Metadata block
    lines.append("## Metadata")
    lines.append("")
    lines.append(f"- **Query:** {session.query}")
    lines.append(f"- **Created:** {session.created_at_iso}")
    if session.duration_seconds:
        mins = int(session.duration_seconds // 60)
        secs = int(session.duration_seconds % 60)
        lines.append(f"- **Duration:** {mins}m {secs}s")
    if session.total_tokens:
        lines.append(f"- **Tokens:** {session.total_tokens:,}")
    if session.agent_name:
        lines.append(f"- **Agent:** {session.agent_name}")
    if session.tags:
        lines.append(f"- **Tags:** {', '.join(session.tags)}")
    if session.notes:
        lines.append(f"- **Notes:** {session.notes}")
    lines.append(f"- **Interaction ID:** `{session.interaction_id}`")
    if session.expires_at_iso:
        lines.append(f"- **Expires:** {session.expires_at_iso}")
    lines.append("")

    # Summary
    if session.summary:
        lines.append("## Summary")
        lines.append("")
        lines.append(session.summary)
        lines.append("")

    # Full report
    if session.report_text:
        lines.append("## Research Report")
        lines.append("")
        lines.append(session.report_text)
        lines.append("")

    # Footer
    lines.append("---")
    lines.append(
        f"*Exported from Gemini Research MCP on "
        f"{datetime.now(tz=UTC).strftime('%Y-%m-%d %H:%M:%S UTC')}*"
    )

    return "\n".join(lines)


def export_to_markdown(session: ResearchSession) -> ExportResult:
    """Export a research session to Markdown format."""
    content = _format_markdown_export(session)
    filename = _generate_filename(session, "md")

    return ExportResult(
        format=ExportFormat.MARKDOWN,
        filename=filename,
        content=content.encode("utf-8"),
        mime_type="text/markdown",
    )


# =============================================================================
# JSON Export
# =============================================================================


def _session_to_export_dict(session: ResearchSession) -> dict[str, Any]:
    """Convert session to export-friendly dictionary."""
    return {
        "interaction_id": session.interaction_id,
        "query": session.query,
        "title": session.title,
        "summary": session.summary,
        "report_text": session.report_text,
        "format_instructions": session.format_instructions,
        "agent_name": session.agent_name,
        "duration_seconds": session.duration_seconds,
        "total_tokens": session.total_tokens,
        "tags": session.tags,
        "notes": session.notes,
        "created_at": session.created_at_iso,
        "expires_at": session.expires_at_iso,
        "export_timestamp": datetime.now(tz=UTC).isoformat(),
    }


def export_to_json(session: ResearchSession) -> ExportResult:
    """Export a research session to JSON format."""
    data = _session_to_export_dict(session)
    content = json.dumps(data, indent=2, ensure_ascii=False)
    filename = _generate_filename(session, "json")

    return ExportResult(
        format=ExportFormat.JSON,
        filename=filename,
        content=content.encode("utf-8"),
        mime_type="application/json",
    )


# =============================================================================
# DOCX Export (using Marko + python-docx)
# =============================================================================


def _configure_document_styles(document: Any) -> None:
    """
    Configure professional document styles for a polished, Pandoc-quality look.

    Sets up:
    - Page margins (1-inch standard)
    - Base typography (Calibri 11pt, professional line spacing)
    - Heading styles with proper hierarchy, colors, and spacing
    - Keep-with-next for headings to prevent orphans
    """
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Cm, Pt, RGBColor

    # Professional color palette
    NAVY_BLUE = RGBColor(0x1F, 0x49, 0x7D)  # #1F497D - Professional navy
    DARK_GRAY = RGBColor(0x33, 0x33, 0x33)  # #333333 - Body text

    # --- Page Margins (1-inch all around for professional look) ---
    for section in document.sections:
        section.top_margin = Cm(2.54)  # 1 inch
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # --- Normal (body text) style ---
    try:
        normal = document.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.font.color.rgb = DARK_GRAY
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        normal.paragraph_format.space_after = Pt(8)
    except KeyError:
        pass  # Style might not exist in all templates

    # --- Heading styles with professional formatting ---
    heading_configs = [
        # (style_name, font_size, bold, color, space_before, space_after)
        ("Heading 1", 18, True, NAVY_BLUE, 24, 12),
        ("Heading 2", 16, True, NAVY_BLUE, 18, 8),
        ("Heading 3", 14, True, NAVY_BLUE, 14, 6),
        ("Heading 4", 12, True, NAVY_BLUE, 12, 4),
    ]

    for style_name, size, bold, color, before, after in heading_configs:
        try:
            style = document.styles[style_name]
            style.font.name = "Calibri"
            style.font.size = Pt(size)
            style.font.bold = bold
            style.font.color.rgb = color
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True  # Prevent orphan headings
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        except KeyError:
            pass

    # --- Title style (for cover page) ---
    try:
        title = document.styles["Title"]
        title.font.name = "Calibri Light"
        title.font.size = Pt(28)
        title.font.color.rgb = NAVY_BLUE
        title.font.bold = False
        title.paragraph_format.space_after = Pt(24)
    except KeyError:
        pass

    # --- Subtitle style ---
    try:
        subtitle = document.styles["Subtitle"]
        subtitle.font.name = "Calibri"
        subtitle.font.size = Pt(14)
        subtitle.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        subtitle.font.italic = True
    except KeyError:
        pass

    # --- TOC styles ---
    for level in range(1, 4):
        toc_style_name = f"TOC {level}"
        try:
            toc_style = document.styles[toc_style_name]
            toc_style.font.name = "Calibri"
            toc_style.font.size = Pt(12 - level + 1)  # 12, 11, 10
            toc_style.paragraph_format.space_before = Pt(4 - level + 1)
            toc_style.paragraph_format.space_after = Pt(2)
        except KeyError:
            pass

    # --- List styles ---
    for list_style in ["List Bullet", "List Number"]:
        try:
            style = document.styles[list_style]
            style.font.name = "Calibri"
            style.font.size = Pt(11)
            style.paragraph_format.space_after = Pt(4)
        except KeyError:
            pass


def _add_hyperlink(paragraph: Any, url: str, text: str) -> Any:
    """
    Add a proper clickable hyperlink to a paragraph.

    Creates a Word hyperlink element with proper relationship registration.
    Returns the created hyperlink element.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Relationship type for hyperlinks
    rel_type = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
    )

    # Get the document part to register the relationship
    part = paragraph.part
    r_id = part.relate_to(url, rel_type, is_external=True)

    # Create the w:hyperlink element
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create a run inside the hyperlink
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Style as link (blue, underlined)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")  # Standard link blue
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    new_run.append(rPr)

    # Add the text
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


def _create_bookmark_id(title: str, index: int) -> str:
    """
    Create a valid Word bookmark ID from a heading title.

    Bookmark IDs must:
    - Start with a letter
    - Contain only letters, digits, and underscores
    - Be at most 40 characters

    We use a simple index-based approach for reliability, with a short
    prefix from the title for human readability when debugging.
    """
    import re

    # Extract just alphanumeric characters, take first 15 chars
    safe = re.sub(r"[^a-zA-Z0-9]", "", title)
    safe = safe[:15] if safe else "heading"

    # Ensure it starts with a letter
    if not safe[0].isalpha():
        safe = "h" + safe

    # Use index as primary identifier for reliability
    return f"_bm_{index}_{safe}"


def _add_bookmark_to_paragraph(paragraph: Any, bookmark_id: str) -> None:
    """
    Add a bookmark to a paragraph for internal document navigation.

    Creates a Word bookmark element that can be targeted by internal hyperlinks.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Generate a unique bookmark ID number
    # Word uses incrementing IDs, we'll use hash of the bookmark name
    bookmark_num = str(abs(hash(bookmark_id)) % 100000)

    # Create bookmark start
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), bookmark_num)
    bookmark_start.set(qn("w:name"), bookmark_id)

    # Create bookmark end
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), bookmark_num)

    # Insert at the beginning of the paragraph
    paragraph._p.insert(0, bookmark_start)
    paragraph._p.append(bookmark_end)


def _add_internal_hyperlink(paragraph: Any, bookmark_id: str, text: str) -> Any:
    """
    Add an internal hyperlink (to a bookmark) to a paragraph.

    Creates a clickable link that navigates to the specified bookmark within the document.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Create the w:hyperlink element with internal anchor
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_id)

    # Create a run inside the hyperlink
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Style as link (blue, underlined)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")  # Standard link blue
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    new_run.append(rPr)

    # Add the text
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


def _add_paragraph_shading(paragraph: Any, color_hex: str) -> None:
    """
    Add background shading to a paragraph.

    Args:
        paragraph: The python-docx Paragraph object
        color_hex: Hex color string without # (e.g., "F5F5F5")
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Get or create pPr element
    pPr = paragraph._p.get_or_add_pPr()

    # Create shading element
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)

    pPr.append(shd)


def _add_paragraph_border(paragraph: Any) -> None:
    """
    Add a subtle border around a paragraph (for code blocks).

    Creates a light grey border similar to GitHub code blocks.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Get or create pPr element
    pPr = paragraph._p.get_or_add_pPr()

    # Create pBdr (paragraph border) element
    pBdr = OxmlElement("w:pBdr")

    # Add borders on all sides
    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")  # 0.5pt border
        border.set(qn("w:space"), "4")  # 4pt padding
        border.set(qn("w:color"), "D1D5DA")  # GitHub border grey
        pBdr.append(border)

    pPr.append(pBdr)


def _render_code_block(
    document: Any,
    code: str,
    language: str | None = None,
) -> None:
    """
    Render a code block with optional syntax highlighting using Pygments.

    Creates a GitHub-style code block with:
    - Light grey background (#F6F8FA)
    - Subtle border
    - Syntax highlighting if Pygments is available and language is recognized

    Args:
        document: The python-docx Document object
        code: The code text to render
        language: Optional language identifier (e.g., "python", "javascript")
    """
    from docx.shared import Pt, RGBColor

    # Try to use Pygments for syntax highlighting
    tokens = None
    if language:
        try:
            from pygments import lex
            from pygments.lexers import get_lexer_by_name

            lexer = get_lexer_by_name(language, stripall=True)
            tokens = list(lex(code, lexer))
        except Exception:
            # Pygments not available or language not recognized - fall back to plain
            tokens = None

    # Create paragraph with GitHub-style formatting
    para = document.add_paragraph()
    para.paragraph_format.left_indent = Pt(12)
    para.paragraph_format.right_indent = Pt(12)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(8)
    _add_paragraph_shading(para, "F6F8FA")  # GitHub code background
    _add_paragraph_border(para)

    if tokens:
        # Render with syntax highlighting
        _render_highlighted_tokens(para, tokens)
    else:
        # Plain monospace rendering
        run = para.add_run(code)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)  # GitHub dark text


def _render_highlighted_tokens(paragraph: Any, tokens: list[tuple[Any, str]]) -> None:
    """
    Render Pygments tokens to a paragraph with appropriate colors.

    Uses a GitHub-inspired color scheme for syntax highlighting.

    Args:
        paragraph: The python-docx Paragraph object
        tokens: List of (token_type, token_value) tuples from Pygments
    """
    from docx.shared import Pt, RGBColor

    # Import Pygments Token types
    try:
        from pygments.token import (
            Comment,
            Error,
            Keyword,
            Literal,
            Name,
            Number,
            Operator,
            Punctuation,
            String,
        )
    except ImportError:
        return

    # GitHub-inspired color scheme
    # Based on github.com's syntax highlighting
    COLOR_MAP = {
        # Keywords (purple)
        Keyword: RGBColor(0xCF, 0x22, 0x2E),  # Red for keywords
        Keyword.Constant: RGBColor(0x00, 0x5C, 0xC5),  # Blue
        Keyword.Declaration: RGBColor(0xCF, 0x22, 0x2E),
        Keyword.Namespace: RGBColor(0xCF, 0x22, 0x2E),
        Keyword.Pseudo: RGBColor(0x00, 0x5C, 0xC5),
        Keyword.Reserved: RGBColor(0xCF, 0x22, 0x2E),
        Keyword.Type: RGBColor(0x00, 0x5C, 0xC5),
        # Names
        Name.Builtin: RGBColor(0x00, 0x5C, 0xC5),  # Blue for builtins
        Name.Class: RGBColor(0x95, 0x3D, 0xAE),  # Purple for classes
        Name.Constant: RGBColor(0x00, 0x5C, 0xC5),
        Name.Decorator: RGBColor(0x95, 0x3D, 0xAE),  # Purple for decorators
        Name.Exception: RGBColor(0x95, 0x3D, 0xAE),
        Name.Function: RGBColor(0x6F, 0x42, 0xC1),  # Purple for functions
        Name.Namespace: RGBColor(0x24, 0x29, 0x2E),
        Name.Tag: RGBColor(0x22, 0x86, 0x3A),  # Green for tags (HTML/XML)
        Name.Variable: RGBColor(0xE3, 0x6D, 0x09),  # Orange for variables
        # Strings (blue)
        String: RGBColor(0x03, 0x2F, 0x62),  # Dark blue for strings
        String.Doc: RGBColor(0x6A, 0x73, 0x7D),  # Grey for docstrings
        String.Escape: RGBColor(0x00, 0x5C, 0xC5),
        String.Interpol: RGBColor(0x00, 0x5C, 0xC5),
        String.Regex: RGBColor(0x03, 0x2F, 0x62),
        # Numbers (blue)
        Number: RGBColor(0x00, 0x5C, 0xC5),
        # Comments (grey)
        Comment: RGBColor(0x6A, 0x73, 0x7D),
        Comment.Multiline: RGBColor(0x6A, 0x73, 0x7D),
        Comment.Single: RGBColor(0x6A, 0x73, 0x7D),
        # Operators
        Operator: RGBColor(0xCF, 0x22, 0x2E),  # Red
        Operator.Word: RGBColor(0xCF, 0x22, 0x2E),
        # Literals
        Literal: RGBColor(0x00, 0x5C, 0xC5),
        # Punctuation
        Punctuation: RGBColor(0x24, 0x29, 0x2E),
        # Error
        Error: RGBColor(0xCB, 0x24, 0x31),  # Red for errors
    }

    # Default color (dark text)
    DEFAULT_COLOR = RGBColor(0x24, 0x29, 0x2E)

    for token_type, token_value in tokens:
        if not token_value:
            continue

        run = paragraph.add_run(token_value)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

        # Find color for this token type (check ancestors if exact match not found)
        color = None
        current_type = token_type
        while current_type is not None:
            if current_type in COLOR_MAP:
                color = COLOR_MAP[current_type]
                break
            # Move up the token hierarchy
            current_type = current_type.parent if hasattr(current_type, "parent") else None

        run.font.color.rgb = color if color else DEFAULT_COLOR


def _add_toc_field(
    document: Any,
    sections: list[tuple[str, int, str]] | None = None,
) -> None:
    """
    Add a native Word Table of Contents field with clickable preview entries.

    This inserts TOC field codes that Word will populate when opened.
    The TOC auto-updates when the user presses F9 or updates fields in Word.
    If sections are provided, includes clickable preview entries that link to bookmarks.

    Args:
        document: The python-docx Document object
        sections: List of (title, level, bookmark_id) tuples for TOC entries
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    run = paragraph.add_run()

    # Create TOC field
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r' TOC \o "1-3" \h \z \u '

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    # Add elements to run
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)

    # Add clickable preview TOC entries if sections provided
    if sections:
        for title, level, bookmark_id in sections:  # All entries
            toc_para = document.add_paragraph()
            indent_amount = 360000 * (level - 1)  # 0.5 inch per level
            toc_para.paragraph_format.left_indent = indent_amount
            toc_para.paragraph_format.space_before = Pt(2)
            toc_para.paragraph_format.space_after = Pt(2)

            # 🎯 Create clickable internal hyperlink to the heading bookmark
            _add_internal_hyperlink(toc_para, bookmark_id, title)
    else:
        paragraph.add_run(
            "Table of Contents\n(Update this field in Word: References → Update Table)"
        )

    # Close the field
    end_run = paragraph.add_run()
    end_run._r.append(fld_char_end)


def _render_inline_to_paragraph(paragraph: Any, element: Any) -> None:
    """Render inline Marko elements to a python-docx paragraph with full formatting."""
    from marko import inline as marko_inline

    if isinstance(element, marko_inline.RawText):
        run = paragraph.add_run(element.children)
    elif isinstance(element, str):
        paragraph.add_run(element)
    elif isinstance(element, marko_inline.Emphasis):
        # Italic
        text = _get_text_content(element)
        run = paragraph.add_run(text)
        run.italic = True
    elif isinstance(element, marko_inline.StrongEmphasis):
        # Bold
        text = _get_text_content(element)
        run = paragraph.add_run(text)
        run.bold = True
    elif isinstance(element, marko_inline.CodeSpan):
        # Inline code - use monospace font with background
        from docx.shared import RGBColor
        run = paragraph.add_run(element.children)
        run.font.name = "Consolas"
        run.font.color.rgb = RGBColor(0x88, 0x00, 0x00)  # Dark red for code
    elif isinstance(element, marko_inline.Link):
        # 🎯 Proper clickable hyperlink with full URL display
        text = _get_text_content(element)
        url = element.dest if hasattr(element, "dest") else ""
        if url and url.startswith(("http://", "https://")):
            # Use full URL as display text for better traceability in sources
            display_text = url
            _add_hyperlink(paragraph, url, display_text)
        else:
            # Fallback for non-URL links
            run = paragraph.add_run(text)
            run.underline = True
    elif isinstance(element, marko_inline.LineBreak):
        paragraph.add_run().add_break()
    elif hasattr(element, "children") and element.children:
        if isinstance(element.children, list):
            for child in element.children:
                _render_inline_to_paragraph(paragraph, child)
        else:
            paragraph.add_run(str(element.children))


def _render_inline_to_run(run: Any, element: Any) -> None:
    """Render inline Marko elements to a python-docx Run with formatting (legacy)."""
    from marko import inline as marko_inline

    if isinstance(element, marko_inline.RawText):
        run.add_text(element.children)
    elif isinstance(element, str):
        run.add_text(element)
    elif isinstance(element, marko_inline.Emphasis):
        # Italic
        for child in element.children:
            _render_inline_to_run(run, child)
        run.italic = True
    elif isinstance(element, marko_inline.StrongEmphasis):
        # Bold
        for child in element.children:
            _render_inline_to_run(run, child)
        run.bold = True
    elif isinstance(element, marko_inline.CodeSpan):
        # Inline code - use monospace font
        run.add_text(element.children)
        run.font.name = "Consolas"
    elif isinstance(element, marko_inline.Link):
        # Hyperlink text - simplified for run context
        for child in element.children:
            _render_inline_to_run(run, child)
        run.underline = True
    elif isinstance(element, marko_inline.LineBreak):
        run.add_break()
    elif hasattr(element, "children") and element.children:
        for child in element.children:
            _render_inline_to_run(run, child)


def _get_text_content(element: Any) -> str:
    """Extract plain text content from a Marko element recursively."""
    from marko import inline as marko_inline

    if isinstance(element, str):
        return element
    if isinstance(element, marko_inline.RawText):
        return element.children if isinstance(element.children, str) else ""
    if hasattr(element, "children"):
        if isinstance(element.children, str):
            return element.children
        if isinstance(element.children, list):
            return "".join(_get_text_content(child) for child in element.children)
    return ""


def _extract_headings(parsed: Any) -> list[tuple[str, int, str]]:
    """
    Extract heading titles, levels, and bookmark IDs from parsed Markdown for TOC.

    Returns a list of (title, level, bookmark_id) tuples. The bookmark_id is used
    to create clickable links in the TOC that navigate to the corresponding heading.
    """
    from marko import block as marko_block

    headings: list[tuple[str, int, str]] = []
    heading_index = 0

    def walk(element: Any) -> None:
        nonlocal heading_index
        if isinstance(element, marko_block.Heading):
            text = _get_text_content(element)
            bookmark_id = _create_bookmark_id(text, heading_index)
            headings.append((text, element.level, bookmark_id))
            heading_index += 1
        if hasattr(element, "children") and isinstance(element.children, list):
            for child in element.children:
                walk(child)

    walk(parsed)
    return headings


def _render_block_to_docx(  # noqa: C901, PLR0912
    document: Any,
    element: Any,
    *,
    list_level: int = 0,
    heading_counter: list[int] | None = None,
    heading_bookmarks: list[str] | None = None,
) -> None:
    """
    Render a Marko block element to a python-docx Document.

    Handles headings, paragraphs, lists, code blocks, blockquotes, and tables.
    Uses professional styling with proper spacing instead of page breaks.

    Args:
        document: The python-docx Document object
        element: The Marko block element to render
        list_level: Current nesting level for lists
        heading_counter: Mutable counter [index] for tracking heading positions
        heading_bookmarks: List of bookmark IDs matching heading order from _extract_headings
    """
    from docx.shared import RGBColor
    from marko import block as marko_block

    if isinstance(element, marko_block.Heading):
        # Add heading with proper level
        text = _get_text_content(element)
        heading = document.add_heading(text, level=element.level)

        # 🎯 Add bookmark for TOC navigation
        if heading_counter is not None and heading_bookmarks is not None:
            idx = heading_counter[0]
            if idx < len(heading_bookmarks):
                bookmark_id = heading_bookmarks[idx]
                _add_bookmark_to_paragraph(heading, bookmark_id)
                heading_counter[0] += 1

        # Professional styling relies on proper heading spacing (space_before/space_after)
        # and keep_with_next instead of page breaks for a flowing document

        # Style H2 with accent color
        if element.level == 2:
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # Professional blue

    elif isinstance(element, marko_block.Paragraph):
        para = document.add_paragraph()
        # Use the new paragraph-level renderer for proper hyperlinks
        for child in element.children:
            _render_inline_to_paragraph(para, child)
        
        # 🎯 Page break before Sources section for professional layout
        text = _get_text_content(element).strip()
        if text.lower().startswith("sources:") or text.lower() == "sources":
            para.paragraph_format.page_break_before = True

    elif isinstance(element, marko_block.List):
        for item_idx, item in enumerate(element.children):
            if isinstance(item, marko_block.ListItem):
                # Create the list paragraph
                if element.ordered:
                    # For ordered lists, manually number starting from 1
                    # (Gemini sometimes outputs lists starting from arbitrary numbers)
                    para = document.add_paragraph()
                    # Add the number manually
                    num_run = para.add_run(f"{item_idx + 1}. ")
                    num_run.bold = False
                else:
                    para = document.add_paragraph(style="List Bullet")

                # Render content with proper formatting (including links)
                for child in item.children:
                    if hasattr(child, "children"):
                        for inline in child.children:
                            _render_inline_to_paragraph(para, inline)

    elif isinstance(element, marko_block.CodeBlock):
        # Code block with GitHub-like styling (no language info available)
        text = _get_text_content(element)
        _render_code_block(document, text, language=None)

    elif isinstance(element, marko_block.FencedCode):
        # Fenced code block (```language ... ```) - with syntax highlighting
        text = _get_text_content(element)
        # Get language from the fenced code info string
        language = getattr(element, "lang", None) or None
        _render_code_block(document, text, language=language)

    elif isinstance(element, marko_block.Quote):
        # Blockquote - indent and italicize
        for child in element.children:
            if isinstance(child, marko_block.Paragraph):
                text = _get_text_content(child)
                para = document.add_paragraph()
                run = para.add_run(text)
                run.italic = True
                para.paragraph_format.left_indent = 720000  # 1 inch

    elif isinstance(element, marko_block.ThematicBreak):
        # Skip horizontal rules - page breaks between H2 sections provide sufficient separation
        pass

    elif isinstance(element, marko_block.BlankLine):
        # Skip blank lines (they're handled by paragraph spacing)
        pass

    elif isinstance(element, marko_block.HTMLBlock):
        # Skip raw HTML blocks in DOCX
        pass

    else:
        # Check for GFM table elements
        element_type = type(element).__name__

        if element_type == "Table":
            # GFM Table - render as proper Word table
            _render_gfm_table(document, element)

        elif hasattr(element, "children") and isinstance(element.children, list):
            # Generic container - recurse into children
            for child in element.children:
                _render_block_to_docx(
                    document,
                    child,
                    list_level=list_level,
                    heading_counter=heading_counter,
                    heading_bookmarks=heading_bookmarks,
                )


def _render_gfm_table(document: Any, table_element: Any) -> None:
    """
    Render a GFM (GitHub Flavored Markdown) table to a Word table.

    Handles:
    - Table headers (first row, bold)
    - Table body rows
    - Cell content with basic formatting
    - Both wrapped (TableHead/TableBody) and unwrapped (direct TableRow) formats

    Args:
        document: The python-docx Document object
        table_element: The GFM Table element from marko parser
    """
    from docx.shared import Pt, RGBColor

    # Extract rows from the table element
    rows_data: list[list[str]] = []
    has_header = False

    for child in table_element.children:
        child_type = type(child).__name__

        if child_type == "TableHead":
            has_header = True
            for row in child.children:
                row_type = type(row).__name__
                if row_type == "TableRow":
                    cells = []
                    for cell in row.children:
                        cells.append(_get_text_content(cell).strip())
                    if cells:
                        rows_data.append(cells)

        elif child_type == "TableBody":
            for row in child.children:
                row_type = type(row).__name__
                if row_type == "TableRow":
                    cells = []
                    for cell in row.children:
                        cells.append(_get_text_content(cell).strip())
                    if cells:
                        rows_data.append(cells)

        elif child_type == "TableRow":
            # Direct TableRow children (no TableHead/TableBody wrappers)
            # First row is treated as header
            if not rows_data:
                has_header = True
            cells = []
            for cell in child.children:
                cells.append(_get_text_content(cell).strip())
            if cells:
                rows_data.append(cells)

    if not rows_data:
        return

    # Determine number of columns
    num_cols = max(len(row) for row in rows_data)
    num_rows = len(rows_data)

    # Create the Word table
    table = document.add_table(rows=num_rows, cols=num_cols)
    table.style = "Table Grid"

    # Professional colors
    HEADER_BG = "E7EFF8"  # Light blue background
    NAVY_BLUE = RGBColor(0x1F, 0x49, 0x7D)

    # Fill the table
    for row_idx, row_data in enumerate(rows_data):
        row = table.rows[row_idx]
        is_header = has_header and row_idx == 0

        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= num_cols:
                break

            cell = row.cells[col_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            run.font.size = Pt(10)
            run.font.name = "Calibri"

            if is_header:
                run.bold = True
                run.font.color.rgb = NAVY_BLUE
                # Set header cell background
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                cell_props = cell._tc.get_or_add_tcPr()
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), HEADER_BG)
                cell_props.append(shading)

    # Space after table is handled by next paragraph's space_before


def _add_cover_page(document: Any, session: ResearchSession) -> None:
    """Add a professional, clean cover page to the document."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    # Professional color palette
    NAVY_BLUE = RGBColor(0x1F, 0x49, 0x7D)
    DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
    SUBTLE_GRAY = RGBColor(0x88, 0x88, 0x88)
    LIGHT_GRAY = RGBColor(0xAA, 0xAA, 0xAA)

    # Extract clean title (handles clarified queries)
    title = _extract_clean_title(session.query, session.title)

    # Add vertical space at the top for balanced layout
    for _ in range(4):
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)

    # Title - elegant, large, centered
    title_para = document.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(36)
    for run in title_para.runs:
        run.font.name = "Calibri Light"
        run.font.size = Pt(32)
        run.font.color.rgb = NAVY_BLUE
        run.bold = False

    # Generous vertical space
    for _ in range(3):
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)

    # Date - prominent, centered
    date_str = datetime.fromtimestamp(session.created_at, tz=UTC).strftime("%B %d, %Y")
    date_para = document.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(date_str)
    date_run.font.name = "Calibri"
    date_run.font.size = Pt(14)
    date_run.bold = True
    date_run.font.color.rgb = DARK_GRAY
    date_para.paragraph_format.space_after = Pt(12)

    # Research metrics in elegant format
    meta_parts = []
    if session.duration_seconds:
        mins = int(session.duration_seconds // 60)
        secs = int(session.duration_seconds % 60)
        meta_parts.append(f"Research Duration: {mins}m {secs}s")

    if session.total_tokens:
        meta_parts.append(f"Tokens: {session.total_tokens:,}")

    if meta_parts:
        meta_para = document.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_para.add_run(" • ".join(meta_parts))
        meta_run.font.name = "Calibri"
        meta_run.font.size = Pt(11)
        meta_run.font.color.rgb = SUBTLE_GRAY
        meta_para.paragraph_format.space_after = Pt(6)

    # Agent name
    if session.agent_name:
        agent_para = document.add_paragraph()
        agent_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        agent_run = agent_para.add_run(f"AI Agent: {session.agent_name}")
        agent_run.font.name = "Calibri"
        agent_run.font.size = Pt(11)
        agent_run.font.color.rgb = SUBTLE_GRAY
        agent_para.paragraph_format.space_after = Pt(6)

    # Push branding to bottom of page
    for _ in range(6):
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)

    # Branding footer on cover page
    brand_para = document.add_paragraph()
    brand_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand_run = brand_para.add_run("Generated by Gemini Research MCP")
    brand_run.font.name = "Calibri"
    brand_run.font.size = Pt(9)
    brand_run.font.color.rgb = LIGHT_GRAY
    brand_run.italic = True

    # Page break after cover
    document.add_page_break()


def _add_metadata_table(document: Any, session: ResearchSession) -> None:
    """Add a professional metadata information table to the document."""
    from docx.shared import Inches, Pt, RGBColor

    # Create table
    rows_data = [
        ("Research Query", session.query),
        ("Created", session.created_at_iso),
    ]

    if session.duration_seconds:
        mins = int(session.duration_seconds // 60)
        secs = int(session.duration_seconds % 60)
        rows_data.append(("Duration", f"{mins}m {secs}s"))

    if session.total_tokens:
        rows_data.append(("Tokens Used", f"{session.total_tokens:,}"))

    if session.agent_name:
        rows_data.append(("AI Agent", session.agent_name))

    if session.tags:
        rows_data.append(("Tags", ", ".join(session.tags)))

    if session.notes:
        rows_data.append(("Notes", session.notes))

    rows_data.append(("Session ID", session.interaction_id))

    if session.expires_at_iso:
        rows_data.append(("Expires", session.expires_at_iso))

    # Create the table with professional styling
    table = document.add_table(rows=len(rows_data), cols=2)
    table.style = "Table Grid"

    # Style each row
    for row_idx, (field, value) in enumerate(rows_data):
        row = table.rows[row_idx]

        # Field name cell - bold, styled
        cell0 = row.cells[0]
        cell0.text = ""
        para0 = cell0.paragraphs[0]
        run0 = para0.add_run(field)
        run0.bold = True
        run0.font.size = Pt(10)
        run0.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # Professional blue
        cell0.width = Inches(1.5)

        # Value cell
        cell1 = row.cells[1]
        cell1.text = ""
        para1 = cell1.paragraphs[0]
        run1 = para1.add_run(str(value))
        run1.font.size = Pt(10)
        cell1.width = Inches(4.5)

    # Space after table is handled by next paragraph's space_before


def export_to_docx(
    session: ResearchSession,
    *,
    include_toc: bool = True,
    include_cover_page: bool = True,
    toc_levels: int = 3,  # noqa: ARG001 - kept for API compatibility
) -> ExportResult:
    """
    Export a research session to DOCX format using Marko + python-docx.

    Creates a professionally formatted Word document with:
    - Cover page with title and metadata (optional)
    - Automatic Table of Contents field (optional)
    - Executive summary as blockquote
    - Metadata table
    - Full research report with proper formatting (including GFM tables)
    - Page breaks before H1 headings (excellence feature)
    - Support for all CommonMark + GFM Markdown features

    Args:
        session: The research session to export
        include_toc: Whether to include a Table of Contents (default: True)
        include_cover_page: Whether to include a cover page (default: True)
        toc_levels: Number of heading levels in TOC (kept for API compat)

    Requires marko and python-docx packages.
    Install with: pip install 'gemini-research-mcp[docx]'
    """
    try:
        import marko
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor
    except ImportError as e:
        raise ImportError(
            "marko and python-docx are required for DOCX export. "
            "Install with: pip install 'gemini-research-mcp[docx]'"
        ) from e

    # Create new Word document
    document = Document()

    # 🎯 Apply professional document styling (margins, fonts, colors, spacing)
    _configure_document_styles(document)

    # Pre-parse the report to extract headings for TOC with bookmark IDs
    toc_sections: list[tuple[str, int, str]] = []
    heading_bookmarks: list[str] = []
    parsed = None
    if session.report_text:
        # Use GFM extension for table support
        md = marko.Markdown(extensions=["gfm"])
        parsed = md.parse(session.report_text)
        toc_sections = _extract_headings(parsed)

        # Extract just the bookmark IDs for use during rendering
        heading_bookmarks = [bookmark_id for _, _, bookmark_id in toc_sections]

    # Add cover page
    if include_cover_page:
        _add_cover_page(document, session)

    # Add Table of Contents (as a Word field with clickable preview entries)
    if include_toc:
        toc_heading = document.add_heading("Table of Contents", level=1)
        toc_heading.paragraph_format.page_break_before = False
        # Style the TOC heading
        for run in toc_heading.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        # Add pre-populated clickable TOC entries from the report
        _add_toc_field(document, toc_sections)
        # Note: No page break needed here - Document Info has page_break_before=True

    # Document Information section - on its own page for professional look
    info_heading = document.add_heading("Document Information", level=1)
    # Style it professionally
    for run in info_heading.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    # Page break before Document Information (new page after TOC)
    info_heading.paragraph_format.page_break_before = True
    _add_metadata_table(document, session)

    # Executive Summary (if available)
    if session.summary:
        summary_heading = document.add_heading("Executive Summary", level=1)
        for run in summary_heading.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        # No page break - flows naturally after Document Information
        summary_heading.paragraph_format.space_before = Pt(24)

        # Render summary in an elegant box-like style
        summary_para = document.add_paragraph()
        summary_para.paragraph_format.left_indent = Pt(18)
        summary_para.paragraph_format.right_indent = Pt(18)
        summary_para.paragraph_format.space_before = Pt(12)
        summary_para.paragraph_format.space_after = Pt(12)
        run = summary_para.add_run(session.summary)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Main Research Report
    if parsed:
        # Render each block element to the Word document
        # Pass bookmark list and mutable counter so headings get bookmarks for TOC links
        heading_counter = [0]  # Mutable counter to track which heading we're on

        for element in parsed.children:
            _render_block_to_docx(
                document,
                element,
                heading_counter=heading_counter,
                heading_bookmarks=heading_bookmarks,
            )

    # Add attribution to Word's footer section (margin area at bottom of pages)
    # This avoids page break issues that occur with body paragraphs
    section = document.sections[-1]  # Last section
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.clear()  # Clear any existing content
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen_run = footer_para.add_run(
        f"Generated by Gemini Research MCP • "
        f"{datetime.now(tz=UTC).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    gen_run.font.name = "Calibri"
    gen_run.italic = True
    gen_run.font.size = Pt(9)
    gen_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Save to bytes
    output = BytesIO()
    document.save(output)
    content = output.getvalue()

    filename = _generate_filename(session, "docx")

    return ExportResult(
        format=ExportFormat.DOCX,
        filename=filename,
        content=content,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# =============================================================================
# Helper Functions
# =============================================================================


def _extract_clean_title(query: str, title: str | None = None) -> str:
    """
    Extract a clean title from a query that may contain clarification context.

    When elicitation occurs, the query is refined to include:
    "Original query\n\nAdditional context:\nQ: ...\nA: ..."

    This function extracts just the original query portion for display.
    """
    if title:
        return title

    # Check for clarification marker
    if "\n\nAdditional context:" in query:
        # Extract just the original query before clarification
        clean = query.split("\n\nAdditional context:")[0].strip()
        return clean[:60] if len(clean) > 60 else clean

    return query[:60] if len(query) > 60 else query


def _generate_filename(session: ResearchSession, extension: str) -> str:
    """Generate a safe filename from session metadata."""
    # Use title or extract clean title from query
    base = _extract_clean_title(session.query, session.title)

    # Clean up for filename
    safe = re.sub(r"[^\w\s-]", "", base)  # Remove special chars
    safe = re.sub(r"\s+", "_", safe)  # Replace spaces with underscores
    safe = safe[:50]  # Limit length

    # Add timestamp
    timestamp = datetime.fromtimestamp(session.created_at, tz=UTC).strftime("%Y%m%d")

    return f"{safe}_{timestamp}.{extension}"


def export_session(
    session: ResearchSession,
    format: ExportFormat | str,
    output_path: Path | str | None = None,
) -> ExportResult:
    """
    Export a research session to the specified format.

    Args:
        session: The research session to export
        format: Export format (markdown, json, docx)
        output_path: Optional path to save the file (if None, returns bytes only)

    Returns:
        ExportResult with format, filename, content bytes, and mime_type
    """
    # Normalize format
    if isinstance(format, str):
        format_str = format.lower()
        if format_str in ("md", "markdown"):
            export_format = ExportFormat.MARKDOWN
        elif format_str == "json":
            export_format = ExportFormat.JSON
        elif format_str in ("docx", "word"):
            export_format = ExportFormat.DOCX
        else:
            raise ValueError(f"Unsupported format: {format}. Use: markdown, json, docx")
    else:
        export_format = format

    # Export
    if export_format == ExportFormat.MARKDOWN:
        result = export_to_markdown(session)
    elif export_format == ExportFormat.JSON:
        result = export_to_json(session)
    elif export_format == ExportFormat.DOCX:
        result = export_to_docx(session)
    else:
        raise ValueError(f"Unsupported format: {export_format}")

    # Save to file if path provided
    if output_path is not None:
        path = Path(output_path)
        path.write_bytes(result.content)
        logger.info("📄 Exported to %s (%s)", path, result.size_human)

    return result


# =============================================================================
# Convenience Functions
# =============================================================================


def get_supported_formats() -> list[str]:
    """Return list of supported export formats."""
    return [f.value for f in ExportFormat]
