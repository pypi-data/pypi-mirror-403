"""Tests for research session export functionality."""

from __future__ import annotations

import json
import time
from pathlib import Path

import pytest

from gemini_research_mcp.export import (
    ExportFormat,
    ExportResult,
    export_session,
    export_to_docx,
    export_to_json,
    export_to_markdown,
    get_supported_formats,
)
from gemini_research_mcp.storage import ResearchSession

# =============================================================================
# Fixtures
# =============================================================================


@pytest.fixture
def sample_session() -> ResearchSession:
    """Create a sample research session for export testing."""
    return ResearchSession(
        interaction_id="test-export-12345-abcdef",
        query="What are the best practices for quantum computing security?",
        created_at=time.time(),
        title="Quantum Computing Security Research",
        summary="This research explores quantum computing security best practices including post-quantum cryptography and quantum key distribution.",
        report_text="""## Executive Summary

Quantum computing poses significant challenges to current cryptographic systems.

### Key Findings

1. **Post-Quantum Cryptography**: NIST has standardized several algorithms
2. **Quantum Key Distribution (QKD)**: Offers theoretically unbreakable encryption
3. **Timeline**: Experts estimate 10-15 years until cryptographically relevant quantum computers

### Recommendations

- Begin migration to post-quantum algorithms
- Implement crypto-agility in systems
- Monitor NIST standards developments

**Sources:**
1. [NIST Post-Quantum Cryptography](https://example.com/nist)
2. [IBM Quantum Research](https://example.com/ibm)
""",
        format_instructions="Create an executive briefing",
        agent_name="deep-research-pro-preview-12-2025",
        duration_seconds=342.5,
        total_tokens=15000,
        tags=["security", "quantum", "cryptography"],
        notes="Priority research for Q1 security review",
    )


@pytest.fixture
def minimal_session() -> ResearchSession:
    """Create a minimal session with only required fields."""
    return ResearchSession(
        interaction_id="minimal-session-123",
        query="Simple query",
        created_at=time.time(),
    )


# =============================================================================
# Export Format Tests
# =============================================================================


class TestExportFormats:
    """Tests for supported export formats."""

    def test_get_supported_formats(self) -> None:
        """Test that we return the expected formats."""
        formats = get_supported_formats()
        assert "markdown" in formats
        assert "json" in formats
        assert "docx" in formats
        assert len(formats) == 3


# =============================================================================
# Markdown Export Tests
# =============================================================================


class TestMarkdownExport:
    """Tests for Markdown export functionality."""

    def test_export_to_markdown_basic(self, sample_session: ResearchSession) -> None:
        """Test basic Markdown export."""
        result = export_to_markdown(sample_session)

        assert result.format == ExportFormat.MARKDOWN
        assert result.filename.endswith(".md")
        assert result.mime_type == "text/markdown"
        assert len(result.content) > 0

    def test_markdown_content_structure(self, sample_session: ResearchSession) -> None:
        """Test Markdown content includes expected sections."""
        result = export_to_markdown(sample_session)
        content = result.content.decode("utf-8")

        # Title
        assert "# Quantum Computing Security Research" in content

        # Metadata section
        assert "## Metadata" in content
        assert "**Query:**" in content
        assert "**Created:**" in content
        assert "**Duration:**" in content
        assert "**Tokens:**" in content
        assert "**Agent:**" in content
        assert "**Tags:**" in content
        assert "security, quantum, cryptography" in content

        # Summary section
        assert "## Summary" in content
        assert "post-quantum cryptography" in content

        # Report section
        assert "## Research Report" in content
        assert "Executive Summary" in content

        # Footer
        assert "Exported from Gemini Research MCP" in content

    def test_markdown_minimal_session(self, minimal_session: ResearchSession) -> None:
        """Test Markdown export with minimal session data."""
        result = export_to_markdown(minimal_session)
        content = result.content.decode("utf-8")

        # Should still have basic structure
        assert "# Simple query" in content
        assert "## Metadata" in content
        assert "**Query:** Simple query" in content

    def test_markdown_size_human(self, sample_session: ResearchSession) -> None:
        """Test human-readable size formatting."""
        result = export_to_markdown(sample_session)
        assert "B" in result.size_human or "KB" in result.size_human


# =============================================================================
# JSON Export Tests
# =============================================================================


class TestJsonExport:
    """Tests for JSON export functionality."""

    def test_export_to_json_basic(self, sample_session: ResearchSession) -> None:
        """Test basic JSON export."""
        result = export_to_json(sample_session)

        assert result.format == ExportFormat.JSON
        assert result.filename.endswith(".json")
        assert result.mime_type == "application/json"

    def test_json_content_valid(self, sample_session: ResearchSession) -> None:
        """Test JSON content is valid and parseable."""
        result = export_to_json(sample_session)
        content = result.content.decode("utf-8")

        # Should be valid JSON
        data = json.loads(content)

        # Check required fields
        assert data["interaction_id"] == sample_session.interaction_id
        assert data["query"] == sample_session.query
        assert data["title"] == sample_session.title
        assert data["summary"] == sample_session.summary
        assert data["report_text"] == sample_session.report_text
        assert data["tags"] == sample_session.tags
        assert data["duration_seconds"] == sample_session.duration_seconds
        assert data["total_tokens"] == sample_session.total_tokens

        # Check timestamp fields
        assert "created_at" in data
        assert "expires_at" in data
        assert "export_timestamp" in data

    def test_json_minimal_session(self, minimal_session: ResearchSession) -> None:
        """Test JSON export with minimal session."""
        result = export_to_json(minimal_session)
        data = json.loads(result.content.decode("utf-8"))

        assert data["interaction_id"] == minimal_session.interaction_id
        assert data["query"] == minimal_session.query
        assert data["title"] is None
        assert data["summary"] is None


# =============================================================================
# DOCX Export Tests
# =============================================================================


class TestDocxExport:
    """Tests for DOCX export functionality."""

    def test_export_to_docx_basic(self, sample_session: ResearchSession) -> None:
        """Test basic DOCX export."""
        # Import check - skip if marko/python-docx not installed
        pytest.importorskip("marko")
        pytest.importorskip("docx")

        from gemini_research_mcp.export import export_to_docx

        result = export_to_docx(sample_session)

        assert result.format == ExportFormat.DOCX
        assert result.filename.endswith(".docx")
        assert result.mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert len(result.content) > 0

    def test_docx_is_valid_zip(self, sample_session: ResearchSession) -> None:
        """Test that DOCX output is a valid ZIP archive (OOXML format)."""
        pytest.importorskip("marko")
        pytest.importorskip("docx")

        from io import BytesIO
        from zipfile import ZipFile

        from gemini_research_mcp.export import export_to_docx

        result = export_to_docx(sample_session)

        # DOCX is a ZIP archive
        with ZipFile(BytesIO(result.content)) as zf:
            namelist = zf.namelist()
            # Should contain standard OOXML files
            assert "[Content_Types].xml" in namelist
            assert "word/document.xml" in namelist

    def test_docx_minimal_session(self, minimal_session: ResearchSession) -> None:
        """Test DOCX export with minimal session."""
        pytest.importorskip("marko")
        pytest.importorskip("docx")

        from gemini_research_mcp.export import export_to_docx

        result = export_to_docx(minimal_session)
        assert len(result.content) > 0

    def test_docx_import_error(self, sample_session: ResearchSession) -> None:
        """Test that missing marko/docx raises helpful error."""
        # This test is informational - we can't easily test ImportError
        # when dependencies are installed in dev environment
        pass

    def test_docx_has_professional_styling(self, sample_session: ResearchSession) -> None:
        """Test DOCX has professional page margins and styles configured."""
        pytest.importorskip("marko")
        pytest.importorskip("docx")

        from io import BytesIO

        from docx import Document
        from docx.shared import Cm

        result = export_to_docx(sample_session)
        doc = Document(BytesIO(result.content))

        # Check page margins are set to 1 inch (2.54 cm)
        for section in doc.sections:
            assert section.top_margin == Cm(2.54)
            assert section.bottom_margin == Cm(2.54)
            assert section.left_margin == Cm(2.54)
            assert section.right_margin == Cm(2.54)

    def test_docx_contains_cover_page(self, sample_session: ResearchSession) -> None:
        """Test DOCX includes cover page with title."""
        pytest.importorskip("marko")
        pytest.importorskip("docx")

        from io import BytesIO

        from docx import Document

        result = export_to_docx(sample_session)
        doc = Document(BytesIO(result.content))

        # First meaningful content should be the title
        text_content = " ".join(p.text for p in doc.paragraphs[:10])
        assert sample_session.title in text_content

    def test_docx_contains_toc(self, sample_session: ResearchSession) -> None:
        """Test DOCX includes Table of Contents."""
        pytest.importorskip("marko")
        pytest.importorskip("docx")

        from io import BytesIO

        from docx import Document

        result = export_to_docx(sample_session)
        doc = Document(BytesIO(result.content))

        text_content = " ".join(p.text for p in doc.paragraphs)
        assert "Table of Contents" in text_content

    def test_docx_contains_document_info(self, sample_session: ResearchSession) -> None:
        """Test DOCX includes Document Information section."""
        pytest.importorskip("marko")
        pytest.importorskip("docx")

        from io import BytesIO
        from zipfile import ZipFile

        result = export_to_docx(sample_session)

        # Check XML content for Document Information and session ID
        with ZipFile(BytesIO(result.content)) as zf:
            doc_xml = zf.read("word/document.xml").decode("utf-8")
            assert "Document Information" in doc_xml
            # Interaction ID should be in the metadata table
            assert sample_session.interaction_id in doc_xml

    def test_docx_toc_has_clickable_links(self, sample_session: ResearchSession) -> None:
        """Test TOC entries are internal hyperlinks."""
        pytest.importorskip("marko")
        pytest.importorskip("docx")

        from io import BytesIO
        from zipfile import ZipFile

        result = export_to_docx(sample_session)

        # Parse the document XML to check for hyperlinks with anchors
        with ZipFile(BytesIO(result.content)) as zf:
            doc_xml = zf.read("word/document.xml").decode("utf-8")
            # Internal hyperlinks use w:anchor attribute
            assert "w:anchor" in doc_xml

    def test_docx_headings_have_bookmarks(self, sample_session: ResearchSession) -> None:
        """Test headings have bookmarks for TOC navigation."""
        pytest.importorskip("marko")
        pytest.importorskip("docx")

        from io import BytesIO
        from zipfile import ZipFile

        result = export_to_docx(sample_session)

        # Parse the document XML to check for bookmarks
        with ZipFile(BytesIO(result.content)) as zf:
            doc_xml = zf.read("word/document.xml").decode("utf-8")
            # Bookmarks use w:bookmarkStart element
            assert "w:bookmarkStart" in doc_xml
            assert "w:bookmarkEnd" in doc_xml

    def test_docx_without_cover_page(self, sample_session: ResearchSession) -> None:
        """Test DOCX export without cover page option."""
        pytest.importorskip("marko")
        pytest.importorskip("docx")

        from io import BytesIO

        from docx import Document

        result = export_to_docx(sample_session, include_cover_page=False)
        doc = Document(BytesIO(result.content))

        # First paragraph should be TOC or Document Info, not spacers
        first_texts = [p.text.strip() for p in doc.paragraphs[:5] if p.text.strip()]
        # Should still have content
        assert len(first_texts) > 0

    def test_docx_without_toc(self, sample_session: ResearchSession) -> None:
        """Test DOCX export without TOC option."""
        pytest.importorskip("marko")
        pytest.importorskip("docx")

        from io import BytesIO

        from docx import Document

        result = export_to_docx(sample_session, include_toc=False)
        doc = Document(BytesIO(result.content))

        text_content = " ".join(p.text for p in doc.paragraphs)
        # Should not have TOC heading (but Document Information should exist)
        assert "Table of Contents" not in text_content
        assert "Document Information" in text_content

    def test_docx_sources_hyperlinks(self, sample_session: ResearchSession) -> None:
        """Test that sources section contains proper hyperlinks."""
        pytest.importorskip("marko")
        pytest.importorskip("docx")

        from io import BytesIO
        from zipfile import ZipFile

        result = export_to_docx(sample_session)

        # Parse the document relationships to check for external hyperlinks
        with ZipFile(BytesIO(result.content)) as zf:
            rels_xml = zf.read("word/_rels/document.xml.rels").decode("utf-8")
            # External hyperlinks should be registered
            assert "example.com" in rels_xml or "hyperlink" in rels_xml.lower()

    def test_docx_ordered_list_starts_from_one(self) -> None:
        """Test ordered lists always start from 1."""
        pytest.importorskip("marko")
        pytest.importorskip("docx")

        from io import BytesIO

        from docx import Document

        # Create session with numbered list that starts from arbitrary number
        session = ResearchSession(
            interaction_id="list-test-123",
            query="Test query",
            created_at=time.time(),
            report_text="""## Test Section

4. First item
5. Second item
6. Third item
""",
        )

        result = export_to_docx(session)
        doc = Document(BytesIO(result.content))

        # Find paragraphs that look like list items
        list_texts = [p.text for p in doc.paragraphs if p.text.strip().startswith(("1.", "2.", "3."))]
        assert len(list_texts) >= 3
        assert any("1." in t for t in list_texts)


# =============================================================================
# Bookmark and Internal Link Tests
# =============================================================================


class TestBookmarkGeneration:
    """Tests for bookmark ID generation."""

    def test_create_bookmark_id_basic(self) -> None:
        """Test basic bookmark ID generation."""
        from gemini_research_mcp.export import _create_bookmark_id

        bookmark = _create_bookmark_id("Introduction", 0)
        assert bookmark.startswith("_bm_0_")
        assert "Introduction" in bookmark

    def test_create_bookmark_id_with_special_chars(self) -> None:
        """Test bookmark ID handles special characters."""
        from gemini_research_mcp.export import _create_bookmark_id

        bookmark = _create_bookmark_id("1.2 The Rise of \"Vibe Coding\"", 5)
        assert bookmark.startswith("_bm_5_")
        # Should not contain special chars
        assert '"' not in bookmark
        assert "." not in bookmark
        assert " " not in bookmark

    def test_create_bookmark_id_unique_per_index(self) -> None:
        """Test same title with different indices creates unique IDs."""
        from gemini_research_mcp.export import _create_bookmark_id

        bm1 = _create_bookmark_id("Section", 0)
        bm2 = _create_bookmark_id("Section", 1)
        assert bm1 != bm2

    def test_create_bookmark_id_starts_with_valid_char(self) -> None:
        """Test bookmark ID starts with letter or underscore (valid for Word)."""
        from gemini_research_mcp.export import _create_bookmark_id

        # Title starting with number
        bookmark = _create_bookmark_id("123 Numbers First", 0)
        # Bookmark IDs starting with underscore are valid in Word
        assert bookmark[0].isalpha() or bookmark[0] == "_"
        # Should contain the index for uniqueness
        assert "_0_" in bookmark


# =============================================================================
# export_session Function Tests
# =============================================================================


class TestExportSession:
    """Tests for the unified export_session function."""

    def test_export_session_markdown_by_string(self, sample_session: ResearchSession) -> None:
        """Test export with format as string."""
        result = export_session(sample_session, "markdown")
        assert result.format == ExportFormat.MARKDOWN

    def test_export_session_md_shorthand(self, sample_session: ResearchSession) -> None:
        """Test export with 'md' shorthand."""
        result = export_session(sample_session, "md")
        assert result.format == ExportFormat.MARKDOWN

    def test_export_session_json(self, sample_session: ResearchSession) -> None:
        """Test JSON export via export_session."""
        result = export_session(sample_session, "json")
        assert result.format == ExportFormat.JSON

    def test_export_session_docx(self, sample_session: ResearchSession) -> None:
        """Test DOCX export via export_session."""
        pytest.importorskip("marko")
        pytest.importorskip("docx")
        result = export_session(sample_session, "docx")
        assert result.format == ExportFormat.DOCX

    def test_export_session_word_alias(self, sample_session: ResearchSession) -> None:
        """Test 'word' alias for DOCX."""
        pytest.importorskip("marko")
        pytest.importorskip("docx")
        result = export_session(sample_session, "word")
        assert result.format == ExportFormat.DOCX

    def test_export_session_enum(self, sample_session: ResearchSession) -> None:
        """Test export with ExportFormat enum."""
        result = export_session(sample_session, ExportFormat.JSON)
        assert result.format == ExportFormat.JSON

    def test_export_session_invalid_format(self, sample_session: ResearchSession) -> None:
        """Test that invalid format raises ValueError."""
        with pytest.raises(ValueError, match="Unsupported format"):
            export_session(sample_session, "pdf")

    def test_export_session_to_file(
        self, sample_session: ResearchSession, tmp_path: Path
    ) -> None:
        """Test export with file output."""
        output_path = tmp_path / "research.md"
        result = export_session(sample_session, "markdown", output_path=output_path)

        assert output_path.exists()
        assert output_path.read_bytes() == result.content


# =============================================================================
# Filename Generation Tests
# =============================================================================


class TestFilenameGeneration:
    """Tests for filename generation."""

    def test_filename_from_title(self, sample_session: ResearchSession) -> None:
        """Test filename uses title when available."""
        result = export_to_markdown(sample_session)
        assert "Quantum" in result.filename or "quantum" in result.filename.lower()

    def test_filename_from_query(self, minimal_session: ResearchSession) -> None:
        """Test filename uses query when no title."""
        result = export_to_markdown(minimal_session)
        assert "Simple" in result.filename or "simple" in result.filename.lower()

    def test_filename_has_date(self, sample_session: ResearchSession) -> None:
        """Test filename includes date."""
        result = export_to_markdown(sample_session)
        # Should contain YYYYMMDD pattern
        import re
        assert re.search(r"\d{8}", result.filename)

    def test_filename_safe_characters(self) -> None:
        """Test filename handles special characters safely."""
        session = ResearchSession(
            interaction_id="test-123",
            query="What's the best <approach> to AI/ML?",
            created_at=time.time(),
        )
        result = export_to_markdown(session)
        # Should not contain unsafe characters
        assert "<" not in result.filename
        assert ">" not in result.filename
        assert "?" not in result.filename
        assert "/" not in result.filename


# =============================================================================
# ExportResult Tests
# =============================================================================


class TestExportResult:
    """Tests for ExportResult dataclass."""

    def test_size_human_bytes(self) -> None:
        """Test size formatting for small files."""
        result = ExportResult(
            format=ExportFormat.MARKDOWN,
            filename="test.md",
            content=b"Hello",
            mime_type="text/markdown",
        )
        assert "B" in result.size_human
        assert "5" in result.size_human

    def test_size_human_kilobytes(self) -> None:
        """Test size formatting for KB files."""
        result = ExportResult(
            format=ExportFormat.MARKDOWN,
            filename="test.md",
            content=b"x" * 2048,
            mime_type="text/markdown",
        )
        assert "KB" in result.size_human


# =============================================================================
# Export Cache Tests
# =============================================================================


class TestExportCache:
    """Tests for ephemeral export cache used by MCP Resources."""

    def test_cache_export_returns_id(self, sample_session: ResearchSession) -> None:
        """Test caching an export returns a unique ID."""
        from gemini_research_mcp.server import _cache_export, _export_cache

        result = export_to_markdown(sample_session)
        export_id = _cache_export(result, sample_session.interaction_id)

        assert export_id is not None
        assert len(export_id) == 12  # UUID[:12]
        assert export_id in _export_cache

        # Cleanup
        del _export_cache[export_id]

    def test_get_cached_export_returns_entry(
        self, sample_session: ResearchSession
    ) -> None:
        """Test retrieving a cached export."""
        from gemini_research_mcp.server import (
            _cache_export,
            _export_cache,
            _get_cached_export,
        )

        result = export_to_markdown(sample_session)
        export_id = _cache_export(result, sample_session.interaction_id)

        entry = _get_cached_export(export_id)
        assert entry is not None
        assert entry.result == result
        assert entry.session_id == sample_session.interaction_id
        assert not entry.is_expired

        # Cleanup
        del _export_cache[export_id]

    def test_get_cached_export_not_found(self) -> None:
        """Test retrieving non-existent export returns None."""
        from gemini_research_mcp.server import _get_cached_export

        entry = _get_cached_export("nonexistent-id")
        assert entry is None

    def test_export_resource_returns_mcp_types(
        self, sample_session: ResearchSession
    ) -> None:
        """Test the resource function returns proper MCP content types."""
        from mcp.types import BlobResourceContents, TextResourceContents

        from gemini_research_mcp.server import (
            _cache_export,
            _export_cache,
            get_export_by_id,
        )

        # Test markdown returns TextResourceContents
        result = export_to_markdown(sample_session)
        export_id = _cache_export(result, sample_session.interaction_id)

        content = get_export_by_id(export_id)
        assert isinstance(content, TextResourceContents)
        assert content.mimeType == "text/markdown"
        assert content.text == result.content.decode("utf-8")

        # Cleanup
        del _export_cache[export_id]

    def test_export_resource_docx_returns_blob(
        self, sample_session: ResearchSession
    ) -> None:
        """Test DOCX export returns BlobResourceContents with base64."""
        import base64

        from mcp.types import BlobResourceContents

        from gemini_research_mcp.server import (
            _cache_export,
            _export_cache,
            get_export_by_id,
        )

        result = export_to_docx(sample_session)
        export_id = _cache_export(result, sample_session.interaction_id)

        content = get_export_by_id(export_id)
        assert isinstance(content, BlobResourceContents)
        assert content.mimeType == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        # Verify base64 decodes to original content
        decoded = base64.b64decode(content.blob)
        assert decoded == result.content

        # Cleanup
        del _export_cache[export_id]

    def test_export_resource_raises_on_invalid_id(self) -> None:
        """Test resource raises ValueError for invalid export ID."""
        from gemini_research_mcp.server import get_export_by_id

        with pytest.raises(ValueError, match="Export not found"):
            get_export_by_id("invalid-id-123")

    def test_list_exports_returns_json(
        self, sample_session: ResearchSession
    ) -> None:
        """Test list_exports returns JSON with export metadata."""
        from gemini_research_mcp.server import (
            _cache_export,
            _export_cache,
            list_exports,
        )

        result = export_to_markdown(sample_session)
        export_id = _cache_export(result, sample_session.interaction_id)

        exports_json = list_exports()
        data = json.loads(exports_json)

        assert "exports" in data
        assert "count" in data
        assert data["count"] >= 1

        # Find our export
        our_export = next(
            (e for e in data["exports"] if e["export_id"] == export_id), None
        )
        assert our_export is not None
        assert our_export["filename"] == result.filename
        assert our_export["format"] == "markdown"
        assert "uri" in our_export
        assert our_export["uri"] == f"research://exports/{export_id}"

        # Cleanup
        del _export_cache[export_id]

    def test_cache_cleans_up_expired_entries(
        self, sample_session: ResearchSession
    ) -> None:
        """Test that caching new exports cleans up expired entries."""
        from datetime import UTC, datetime, timedelta

        from gemini_research_mcp.server import (
            EXPORT_TTL_SECONDS,
            ExportCacheEntry,
            _cache_export,
            _export_cache,
        )

        # Create an expired entry manually
        old_result = export_to_markdown(sample_session)
        expired_entry = ExportCacheEntry(
            result=old_result,
            session_id="old-session",
            created_at=datetime.now(UTC) - timedelta(seconds=EXPORT_TTL_SECONDS + 100),
        )
        _export_cache["expired-id"] = expired_entry
        assert expired_entry.is_expired

        # Cache a new export - should trigger cleanup
        new_result = export_to_markdown(sample_session)
        new_id = _cache_export(new_result, sample_session.interaction_id)

        # Expired entry should be gone
        assert "expired-id" not in _export_cache
        assert new_id in _export_cache

        # Cleanup
        del _export_cache[new_id]
