"""
Word Parse Module
Extract text and content from Word documents (docx)
"""
import asyncio
import logging
import os
from typing import Any, Dict, List

from ...registry import register_module
from ...schema import compose, presets


logger = logging.getLogger(__name__)


@register_module(
    module_id='word.parse',
    version='1.0.0',
    category='document',
    subcategory='word',
    tags=['word', 'docx', 'parse', 'extract', 'document', 'path_restricted'],
    label='Parse Word Document',
    label_key='modules.word.parse.label',
    description='Extract text and content from Word documents (.docx)',
    description_key='modules.word.parse.description',
    icon='FileText',
    color='#2B579A',

    input_types=['file'],
    output_types=['text', 'object'],
    can_connect_to=['data.*', 'string.*', 'ai.*'],
    can_receive_from=['file.*', 'data.*', 'api.*', 'flow.*', 'start'],

    timeout_ms=60000,
    retryable=True,
    max_retries=2,
    concurrent_safe=True,

    requires_credentials=False,
    handles_sensitive_data=False,
    required_permissions=[],

    params_schema=compose(
        presets.WORD_FILE_PATH(),
        presets.DOC_EXTRACT_TABLES(default=True),
        presets.DOC_EXTRACT_IMAGES(),
        presets.DOC_IMAGES_OUTPUT_DIR(),
        presets.DOC_PRESERVE_FORMATTING(default=False),
    ),
    output_schema={
        'text': {
            'type': 'string',
            'description': 'Full text content of the document'
        ,
                'description_key': 'modules.word.parse.output.text.description'},
        'paragraphs': {
            'type': 'array',
            'description': 'List of paragraphs'
        ,
                'description_key': 'modules.word.parse.output.paragraphs.description'},
        'tables': {
            'type': 'array',
            'description': 'Extracted tables as arrays'
        ,
                'description_key': 'modules.word.parse.output.tables.description'},
        'images': {
            'type': 'array',
            'description': 'Paths to extracted images'
        ,
                'description_key': 'modules.word.parse.output.images.description'},
        'metadata': {
            'type': 'object',
            'description': 'Document metadata'
        ,
                'description_key': 'modules.word.parse.output.metadata.description'}
    },
    examples=[
        {
            'title': 'Extract text from Word',
            'title_key': 'modules.word.parse.examples.basic.title',
            'params': {
                'file_path': '/path/to/document.docx'
            }
        },
        {
            'title': 'Extract with tables and images',
            'title_key': 'modules.word.parse.examples.full.title',
            'params': {
                'file_path': '/path/to/document.docx',
                'extract_tables': True,
                'extract_images': True,
                'images_output_dir': '/path/to/images/'
            }
        }
    ],
    author='Flyto2 Team',
    license='MIT'
)
async def word_parse(context: Dict[str, Any]) -> Dict[str, Any]:
    """Parse Word document and extract content"""
    try:
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError
    except ImportError:
        raise ImportError("python-docx is required for word.parse. Install with: pip install python-docx")

    params = context['params']
    file_path = params['file_path']
    extract_tables = params.get('extract_tables', True)
    extract_images = params.get('extract_images', False)
    images_output_dir = params.get('images_output_dir')
    preserve_formatting = params.get('preserve_formatting', False)

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    def _parse():
        try:
            doc = Document(file_path)
        except PackageNotFoundError:
            raise ValueError(f"Invalid or corrupted Word document: {file_path}")

        paragraphs = []
        full_text_parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                if preserve_formatting:
                    style_name = para.style.name if para.style else 'Normal'
                    paragraphs.append({
                        'text': text,
                        'style': style_name
                    })
                else:
                    paragraphs.append(text)
                full_text_parts.append(text)

        tables = []
        if extract_tables:
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                if table_data:
                    tables.append(table_data)

        images = []
        if extract_images and images_output_dir:
            os.makedirs(images_output_dir, exist_ok=True)
            image_count = 0

            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    image_count += 1
                    image_data = rel.target_part.blob
                    content_type = rel.target_part.content_type

                    ext = 'png'
                    if 'jpeg' in content_type or 'jpg' in content_type:
                        ext = 'jpg'
                    elif 'gif' in content_type:
                        ext = 'gif'

                    image_path = os.path.join(
                        images_output_dir,
                        f"image_{image_count}.{ext}"
                    )
                    with open(image_path, 'wb') as f:
                        f.write(image_data)
                    images.append(image_path)

        metadata = {}
        try:
            core_props = doc.core_properties
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'last_modified_by': core_props.last_modified_by or ''
            }
        except Exception:
            pass

        return {
            'text': '\n\n'.join(full_text_parts),
            'paragraphs': paragraphs,
            'tables': tables,
            'images': images,
            'metadata': metadata
        }

    result = await asyncio.to_thread(_parse)

    logger.info(
        f"Parsed Word document: {len(result['paragraphs'])} paragraphs, "
        f"{len(result['tables'])} tables, {len(result['images'])} images"
    )

    return {
        'ok': True,
        **result
    }
