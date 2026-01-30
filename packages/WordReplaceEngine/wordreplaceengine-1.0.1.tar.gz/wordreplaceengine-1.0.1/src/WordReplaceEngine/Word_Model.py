#coding=UTF-8
"""
    WordReplaceEngine
    Copyright (C) 2026 Yun-De Huang <startime_electronics@orztrickster.com>

    This program is free software: you can redistribute it and/or modify
    it under the terms of the GNU Affero General Public License as
    published by the Free Software Foundation, either version 3 of the
    License, or (at your option) any later version.

    This program is distributed in the hope that it will be useful,
    but WITHOUT ANY WARRANTY; without even the implied warranty of
    MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
    GNU Affero General Public License for more details.

    You should have received a copy of the GNU Affero General Public License
    along with this program.  If not, see <https://www.gnu.org/licenses/>.
"""




import os
from docx import Document as Docu
from docx.shared import Pt,RGBColor,Cm,Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from PIL import Image
try:
    from xml.etree.cElementTree import XML
except ImportError:
    from xml.etree.ElementTree import XML
from .PDF_Model import PDF_Model
from .PDF_Model import PDF_Model_Syncfusion


class Document():
    def __init__(self, syncfusion_key = None):
        self.force_chinese_font_symbols = {"︵", "︶", "（", "）"}
        self.force_english_font_symbols = {"0", "1", "2", "3", "4", "5", "6", "7", "8", "9"}
        self.version_path = {}
        self.replace_value = {}
        self.pdf_model = PDF_Model()
        self.syncfusion_key = syncfusion_key
        if self.syncfusion_key is not None:
            try:
                self.pdf_model_syncfusion = PDF_Model_Syncfusion(syncfusion_key = self.syncfusion_key)
            except Exception as e:
                pass
                # print("")
                # print(rf"載入PDF_Model_Syncfusion()錯誤 --> {e}")
                # print("")

    def set_word_version_path(self,name,path):
        self.version_path[str(name)] = path
    def run_version(self,version,save_path = "C:/WordReplaceEngine/"):
        try:
            os.makedirs(save_path, exist_ok=True)
        except:
            pass
        self.save_path = save_path
        if str(version) in self.version_path:
            self.version = version
            self.doc = Docu(self.version_path[str(version)])
        else:
            print("word的檔案格式名稱並沒有定義")

    def get_accepted_text(self, p):
        WORD_NAMESPACE = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
        xml = p._p.xml
        if "w:del" in xml or "w:ins" in xml or "w:br" in xml:
            tree = XML(xml)
            result = []

            for r in tree.iterfind('.//w:r', namespaces=ns):
                for node in r:
                    if node.tag == f"{WORD_NAMESPACE}br" and node.get(f"{WORD_NAMESPACE}type") == "page":
                        result.append("$$$page$$$")
                    elif node.tag == f"{WORD_NAMESPACE}t" and node.text:
                        result.append(node.text)
            return ''.join(result)
        else:
            return p.text
    def delete(self):
        self.replace_value = {}
      
    def set(self, old_text, new_text, **kwargs):

        self.replace_value[str(old_text)] = {}
        self.replace_value[str(old_text)]["替換值"] = str(new_text)
        self.replace_value[str(old_text)]["英文字型"] = kwargs.setdefault("new_font_name", 'Calibri')
        self.replace_value[str(old_text)]["中文字型"] = kwargs.setdefault("new_font_name_Chinese", '標楷體')
        self.replace_value[str(old_text)]["字體大小"] = kwargs.setdefault("new_font_size", Pt(8))
        self.replace_value[str(old_text)]["對其位置"] = kwargs.setdefault("new_alignment", WD_PARAGRAPH_ALIGNMENT.CENTER)
        self.replace_value[str(old_text)]["粗體"] = kwargs.setdefault("bold", False)
        self.replace_value[str(old_text)]["left_indent"] = kwargs.setdefault("left_indent", None)
        self.replace_value[str(old_text)]["line_spacing"] = kwargs.setdefault("line_spacing", None)
        self.replace_value[str(old_text)]["font_color"] = kwargs.setdefault("font_color", None)
        
        self.replace_value[str(old_text)]["__remark__"] = {}
        self.replace_value[str(old_text)]["__remark__"]["字體大小"] = kwargs.setdefault("remark_new_font_size", Pt(8))
        

        self.replace_value[str(old_text)]["image_path"] = kwargs.setdefault("image_path", None)
        self.replace_value[str(old_text)]["image_height"] = kwargs.setdefault("image_height", None)
        self.replace_value[str(old_text)]["image_width"] = kwargs.setdefault("image_width", None)
        self.replace_value[str(old_text)]["固定圖片比例"] = kwargs.setdefault("image_fixed_ratio", True)

            








    def execute_set(self):
        # 遍歷文檔中的所有段落
        for para in self.doc.paragraphs:
            XXX = self.get_accepted_text(para)
                
            for kk in self.replace_value:
                if kk in XXX:
                        


                    old_text = str(self.get_accepted_text(para))
                    try:
                        remaining_string_A = old_text.split(kk)[0]
                        remaining_string_B = old_text.split(kk)[1]
                    except:
                        remaining_string_A = ""
                        remaining_string_B = ""
                    new_text = remaining_string_A + old_text.replace(old_text, self.replace_value[kk]["替換值"]) + remaining_string_B




                    new_text_list = new_text.split("$$$page$$$")


                    #new_text = str(para.text)


                    para.clear()  # 清空段落內容


                    for i in range(len(new_text_list)):
                        new_text = new_text_list[i]

                        run = para.add_run(new_text)  # 添加新的 Run 對象




                        font = run.font
                        #font.name = self.replace_value[kk]["英文字型"]
                        font.size = self.replace_value[kk]["字體大小"]
                        font.bold = self.replace_value[kk]["粗體"]
                        para.alignment = self.replace_value[kk]["對其位置"]

                        #run._element.rPr.rFonts.set(qn('w:eastAsia'), self.replace_value[kk]["中文字型"])  # 用於設置中文時的字型。



                        # 🔹只讓「︵」與「︶」兩個符號使用中文字型
                        text_this_run = run.text or ""

                        # 先檢查是否屬於強制中或英的符號清單
                        if any(sym in text_this_run for sym in getattr(self, "force_chinese_font_symbols", set())):
                            font.name = self.replace_value[kk]["中文字型"]
                        elif any(sym in text_this_run for sym in getattr(self, "force_english_font_symbols", set())):
                            font.name = self.replace_value[kk]["英文字型"]
                        else:
                            # 依一般中英文預設處理
                            font.name = self.replace_value[kk]["英文字型"]

                        # 保持 eastAsia 一律指到中文字型（避免 fallback）
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.replace_value[kk]["中文字型"])



                        # 設置左縮進
                        if self.replace_value[kk]["left_indent"] is not None:
                            para.paragraph_format.left_indent = Pt(self.replace_value[kk]["left_indent"])

                        # 設置固定行高
                        if self.replace_value[kk]["line_spacing"] is not None:
                            para.paragraph_format.line_spacing = Pt(self.replace_value[kk]["line_spacing"])

                        # 設置字體顏色
                        if self.replace_value[kk]["font_color"] is not None:
                            font.color.rgb = RGBColor(*self.replace_value[kk]["font_color"])



                        if i != len(new_text_list)-1:
                            run.add_break(WD_BREAK.PAGE)
                            #self.doc.add_page_break()  


            

        # 遍歷文檔中的所有表格
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    old_text = str(cell.text)

                    for kk in self.replace_value:
                        if kk in old_text:
                            try:
                                remaining_string_A = cell.text.split(kk)[0]
                                remaining_string_B = cell.text.split(kk)[1]
                            except:
                                remaining_string_A = ""
                                remaining_string_B = ""
                            cell.text = remaining_string_A + cell.text.replace(cell.text, self.replace_value[kk]["替換值"]) + remaining_string_B

                        
                            #設置新的內容格式
                            for paragraph in cell.paragraphs:
                            
                                for run in paragraph.runs:
                                    #__※_※__
                                    if("__remark__" in run.text):

                                        run_list = run.text.split("__remark__")

                                        if len(run_list) >= 2:
                                            # 將第二個字分成新的 run，以單獨設定其格式
                                            run.text = run_list[0]  # 保留第一個字
                                            new_run = paragraph.add_run(run_list[1])  # 創建新 run 給第二個字
                                            new_run.font.size = self.replace_value[kk]["__remark__"]["字體大小"]  # 設定第二個字的字體大小為 24 點
                                            new_font = new_run.font
                                            new_font.name = self.replace_value[kk]["英文字型"]
                                            new_run._element.rPr.rFonts.set(qn('w:eastAsia'), self.replace_value[kk]["中文字型"])

                                            # 設置字體顏色
                                            if self.replace_value[kk]["font_color"] is not None:
                                                new_font.color.rgb = RGBColor(*self.replace_value[kk]["font_color"])



                                    font = run.font
                                    #font.name = self.replace_value[kk]["英文字型"]
                                    font.size = self.replace_value[kk]["字體大小"]
                                    font.bold = self.replace_value[kk]["粗體"]
                                    #run._element.rPr.rFonts.set(qn('w:eastAsia'), self.replace_value[kk]["中文字型"])


                                    # 只針對符號強制中文，其餘保持原設定
                                    text_this_run = run.text or ""

                                    # 先檢查是否屬於強制中或英的符號清單
                                    if any(sym in text_this_run for sym in getattr(self, "force_chinese_font_symbols", set())):
                                        font.name = self.replace_value[kk]["中文字型"]
                                    elif any(sym in text_this_run for sym in getattr(self, "force_english_font_symbols", set())):
                                        font.name = self.replace_value[kk]["英文字型"]
                                    else:
                                        # 依一般中英文預設處理
                                        font.name = self.replace_value[kk]["英文字型"]

                                    # 保持 eastAsia 一律指到中文字型（避免 fallback）
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), self.replace_value[kk]["中文字型"])













                                    # 設定字體顏色
                                    if self.replace_value[kk]["font_color"] is not None:
                                        font.color.rgb = RGBColor(*self.replace_value[kk]["font_color"])


                                if self.replace_value[kk]["image_path"] != None:
                                    try:
                                        image = Image.open(self.replace_value[kk]["image_path"])
                                        image_size = image.size
                                        w = float(image_size[0])
                                        h = float(image_size[1])
                                        w_max = self.replace_value[kk]["image_width"]
                                        h_max = self.replace_value[kk]["image_height"]
                                        ratio = w / h

                                        if self.replace_value[kk]["固定圖片比例"]:
                                            if w < w_max and h < h_max:
                                                while True:
                                                    if w < w_max:
                                                        w = w_max
                                                        h = w / ratio
                                                    if h < h_max:
                                                        h = h_max
                                                        w = h * ratio
                                                    if w >= w_max and h >= h_max:
                                                        break

                                            if w > w_max or h > h_max:
                                                while True:
                                                    if w > w_max:
                                                        w = w_max
                                                        h = w / ratio
                                                    if h > h_max:
                                                        h = h_max
                                                        w = h * ratio
                                                    if w <= w_max and h <= h_max:
                                                        break
                                        paragraph.add_run().add_picture(self.replace_value[kk]["image_path"], width = Inches(w), height = Inches(h))
                                    except Exception as e:
                                        print("execute_set設定照片錯誤 --> " + str(e))



                                    

                            for paragraph in cell.paragraphs:
                                paragraph.alignment = self.replace_value[kk]["對其位置"]
                                # 設置左縮進
                                if self.replace_value[kk]["left_indent"] is not None:
                                    paragraph.paragraph_format.left_indent = Pt(self.replace_value[kk]["left_indent"])

                                # 設置固定行高
                                if self.replace_value[kk]["line_spacing"] is not None:
                                    paragraph.paragraph_format.line_spacing = Pt(self.replace_value[kk]["line_spacing"])



    def show_execute_set(self,specify_range = False):
        self.check_dict = {}

        data = {}
        data["許可元件開始"] = {}
        data["全部許可元件停止"] = False
        for old_text in self.replace_value:
            if specify_range:
                data["許可元件開始"][old_text] = False
            else:
                data["許可元件開始"][old_text] = True


        for para in self.doc.paragraphs:
            para.text = self.get_accepted_text(para)
            
            
            for old_text in self.replace_value:
                new_text = self.replace_value[old_text]["替換值"]

                if specify_range:
                    if new_text in para.text:
                        data["許可元件開始"][old_text] = True

                    if("【符號說明】" in para.text):
                        data["全部許可元件停止"] = True


                if data["許可元件開始"][old_text] and data["全部許可元件停止"] == False:
                    
                    if old_text in para.text:

                        save = False
                        if old_text in new_text:
                            text = para.text
                            old_text_N = 0         #目標old_text用詞在text的位置
                        
                            while True:
                                NN = text[old_text_N:len(text)].find(old_text)
                                old_text_N += NN
                                if(NN != -1):
                                    if(str(text[old_text_N - new_text.find(old_text):old_text_N + len(new_text)]) != new_text):
                                        save = True
                                        for old_text_other in self.replace_value:
                                            if ((old_text in old_text_other) and old_text != old_text_other):
                                                if(str(text[old_text_N - new_text.find(old_text) - old_text_other.find(old_text):old_text_N - new_text.find(old_text) - old_text_other.find(old_text) + len(self.replace_value[old_text_other]["替換值"])]) == self.replace_value[old_text_other]["替換值"]):
                                                    save = False
                                                    break
                                    old_text_N += len(str(old_text))
                                else:
                                    break



                        else:
                        
                            save = True
                        if save:
                            if str(old_text) not in self.check_dict:
                                self.check_dict[str(old_text)] = {}
                                self.check_dict[str(old_text)]["text"] = []
                            self.check_dict[str(old_text)]["text"].append(str(para.text)) 

    def show_all_text(self):   #顯示所有文字
        aa = []
        for para in self.doc.paragraphs:
            para.text = self.get_accepted_text(para)
            aa.append(para.text)
        return aa
    def get_numId_text(self):
        word_list = []
        for para in self.doc.paragraphs:
            p = para._p  # 取得段落的 XML
            numPr = p.pPr.numPr if p.pPr is not None else None
            if numPr is not None:
                ilvl = numPr.ilvl.val  # 編號層級
                numId = numPr.numId.val  # 編號 ID
                word_list.append((numId, ilvl, self.get_accepted_text(para)))
            else:
                word_list.append((None, None, self.get_accepted_text(para)))
        return word_list
    def automatic_get_numId_text(self):
        word_list = []
        aa = {}
        for numId, ilvl, text in self.get_numId_text():
            if numId is not None:
                if numId not in aa:
                    aa[numId] = 1
                else:
                    aa[numId] += 1
                word_list.append((aa[numId], text))
            else:
                word_list.append((None, text))
        return word_list


    def output(self,name,pdf = False,docx = True):

            
        self.name = name

        

        name_pdf = self.save_path + self.name
        name_docx = self.save_path + self.name



        self.doc.save(name_docx+'.docx')
        # 指定輸入的 Word 文件和輸出的 PDF 文件
        input_docx = name_docx + '.docx'
        output_pdf = name_pdf + '.pdf'

        if pdf:
            try:
                if self.syncfusion_key is not None:
                    self.pdf_model_syncfusion.convert(input_docx, output_pdf)
                else:
                    raise
            except Exception as e:
                # print(rf"pdf_model_syncfusion錯誤 --> {e}")
                self.pdf_model.convert(input_docx, output_pdf, keep_active = True)
        if(docx != True):
            os.remove(input_docx) 




        #print("[OK！] "+str(output_pdf))
      
