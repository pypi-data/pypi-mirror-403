#!/usr/bin/env python3
"""Set of utilities to use wit hpython-docx."""
import os
import tempfile

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import ns
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
from docx.shared import Cm


def create_element(name):
    """Create a new XML element."""
    return OxmlElement(name)


def create_attribute(element, name, value):
    """Create an attribute."""
    element.set(ns.qn(name), value)


def add_page_number(paragraph):
    """Add page number."""
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    page_run = paragraph.add_run()
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    t1.text = 'Page '
    page_run._r.append(t1)

    page_num_run = paragraph.add_run()

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)

    of_run = paragraph.add_run()
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    t2.text = ' of '
    of_run._r.append(t2)

    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'begin')

    instrText2 = create_element('w:instrText')
    create_attribute(instrText2, 'xml:space', 'preserve')
    instrText2.text = "NUMPAGES"

    fldChar4 = create_element('w:fldChar')
    create_attribute(fldChar4, 'w:fldCharType', 'end')

    num_pages_run = paragraph.add_run()
    num_pages_run._r.append(fldChar3)
    num_pages_run._r.append(instrText2)
    num_pages_run._r.append(fldChar4)


def paragraph_align_center():
    """Align center."""
    return WD_ALIGN_PARAGRAPH.CENTER


class Document(object):
    """Create a document."""

    def __init__(self):
        """Initialize."""
        self.doc = docx.Document()
        self.doc_fign = 1
        self.doc_tbln = 1

    def __getattr__(self, __name):
        """Call docx.document.Document stuff."""
        try:
            return getattr(self.doc, __name)

        except Exception:
            return object.__getattribute__(self, __name)

    def add_page_numbers(self):
        """Add page numbers."""
        add_page_number(self.doc.sections[0].footer.paragraphs[0])

    def add_picture(self, fig_list, center=True, size=10, caption=None):
        """Add a picture in the document.

        Args:
        ----
            doc: the document
            fig : The matplotlib figure or list of figures.
            center (bool, optional): If picture will be centerd. Defaults to True.
            size (int, optional): Size of picture in cm. Defaults to 10.
            caption (str, optional): The text of the caption. Defaults to None.

        Returns
        -------
            Index of Figure if caption is True, otherwise -1.

        """
    
        if not isinstance(fig_list, (list, tuple)):
            fig_list = [fig_list]

        P = self.add_paragraph()
        if center:
            P.alignment = WD_ALIGN_PARAGRAPH.CENTER

        R = P.add_run()
        for fig in fig_list:
            png_file = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
            fig.savefig(png_file, dpi=300)
            png_file.close()

            R.add_text("")
            R.add_picture(png_file.name, width=Cm(size))

        if len(fig_list)>0:
            R.add_text(" ")

        rc = -1
        if caption:
            P = self.add_paragraph()
            R = P.add_run("Figure ")

            p = R.element  # this is the actual lxml element for a paragraph
            fld_xml = r"""<w:fldSimple %s w:instr="SEQ Figure \* ARABIC ">
                            <w:r><w:rPr>
                            <w:noProof/>
                            </w:rPr>
                            <w:t>%d</w:t>
                        </w:r></w:fldSimple>""" % (ns.nsdecls('w'), self.doc_fign)
            rc = self.doc_fign
            self.doc_fign += 1

            # fld_xml = '<w:fldSimple %s w:instr=" SEQ Figure \* ARABIC "/>' % nsdecls('w')
            fldSimple = parse_xml(fld_xml)
            p.addnext(fldSimple)
            P.add_run(". {}".format(caption))
            P.style = self.styles['Caption']
            if center:
                P.alignment = paragraph_align_center()

        os.remove(png_file.name)
        return rc

    def insert_table(self, rows=1, cols=1, caption=None, center=False):
        """Adds a table to the document.

        Args:
        ----
            nrow: The number of rows in the table
            ncol: The number of columns in the table
            caption: The caption if not NOne. Defaults to None.
            center: wheter to center the table and caption or not.

        Returns
        -------
            The table object.

        """
        if caption:
            P = self.add_paragraph()
            R = P.add_run("Table ")

            p = R.element  # this is the actual lxml element for a paragraph
            fld_xml = r"""<w:fldSimple %s w:instr="SEQ Table \* ARABIC ">
                            <w:r><w:rPr>
                            <w:noProof/>
                            </w:rPr>
                            <w:t>%d</w:t>
                        </w:r></w:fldSimple>""" % (ns.nsdecls('w'), self.doc_tbln)
            self.doc_tbln += 1

            # fld_xml = '<w:fldSimple %s w:instr=" SEQ Figure \* ARABIC "/>' % nsdecls('w')
            fldSimple = parse_xml(fld_xml)
            p.addnext(fldSimple)
            P.add_run(". {}".format(caption))
            P.style = self.styles['Caption']
            if center:
                P.alignment = paragraph_align_center()

        table = self.add_table(rows=rows, cols=cols)
        return table
