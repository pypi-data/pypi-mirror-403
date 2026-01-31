from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from m3cli.docs.commons import create_logger

_LOG = create_logger()


def add_page_breaks(doc):
    for para in doc.paragraphs:
        if '{{ pagebreak }}' in para.text:
            para.clear()  # Clear the paragraph text
            run = para.add_run()  # Add a page break
            run.add_break(WD_BREAK.PAGE)


def set_normal_text_style(doc):
    try:
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)
        style.font.color.rgb = RGBColor(0, 0, 0)
    except Exception as e:
        _LOG.error(f"Failed to set normal text style: {e}")


def create_border_element():
    border = OxmlElement('w:pBdr')
    for border_name in ('top', 'left', 'bottom', 'right'):
        border_elem = OxmlElement(f'w:{border_name}')
        border_elem.set(qn('w:val'), 'single')
        border_elem.set(qn('w:sz'), '4')
        border_elem.set(qn('w:space'), '0')
        border_elem.set(qn('w:color'), '000000')
        border.append(border_elem)
    return border


def create_code_block_style(doc):
    try:
        styles = doc.styles
        if 'Code Block' not in styles:
            code_style = styles.add_style('Code Block', 1)
            code_style.font.name = 'Consolas'
            code_style.font.size = Pt(10)
            code_style.font.color.rgb = RGBColor(0, 0, 0)
            code_style.paragraph_format.left_indent = Pt(10)
            code_style.paragraph_format.right_indent = Pt(10)
            code_style.paragraph_format.space_before = Pt(5)
            code_style.paragraph_format.space_after = Pt(5)
            p = code_style.paragraph_format
            p_element = p._element
            p_element.get_or_add_pPr().append(create_border_element())
    except Exception as e:
        _LOG.error(f"Failed to create code block style: {e}")


def set_table_styles(doc):
    for table in doc.tables:
        for row in table.rows:  # Set borders for each cell
            for cell in row.cells:
                set_cell_borders(cell)
        if table.rows:  # Make the header row bold
            header_row = table.rows[0]
            for cell in header_row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True


def set_cell_borders(cell):
    # Define the properties for each border side
    properties = {
        "top": {"sz": 12, "val": "single", "color": "000000", "space": "0"},
        "left": {"sz": 12, "val": "single", "color": "000000", "space": "0"},
        "bottom": {"sz": 12, "val": "single", "color": "000000", "space": "0"},
        "right": {"sz": 12, "val": "single", "color": "000000", "space": "0"},
    }
    # Create a new XML element for the cell borders
    tc_pr = cell._element.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    # Apply the properties to each side of the border
    for side, attrs in properties.items():
        tag = f'w:{side}'
        element = OxmlElement(tag)
        for key, value in attrs.items():
            element.set(qn(f'w:{key}'), str(value))
        tc_borders.append(element)
    tc_pr.append(tc_borders)


def apply_custom_styles(docx_file):
    doc = Document(docx_file)

    add_page_breaks(doc)
    set_normal_text_style(doc)
    create_code_block_style(doc)
    set_table_styles(doc)

    sections = doc.sections
    for section in sections:
        margin = 0.5
        section.top_margin = Inches(margin)
        section.bottom_margin = Inches(margin)
        section.left_margin = Inches(margin)
        section.right_margin = Inches(margin)

    for para in doc.paragraphs:
        style = para.style

        if style.name == 'Normal':
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif style.name == 'Heading 1':
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(20)
                run.font.bold = True
                run.font.color.rgb = RGBColor(70, 69, 71)

        elif style.name == 'Heading 2':
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = RGBColor(26, 156, 176)

        elif style.name == 'Heading 3':
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(44, 196, 201)

        elif style.name in ('Heading 4', 'Heading 5', 'Heading 6'):
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

        # Handle code blocks - apply both style AND run-level formatting
        elif 'code' in style.name.lower() or 'source' in style.name.lower():
            para.style = doc.styles['Code Block']
            for run in para.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(10)  # This overrides Pandoc's 11pt
                run.font.color.rgb = RGBColor(0, 0, 0)

    doc.save(docx_file)
