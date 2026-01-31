from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.text.paragraph import Paragraph
from m3cli.docs.commons import create_logger

_LOG = create_logger()


def add_page_to_docx(
        docx_path: str,
        title_text: str,
        secondary_title_text: str,
        date_text: str,
        version_text: str,
) -> None:
    try:
        doc = Document(docx_path)
        # Check if the first block in the document is a table
        first_element = doc.element.body[0]
        if first_element.tag == qn('w:tbl'):
            # Create a new paragraph element
            new_paragraph = OxmlElement('w:p')
            # Insert the new paragraph before the first table
            doc.element.body.insert(0, new_paragraph)

        first_para: Paragraph = doc.paragraphs[0]

        for _ in range(12):
            first_para.insert_paragraph_before()

        # Add title in the center of the page
        title = first_para.insert_paragraph_before()
        run = title.add_run(title_text)
        run.font.name = 'Arial'
        run.font.size = Pt(28)
        run.font.bold = True
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add new lines
        for _ in range(14):
            first_para.insert_paragraph_before()

        # Add secondary title at the bottom of the page
        secondary_title = first_para.insert_paragraph_before()
        run = secondary_title.add_run(secondary_title_text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.bold = True
        secondary_title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Add date
        date = first_para.insert_paragraph_before()
        run = date.add_run(date_text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        date.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Add new line
        first_para.insert_paragraph_before()

        # Add version
        version = first_para.insert_paragraph_before()
        run = version.add_run(version_text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        version.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Add a page break to create a new page
        run = version.add_run()
        run.add_break(WD_BREAK.PAGE)

        # Save the docx
        doc.save(docx_path)
        _LOG.info(
            f"Title page successfully added and document saved to: {docx_path}"
        )
    except Exception as e:
        _LOG.error(f"Failed to add title page to {docx_path}: {e}")
        raise
