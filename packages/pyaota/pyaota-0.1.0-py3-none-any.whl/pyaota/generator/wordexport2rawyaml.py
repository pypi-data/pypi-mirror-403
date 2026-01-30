#!/usr/bin/env python3
"""
Convert a ZyBooks-exported zip (containing three Word documents) into a
YAML multiple-choice question bank.

Expected zip contents (names may vary but pattern is similar):
    * <something>.docx                    # questions
    * <something>_with_answers.docx       # full test with answers inline (ignored)
    * <something>_answer_key.docx         # answer key

Usage:
    python zybooks_zip_to_yaml.py export.zip output.yaml

Requirements:
    pip install python-docx pyyaml
"""

import sys
import re
import zipfile
import tempfile
from pathlib import Path

from docx import Document
import yaml


# ---------- Utility: detect files in extracted zip ----------

def find_docs_in_extracted_dir(extract_dir: Path):
    """Return (questions_doc_path, answer_key_doc_path) from an extracted ZyBooks zip."""
    docx_files = list(extract_dir.glob("*.docx"))
    if not docx_files:
        raise FileNotFoundError("No .docx files found in the extracted zip directory.")

    questions_doc = None
    answer_key_doc = None

    # Heuristics based on ZyBooks naming conventions
    for p in docx_files:
        name_lower = p.name.lower()
        if "answer_key" in name_lower:
            answer_key_doc = p
        elif "with_answers" in name_lower:
            # We typically don't need this; ignore
            continue
        else:
            # Likely the base questions file (e.g., Test_Test.docx)
            questions_doc = p

    if questions_doc is None:
        raise FileNotFoundError("Could not find the questions .docx (without 'answer_key' or 'with_answers' in the name).")

    if answer_key_doc is None:
        raise FileNotFoundError("Could not find the answer key .docx (name containing 'answer_key').")

    return questions_doc, answer_key_doc


# ---------- Parsing answer key ----------

def parse_answer_key(answer_doc_path: Path):
    """
    Parse an answer-key .docx where each answer is on a line like:
        '1) b'
        '2) a'
    Returns: dict[int -> 'a'..'d']
    """
    doc = Document(str(answer_doc_path))
    answers = {}

    for p in doc.paragraphs:
        text = p.text.strip()
        m = re.match(r'^(\d+)\)\s*([a-dA-D])', text)
        if m:
            qnum = int(m.group(1))
            letter = m.group(2).lower()
            answers[qnum] = letter

    return answers


# ---------- Parsing questions ----------

def is_question_start(text: str) -> bool:
    return re.match(r'^\d+\)\s', text.strip()) is not None


def is_choice_line(text: str) -> bool:
    return re.match(r'^[a-dA-D]\.\s', text.strip()) is not None


def parse_questions(questions_doc_path: Path, answers: dict):
    """
    Parse the questions .docx into a list of dicts matching the YAML schema:

    - id: Q1
      points: 1
      type: mcq
      stem: [ {type: text|code, text: "...", ...}, ... ]
      choices: [ {key: 'a', text: '...'}, ... ]
      correct: 'b'
    """
    doc = Document(str(questions_doc_path))
    paras = [p.text for p in doc.paragraphs]
    n = len(paras)
    i = 0
    questions = []

    # Skip title or other non-question paragraphs at top
    while i < n and not is_question_start(paras[i]):
        i += 1

    while i < n:
        line = paras[i].strip()
        if not is_question_start(line):
            i += 1
            continue

        # Parse "N) ..." line
        m = re.match(r'^(\d+)\)\s*(.*)$', line)
        if not m:
            i += 1
            continue

        qnum = int(m.group(1))
        first_rest = m.group(2).strip()

        stem_blocks = []
        choices = []

        # If the first line has text after "N)", treat as stem text
        if first_rest:
            stem_blocks.append({
                "type": "text",
                "text": first_rest,
            })

        i += 1

        # Collect stem paragraphs until we hit choices or next question
        while i < n:
            raw = paras[i]
            text = raw.strip()

            # Skip blank lines
            if not text:
                i += 1
                continue

            # Next question?
            if is_question_start(text):
                break

            # Choices start?
            if is_choice_line(text):
                break

            print(raw)
            # Otherwise, part of stem
            if "\n" in raw:
                # Treat as code block with literal newlines
                stem_blocks.append({
                    "type": "code",
                    "language": "python",   # assumption for ENGR 131
                    "style": "mypython",    # matches your LaTeX listings style
                    "text": raw,
                })
            else:
                stem_blocks.append({
                    "type": "text",
                    "text": text,
                })

            i += 1

        # Collect choices if present
        while i < n and is_choice_line(paras[i].strip()):
            line = paras[i].strip()
            m = re.match(r'^([a-dA-D])\.\s*(.*)$', line)
            if m:
                key = m.group(1).lower()
                text = m.group(2)
                choices.append({
                    "key": key,
                    "text": text,
                })
            i += 1

        # Build question dict
        qid = f"Q{qnum}"
        correct = answers.get(qnum)

        question = {
            "id": qid,
            "points": 1,
            "type": "mcq",
            "stem": stem_blocks,
            "choices": choices,
            "correct": correct,
        }

        questions.append(question)

        # After choices, loop continues; top of while will detect next question start

    return questions


# ---------- Main entry point ----------

def main():
    if len(sys.argv) != 3:
        print("Usage: python zybooks_zip_to_yaml.py export.zip output.yaml")
        sys.exit(1)

    zip_path = Path(sys.argv[1])
    output_yaml_path = Path(sys.argv[2])

    if not zip_path.is_file():
        print(f"Error: {zip_path} does not exist or is not a file.")
        sys.exit(1)

    # Extract zip to temp dir
    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)

        with zipfile.ZipFile(zip_path, 'r') as zf:
            zf.extractall(tmpdir_path)

        # Find the question and answer-key docs
        questions_doc, answer_key_doc = find_docs_in_extracted_dir(tmpdir_path)

        print(f"Using questions doc:   {questions_doc.name}")
        print(f"Using answer key doc:  {answer_key_doc.name}")

        # Parse
        answers = parse_answer_key(answer_key_doc)
        questions = parse_questions(questions_doc, answers)

        data = {"questions": questions}

        # Write YAML
        with output_yaml_path.open("w", encoding="utf-8") as f:
            yaml.safe_dump(
                data,
                f,
                sort_keys=False,
                allow_unicode=True,
                width=80,
            )

        print(f"Wrote {len(questions)} questions to {output_yaml_path}")


if __name__ == "__main__":
    main()
