import os
import sys
import re
from typing import Optional,Dict,List,Tuple
from AnyQt.QtWidgets import QTableWidgetItem

from Orange.widgets.settings import Setting
import Orange.data
from Orange.data import Table
from AnyQt.QtWidgets import QApplication
from Orange.widgets import widget

# Import des bibliothèques Word
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
if "site-packages/Orange/widgets" in os.path.dirname(os.path.abspath(__file__)).replace("\\", "/"):
    from Orange.widgets.orangecontrib.AAIT.utils.import_uic import uic

else:
    from orangecontrib.AAIT.utils.import_uic import uic

class OWDocumentGenerator(widget.OWWidget):
    """
    Widget simplifié pour générer des documents Word
    - Input 1 : Chemins (template in/out)
    - Input 2 : Remplacements texte (balise -> valeur)
    - Input 3 : Tables à insérer (nom = balise cible)
    """
    name = "Document Generator"
    description = "Modify Word documents with replacements and tables."
    category = "AAIT - TOOLBOX"
    icon = "icons/document_generator.png"
    if "site-packages/Orange/widgets" in os.path.dirname(os.path.abspath(__file__)).replace("\\", "/"):
        icon = "icons_dev/document_generator.png"

    gui = os.path.join(os.path.dirname(os.path.abspath(__file__)), "designer/owgenerate_word.ui")
    want_control_area = False
    priority = 1070

    key_remplace: str = Setting("True")# "True" => replace [col] with row0 value else   => insert full table at [TableKey]


    class Inputs:
        data = Orange.widgets.widget.Input("Data", Orange.data.Table)

    class Outputs:
        data = Orange.widgets.widget.Output("Data", Orange.data.Table)

    # -----------------------------
    # Colors: %!color!% prefix (headers + cells)
    # -----------------------------
    COLOR_MAP = {
        # neutres
        "blanc": "FFFFFF", "white": "FFFFFF",
        "gris": "E0E0E0", "gray": "E0E0E0", "grey": "E0E0E0",
        "gris_clair": "F2F2F2",

        # pastel lisibles
        "jaune": "FFF2CC", "yellow": "FFF2CC",
        "vert": "E2F0D9", "green": "E2F0D9",
        "rouge": "FCE4D6", "red": "FCE4D6",
        "bleu": "DDEBF7", "blue": "DDEBF7",
        "violet": "EDE7F6", "purple": "EDE7F6",
        "orange": "FBE5D6",
        "cyan": "E7F3F8", "turquoise": "E7F3F8",
        "rose": "FCE4EC", "pink": "FCE4EC",
    }

    def __init__(self):
        super().__init__()
        self.setFixedWidth(470)
        self.setFixedHeight(300)
        uic.loadUi(self.gui, self)

        self.tableWidget.insertRow(self.tableWidget.rowCount())


        self.tableWidget.setItem(0, 0, QTableWidgetItem("C:/path.docx"))
        self.tableWidget.setItem(0, 1, QTableWidgetItem("Value1"))
        self.tableWidget.setItem(0, 2, QTableWidgetItem("Value2"))
        self.tableWidget.setItem(0, 3, QTableWidgetItem("Value3"))

        self.tableWidget_2.insertRow(self.tableWidget_2.rowCount())

        self.tableWidget_2.setItem(0, 0, QTableWidgetItem("C:/path.docx"))
        self.tableWidget_2.setItem(0, 1, QTableWidgetItem("Value1"))
        self.tableWidget_2.setItem(0, 2, QTableWidgetItem("%!red!%..."))
        self.tableWidget_2.setItem(0, 3, QTableWidgetItem("..."))




        self.data: Optional[Table] = None

        # Word table style (optional)
        self.table_style_name = "Table Grid"
        if self.key_remplace=="True":
            self.radioButton.setChecked(True)
        else:
            self.radioButton_2.setChecked(True)


        self.radioButton.toggled.connect(self.on_radio_changed)
        self.radioButton_2.toggled.connect(self.on_radio_changed)




    def on_radio_changed(self, checked):
        if not checked:
            return  # on ignore le "décoché"

        sender = self.sender()

        if sender == self.radioButton:
            self.key_remplace="True"
        elif sender == self.radioButton_2:
            self.key_remplace="False"

    # -----------------------------
    # Input
    # -----------------------------
    @Inputs.data
    def set_data(self, data: Optional[Table]):
        self.data = data
        self.process()

    # -----------------------------
    # Helpers
    # -----------------------------
    def _clean_orange_value(self, s) -> str:
        if s is None:
            return ""
        v = str(s).strip()
        if v in ("?", "nan", "NaN", "None"):
            return ""
        return v

    def _normalize_hex(self, color: str) -> str:
        if not color:
            return "FFFFFF"
        c = str(color).strip().lower()
        if c.startswith("#"):
            c = c[1:]
        if re.fullmatch(r"[0-9a-f]{6}", c):
            return c.upper()
        return self.COLOR_MAP.get(c, "FFFFFF")

    def parse_color_prefix(self, value: str) -> Tuple[str, Optional[str]]:
        """
        Detect prefix %!color!% or %!#RRGGBB!%
        Returns (text_without_prefix, hex_color_or_None)
        """
        if value is None:
            return "", None

        s = str(value)
        m = re.match(r"^%!(.*?)!%", s)
        if not m:
            return s, None

        token = m.group(1).strip()
        text = s[m.end():]
        return text, self._normalize_hex(token)

    def set_cell_bg(self, cell, hex_rgb: str):
        hex_rgb = self._normalize_hex(hex_rgb)

        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tcPr.append(shd)

        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_rgb)

    def _get_first_row_mapping(self, table: Table) -> Dict[str, str]:
        if table is None or len(table) == 0:
            return {}

        row0 = table[0]
        mapping: Dict[str, str] = {}

        for var in table.domain.attributes:
            try:
                mapping[var.name] = self._clean_orange_value(row0[var])
            except Exception:
                mapping[var.name] = ""

        for var in table.domain.class_vars:
            try:
                mapping[var.name] = self._clean_orange_value(row0[var])
            except Exception:
                mapping[var.name] = ""

        for j, var in enumerate(table.domain.metas):
            try:
                mapping[var.name] = self._clean_orange_value(table.metas[0, j])
            except Exception:
                mapping[var.name] = ""

        return mapping

    def _bracketed_only(self, key: str) -> str:
        k = (key or "").strip()
        if not k:
            return ""
        if k.startswith("[") and k.endswith("]") and len(k) >= 2:
            inner = k[1:-1].strip()
            return f"[{inner}]" if inner else ""
        return f"[{k}]"

    def _extract_full_table_matrix(self, table: Table, exclude_names: set) -> Tuple[List[str], List[List[str]]]:
        """
        Extract headers + rows from Orange Table (all rows),
        using attrs + targets + metas, excluding certain names.
        NOTE: headers keep raw names (can include %!color!% prefix).
        """
        headers: List[str] = []
        specs = []  # list of tuples describing how to fetch values (kind, var, meta_index)

        for var in table.domain.attributes:
            if var.name not in exclude_names:
                headers.append(var.name)
                specs.append(("attr", var, None))

        for var in table.domain.class_vars:
            if var.name not in exclude_names:
                headers.append(var.name)
                specs.append(("target", var, None))

        for j, var in enumerate(table.domain.metas):
            if var.name not in exclude_names:
                headers.append(var.name)
                specs.append(("meta", None, j))

        rows: List[List[str]] = []
        for i in range(len(table)):
            row_vals: List[str] = []
            for kind, var, j in specs:
                try:
                    if kind in ("attr", "target"):
                        row_vals.append(self._clean_orange_value(table[i][var]))
                    else:
                        row_vals.append(self._clean_orange_value(table.metas[i, j]))
                except Exception:
                    row_vals.append("")
            rows.append(row_vals)

        return headers, rows

    # -----------------------------
    # Format-preserving replacement in paragraph runs
    # -----------------------------
    def _replace_in_paragraph_preserve_format(self, paragraph, placeholder: str, replacement: str) -> int:
        if not placeholder:
            return 0

        runs = paragraph.runs
        if not runs:
            return 0

        full = "".join(r.text for r in runs)
        if placeholder not in full:
            return 0

        repl_count = 0

        while True:
            runs = paragraph.runs
            full = "".join(r.text for r in runs)

            start = full.find(placeholder)
            if start < 0:
                break
            end = start + len(placeholder)

            spans = []
            pos = 0
            for idx, r in enumerate(runs):
                t = r.text or ""
                spans.append((idx, pos, pos + len(t)))
                pos += len(t)

            start_run = None
            for (i, a, b) in spans:
                if a <= start < b:
                    start_run = (i, a, b)
                    break
            if start_run is None:
                break

            end_run = None
            for (i, a, b) in spans:
                if a < end <= b:
                    end_run = (i, a, b)
                    break
            if end_run is None:
                for (i, a, b) in spans:
                    if end == a:
                        end_run = (i, a, b)
                        break

            si, sa, sb = start_run
            if end_run is None:
                ei, ea, eb = si, sa, sb
            else:
                ei, ea, eb = end_run

            start_off = start - sa

            if si == ei:
                r = runs[si]
                text = r.text or ""
                before = text[:start_off]
                after = text[start_off + len(placeholder):]
                r.text = before + replacement + after
                repl_count += 1
                continue

            first = runs[si]
            last = runs[ei]

            first_text = first.text or ""
            last_text = last.text or ""

            before = first_text[:start_off]
            after = last_text[end - ea:]

            first.text = before

            insert_run = paragraph.add_run(replacement)
            insert_run.bold = first.bold
            insert_run.italic = first.italic
            insert_run.underline = first.underline
            insert_run.style = first.style

            try:
                if first.font is not None:
                    insert_run.font.name = first.font.name
                    insert_run.font.size = first.font.size
                    if getattr(first.font.color, "rgb", None) is not None:
                        insert_run.font.color.rgb = first.font.color.rgb
                    insert_run.font.highlight_color = first.font.highlight_color
                    insert_run.font.all_caps = first.font.all_caps
                    insert_run.font.small_caps = first.font.small_caps
                    insert_run.font.strike = first.font.strike
                    insert_run.font.double_strike = first.font.double_strike
                    insert_run.font.subscript = first.font.subscript
                    insert_run.font.superscript = first.font.superscript
            except Exception:
                pass

            first._element.addnext(insert_run._element)

            for idx in range(si + 1, ei + 1):
                runs[idx].text = ""

            runs[ei].text = after

            repl_count += 1

        return repl_count

    # -----------------------------
    # Replace everywhere (body + tables + headers/footers)
    # -----------------------------
    def _replace_everywhere_in_container(self, container, placeholder: str, replacement: str) -> int:
        count = 0

        for p in getattr(container, "paragraphs", []):
            count += self._replace_in_paragraph_preserve_format(p, placeholder, replacement)

        for t in getattr(container, "tables", []):
            for row in t.rows:
                for cell in row.cells:
                    count += self._replace_everywhere_in_container(cell, placeholder, replacement)

        return count

    def _replace_everywhere(self, doc: Document, placeholder: str, replacement: str) -> int:
        total = 0
        total += self._replace_everywhere_in_container(doc, placeholder, replacement)

        for section in doc.sections:
            if section.header is not None:
                total += self._replace_everywhere_in_container(section.header, placeholder, replacement)
            if section.footer is not None:
                total += self._replace_everywhere_in_container(section.footer, placeholder, replacement)

        return total

    # -----------------------------
    # Insert Word table at a [Tag] (body + headers/footers)
    # -----------------------------
    def _apply_table_style_if_possible(self, table_obj) -> None:
        try:
            if self.table_style_name:
                table_obj.style = self.table_style_name
        except Exception:
            pass

    def _insert_table_after_paragraph(self, doc: Document, paragraph, headers: List[str], rows: List[List[str]]) -> None:
        n_cols = max(1, len(headers))
        n_rows = 1 + len(rows)

        tbl = doc.add_table(rows=n_rows, cols=n_cols)
        self._apply_table_style_if_possible(tbl)

        # Header row (support %!color!% prefix)
        for j, raw_h in enumerate(headers):
            cell_h = tbl.cell(0, j)
            clean_h, hex_bg = self.parse_color_prefix(raw_h)
            cell_h.text = clean_h
            if hex_bg:
                self.set_cell_bg(cell_h, hex_bg)

        # Data rows (support %!color!% prefix per cell)
        for i, row_vals in enumerate(rows, start=1):
            for j in range(n_cols):
                v = row_vals[j] if j < len(row_vals) else ""
                cell_d = tbl.cell(i, j)

                clean_v, hex_bg = self.parse_color_prefix(v)
                cell_d.text = "" if clean_v is None else str(clean_v)

                if hex_bg:
                    self.set_cell_bg(cell_d, hex_bg)

        paragraph._element.addnext(tbl._element)

    def _insert_table_at_tag_in_container(self, doc: Document, container, tag: str,
                                         headers: List[str], rows: List[List[str]]) -> bool:
        for p in getattr(container, "paragraphs", []):
            try:
                if tag in (p.text or ""):
                    self._replace_in_paragraph_preserve_format(p, tag, "")
                    self._insert_table_after_paragraph(doc, p, headers, rows)
                    return True
            except Exception as e:
                import traceback
                print("Insert table error:", e)
                print(traceback.format_exc())
                continue

        for t in getattr(container, "tables", []):
            for row in t.rows:
                for cell in row.cells:
                    if self._insert_table_at_tag_in_container(doc, cell, tag, headers, rows):
                        return True

        return False

    def _insert_table_at_tag(self, doc: Document, tag: str,
                             headers: List[str], rows: List[List[str]]) -> bool:
        if self._insert_table_at_tag_in_container(doc, doc, tag, headers, rows):
            return True

        for section in doc.sections:
            if section.header is not None:
                if self._insert_table_at_tag_in_container(doc, section.header, tag, headers, rows):
                    return True
            if section.footer is not None:
                if self._insert_table_at_tag_in_container(doc, section.footer, tag, headers, rows):
                    return True

        return False

    # -----------------------------
    # Main
    # -----------------------------
    def process(self):
        self.error("")
        self.warning("")

        if self.data is None:
            self.Outputs.data.send(None)
            return

        if len(self.data) == 0:
            self.error("La table d'entrée est vide (aucune ligne).")
            self.Outputs.data.send(None)
            return

        try:
            mapping = self._get_first_row_mapping(self.data)

            if "DocxPath" not in mapping:
                self.error("Colonne 'DocxPath' introuvable (attributes/targets/metas).")
                self.Outputs.data.send(None)
                return

            docx_path = self._clean_orange_value(mapping.get("DocxPath", "")).strip()
            if not docx_path:
                self.error("La valeur de 'DocxPath' est vide sur la première ligne.")
                self.Outputs.data.send(None)
                return

            docx_path = docx_path.replace("\\", os.sep).replace("/", os.sep)
            if not os.path.exists(docx_path):
                self.error(f"Fichier DOCX introuvable : {docx_path}")
                self.Outputs.data.send(None)
                return

            doc = Document(docx_path)

            # MODE 1: key/value replacements in [tags] from row0
            if str(self.key_remplace) == "True":
                replacements = {k: v for k, v in mapping.items() if k != "DocxPath"}
                for key, val in replacements.items():
                    tag = self._bracketed_only(key)  # only [key]
                    if not tag:
                        continue
                    rep = self._clean_orange_value(val)
                    self._replace_everywhere(doc, tag, rep)

                doc.save(docx_path)
                self.Outputs.data.send(self.data)
                return

            # MODE 2: insert full table at [TableKey]
            if "TableKey" not in mapping:
                self.error("Mode table: colonne 'TableKey' introuvable sur la première ligne.")
                self.Outputs.data.send(None)
                return

            table_key = self._clean_orange_value(mapping.get("TableKey", "")).strip()
            table_key, _ = self.parse_color_prefix(table_key)  # enlève %!couleur!%
            table_key = (table_key or "").strip()
            if not table_key:
                self.error("Mode table: la valeur de 'TableKey' est vide sur la première ligne.")
                self.Outputs.data.send(None)
                return

            insert_tag = self._bracketed_only(table_key)  # => [Tableau1]
            if not insert_tag:
                self.error("Mode table: TableKey invalide.")
                self.Outputs.data.send(None)
                return
            exclude = {"DocxPath", "TableKey"}
            headers, rows = self._extract_full_table_matrix(self.data, exclude_names=exclude)

            if not headers:
                self.error("Mode table: aucune colonne à insérer (tout est exclu).")
                self.Outputs.data.send(None)
                return

            inserted = self._insert_table_at_tag(doc, insert_tag, headers, rows)
            if not inserted:
                self.error(f"Balise d'insertion non trouvée dans le DOCX : {insert_tag}")
                self.Outputs.data.send(None)
                return

            doc.save(docx_path)
            self.Outputs.data.send(self.data)

        except Exception as e:
            self.error(f"Erreur : {e}")
            self.Outputs.data.send(None)

if __name__ == "__main__":
    app = QApplication(sys.argv)
    widget = OWDocumentGenerator()
    widget.show()
    if hasattr(app, "exec"):
        app.exec()
    else:
        app.exec_()
