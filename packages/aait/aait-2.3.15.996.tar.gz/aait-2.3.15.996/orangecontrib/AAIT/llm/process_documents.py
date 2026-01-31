import os
import json
import ntpath
from pathlib import Path

from Orange.data import Table, Domain, StringVariable

import fitz
import docx



def process_documents(dirpath):
    if dirpath is None or not os.path.exists(dirpath):
        return None, None

    # get path from user selection
    embeddings = check_for_embeddings(dirpath)
    dirpath = dirpath.replace("\\","/")

    # Set selected path in the saved embeddings
    if embeddings is not None:
        common_path = find_common_root(embeddings).replace("\\","/")
        for row in embeddings:
            row["path"] = row["path"].value.replace("\\","/").replace(common_path, dirpath)

    # Verify which files are already processed
    files_to_process = get_files_to_process(dirpath, embeddings)

    rows = []
    for file in files_to_process:
        # Get the text content from the file
        content = extract_text(file)
        filename = ntpath.basename(file)
        # Build a row containing dirpath | filename | content
        row = [file, filename, content]
        rows.append(row)

    # Build a table with the constructed rows
    path_var = StringVariable("path")
    name_var = StringVariable("name")
    content_var = StringVariable("content")
    domain = Domain(attributes=[], metas=[path_var, name_var, content_var])
    out_data = Table.from_list(domain=domain, rows=rows)
    return out_data, embeddings


def find_common_root(data_table, column_name="path"):
    """Finds the common root path from a column of file paths in an Orange Data Table."""
    paths = [str(row[column_name]) for row in data_table if row[column_name] is not None]
    if not paths:
        return ""
    return os.path.commonpath(paths)


def get_files_to_process(folder_path, table=None):
    """
    Finds all PDF files in a folder (including subfolders) that are not already in the table.
    The comparison is based on "name" (relative path from the main folder) instead of full paths.

    :param folder_path: Path to the folder to scan for documents.
    :param table: Orange Data Table with columns "path", "name", and "content".
    :return: List of paths to files not present in the table (by name, including subfolder structure).
    """
    #TODO
    # Supported file extensions
    supported_extensions = [".pdf", ".docx"]

    # Read the json containing file sizes
    filepath_sizes = os.path.join(folder_path, "sizes.json")
    if os.path.exists(filepath_sizes):
        with open(filepath_sizes, "r") as json_file:
            sizes = json.load(json_file)
    else:
        sizes = dict()

    # Extract the existing file names from the Orange Data Table
    if table:
        existing_paths = set(table[:, "path"].metas.flatten())  # Extract names from the table
    else:
        existing_paths = set()


    # Walk through the folder and its subfolders
    new_files = []
    for root, _, files in os.walk(folder_path):
        for file in files:
            # Check if the file has a supported extension
            if os.path.splitext(file)[1].lower() in supported_extensions:
                # Add the file if it is not already in the table
                filepath = os.path.join(root, file).replace("\\","/")
                if filepath not in existing_paths:
                    new_files.append(filepath)
                    sizes[filepath] = os.path.getsize(filepath)
                # If the file is in the table, verify if the file has been modified (comparing the size)
                else:
                    new_size = os.path.getsize(filepath)
                    if filepath not in sizes.keys():
                        sizes[filepath] = new_size
                    else:
                        old_size = sizes[filepath]
                        if old_size != new_size:
                            new_files.append(filepath)
                            table = remove_from_table(filepath, table)
                            sizes[filepath] = new_size
    with open(filepath_sizes, "w") as json_file:
        json.dump(sizes, json_file, indent=4)
    return new_files


def remove_from_table(filepath, table):
    filtered_table = Table.from_list(domain=table.domain,
                                      rows=[row for row in table if row["path"].value != filepath])
    return filtered_table


def check_for_embeddings(folder_path):
    """
    Check for an embeddings.pkl file in a given folder. Return its content if it exists.

    Parameters:
        folder_path (str): The path to the folder where embeddings.pkl may exist.

    Returns:
        Table or None: The content of embeddings.pkl.
    """
    filepaths = [os.path.join(folder_path, "embeddings_question.pkl"),
                 os.path.join(folder_path, "embeddings.pkl")]
    for filepath in filepaths:
        if os.path.exists(filepath):
            data = Table.from_file(filepath)
            return data
    else:
        return None


def load_documents_in_table(table, progress_callback=None, argself=None):
    """
    Load the text content of each document listed in a table and add it
    as a new column "content".

    :param table: Orange.data.Table containing file paths in a column named "path".
    :return: Orange.data.Table with an added meta column "content" containing the extracted text.
    """
    # Make a copy of the table to avoid modifying the original
    data = table.copy()

    # List to store text from each document
    texts = []
    names = []
    # Iterate over all rows in the table
    for i, row in enumerate(data):
        # Get file path from the "path" column
        filepath = row["path"].value
        # Get text and name
        name = Path(filepath).name
        text = extract_text(filepath)
        # Store results
        names.append(name)
        texts.append(text)
        # Update progress if a callback is provided
        if progress_callback is not None:
            progress_value = float(100 * (i + 1) / len(data))
            progress_callback(progress_value)
        # Check if processing should be stopped
        if argself is not None and getattr(argself, "stop", False):
            break

    # Create a StringVariable for the new column
    var_content = StringVariable("content")
    var_name = StringVariable("name")

    # Add the column as a meta-column in the table
    data = data.add_column(variable=var_name, data=names, to_metas=True)
    data = data.add_column(variable=var_content, data=texts, to_metas=True)
    return data



def extract_text(filepath):
    """
    Extrait le texte d'un fichier en fonction de son type (PDF ou DOCX).

    :param filepath: Chemin vers le fichier.
    :return: Texte extrait du fichier sous forme de chaîne.
    """
    try:
        # Vérifie l'extension du fichier
        file_extension = os.path.splitext(filepath)[1].lower()

        if file_extension == ".pdf":
            return extract_text_from_pdf(filepath)
        elif file_extension == ".docx":
            return extract_text_from_docx(filepath)
        elif file_extension in [".txt", ".md", ".py", ".html", ".json", ".ows"]:
            return extract_text_from_txt(filepath)
        else:
            return "ERROR: Unsupported file format. Please use a .pdf, .docx, .txt, .md, .py, .html, .ows or .json file."
    except Exception as e:
        print(f"Erreur lors de l'extraction de texte depuis {filepath}: {e}")
        return f"ERROR: Extraction Error ({e})"


def extract_text_from_pdf(pdf_path):
    """
    Extrait le texte d'un fichier PDF.

    :param pdf_path: Chemin vers le fichier PDF.
    :return: Texte extrait du PDF sous forme de chaîne.
    """
    try:
        # Ouvre le fichier PDF
        pdf_document = fitz.open(pdf_path)
        extracted_text = ""

        # Parcourt toutes les pages et extrait le texte
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            extracted_text += page.get_text()

        pdf_document.close()
        return extracted_text
    except Exception as e:
        print(f"Erreur lors de l'extraction de texte depuis {pdf_path}: {e}")
        return f"ERROR: Extraction Error ({e})"


def extract_text_from_docx(docx_path):
    """
    Extrait le texte d'un fichier DOCX en conservant l'ordre des éléments (paragraphes, tableaux et titres).

    :param docx_path: Chemin vers le fichier DOCX.
    :return: Texte extrait du document sous forme de chaîne.
    """
    try:
        doc = docx.Document(docx_path)
        extracted_text = []
        title_numbers = {}  # Dictionary to track numbering per heading level

        for para in doc.paragraphs:
            # Vérifie si c'est un titre
            if para.style.name.startswith('Heading'):
                heading_level = int(para.style.name.split()[-1])  # Niveau du titre (1, 2, 3, etc.)
                heading_text = para.text.strip()

                # Met à jour la numérotation des titres
                if heading_level not in title_numbers:
                    title_numbers[heading_level] = 1  # Nouveau niveau
                else:
                    title_numbers[heading_level] += 1  # Incrémente niveau actuel

                # Réinitialise les niveaux inférieurs
                for level in list(title_numbers.keys()):
                    if level > heading_level:
                        del title_numbers[level]

                # Forme le numéro du titre (ex: "1", "1.1", "1.2.1")
                full_title = ".".join(str(title_numbers[i]) for i in sorted(title_numbers.keys()))
                extracted_text.append(f"\n{full_title} {heading_text}")  # Ajoute le titre formaté
            else:
                extracted_text.append(para.text.strip())  # Ajoute le paragraphe

        # Parcourt les tableaux du document
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]  # Extrait le texte de chaque cellule
                table_text.append("\t".join(row_text))  # Sépare les colonnes par des tabulations
            extracted_text.append("\n".join(table_text))  # Ajoute le tableau sous forme de texte
        return "\n".join(filter(None, extracted_text))  # Retourne le texte en filtrant les vides

    except Exception as e:
        print(f"Erreur lors de l'extraction de texte depuis {docx_path}: {e}")
        return f"ERROR: Extraction Error ({e})"


def extract_text_from_txt(filepath):
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"Erreur lors de l'extraction de texte depuis {filepath}: {e}")
        return f"ERROR: Extraction Error ({e})"


def get_pages_of_extract(pdf_path, extract):
    """
    Identify the pages that a given extract belongs to.

    :param pdf_path: The path of the pdf to search in.
    :param extract: The text snippet to locate.
    :return: A list of page numbers the extract spans.
    """
    full_text, page_mapping = load_pdf_with_sparse_mapping(pdf_path)
    # Find the start index of the extract in the full text
    start_index = full_text.find(extract)
    if start_index == -1:
        return []  # Extract not found

    # Determine the end index of the extract
    end_index = start_index + len(extract) - 1

    # Find all pages the extract spans
    pages = []
    for page, (start, end) in page_mapping.items():
        if start <= end_index and end >= start_index:
            pages.append(page)

    if pages == []:
        return [1]
    return pages


def load_pdf_with_sparse_mapping(pdf_path):
    doc = fitz.open(pdf_path)
    full_text = ""
    page_mapping = {}  # Sparse mapping: {page_num: (start_index, end_index)}

    for page_num in range(len(doc)):
        page_text = doc[page_num].get_text()
        start_index = len(full_text)
        full_text += page_text
        end_index = len(full_text) - 1
        page_mapping[page_num + 1] = (start_index, end_index)

    doc.close()
    return full_text, page_mapping
